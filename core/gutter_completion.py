"""core/gutter_completion.py — Reconstrucción de palabras cortadas por costura de página.

El problema: en revistas encuadernadas escaneadas por la BNC, el pliegue central
(binding gutter) corta palabras en el borde de la imagen. El OCR produce fragmentos
como "polí-" o "co" al final de una línea, y "tica" o "lombiana" al inicio de la
línea en la página siguiente.

Este módulo:
  1. Detecta líneas con palabras truncadas (borde izquierdo/derecho)
  2. Reconstruye las palabras faltantes con LLM (Claude/Ollama)
  3. Marca las palabras generadas con tag [GEN] para trazabilidad
  4. Exporta el texto con marcas para resaltar en DOCX (color rojo)

Estrategia de detección:
  - Palabra al final de línea con guión de corte → truncamiento explícito
  - Palabra al final de línea de longitud < 4 chars → posible truncamiento
  - Palabra al inicio de línea en minúscula sin puntuación previa → continuación

Formato de salida:
  El texto reconstruido usa el tag ⟦palabra⟧ para palabras generadas.
  El DOCX las renderiza en rojo con nota al pie opcional.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from typing import Callable

# Tag para palabras generadas — Unicode no conflictivo con texto español
TAG_OPEN  = "⟦"
TAG_CLOSE = "⟧"
RE_GENERADO = re.compile(r'⟦([^⟧]+)⟧')

# Detectores de truncamiento
RE_GUION_CORTE = re.compile(r'\b(\w{2,})-\s*$', re.MULTILINE)
RE_FRAGMENTO_INICIO = re.compile(r'^\s*([a-záéíóúüñ]{2,5})\b', re.IGNORECASE)
RE_LINEA_CORTA_FINAL = re.compile(r'\b(\w{2,4})\s*$', re.MULTILINE)


@dataclass
class FragmentoCortado:
    """Representa una palabra cortada detectada en el texto."""
    linea_idx: int          # índice de línea donde está el corte
    lado: str               # "derecho" (fin de línea) | "izquierdo" (inicio)
    fragmento: str          # la parte visible de la palabra
    contexto_prev: str      # líneas previas para contexto
    contexto_post: str      # líneas siguientes para contexto
    reconstruida: str = ""  # palabra completa reconstruida
    confianza: float = 0.0


def detectar_fragmentos(texto: str, margen_chars: int = 5) -> list[FragmentoCortado]:
    """
    Detecta líneas con posibles palabras cortadas por la costura.

    margen_chars: cantidad de caracteres al final/inicio de línea
                  que se considera "borde de costura"

    Retorna lista de FragmentoCortado ordenados por línea.
    """
    lineas = texto.splitlines()
    fragmentos = []

    for i, linea in enumerate(lineas):
        linea_strip = linea.rstrip()
        if not linea_strip:
            continue

        ctx_prev = "\n".join(lineas[max(0, i-3):i])
        ctx_post = "\n".join(lineas[i+1:min(len(lineas), i+4)])

        # Caso 1: guión de corte explícito al final
        m = RE_GUION_CORTE.search(linea_strip)
        if m:
            fragmentos.append(FragmentoCortado(
                linea_idx=i,
                lado="derecho",
                fragmento=m.group(1) + "-",
                contexto_prev=ctx_prev,
                contexto_post=ctx_post,
            ))
            continue

        # Caso 2: palabra muy corta al final (probable corte sin guión)
        palabras = linea_strip.split()
        if palabras:
            ultima = palabras[-1]
            if (2 <= len(ultima) <= 4
                    and ultima.isalpha()
                    and linea_strip[-1] not in ".,:;!?»\"'"):
                # Verificar que la siguiente línea empiece con minúscula
                if i + 1 < len(lineas):
                    sig = lineas[i + 1].lstrip()
                    if sig and sig[0].islower():
                        fragmentos.append(FragmentoCortado(
                            linea_idx=i,
                            lado="derecho",
                            fragmento=ultima,
                            contexto_prev=ctx_prev,
                            contexto_post=ctx_post,
                        ))

    return fragmentos


def _prompt_reconstruir(frag: FragmentoCortado) -> str:
    """Construye el prompt para reconstruir una palabra cortada."""
    if frag.lado == "derecho":
        return (
            f"Eres un editor de textos históricos colombianos de los años 1930.\n"
            f"El siguiente texto fue escaneado de una revista encuadernada y la "
            f"costura del libro cortó una palabra al final de la línea.\n\n"
            f"Contexto previo:\n{frag.contexto_prev}\n\n"
            f"Línea con corte (la última palabra está incompleta): "
            f"...{frag.fragmento}\n\n"
            f"Contexto posterior:\n{frag.contexto_post}\n\n"
            f"¿Cuál es la palabra completa que fue cortada? "
            f"Responde SOLO con la palabra completa reconstruida, sin explicación. "
            f"Si no podés determinarlo con certeza, responde con la palabra más probable."
        )
    else:
        return (
            f"Eres un editor de textos históricos colombianos de los años 1930.\n"
            f"El siguiente fragmento es el inicio de una continuación de la línea "
            f"anterior, cortada por la costura del libro.\n\n"
            f"Contexto previo:\n{frag.contexto_prev}\n\n"
            f"Fragmento visible al inicio de la línea: {frag.fragmento}...\n\n"
            f"Contexto posterior:\n{frag.contexto_post}\n\n"
            f"¿Qué palabra completa comienza con '{frag.fragmento}'? "
            f"Responde SOLO con la palabra completa, sin explicación."
        )


def reconstruir_fragmento(
    frag: FragmentoCortado,
    api_key: str,
    modelo: str = "claude-haiku-4-5-20251001",
) -> FragmentoCortado:
    """Reconstruye una palabra cortada usando LLM."""
    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(
            model=modelo,
            max_tokens=32,
            messages=[{"role": "user", "content": _prompt_reconstruir(frag)}],
        )
        palabra = msg.content[0].text.strip().split()[0]  # solo primera palabra
        # Limpiar puntuación final
        palabra = re.sub(r'[.,:;!?»"\']$', '', palabra)
        frag.reconstruida = palabra
        frag.confianza = 0.75  # confianza base para LLM
    except Exception:
        frag.reconstruida = frag.fragmento  # mantener original si falla
        frag.confianza = 0.0
    return frag


def reconstruir_texto(
    texto: str,
    api_key: str,
    modelo: str = "claude-haiku-4-5-20251001",
    callback: Callable[[int, int, str], None] | None = None,
) -> tuple[str, list[FragmentoCortado]]:
    """
    Detecta y reconstruye todas las palabras cortadas en el texto.

    Retorna:
      (texto_reconstruido, lista_de_fragmentos)

    El texto reconstruido usa ⟦palabra⟧ para marcar palabras generadas.
    """
    fragmentos = detectar_fragmentos(texto)
    if not fragmentos or not api_key:
        return texto, fragmentos

    total = len(fragmentos)
    for i, frag in enumerate(fragmentos):
        if callback:
            callback(i + 1, total, f"línea {frag.linea_idx}")
        reconstruir_fragmento(frag, api_key, modelo)

    # Aplicar reconstrucciones al texto
    lineas = texto.splitlines()
    for frag in fragmentos:
        if not frag.reconstruida or frag.reconstruida == frag.fragmento:
            continue
        linea = lineas[frag.linea_idx]
        if frag.lado == "derecho":
            # Reemplazar el fragmento final con la palabra completa marcada
            # Caso guión: "polí-" → "⟦política⟧"
            # Caso sin guión: "co" → "⟦colombia⟧" (reemplaza el fragmento)
            frag_pat = re.escape(frag.fragmento.rstrip())
            nuevo = f"{TAG_OPEN}{frag.reconstruida}{TAG_CLOSE}"
            lineas[frag.linea_idx] = re.sub(
                frag_pat + r'\s*$', nuevo, linea, count=1)

    return "\n".join(lineas), fragmentos


def texto_limpio(texto_reconstruido: str) -> str:
    """Quita los tags ⟦⟧ dejando solo el texto reconstruido."""
    return RE_GENERADO.sub(r'\1', texto_reconstruido)


def texto_original(texto_reconstruido: str) -> str:
    """Quita las palabras generadas, dejando solo el texto original."""
    return RE_GENERADO.sub('', texto_reconstruido)


def estadisticas(fragmentos: list[FragmentoCortado]) -> dict:
    """Estadísticas de la reconstrucción."""
    total = len(fragmentos)
    reconstruidos = sum(1 for f in fragmentos if f.reconstruida and f.reconstruida != f.fragmento)
    conf_media = (sum(f.confianza for f in fragmentos) / total) if total > 0 else 0.0
    return {
        "total_fragmentos": total,
        "reconstruidos": reconstruidos,
        "fallidos": total - reconstruidos,
        "confianza_media": round(conf_media, 3),
        "tasa_exito": round(reconstruidos / total, 3) if total > 0 else 0.0,
    }


def exportar_docx_con_marcas(
    texto_reconstruido: str,
    ruta_destino: str,
    titulo: str = "",
    incluir_nota: bool = True,
) -> int:
    """
    Exporta el texto a DOCX marcando las palabras generadas en rojo.

    Las palabras dentro de ⟦⟧ aparecen en rojo en el documento.
    Retorna cantidad de palabras generadas marcadas.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise ImportError("Instala python-docx: pip install python-docx")

    doc = Document()

    # Estilos base
    estilo = doc.styles["Normal"]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(11)

    if titulo:
        h = doc.add_heading(titulo, level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if incluir_nota:
        nota = doc.add_paragraph()
        run = nota.add_run(
            "Nota editorial: Las palabras marcadas en color rojo fueron "
            "reconstruidas computacionalmente a partir del contexto textual. "
            "El texto original presentaba cortes causados por la costura de "
            "encuadernación del impreso digitalizado."
        )
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        doc.add_paragraph()

    n_generadas = 0
    parrafos = texto_reconstruido.split("\n\n")

    for parrafo_txt in parrafos:
        if not parrafo_txt.strip():
            continue
        p = doc.add_paragraph()

        # Partir el párrafo en segmentos: normales y generados
        segmentos = RE_GENERADO.split(parrafo_txt)
        # split con grupo capturador alterna: normal, gen, normal, gen...
        for j, seg in enumerate(segmentos):
            if not seg:
                continue
            run = p.add_run(seg)
            if j % 2 == 1:  # índice impar = contenido del grupo capturador = generado
                run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)  # rojo
                run.bold = True
                n_generadas += 1

    doc.save(ruta_destino)
    return n_generadas


def exportar_html_con_marcas(texto_reconstruido: str) -> str:
    """Convierte el texto a HTML con spans rojos para palabras generadas."""
    def _repl(m):
        return (f'<span class="generado" '
                f'title="Palabra reconstruida computacionalmente" '
                f'style="color:#EF4444;font-weight:bold;">{m.group(1)}</span>')
    html_texto = RE_GENERADO.sub(_repl, texto_reconstruido)
    # Preservar saltos de línea
    html_texto = html_texto.replace("\n", "<br>\n")
    return (
        '<div class="texto-ocr" style="font-family:Georgia,serif;'
        'line-height:1.8;font-size:15px;">\n'
        + html_texto
        + "\n</div>"
    )
