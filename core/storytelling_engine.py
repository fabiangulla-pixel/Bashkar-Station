"""core/storytelling_engine.py — Narrativas académicas y reportes del corpus.

Genera:
  - Narrativas académicas automáticas sobre redes y tópicos (Claude)
  - Reporte HTML scrollytelling completo con estadísticas del corpus
  - Exportación Word (.docx)
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Callable

# ── Narrativa académica con Claude ───────────────────────────────────────────

_PROMPT_NARRATIVA = """\
Eres un historiador especializado en cultura colombiana del siglo XX.
Basándote en los siguientes datos cuantitativos del corpus de la revista *Estampa* (1930-1940),
redacta un párrafo académico en español que interprete los hallazgos.

Estilo: prosa académica, formal, sin tecnicismos computacionales.
Extensión: 150-200 palabras.
No uses listas ni viñetas; solo prosa continua.

Datos del análisis:
{datos}

Escribe solo el párrafo, sin título ni encabezado.
"""


def generar_narrativa(
    datos: dict,
    api_key: str,
    modelo: str = "claude-sonnet-4-6",
    seccion: str = "corpus",
) -> str:
    """
    Genera un párrafo narrativo académico sobre los datos del análisis.

    datos: dict con estadísticas (entidades top, tonos, métricas de red, etc.)
    seccion: "corpus" | "red" | "tono" | "lexico"
    """
    try:
        import anthropic
    except ImportError:
        raise ImportError("Instala anthropic: pip install anthropic>=0.25.0")

    from core import inference_provider as _ip

    datos_str = json.dumps(datos, ensure_ascii=False, indent=2)[:3000]
    client = anthropic.Anthropic(api_key=api_key)
    try:
        msg = client.messages.create(
            model=modelo,
            max_tokens=600,
            messages=[{"role": "user",
                        "content": _PROMPT_NARRATIVA.replace("{datos}", datos_str)}],
        )
        return _ip.texto_de_respuesta_claude(msg)
    except Exception as e:
        return f"[Error al generar narrativa: {e}]"


# ── Reporte HTML scrollytelling ──────────────────────────────────────────────

def generar_reporte_html(
    proyecto_nombre: str,
    stats_corpus: dict,
    indice_ner: dict,
    stats_tono: dict | None,
    metricas_red: dict | None,
    narrativas: dict | None,
    ruta: Path,
    callback: Callable[[str], None] | None = None,
) -> Path:
    """
    Genera reporte HTML completo con todas las estadísticas del corpus.

    stats_corpus: {n_pdfs, n_paginas, n_articulos, n_palabras_total, ...}
    indice_ner:   {categoria: {entidad: [arts]}}
    stats_tono:   resultado de sentiment_engine.estadisticas_tono()
    metricas_red: resultado de network_engine.metricas_red()
    narrativas:   {seccion: texto_narrativo}
    """
    def log(m):
        if callback:
            callback(m)

    ruta = Path(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    fecha = datetime.now().strftime("%d de %B de %Y")

    # ── Sección NER ───────────────────────────────────────────────────────────
    def top_entidades(cat, n=10):
        ents = indice_ner.get(cat, {})
        if not isinstance(ents, dict):
            return ""
        top = sorted(ents.items(), key=lambda x: len(x[1]), reverse=True)[:n]
        rows = "".join(
            f"<tr><td>{ent}</td><td>{len(arts)}</td></tr>"
            for ent, arts in top
        )
        return f"<table class='data-table'><tr><th>Entidad</th><th>Artículos</th></tr>{rows}</table>"

    ner_html = ""
    cat_labels = {
        "personas": "Personas",
        "lugares": "Lugares",
        "organizaciones": "Organizaciones",
        "obras_publicaciones": "Obras y publicaciones",
        "eventos_historicos": "Eventos históricos",
    }
    for cat, label in cat_labels.items():
        if cat in indice_ner and indice_ner[cat]:
            ner_html += f"<h3>{label}</h3>{top_entidades(cat)}"

    # ── Sección tono ──────────────────────────────────────────────────────────
    tono_html = ""
    if stats_tono and "distribucion" in stats_tono:
        rows = "".join(
            f"<tr><td>{t}</td><td>{d['n']}</td><td>{d['porcentaje']}%</td><td>{d['confianza_media']}</td></tr>"
            for t, d in stats_tono["distribucion"].items()
        )
        tono_html = (
            "<table class='data-table'>"
            "<tr><th>Tono</th><th>Artículos</th><th>%</th><th>Confianza</th></tr>"
            f"{rows}</table>"
        )

    # ── Sección red ───────────────────────────────────────────────────────────
    red_html = ""
    if metricas_red:
        items = "".join(
            f"<tr><td>{k.replace('_',' ').capitalize()}</td><td>{v}</td></tr>"
            for k, v in metricas_red.items() if k != "top_centralidad"
        )
        top_c = metricas_red.get("top_centralidad", [])
        top_rows = "".join(f"<tr><td>{i+1}</td><td>{n}</td><td>{v}</td></tr>"
                           for i, (n, v) in enumerate(top_c))
        red_html = (
            f"<table class='data-table'><tr><th>Métrica</th><th>Valor</th></tr>{items}</table>"
            f"<h3>Nodos más centrales</h3>"
            f"<table class='data-table'><tr><th>#</th><th>Entidad</th><th>Centralidad</th></tr>"
            f"{top_rows}</table>"
        )

    # ── Narrativas ────────────────────────────────────────────────────────────
    nar = narrativas or {}
    nar_corpus = nar.get("corpus", "")
    nar_red    = nar.get("red", "")
    nar_tono   = nar.get("tono", "")

    # ── Stats rápidas ─────────────────────────────────────────────────────────
    n_pdfs   = stats_corpus.get("n_pdfs", "—")
    n_pags   = stats_corpus.get("n_paginas", "—")
    n_arts   = stats_corpus.get("n_articulos", "—")
    n_words  = stats_corpus.get("n_palabras_total", "—")
    n_ents   = sum(len(v) for v in indice_ner.values())

    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Reporte — {proyecto_nombre}</title>
<style>
  *, *::before, *::after {{ box-sizing: border-box; }}
  body {{ font-family: 'Georgia', serif; background: #f8f6f0; color: #2c2c2c;
          margin: 0; padding: 0; }}
  header {{ background: #1a1a2e; color: #e0e0e0; padding: 3rem 4rem 2rem; }}
  header h1 {{ font-size: 2.2rem; margin: 0 0 0.3rem; color: #a78bfa; }}
  header p {{ margin: 0; color: #94a3b8; font-size: 1rem; }}
  .container {{ max-width: 900px; margin: 0 auto; padding: 2rem 1.5rem; }}
  section {{ margin-bottom: 3rem; padding-bottom: 2rem;
              border-bottom: 1px solid #e0d9d0; }}
  h2 {{ font-size: 1.5rem; color: #1a1a2e; border-left: 4px solid #a78bfa;
        padding-left: 0.8rem; margin-top: 2rem; }}
  h3 {{ font-size: 1.1rem; color: #374151; margin-top: 1.5rem; }}
  .stat-grid {{ display: grid; grid-template-columns: repeat(auto-fit, minmax(150px,1fr));
                gap: 1rem; margin: 1.5rem 0; }}
  .stat-card {{ background: #1a1a2e; color: #e0e0e0; border-radius: 8px;
                padding: 1rem; text-align: center; }}
  .stat-card .num {{ font-size: 2rem; font-weight: bold; color: #a78bfa; }}
  .stat-card .lbl {{ font-size: 0.8rem; color: #94a3b8; margin-top: 4px; }}
  .data-table {{ width: 100%; border-collapse: collapse; margin: 1rem 0;
                  font-size: 0.9rem; }}
  .data-table th {{ background: #1a1a2e; color: #a78bfa; padding: 0.5rem 0.8rem;
                     text-align: left; }}
  .data-table td {{ padding: 0.4rem 0.8rem; border-bottom: 1px solid #e0d9d0; }}
  .data-table tr:hover td {{ background: #f0ece4; }}
  blockquote {{ background: #ede9fe; border-left: 4px solid #7c3aed;
                padding: 1rem 1.2rem; margin: 1rem 0; border-radius: 0 8px 8px 0;
                font-style: italic; color: #374151; }}
  footer {{ text-align: center; padding: 2rem; color: #9ca3af;
             font-size: 0.85rem; background: #1a1a2e; }}
</style>
</head>
<body>
<header>
  <h1>{proyecto_nombre}</h1>
  <p>Análisis computacional · Revista <em>Estampa</em> 1930-1940 · {fecha}</p>
</header>

<div class="container">

  <!-- Estadísticas generales -->
  <section>
    <h2>Resumen del corpus</h2>
    <div class="stat-grid">
      <div class="stat-card"><div class="num">{n_pdfs}</div><div class="lbl">PDFs procesados</div></div>
      <div class="stat-card"><div class="num">{n_pags}</div><div class="lbl">Páginas</div></div>
      <div class="stat-card"><div class="num">{n_arts}</div><div class="lbl">Artículos</div></div>
      <div class="stat-card"><div class="num">{n_words}</div><div class="lbl">Palabras</div></div>
      <div class="stat-card"><div class="num">{n_ents}</div><div class="lbl">Entidades NER</div></div>
    </div>
    {f'<blockquote>{nar_corpus}</blockquote>' if nar_corpus else ''}
  </section>

  <!-- Índice NER -->
  <section>
    <h2>Índice de entidades nombradas</h2>
    {ner_html or '<p>Sin datos NER.</p>'}
  </section>

  <!-- Tono editorial -->
  {'<section><h2>Tono editorial</h2>' + tono_html + (f'<blockquote>{nar_tono}</blockquote>' if nar_tono else '') + '</section>' if tono_html else ''}

  <!-- Redes -->
  {'<section><h2>Redes de co-ocurrencia</h2>' + red_html + (f'<blockquote>{nar_red}</blockquote>' if nar_red else '') + '</section>' if red_html else ''}

</div>

<footer>
  Generado por Bashkar Station · Instituto Caro y Cuervo · {fecha}
</footer>
</body>
</html>"""

    log(f"Escribiendo reporte HTML: {ruta}")
    ruta.write_text(html, encoding="utf-8")
    return ruta


# ── Exportación Word (.docx) ─────────────────────────────────────────────────

def exportar_word(
    proyecto_nombre: str,
    stats_corpus: dict,
    indice_ner: dict,
    narrativas: dict | None,
    ruta: Path,
) -> Path:
    """
    Exporta un reporte básico en formato Word (.docx).
    Requiere: python-docx.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise ImportError("Instala python-docx: pip install python-docx>=1.1.0")

    ruta = Path(ruta)
    ruta.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Título
    doc.add_heading(proyecto_nombre, 0)
    doc.add_paragraph(
        f"Análisis computacional · Revista Estampa 1930-1940 · "
        f"{datetime.now().strftime('%d/%m/%Y')}"
    )
    doc.add_paragraph()

    # Estadísticas
    doc.add_heading("Resumen del corpus", level=1)
    tabla = doc.add_table(rows=1, cols=2)
    tabla.style = "Table Grid"
    hdr = tabla.rows[0].cells
    hdr[0].text = "Indicador"
    hdr[1].text = "Valor"
    for key, val in stats_corpus.items():
        row = tabla.add_row().cells
        row[0].text = str(key).replace("_", " ").capitalize()
        row[1].text = str(val)
    doc.add_paragraph()

    # Narrativa corpus
    nar = narrativas or {}
    if nar.get("corpus"):
        doc.add_heading("Interpretación del corpus", level=2)
        doc.add_paragraph(nar["corpus"])
        doc.add_paragraph()

    # NER
    doc.add_heading("Entidades más frecuentes", level=1)
    cat_labels = {
        "personas": "Personas",
        "lugares": "Lugares",
        "organizaciones": "Organizaciones",
    }
    for cat, label in cat_labels.items():
        if cat not in indice_ner or not isinstance(indice_ner[cat], dict) or not indice_ner[cat]:
            continue
        doc.add_heading(label, level=2)
        top = sorted(indice_ner[cat].items(), key=lambda x: len(x[1]), reverse=True)[:10]
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Entidad"
        t.rows[0].cells[1].text = "N° artículos"
        for ent, arts in top:
            row = t.add_row().cells
            row[0].text = ent
            row[1].text = str(len(arts))
        doc.add_paragraph()

    doc.save(str(ruta))
    return ruta
