#!/usr/bin/env python3
"""
================================================================================
 Conversor masivo de PDF (con OCR ya incrustado) a Word + TXT
 Módulo para Bashkar Station
================================================================================
Toma uno o varios PDF que YA traen capa de texto (el OCR ya está hecho) y, de
forma automática, extrae ese texto y lo vuelca a Word y TXT, organizando todo
por carpetas.

  • NO rasteriza páginas ni vuelve a hacer OCR.
  • El consumo de memoria es mínimo y CONSTANTE aunque los escaneos pesen mucho:
    se procesa una página a la vez y solo se manejan cadenas de texto.
  • 900 páginas se convierten en segundos/minutos, sin estar dando clics.

IMPORTANTE: el Word contiene el TEXTO de la capa OCR (una transcripción), no una
copia visual de las páginas escaneadas. Reproducir las imágenes dentro de Word
es justamente lo que satura la máquina y genera archivos enormes; por eso no se
hace. (Si algún día necesitas el Word "con la pinta" del original, eso es otra
ruta, más pesada, y conviene hacerla por lotes pequeños.)

Estructura generada por cada PDF de entrada:

    <salida>/<nombre_del_pdf>/
        01_PDF_Original/      copia del PDF de origen
        02_PDF_Fragmentado/   un PDF por página           (opcional)
        03_Word_OCR/          Word con el texto: 1 por PDF (+ 1 por página opc.)
        04_TXT_Bashkar/       TXT con el texto:  1 por PDF (+ 1 por página opc.)

Dependencias (¡sin Tesseract!):
    pip install pymupdf python-docx

Uso como script:
    python conversor_pdf_a_word.py "C:/ruta/pdfs" "C:/ruta/salida"
    python conversor_pdf_a_word.py entrada/ salida/ --por-pagina
    python conversor_pdf_a_word.py entrada/ salida/ --sin-fragmentar --solo-txt

Uso como módulo (integración en Bashkar Station, ver el final del archivo):
    from conversor_pdf_a_word import ConfiguracionConversor, ConversorPDFaWord
    cfg = ConfiguracionConversor(carpeta_entrada="...", carpeta_salida="...")
    ConversorPDFaWord(cfg, callback_progreso=mi_callback).procesar_todo()
================================================================================
"""

from __future__ import annotations

import argparse
import json
import logging
import re
import shutil
import sys
import time
import traceback
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Callable

try:
    import fitz  # PyMuPDF
except ImportError:
    raise SystemExit("Falta PyMuPDF. Instala con:  pip install pymupdf")

try:
    from core.ocr_normalizer import limpiar_coordenadas_bnc, normalizar_texto_ocr
    _NORMALIZER_OK = True
except ImportError:
    _NORMALIZER_OK = False

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    raise SystemExit("Falta python-docx. Instala con:  pip install python-docx")


# =============================================================================
#  Configuración
# =============================================================================
@dataclass
class ConfiguracionConversor:
    """Parámetros del proceso. Por defecto genera UN Word y UN TXT por cada PDF
    (lo natural al 'convertir el PDF a Word'); con `por_pagina` añade además un
    archivo por página, como en tu flujo manual."""

    carpeta_entrada: Path                       # carpeta con los PDF de origen
    carpeta_salida: Path                        # raíz donde se crea la estructura

    # --- Extracción de texto ---
    # "texto"    : respeta el orden de la capa de texto (recomendado).
    # "ordenado" : reordena bloques por posición (a veces ayuda, a veces
    #              entrevera columnas; probar solo si el orden sale raro).
    modo_texto: str = "texto"

    # --- Qué generar ---
    fragmentar_pdf: bool = True                 # crea un PDF por página
    word_consolidado: bool = True               # un .docx por PDF
    word_por_pagina: bool = False               # además, un .docx por página
    txt_consolidado: bool = True                # un .txt por PDF (ideal Bashkar)
    txt_por_pagina: bool = False                # además, un .txt por página
    copiar_original: bool = True                # copia el PDF a 01_PDF_Original

    # --- Limpieza de texto post-extracción ---
    # True: elimina coordenadas BNC, sello de biblioteca, une guiones de columna
    # y aplica normalización Unicode. Recomendado para corpus BNC / Paper Capture.
    limpiar_texto: bool = True

    # --- Integración con pipeline Bashkar Station ---
    # Si True, además de los archivos en 04_TXT_Bashkar/, escribe una copia
    # de cada página como p0001.txt, p0002.txt... en:
    #   carpeta_out_dir / "03_ocr" / <nombre_pdf> /
    # Esto permite abrir el resultado directamente en el módulo Normalizar.
    exportar_para_normalizar: bool = False
    carpeta_out_dir: Path | None = None   # = ST.out_dir de la app

    # --- Filtros opcionales (1-indexados, inclusivos) ---
    paginas_desde: int | None = None
    paginas_hasta: int | None = None

    # --- Nombres de subcarpetas ---
    subcarpetas: dict = field(default_factory=lambda: {
        "original": "01_PDF_Original",
        "fragmentado": "02_PDF_Fragmentado",
        "word": "03_Word_OCR",
        "txt": "04_TXT_Bashkar",
    })

    def __post_init__(self):
        self.carpeta_entrada = Path(self.carpeta_entrada)
        self.carpeta_salida  = Path(self.carpeta_salida)
        if self.carpeta_out_dir is not None:
            self.carpeta_out_dir = Path(self.carpeta_out_dir)


# =============================================================================
#  Utilidades de texto / Word
# =============================================================================
def _slug(texto: str) -> str:
    """Nombre de carpeta seguro a partir del nombre del PDF."""
    texto = re.sub(r"[^\w\s.\-]", "", texto, flags=re.UNICODE).strip()
    texto = re.sub(r"\s+", "_", texto)
    return texto or "revista"


def _config_estilo(doc: "Document") -> None:
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)


def _agregar_texto(doc: "Document", texto: str) -> None:
    """Añade el texto al documento preservando los saltos de línea.

    Cada bloque separado por una línea en blanco se vuelve un párrafo; dentro
    de cada párrafo los saltos de línea se mantienen como saltos suaves. (No se
    vuelca '\\n' crudo en un único run, que Word ignoraría.)"""
    if not texto.strip():
        doc.add_paragraph("[Página sin texto en la capa OCR]")
        return
    for bloque in re.split(r"\n[ \t]*\n", texto):
        p = doc.add_paragraph()
        for i, linea in enumerate(bloque.split("\n")):
            if i:
                p.add_run().add_break()
            p.add_run(linea)


def callback_consola(evento: dict) -> None:
    """Callback por defecto: imprime el progreso en consola."""
    pct = evento.get("porcentaje")
    pref = f"[{pct:5.1f}%] " if pct is not None else ""
    print(f"{pref}{evento.get('mensaje', '')}", flush=True)


# =============================================================================
#  Conversor principal
# =============================================================================
class ConversorPDFaWord:
    """Orquesta la conversión por lotes.

    callback_progreso(evento: dict) recibe en cada página un diccionario con:
        pdf, indice_pdf, total_pdfs, pagina, total_paginas, porcentaje, mensaje
    Pensado para refrescar una barra de progreso de Tkinter desde un hilo
    (ver el ejemplo de integración al final del archivo)."""

    def __init__(
        self,
        config: ConfiguracionConversor,
        callback_progreso: Callable[[dict], None] | None = None,
        logger: logging.Logger | None = None,
    ):
        self.cfg = config
        self.cb = callback_progreso or callback_consola
        self.log = logger or self._crear_logger()
        self._cancelado = False  # la GUI puede llamar proc.cancelar()

    # ----------------------------------------------------------------- API ---
    def cancelar(self) -> None:
        """Cancelación ordenada: se detiene tras la página en curso."""
        self._cancelado = True

    def procesar_todo(self) -> dict:
        """Procesa todos los PDF de la carpeta de entrada. Devuelve un reporte."""
        self.cfg.carpeta_salida.mkdir(parents=True, exist_ok=True)
        pdfs = sorted(p for p in self.cfg.carpeta_entrada.glob("*.pdf") if p.is_file())
        if not pdfs:
            raise FileNotFoundError(f"No se encontraron PDF en: {self.cfg.carpeta_entrada}")

        self.log.info("Se procesarán %d PDF.", len(pdfs))
        reporte = {
            "carpeta_entrada": str(self.cfg.carpeta_entrada),
            "carpeta_salida": str(self.cfg.carpeta_salida),
            "config": {k: (str(v) if isinstance(v, Path) else v)
                       for k, v in asdict(self.cfg).items()},
            "pdfs": [],
            "inicio": time.strftime("%Y-%m-%d %H:%M:%S"),
        }
        t0 = time.time()

        for i, ruta_pdf in enumerate(pdfs, start=1):
            if self._cancelado:
                self.log.warning("Proceso cancelado por el usuario.")
                break
            try:
                reporte["pdfs"].append(self._procesar_pdf(ruta_pdf, i, len(pdfs)))
            except Exception as exc:  # un PDF dañado no debe tumbar el lote
                self.log.error("Error procesando %s: %s", ruta_pdf.name, exc)
                self.log.debug(traceback.format_exc())
                reporte["pdfs"].append({"archivo": ruta_pdf.name, "error": str(exc)})

        reporte["fin"] = time.strftime("%Y-%m-%d %H:%M:%S")
        reporte["segundos_totales"] = round(time.time() - t0, 1)
        reporte["cancelado"] = self._cancelado

        ruta_reporte = self.cfg.carpeta_salida / "_reporte_conversion.json"
        ruta_reporte.write_text(json.dumps(reporte, ensure_ascii=False, indent=2),
                                encoding="utf-8")
        self.log.info("Reporte guardado en %s", ruta_reporte)
        return reporte

    # -------------------------------------------------------- por documento ---
    def _procesar_pdf(self, ruta_pdf: Path, indice: int, total: int) -> dict:
        nombre_base = _slug(ruta_pdf.stem)
        carpetas = self._preparar_estructura(nombre_base)

        if self.cfg.copiar_original:
            shutil.copy2(ruta_pdf, carpetas["original"] / ruta_pdf.name)

        doc = fitz.open(str(ruta_pdf))
        n_paginas = doc.page_count
        ini = max(0, (self.cfg.paginas_desde or 1) - 1)
        fin = min(n_paginas, self.cfg.paginas_hasta or n_paginas)
        self.log.info("[%d/%d] %s → %d páginas (procesando %d..%d)",
                      indice, total, ruta_pdf.name, n_paginas, ini + 1, fin)

        paginas_texto: list[tuple[int, str]] = []
        n_vacias = n_fallo = 0

        # Carpeta 03_ocr para integración con módulo Normalizar
        ocr_dir: Path | None = None
        if (self.cfg.exportar_para_normalizar
                and self.cfg.carpeta_out_dir is not None):
            ocr_dir = self.cfg.carpeta_out_dir / "03_ocr" / nombre_base
            ocr_dir.mkdir(parents=True, exist_ok=True)

        for p in range(ini, fin):
            if self._cancelado:
                break
            num = p + 1
            etiqueta = f"pagina_{num:04d}"
            try:
                pagina = doc.load_page(p)

                # 1) Fragmento PDF (una página) — opcional, muy ligero
                if self.cfg.fragmentar_pdf:
                    self._guardar_pagina_pdf(doc, p,
                                             carpetas["fragmentado"] / f"{etiqueta}.pdf")

                # 2) Extraer texto de la capa OCR existente
                texto = self._extraer_texto(pagina)
                if not texto.strip():
                    n_vacias += 1
                paginas_texto.append((num, texto))

                # 3) Word por página (opcional)
                if self.cfg.word_por_pagina:
                    d = Document(); _config_estilo(d)
                    d.add_heading(f"{ruta_pdf.stem} — Página {num}", level=1)
                    _agregar_texto(d, texto)
                    d.save(str(carpetas["word"] / f"{etiqueta}.docx"))

                # 4) TXT por página (opcional)
                if self.cfg.txt_por_pagina:
                    (carpetas["txt"] / f"{etiqueta}.txt").write_text(texto, encoding="utf-8")

                # 5) Exportar a 03_ocr/ para módulo Normalizar (p0001.txt ...)
                if ocr_dir is not None:
                    (ocr_dir / f"p{num:04d}.txt").write_text(texto, encoding="utf-8")

            except Exception as exc:
                n_fallo += 1
                self.log.error("  Página %d falló: %s", num, exc)
                self.log.debug(traceback.format_exc())

            pct = ((indice - 1) + (p - ini + 1) / max(1, fin - ini)) / total * 100
            self.cb({
                "pdf": ruta_pdf.name, "indice_pdf": indice, "total_pdfs": total,
                "pagina": num, "total_paginas": n_paginas, "porcentaje": round(pct, 1),
                "mensaje": f"{ruta_pdf.name} — pág. {num}/{n_paginas}",
            })

        # 5) Consolidados por revista (un solo Word y un solo TXT)
        if self.cfg.word_consolidado:
            d = Document(); _config_estilo(d)
            d.add_heading(f"{ruta_pdf.stem} — Texto OCR", level=0)
            for idx, (num, texto) in enumerate(paginas_texto):
                if idx:
                    d.add_page_break()
                d.add_heading(f"Página {num}", level=2)
                _agregar_texto(d, texto)
            d.save(str(carpetas["word"] / f"{nombre_base}.docx"))

        if self.cfg.txt_consolidado:
            partes = [f"===== Página {num} =====\n\n{texto}" for num, texto in paginas_texto]
            (carpetas["txt"] / f"{nombre_base}.txt").write_text(
                "\n\n".join(partes), encoding="utf-8")

        doc.close()
        return {
            "archivo": ruta_pdf.name,
            "carpeta": str(carpetas["original"].parent),
            "paginas_totales": n_paginas,
            "paginas_procesadas": fin - ini,
            "paginas_sin_texto": n_vacias,
            "paginas_con_error": n_fallo,
        }

    # ----------------------------------------------------------- auxiliares ---
    def _preparar_estructura(self, nombre_base: str) -> dict:
        raiz = self.cfg.carpeta_salida / nombre_base
        carpetas = {}
        for clave, sub in self.cfg.subcarpetas.items():
            if clave == "fragmentado" and not self.cfg.fragmentar_pdf:
                continue
            ruta = raiz / sub
            ruta.mkdir(parents=True, exist_ok=True)
            carpetas[clave] = ruta
        # 'original' siempre existe para devolver la raíz aunque no se copie
        carpetas.setdefault("original", raiz / self.cfg.subcarpetas["original"])
        carpetas["original"].mkdir(parents=True, exist_ok=True)
        return carpetas

    def _guardar_pagina_pdf(self, doc, indice_pagina: int, ruta: Path) -> None:
        nuevo = fitz.open()
        nuevo.insert_pdf(doc, from_page=indice_pagina, to_page=indice_pagina)
        nuevo.save(str(ruta), garbage=4, deflate=True)
        nuevo.close()

    def _extraer_texto(self, pagina) -> str:
        if self.cfg.modo_texto == "ordenado":
            texto = pagina.get_text("text", sort=True).strip()
        else:
            texto = pagina.get_text("text").strip()

        if self.cfg.limpiar_texto and _NORMALIZER_OK:
            # 1. Eliminar coordenadas XY de Adobe Paper Capture (BNC)
            texto = limpiar_coordenadas_bnc(texto)
            # 2. Normalización completa: Unicode NFC, guiones de columna,
            #    chars OCR, sello BNC, dígitos en palabras, ruido tipográfico
            texto = normalizar_texto_ocr(
                texto,
                unir_silabas=True,
                corregir_chars=True,
                corregir_digitos=True,
                corregir_vocab=True,
                limpiar_ruido=True,
                normalizar_unicode=True,
                spell_check=False,
            )
        return texto

    @staticmethod
    def _crear_logger() -> logging.Logger:
        logger = logging.getLogger("ConversorPDFaWord")
        if not logger.handlers:
            logger.setLevel(logging.INFO)
            h = logging.StreamHandler(sys.stdout)
            h.setFormatter(logging.Formatter("%(asctime)s  %(levelname)s  %(message)s", "%H:%M:%S"))
            logger.addHandler(h)
        return logger


# =============================================================================
#  CLI
# =============================================================================
def _construir_parser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(
        description="Convierte por lotes PDF (con OCR ya hecho) a Word + TXT, sin saturar la PC.")
    p.add_argument("entrada", help="Carpeta con los PDF originales.")
    p.add_argument("salida", help="Carpeta raíz donde se crea la estructura.")
    p.add_argument("--por-pagina", action="store_true",
                   help="Generar además un Word y un TXT por cada página.")
    p.add_argument("--sin-fragmentar", action="store_true",
                   help="No crear los PDF página por página.")
    p.add_argument("--solo-word", action="store_true", help="Generar solo Word (no TXT).")
    p.add_argument("--solo-txt", action="store_true", help="Generar solo TXT (no Word).")
    p.add_argument("--orden", choices=["texto", "ordenado"], default="texto",
                   help="Modo de extracción de texto (def: texto).")
    p.add_argument("--desde", type=int, default=None, help="Página inicial (1-indexada).")
    p.add_argument("--hasta", type=int, default=None, help="Página final (1-indexada).")
    return p


def main(argv: list | None = None) -> int:
    a = _construir_parser().parse_args(argv)
    cfg = ConfiguracionConversor(
        carpeta_entrada=a.entrada,
        carpeta_salida=a.salida,
        modo_texto=a.orden,
        fragmentar_pdf=not a.sin_fragmentar,
        word_consolidado=not a.solo_txt,
        word_por_pagina=a.por_pagina and not a.solo_txt,
        txt_consolidado=not a.solo_word,
        txt_por_pagina=a.por_pagina and not a.solo_word,
        paginas_desde=a.desde,
        paginas_hasta=a.hasta,
    )
    rep = ConversorPDFaWord(cfg).procesar_todo()
    print("\n" + "=" * 60 + "\nRESUMEN")
    for d in rep["pdfs"]:
        if "error" in d:
            print(f"  ✗ {d['archivo']}: {d['error']}")
        else:
            print(f"  ✓ {d['archivo']}: {d['paginas_procesadas']} pág. "
                  f"(sin texto: {d['paginas_sin_texto']}, errores: {d['paginas_con_error']})")
    print(f"  Tiempo total: {rep['segundos_totales']} s\n" + "=" * 60)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
