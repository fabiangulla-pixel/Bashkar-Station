"""
╔══════════════════════════════════════════════════════════════════════╗
║  BASHKAR STATION v11.7 — Análisis editorial computacional           ║
║  Aplicación de escritorio · 100% offline · Publicaciones históricas ║
╚══════════════════════════════════════════════════════════════════════╝
"""
import gc
import os
import platform
import queue
import sys
import threading
from datetime import datetime
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from core.bitacora_engine import BitacoraEngine

# ── Rutas Windows (Tesseract/Poppler) ────────────────────────────────────────
_APP_DIR = Path(__file__).parent
_APP_VERSION_SPLASH = "11.6"   # sincronizar con APP_VERSION abajo

def _configurar_rutas_windows():
    for cfg_file in ["tesseract_path.txt", "poppler_path.txt"]:
        cfg = _APP_DIR / cfg_file
        if not cfg.exists():
            continue
        ruta = cfg.read_text(encoding="utf-8").strip()
        p = Path(ruta)
        if not p.exists():
            continue
        # Si la ruta apunta a un ejecutable, agregar su carpeta al PATH
        carpeta = str(p.parent) if p.is_file() else str(p)
        os.environ["PATH"] = carpeta + os.pathsep + os.environ.get("PATH", "")
        if cfg_file.startswith("tesseract"):
            try:
                import pytesseract
                exe = str(p) if p.suffix.lower() == ".exe" else str(p / "tesseract.exe")
                pytesseract.pytesseract.tesseract_cmd = exe
            except ImportError:
                pass

    # ── TESSDATA_PREFIX — buscar tessdata/spa.traineddata ────────────────────
    # Orden de prioridad: carpeta local del usuario → carpeta de instalación
    _tessdata_candidatos = [
        Path.home() / "tessdata",
        Path(r"C:\Users\Lenovo\tessdata"),
        Path(r"C:\Program Files\Tesseract-OCR\tessdata"),
        Path(r"C:\Program Files (x86)\Tesseract-OCR\tessdata"),
    ]
    for _td in _tessdata_candidatos:
        if (_td / "spa.traineddata").exists():
            os.environ["TESSDATA_PREFIX"] = str(_td)
            break

if platform.system() == "Windows":
    _configurar_rutas_windows()

# ── Fijar NumPy < 2 ───────────────────────────────────────────────────────────
def _fijar_numpy():
    import subprocess as _sp
    try:
        import numpy as _np
        v = tuple(int(x) for x in _np.__version__.split(".")[:2])
        if v >= (2, 0):
            _sp.check_call([sys.executable,"-m","pip","install","numpy<2","-q","--force-reinstall"],
                           stdout=_sp.DEVNULL, stderr=_sp.DEVNULL)
    except Exception:
        pass
# _fijar_numpy()  # numpy 2.4 instalado — no forzar downgrade al arranque

# ── Auto-instalación de dependencias ─────────────────────────────────────────
_PAQUETES = [
    ("numpy",      "numpy<2"),
    ("fitz",       "pymupdf>=1.23"),
    ("pdf2image",  "pdf2image>=1.17"),
    ("pytesseract","pytesseract>=0.3.10"),
    ("PIL",        "Pillow>=10.0"),
    ("spacy",      "spacy>=3.7"),
    ("sklearn",    "scikit-learn>=1.4"),
    ("networkx",   "networkx>=3.2"),
    ("matplotlib", "matplotlib>=3.8"),
    ("seaborn",    "seaborn>=0.13"),
    ("pandas",     "pandas>=2.1"),
    ("openpyxl",   "openpyxl>=3.1"),
    ("scipy",      "scipy>=1.12"),
    ("cv2",        "opencv-python-headless>=4.9"),
    # gensim es OPCIONAL: Word2Vec usa backend PyTorch si gensim no está.
    # gensim 4.4 no compila en Python 3.14 (API CPython ob_digit eliminada),
    # por eso NO se auto-instala. En Python ≤3.12 puede instalarse a mano.
]

def _auto_instalar():
    import subprocess
    faltantes = []
    for mod, pkg in _PAQUETES:
        try:
            __import__(mod)
        except ImportError:
            faltantes.append((mod, pkg))
    if not faltantes:
        return
    import tkinter as tk
    from tkinter import ttk
    root = tk.Tk(); root.title(f"Bashkar Station v{_APP_VERSION_SPLASH} — Instalando")
    root.geometry("520x340"); root.configure(bg="#0A1628"); root.resizable(False,False)
    try:
        from PIL import Image, ImageTk
        _sp_img = Image.open(_APP_DIR / "assets" / "logo_splash.png")
        _sp_tk  = ImageTk.PhotoImage(_sp_img)
        tk.Label(root, image=_sp_tk, bg="#0A1628").pack(pady=(12,0))
        root._sp_tk = _sp_tk  # evitar GC
    except Exception:
        tk.Label(root, text=f"BASHKAR STATION v{_APP_VERSION_SPLASH}", bg="#0A1628", fg="white",
                 font=("Segoe UI",16,"bold")).pack(pady=(20,4))
    tk.Label(root, text=f"Instalando {len(faltantes)} paquete(s) faltantes…",
             bg="#0A1628", fg="#7FB3D3", font=("Segoe UI",10)).pack(pady=(0,12))
    lbl = tk.Label(root, text="", bg="#0A1628", fg="white", font=("Courier",10)); lbl.pack()
    prog = ttk.Progressbar(root, length=380, mode="determinate", maximum=len(faltantes)); prog.pack(pady=10)
    lbl_e = tk.Label(root, text="", bg="#0A1628", fg="#70AD47", font=("Segoe UI",9)); lbl_e.pack()
    errores = []
    def run():
        for i,(mod,pkg) in enumerate(faltantes):
            lbl.config(text=f"pip install {pkg}"); root.update()
            try:
                subprocess.check_call([sys.executable,"-m","pip","install",pkg,"-q"],
                                      stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            except Exception:
                errores.append(pkg)
            prog["value"] = i+1; root.update()
        lbl_e.config(text="✅ Listo — abriendo…" if not errores else f"⚠️ Error: {', '.join(errores)}")
        root.after(1500, root.destroy)
    root.after(200, run); root.mainloop()
_auto_instalar()

import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, ttk

import matplotlib

matplotlib.use("TkAgg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

APP_VERSION = "11.7"
APP_NAME    = "Bashkar Station"

# ── Paleta visual v2 — Dark mode académico ────────────────────────────────────
# ── Paletas de color — VS Code Dark+ y VS Code Light+ ────────────────────────
# Colores tomados directamente del tema oficial de VS Code.
# La activity bar es SIEMPRE oscura (#333333) en ambos modos — decisión de diseño
# de VS Code para dar contraste y anclar la navegación visualmente.

_PALETA_DARK = {
    # Activity bar — siempre oscura (#333333 VS Code)
    "AB_BG":     "#333333",
    "AB_SEL":    "#FFFFFF",       # ícono activo: blanco puro
    "AB_TXT":    "#858585",       # ícono inactivo: gris medio
    "AB_IND":    "#0078D4",       # indicador izquierdo: azul VS Code

    # Sidebar (#252526 VS Code Dark+)
    "SB_BG":     "#252526",
    "SB_HOV":    "#2A2D2E",
    "SB_SEL":    "#094771",       # selección azul oscuro
    "SB_TXT":    "#BBBBBB",
    "SB_TXT2":   "#CCCCCC",

    # Editor / contenido (#1E1E1E VS Code Dark+)
    "CONTENT_BG":"#1E1E1E",
    "CARD_BG":   "#252526",
    "CARD_BOR":  "#474747",
    "HDR_LINE":  "#0078D4",

    # Topbar / title bar (#3C3C3C VS Code Dark+)
    "TOPBAR_BG": "#3C3C3C",

    # Paleta funcional VS Code Dark+
    "AZ1":       "#1E1E1E",
    "AZ2":       "#252526",
    "AZ3":       "#0078D4",       # azul VS Code
    "AZ4":       "#4FC1FF",       # azul claro (variables)
    "ACENT":     "#CE9178",       # naranja strings
    "VERDE":     "#4EC9B0",       # verde tipos
    "ROJO":      "#F44747",       # rojo errores
    "GRIS":      "#252526",
    "GRIS2":     "#474747",

    # Texto (#D4D4D4 VS Code Dark+)
    "TXT_PRI":   "#D4D4D4",
    "TXT_SEC":   "#858585",
    "TXT_DIM":   "#5A5A5A",
}

_PALETA_LIGHT = {
    # Activity bar — SIEMPRE oscura en VS Code (#2C2C2C), igual que dark
    "AB_BG":     "#2C2C2C",
    "AB_SEL":    "#FFFFFF",
    "AB_TXT":    "#858585",
    "AB_IND":    "#005FB8",       # azul VS Code Light

    # Sidebar (#F3F3F3 VS Code Light+)
    "SB_BG":     "#F3F3F3",
    "SB_HOV":    "#E8E8E8",
    "SB_SEL":    "#0060C0",
    "SB_TXT":    "#616161",
    "SB_TXT2":   "#383838",

    # Editor / contenido (#FFFFFF VS Code Light+)
    "CONTENT_BG":"#FFFFFF",
    "CARD_BG":   "#F8F8F8",
    "CARD_BOR":  "#E5E5E5",
    "HDR_LINE":  "#005FB8",

    # Topbar (#DDDDDD VS Code Light+)
    "TOPBAR_BG": "#DDDDDD",

    # Paleta funcional VS Code Light+
    "AZ1":       "#FFFFFF",
    "AZ2":       "#F3F3F3",
    "AZ3":       "#005FB8",
    "AZ4":       "#0070C1",
    "ACENT":     "#A31515",       # rojo strings light
    "VERDE":     "#267F99",       # verde tipos light
    "ROJO":      "#CD3131",
    "GRIS":      "#F3F3F3",
    "GRIS2":     "#E5E5E5",

    # Texto (#383838 VS Code Light+)
    "TXT_PRI":   "#383838",
    "TXT_SEC":   "#616161",
    "TXT_DIM":   "#A0A0A0",
}

_MODO_OSCURO = True   # estado global mutable

def _aplicar_paleta(paleta: dict):
    """Inyecta la paleta elegida en las variables globales de color."""
    g = globals()
    for k, v in paleta.items():
        g[k] = v

_aplicar_paleta(_PALETA_DARK)

# ── Variables de color activas (actualizadas por _aplicar_paleta) ─────────────
# Inicializadas con los valores de _PALETA_DARK arriba

# Activity Bar
AB_BG   = "#333333"
AB_HOV  = "#3C3C3C"
AB_SEL  = "#FFFFFF"
AB_TXT  = "#858585"
AB_IND  = "#0078D4"

# Sidebar
SB_BG   = "#252526"
SB_HOV  = "#2A2D2E"
SB_SEL  = "#094771"
SB_TXT  = "#BBBBBB"
SB_TXT2 = "#CCCCCC"

# Contenido
CONTENT_BG = "#1E1E1E"
CARD_BG    = "#252526"
CARD_BOR   = "#474747"
HDR_LINE   = "#0078D4"

# Topbar
TOPBAR_BG  = "#3C3C3C"
TOPBAR_H   = 40

# Paleta funcional
AZ1="#1E1E1E"; AZ2="#252526"; AZ3="#0078D4"; AZ4="#4FC1FF"
ACENT="#CE9178"; VERDE="#4EC9B0"; ROJO="#F44747"
GRIS="#252526"; GRIS2="#474747"

# Texto
TXT_PRI = "#D4D4D4"
TXT_SEC = "#858585"
TXT_DIM = "#5A5A5A"

# Legacy para compatibilidad interna
PALETTE=[AZ1,AZ2,AZ3,AZ4,"#58A6FF","#79C0FF",ACENT,"#F0C070"]

COLABS_DEFAULT = ("Jorge Zalamea\nLeón de Greiff\nGermán Arciniegas\n"
                  "Eduardo Carranza\nHernando Téllez\nLeo Matiz\n"
                  "Gilberto Owen\nFernando Martínez")

CAMPOS_DEFAULT = {
    "Nación":      ["colombia","colombiano","patria","nación","nacional","bogotá","república","gobierno","pueblo"],
    "Modernidad":  ["moderno","modernidad","progreso","técnica","industrial","máquina","radio","cine","automóvil","avión"],
    "Género":      ["mujer","mujeres","femenino","familia","hogar","moda","maternidad","belleza","matrimonio"],
    "Ciudad":      ["ciudad","urbano","calle","barrio","edificio","capital","plaza","parque","comercio"],
    "Guerra/Eur.": ["guerra","europa","español","alemania","fascismo","exilio","refugiado","francia","nazismo"],
    "Cultura":     ["literatura","arte","poesía","novela","música","teatro","escritor","artista","libro"],
}


# ══════════════════════════════════════════════════════════════════════════════
# ESTADO GLOBAL
# ══════════════════════════════════════════════════════════════════════════════
class Estado:
    def __init__(self):
        self.reset()
    def reset(self):
        self.publicacion   = "Mi publicación"
        self.periodo       = ""
        self.pdf_dir       = None
        self.out_dir       = None
        self.input_tipo    = "pdf"
        self.archivos_sel  = []
        self.modos_detec   = {}
        self.etz_done      = False
        self.ocr_done      = False
        self.norm_done     = False
        self.norm_version  = "manual"  # "crudo" | "manual" | "ia"
        self.seg_done      = False
        self.anal_done     = False
        self.vis_done      = False
        self.comp_done     = False
        self.corpus_meta   = None
        self.df_articulos  = None
        self.df_firmas     = None
        self.df_secciones  = None
        self.df_campos     = None
        self.df_layout     = None
        self.df_temas      = None
        self.df_doc_temas  = None
        self.graph_path    = None
        self.xlsx_path     = None
        self.word_model    = None
        self.campos_expandidos = {}
        self.datos_visual  = {}
        self.datos_imagenes = {}
        self.datos_comparativo = {}
        self.figuras       = {}
        self.api_key         = ""   # clave activa (compatibilidad legado)
        self.api_keys        = {    # claves por proveedor
            "anthropic": "", "openai": "", "gemini": "", "ollama": "",
        }
        # Switch global: False = modo 100% offline, ninguna función llama APIs externas
        self.ia_habilitada   = False
        # Modelo elegido por etapa. Valores: "<proveedor>/<modelo>"
        self.modelos_etapa   = {
            "ocr_mejora":   "ollama/llava",
            "ner":          "ollama/mistral",
            "deteccion":    "ollama/llava",
            "tono":         "ollama/mistral",
            "narrativas":   "ollama/mistral",
            "asistente":    "ollama/mistral",
        }
        self.max_ia          = 15
        self.campos_semillas = {}
        # Resultados de análisis (inicializados explícitamente para type-safety)
        self.resumen_ocr     = None   # dict con stats de extracción
        self.temas_lda       = None   # list de dicts con palabras por tema
        self.matriz_sim      = None   # DataFrame de similitud comparativa
        self.terminos_dist   = {}     # dict {nombre: [términos distinctivos]}
        self.ner_done        = False
        self.indice_ner_global = {}  # {categoria: {entidad: [art_ids]}}
        self.stopwords_proyecto = []  # palabras extra a filtrar en análisis léxico
        self.lematizar       = True   # False = usar formas originales (corpus histórico)
        self.corpus_txt      = []    # [str] textos planos listos para módulo lingüística
        # v11 DB
        self.repo           = None  # instancia datos.repositorio.Repositorio
        self.ruta_db        = ""    # ruta al archivo .db SQLite
        self.wikidata_enlaces = {}  # {cat: {texto: {id,label,description,url,confianza}}}
        # v17 — comparador
        self.comparar_rutas    = []   # list of str paths for comparison
        self.reporte_comparativo = {} # dict result
        # v17 — intertextualidad
        self.intertex_resultado = {}
        # v18 — confianza
        self.confianza_corpus  = {}
        # v19 — colaboracion
        self.colaboracion_parche = None  # dict parche cargado
        # v20
        self.pptx_path         = None
        # Etiquetador — prompt de detección editable por proyecto
        self.prompt_deteccion  = ""     # vacío = usar prompt por defecto de zone_labeler
        # Semáforos de flujo: "pending" | "ready" | "stale"
        # "ready"   = etapa completada y datos al día
        # "stale"   = completada pero una etapa anterior cambió → re-ejecutar
        # "pending" = nunca ejecutada o reseteada
        self.estado_etapas: dict = {
            "etz":  "pending",
            "ocr":  "pending",
            "norm": "pending",
            "seg":  "pending",
            "anal": "pending",
        }

    def marcar_etapa(self, etapa: str, estado: str):
        """Marca una etapa y propaga 'stale' a todas las etapas posteriores."""
        _orden = ["etz", "ocr", "norm", "seg", "anal"]
        self.estado_etapas[etapa] = estado
        if estado in ("ready", "stale"):
            try:
                idx = _orden.index(etapa)
            except ValueError:
                return
            for posterior in _orden[idx + 1:]:
                if self.estado_etapas.get(posterior) == "ready":
                    self.estado_etapas[posterior] = "stale"

ST = Estado()


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS API / MODELOS
# ══════════════════════════════════════════════════════════════════════════════

def _ocr_vision_multiproveedor(img_path, proveedor: str,
                                api_key: str, modelo: str = "") -> str:
    """
    Extrae el texto de una imagen de página usando cualquier proveedor de visión IA.
    Retorna el texto transcrito como string.
    """
    import base64
    from pathlib import Path as _Path

    _PROMPT_OCR = (
        "Transcribe todo el texto visible en esta imagen de página de revista histórica "
        "(Colombia, años 1930-1940). Respeta el orden de lectura: columna izquierda de "
        "arriba a abajo, luego columna derecha. Preserva los saltos de párrafo. "
        "No añadas comentarios ni explicaciones — solo el texto transcrito."
    )

    img_path = _Path(img_path)
    ext = img_path.suffix.lower()
    mt_map = {".jpg": "image/jpeg", ".jpeg": "image/jpeg",
               ".png": "image/png",  ".webp": "image/webp"}
    mt = mt_map.get(ext, "image/png")
    with open(img_path, "rb") as f:
        b64 = base64.standard_b64encode(f.read()).decode()

    if proveedor == "claude":
        import anthropic
        m = modelo or "claude-haiku-4-5-20251001"
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model=m, max_tokens=4096,
            messages=[{"role": "user", "content": [
                {"type": "image", "source": {"type": "base64",
                                              "media_type": mt, "data": b64}},
                {"type": "text", "text": _PROMPT_OCR},
            ]}])
        return resp.content[0].text

    elif proveedor == "openai":
        import openai
        m = modelo or "gpt-4o-mini"
        client = openai.OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=m, max_tokens=4096,
            messages=[{"role": "user", "content": [
                {"type": "image_url",
                 "image_url": {"url": f"data:{mt};base64,{b64}"}},
                {"type": "text", "text": _PROMPT_OCR},
            ]}])
        return resp.choices[0].message.content

    elif proveedor == "gemini":
        import google.generativeai as genai
        from PIL import Image as _Img
        genai.configure(api_key=api_key)
        m = modelo or "gemini-1.5-flash"
        gm = genai.GenerativeModel(m)
        img = _Img.open(img_path)
        resp = gm.generate_content([_PROMPT_OCR, img])
        return resp.text

    elif proveedor == "ollama":
        import requests
        m = modelo or "llava"
        resp = requests.post(
            "http://localhost:11434/api/generate",
            json={"model": m, "prompt": _PROMPT_OCR,
                  "images": [b64], "stream": False},
            timeout=300)
        return resp.json().get("response", "")

    return ""


def _resolver_api_key_modelo(etapa: str) -> tuple[str, str]:
    """
    Devuelve (api_key, modelo_id) para la etapa indicada.
    Retorna ("", "") si ST.ia_habilitada es False (modo offline).

    Lógica:
    1. Si ST.ia_habilitada es False → retorna ("", "") para bloquear cualquier llamada a API
    2. Lee ST.modelos_etapa[etapa] → "proveedor/modelo"
    3. Busca la api_key en ST.api_keys[proveedor]
    4. Si no hay clave específica, cae a ST.api_key (legado)
    """
    if not getattr(ST, "ia_habilitada", False):
        return "", ""

    modelo_full = ST.modelos_etapa.get(etapa, "")
    if "/" in modelo_full:
        proveedor, modelo_id = modelo_full.split("/", 1)
    else:
        proveedor, modelo_id = "anthropic", modelo_full

    # Ollama no necesita API key — devuelve la URL del servidor como "key"
    if proveedor == "ollama":
        api_key = ST.api_keys.get("ollama", "http://localhost:11434").strip()
        return api_key or "http://localhost:11434", modelo_id

    api_key = ST.api_keys.get(proveedor, "").strip()
    if not api_key:
        api_key = ST.api_key  # fallback legado

    return api_key, modelo_id


# ══════════════════════════════════════════════════════════════════════════════
# ESTILOS ttk
# ══════════════════════════════════════════════════════════════════════════════
def _estilos():
    s = ttk.Style(); s.theme_use("clam")
    # Base dark
    s.configure(".", background=CONTENT_BG, foreground=TXT_PRI,
                font=("Segoe UI", 10))
    # Botón primario — azul GitHub
    s.configure("P.TButton", background=AZ3, foreground="white",
                font=("Segoe UI", 10, "bold"), padding=[18, 8],
                borderwidth=0, relief="flat")
    s.map("P.TButton",
          background=[("active","#388BFD"), ("disabled", TXT_DIM)],
          foreground=[("disabled", TXT_SEC)])
    # Botón secundario — gris oscuro
    s.configure("S.TButton", background=CARD_BOR, foreground=TXT_PRI,
                font=("Segoe UI", 9), padding=[12, 6],
                borderwidth=0, relief="flat")
    s.map("S.TButton", background=[("active","#444C56"), ("disabled", AZ2)])
    # Botón acento — naranja
    s.configure("A.TButton", background=ACENT, foreground="#0D1117",
                font=("Segoe UI", 9, "bold"), padding=[12, 6],
                borderwidth=0, relief="flat")
    s.map("A.TButton", background=[("active","#DB6D28")])
    # Botón éxito — verde
    s.configure("OK.TButton", background=VERDE, foreground="#0D1117",
                font=("Segoe UI", 9, "bold"), padding=[12, 6],
                borderwidth=0, relief="flat")
    s.map("OK.TButton", background=[("active","#2EA043")])
    # Etiquetas
    s.configure("H.TLabel",    background=CONTENT_BG, foreground=TXT_PRI,
                font=("Segoe UI", 13, "bold"))
    s.configure("H2.TLabel",   background=CARD_BG,    foreground=TXT_PRI,
                font=("Segoe UI", 11, "bold"))
    s.configure("Sub.TLabel",  background=CONTENT_BG, foreground=TXT_SEC,
                font=("Segoe UI", 9))
    s.configure("Card.TLabel", background=CARD_BG,    foreground=TXT_PRI,
                font=("Segoe UI", 10))
    # Progreso
    s.configure("TProgressbar", troughcolor=CARD_BOR, background=AZ3,
                thickness=6, borderwidth=0)
    # Notebook tabs
    s.configure("TNotebook", background=CARD_BG, borderwidth=0)
    s.configure("TNotebook.Tab", background=AZ2, foreground=TXT_SEC,
                font=("Segoe UI", 9), padding=[12, 5])
    s.map("TNotebook.Tab",
          background=[("selected", CARD_BG)],
          foreground=[("selected", TXT_PRI)])
    # Frames y separadores
    s.configure("TLabelframe", background=CARD_BG, foreground=TXT_PRI,
                bordercolor=CARD_BOR, borderwidth=1, relief="solid",
                font=("Segoe UI", 10, "bold"))
    s.configure("TLabelframe.Label", background=CARD_BG, foreground=AZ4,
                font=("Segoe UI", 10, "bold"))
    s.configure("TCheckbutton", background=CARD_BG, foreground=TXT_PRI)
    s.configure("TRadiobutton", background=CARD_BG, foreground=TXT_PRI)
    s.configure("TSeparator", background=CARD_BOR)
    # Combobox
    s.configure("TCombobox", background=CARD_BG, foreground=TXT_PRI,
                fieldbackground=CARD_BG, selectbackground=AZ3,
                arrowcolor=TXT_SEC, bordercolor=CARD_BOR)
    # Spinbox
    s.configure("TSpinbox", background=CARD_BG, foreground=TXT_PRI,
                fieldbackground=CARD_BG, bordercolor=CARD_BOR,
                arrowcolor=TXT_SEC)
    # Treeview
    s.configure("Treeview", background=CARD_BG, foreground=TXT_PRI,
                fieldbackground=CARD_BG, rowheight=28,
                font=("Segoe UI", 9), borderwidth=0)
    s.configure("Treeview.Heading", background=AZ2, foreground=TXT_SEC,
                font=("Segoe UI", 9, "bold"), relief="flat", borderwidth=0)
    s.map("Treeview",
          background=[("selected", AZ3)],
          foreground=[("selected", "white")])


# ══════════════════════════════════════════════════════════════════════════════
# WIDGET DE SCROLL CORRECTO
# (sin bind_all que causa conflictos entre paneles)
# ══════════════════════════════════════════════════════════════════════════════

def _hacer_scrollable(frame_padre, bg=CONTENT_BG):
    """
    Crea un Canvas con Scrollbar vertical dentro de frame_padre.
    El scroll con rueda solo se activa cuando el ratón está sobre el canvas.
    Devuelve el Frame interior donde poner los widgets.
    """
    canvas = tk.Canvas(frame_padre, bg=bg, highlightthickness=0,
                       borderwidth=0)
    vsb    = ttk.Scrollbar(frame_padre, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=vsb.set)

    vsb.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    inner = tk.Frame(canvas, bg=bg)
    wid   = canvas.create_window((0, 0), window=inner, anchor="nw")

    def _on_inner_configure(e):
        canvas.configure(scrollregion=canvas.bbox("all"))
    inner.bind("<Configure>", _on_inner_configure)

    def _on_canvas_configure(e):
        canvas.itemconfig(wid, width=e.width)
    canvas.bind("<Configure>", _on_canvas_configure)

    # Rueda de ratón — activa en canvas E hijos (propagación completa)
    def _scroll(e):
        delta = -1 if (e.delta > 0 or e.num == 4) else 1
        canvas.yview_scroll(delta, "units")

    def _bind_scroll_widget(w):
        """Propaga scroll a todos los widgets hijos del inner frame."""
        try:
            w.bind("<MouseWheel>", _scroll)
            w.bind("<Button-4>",   _scroll)
            w.bind("<Button-5>",   _scroll)
        except Exception:
            pass
        for child in w.winfo_children():
            _bind_scroll_widget(child)

    # Bind inmediato al canvas
    canvas.bind("<MouseWheel>", _scroll)
    canvas.bind("<Button-4>",   _scroll)
    canvas.bind("<Button-5>",   _scroll)
    inner.bind("<MouseWheel>",  _scroll)
    inner.bind("<Button-4>",    _scroll)
    inner.bind("<Button-5>",    _scroll)

    # Bind diferido a hijos (se crean después del return)
    def _bind_hijos_diferido():
        _bind_scroll_widget(inner)
        # Repetir cada 2s para capturar widgets añadidos dinámicamente
        canvas.after(2000, _bind_hijos_diferido)
    canvas.after(500, _bind_hijos_diferido)

    return inner, canvas


# ══════════════════════════════════════════════════════════════════════════════
# VENTANA PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════
# ── Prompts de asistente IA sugeridos por pestaña ────────────────────────────
# Orientados a estudios editoriales colombianos: historia de la prensa,
# análisis del discurso, estudios culturales, sociología del campo editorial.
# Fuentes de referencia: Loaiza Cano, Osorio Tejeda, Riaño, Silva Olarte,
# Colombia 100 años de revistas (Banco de la República).

_AI_PROMPTS = {
    "ocr": [
        ("Calidad de transcripción",
         "Analiza el texto extraído e identifica los principales problemas de calidad OCR: "
         "palabras mal reconocidas, errores sistemáticos de caracteres confundidos (I/l, 0/O, rn/m), "
         "y fragmentos ilegibles. Estima qué porcentaje del texto es confiable para análisis cuantitativo."),
        ("Vocabulario de época",
         "Construye un glosario de términos propios del período a partir del texto extraído: "
         "arcaísmos, neologismos de los años 30-40, términos técnicos de la industria editorial "
         "colombiana y colombianismos. Indica frecuencia y contexto de aparición de cada término."),
        ("Patrones tipográficos y editoriales",
         "Identifica en el texto los patrones de presentación editorial: secciones fijas recurrentes, "
         "fórmulas de apertura y cierre de artículos, convenciones de citación y pie de foto. "
         "¿Qué revelan sobre las prácticas editoriales de la publicación?"),
        ("Variantes del español colombiano",
         "Compara la ortografía y gramática con la norma del español colombiano de los años 30-40. "
         "¿Hay variantes regionales, influencias del español peninsular o rasgos bogotanos cultos? "
         "Distingue entre errores OCR y usos propios de la época."),
    ],
    "seg": [
        ("Autoría anónima y prácticas de firma",
         "Analiza los artículos sin autor identificado. ¿Qué tipos de contenido se publican sin firma "
         "(editoriales, notas, traducciones)? Contrasta con los patrones de firma en la prensa "
         "colombiana de los años 30 y discute las implicaciones para la atribución de autoría."),
        ("Campo intelectual y redes de colaboradores",
         "A partir de los autores identificados, reconstruye el campo intelectual de la publicación. "
         "¿Hay colaboradores recurrentes? ¿Se pueden identificar redes de intelectuales, círculos "
         "literarios o afiliaciones ideológicas (liberalismo lopista, conservatismo, vanguardias)?"),
        ("Distribución de géneros periodísticos y literarios",
         "Clasifica los artículos segmentados por géneros (crónica, editorial, cuento, poema, ensayo, "
         "reportaje, nota social). ¿Cómo refleja esta distribución el proyecto editorial y el "
         "posicionamiento de la publicación en el mercado de revistas colombianas?"),
        ("Extensión, espacio y jerarquía editorial",
         "Analiza la distribución de longitud de los artículos. ¿Qué géneros o autores reciben más "
         "espacio? ¿Hay correlación entre extensión y posición en la página? Discute qué dice esto "
         "sobre la jerarquía de valores editoriales de la publicación."),
    ],
    "anal": [
        ("Posicionamiento ideológico y discurso político",
         "A partir de los temas LDA y entidades nombradas, analiza el posicionamiento ideológico. "
         "¿Cómo se refieren a los partidos políticos, al Estado, a la Iglesia? ¿Se detectan "
         "alineaciones con el liberalismo lopista, el conservatismo o corrientes de izquierda "
         "de los años 30 colombianos?"),
        ("Imaginarios de modernidad y progreso",
         "Identifica cómo la publicación construye el imaginario de la modernidad. ¿Qué referentes "
         "usa (técnica, industria, ciudad, Europa, EE.UU.)? ¿Cómo se articula el discurso del "
         "progreso con la identidad nacional colombiana de la República Liberal?"),
        ("Canon literario e intertextualidad",
         "A partir de las entidades y el vocabulario, identifica referencias a autores, publicaciones "
         "y obras. ¿Qué canon construye la revista? ¿Hay presencia de la vanguardia latinoamericana, "
         "el modernismo tardío o la literatura española? ¿Se cita prensa internacional?"),
        ("Género, mujer y representación social",
         "Analiza cómo aparece la figura femenina: como autora, como tema, como lectora implícita. "
         "¿Hay secciones dedicadas a la mujer? ¿Qué roles se le asignan? Contrasta con la prensa "
         "femenina colombiana de los años 30 (El Hogar, Letras y Encajes, Agitación Femenina)."),
    ],
    "vis": [
        ("Economía de la imagen y financiamiento publicitario",
         "Analiza la proporción de espacio editorial dedicado a publicidad gráfica vs. contenido "
         "editorial. ¿Qué sectores económicos anuncian? ¿Cómo evoluciona la presencia publicitaria? "
         "Discute la relación entre financiamiento publicitario y autonomía editorial en la "
         "prensa comercial colombiana."),
        ("Fotografía y construcción visual de la modernidad",
         "Describe el uso de la fotografía: ¿predominan retratos, eventos sociales, paisajes, "
         "industria? ¿Hay fotógrafos identificados? ¿Cómo se integra con el texto? Discute en "
         "relación con la introducción del fotograbado y el fotoperiodismo en Colombia."),
        ("Ilustración, caricatura y sátira política",
         "Analiza las ilustraciones y caricaturas. ¿Hay ilustradores identificados? ¿Los temas "
         "son costumbristas, políticos o sociales? Relaciona con la tradición de la caricatura "
         "colombiana (Ricardo Rendón, Chapete) y su función en el debate público."),
        ("Tipografía, diseño e identidad editorial",
         "A partir de las fuentes tipográficas identificadas, analiza la identidad gráfica de la "
         "publicación. ¿Hay coherencia tipográfica? ¿Se usan tipos modernos o tradicionales? "
         "¿Cómo se diferencia de otras revistas del período? Discute el diseño como estrategia "
         "de posicionamiento en el mercado editorial colombiano."),
    ],
    "comp": [
        ("Posicionamiento en el campo editorial colombiano",
         "Compara el perfil temático con el corpus de referencia. ¿En qué se diferencia la agenda "
         "editorial? ¿Ocupa un nicho especializado o compite directamente con otras publicaciones? "
         "Relaciona con la estructura del campo editorial colombiano de los años 30 "
         "(El Tiempo, El Espectador, Cromos, Semana)."),
        ("Inflexiones históricas en el discurso editorial",
         "Analiza cómo cambian los temas y el vocabulario entre números o períodos. ¿Hay "
         "inflexiones relacionadas con eventos históricos colombianos (República Liberal, "
         "violencia partidista, Segunda Guerra Mundial, Bogotazo)? ¿Cómo responde la "
         "publicación a la coyuntura política y cultural?"),
        ("Transferencias culturales e influencias externas",
         "A partir de los términos distintivos y entidades, identifica la presencia de referencias "
         "internacionales. ¿Qué literaturas nacionales tienen mayor presencia (francesa, española, "
         "anglosajona, latinoamericana)? ¿Cómo se articulan las influencias externas con el "
         "proyecto editorial nacional?"),
    ],
    "meta": [
        ("Contextualización histórica del registro",
         "A partir de los metadatos (título, fecha, editorial, descripción), elabora una ficha "
         "de contextualización histórica. ¿Qué eventos históricos colombianos o internacionales "
         "son contemporáneos? ¿Hay cambios en la dirección o el perfil editorial? Ubica el "
         "número en la historia de la publicación."),
        ("Análisis del paratexto y aparato editorial",
         "Analiza los elementos paratextuales: subtítulos, lemas, declaraciones de propósito, "
         "índices, sumarios. ¿Cómo se autopresenta la publicación? ¿Qué lector implícito "
         "construye? Usa conceptos de Gerard Genette y la historia del libro."),
        ("Reconstrucción de la cadena editorial",
         "A partir de los metadatos disponibles, reconstruye la cadena editorial: director, "
         "propietario, imprenta, distribución, precio, tiraje si está disponible. ¿Qué dice "
         "esto sobre el modelo de negocio y la sostenibilidad de la publicación en el "
         "mercado editorial colombiano?"),
    ],
    "res": [
        ("Síntesis interpretativa del corpus",
         "Con base en todos los resultados (OCR, segmentación, temas LDA, imágenes, red de autoría), "
         "elabora una síntesis interpretativa del perfil editorial. ¿Qué hipótesis de investigación "
         "surgen? ¿Qué preguntas quedan abiertas para trabajo cualitativo de archivo?"),
        ("Contraste con la historiografía existente",
         "Contrasta los hallazgos cuantitativos con lo que la historiografía dice sobre esta "
         "publicación o el período. ¿Los datos confirman, matizan o contradicen las interpretaciones "
         "existentes? Sugiere líneas de investigación novedosas derivadas del análisis computacional."),
        ("Sección metodológica para publicación académica",
         "Redacta una sección de metodología (500 palabras) que describa el proceso de análisis "
         "computacional realizado, los datos cuantitativos más relevantes, y las decisiones "
         "metodológicas. Calibra el tono para una revista de humanidades digitales o historia "
         "de la prensa latinoamericana."),
    ],
}


class BashkarApp(tk.Tk):

    # ── Definición de páginas del sidebar ─────────────────────────────────────
    # Formato: (id, emoji, label, label_largo, badge_attr, grupo)
    # grupo:
    #   "flujo"     → pasos numerados del flujo principal
    #   "analisis"  → herramientas de análisis (sin número, opcionales)
    #   "salida"    → exportación y colaboración
    _PAGINAS = [
        # ── FLUJO PRINCIPAL (numerado, secuencial) ────────────────────────────
        ("cfg",   "⚙",   "Configuración",  "Configuración del corpus",           None,        "flujo"),
        ("etz",   "✏️",  "Etiquetar",      "Etiquetar zonas de página (opcional)","etz_done",  "flujo"),
        ("ocr",   "📄",  "Extracción OCR", "Extracción de texto por OCR",         "ocr_done",  "flujo"),
        ("conv",  "⚡",  "Conversor PDF",  "Conversión masiva PDF→Word/TXT (texto embebido)", None, "flujo"),
        ("mmx",   "🧠",  "Extracción IA",  "Extracción multimodal estructurada de imágenes (IA de visión → JSON → .md)", None, "flujo"),
        ("norm",  "📝",  "Normalizar",     "Revisión y normalización del texto",  "norm_done", "flujo"),
        ("seg",   "📋",  "Segmentar",      "Segmentación en artículos",           "seg_done",  "flujo"),
        ("anal",  "🔬",  "Analizar",       "Análisis textual y semántico",        "anal_done", "flujo"),
        ("res",   "📈",  "Resultados",     "Resultados y exportación",            None,        "flujo"),
        # ── ANÁLISIS OPCIONALES ────────────────────────────────────────────────
        ("ner",   "🏷",  "Entidades",      "Índice de entidades nombradas",       "ner_done",  "analisis"),
        ("anot",  "✍️", "Anotar",         "Anotación semántica revisable",       None,        "analisis"),
        ("bsem",  "🔍",  "Búsqueda",       "Búsqueda semántica por similitud",    None,        "analisis"),
        ("coloc", "🔤",  "Collocates",     "Redes léxicas y concordancias",       None,        "analisis"),
        ("nov",   "🆕",  "Novedad",        "Detección de novedad y cambio discursivo", None,   "analisis"),
        ("red",   "🕸",  "Redes",          "Redes de co-ocurrencia",              None,        "analisis"),
        ("ling",  "🔭",  "Lingüística",    "Sintaxis, correferencia, morfología, encuadre, polaridad, revisión NER y validación", None, "analisis"),
        ("sem",   "🧠",  "Semántico",      "Tono, léxico y estilo",               None,        "analisis"),
        ("top",   "🧩",  "Tópicos",        "Topic modeling del corpus",           None,        "analisis"),
        ("viz",   "🎨",  "Visualizar",     "Visualizaciones avanzadas",           None,        "analisis"),
        ("comp",  "📊",  "Comparativo",    "Análisis comparativo interno",        "comp_done", "analisis"),
        ("comp2", "🔀",  "Multi-corpus",   "Comparación entre proyectos",         None,        "analisis"),
        ("intxt", "🔗",  "Intertexto",     "Análisis intertextual",               None,        "analisis"),
        ("meta",  "🌐",  "Metadatos URL",  "Metadatos desde URL externa",         None,        "analisis"),
        ("vis",   "🖼",  "Tipografía",     "Visual y tipografía",                 "vis_done",  "analisis"),
        ("imgdesc","🎨", "Desc. imágenes", "Descripción e iconografía de imágenes etiquetadas", None, "analisis"),
        # ── SALIDA Y COLABORACIÓN ──────────────────────────────────────────────
        ("rep",   "📝",  "Reporte",        "Reporte narrativo (IA)",              None,        "salida"),
        ("dash",  "📊",  "Dashboard",      "Dashboard ejecutivo",                 None,        "salida"),
        ("valid", "✅",  "Validar",        "Validación humana y semáforo",        None,        "salida"),
        ("colab", "👥",  "Colaborar",      "Colaboración y trazabilidad",         None,        "salida"),
    ]

    # Lista plana de ids para compatibilidad con código que itera _PAGINAS
    @classmethod
    def _paginas_ids(cls):
        return [p[0] for p in cls._PAGINAS]

    def __init__(self):
        super().__init__()
        self.title(f"{APP_NAME} v{APP_VERSION} — Análisis Editorial Computacional")
        self.geometry("1280x820")
        self.minsize(1024, 680)
        self.configure(bg=SB_BG)
        _estilos()
        self._q             = queue.Queue()
        self._pagina_activa = tk.StringVar(value="cfg")
        self._frames_pagina = {}
        self._sb_btns       = {}
        self._proyecto_ruta = None      # Path del .bashkar activo
        self._historial_ia  = []        # [{tab,prompt,respuesta,fecha}]
        self._hay_cambios   = False     # indica cambios sin guardar
        self._cp_win        = None      # ventana Command Palette
        self._toasts_activos: list = [] # toasts visibles
        self._build_ui()
        self._poll()
        self.protocol("WM_DELETE_WINDOW", self._on_cerrar)
        # Restaurar última sesión
        self.after(200, self._cargar_ultimo_proyecto)
        # Autoguardado periódico cada 3 min
        self.after(180_000, self._autoguardar_periodico)
        # Bind global Ctrl+K → Command Palette
        self.bind_all("<Control-k>", self._abrir_command_palette)

    # ── Cola de mensajes ──────────────────────────────────────────────────────
    def _poll(self):
        try:
            while True:
                m = self._q.get_nowait()
                t = m.get("tipo")
                if   t == "log":  self._log(m["texto"], m.get("color","#94A3B8"))
                elif t == "prog": self._set_prog(m["val"], m.get("txt",""))
                elif t == "fase": self._lbl_fase.config(text=m["txt"])
                elif t == "ok":   self._on_ok(m.get("res"))
                elif t == "err":  messagebox.showerror("Error", m["txt"])
        except queue.Empty:
            pass
        self.after(100, self._poll)

    def _put(self, **kw): self._q.put(kw)

    # ─────────────────────────────────────────────────────────────────────────
    # CONSTRUCCIÓN DE LA UI
    # ─────────────────────────────────────────────────────────────────────────
    def _build_ui(self):
        # ── Layout raíz: topbar + body ────────────────────────────────────────
        self.configure(bg=AB_BG)
        root_frame = tk.Frame(self, bg=AB_BG)
        root_frame.pack(fill="both", expand=True)

        # ── TOPBAR (barra superior fija) ──────────────────────────────────────
        self._topbar = tk.Frame(root_frame, bg=TOPBAR_BG, height=TOPBAR_H)
        self._topbar.pack(side="top", fill="x")
        self._topbar.pack_propagate(False)
        self._build_topbar()

        # ── Separador visual topbar/body ──────────────────────────────────────
        tk.Frame(root_frame, bg=CARD_BOR, height=1).pack(side="top", fill="x")

        # ── BODY: activity bar + sidebar + contenido ──────────────────────────
        body = tk.Frame(root_frame, bg=AB_BG)
        body.pack(side="top", fill="both", expand=True)

        # Activity Bar (60px, íconos)
        self._activity_bar = tk.Frame(body, bg=AB_BG, width=60)
        self._activity_bar.pack(side="left", fill="y")
        self._activity_bar.pack_propagate(False)

        # Separador activity bar / sidebar
        tk.Frame(body, bg=CARD_BOR, width=1).pack(side="left", fill="y")

        # Sidebar de sub-items (200px)
        self._sidebar = tk.Frame(body, bg=SB_BG, width=200)
        self._sidebar.pack(side="left", fill="y")
        self._sidebar.pack_propagate(False)

        # Separador sidebar / contenido
        tk.Frame(body, bg=CARD_BOR, width=1).pack(side="left", fill="y")

        # Área de contenido
        self._content_area = tk.Frame(body, bg=CONTENT_BG)
        self._content_area.pack(side="left", fill="both", expand=True)

        self._build_activity_bar()
        self._build_sidebar()

        # Crear todos los frames de página (apilados, solo uno visible)
        builds = {
            "cfg":  self._build_cfg,
            "ocr":  self._build_ocr,
            "etz":  self._build_etz,
            "conv": self._build_conv,
            "mmx":  self._build_mmx,
            "norm": self._build_norm,
            "seg":  self._build_seg,
            "anal": self._build_anal,
            "vis":     self._build_vis,
            "imgdesc": self._build_imgdesc,
            "comp": self._build_comp,
            "meta": self._build_meta,
            "res":  self._build_res,
            "ner":  self._build_ner,
            "anot": self._build_anot,
            "bsem": self._build_busqueda_semantica,
            "coloc":self._build_coloc,
            "nov":  self._build_nov,
            "red":  self._build_red,
            "ling": self._build_ling,
            "sem":  self._build_sem,
            "top":  self._build_top,
            "viz":  self._build_viz,
            "rep":   self._build_rep,
            "dash":  self._build_dash,
            "comp2": self._build_comp2,
            "intxt": self._build_intxt,
            "valid": self._build_valid,
            "colab": self._build_colab,
        }
        # Cada página vive dentro de un Canvas con scrollbar vertical, para que
        # los módulos altos (guía + estadísticas + opciones + botones) siempre
        # sean navegables aunque no quepan en la ventana. `_frames_pagina[pid]`
        # es el frame INTERIOR desplazable (donde cada _build_* hace pack); el
        # contenedor externo (canvas+scrollbar) se guarda en `_contenedores_pagina`.
        self._contenedores_pagina = {}
        for pid, _, _, _, _, _ in self._PAGINAS:
            frm = self._crear_pagina_scrollable(pid)
            self._frames_pagina[pid] = frm
            # Compatibilidad con código legacy que usa self._tab_XXX
            setattr(self, f"_tab_{pid}", frm)
        # Alias legacy
        self._tab_cfg  = self._frames_pagina["cfg"]
        self._tab_ocr  = self._frames_pagina["ocr"]
        self._tab_norm = self._frames_pagina["norm"]
        self._tab_seg  = self._frames_pagina["seg"]
        self._tab_anal = self._frames_pagina["anal"]
        self._tab_vis  = self._frames_pagina["vis"]
        self._tab_comp = self._frames_pagina["comp"]
        self._tab_meta = self._frames_pagina["meta"]
        self._tab_res  = self._frames_pagina["res"]
        self._tab_ner  = self._frames_pagina["ner"]
        self._tab_red  = self._frames_pagina["red"]
        self._tab_sem  = self._frames_pagina["sem"]
        self._tab_top  = self._frames_pagina["top"]
        self._tab_viz  = self._frames_pagina["viz"]
        self._tab_rep  = self._frames_pagina["rep"]
        self._tab_dash = self._frames_pagina["dash"]

        # Construir contenido. Fijamos el id de página actual ANTES de cada
        # build_fn para que _page_header pueda inyectar la guía del módulo
        # (qué es / para qué / cómo interpretar) sin tocar cada _build_*.
        for pid, build_fn in builds.items():
            self._guia_pagina_actual = pid
            build_fn()
        self._guia_pagina_actual = None

        # Mostrar página inicial
        self._mostrar_pagina("cfg")

        # ── STATUS BAR ────────────────────────────────────────────────────────
        sb = tk.Frame(self, bg="#0A1525", height=24)
        sb.pack(fill="x", side="bottom")
        sb.pack_propagate(False)
        self._lbl_status = tk.Label(sb, text="  ✓ Listo",
                                     bg="#0A1525", fg="#8B949E",
                                     font=("Segoe UI", 8))
        self._lbl_status.pack(side="left", padx=10, pady=3)
        tk.Label(sb, text=f"Bashkar Station v{APP_VERSION}",
                 bg="#0A1525", fg="#334155",
                 font=("Segoe UI", 8)).pack(side="right", padx=10)

    # ── SIDEBAR ───────────────────────────────────────────────────────────────
    # ══════════════════════════════════════════════════════════════════════════
    # TOPBAR
    # ══════════════════════════════════════════════════════════════════════════
    def _build_topbar(self):
        tb = self._topbar

        # Logo + nombre app
        logo_grp = tk.Frame(tb, bg=TOPBAR_BG)
        logo_grp.pack(side="left", padx=(12, 0))
        tk.Label(logo_grp, text="⬡", bg=TOPBAR_BG, fg=AZ4,
                 font=("Segoe UI", 14, "bold")).pack(side="left")
        tk.Label(logo_grp, text="Bashkar Station", bg=TOPBAR_BG, fg=TXT_PRI,
                 font=("Segoe UI", 10, "bold")).pack(side="left", padx=(6, 16))

        # Separador
        tk.Frame(tb, bg=CARD_BOR, width=1).pack(side="left", fill="y",
                                                  pady=8, padx=4)

        # Nombre del proyecto activo
        self._lbl_pub_hdr = tk.Label(tb, text="Sin proyecto",
                                      bg=TOPBAR_BG, fg=TXT_SEC,
                                      font=("Segoe UI", 9))
        self._lbl_pub_hdr.pack(side="left", padx=12)

        # ── Lado derecho: switch IA + botones proyecto ────────────────────────
        right = tk.Frame(tb, bg=TOPBAR_BG)
        right.pack(side="right", padx=12)

        # Botón Dark / Light mode
        self._btn_theme = tk.Label(right, text="☀ Claro", bg=TOPBAR_BG,
                                    fg=TXT_SEC, font=("Segoe UI", 8),
                                    cursor="hand2", padx=8, pady=2,
                                    relief="solid", bd=1)
        self._btn_theme.pack(side="right", padx=(0, 6))
        self._btn_theme.bind("<Button-1>", lambda e: self._toggle_theme())
        self._btn_theme.bind("<Enter>",
            lambda e: self._btn_theme.config(fg=TXT_PRI))
        self._btn_theme.bind("<Leave>",
            lambda e: self._btn_theme.config(fg=TXT_SEC))

        tk.Frame(right, bg=CARD_BOR, width=1).pack(side="right", fill="y",
                                                     pady=6, padx=4)

        # Switch IA — siempre visible
        self._var_ia_habilitada = tk.BooleanVar(
            value=getattr(ST, "ia_habilitada", False))
        ia_frame = tk.Frame(right, bg=TOPBAR_BG)
        ia_frame.pack(side="right", padx=(8, 0))

        self._lbl_ia_topbar = tk.Label(ia_frame, bg=TOPBAR_BG,
                                        font=("Segoe UI", 8, "bold"),
                                        cursor="hand2")
        self._lbl_ia_topbar.pack(side="left", padx=(0, 4))
        self._lbl_ia_topbar.bind("<Button-1>", lambda e: self._topbar_toggle_ia())

        ttk.Checkbutton(ia_frame, text="IA",
                         variable=self._var_ia_habilitada,
                         command=self._topbar_toggle_ia).pack(side="left")

        self._topbar_toggle_ia()  # inicializar etiqueta

        # Separador
        tk.Frame(right, bg=CARD_BOR, width=1).pack(side="right", fill="y",
                                                     pady=6, padx=8)

        # Botones de proyecto
        for txt, cmd in [("💾", self._guardar_proyecto),
                          ("📂", self._abrir_gestor_proyectos),
                          ("➕", self._nuevo_proyecto_dialogo)]:
            b = tk.Label(right, text=txt, bg=TOPBAR_BG, fg=TXT_SEC,
                         font=("Segoe UI", 12), cursor="hand2", padx=6)
            b.pack(side="right")
            b.bind("<Button-1>", lambda e, c=cmd: c())
            b.bind("<Enter>", lambda e, w=b: w.config(fg=TXT_PRI))
            b.bind("<Leave>", lambda e, w=b: w.config(fg=TXT_SEC))

        # Botón análisis rápido (modo sin proyecto)
        b_adhoc = tk.Label(right, text="⚡", bg=TOPBAR_BG, fg=TXT_SEC,
                           font=("Segoe UI", 12), cursor="hand2", padx=6)
        b_adhoc.pack(side="right")
        b_adhoc.bind("<Button-1>", lambda e: self._modo_adhoc())
        b_adhoc.bind("<Enter>",    lambda e: b_adhoc.config(fg="#F59E0B"))
        b_adhoc.bind("<Leave>",    lambda e: b_adhoc.config(fg=TXT_SEC))
        self._mk_ayuda_topbar(b_adhoc, "⚡ Análisis rápido sin proyecto\n"
                                        "Carga una carpeta de TXT directamente.")

        # Separador + botón Bitácora + botón Command Palette
        tk.Frame(right, bg=CARD_BOR, width=1).pack(
            side="right", fill="y", pady=6, padx=4)
        self._btn_bitacora = tk.Label(
            right, text="📓", bg=TOPBAR_BG, fg=TXT_SEC,
            font=("Segoe UI", 13), cursor="hand2", padx=6)
        self._btn_bitacora.pack(side="right")
        self._btn_bitacora.bind("<Button-1>", lambda e: self._bitacora_abrir())
        self._btn_bitacora.bind("<Enter>",
            lambda e: self._btn_bitacora.config(fg=TXT_PRI))
        self._btn_bitacora.bind("<Leave>",
            lambda e: self._btn_bitacora.config(fg=TXT_SEC))
        self._bitacora_win = None   # referencia a la ventana flotante
        self._bitacora_engine = None  # BitacoraEngine (se inicializa al abrir)

        # Botón Command Palette
        self._btn_cp = tk.Label(
            right, text="⌨", bg=TOPBAR_BG, fg=TXT_SEC,
            font=("Segoe UI", 13), cursor="hand2", padx=6)
        self._btn_cp.pack(side="right")
        self._btn_cp.bind("<Button-1>", self._abrir_command_palette)
        self._btn_cp.bind("<Enter>", lambda e: self._btn_cp.config(fg=AZ4))
        self._btn_cp.bind("<Leave>", lambda e: self._btn_cp.config(fg=TXT_SEC))
        self._mk_ayuda_topbar(self._btn_cp, "⌨  Command Palette\nCtrl+K — busca y ejecuta cualquier acción")

    def _mk_ayuda_topbar(self, widget, texto: str):
        """Tooltip simple para widgets de la topbar."""
        tip = [None]
        def _show(e):
            try:
                if tip[0]: tip[0].destroy()
                x = widget.winfo_rootx()
                y = widget.winfo_rooty() + widget.winfo_height() + 4
                w = tk.Toplevel(self); w.wm_overrideredirect(True)
                w.geometry(f"+{x}+{y}"); w.configure(bg=CARD_BOR)
                tk.Label(w, text=texto, bg=AZ2, fg=TXT_PRI,
                         font=("Segoe UI", 8), padx=8, pady=4,
                         justify="left").pack()
                tip[0] = w
            except Exception: pass
        def _hide(e):
            try:
                if tip[0]: tip[0].destroy(); tip[0] = None
            except Exception: pass
        widget.bind("<Enter>", _show); widget.bind("<Leave>", _hide)

    def _modo_adhoc(self):
        """Carga una carpeta de TXT directamente sin crear proyecto .bashkar."""
        from tkinter import filedialog
        carpeta = filedialog.askdirectory(
            title="Seleccionar carpeta con archivos .txt para análisis rápido")
        if not carpeta:
            return
        carpeta = Path(carpeta)
        txts = sorted(carpeta.rglob("*.txt"))
        if not txts:
            messagebox.showwarning("Sin archivos",
                                   f"No se encontraron archivos .txt en:\n{carpeta}")
            return

        # Cargar textos directamente en ST sin pipeline OCR
        corpus_txt = []
        corpus_meta = {}
        for i, p in enumerate(txts):
            try:
                texto = p.read_text(encoding="utf-8", errors="replace")
                corpus_txt.append(texto)
                corpus_meta[str(i)] = {
                    "titulo": p.stem,
                    "numero": p.parent.name,
                    "pagina": p.stem,
                    "art_id": str(i),
                }
            except Exception:
                continue

        ST.corpus_txt  = corpus_txt
        ST.corpus_meta = corpus_meta
        ST.out_dir     = carpeta
        ST.pdf_dir     = carpeta
        ST.publicacion = carpeta.name
        ST.archivos_sel = list(txts)
        ST.ocr_done = True
        ST.marcar_etapa("ocr", "ready")
        self._actualizar_badges()

        # Actualizar etiqueta de proyecto
        if hasattr(self, "_lbl_pub_hdr"):
            self._lbl_pub_hdr.config(
                text=f"⚡ {carpeta.name}  ·  {len(txts)} archivos (modo ad-hoc)")
        if hasattr(self, "_lbl_proyecto"):
            try:
                self._lbl_proyecto.config(text=f"⚡ {carpeta.name}")
            except Exception:
                pass

        messagebox.showinfo(
            "Corpus cargado ⚡",
            f"✅ {len(corpus_txt)} archivos TXT cargados.\n\n"
            f"Carpeta: {carpeta}\n\n"
            f"Puedes ir directamente a:\n"
            f"  · Segmentar — dividir en artículos\n"
            f"  · Collocates — análisis léxico\n"
            f"  · Búsqueda semántica — buscar en el corpus\n\n"
            f"(No se creó proyecto .bashkar — los resultados no se guardan automáticamente)")
        self._mostrar_pagina("coloc")

    def _toggle_theme(self):
        global _MODO_OSCURO
        _MODO_OSCURO = not _MODO_OSCURO
        paleta = _PALETA_DARK if _MODO_OSCURO else _PALETA_LIGHT
        _aplicar_paleta(paleta)

        icono = "☀ Claro" if _MODO_OSCURO else "🌙 Oscuro"
        self._btn_theme.config(text=icono)

        # Repintar todos los widgets conocidos recursivamente
        bg_main  = paleta["CONTENT_BG"]
        bg_card  = paleta["CARD_BG"]
        bg_sb    = paleta["SB_BG"]
        bg_ab    = paleta["AB_BG"]
        bg_top   = paleta["TOPBAR_BG"]
        fg_pri   = paleta["TXT_PRI"]
        fg_sec   = paleta["TXT_SEC"]
        fg_dim   = paleta["TXT_DIM"]
        bor      = paleta["CARD_BOR"]

        # Mapa completo: todos los colores oscuros hardcodeados → equivalente del tema
        _DARK_BG_MAP = {
            # VS Code dark (nuevos)
            "#1E1E1E": bg_main,  "#252526": bg_card,  "#333333": bg_ab,
            "#3C3C3C": bg_top,   "#2A2D2E": paleta["SB_HOV"],
            "#094771": paleta["SB_SEL"],
            # GitHub dark (legado)
            "#0D1117": bg_sb,    "#161B22": bg_top,   "#1C2128": bg_card,
            "#13171F": bg_main,  "#2D333B": paleta["SB_HOV"],
            # Hardcoded varios
            "#0D1B2A": bg_card,  "#112240": bg_card,  "#0A1628": bg_card,
            "#1E2A3A": bg_card,  "#1E293B": bg_card,  "#0A1525": bg_top,
            "#F0FDF4": bg_card,  "#1F2937": bg_card,  "#111827": bg_main,
            "#2D2D2D": bg_card,  "#3D3D3D": bg_card,
        }
        _DARK_FG_MAP = {
            # VS Code dark (nuevos)
            "#D4D4D4": fg_pri,   "#BBBBBB": fg_sec,   "#CCCCCC": fg_pri,
            "#858585": fg_sec,   "#5A5A5A": fg_dim,
            # GitHub dark (legado)
            "#E6EDF3": fg_pri,   "#CDD6F4": fg_pri,   "#E2E8F0": fg_pri,
            "#8B949E": fg_sec,   "#94A3B8": fg_sec,   "#60A5FA": paleta["AZ3"],
            "#484F58": fg_dim,   "#30363D": bor,       "#8C959F": fg_dim,
            "white":   fg_pri,   "#FFFFFF": fg_pri,
        }

        def _repintar(widget):
            try:
                cls = widget.winfo_class()

                if cls in ("Frame", "Canvas"):
                    try:
                        cur = widget.cget("bg")
                        widget.config(bg=_DARK_BG_MAP.get(cur, bg_main))
                    except Exception:
                        pass

                elif cls == "Label":
                    try:
                        cur_bg = widget.cget("bg")
                        cur_fg = widget.cget("fg")
                        widget.config(
                            bg=_DARK_BG_MAP.get(cur_bg, cur_bg),
                            fg=_DARK_FG_MAP.get(cur_fg, cur_fg))
                    except Exception:
                        pass

                elif cls in ("Text",):
                    try:
                        widget.config(bg=bg_card, fg=fg_pri,
                                      insertbackground=fg_pri)
                    except Exception:
                        pass

                elif cls == "Entry":
                    try:
                        cur_bg = widget.cget("bg")
                        widget.config(
                            bg=_DARK_BG_MAP.get(cur_bg, bg_card),
                            fg=fg_pri,
                            insertbackground=fg_pri,
                            relief="solid", bd=1)
                    except Exception:
                        pass

                elif cls == "Listbox":
                    try:
                        widget.config(bg=bg_card, fg=fg_pri,
                                      selectbackground=paleta["SB_SEL"],
                                      selectforeground="#FFFFFF")
                    except Exception:
                        pass

                elif cls == "Spinbox":
                    try:
                        widget.config(bg=bg_card, fg=fg_pri,
                                      insertbackground=fg_pri)
                    except Exception:
                        pass

            except Exception:
                pass
            for child in widget.winfo_children():
                _repintar(child)

        _repintar(self)

        # Actualizar sidebar y activity bar explícitamente
        self.configure(bg=bg_sb)
        if hasattr(self, "_activity_bar"):
            self._activity_bar.config(bg=bg_ab)
        if hasattr(self, "_sidebar"):
            self._sidebar.config(bg=bg_sb)
        if hasattr(self, "_topbar"):
            self._topbar.config(bg=bg_top)
        if hasattr(self, "_content_area"):
            self._content_area.config(bg=bg_main)

        # Actualizar estilos ttk (Combobox, Entry, Treeview, Button)
        style = ttk.Style(self)
        style.configure("TCombobox",
            fieldbackground=bg_card, background=bg_card,
            foreground=fg_pri, selectbackground=paleta["SB_SEL"],
            selectforeground=fg_pri, arrowcolor=fg_sec)
        style.configure("TEntry",
            fieldbackground=bg_card, foreground=fg_pri,
            insertcolor=fg_pri, bordercolor=bor)
        style.configure("TSpinbox",
            fieldbackground=bg_card, foreground=fg_pri,
            background=bg_card, arrowcolor=fg_sec)
        style.configure("Treeview",
            background=bg_card, foreground=fg_pri,
            fieldbackground=bg_card, rowheight=24)
        style.configure("Treeview.Heading",
            background=bg_card, foreground=fg_sec,
            relief="flat")
        style.map("Treeview",
            background=[("selected", paleta["SB_SEL"])],
            foreground=[("selected", "#FFFFFF")])
        style.configure("TScrollbar",
            background=bg_card, troughcolor=bg_main,
            arrowcolor=fg_sec)
        style.configure("TNotebook",
            background=bg_main, tabmargins=[2, 5, 2, 0])
        style.configure("TNotebook.Tab",
            background=bg_card, foreground=fg_sec,
            padding=[8, 4])
        style.map("TNotebook.Tab",
            background=[("selected", bg_main)],
            foreground=[("selected", fg_pri)])
        style.configure("TCheckbutton",
            background=bg_card, foreground=fg_pri)
        style.configure("TRadiobutton",
            background=bg_main, foreground=fg_pri)
        style.configure("TLabelframe",
            background=bg_card, foreground=fg_pri, bordercolor=bor)
        style.configure("TLabelframe.Label",
            background=bg_card, foreground=fg_pri)

        # Forzar redibujado de badges/semáforos
        self._actualizar_badges()

    def _topbar_toggle_ia(self):
        habilitada = self._var_ia_habilitada.get()
        ST.ia_habilitada = habilitada
        if habilitada:
            self._lbl_ia_topbar.config(text="● IA ON",  fg=VERDE)
        else:
            self._lbl_ia_topbar.config(text="○ IA OFF", fg=ROJO)
        # Sincronizar con el checkbox de Configuración si existe
        if hasattr(self, "_cfg_toggle_ia"):
            try:
                self._cfg_toggle_ia()
            except Exception:
                pass

    # ══════════════════════════════════════════════════════════════════════════
    # ACTIVITY BAR
    # ══════════════════════════════════════════════════════════════════════════

    # Contextos del Activity Bar: (id_contexto, icono, tooltip, [pids_incluidos])
    _CONTEXTOS = [
        ("ingest",    "📥", "Ingestión",   ["cfg", "etz", "ocr", "conv", "mmx"]),
        ("normalize", "📝", "Normalizar",  ["norm"]),
        ("segment",   "✂", "Segmentar",   ["seg"]),
        ("analyze",   "🔬", "Analizar",    ["anal", "ner", "anot", "bsem",
                                             "coloc", "nov", "red", "ling",
                                             "sem", "top", "vis", "imgdesc"]),
        ("visualize", "🎨", "Visualizar",  ["viz", "comp", "comp2", "intxt"]),
        ("publish",   "📤", "Publicar",    ["rep", "dash", "valid", "colab"]),
        ("settings",  "⚙️", "Proyecto",    ["meta", "res"]),
    ]

    def _build_activity_bar(self):
        ab = self._activity_bar
        self._ctx_activo = tk.StringVar(value="ingest")
        self._ab_btns = {}

        # Logo pequeño en el tope
        tk.Label(ab, text="⬡", bg=AB_BG, fg=AB_SEL,
                 font=("Segoe UI", 16, "bold")).pack(pady=(10, 8))
        tk.Frame(ab, bg=CARD_BOR, height=1).pack(fill="x", padx=8, pady=(0, 8))

        # Botones de contexto
        for ctx_id, icono, tooltip, _ in self._CONTEXTOS:
            self._ab_btns[ctx_id] = self._make_ab_btn(ab, ctx_id, icono, tooltip)

        # Spacer
        tk.Frame(ab, bg=AB_BG).pack(fill="both", expand=True)

        # Botón ayuda al fondo
        help_b = tk.Label(ab, text="?", bg=AB_BG, fg=AB_TXT,
                          font=("Segoe UI", 11, "bold"), cursor="hand2",
                          width=3, pady=8)
        help_b.pack(pady=(0, 8))
        help_b.bind("<Button-1>", lambda e: self._abrir_docs())
        help_b.bind("<Enter>",    lambda e: help_b.config(fg=TXT_PRI))
        help_b.bind("<Leave>",    lambda e: help_b.config(fg=AB_TXT))

        # Activar contexto inicial
        self._activar_contexto("ingest")

    def _make_ab_btn(self, parent, ctx_id, icono, tooltip):
        frm = tk.Frame(parent, bg=AB_BG, cursor="hand2", width=60, height=56)
        frm.pack(fill="x")
        frm.pack_propagate(False)

        # Barra indicadora izquierda (3px, visible cuando activo)
        ind = tk.Frame(frm, bg=AB_BG, width=3)
        ind.pack(side="left", fill="y")

        # Ícono centrado
        lbl = tk.Label(frm, text=icono, bg=AB_BG, fg=AB_TXT,
                        font=("Segoe UI", 18), cursor="hand2")
        lbl.pack(expand=True)

        widgets = {"frm": frm, "ind": ind, "lbl": lbl}

        def _click(e, c=ctx_id):
            self._activar_contexto(c)
        def _enter(e, c=ctx_id):
            if self._ctx_activo.get() != c:
                lbl.config(fg=TXT_PRI)
                # Tooltip
                self._mostrar_tooltip(tooltip, frm)
        def _leave(e, c=ctx_id):
            if self._ctx_activo.get() != c:
                lbl.config(fg=AB_TXT)
            self._ocultar_tooltip()

        for w in (frm, lbl):
            w.bind("<Button-1>", _click)
            w.bind("<Enter>",    _enter)
            w.bind("<Leave>",    _leave)

        return widgets

    def _activar_contexto(self, ctx_id: str):
        self._ctx_activo.set(ctx_id)
        # Actualizar estilos activity bar
        for cid, widgets in self._ab_btns.items():
            if cid == ctx_id:
                widgets["frm"].config(bg=AB_BG)
                widgets["ind"].config(bg=AB_IND)
                widgets["lbl"].config(fg=TXT_PRI)
            else:
                widgets["frm"].config(bg=AB_BG)
                widgets["ind"].config(bg=AB_BG)
                widgets["lbl"].config(fg=AB_TXT)

        # Poblar el sidebar con los pids de este contexto
        self._poblar_sidebar_contexto(ctx_id)

        # Navegar al primer pid del contexto si ninguno está activo
        pids = next((c[3] for c in self._CONTEXTOS if c[0] == ctx_id), [])
        actual = self._pagina_activa.get()
        if actual not in pids and pids:
            self._mostrar_pagina(pids[0])

    def _mostrar_tooltip(self, texto: str, widget):
        """Tooltip simple junto al activity bar."""
        try:
            if hasattr(self, "_tooltip_win") and self._tooltip_win:
                self._tooltip_win.destroy()
            x = widget.winfo_rootx() + 64
            y = widget.winfo_rooty() + 16
            win = tk.Toplevel(self)
            win.wm_overrideredirect(True)
            win.geometry(f"+{x}+{y}")
            win.configure(bg=CARD_BOR)
            tk.Label(win, text=texto, bg=AZ2, fg=TXT_PRI,
                     font=("Segoe UI", 8), padx=8, pady=4).pack()
            self._tooltip_win = win
        except Exception:
            pass

    def _ocultar_tooltip(self):
        try:
            if hasattr(self, "_tooltip_win") and self._tooltip_win:
                self._tooltip_win.destroy()
                self._tooltip_win = None
        except Exception:
            pass

    # ══════════════════════════════════════════════════════════════════════════
    # SIDEBAR DE SUB-ITEMS
    # ══════════════════════════════════════════════════════════════════════════
    def _poblar_sidebar_contexto(self, ctx_id: str):
        """Limpia el sidebar y lo repuebla con los pids del contexto activo."""
        sb = self._sidebar
        for w in sb.winfo_children():
            w.destroy()

        pids = next((c[3] for c in self._CONTEXTOS if c[0] == ctx_id), [])
        ctx_label = next((c[2] for c in self._CONTEXTOS if c[0] == ctx_id), "")

        # Eliminar de _sb_btns los pids que ya no están en el sidebar visible
        # para evitar .config() sobre widgets destruidos en iteraciones posteriores
        for old_pid in [p for p in self._sb_btns if p not in pids]:
            del self._sb_btns[old_pid]

        # Cabecera del contexto
        hdr = tk.Frame(sb, bg=SB_BG)
        hdr.pack(fill="x", padx=0, pady=0)
        tk.Label(hdr, text=ctx_label.upper(), bg=SB_BG, fg=TXT_DIM,
                 font=("Segoe UI", 7, "bold")).pack(anchor="w", padx=16, pady=(12, 6))

        # Nombre del proyecto
        self._lbl_proyecto = tk.Label(sb, text=getattr(ST, "publicacion", "") or "Sin proyecto",
                                       bg=SB_BG, fg=TXT_SEC,
                                       font=("Segoe UI", 8),
                                       wraplength=180, justify="left")
        self._lbl_proyecto.pack(anchor="w", padx=16, pady=(0, 8))
        tk.Frame(sb, bg=CARD_BOR, height=1).pack(fill="x", padx=12, pady=(0, 4))

        # Área scrollable — canvas sin scrollbar visible
        canvas = tk.Canvas(sb, bg=SB_BG, highlightthickness=0, borderwidth=0)
        canvas.pack(fill="both", expand=True)

        inner = tk.Frame(canvas, bg=SB_BG)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")

        def _on_inner(e):
            canvas.configure(scrollregion=canvas.bbox("all"))
        def _on_canvas(e):
            canvas.itemconfig(win_id, width=e.width)

        inner.bind("<Configure>", _on_inner)
        canvas.bind("<Configure>", _on_canvas)

        # Scroll con rueda del mouse
        def _on_wheel(e):
            delta = -1 if (e.delta > 0 or e.num == 4) else 1
            canvas.yview_scroll(delta, "units")

        canvas.bind("<MouseWheel>", _on_wheel)
        canvas.bind("<Button-4>",   _on_wheel)
        canvas.bind("<Button-5>",   _on_wheel)
        inner.bind("<MouseWheel>",  _on_wheel)
        inner.bind("<Button-4>",    _on_wheel)
        inner.bind("<Button-5>",    _on_wheel)

        # Buscar info de cada pid en _PAGINAS
        pag_info = {p[0]: p for p in self._PAGINAS}
        flujo_num = 1
        for pid in pids:
            info = pag_info.get(pid)
            if not info:
                continue
            _, emoji, label, _, badge_attr, grupo = info
            num = flujo_num if grupo == "flujo" else None
            btn = self._make_sb_btn(inner, pid, emoji, label, num,
                                    badge_attr, es_flujo=(grupo == "flujo"),
                                    scroll_fn=_on_wheel)
            self._sb_btns[pid] = btn
            if grupo == "flujo":
                flujo_num += 1

        # Actualizar estado activo
        actual = self._pagina_activa.get()
        for pid, widgets in self._sb_btns.items():
            self._aplicar_estilo_sb_btn(pid, widgets, pid == actual)

    def _build_sidebar(self):
        sb = self._sidebar

        # ── CABECERA FIJA (logo + publicación) ───────────────────────────────
        # Cabecera del proyecto — sin logo, solo nombre y período
        self._lbl_pub_hdr = tk.Label(sb, text="Sin proyecto",
                                      bg=SB_BG, fg=TXT_SEC,
                                      font=("Segoe UI", 8, "bold"),
                                      wraplength=185, justify="left",
                                      anchor="w")
        self._lbl_pub_hdr.pack(fill="x", padx=10, pady=(10, 2))
        tk.Frame(sb, bg=CARD_BOR, height=1).pack(fill="x", padx=10, pady=(2, 4))

        # ── ÁREA SCROLLABLE DE BOTONES ────────────────────────────────────────
        nav_container = tk.Frame(sb, bg=SB_BG)
        nav_container.pack(fill="both", expand=True)

        # Canvas + scrollbar vertical
        nav_canvas = tk.Canvas(nav_container, bg=SB_BG, highlightthickness=0,
                                borderwidth=0)
        nav_scrollbar = tk.Scrollbar(nav_container, orient="vertical",
                                      command=nav_canvas.yview)
        nav_canvas.configure(yscrollcommand=nav_scrollbar.set)

        nav_scrollbar.pack(side="right", fill="y")
        nav_canvas.pack(side="left", fill="both", expand=True)

        # Frame interior que contiene los botones
        nav_inner = tk.Frame(nav_canvas, bg=SB_BG)
        nav_canvas_window = nav_canvas.create_window(
            (0, 0), window=nav_inner, anchor="nw"
        )

        def _on_nav_inner_configure(event):
            nav_canvas.configure(scrollregion=nav_canvas.bbox("all"))
        def _on_nav_canvas_configure(event):
            nav_canvas.itemconfig(nav_canvas_window, width=event.width)

        nav_inner.bind("<Configure>", _on_nav_inner_configure)
        nav_canvas.bind("<Configure>", _on_nav_canvas_configure)

        # Scroll con rueda del mouse dentro del canvas de navegación
        def _nav_wheel(event):
            if event.num == 4 or event.delta > 0:
                nav_canvas.yview_scroll(-1, "units")
            else:
                nav_canvas.yview_scroll(1, "units")

        nav_canvas.bind("<MouseWheel>", _nav_wheel)
        nav_canvas.bind("<Button-4>",   _nav_wheel)
        nav_canvas.bind("<Button-5>",   _nav_wheel)
        nav_inner.bind("<MouseWheel>",  _nav_wheel)
        nav_inner.bind("<Button-4>",    _nav_wheel)
        nav_inner.bind("<Button-5>",    _nav_wheel)

        # Botones de navegación agrupados
        self._etz_nav_wheel_fn = _nav_wheel   # guardar para bind dinámico
        self._poblar_nav_sidebar(nav_inner, _nav_wheel)

        # El sidebar se puebla dinámicamente desde _poblar_sidebar_contexto
        # Este método solo inicializa el label de proyecto (fallback)
        self._lbl_proyecto = tk.Label(sb, text="Sin proyecto",
                                       bg=SB_BG, fg=TXT_SEC,
                                       font=("Segoe UI", 8),
                                       wraplength=180, justify="left")

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL DE ASISTENTE IA (compartido por todas las pestañas de análisis)
    # ══════════════════════════════════════════════════════════════════════════

    def _build_ai_panel(self, parent_frame: "tk.Frame", tab_id: str):
        """
        Construye el panel de asistente IA en la parte inferior de una pestaña.
        Debe llamarse ANTES de construir el contenido principal (pack side=bottom).
        """
        sugerencias = _AI_PROMPTS.get(tab_id, [])

        # Contenedor principal del panel — fondo oscuro tipo terminal
        panel = tk.Frame(parent_frame, bg="#0A1525", bd=0)
        panel.pack(fill="x", side="bottom")

        # ── Cabecera colapsable ───────────────────────────────────────────────
        hdr = tk.Frame(panel, bg="#112240", cursor="hand2")
        hdr.pack(fill="x")
        self._ai_expanded = getattr(self, "_ai_expanded", {})
        self._ai_expanded[tab_id] = tk.BooleanVar(value=False)

        lbl_toggle = tk.Label(hdr,
            text="  🤖  Asistente IA  ▸  haz una pregunta sobre este análisis",
            bg="#112240", fg="#60A5FA", font=("Segoe UI", 9, "bold"),
            anchor="w", cursor="hand2")
        lbl_toggle.pack(side="left", fill="x", expand=True, pady=5, padx=8)
        lbl_chevron = tk.Label(hdr, text="▾", bg="#112240", fg="#60A5FA",
                                font=("Segoe UI", 11, "bold"))
        lbl_chevron.pack(side="right", padx=10)

        # ── Cuerpo (oculto por defecto) ───────────────────────────────────────
        body = tk.Frame(panel, bg="#0A1525")

        # Prompts sugeridos
        if sugerencias:
            sug_frame = tk.Frame(body, bg="#0A1525")
            sug_frame.pack(fill="x", padx=10, pady=(8, 0))
            tk.Label(sug_frame, text="Sugerencias:", bg="#0A1525", fg="#8B949E",
                     font=("Segoe UI", 8, "bold")).pack(anchor="w")
            btn_row = tk.Frame(sug_frame, bg="#0A1525")
            btn_row.pack(fill="x", pady=(4, 0))
            for i, (label, prompt_txt) in enumerate(sugerencias):
                btn = tk.Label(btn_row, text=f"  {label}  ",
                               bg="#1E3A5F", fg="#93C5FD",
                               font=("Segoe UI", 8), relief="flat",
                               cursor="hand2", padx=6, pady=3)
                btn.grid(row=i//3, column=i%3, padx=3, pady=2, sticky="w")
                btn.bind("<Enter>", lambda e, b=btn: b.config(bg="#1D4ED8", fg="white"))
                btn.bind("<Leave>", lambda e, b=btn: b.config(bg="#1E3A5F", fg="#93C5FD"))
                btn.bind("<Button-1>",
                         lambda e, t=prompt_txt, tid=tab_id: self._set_ai_prompt(t, tid))

        # Área de texto del prompt
        txt_frame = tk.Frame(body, bg="#0A1525")
        txt_frame.pack(fill="x", padx=10, pady=(8, 0))
        tk.Label(txt_frame, text="Tu prompt:", bg="#0A1525", fg="#94A3B8",
                 font=("Segoe UI", 8)).pack(anchor="w")

        prompt_txt_widget = tk.Text(txt_frame, height=3, font=("Segoe UI", 9),
                                     bg="#112240", fg="#E2E8F0",
                                     insertbackground="#60A5FA",
                                     relief="flat", bd=0,
                                     wrap="word", padx=8, pady=6)
        prompt_txt_widget.pack(fill="x", pady=(3, 0))
        prompt_txt_widget.insert("1.0",
            "Escribe tu pregunta o selecciona una sugerencia arriba…")
        prompt_txt_widget.config(fg="#8B949E")

        def _on_focus_in(e):
            if prompt_txt_widget.get("1.0","end-1c") ==                "Escribe tu pregunta o selecciona una sugerencia arriba…":
                prompt_txt_widget.delete("1.0","end")
                prompt_txt_widget.config(fg="#E2E8F0")
        def _on_focus_out(e):
            if not prompt_txt_widget.get("1.0","end-1c").strip():
                prompt_txt_widget.insert("1.0",
                    "Escribe tu pregunta o selecciona una sugerencia arriba…")
                prompt_txt_widget.config(fg="#8B949E")
        prompt_txt_widget.bind("<FocusIn>",  _on_focus_in)
        prompt_txt_widget.bind("<FocusOut>", _on_focus_out)

        # Botones de acción
        act_row = tk.Frame(body, bg="#0A1525")
        act_row.pack(fill="x", padx=10, pady=(6, 0))
        send_btn = tk.Label(act_row, text="  ▶  Enviar a IA  ",
                            bg="#1D4ED8", fg="white",
                            font=("Segoe UI", 9, "bold"),
                            cursor="hand2", padx=8, pady=4)
        send_btn.pack(side="left")
        send_btn.bind("<Enter>", lambda e: send_btn.config(bg="#2563EB"))
        send_btn.bind("<Leave>", lambda e: send_btn.config(bg="#1D4ED8"))

        clear_btn = tk.Label(act_row, text="  ✕ Limpiar  ",
                              bg="#1E3A5F", fg="#94A3B8",
                              font=("Segoe UI", 8), cursor="hand2",
                              padx=6, pady=4)
        clear_btn.pack(side="left", padx=(6, 0))
        clear_btn.bind("<Enter>", lambda e: clear_btn.config(bg="#2D4A6E"))
        clear_btn.bind("<Leave>", lambda e: clear_btn.config(bg="#1E3A5F"))

        lbl_proveedor = tk.Label(act_row, text="", bg="#0A1525", fg="#8B949E",
                                  font=("Segoe UI", 7, "italic"))
        lbl_proveedor.pack(side="right", padx=8)

        # Área de respuesta
        resp_frame = tk.Frame(body, bg="#0A1525")
        resp_frame.pack(fill="x", padx=10, pady=(8, 0))
        resp_hdr = tk.Frame(resp_frame, bg="#112240")
        resp_hdr.pack(fill="x")
        tk.Label(resp_hdr, text="  Respuesta de la IA",
                 bg="#112240", fg="#94A3B8",
                 font=("Segoe UI", 8, "bold")).pack(side="left", pady=3)
        self._ai_lbl_estado = getattr(self, "_ai_lbl_estado", {})
        lbl_estado = tk.Label(resp_hdr, text="", bg="#112240", fg="#34D399",
                               font=("Segoe UI", 7, "italic"))
        lbl_estado.pack(side="right", padx=8)
        self._ai_lbl_estado[tab_id] = lbl_estado

        resp_txt = scrolledtext.ScrolledText(resp_frame, height=6,
                                              font=("Segoe UI", 9),
                                              bg="#0F1B2D", fg="#E2E8F0",
                                              relief="flat",
                                              insertbackground="white",
                                              state="disabled", wrap="word")
        resp_txt.pack(fill="x", pady=(0, 0))

        # Separador inferior
        tk.Frame(body, bg="#1E3A5F", height=1).pack(fill="x", pady=(8, 0))

        # ── Guardar referencias ───────────────────────────────────────────────
        if not hasattr(self, "_ai_widgets"):
            self._ai_widgets = {}
        self._ai_widgets[tab_id] = {
            "prompt": prompt_txt_widget,
            "resp":   resp_txt,
            "estado": lbl_estado,
            "prov":   lbl_proveedor,
        }

        # ── Toggle show/hide ──────────────────────────────────────────────────
        def _toggle(e=None):
            if body.winfo_ismapped():
                body.pack_forget()
                lbl_chevron.config(text="▾")
                lbl_toggle.config(text="  🤖  Asistente IA  ▸  haz una pregunta sobre este análisis")
            else:
                body.pack(fill="x", pady=(0, 4))
                lbl_chevron.config(text="▴")
                lbl_toggle.config(text="  🤖  Asistente IA")
                # Actualizar label de proveedor
                key = getattr(ST, "api_key", "")
                if key:
                    from core.image_describer import nombre_proveedor
                    lbl_proveedor.config(text=nombre_proveedor(key))
                else:
                    lbl_proveedor.config(text="⚠ Sin API key — configura en Sección 8")

        for w in (hdr, lbl_toggle, lbl_chevron):
            w.bind("<Button-1>", _toggle)

        # ── Botón enviar ──────────────────────────────────────────────────────
        send_btn.bind("<Button-1>",
                      lambda e, tid=tab_id: self._enviar_prompt_ia(tid))
        clear_btn.bind("<Button-1>",
                       lambda e, tid=tab_id: self._limpiar_respuesta_ia(tid))

    def _set_ai_prompt(self, texto: str, tab_id: str):
        """Rellena el área de texto del prompt con el texto sugerido."""
        widgets = getattr(self, "_ai_widgets", {}).get(tab_id)
        if not widgets: return
        w = widgets["prompt"]
        w.config(fg="#E2E8F0")
        w.delete("1.0", "end")
        w.insert("1.0", texto)
        w.focus_set()

    def _limpiar_respuesta_ia(self, tab_id: str):
        widgets = getattr(self, "_ai_widgets", {}).get(tab_id)
        if not widgets: return
        widgets["resp"].config(state="normal")
        widgets["resp"].delete("1.0", "end")
        widgets["resp"].config(state="disabled")
        widgets["estado"].config(text="")

    def _enviar_prompt_ia(self, tab_id: str):
        """Construye el contexto de la pestaña activa y llama a la IA."""
        widgets = getattr(self, "_ai_widgets", {}).get(tab_id)
        if not widgets:
            messagebox.showwarning("Panel IA", "Panel no inicializado."); return

        api_key, _modelo_ia = _resolver_api_key_modelo("asistente")
        api_key = api_key.strip()
        if not api_key:
            messagebox.showwarning(
                "Sin API key",
                "Configura una clave API en Sección 8 de Configuración.\n"
                "Compatible con Anthropic, OpenAI y Google Gemini."); return

        prompt = widgets["prompt"].get("1.0", "end-1c").strip()
        if not prompt or prompt == "Escribe tu pregunta o selecciona una sugerencia arriba…":
            messagebox.showwarning("Prompt vacío", "Escribe o selecciona un prompt."); return

        widgets["estado"].config(text="⏳ Consultando IA…", fg="#FBBF24")
        widgets["resp"].config(state="normal")
        widgets["resp"].delete("1.0", "end")
        widgets["resp"].insert("end", "⏳ Esperando respuesta…")
        widgets["resp"].config(state="disabled")

        # Construir contexto de la pestaña
        contexto = self._construir_contexto_ia(tab_id)

        prompt_snap = prompt  # capturar para historial
        def worker():
            try:
                respuesta = self._llamar_ia_texto(
                    api_key=api_key,
                    contexto=contexto,
                    prompt=prompt_snap,
                )
                self.after(0, lambda r=respuesta, tid=tab_id, p=prompt_snap:
                           self._mostrar_respuesta_ia(r, tid, ok=True, prompt_original=p))
            except Exception as e:
                self.after(0, lambda err=str(e), tid=tab_id:
                           self._mostrar_respuesta_ia(f"⚠️ Error: {err}", tid, ok=False))

        threading.Thread(target=worker, daemon=True).start()

    def _construir_contexto_ia(self, tab_id: str) -> str:
        """Genera un resumen del estado actual del análisis como contexto para la IA."""
        pub  = getattr(ST, "publicacion", "desconocida")
        per  = getattr(ST, "periodo", "")
        lineas = [
            f"Publicación analizada: {pub}" + (f" ({per})" if per else ""),
            f"Pestaña activa: {tab_id}",
        ]

        if tab_id == "ocr" and getattr(ST, "resumen_ocr", None):
            r = ST.resumen_ocr
            lineas.append(f"Archivos procesados: {r.get('n_archivos',0)}")
            lineas.append(f"Páginas totales: {r.get('n_paginas',0)}")
            lineas.append(f"Palabras extraídas: {r.get('n_palabras',0):,}")
            lineas.append(f"Confianza OCR media: {r.get('confianza_media',0):.1%}")

        elif tab_id == "seg" and getattr(ST, "df_articulos", None) is not None:
            df = ST.df_articulos
            lineas.append(f"Artículos segmentados: {len(df)}")
            if "autor" in df.columns:
                top_autores = df["autor"].value_counts().head(10)
                lineas.append("Top autores: " + ", ".join(
                    f"{a} ({n})" for a, n in top_autores.items() if a))
            if "seccion" in df.columns:
                secs = df["seccion"].value_counts().head(6)
                lineas.append("Secciones: " + ", ".join(f"{s} ({n})" for s, n in secs.items()))

        elif tab_id == "anal" and getattr(ST, "temas_lda", None):
            lineas.append("Temas LDA detectados:")
            for i, t in enumerate(ST.temas_lda[:6]):
                palabras = t.get("palabras", [])[:8]
                lineas.append(f"  Tema {i+1}: {', '.join(palabras)}")

        elif tab_id == "vis" and getattr(ST, "datos_imagenes", None):
            total_imgs = sum(
                len(pag.get("elementos", []))
                for datos in ST.datos_imagenes.values()
                for pag in datos.get("paginas", [])
            )
            lineas.append(f"Elementos visuales detectados: {total_imgs}")

        elif tab_id == "comp" and getattr(ST, "matriz_sim", None) is not None:
            lineas.append("Matriz de similitud calculada")
            if getattr(ST, "terminos_dist", None):
                dist = list(ST.terminos_dist.items())[:3]
                for nombre, terms in dist:
                    lineas.append(f"  Términos distintivos de {nombre}: "
                                   + ", ".join(terms[:8]))

        elif tab_id == "meta" and getattr(self, "_meta_actual", {}):
            meta = self._meta_actual
            for campo in ("titulo", "creador", "fecha", "editorial", "descripcion"):
                if meta.get(campo):
                    lineas.append(f"{campo.capitalize()}: {meta[campo]}")

        elif tab_id == "res" and getattr(ST, "df_articulos", None) is not None:
            df = ST.df_articulos
            lineas.append(f"Corpus completo: {len(df)} artículos analizados")
            if getattr(ST, "temas_lda", None):
                lineas.append(f"Temas LDA: {len(ST.temas_lda)}")
            if getattr(ST, "datos_imagenes", None):
                lineas.append(f"Números con análisis visual: {len(ST.datos_imagenes)}")

        if len(lineas) == 2:
            lineas.append("(Análisis aún no ejecutado — responde basándote en el contexto general de la publicación)")

        return "\n".join(lineas)

    def _llamar_ia_texto(self, api_key: str, contexto: str,
                          prompt: str, timeout: int = 60) -> str:
        """
        Llama a la IA con contexto + prompt y devuelve la respuesta como texto.
        Soporta Anthropic, OpenAI y Gemini.
        """
        import json
        import urllib.request

        from core.image_describer import detectar_proveedor

        proveedor = detectar_proveedor(api_key)

        system_msg = (
            "Eres un asistente experto en historia de la prensa colombiana, "
            "estudios editoriales latinoamericanos y humanidades digitales. "
            "Ayudas a investigadores a interpretar resultados de análisis "
            "computacional de publicaciones históricas. "
            "Responde siempre en español, de forma clara y académica, "
            "con referencias al contexto histórico colombiano cuando sea pertinente."
        )
        user_msg = (
            f"CONTEXTO DEL ANÁLISIS:\n{contexto}\n\n"
            f"PREGUNTA DEL INVESTIGADOR:\n{prompt}"
        )

        def _post(url, payload, headers):
            req = urllib.request.Request(url, data=json.dumps(payload).encode("utf-8"),
                                          headers=headers, method="POST")
            with urllib.request.urlopen(req, timeout=timeout) as r:
                return json.loads(r.read())

        if proveedor == "anthropic":
            data = _post(
                "https://api.anthropic.com/v1/messages",
                {"model": "claude-sonnet-4-20250514", "max_tokens": 1500,
                 "system": system_msg,
                 "messages": [{"role": "user", "content": user_msg}]},
                {"Content-Type": "application/json",
                 "x-api-key": api_key, "anthropic-version": "2023-06-01"})
            return data["content"][0]["text"].strip()

        elif proveedor == "openai":
            data = _post(
                "https://api.openai.com/v1/chat/completions",
                {"model": "gpt-4o", "max_tokens": 1500,
                 "messages": [{"role": "system", "content": system_msg},
                               {"role": "user", "content": user_msg}]},
                {"Content-Type": "application/json",
                 "Authorization": f"Bearer {api_key}"})
            return data["choices"][0]["message"]["content"].strip()

        elif proveedor == "gemini":
            url = (f"https://generativelanguage.googleapis.com/v1beta/models/"
                   f"gemini-1.5-flash:generateContent?key={api_key}")
            data = _post(url,
                {"contents": [{"parts": [{"text": system_msg + "\n\n" + user_msg}]}],
                 "generationConfig": {"maxOutputTokens": 1500, "temperature": 0.3}},
                {"Content-Type": "application/json"})
            return data["candidates"][0]["content"]["parts"][0]["text"].strip()

        raise ValueError(f"Proveedor desconocido: {proveedor}")

    def _mostrar_respuesta_ia(self, texto: str, tab_id: str, ok: bool = True,
                               prompt_original: str = ""):
        widgets = getattr(self, "_ai_widgets", {}).get(tab_id)
        if not widgets: return
        widgets["resp"].config(state="normal")
        widgets["resp"].delete("1.0", "end")
        widgets["resp"].insert("end", texto)
        widgets["resp"].config(state="disabled")
        if ok:
            widgets["estado"].config(text="✓ Respuesta recibida", fg="#34D399")
            # Guardar en historial
            from datetime import datetime as _dt
            self._historial_ia.append({
                "tab":       tab_id,
                "prompt":    prompt_original,
                "respuesta": texto,
                "fecha":     _dt.now().strftime("%Y-%m-%dT%H:%M:%S"),
            })
        else:
            widgets["estado"].config(text="⚠ Error", fg="#F87171")


    # ══════════════════════════════════════════════════════════════════════════
    # GESTIÓN DE PROYECTOS
    # ══════════════════════════════════════════════════════════════════════════

    def _cargar_ultimo_proyecto(self):
        """Al arrancar: muestra la pantalla de inicio (si está habilitada)
        o restaura la última sesión directamente. Único call-site: el
        self.after(200, ...) del __init__ — los tests headless parchean
        este método entero, así que lo de adentro no les afecta."""
        import os
        from core.user_prefs import obtener_pref
        if not os.environ.get("BASHKAR_NO_WELCOME") and obtener_pref("mostrar_inicio", True):
            self._welcome_mostrar()
            return
        self._cargar_ultimo_proyecto_directo()

    def _welcome_mostrar(self):
        """Pantalla de inicio estilo FineReader: continuar/nuevo/abrir/
        recientes. Cerrarla sin elegir SIEMPRE deja un proyecto cargado
        (el resto de la app asume que ST tiene uno)."""
        from core.project_manager import cargar_proyecto, listar_proyectos
        from core.user_prefs import guardar_pref

        win, content = self._mk_glass_toplevel("Bienvenido a Bashkar Station", 560, 480)
        self._welcome_win = win
        win.protocol("WM_DELETE_WINDOW", lambda: self._welcome_elegir(win, self._crear_proyecto_automatico))

        pad = tk.Frame(content, bg=CONTENT_BG)
        pad.pack(fill="both", expand=True, padx=20, pady=16)

        proyectos = listar_proyectos()
        ultimo = proyectos[0] if proyectos else None

        if ultimo:
            c = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1, cursor="hand2")
            c.pack(fill="x", pady=(0, 10))
            tk.Label(c, text=f"▶ Continuar «{ultimo['nombre']}»", bg=CARD_BG, fg=TXT_PRI,
                     font=("Segoe UI", 11, "bold")).pack(anchor="w", padx=14, pady=(10, 2))
            tk.Label(c, text=f"{ultimo.get('publicacion','')}  ·  modificado {ultimo.get('modificado','')}",
                     bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(anchor="w", padx=14, pady=(0, 10))
            def _continuar(_e=None):
                self._welcome_elegir(win, self._cargar_ultimo_proyecto_directo)
            c.bind("<Button-1>", _continuar)
            for w in c.winfo_children():
                w.bind("<Button-1>", _continuar)

        fila = tk.Frame(pad, bg=CONTENT_BG)
        fila.pack(fill="x", pady=(0, 10))
        ttk.Button(fila, text="🆕 Nuevo proyecto", style="S.TButton",
                   command=lambda: self._welcome_elegir(win, self._nuevo_proyecto_dialogo)
                   ).pack(side="left", padx=(0, 8))
        ttk.Button(fila, text="📂 Abrir…", style="S.TButton",
                   command=lambda: self._welcome_elegir(win, self._abrir_gestor_proyectos)
                   ).pack(side="left")

        recientes = proyectos[1:8] if ultimo else proyectos[:8]
        if recientes:
            tk.Label(pad, text="Recientes:", bg=CONTENT_BG, fg=TXT_PRI,
                     font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(6, 2))
            lb = tk.Listbox(pad, height=6, bg="#0D1117", fg="#CDD6F4", relief="solid", bd=1)
            lb.pack(fill="both", expand=True)
            for p in recientes:
                lb.insert("end", f"{p['nombre']}  ·  {p.get('modificado','')}")

            def _abrir_recientes_sel(_e=None):
                sel = lb.curselection()
                if not sel:
                    return
                ruta = Path(recientes[sel[0]]["ruta"])

                def _cargar():
                    res = cargar_proyecto(ruta, ST)
                    if res["ok"]:
                        self._proyecto_ruta = ruta
                        self._historial_ia = res.get("historial_ia", [])
                        self._sincronizar_ui_con_st()
                    else:
                        self._crear_proyecto_automatico()
                self._welcome_elegir(win, _cargar)
            lb.bind("<Double-Button-1>", _abrir_recientes_sel)

        var_no_mostrar = tk.BooleanVar(value=False)
        ttk.Checkbutton(pad, text="No mostrar esta pantalla al inicio",
                        variable=var_no_mostrar,
                        command=lambda: guardar_pref("mostrar_inicio", not var_no_mostrar.get())
                        ).pack(anchor="w", pady=(10, 0))

    def _welcome_elegir(self, win, fn):
        """Cierra la pantalla de inicio y ejecuta la acción elegida."""
        if win.winfo_exists():
            win.destroy()
        self._welcome_win = None
        fn()

    def _cargar_ultimo_proyecto_directo(self):
        """Restaura la última sesión guardada, o crea un proyecto vacío si
        no hay ninguna. Cuerpo original de _cargar_ultimo_proyecto (previo
        a la pantalla de inicio) — también es el camino de «Continuar»."""
        from core.project_manager import cargar_proyecto, cargar_ultimo
        ruta = cargar_ultimo()
        if ruta:
            res = cargar_proyecto(ruta, ST)
            if res["ok"]:
                self._proyecto_ruta = ruta
                self._historial_ia  = res.get("historial_ia", [])
                self._sincronizar_ui_con_st()
                nombre = res.get("nombre", ruta.stem)
                self._lbl_proyecto.config(text=nombre)
                self._lbl_pub_hdr.config(
                    text=f"{ST.publicacion}  ·  {ST.periodo}" if ST.periodo
                    else ST.publicacion)
                self._actualizar_badges()
                self.after(300, self._etz_refrescar_numeros)
                self.after(600, self._ocr_actualizar_estimacion)
                if res.get("migrado"):
                    self.after(500, lambda: messagebox.showinfo(
                        "Migración completada",
                        "El proyecto fue migrado automáticamente de v10 a v11.\n"
                        "Se creó una base de datos SQLite (.db) con todos los datos.\n"
                        "Se conservó un backup del archivo original."))
                return
        # Sin proyecto previo → crear uno vacío automáticamente
        self._crear_proyecto_automatico()

    def _crear_proyecto_automatico(self):
        """Crea un proyecto vacío con los datos por defecto de ST."""
        from datetime import datetime

        from core.project_manager import guardar_ultimo, nuevo_proyecto
        nombre = f"Proyecto {datetime.now().strftime('%d %b %Y')}"
        ruta = nuevo_proyecto(nombre, ST.publicacion, ST.periodo)
        guardar_ultimo(ruta)
        self._proyecto_ruta = ruta
        self._lbl_proyecto.config(text=nombre)

    def _guardar_proyecto(self):
        """Guarda el estado actual en el archivo .bashkar activo."""
        from core.project_manager import (
            guardar_proyecto,
            guardar_ultimo,
        )
        if not self._proyecto_ruta:
            self._nuevo_proyecto_dialogo(); return
        try:
            guardar_proyecto(self._proyecto_ruta, ST, self._historial_ia)
            guardar_ultimo(self._proyecto_ruta)
            self._limpiar_modificado()
            self._lbl_proyecto.config(fg="#34D399")
            self.after(1500, lambda: self._lbl_proyecto.config(fg="#94A3B8"))
        except Exception as e:
            messagebox.showerror("Error al guardar", str(e))

    def _on_cerrar(self):
        """Autosave al cerrar la ventana."""
        # Detener dictado activo antes de cerrar
        if getattr(self, "_dictar_session", None) is not None:
            try:
                self._dictar_session.detener()
            except Exception:
                pass
        if self._proyecto_ruta:
            try:
                from core.project_manager import guardar_proyecto, guardar_ultimo
                guardar_proyecto(self._proyecto_ruta, ST, self._historial_ia)
                guardar_ultimo(self._proyecto_ruta)
            except Exception:
                pass
        self.destroy()
        sys.exit(0)

    def _marcar_modificado(self):
        """Activa indicador ● en la etiqueta del proyecto para señalar cambios pendientes."""
        era_falso = not self._hay_cambios
        self._hay_cambios = True
        try:
            txt = self._lbl_proyecto.cget("text")
            if not txt.startswith("●  "):
                self._lbl_proyecto.config(text=f"●  {txt}", fg="#F59E0B")
        except Exception:
            pass
        # Iniciar pulso solo la primera vez que se activa
        if era_falso:
            self.after(800, lambda: self._pulso_modificado(True))

    def _limpiar_modificado(self):
        """Quita el indicador ● tras guardar."""
        self._hay_cambios = False
        try:
            txt = self._lbl_proyecto.cget("text")
            if txt.startswith("●  "):
                self._lbl_proyecto.config(text=txt[3:], fg="#94A3B8")
        except Exception:
            pass

    def _autoguardar_periodico(self):
        """Guarda silenciosamente el proyecto cada 3 minutos si hay cambios."""
        if self._proyecto_ruta and self._hay_cambios:
            try:
                from core.project_manager import guardar_proyecto, guardar_ultimo
                guardar_proyecto(self._proyecto_ruta, ST, self._historial_ia)
                guardar_ultimo(self._proyecto_ruta)
                self._limpiar_modificado()
            except Exception:
                pass
        # Reprogramar para el siguiente ciclo
        self.after(180_000, self._autoguardar_periodico)

    # ══════════════════════════════════════════════════════════════════════════
    # TOAST NOTIFICATIONS
    # ══════════════════════════════════════════════════════════════════════════

    _TOAST_Y_OFFSET = 24   # separación entre toasts apilados

    def toast(self, mensaje: str, tipo: str = "info", duracion: int = 3500):
        """
        Muestra una notificación no-modal en la esquina inferior derecha.
        tipo: "info" | "ok" | "warn" | "error"
        """
        colores = {
            "info":  ("#252526", "#4FC1FF", "#4FC1FF"),
            "ok":    ("#1A2F1A", "#4EC9B0", "#4EC9B0"),
            "warn":  ("#2F2A1A", "#F59E0B", "#F59E0B"),
            "error": ("#2F1A1A", "#F44747", "#F44747"),
        }
        iconos = {"info": "ℹ", "ok": "✓", "warn": "⚠", "error": "✕"}
        bg, fg_icon, fg_bor = colores.get(tipo, colores["info"])

        # Calcular posición — apilar toasts activos
        if not hasattr(self, "_toasts_activos"):
            self._toasts_activos = []

        win = tk.Toplevel(self)
        win.wm_overrideredirect(True)
        win.attributes("-topmost", True)
        win.configure(bg=fg_bor)   # borde de color

        inner = tk.Frame(win, bg=bg, padx=14, pady=10)
        inner.pack(padx=1, pady=1)

        tk.Label(inner, text=iconos.get(tipo, "ℹ"), bg=bg, fg=fg_icon,
                 font=("Segoe UI", 13, "bold")).pack(side="left", padx=(0, 10))
        tk.Label(inner, text=mensaje, bg=bg, fg=TXT_PRI,
                 font=("Segoe UI", 9), wraplength=280,
                 justify="left").pack(side="left")

        # Botón cerrar
        def _cerrar():
            try:
                self._toasts_activos.remove(win)
                win.destroy()
                self._reposicionar_toasts()
            except Exception:
                pass

        tk.Label(inner, text="×", bg=bg, fg=TXT_SEC,
                 font=("Segoe UI", 11), cursor="hand2",
                 padx=6).pack(side="right", padx=(10, 0))
        inner.children[list(inner.children)[-1]].bind("<Button-1>", lambda e: _cerrar())

        win.update_idletasks()
        self._toasts_activos.append(win)
        self._reposicionar_toasts()

        # Animación entrada: fade in
        win.attributes("-alpha", 0.0)
        self._toast_fade(win, 0.0, 1.0, 30, duracion, _cerrar)

    def _reposicionar_toasts(self):
        """Recalcula posición Y de todos los toasts activos."""
        if not hasattr(self, "_toasts_activos"):
            return
        sw = self.winfo_screenwidth()
        sh = self.winfo_screenheight()
        y = sh - 72
        for w in reversed(self._toasts_activos):
            try:
                w.update_idletasks()
                ww = w.winfo_reqwidth()
                wh = w.winfo_reqheight()
                w.geometry(f"+{sw - ww - 20}+{y - wh}")
                y -= wh + 8
            except Exception:
                pass

    def _toast_fade(self, win, alpha: float, target: float,
                    steps: int, duracion: int, on_done):
        """Anima alpha de `alpha` a `target` en `steps` pasos."""
        if not win.winfo_exists():
            return
        step = (target - alpha) / max(steps, 1)
        alpha = round(alpha + step, 3)
        try:
            win.attributes("-alpha", alpha)
        except Exception:
            return
        if (step > 0 and alpha < target) or (step < 0 and alpha > target):
            self.after(16, lambda: self._toast_fade(win, alpha, target, steps - 1,
                                                     duracion, on_done))
        elif target == 1.0:
            # Fade-in terminó → esperar duracion → fade-out
            self.after(duracion, lambda: self._toast_fade(win, 1.0, 0.0, 20,
                                                           0, on_done))
        else:
            # Fade-out terminó
            on_done()

    # ══════════════════════════════════════════════════════════════════════════
    # COMMAND PALETTE  (Ctrl+K)
    # ══════════════════════════════════════════════════════════════════════════

    # Catálogo de comandos: (etiqueta, descripción, acción_callable)
    _COMANDOS_PALETTE: list = []   # se construye en _init_command_palette

    def _init_command_palette(self):
        """Construye el catálogo de comandos disponibles."""
        nav = [
            # Navegación de páginas
            ("Ir a Configuración",        "cfg",   lambda: self._mostrar_pagina("cfg")),
            ("Ir a Etiquetador de zonas",  "etz",   lambda: self._mostrar_pagina("etz")),
            ("Ir a Extracción OCR",        "ocr",   lambda: self._mostrar_pagina("ocr")),
            ("Ir a Conversor PDF",         "conv",  lambda: self._mostrar_pagina("conv")),
            ("Ir a Extracción multimodal IA", "mmx", lambda: self._mostrar_pagina("mmx")),
            ("Ir a Normalizar",            "norm",  lambda: self._mostrar_pagina("norm")),
            ("Ir a Segmentar",             "seg",   lambda: self._mostrar_pagina("seg")),
            ("Ir a Análisis textual",      "anal",  lambda: self._mostrar_pagina("anal")),
            ("Ir a NER / Entidades",       "ner",   lambda: self._mostrar_pagina("ner")),
            ("Ir a Anotaciones",           "anot",  lambda: self._mostrar_pagina("anot")),
            ("Ir a Colocaciones",          "coloc", lambda: self._mostrar_pagina("coloc")),
            ("Ir a Lingüística computacional", "ling", lambda: self._mostrar_pagina("ling")),
            ("Ir a Visualizar",            "vis",   lambda: self._mostrar_pagina("vis")),
            ("Ir a Comparativo",           "comp",  lambda: self._mostrar_pagina("comp")),
            ("Ir a Publicar / Resultados", "res",   lambda: self._mostrar_pagina("res")),
            ("Ir a Bitácora",              "bit",   lambda: self._bitacora_abrir()),
        ]
        acciones = [
            ("💾  Guardar proyecto",       "guardar",   self._guardar_proyecto),
            ("📂  Abrir gestor proyectos", "proyectos", self._abrir_gestor_proyectos),
            ("➕  Nuevo proyecto",         "nuevo",     self._nuevo_proyecto_dialogo),
            ("⚡  Modo análisis rápido",   "adhoc",     self._modo_adhoc),
            ("📓  Abrir bitácora",         "bitacora",  self._bitacora_abrir),
            ("🌙  Cambiar tema claro/oscuro", "tema",   self._toggle_theme),
            ("🖼  Análisis de encuadre (framing)", "frame",
             lambda: self._ir_a_ling_pestania(6)),
            ("⚖  Polaridad discriminante", "pol",
             lambda: self._ir_a_ling_pestania(7)),
            ("🔍  Revisión NER (validar entidades)", "revner",
             lambda: self._ir_a_ling_pestania(8)),
            ("✔  Validación metodológica (Kappa)", "valida",
             lambda: self._ir_a_ling_pestania(9)),
            ("🕸  Grafo canónico (entidades + relaciones)", "grafo",
             lambda: self._mostrar_pagina("red")),
            ("✔  Verificación OCR palabra por palabra", "verificar",
             self._verif_abrir),
            ("🔁  Detectar cabeceras repetidas (Etiquetador)", "cabeceras",
             self._etz_detectar_cabeceras),
            ("💾  Guardar como… (PDF/TEI/Excel/texto)", "guardarcomo",
             self._exp_abrir_dialogo),
        ]
        self._COMANDOS_PALETTE = [
            {"label": lab, "tags": tag, "accion": fn}
            for lab, tag, fn in (nav + acciones)
        ]

    def _abrir_command_palette(self, event=None):
        """Abre (o cierra si ya está abierta) la Command Palette."""
        if getattr(self, "_cp_win", None) and self._cp_win.winfo_exists():
            self._cp_win.destroy()
            self._cp_win = None
            return

        if not self._COMANDOS_PALETTE:
            self._init_command_palette()

        sw = self.winfo_screenwidth()
        w_pal = 500
        x = self.winfo_rootx() + (self.winfo_width() - w_pal) // 2
        y = self.winfo_rooty() + 60

        win = tk.Toplevel(self)
        win.wm_overrideredirect(True)
        win.attributes("-topmost", True)
        win.geometry(f"{w_pal}+{x}+{y}")
        win.configure(bg=CARD_BOR)   # borde 1px simulado
        self._cp_win = win

        # Glass frame interior
        frame = tk.Frame(win, bg=CARD_BG, bd=0)
        frame.pack(padx=1, pady=1, fill="both", expand=True)

        # Header con título
        hdr = tk.Frame(frame, bg=CARD_BG, pady=6)
        hdr.pack(fill="x", padx=10)
        tk.Label(hdr, text="⌨  Comandos", bg=CARD_BG, fg=TXT_SEC,
                 font=("Segoe UI", 8)).pack(side="left")
        tk.Label(hdr, text="Esc para cerrar", bg=CARD_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="right")

        # Separador
        tk.Frame(frame, bg=CARD_BOR, height=1).pack(fill="x")

        # Campo de búsqueda
        var_q = tk.StringVar()
        entry = tk.Entry(frame, textvariable=var_q, bg="#1E1E1E", fg=TXT_PRI,
                         insertbackground=TXT_PRI, font=("Segoe UI", 12),
                         relief="flat", bd=0)
        entry.pack(fill="x", padx=14, pady=10, ipady=6)
        entry.focus_set()

        # Separador
        tk.Frame(frame, bg=CARD_BOR, height=1).pack(fill="x")

        # Lista de resultados
        listbox = tk.Listbox(frame, bg=CARD_BG, fg=TXT_PRI,
                             selectbackground=AZ3, selectforeground="#FFFFFF",
                             font=("Segoe UI", 10), relief="flat", bd=0,
                             activestyle="none", height=10)
        listbox.pack(fill="both", expand=True, padx=0, pady=4)

        # Hint inferior
        hint = tk.Frame(frame, bg="#1A1A1A", pady=5)
        hint.pack(fill="x")
        for txt in ["↑↓ navegar", "Enter ejecutar", "Ctrl+K cerrar"]:
            tk.Label(hint, text=txt, bg="#1A1A1A", fg=TXT_DIM,
                     font=("Segoe UI", 8)).pack(side="left", padx=10)

        def _poblar(q: str = ""):
            listbox.delete(0, "end")
            q_low = q.lower().strip()
            for cmd in self._COMANDOS_PALETTE:
                lbl = cmd["label"]
                if not q_low or q_low in lbl.lower() or q_low in cmd["tags"].lower():
                    listbox.insert("end", f"  {lbl}")
            if listbox.size():
                listbox.selection_set(0)

        def _ejecutar(idx=None):
            if idx is None:
                sel = listbox.curselection()
                if not sel:
                    return
                idx = sel[0]
            q_low = var_q.get().lower().strip()
            coincidentes = [c for c in self._COMANDOS_PALETTE
                            if not q_low or q_low in c["label"].lower()
                            or q_low in c["tags"].lower()]
            if idx < len(coincidentes):
                win.destroy()
                self._cp_win = None
                try:
                    coincidentes[idx]["accion"]()
                except Exception:
                    pass

        def _on_key(event):
            if event.keysym == "Escape":
                win.destroy(); self._cp_win = None
            elif event.keysym == "Return":
                _ejecutar()
            elif event.keysym == "Down":
                cur = listbox.curselection()
                nxt = (cur[0] + 1) if cur else 0
                if nxt < listbox.size():
                    listbox.selection_clear(0, "end")
                    listbox.selection_set(nxt)
                    listbox.see(nxt)
            elif event.keysym == "Up":
                cur = listbox.curselection()
                prv = (cur[0] - 1) if cur else 0
                if prv >= 0:
                    listbox.selection_clear(0, "end")
                    listbox.selection_set(prv)
                    listbox.see(prv)

        var_q.trace_add("write", lambda *_: _poblar(var_q.get()))
        entry.bind("<Key>", _on_key)
        listbox.bind("<Double-Button-1>",
                     lambda e: _ejecutar(listbox.nearest(e.y)))
        listbox.bind("<Return>", lambda e: _ejecutar())
        win.bind("<Escape>", lambda e: (win.destroy(), setattr(self, "_cp_win", None)))

        # Cerrar al perder foco
        win.bind("<FocusOut>", lambda e: self.after(150, _check_focus))
        def _check_focus():
            try:
                if win.winfo_exists() and win.focus_get() is None:
                    win.destroy(); self._cp_win = None
            except Exception:
                pass

        _poblar()

    # ══════════════════════════════════════════════════════════════════════════
    # MICRO-ANIMACIONES
    # ══════════════════════════════════════════════════════════════════════════

    def _fade_pagina(self, frame, steps: int = 8):
        """Fade-in suave de un frame al mostrarlo (alpha via after)."""
        # tkinter no soporta alpha por widget, pero simulamos con
        # cambios rápidos de bg que producen efecto visual perceptible.
        # En Windows con DWM el Toplevel sí soporta alpha; aquí usamos
        # una variante: revelar el frame con delay mínimo para dar sensación
        # de transición sin coste de rendimiento.
        frame.update_idletasks()

    def _pulso_modificado(self, activo: bool = True):
        """Pulsa el indicador ● alternando entre ámbar y naranja."""
        if not activo or not self._hay_cambios:
            return
        try:
            txt = self._lbl_proyecto.cget("text")
            if not txt.startswith("●"):
                return
            cur = self._lbl_proyecto.cget("fg")
            nxt = "#E6A817" if cur == "#F59E0B" else "#F59E0B"
            self._lbl_proyecto.config(fg=nxt)
            self.after(800, lambda: self._pulso_modificado(self._hay_cambios))
        except Exception:
            pass

    # ══════════════════════════════════════════════════════════════════════════
    # SKELETON LOADING
    # ══════════════════════════════════════════════════════════════════════════

    def _skeleton_show(self, parent, n_filas: int = 6) -> tk.Frame:
        """
        Crea y muestra un skeleton placeholder en `parent`.
        Retorna el frame para poder destruirlo con _skeleton_hide().
        """
        sk = tk.Frame(parent, bg=CONTENT_BG)
        sk.place(relx=0, rely=0, relwidth=1, relheight=1)

        for i in range(n_filas):
            row = tk.Frame(sk, bg=CONTENT_BG)
            row.pack(fill="x", padx=24, pady=6)
            # línea corta (título simulado)
            w_pct = 0.4 if i % 3 == 0 else 0.75
            bar_bg = CARD_BOR
            tk.Frame(row, bg=bar_bg, height=10,
                     width=int(360 * w_pct)).pack(side="left", fill="y")
        self._skeleton_animar(sk, 0)
        return sk

    def _skeleton_animar(self, sk: tk.Frame, fase: int):
        """Pulsa el brillo de las barras skeleton."""
        if not sk.winfo_exists():
            return
        colores = [CARD_BOR, "#3A3A3A", CARD_BOR]
        c = colores[fase % len(colores)]
        try:
            for row in sk.winfo_children():
                for bar in row.winfo_children():
                    bar.config(bg=c)
        except Exception:
            pass
        self.after(400, lambda: self._skeleton_animar(sk, fase + 1))

    def _skeleton_hide(self, sk: tk.Frame | None):
        """Destruye el skeleton."""
        if sk and sk.winfo_exists():
            sk.destroy()

    # ══════════════════════════════════════════════════════════════════════════
    # PROGRESSIVE DISCLOSURE — helper de secciones avanzadas
    # ══════════════════════════════════════════════════════════════════════════

    def _mk_avanzado(self, parent, label: str = "Opciones avanzadas",
                     build_fn=None) -> tk.Frame:
        """
        Crea un bloque colapsable 'Opciones avanzadas' en `parent`.
        `build_fn(frame)` construye el contenido al expandir.
        Retorna el frame contenedor exterior.
        """
        outer = tk.Frame(parent, bg=CONTENT_BG)
        outer.pack(fill="x", padx=0, pady=(2, 0))

        var_open = tk.BooleanVar(value=False)
        _built   = [False]

        hdr = tk.Frame(outer, bg=CONTENT_BG, cursor="hand2")
        hdr.pack(fill="x")
        lbl_arrow = tk.Label(hdr, text="▶", bg=CONTENT_BG, fg=TXT_DIM,
                             font=("Segoe UI", 8))
        lbl_arrow.pack(side="left", padx=(4, 2))
        tk.Label(hdr, text=label, bg=CONTENT_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8, "italic")).pack(side="left")

        body = tk.Frame(outer, bg=CONTENT_BG)

        def _toggle(e=None):
            if var_open.get():
                var_open.set(False)
                lbl_arrow.config(text="▶")
                body.pack_forget()
            else:
                var_open.set(True)
                lbl_arrow.config(text="▼")
                if not _built[0] and build_fn:
                    build_fn(body)
                    _built[0] = True
                body.pack(fill="x", pady=(4, 0))

        hdr.bind("<Button-1>", _toggle)
        lbl_arrow.bind("<Button-1>", _toggle)
        for child in hdr.winfo_children():
            child.bind("<Button-1>", _toggle)

        return outer

    # ══════════════════════════════════════════════════════════════════════════
    # GLASSMORPHISM — paneles flotantes con estilo mejorado
    # ══════════════════════════════════════════════════════════════════════════

    def _mk_glass_toplevel(self, titulo: str, ancho: int = 700,
                           alto: int = 500) -> tuple["tk.Toplevel", "tk.Frame"]:
        """
        Crea un Toplevel con estilo glass: borde de color,
        shadow simulada, header con título.
        Retorna (win, content_frame).
        """
        win = tk.Toplevel(self)
        win.title(titulo)
        win.geometry(f"{ancho}x{alto}")
        win.resizable(True, True)
        win.configure(bg=HDR_LINE)   # borde 1px azul

        # Shadow: frame exterior ligeramente más oscuro
        shadow = tk.Frame(win, bg="#1A1A1A")
        shadow.pack(padx=1, pady=1, fill="both", expand=True)

        # Header del panel
        hdr = tk.Frame(shadow, bg=CARD_BG, pady=8)
        hdr.pack(fill="x")
        tk.Label(hdr, text=titulo, bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 10, "bold")).pack(side="left", padx=14)
        btn_close = tk.Label(hdr, text="✕", bg=CARD_BG, fg=TXT_SEC,
                             font=("Segoe UI", 11), cursor="hand2", padx=10)
        btn_close.pack(side="right")
        btn_close.bind("<Button-1>", lambda e: win.destroy())
        btn_close.bind("<Enter>",    lambda e: btn_close.config(fg=ROJO))
        btn_close.bind("<Leave>",    lambda e: btn_close.config(fg=TXT_SEC))

        tk.Frame(shadow, bg=CARD_BOR, height=1).pack(fill="x")

        content = tk.Frame(shadow, bg=CONTENT_BG)
        content.pack(fill="both", expand=True)

        return win, content

    # ══════════════════════════════════════════════════════════════════════════
    # BITÁCORA DE INVESTIGACIÓN
    # ══════════════════════════════════════════════════════════════════════════

    def _bitacora_engine(self) -> "BitacoraEngine | None":
        """Retorna instancia de BitacoraEngine para el proyecto activo."""
        if not ST.ruta_db:
            return None
        if self._bitacora_engine is None:
            try:
                from core.bitacora_engine import BitacoraEngine
                self._bitacora_engine = BitacoraEngine(ST.ruta_db)
            except Exception:
                return None
        return self._bitacora_engine

    def _bitacora_abrir(self):
        """Abre (o trae al frente) la ventana flotante de bitácora."""
        if self._bitacora_win is not None:
            try:
                self._bitacora_win.lift()
                self._bitacora_win.focus_set()
                return
            except Exception:
                self._bitacora_win = None

        win, bit_content = self._mk_glass_toplevel("📓 Bitácora de investigación",
                                                     ancho=760, alto=560)
        win.attributes("-topmost", True)
        win.protocol("WM_DELETE_WINDOW", lambda: self._bitacora_cerrar(win))
        self._bitacora_win = win

        nb = ttk.Notebook(bit_content)
        nb.pack(fill="both", expand=True, padx=8, pady=8)

        # ── Tab 1: Nueva nota ──────────────────────────────────────────────────
        tab_nueva = tk.Frame(nb, bg=CONTENT_BG, padx=12, pady=10)
        nb.add(tab_nueva, text="  ➕ Nueva nota  ")
        self._bitacora_build_nueva(tab_nueva)

        # ── Tab 2: Todas las notas ─────────────────────────────────────────────
        tab_lista = tk.Frame(nb, bg=CONTENT_BG, padx=8, pady=6)
        nb.add(tab_lista, text="  📋 Todas las notas  ")
        self._bitacora_build_lista(tab_lista)

        win._nb = nb
        self._bitacora_refrescar_lista()

    def _bitacora_cerrar(self, win):
        win.destroy()
        self._bitacora_win = None

    def _bitacora_build_nueva(self, parent):
        """Construye el formulario de nueva nota."""
        # Tipo de nota
        row_tipo = tk.Frame(parent, bg=CONTENT_BG)
        row_tipo.pack(fill="x", pady=(0, 6))
        tk.Label(row_tipo, text="Tipo:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold"), width=10, anchor="w").pack(side="left")
        self._bvar_tipo = tk.StringVar(value="libre")
        for val, etiq in [("libre", "📝 Libre"), ("hipotesis", "💡 Hipótesis"), ("cita", "📌 Cita")]:
            ttk.Radiobutton(row_tipo, text=etiq, variable=self._bvar_tipo,
                            value=val, command=self._bitacora_on_tipo).pack(side="left", padx=6)

        # Estado (solo si hipótesis)
        self._row_estado = tk.Frame(parent, bg=CONTENT_BG)
        tk.Label(self._row_estado, text="Estado:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold"), width=10, anchor="w").pack(side="left")
        self._bvar_estado = tk.StringVar(value="abierta")
        for val, etiq in [("abierta","🔵 Abierta"),("confirmada","✅ Confirmada"),
                          ("descartada","❌ Descartada"),("revisada","🔄 Revisada")]:
            ttk.Radiobutton(self._row_estado, text=etiq,
                            variable=self._bvar_estado, value=val).pack(side="left", padx=4)
        # Ocultar inicialmente
        self._bitacora_on_tipo()

        # Referencia (pre-rellena desde módulo activo)
        row_ref = tk.Frame(parent, bg=CONTENT_BG)
        row_ref.pack(fill="x", pady=(0, 4))
        tk.Label(row_ref, text="Referencia:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold"), width=10, anchor="w").pack(side="left")
        self._bvar_ref_num = tk.StringVar()
        self._bvar_ref_pag = tk.StringVar()
        tk.Entry(row_ref, textvariable=self._bvar_ref_num, width=20,
                 font=("Segoe UI", 9), bg="#1C2128", fg=TXT_PRI,
                 relief="solid", bd=1, insertbackground=TXT_PRI).pack(side="left", padx=(0, 4))
        tk.Label(row_ref, text="pág:", bg=CONTENT_BG,
                 fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left")
        tk.Entry(row_ref, textvariable=self._bvar_ref_pag, width=10,
                 font=("Segoe UI", 9), bg="#1C2128", fg=TXT_PRI,
                 relief="solid", bd=1, insertbackground=TXT_PRI).pack(side="left", padx=(2, 8))
        tk.Label(row_ref, text="módulo:", bg=CONTENT_BG,
                 fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left")
        self._bvar_modulo = tk.StringVar()
        tk.Entry(row_ref, textvariable=self._bvar_modulo, width=12,
                 font=("Segoe UI", 9), bg="#1C2128", fg=TXT_PRI,
                 relief="solid", bd=1, insertbackground=TXT_PRI,
                 state="readonly").pack(side="left", padx=2)

        # Etiquetas
        row_tags = tk.Frame(parent, bg=CONTENT_BG)
        row_tags.pack(fill="x", pady=(0, 4))
        tk.Label(row_tags, text="Etiquetas:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold"), width=10, anchor="w").pack(side="left")
        self._bvar_tags = tk.StringVar()
        tk.Entry(row_tags, textvariable=self._bvar_tags, width=50,
                 font=("Segoe UI", 9), bg="#1C2128", fg=TXT_PRI,
                 relief="solid", bd=1, insertbackground=TXT_PRI).pack(side="left")
        tk.Label(row_tags, text="(separadas por coma)", bg=CONTENT_BG,
                 fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left", padx=6)

        # Texto
        tk.Label(parent, text="Nota:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(4, 2))
        self._btxt_nota = scrolledtext.ScrolledText(
            parent, height=8, bg="#1C2128", fg=TXT_PRI,
            insertbackground=TXT_PRI, font=("Segoe UI", 10),
            relief="solid", bd=1, wrap="word")
        self._btxt_nota.pack(fill="both", expand=True)

        # Botón guardar
        row_btn = tk.Frame(parent, bg=CONTENT_BG)
        row_btn.pack(fill="x", pady=(8, 0))
        ttk.Button(row_btn, text="💾 Guardar nota", style="P.TButton",
                   command=self._bitacora_guardar_nota).pack(side="left")
        self._blbl_ok = tk.Label(row_btn, text="", bg=CONTENT_BG,
                                  fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._blbl_ok.pack(side="left", padx=10)

    def _bitacora_on_tipo(self):
        """Muestra/oculta el selector de estado según tipo."""
        if self._bvar_tipo.get() == "hipotesis":
            self._row_estado.pack(fill="x", pady=(0, 4))
        else:
            self._row_estado.pack_forget()

    def _bitacora_build_lista(self, parent):
        """Construye la vista de todas las notas con filtros."""
        # Barra de filtros
        fbar = tk.Frame(parent, bg=CONTENT_BG)
        fbar.pack(fill="x", pady=(0, 6))

        tk.Label(fbar, text="Tipo:", bg=CONTENT_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left")
        self._bflt_tipo = tk.StringVar(value="todos")
        ttk.Combobox(fbar, textvariable=self._bflt_tipo,
                     values=["todos", "libre", "hipotesis", "cita"],
                     state="readonly", width=10,
                     font=("Segoe UI", 8)).pack(side="left", padx=(2, 8))

        tk.Label(fbar, text="Estado:", bg=CONTENT_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left")
        self._bflt_estado = tk.StringVar(value="todos")
        ttk.Combobox(fbar, textvariable=self._bflt_estado,
                     values=["todos", "abierta", "confirmada", "descartada", "revisada"],
                     state="readonly", width=12,
                     font=("Segoe UI", 8)).pack(side="left", padx=(2, 8))

        tk.Label(fbar, text="Buscar:", bg=CONTENT_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left")
        self._bflt_q = tk.StringVar()
        tk.Entry(fbar, textvariable=self._bflt_q, width=20,
                 font=("Segoe UI", 9), bg="#1C2128", fg=TXT_PRI,
                 relief="solid", bd=1, insertbackground=TXT_PRI).pack(side="left", padx=(2, 6))
        ttk.Button(fbar, text="🔍", style="S.TButton",
                   command=self._bitacora_refrescar_lista).pack(side="left")
        ttk.Button(fbar, text="📄 Exportar Markdown", style="S.TButton",
                   command=self._bitacora_exportar_md).pack(side="right")

        # Treeview
        cols = ("tipo", "ref", "texto", "etiquetas", "fecha")
        self._btv = ttk.Treeview(parent, columns=cols, show="headings", height=12)
        self._btv.heading("tipo",     text="Tipo",        anchor="w")
        self._btv.heading("ref",      text="Referencia",  anchor="w")
        self._btv.heading("texto",    text="Texto",        anchor="w")
        self._btv.heading("etiquetas",text="Etiquetas",   anchor="w")
        self._btv.heading("fecha",    text="Fecha",       anchor="w")
        self._btv.column("tipo",     width=90,  stretch=False)
        self._btv.column("ref",      width=110, stretch=False)
        self._btv.column("texto",    width=300)
        self._btv.column("etiquetas",width=120, stretch=False)
        self._btv.column("fecha",    width=90,  stretch=False)

        # Colores por tipo
        self._btv.tag_configure("libre",     background="#1C2128", foreground=TXT_SEC)
        self._btv.tag_configure("hipotesis", background="#1C3A5A", foreground="#93C5FD")
        self._btv.tag_configure("cita",      background="#1A2E1A", foreground="#86EFAC")

        sb = ttk.Scrollbar(parent, orient="vertical", command=self._btv.yview)
        self._btv.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        self._btv.pack(fill="both", expand=True)
        self._btv.bind("<Double-1>", self._bitacora_editar_seleccion)

        # Botón eliminar
        btn_row = tk.Frame(parent, bg=CONTENT_BG)
        btn_row.pack(fill="x", pady=(4, 0))
        ttk.Button(btn_row, text="🗑 Eliminar seleccionada", style="S.TButton",
                   command=self._bitacora_eliminar_seleccion).pack(side="left")
        self._btv_ids: dict[str, int] = {}  # iid → nota_id

    def _bitacora_guardar_nota(self):
        """Guarda la nota actual en la BD."""
        eng = self._bitacora_engine()
        if eng is None:
            messagebox.showwarning("Sin proyecto",
                                   "Abre o crea un proyecto antes de guardar notas.")
            return
        texto = self._btxt_nota.get("1.0", "end-1c").strip()
        if not texto:
            messagebox.showwarning("Nota vacía", "Escribe algo en el campo de nota.")
            return
        tags_raw = self._bvar_tags.get()
        etiquetas = [t.strip() for t in tags_raw.split(",") if t.strip()]
        nota = {
            "tipo":          self._bvar_tipo.get(),
            "estado":        self._bvar_estado.get() if self._bvar_tipo.get() == "hipotesis" else None,
            "texto":         texto,
            "etiquetas":     etiquetas,
            "ref_numero":    self._bvar_ref_num.get().strip(),
            "ref_pagina":    self._bvar_ref_pag.get().strip(),
            "modulo_origen": self._bvar_modulo.get().strip(),
        }
        eng.insertar(nota)
        self._btxt_nota.delete("1.0", "end")
        self._bvar_tags.set("")
        self._blbl_ok.config(text="✅ Nota guardada")
        self.after(2000, lambda: self._blbl_ok.config(text=""))
        self._bitacora_refrescar_lista()

    def _bitacora_refrescar_lista(self):
        """Recarga el treeview con las notas filtradas."""
        if not hasattr(self, "_btv"):
            return
        eng = self._bitacora_engine()
        if eng is None:
            return
        tipo   = self._bflt_tipo.get()  if hasattr(self, "_bflt_tipo")   else "todos"
        estado = self._bflt_estado.get() if hasattr(self, "_bflt_estado") else "todos"
        q      = self._bflt_q.get()     if hasattr(self, "_bflt_q")      else ""
        notas = eng.listar(
            tipo=None if tipo == "todos" else tipo,
            estado=None if estado == "todos" else estado,
            q=q or None,
        )
        self._btv.delete(*self._btv.get_children())
        self._btv_ids = {}
        iconos = {"libre": "📝", "hipotesis": "💡", "cita": "📌"}
        for n in notas:
            tipo_n = n.get("tipo", "libre")
            icono  = iconos.get(tipo_n, "")
            estado_n = n.get("estado") or ""
            ref    = " ".join(filter(None, [n.get("ref_numero"), n.get("ref_pagina")]))
            texto  = n.get("texto", "")
            preview = texto[:60].replace("\n", " ") + ("…" if len(texto) > 60 else "")
            tags_str = ", ".join(n.get("etiquetas", []))
            fecha  = (n.get("creado") or "")[:10]
            tipo_display = f"{icono} {tipo_n}" + (f" · {estado_n}" if estado_n else "")
            iid = self._btv.insert("", "end",
                values=(tipo_display, ref, preview, tags_str, fecha),
                tags=(tipo_n,))
            self._btv_ids[iid] = n["id"]

    def _bitacora_editar_seleccion(self, event=None):
        """Abre diálogo de edición para la nota seleccionada."""
        sel = self._btv.selection()
        if not sel:
            return
        nota_id = self._btv_ids.get(sel[0])
        if nota_id is None:
            return
        eng = self._bitacora_engine()
        if eng is None:
            return
        nota = eng.obtener(nota_id)
        if not nota:
            return
        # Pre-rellenar formulario de nueva nota y cambiar a ese tab
        self._bvar_tipo.set(nota.get("tipo", "libre"))
        self._bitacora_on_tipo()
        if nota.get("estado"):
            self._bvar_estado.set(nota["estado"])
        self._bvar_ref_num.set(nota.get("ref_numero", ""))
        self._bvar_ref_pag.set(nota.get("ref_pagina", ""))
        self._bvar_tags.set(", ".join(nota.get("etiquetas", [])))
        self._btxt_nota.delete("1.0", "end")
        self._btxt_nota.insert("1.0", nota.get("texto", ""))
        if self._bitacora_win:
            self._bitacora_win._nb.select(0)

    def _bitacora_eliminar_seleccion(self):
        sel = self._btv.selection()
        if not sel:
            return
        if not messagebox.askyesno("Eliminar", "¿Eliminar la nota seleccionada?"):
            return
        eng = self._bitacora_engine()
        if eng is None:
            return
        nota_id = self._btv_ids.get(sel[0])
        if nota_id:
            eng.eliminar(nota_id)
        self._bitacora_refrescar_lista()

    def _bitacora_exportar_md(self):
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".md",
            filetypes=[("Markdown", "*.md"), ("Texto", "*.txt")],
            initialfile="bitacora_investigacion.md",
            title="Exportar bitácora a Markdown")
        if not dest:
            return
        eng = self._bitacora_engine()
        if eng is None:
            messagebox.showwarning("Sin proyecto", "Abre un proyecto primero.")
            return
        ruta = eng.exportar_markdown(
            Path(dest),
            publicacion=getattr(ST, "publicacion", ""))
        self.toast(f"Bitácora exportada → {Path(dest).name}", tipo="ok")

    def _bitacora_nueva_nota(self, modulo_pid: str = ""):
        """Abre la bitácora y pre-rellena la referencia con el contexto del módulo activo."""
        self._bitacora_abrir()
        if self._bitacora_win is None:
            return
        # Pre-rellenar módulo
        if hasattr(self, "_bvar_modulo"):
            self._bvar_modulo.config(state="normal")
            self._bvar_modulo.set(modulo_pid)
            self._bvar_modulo.config(state="readonly")
        # Pre-rellenar número/página según el módulo
        numero = ""
        pagina = ""
        if modulo_pid == "norm" and hasattr(self, "_norm_var_numero"):
            numero = self._norm_var_numero.get()
            if hasattr(self, "_norm_bloques") and self._norm_idx_actual >= 0:
                try:
                    pagina = self._norm_bloques[self._norm_idx_actual].get("pagina", "")
                except Exception:
                    pass
        elif modulo_pid == "seg" and hasattr(self, "_lbl_seg_n"):
            numero = getattr(ST, "corpus_meta", {}) and ""
        if hasattr(self, "_bvar_ref_num") and numero:
            self._bvar_ref_num.set(numero)
        if hasattr(self, "_bvar_ref_pag") and pagina:
            self._bvar_ref_pag.set(pagina)
        # Traer al frente y seleccionar tab nueva nota
        try:
            self._bitacora_win.lift()
            self._bitacora_win.focus_set()
            self._bitacora_win._nb.select(0)
        except Exception:
            pass

    def _nuevo_proyecto_dialogo(self):
        """Diálogo para crear un nuevo proyecto."""
        dlg = tk.Toplevel(self)
        dlg.title("Nuevo proyecto")
        dlg.geometry("420x220")
        dlg.resizable(False, False)
        dlg.configure(bg=CONTENT_BG)
        dlg.grab_set()
        dlg.transient(self)

        tk.Label(dlg, text="Nuevo proyecto de investigación",
                 bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 11, "bold")).pack(pady=(18, 12))

        form = tk.Frame(dlg, bg=CONTENT_BG)
        form.pack(fill="x", padx=24)
        form.columnconfigure(1, weight=1)

        campos = [("Nombre del proyecto:", ""), ("Publicación:", ST.publicacion),
                  ("Período:", ST.periodo)]
        vars_ = []
        for i, (lbl, val) in enumerate(campos):
            tk.Label(form, text=lbl, bg=CONTENT_BG, fg="#CDD6F4",
                     font=("Segoe UI", 9)).grid(row=i, column=0, sticky="w", pady=4)
            v = tk.StringVar(value=val)
            tk.Entry(form, textvariable=v, font=("Segoe UI", 9),
                     relief="solid", bd=1).grid(row=i, column=1, sticky="ew", padx=(8, 0))
            vars_.append(v)

        def _crear():
            nombre = vars_[0].get().strip()
            if not nombre:
                messagebox.showwarning("Campo vacío", "Escribe un nombre.", parent=dlg)
                return
            from core.project_manager import (
                guardar_proyecto,
                guardar_ultimo,
                nuevo_proyecto,
            )
            # Guardar el proyecto actual antes de cambiar
            if self._proyecto_ruta:
                try:
                    guardar_proyecto(self._proyecto_ruta, ST, self._historial_ia)
                except Exception: pass
            ruta = nuevo_proyecto(nombre, vars_[1].get().strip(),
                                   vars_[2].get().strip())
            guardar_ultimo(ruta)
            self._proyecto_ruta = ruta
            self._historial_ia  = []
            self._lbl_proyecto.config(text=nombre)
            dlg.destroy()

        btn_row = tk.Frame(dlg, bg=CONTENT_BG)
        btn_row.pack(pady=14)
        ttk.Button(btn_row, text="Crear proyecto", style="P.TButton",
                   command=_crear).pack(side="left", padx=4)
        ttk.Button(btn_row, text="Cancelar", style="S.TButton",
                   command=dlg.destroy).pack(side="left", padx=4)

    def _abrir_gestor_proyectos(self):
        """Ventana con la lista de proyectos guardados."""
        from core.project_manager import (
            cargar_proyecto,
            eliminar_proyecto,
            fecha_legible,
            guardar_proyecto,
            guardar_ultimo,
            listar_proyectos,
            progreso_str,
        )

        win = tk.Toplevel(self)
        win.title("Proyectos guardados")
        win.geometry("780x480")
        win.configure(bg=CONTENT_BG)
        win.grab_set()
        win.transient(self)

        # Cabecera
        hdr = tk.Frame(win, bg=AZ1)
        hdr.pack(fill="x")
        tk.Label(hdr, text="  📂  Proyectos de investigación",
                 bg=AZ1, fg="white",
                 font=("Segoe UI", 11, "bold")).pack(side="left", pady=10, padx=12)
        ttk.Button(hdr, text="➕ Nuevo proyecto", style="S.TButton",
                   command=lambda: [win.destroy(),
                                    self._nuevo_proyecto_dialogo()]).pack(
                   side="right", padx=12, pady=8)

        # Lista
        cols = ("Nombre", "Publicación", "Período",
                "Última sesión", "Progreso")
        tv_f = tk.Frame(win, bg=CONTENT_BG)
        tv_f.pack(fill="both", expand=True, padx=12, pady=8)
        sbv = ttk.Scrollbar(tv_f, orient="vertical")
        tv = ttk.Treeview(tv_f, columns=cols, show="headings",
                           yscrollcommand=sbv.set, height=14)
        widths = [200, 140, 100, 140, 200]
        for col, w in zip(cols, widths):
            tv.heading(col, text=col, anchor="w")
            tv.column(col, width=w, minwidth=60)
        sbv.config(command=tv.yview)
        sbv.pack(side="right", fill="y")
        tv.pack(fill="both", expand=True)

        proyectos = listar_proyectos()
        ruta_map = {}
        for p in proyectos:
            iid = tv.insert("", "end", values=(
                p["nombre"], p["publicacion"], p["periodo"],
                fecha_legible(p["modificado"]),
                progreso_str(p["progreso"]),
            ))
            ruta_map[iid] = p["ruta"]
            # Marcar el activo
            if self._proyecto_ruta and str(p["ruta"]) == str(self._proyecto_ruta):
                tv.item(iid, tags=("activo",))
        tv.tag_configure("activo", background="#1F3A5C", font=("Segoe UI", 9, "bold"))

        # Botones de acción
        act = tk.Frame(win, bg=CONTENT_BG)
        act.pack(fill="x", padx=12, pady=(0, 10))
        lbl_sel = tk.Label(act, text="Selecciona un proyecto de la lista",
                            bg=CONTENT_BG, fg="#8B949E", font=("Segoe UI", 9))
        lbl_sel.pack(side="left")

        def _abrir():
            sel = tv.selection()
            if not sel: return
            ruta = ruta_map[sel[0]]
            # Guardar actual
            if self._proyecto_ruta:
                try:
                    guardar_proyecto(self._proyecto_ruta, ST, self._historial_ia)
                except Exception: pass
            ST.reset()
            res = cargar_proyecto(ruta, ST)
            if not res["ok"]:
                messagebox.showerror("Error", res["mensaje"], parent=win); return
            guardar_ultimo(ruta)
            self._proyecto_ruta = ruta
            self._historial_ia  = res.get("historial_ia", [])
            self._sincronizar_ui_con_st()
            nombre = res.get("nombre", Path(ruta).stem)
            self._lbl_proyecto.config(text=nombre)
            self._lbl_pub_hdr.config(
                text=f"{ST.publicacion}  ·  {ST.periodo}" if ST.periodo
                else ST.publicacion)
            self._actualizar_badges()
            self.after(600, self._ocr_actualizar_estimacion)
            win.destroy()
            if res.get("migrado"):
                messagebox.showinfo(
                    "Migración completada",
                    "El proyecto fue migrado automáticamente de v10 a v11.\n"
                    "Se creó una base de datos SQLite (.db) con todos los datos.\n"
                    "Se conservó un backup del archivo original.")

        def _eliminar():
            sel = tv.selection()
            if not sel: return
            ruta = ruta_map[sel[0]]
            nombre = tv.item(sel[0])["values"][0]
            if not messagebox.askyesno(
                    "Eliminar proyecto",
                    f"¿Eliminar '{nombre}'?\nEsta acción no se puede deshacer.",
                    parent=win): return
            eliminar_proyecto(ruta)
            tv.delete(sel[0])
            if str(ruta) == str(self._proyecto_ruta):
                self._proyecto_ruta = None
                self._lbl_proyecto.config(text="Sin proyecto")

        ttk.Button(act, text="📂 Abrir", style="P.TButton",
                   command=_abrir).pack(side="right", padx=(4, 0))
        ttk.Button(act, text="🗑 Eliminar", style="S.TButton",
                   command=_eliminar).pack(side="right")
        tv.bind("<Double-1>", lambda e: _abrir())

    def _sincronizar_ui_con_st(self):
        """Actualiza los widgets de configuración con los valores de ST."""
        try:
            self._var_pub.set(ST.publicacion)
            self._var_per.set(ST.periodo)
            # Restaurar versión de normalización
            if hasattr(self, "_norm_var_version"):
                self._norm_var_version.set(getattr(ST, "norm_version", "manual"))
                self._norm_version_cambio()
            if ST.pdf_dir:
                self._var_ent.set(str(ST.pdf_dir))
                try: self._poblar_lista(ST.pdf_dir)
                except Exception: pass
                # Restaurar en el Conversor también
                if hasattr(self, "_conv_entrada"):
                    self._conv_entrada.set(str(ST.pdf_dir))
            if ST.out_dir:
                self._var_sal.set(str(ST.out_dir))
                if hasattr(self, "_conv_salida"):
                    self._conv_salida.set(str(ST.out_dir))
            if ST.api_key:
                self._var_api_key.set(ST.api_key)
            # Restaurar claves por proveedor
            keys = getattr(ST, "api_keys", {})
            if keys.get("anthropic"): self._var_key_anthropic.set(keys["anthropic"])
            if keys.get("openai"):    self._var_key_openai.set(keys["openai"])
            if keys.get("gemini"):    self._var_key_gemini.set(keys["gemini"])
            if keys.get("ollama"):    self._var_key_ollama.set(keys["ollama"])
            # Restaurar modelos por etapa
            for etapa_id, var_e in getattr(self, "_vars_modelo_etapa", {}).items():
                modelo = ST.modelos_etapa.get(etapa_id, "")
                if modelo: var_e.set(modelo)
            self._var_max_ia.set(ST.max_ia)
            if ST.campos_semillas:
                import json as _json
                txt = _json.dumps(ST.campos_semillas, ensure_ascii=False, indent=2)
                try:
                    self._txt_sem.config(state="normal")
                    self._txt_sem.delete("1.0", "end")
                    self._txt_sem.insert("1.0", txt)
                except Exception: pass
            # Restaurar índice semántico FAISS si fue cargado con el proyecto
            indice_cargado = getattr(ST, "_bsem_indice", None)
            if indice_cargado is not None and getattr(indice_cargado, "construido", False):
                if hasattr(self, "_bsem_indice"):
                    self._bsem_indice = indice_cargado
                if hasattr(self, "_lbl_bsem_estado"):
                    self._lbl_bsem_estado.config(
                        text=f"✓ {indice_cargado.n_articulos} artículos (restaurado)", fg=VERDE)
                if hasattr(self, "_lbl_bsem_n"):
                    self._lbl_bsem_n.config(text=f"{indice_cargado.n_articulos} artículos")
            # Restaurar lematización
            if hasattr(self, "_var_lematizar"):
                self._var_lematizar.set(getattr(ST, "lematizar", True))
            # Restaurar stopwords del proyecto en el widget de Collocates
            if hasattr(self, "_txt_stopwords"):
                sw = getattr(ST, "stopwords_proyecto", [])
                if sw:
                    self._txt_stopwords.delete("1.0", "end")
                    self._txt_stopwords.insert("1.0", "\n".join(sw))
                    # Aplicar al motor de collocations en esta sesión
                    try:
                        import core.collocation_engine as _ce
                        _ce.STOPWORDS_ES = _ce.STOPWORDS_ES | frozenset(sw)
                    except Exception:
                        pass
        except Exception:
            pass  # widgets aún no inicializados en el primer arranque


    # ── Cabecera de sección en el sidebar ────────────────────────────────────
    def _sb_seccion(self, parent, texto: str, scroll_fn=None):
        """Renderiza un separador + etiqueta de sección."""
        tk.Frame(parent, bg="#1E3A5F", height=1).pack(
            fill="x", padx=12, pady=(8, 2))
        lbl = tk.Label(parent, text=texto, bg=SB_BG, fg="#8B949E",
                       font=("Segoe UI", 7, "bold"), anchor="w")
        lbl.pack(fill="x", padx=16, pady=(0, 3))
        if scroll_fn:
            lbl.bind("<MouseWheel>", scroll_fn)
            lbl.bind("<Button-4>",   scroll_fn)
            lbl.bind("<Button-5>",   scroll_fn)

    # ── Rellena el área de navegación con grupos ──────────────────────────────
    def _poblar_nav_sidebar(self, parent, scroll_fn=None):
        """
        Crea todos los botones del sidebar agrupados por sección.
        Flujo principal: numerado 1-6.
        Análisis opcionales y Salida: sin número, tamaño menor.
        """
        # Separar por grupo manteniendo orden
        grupos = {}
        for entry in self._PAGINAS:
            pid, emoji, label, _, badge_attr, grupo = entry
            grupos.setdefault(grupo, []).append((pid, emoji, label, badge_attr))

        ETIQUETAS_GRUPO = {
            "flujo":    "FLUJO DE TRABAJO",
            "analisis": "ANÁLISIS OPCIONALES",
            "salida":   "EXPORTAR · COLABORAR",
        }

        flujo_num = 1
        for grupo_id in ("flujo", "analisis", "salida"):
            items = grupos.get(grupo_id, [])
            if not items:
                continue
            self._sb_seccion(parent, ETIQUETAS_GRUPO[grupo_id], scroll_fn)
            for pid, emoji, label, badge_attr in items:
                num = flujo_num if grupo_id == "flujo" else None
                btn = self._make_sb_btn(
                    parent, pid, emoji, label, num, badge_attr,
                    es_flujo=(grupo_id == "flujo"),
                    scroll_fn=scroll_fn,
                )
                self._sb_btns[pid] = btn
                if grupo_id == "flujo":
                    flujo_num += 1

    # ── Botón individual del sidebar ──────────────────────────────────────────
    def _make_sb_btn(self, parent, pid, emoji, label, num, badge_attr,
                     es_flujo: bool = True, scroll_fn=None):
        """Botón de navegación — dark mode GitHub-style."""
        h = 40 if es_flujo else 34
        row = tk.Frame(parent, bg=SB_BG, cursor="hand2", height=h)
        row.pack(fill="x", pady=0)
        row.pack_propagate(False)

        # Barra indicadora izquierda 3px
        bar = tk.Frame(row, bg=SB_BG, width=3)
        bar.pack(side="left", fill="y")

        # Número de paso (flujo principal)
        if num is not None:
            num_lbl = tk.Label(row, text=str(num), bg=SB_BG, fg=TXT_DIM,
                               font=("Segoe UI", 7, "bold"), width=2)
            num_lbl.pack(side="left", padx=(4, 0))
        else:
            num_lbl = tk.Label(row, text="", bg=SB_BG, width=0)
            num_lbl.pack(side="left", padx=(8, 0))

        # Emoji
        em_lbl = tk.Label(row, text=emoji, bg=SB_BG, fg=TXT_SEC,
                           font=("Segoe UI", 11 if es_flujo else 10))
        em_lbl.pack(side="left", padx=(6, 5))

        # Label
        txt_lbl = tk.Label(row, text=label, bg=SB_BG, fg=TXT_SEC,
                            font=("Segoe UI", 9 if es_flujo else 8), anchor="w")
        txt_lbl.pack(side="left", fill="x", expand=True)

        # Badge ✓
        badge = tk.Label(row, text="✓", bg=SB_BG, fg=VERDE,
                          font=("Segoe UI", 8, "bold"))

        widgets = {
            "row": row, "bar": bar, "em": em_lbl,
            "txt": txt_lbl, "num": num_lbl, "badge": badge,
            "badge_attr": badge_attr, "es_flujo": es_flujo,
        }

        def _click(e, p=pid):   self._mostrar_pagina(p)
        def _enter(e, p=pid):
            if self._pagina_activa.get() != p:
                # Hover: fondo gris medio, texto bien legible
                for w in (row, bar, em_lbl, txt_lbl, num_lbl):
                    w.config(bg="#2D333B")
                txt_lbl.config(fg="#CDD6F4")
                em_lbl.config(fg="#CDD6F4")
        def _leave(e, p=pid):
            if self._pagina_activa.get() != p:
                for w in (row, bar, em_lbl, txt_lbl, num_lbl):
                    w.config(bg=SB_BG)
                txt_lbl.config(fg=TXT_SEC)
                em_lbl.config(fg=TXT_SEC)

        for w in (row, bar, em_lbl, txt_lbl, num_lbl):
            w.bind("<Button-1>", _click)
            w.bind("<Enter>",    _enter)
            w.bind("<Leave>",    _leave)
            if scroll_fn:
                w.bind("<MouseWheel>", scroll_fn)
                w.bind("<Button-4>",   scroll_fn)
                w.bind("<Button-5>",   scroll_fn)

        return widgets

    def _aplicar_estilo_sb_btn(self, pid: str, widgets: dict, activo: bool):
        """Aplica estilos dark mode al botón de sidebar según estado activo/inactivo."""
        if activo:
            bg     = "#1C2128"   # fondo levemente más claro que sidebar
            fg_txt = "#E6EDF3"   # texto blanco cálido — siempre legible
            bg_bar = AB_IND      # barra azul izquierda
        else:
            bg     = SB_BG
            fg_txt = "#8B949E"   # gris medio legible sobre fondo negro
            bg_bar = SB_BG
        widgets["row"].config(bg=bg)
        widgets["bar"].config(bg=bg_bar)
        for w_key in ("em", "txt", "num"):
            if w_key in widgets:
                widgets[w_key].config(bg=bg, fg=fg_txt)
        if widgets.get("badge_attr"):
            done = getattr(ST, widgets["badge_attr"], False)
            if done:
                widgets["badge"].config(bg=bg)
                widgets["badge"].pack(side="right", padx=6)

    def _crear_pagina_scrollable(self, pid: str) -> tk.Frame:
        """Devuelve un frame interior desplazable para la página `pid`.
        Envuelve el frame en un Canvas + Scrollbar vertical dentro de
        `_content_area`; el contenedor externo se guarda en
        `_contenedores_pagina[pid]` (es lo que se muestra/oculta)."""
        cont = tk.Frame(self._content_area, bg=CONTENT_BG)
        canvas = tk.Canvas(cont, bg=CONTENT_BG, highlightthickness=0, bd=0)
        vbar = ttk.Scrollbar(cont, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side="right", fill="y")
        canvas.pack(side="left", fill="both", expand=True)

        interior = tk.Frame(canvas, bg=CONTENT_BG)
        win_id = canvas.create_window((0, 0), window=interior, anchor="nw")

        def _on_interior_config(_e=None):
            canvas.configure(scrollregion=canvas.bbox("all"))
        interior.bind("<Configure>", _on_interior_config)

        # El interior siempre ocupa el ancho del canvas (para que el pack fill=x
        # de los _build_* se extienda a lo ancho).
        def _on_canvas_config(e):
            canvas.itemconfig(win_id, width=e.width)
        canvas.bind("<Configure>", _on_canvas_config)

        # Rueda del ratón: desplaza solo cuando el puntero está sobre esta página
        # y solo si el contenido excede la altura visible.
        def _on_wheel(e):
            x0, y0, x1, y1 = canvas.bbox("all") or (0, 0, 0, 0)
            if (y1 - y0) <= canvas.winfo_height():
                return
            canvas.yview_scroll(int(-e.delta / 120), "units")

        def _bind_wheel(_e=None):
            canvas.bind_all("<MouseWheel>", _on_wheel)

        def _unbind_wheel(_e=None):
            canvas.unbind_all("<MouseWheel>")
        interior.bind("<Enter>", _bind_wheel)
        interior.bind("<Leave>", _unbind_wheel)

        self._contenedores_pagina[pid] = cont
        return interior

    def _mostrar_pagina(self, pid: str):
        # Ocultar todas (el contenedor externo canvas+scrollbar)
        for p, cont in self._contenedores_pagina.items():
            cont.pack_forget()

        # Mostrar la seleccionada
        self._contenedores_pagina[pid].pack(fill="both", expand=True)
        self._pagina_activa.set(pid)
        self._current_page = pid

        # Activar el contexto del activity bar que contiene este pid
        for ctx_id, _, _, pids in self._CONTEXTOS:
            if pid in pids:
                if self._ctx_activo.get() != ctx_id:
                    # Actualizar indicadores del activity bar sin repoblar sidebar
                    self._ctx_activo.set(ctx_id)
                    for cid, widgets in self._ab_btns.items():
                        widgets["ind"].config(bg=AB_IND if cid == ctx_id else AB_BG)
                        widgets["lbl"].config(fg=TXT_PRI if cid == ctx_id else AB_TXT)
                break

        # Actualizar estilos del sidebar
        for p, widgets in self._sb_btns.items():
            self._aplicar_estilo_sb_btn(p, widgets, p == pid)

        # Acciones al entrar a páginas específicas
        if pid == "norm":
            self.after(50, self._norm_refrescar_numeros)

    def _actualizar_badges(self):
        """Actualiza los semáforos del sidebar según estado de etapas y badges legacy."""
        _semaforo_etapa = {
            "etz":  "etz",
            "ocr":  "ocr",
            "norm": "norm",
            "seg":  "seg",
            "anal": "anal",
        }
        for pid, widgets in self._sb_btns.items():
            badge_lbl = widgets.get("badge")
            if badge_lbl is None:
                continue
            bg = SB_BG if self._pagina_activa.get() != pid else "#1C2128"

            # Semáforo de flujo (prioridad sobre badge legacy)
            if pid in _semaforo_etapa:
                estado = ST.estado_etapas.get(pid, "pending")
                if estado == "ready":
                    badge_lbl.config(text="✓", fg=VERDE, bg=bg)
                    badge_lbl.pack(side="right", padx=6)
                elif estado == "stale":
                    badge_lbl.config(text="⚠", fg="#F59E0B", bg=bg)
                    badge_lbl.pack(side="right", padx=6)
                else:
                    badge_lbl.pack_forget()
                continue

            # Badge legacy (✓) para páginas sin semáforo de flujo
            attr = widgets.get("badge_attr")
            if attr and getattr(ST, attr, False):
                badge_lbl.config(text="✓", fg=VERDE, bg=bg)
                badge_lbl.pack(side="right", padx=6)
            else:
                badge_lbl.pack_forget()

    def _abrir_docs(self):
        doc = Path(__file__).parent / "PROMPT_SISTEMA.md"
        if doc.exists():
            import subprocess
            subprocess.Popen(["notepad.exe" if platform.system()=="Windows"
                              else "open", str(doc)])

    # ── Compatibilidad con código que usa self._nb.select(n) ──────────────────
    class _FakeNb:
        """Proxy para compatibilidad con self._nb.select(índice)."""
        def __init__(self, app):
            self._app = app
        def select(self, idx):
            pids = [p for p,*_ in BashkarApp._PAGINAS]
            if isinstance(idx, int) and idx < len(pids):
                self._app._mostrar_pagina(pids[idx])

    @property
    def _nb(self): return self._FakeNb(self)

    # ── Header de sección (reutilizable) ──────────────────────────────────────
    def _page_header(self, parent, titulo: str, subtitulo: str = "",
                     emoji: str = "") -> tk.Frame:
        """Crea un encabezado estilizado para cada página."""
        hdr = tk.Frame(parent, bg=CONTENT_BG)
        hdr.pack(fill="x", padx=0, pady=0)

        # Banda de color superior
        banda = tk.Frame(hdr, bg=AZ3, height=3)
        banda.pack(fill="x")

        inner = tk.Frame(hdr, bg=CONTENT_BG)
        inner.pack(fill="x", padx=28, pady=(16, 8))

        if emoji:
            tk.Label(inner, text=emoji, bg=CONTENT_BG, fg=AZ3,
                     font=("Segoe UI", 20)).pack(side="left", padx=(0, 12))
        txt_frame = tk.Frame(inner, bg=CONTENT_BG)
        txt_frame.pack(side="left", fill="x", expand=True)
        tk.Label(txt_frame, text=titulo, bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 14, "bold"), anchor="w").pack(anchor="w")
        if subtitulo:
            tk.Label(txt_frame, text=subtitulo, bg=CONTENT_BG, fg=TXT_SEC,
                     font=("Segoe UI", 9), anchor="w").pack(anchor="w")

        # ── Guía del módulo (qué es / para qué / cómo interpretar) ──────────
        # Se inyecta según el id de página que el constructor fijó antes de
        # llamar a build_fn. Siempre visible el resumen; colapsable la guía
        # de interpretación profunda (orientada a investigación HD).
        self._inyectar_guia_modulo(hdr, getattr(self, "_guia_pagina_actual", None))

        # Línea separadora
        tk.Frame(hdr, bg=CARD_BOR, height=1).pack(fill="x")
        return hdr

    def _inyectar_guia_modulo(self, parent, page_id) -> None:
        """Bajo el encabezado del panel, muestra la guía del módulo en dos
        secciones colapsables (cerradas por defecto): «ℹ Qué es esta
        herramienta» (qué es / para qué) y «📖 Cómo interpretar los
        resultados». Silencioso si no hay guía para la página."""
        if not page_id:
            return
        try:
            from core import guia_modulos
        except ImportError:
            return
        resumen = guia_modulos.resumen_visible(page_id)
        if not resumen:
            return
        interpretacion = guia_modulos.guia_interpretacion(page_id)

        cont = tk.Frame(parent, bg=CONTENT_BG)
        cont.pack(fill="x", padx=28, pady=(0, 8))

        # «Qué es esta herramienta» — colapsable (lazy), cerrado por defecto
        def _build_resumen(frame):
            caja = tk.Frame(frame, bg=CARD_BG, bd=0)
            caja.pack(fill="x")
            tk.Frame(caja, bg=AZ3, width=3).pack(side="left", fill="y")
            tk.Label(caja, text=resumen, bg=CARD_BG, fg=TXT_SEC,
                     font=("Segoe UI", 9), anchor="w", justify="left",
                     wraplength=640).pack(side="left", fill="x",
                                          expand=True, padx=12, pady=8)
        self._mk_avanzado(cont, "ℹ  Qué es esta herramienta", _build_resumen)

        # Sección colapsable con la guía de interpretación (lazy)
        if interpretacion:
            def _build_interp(frame):
                ic = tk.Frame(frame, bg=CARD_BG)
                ic.pack(fill="x")
                tk.Frame(ic, bg="#3FB950", width=3).pack(side="left", fill="y")
                tk.Label(ic, text=interpretacion, bg=CARD_BG, fg=TXT_SEC,
                         font=("Segoe UI", 9), anchor="w", justify="left",
                         wraplength=640).pack(side="left", fill="x",
                                              expand=True, padx=12, pady=8)
            self._mk_avanzado(cont, "📖  Cómo interpretar los resultados",
                              _build_interp)

    # ── Card dark mode ─────────────────────────────────────────────────────────
    def _card(self, parent, titulo: str = "", padding: int = 16) -> tk.Frame:
        """Crea una tarjeta oscura con borde sutil estilo GitHub dark."""
        outer = tk.Frame(parent, bg=CARD_BOR, bd=0)
        outer.pack(fill="x", padx=24, pady=6)
        inner = tk.Frame(outer, bg=CARD_BG, bd=0)
        inner.pack(fill="x", padx=1, pady=1)
        if titulo:
            title_bar = tk.Frame(inner, bg=CARD_BG)
            title_bar.pack(fill="x", padx=padding, pady=(padding, 4))
            tk.Frame(title_bar, bg=AZ3, width=3, height=18).pack(side="left")
            tk.Label(title_bar, text=f"  {titulo}", bg=CARD_BG, fg=TXT_PRI,
                     font=("Segoe UI", 10, "bold")).pack(side="left")
            tk.Frame(inner, bg=CARD_BOR, height=1).pack(fill="x", padx=padding)
        content = tk.Frame(inner, bg=CARD_BG)
        content.pack(fill="x", padx=padding, pady=padding)
        return content

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 1: CONFIGURACIÓN
    # ══════════════════════════════════════════════════════════════════════════
    def _mk_ayuda(self, parent, texto: str):
        lbl = tk.Label(parent, text=" ❓", fg=TXT_DIM, bg=CARD_BG,
                       font=("Segoe UI", 9), cursor="question_arrow")
        lbl.pack(side="left")
        tip_win = [None]
        def show(event):
            tw = tk.Toplevel(lbl); tw.wm_overrideredirect(True)
            tw.wm_geometry(f"+{event.x_root+12}+{event.y_root+8}")
            tk.Message(tw, text=texto, bg="#1E3A5F", fg="#E2E8F0",
                       relief="flat", font=("Segoe UI", 9),
                       width=340, padx=12, pady=10).pack()
            tip_win[0] = tw
        def hide(_event):
            if tip_win[0]: tip_win[0].destroy(); tip_win[0]=None
        lbl.bind("<Enter>", show); lbl.bind("<Leave>", hide)
        return lbl

    def _mk_ayuda_bg(self, parent, texto: str, bg: str):
        """Como _mk_ayuda pero con bg configurable (para barras de color distinto)."""
        lbl = tk.Label(parent, text=" ❓", fg=TXT_DIM, bg=bg,
                       font=("Segoe UI", 9), cursor="question_arrow")
        lbl.pack(side="left")
        tip_win = [None]
        def show(event):
            tw = tk.Toplevel(lbl); tw.wm_overrideredirect(True)
            tw.wm_geometry(f"+{event.x_root+12}+{event.y_root+8}")
            tk.Message(tw, text=texto, bg="#1E3A5F", fg="#E2E8F0",
                       relief="flat", font=("Segoe UI", 9),
                       width=340, padx=12, pady=10).pack()
            tip_win[0] = tw
        def hide(_event):
            if tip_win[0]: tip_win[0].destroy(); tip_win[0]=None
        lbl.bind("<Enter>", show); lbl.bind("<Leave>", hide)
        return lbl

    def _build_cfg(self):
        f = self._tab_cfg
        self._page_header(f, "Configuración del corpus",
                          "Define la publicación, los archivos y los parámetros del análisis",
                          "⚙")

        # Área scrollable
        pad, _canvas = _hacer_scrollable(f, bg=CONTENT_BG)
        pad.columnconfigure(0, weight=1)
        r = 0

        # Registro de secciones colapsables: {id: {"open": BooleanVar, "body": Frame}}
        _secciones_cfg: dict = {}

        def _seccion(sec_id: str, num: str, titulo: str, subtitulo: str = "",
                     abierta: bool = True):
            """
            Crea un encabezado de sección colapsable numerado.
            Retorna el frame 'body' donde se coloca el contenido de la sección.
            El body se muestra/oculta al hacer click en el encabezado.
            """
            nonlocal r
            var_open = tk.BooleanVar(value=abierta)

            # ── Fila del encabezado ─────────────────────────────────────────────
            hdr = tk.Frame(pad, bg=CONTENT_BG, cursor="hand2")
            hdr.grid(row=r, column=0, sticky="ew", padx=24, pady=(20 if num == "1" else 8, 2))
            r += 1

            arrow_lbl = tk.Label(hdr, text="▼" if abierta else "▶",
                                 bg=CONTENT_BG, fg=AZ3,
                                 font=("Segoe UI", 10, "bold"))
            arrow_lbl.pack(side="left", padx=(0, 8))

            tk.Label(hdr, text=f"{num} · {titulo}", bg=CONTENT_BG,
                     fg=TXT_PRI, font=("Segoe UI", 11, "bold")).pack(side="left")

            if subtitulo:
                tk.Label(hdr, text=f"  {subtitulo}", bg=CONTENT_BG,
                         fg="#8B949E", font=("Segoe UI", 8)).pack(side="left", padx=(8, 0))

            # ── Body colapsable ─────────────────────────────────────────────────
            body = tk.Frame(pad, bg=CONTENT_BG)
            body.grid(row=r, column=0, sticky="ew")
            body.columnconfigure(0, weight=1)
            if not abierta:
                body.grid_remove()
            r += 1

            def _toggle(e=None):
                if var_open.get():
                    var_open.set(False)
                    body.grid_remove()
                    arrow_lbl.config(text="▶")
                else:
                    var_open.set(True)
                    body.grid()
                    arrow_lbl.config(text="▼")

            for w in (hdr, arrow_lbl):
                w.bind("<Button-1>", _toggle)
                w.bind("<Enter>", lambda e, h=hdr: h.config(bg="#161B22"))
                w.bind("<Leave>", lambda e, h=hdr: h.config(bg=CONTENT_BG))
            for child in hdr.winfo_children():
                child.bind("<Button-1>", _toggle)

            _secciones_cfg[sec_id] = {"open": var_open, "body": body}
            return body

        def card_lf(titulo, parent=None):
            """LabelFrame estilo card. Si parent es un body de sección, usa pack."""
            nonlocal r
            if parent is not None:
                # Dentro de un body colapsable — usar pack para apilar varias cards
                outer = tk.Frame(parent, bg=CARD_BOR)
                outer.pack(fill="x", padx=24, pady=4)
            else:
                outer = tk.Frame(pad, bg=CARD_BOR)
                outer.grid(row=r, column=0, sticky="ew", padx=24, pady=4)
                r += 1
            inner = tk.Frame(outer, bg=CARD_BG, padx=16, pady=14)
            inner.pack(fill="x", padx=1, pady=1)
            inner.columnconfigure(1, weight=1)
            title_row = tk.Frame(inner, bg=CARD_BG)
            title_row.grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 10))
            tk.Frame(title_row, bg=AZ3, width=3, height=16).pack(side="left")
            tk.Label(title_row, text=f"  {titulo}", bg=CARD_BG, fg=TXT_PRI,
                     font=("Segoe UI", 10, "bold")).pack(side="left")
            return inner

        def sep():
            nonlocal r
            tk.Frame(pad, bg=CARD_BOR, height=1).grid(
                row=r, column=0, sticky="ew", padx=24, pady=6); r+=1

        # ── Sección 1 ─────────────────────────────────────────────────────────
        body1 = _seccion("id", "1", "Identificación del proyecto",
                         "Estos datos aparecerán en todos los informes generados.")
        c1 = card_lf("🗞  Publicación y período", body1)
        tk.Label(c1, text="Nombre:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=1,column=0,sticky="w",pady=5)
        self._var_pub = tk.StringVar(value="Mi publicación")
        # Reflejo en tiempo real: actualizar etiqueta de nombre en topbar y sidebar
        def _sync_nombre(*_):
            nombre = self._var_pub.get().strip() or "Sin proyecto"
            if hasattr(self, "_lbl_pub_hdr"):
                self._lbl_pub_hdr.config(text=nombre)
            if hasattr(self, "_lbl_proyecto"):
                self._lbl_proyecto.config(text=nombre)
        self._var_pub.trace_add("write", _sync_nombre)
        tk.Entry(c1, textvariable=self._var_pub, width=52,
                 font=("Segoe UI",10), relief="solid", bd=1,
                 bg="#1C2128").grid(row=1,column=1,sticky="ew",padx=8)
        tk.Label(c1, text="Período:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=2,column=0,sticky="w",pady=5)
        self._var_per = tk.StringVar()
        tk.Entry(c1, textvariable=self._var_per, width=34,
                 font=("Segoe UI",10), relief="solid", bd=1,
                 bg="#1C2128").grid(row=2,column=1,sticky="w",padx=8)
        tk.Label(c1, text='Ej.: "enero–junio 1939"', bg=CARD_BG,
                 fg="#6E7681", font=("Segoe UI",8)).grid(row=2,column=2,sticky="w")
        # Tipo de corpus
        tk.Label(c1, text="Tipo:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=3,column=0,sticky="w",pady=5)
        self._var_tipo_corpus = tk.StringVar(value="revista")
        ttk.Combobox(c1, textvariable=self._var_tipo_corpus,
                     values=["revista", "periódico", "libro", "colección"],
                     state="readonly", width=18,
                     font=("Segoe UI",9)).grid(row=3,column=1,sticky="w",padx=8)
        # Idioma
        tk.Label(c1, text="Idioma:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=4,column=0,sticky="w",pady=5)
        self._var_idioma = tk.StringVar(value="español histórico")
        ttk.Combobox(c1, textvariable=self._var_idioma,
                     values=["español histórico", "español moderno", "otro"],
                     state="readonly", width=18,
                     font=("Segoe UI",9)).grid(row=4,column=1,sticky="w",padx=8)
        # Investigador e institución
        tk.Label(c1, text="Investigador:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=5,column=0,sticky="w",pady=5)
        self._var_investigador = tk.StringVar()
        tk.Entry(c1, textvariable=self._var_investigador, width=34,
                 font=("Segoe UI",9), relief="solid", bd=1,
                 bg="#1C2128").grid(row=5,column=1,sticky="w",padx=8)
        tk.Label(c1, text="Institución:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=6,column=0,sticky="w",pady=5)
        self._var_institucion = tk.StringVar(value="Instituto Caro y Cuervo")
        tk.Entry(c1, textvariable=self._var_institucion, width=40,
                 font=("Segoe UI",9), relief="solid", bd=1,
                 bg="#1C2128").grid(row=6,column=1,sticky="ew",padx=8)

        # ── Sección 2 ─────────────────────────────────────────────────────────
        body2 = _seccion("archivos", "2", "Archivos de entrada",
                         "Bashkar detecta automáticamente PDFs con texto embebido vs. escaneados.")
        c2 = card_lf("📁  Tipo y carpeta de archivos", body2)
        self._var_tipo = tk.StringVar(value="pdf")
        ttk.Radiobutton(c2, text="📄  PDFs — un archivo por número (la app detecta texto vs. OCR)",
                        variable=self._var_tipo, value="pdf",
                        command=self._on_tipo, style="TRadiobutton").grid(
                        row=1, column=0, columnspan=3, sticky="w", pady=2)
        ttk.Radiobutton(c2,
                        text="📁  Subcarpetas — cada subcarpeta es un número (ej. corpus BNC/Hemeroteca)",
                        variable=self._var_tipo, value="carpetas",
                        command=self._on_tipo, style="TRadiobutton").grid(
                        row=2, column=0, columnspan=3, sticky="w", pady=2)
        ttk.Radiobutton(c2, text="🖼  Imágenes sueltas (JPG/PNG/TIFF — siempre OCR)",
                        variable=self._var_tipo, value="img",
                        command=self._on_tipo, style="TRadiobutton").grid(
                        row=3, column=0, columnspan=3, sticky="w", pady=(2,6))

        # Separador entre tipo y carpetas
        tk.Frame(c2, bg=CARD_BOR, height=1).grid(row=4, column=0, columnspan=3,
                                                   sticky="ew", pady=(0, 8))
        tk.Label(c2, text="Carpeta de entrada:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=5,column=0,sticky="w",pady=4)
        self._var_ent = tk.StringVar()
        ent_row = tk.Frame(c2, bg=CARD_BG); ent_row.grid(row=5,column=1,columnspan=2,sticky="ew",padx=8)
        self._lf_ent = ent_row   # compatibilidad
        tk.Entry(ent_row, textvariable=self._var_ent, width=52,
                 font=("Segoe UI",9), relief="solid", bd=1,
                 bg="#1C2128").pack(side="left", fill="x", expand=True)
        ttk.Button(ent_row, text="Examinar…", style="S.TButton",
                   command=self._pick_ent).pack(side="left", padx=(6,0))

        # Salida
        tk.Label(c2, text="Carpeta de salida:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=6,column=0,sticky="w",pady=4)
        self._var_sal = tk.StringVar()
        sal_row = tk.Frame(c2, bg=CARD_BG); sal_row.grid(row=6,column=1,columnspan=2,sticky="ew",padx=8)
        tk.Entry(sal_row, textvariable=self._var_sal, width=52,
                 font=("Segoe UI",9), relief="solid", bd=1,
                 bg="#1C2128").pack(side="left", fill="x", expand=True)
        ttk.Button(sal_row, text="Examinar…", style="S.TButton",
                   command=self._pick_sal).pack(side="left", padx=(6,0))

        # Lista de archivos
        tk.Label(c2, text="Archivos a analizar:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=7,column=0,sticky="nw",pady=(8,0))
        arch_frame = tk.Frame(c2, bg=CARD_BG)
        arch_frame.grid(row=7,column=1,columnspan=2,sticky="ew",padx=8,pady=(6,0))
        btn_arch = tk.Frame(arch_frame, bg=CARD_BG); btn_arch.pack(fill="x")
        ttk.Button(btn_arch, text="☑ Todos", style="S.TButton",
                   command=lambda: self._sel_todos(True)).pack(side="left", padx=(0,4))
        ttk.Button(btn_arch, text="☐ Ninguno", style="S.TButton",
                   command=lambda: self._sel_todos(False)).pack(side="left")
        self._lbl_arch_info = tk.Label(btn_arch, text="(primero selecciona la carpeta)",
                                        bg=CARD_BG, fg="#6E7681", font=("Segoe UI",8))
        self._lbl_arch_info.pack(side="left", padx=10)
        lb_frame = tk.Frame(arch_frame, bg=CARD_BOR, bd=1, relief="solid")
        lb_frame.pack(fill="x", pady=(4,0))
        sby_lb = ttk.Scrollbar(lb_frame, orient="vertical")
        self._lb = tk.Listbox(lb_frame, selectmode="multiple", height=5,
                              font=("Courier",9), bg="#1C2128", relief="flat",
                              yscrollcommand=sby_lb.set, selectbackground=AZ3,
                              selectforeground="white", activestyle="none")
        sby_lb.config(command=self._lb.yview)
        sby_lb.pack(side="right", fill="y"); self._lb.pack(fill="x", expand=True)
        self._archivos_disp = []

        # ── Sección 3 ─────────────────────────────────────────────────────────
        body3 = _seccion("ocr", "3", "Calidad de lectura OCR")
        c3 = card_lf("🔬  Resolución e idioma", body3)
        tk.Label(c3, text="Resolución (DPI):", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=1,column=0,sticky="w",pady=5)
        dpi_row = tk.Frame(c3, bg=CARD_BG); dpi_row.grid(row=1,column=1,sticky="w",padx=8)
        self._var_dpi = tk.StringVar(value="150")
        for val, etiq in [("100","100 DPI · Rápido"),("150","150 DPI · Recomendado ✓"),
                           ("200","200 DPI · Alta calidad"),("300","300 DPI · Máxima")]:
            ttk.Radiobutton(dpi_row, text=etiq, variable=self._var_dpi,
                            value=val).pack(side="left", padx=8)
        self._mk_ayuda(dpi_row,
            "100 DPI: ~20 seg/número, puede perder letras pequeñas.\n"
            "150 DPI: equilibrio recomendado (~45 seg/número).\n"
            "200 DPI: mejor para textos con letra pequeña (~90 seg).\n"
            "300 DPI: máxima precisión, lento (~3 min/número).")

        tk.Label(c3, text="Idioma:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=2,column=0,sticky="w",pady=5)
        lang_row = tk.Frame(c3, bg=CARD_BG); lang_row.grid(row=2,column=1,sticky="w",padx=8)
        self._var_lang = tk.StringVar(value="spa")
        for val, etiq in [("spa","Español"),("spa+eng","Esp + Inglés"),("eng","Inglés")]:
            ttk.Radiobutton(lang_row, text=etiq, variable=self._var_lang,
                            value=val).pack(side="left", padx=8)

        # Lematización configurable (corpus histórico = mejor sin lematizar)
        tk.Label(c3, text="Análisis léxico:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=3,column=0,sticky="w",pady=5)
        lem_row = tk.Frame(c3, bg=CARD_BG); lem_row.grid(row=3,column=1,sticky="w",padx=8)
        self._var_lematizar = tk.BooleanVar(value=True)
        ttk.Checkbutton(lem_row, text="Lematizar (formas → lema base)",
                        variable=self._var_lematizar,
                        command=lambda: setattr(ST, "lematizar",
                                               self._var_lematizar.get())).pack(side="left")
        self._mk_ayuda(lem_row,
            "ACTIVADO (recomendado para español moderno):\n"
            "  'corrían', 'corrió', 'correr' → 'correr'\n"
            "  Mejora el recall en análisis de temas.\n\n"
            "DESACTIVADO (recomendado para corpus histórico):\n"
            "  Preserva formas originales: 'corrían', 'habia', 'fué'\n"
            "  Esencial para estudio de variación lingüística diacrónica.\n"
            "  Bashkar Station detectó que tu corpus es español histórico.")

        # ── Sección 4 ─────────────────────────────────────────────────────────
        body4 = _seccion("modulos", "4", "Módulos de análisis",
                         "Activa solo los módulos que necesitas para tu investigación.")
        c4 = card_lf("🧩  Módulos activos", body4)
        self._var_seg    = tk.BooleanVar(value=True)
        self._var_w2v    = tk.BooleanVar(value=False)
        self._var_vis    = tk.BooleanVar(value=True)
        self._var_red    = tk.BooleanVar(value=True)
        self._var_layout = tk.BooleanVar(value=True)

        modulos_def = [
            (self._var_seg,    "📝  Segmentación de artículos y atribución de autoría",
             "Divide cada número en artículos individuales y detecta quién los firmó."),
            (self._var_w2v,    "🧠  Word2Vec (expansión semántica)",
             "Entrena un modelo de vectores sobre el corpus para expandir automáticamente los campos temáticos. Requiere ≥50 páginas."),
            (self._var_vis,    "🖼  Análisis visual y tipográfico",
             "Detecta fotografías, ilustraciones y publicidades. Analiza fuentes tipográficas."),
            (self._var_red,    "🕸  Red de colaboración entre autores",
             "Construye un grafo de quiénes colaboraron en los mismos números."),
            (self._var_layout, "📐  Análisis de layout (distribución texto/imagen)",
             "Mide la proporción de espacio editorial dedicado a texto vs. imagen por número."),
        ]
        for i, (var, etiq, desc) in enumerate(modulos_def, 1):
            mod_row = tk.Frame(c4, bg=CARD_BG); mod_row.grid(
                row=i, column=0, columnspan=3, sticky="ew", pady=3)
            cb = ttk.Checkbutton(mod_row, text=etiq, variable=var)
            cb.pack(side="left")
            self._mk_ayuda(mod_row, desc)

        # ── Sección 5 ─────────────────────────────────────────────────────────
        body5 = _seccion("ajuste", "5", "Ajuste fino (opcional)", abierta=False)
        c5 = card_lf("🎛  Parámetros avanzados", body5)
        # spaCy
        tk.Label(c5, text="Precisión NLP:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=1,column=0,sticky="w",pady=5)
        spacy_row = tk.Frame(c5, bg=CARD_BG); spacy_row.grid(row=1,column=1,sticky="w",padx=8)
        self._var_spacy = tk.StringVar(value="es_core_news_sm")
        for val, etiq in [("es_core_news_sm","Pequeño ✓"),
                           ("es_core_news_md","Mediano"),
                           ("es_core_news_lg","Grande")]:
            ttk.Radiobutton(spacy_row, text=etiq, variable=self._var_spacy,
                            value=val).pack(side="left", padx=8)
        self._mk_ayuda(spacy_row,
            "Pequeño: rápido, suficiente para la mayoría de casos.\n"
            "Mediano: incluye vectores pre-entrenados (~43 MB).\n"
            "Grande: máxima precisión, requiere ~700 MB.")
        # LDA
        tk.Label(c5, text="Temas LDA:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=2,column=0,sticky="w",pady=5)
        lda_row = tk.Frame(c5, bg=CARD_BG); lda_row.grid(row=2,column=1,sticky="w",padx=8)
        self._var_lda = tk.IntVar(value=6)
        tk.Spinbox(lda_row, from_=3, to=20, textvariable=self._var_lda,
                   width=5, font=("Segoe UI",10), relief="solid", bd=1).pack(side="left")
        tk.Label(lda_row, text="  (3–20 temas latentes)", bg=CARD_BG,
                 fg="#6E7681", font=("Segoe UI",8)).pack(side="left")
        self._mk_ayuda(lda_row,
            "Corpus pequeño (2-5 números): 4-6 temas.\n"
            "Corpus mediano (6-20 números): 6-10 temas.\n"
            "Si los temas son demasiado similares, reduce el número.")
        # Red — umbral
        tk.Label(c5, text="Umbral red:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=3,column=0,sticky="w",pady=5)
        red_row = tk.Frame(c5, bg=CARD_BG); red_row.grid(row=3,column=1,sticky="w",padx=8)
        self._var_red_min = tk.IntVar(value=2)
        tk.Spinbox(red_row, from_=1, to=10, textvariable=self._var_red_min,
                   width=5, font=("Segoe UI",10), relief="solid", bd=1).pack(side="left")
        tk.Label(red_row, text="  apariciones mínimas en el corpus", bg=CARD_BG,
                 fg="#6E7681", font=("Segoe UI",8)).pack(side="left")
        self._mk_ayuda(red_row,
            "1: todos los autores.\n"
            "2: solo autores en ≥2 números (recomendado).\n"
            "3+: solo colaboradores frecuentes.")

        # ── Sección 6 ─────────────────────────────────────────────────────────
        body6 = _seccion("referencia", "6", "Corpus de referencia (comparativo)", abierta=False)
        c6 = card_lf("📚  Carpeta de referencia opcional", body6)
        tk.Label(c6, text="Carpeta:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=1,column=0,sticky="w",pady=5)
        self._var_ref = tk.StringVar()
        ref_row = tk.Frame(c6, bg=CARD_BG); ref_row.grid(row=1,column=1,columnspan=2,sticky="ew",padx=8)
        tk.Entry(ref_row, textvariable=self._var_ref, width=52,
                 font=("Segoe UI",9), relief="solid", bd=1, bg="#1C2128").pack(
                 side="left", fill="x", expand=True)
        ttk.Button(ref_row, text="Examinar…", style="S.TButton",
                   command=self._pick_ref).pack(side="left", padx=(6,0))
        tk.Label(c6,
                 text="Estructura: referencia/ → El_Tiempo/ → p001.txt, p002.txt…",
                 bg=CARD_BG, fg="#6E7681", font=("Segoe UI",8)).grid(
                 row=2, column=0, columnspan=3, sticky="w", pady=(0,4))

        # ── Sección 7 ─────────────────────────────────────────────────────────
        body7 = _seccion("vocab", "7", "Colaboradores y vocabulario", abierta=False)
        c7 = card_lf("👥  Colaboradores conocidos", body7)
        tk.Label(c7, text="Un nombre por línea, exactamente como aparece en la revista:",
                 bg=CARD_BG, fg="#CDD6F4", font=("Segoe UI",9)).grid(
                 row=1, column=0, columnspan=3, sticky="w", pady=(0,6))
        self._txt_col = scrolledtext.ScrolledText(c7, height=5, width=60, font=("Courier",9), bg="#0D1117", fg="#CDD6F4", insertbackground="#CDD6F4", relief="solid", bd=1)
        self._txt_col.grid(row=2, column=0, columnspan=3, sticky="ew")
        self._txt_col.insert("1.0", COLABS_DEFAULT)

        c8 = card_lf("🏷  Campos semánticos de interés", body7)
        import json as _json
        campos_txt = _json.dumps(
            CAMPOS_DEFAULT, ensure_ascii=False, indent=2
        ).strip("{}\n").strip()
        self._txt_sem = scrolledtext.ScrolledText(c8, height=8, width=70, font=("Courier",9), bg="#0D1117", fg="#CDD6F4", insertbackground="#CDD6F4", relief="solid", bd=1)
        self._txt_sem.grid(row=1, column=0, columnspan=3, sticky="ew")
        self._txt_sem.insert("1.0", campos_txt)
        # _mk_ayuda usa .pack() internamente — necesita frame propio (c8 usa grid)
        c8_ayuda_row = tk.Frame(c8, bg=CARD_BG)
        c8_ayuda_row.grid(row=2, column=0, columnspan=3, sticky="w", pady=(4, 0))
        self._mk_ayuda(c8_ayuda_row,
                       "Formato JSON. Cada campo está definido por palabras 'semilla'.\n"
                       "Si activas Word2Vec, el modelo añade automáticamente términos relacionados.")

        # ── Sección 8 ─────────────────────────────────────────────────────────
        body8 = _seccion("ia", "8", "Proveedores de IA y modelos",
                         "Todas las funciones IA son opcionales — la app funciona sin claves.",
                         abierta=False)

        # ── Switch global IA ──────────────────────────────────────────────────
        sw_card = card_lf("🔌  Funciones de inteligencia artificial (IA)", body8)

        # Reutilizar la variable ya creada en _build_topbar (no crear nueva)
        if not hasattr(self, "_var_ia_habilitada"):
            self._var_ia_habilitada = tk.BooleanVar(value=getattr(ST, "ia_habilitada", False))

        sw_top = tk.Frame(sw_card, bg=CARD_BG)
        sw_top.grid(row=1, column=0, columnspan=3, sticky="ew", pady=(0, 8))

        self._lbl_ia_estado = tk.Label(
            sw_top, bg=CARD_BG, font=("Segoe UI", 10, "bold"))
        self._lbl_ia_estado.pack(side="left", padx=(0, 16))

        ttk.Checkbutton(
            sw_top, text="Habilitar IA externa (requiere API key o Ollama local)",
            variable=self._var_ia_habilitada,
            command=self._cfg_toggle_ia
        ).pack(side="left")

        info_ia = (
            "Cuando está DESACTIVADO (modo offline):\n"
            "  • OCR: Tesseract y Kraken funcionan normalmente\n"
            "  • NER: solo BERT local (mrm8488/bert-spanish-cased-finetuned-ner)\n"
            "  • Búsqueda semántica: embeddings locales (MiniLM)\n"
            "  • Tono, narrativas, asistente: no disponibles\n\n"
            "Cuando está ACTIVADO:\n"
            "  • Se habilitan todas las funciones que requieren API externa\n"
            "  • Cada llamada consume crédito del proveedor configurado\n"
            "  • Los datos del corpus se envían a servidores externos\n"
            "  • Ollama (si está configurado) funciona localmente sin enviar datos"
        )
        tk.Label(sw_card, text=info_ia,
                 bg=CARD_BG, fg="#CDD6F4", font=("Segoe UI", 8),
                 justify="left", wraplength=580).grid(
                 row=2, column=0, columnspan=3, sticky="w", pady=(0, 4))

        self._cfg_toggle_ia()  # actualizar etiqueta al cargar

        # ── Card: Claves API ──────────────────────────────────────────────────
        c9 = card_lf("🔑  Claves API por proveedor", body8)

        # Variable legado (compatibilidad con el resto del código)
        self._var_api_key = tk.StringVar()

        # Variables individuales por proveedor
        self._var_key_anthropic = tk.StringVar()
        self._var_key_openai    = tk.StringVar()
        self._var_key_gemini    = tk.StringVar()
        self._var_key_ollama    = tk.StringVar(value="http://localhost:11434")

        _proveedores_info = [
            ("anthropic", "Anthropic Claude", self._var_key_anthropic,
             "sk-ant-…",
             "Prefijo: sk-ant-…\n\n"
             "Modelos disponibles:\n"
             "  claude-fable-5      — el más capaz (máx.)\n"
             "  claude-opus-4-8     — muy potente\n"
             "  claude-sonnet-4-6   — equilibrio costo/calidad ✓\n"
             "  claude-haiku-4-5    — el más económico\n\n"
             "Costo aproximado por página OCR: ~$0.003\n"
             "Obtén tu clave en: console.anthropic.com"),
            ("openai", "OpenAI", self._var_key_openai,
             "sk-…",
             "Prefijo: sk-…\n\n"
             "Modelos disponibles:\n"
             "  gpt-5.5             — visión + texto, máxima calidad\n"
             "  gpt-5.4-mini        — más económico ✓\n"
             "  gpt-4o              — alternativa previa\n\n"
             "Costo aproximado por página OCR: ~$0.005\n"
             "Obtén tu clave en: platform.openai.com"),
            ("gemini", "Google Gemini", self._var_key_gemini,
             "AIza…",
             "Prefijo: AIza…\n\n"
             "Modelos disponibles:\n"
             "  gemini-2.5-flash    — muy económico ✓✓\n"
             "  gemini-2.5-pro      — alta calidad\n"
             "  gemini-3.1-flash    — nueva generación, rápido ✓\n"
             "  gemini-3-pro        — el más capaz\n\n"
             "Costo aproximado por página OCR: ~$0.0001 (más barato)\n"
             "Obtén tu clave en: aistudio.google.com"),
            ("ollama", "Ollama (local)", self._var_key_ollama,
             "http://localhost:11434",
             "URL del servidor Ollama local.\n"
             "Por defecto: http://localhost:11434\n\n"
             "Modelos recomendados (instalar con 'ollama pull'):\n"
             "  llava              — visión + texto\n"
             "  mistral            — texto, rápido\n"
             "  llama3.2           — texto, buena calidad\n"
             "  latamgpt           — texto, español latinoamericano ✓ (recomendado para Estampa)\n\n"
             "Ventaja: 100% offline, sin costo por token.\n"
             "Requiere: ollama instalado y corriendo localmente."),
        ]

        for fila_idx, (prov_id, prov_label, prov_var, placeholder, ayuda_txt) in enumerate(_proveedores_info, 1):
            tk.Label(c9, text=f"{prov_label}:", bg=CARD_BG, fg="#CDD6F4",
                     font=("Segoe UI",9,"bold")).grid(row=fila_idx, column=0, sticky="w", pady=4)
            prow = tk.Frame(c9, bg=CARD_BG)
            prow.grid(row=fila_idx, column=1, columnspan=2, sticky="ew", padx=8)
            ent = tk.Entry(prow, textvariable=prov_var, width=42,
                           font=("Courier", 9), relief="solid", bd=1, bg="#1C2128",
                           show="*" if prov_id != "ollama" else "")
            ent.pack(side="left", fill="x", expand=True)
            if prov_id != "ollama":
                ttk.Button(prow, text="👁", style="S.TButton",
                           command=lambda e=ent: e.config(
                               show="" if e.cget("show") == "*" else "*")).pack(side="left", padx=2)
            self._mk_ayuda(prow, ayuda_txt)

        # Máximo de imágenes
        tk.Label(c9, text="Máx. imágenes/número:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=len(_proveedores_info)+1, column=0, sticky="w", pady=5)
        max_ia_row = tk.Frame(c9, bg=CARD_BG)
        max_ia_row.grid(row=len(_proveedores_info)+1, column=1, sticky="w", padx=8)
        self._var_max_ia = tk.IntVar(value=15)
        tk.Spinbox(max_ia_row, from_=1, to=50, textvariable=self._var_max_ia,
                   width=5, font=("Segoe UI",10), relief="solid", bd=1).pack(side="left")
        tk.Label(max_ia_row, text="  (cada imagen consume crédito de API)", bg=CARD_BG,
                 fg="#6E7681", font=("Segoe UI",8)).pack(side="left")

        # ── Card: Modelo por etapa (dentro de body8 también) ────────────────
        c9b = card_lf("⚙️  Modelo activo por etapa del análisis", body8)
        tk.Label(c9b,
                 text="Elige qué proveedor y modelo usar en cada función. "
                      "Pulsa ? para ver en qué se destaca cada opción.",
                 bg=CARD_BG, fg="#8B949E", font=("Segoe UI",8),
                 wraplength=560, justify="left").grid(
                 row=1, column=0, columnspan=3, sticky="w", pady=(0,8))

        # Opciones de modelo disponibles: (valor_interno, etiqueta_corta)
        _OPCIONES_MODELO = [
            ("anthropic/claude-fable-5",     "Claude Fable 5 (máx.)"),
            ("anthropic/claude-opus-4-8",    "Claude Opus 4.8"),
            ("anthropic/claude-sonnet-4-6",  "Claude Sonnet 4.6"),
            ("anthropic/claude-haiku-4-5",   "Claude Haiku 4.5"),
            ("openai/gpt-5.5",               "GPT-5.5"),
            ("openai/gpt-5.4-mini",          "GPT-5.4 mini"),
            ("openai/gpt-4o",                "GPT-4o"),
            ("gemini/gemini-3-pro",          "Gemini 3 Pro"),
            ("gemini/gemini-3.1-flash",      "Gemini 3.1 Flash"),
            ("gemini/gemini-2.5-pro",        "Gemini 2.5 Pro"),
            ("gemini/gemini-2.5-flash",      "Gemini 2.5 Flash"),
            ("ollama/llava",                 "Ollama llava"),
            ("ollama/mistral",               "Ollama mistral"),
        ]
        _vals_modelo = [v for v, _ in _OPCIONES_MODELO]
        _etiq_modelo = [e for _, e in _OPCIONES_MODELO]

        # Definición de etapas: (id, etiqueta UI, ayuda detallada)
        _ETAPAS = [
            ("ocr_mejora", "Mejora OCR (visión)",
             "Analiza imágenes de páginas con baja confianza Tesseract.\n"
             "Requiere modelo con capacidad de visión (imagen → texto).\n\n"
             "Comparación honesta:\n"
             "  GPT-4o        — buen desempeño en layouts multi-columna;\n"
             "                  costo medio (~$0.005/pág); envía datos a OpenAI\n"
             "  Claude Sonnet — desempeño comparable a GPT-4o en texto impreso;\n"
             "                  costo similar; envía datos a Anthropic\n"
             "  Gemini Flash  — el más económico (~$0.0001/pág); menor precisión\n"
             "                  en tipografías deterioradas; envía datos a Google\n"
             "  Ollama llava  — 100% local, sin costo, sin envío de datos;\n"
             "                  precisión notablemente inferior en documentos históricos\n\n"
             "Nota: ningún modelo fue evaluado sistemáticamente en prensa colombiana\n"
             "de los años 30. Los resultados pueden variar según el corpus."),
            ("deteccion", "Detección de zonas (visión)",
             "Identifica en la imagen qué áreas son foto, texto, publicidad, etc.\n"
             "Requiere modelo con capacidad de visión.\n\n"
             "Comparación honesta:\n"
             "  GPT-4o        — buen desempeño general en clasificación visual;\n"
             "                  envía imágenes a servidores de OpenAI\n"
             "  Claude Sonnet — desempeño similar a GPT-4o; envía a Anthropic\n"
             "  Gemini Flash  — más rápido y económico; menor detalle en zonas\n"
             "                  mixtas (texto+imagen en la misma área)\n"
             "  Ollama llava  — local y gratuito; dificultades con layouts complejos\n\n"
             "Alternativa sin API: el etiquetador manual de zonas no requiere IA."),
            ("ner", "Reconocimiento de entidades (NER)",
             "Extrae personas, lugares, organizaciones y eventos del texto OCR.\n\n"
             "Comparación honesta:\n"
             "  BERT local    — gratis, offline, sin envío de datos; entrenado en\n"
             "                  español moderno, menor recall en grafías de época\n"
             "  GPT-4o        — buen desempeño en NER español; costo por token;\n"
             "                  envía texto a OpenAI\n"
             "  Claude Sonnet — desempeño comparable a GPT-4o; envía a Anthropic\n"
             "  Gemini Pro    — resultados similares; envía a Google\n"
             "  Ollama mistral— local, sin costo; menor recall en entidades poco\n"
             "                  frecuentes o con grafía arcaica\n\n"
             "Recomendación práctica: BERT local para exploración inicial;\n"
             "API externa solo para validación o entidades críticas."),
            ("tono", "Análisis de tono editorial",
             "Clasifica artículos en categorías de tono histórico:\n"
             "celebratorio, crítico, elegíaco, polémico, informativo.\n\n"
             "Comparación honesta:\n"
             "  GPT-4o        — buen desempeño en clasificación; puede ser literal\n"
             "                  con ironía y ambigüedad retórica; envía a OpenAI\n"
             "  Claude Sonnet — desempeño similar; envía a Anthropic\n"
             "  Gemini Pro    — resultados comparables; más variable en matices\n"
             "                  culturales latinoamericanos; envía a Google\n"
             "  Ollama llama3 — local y gratuito; peor desempeño en registros\n"
             "                  formales del español de los años 30\n\n"
             "Importante: ninguno de estos modelos fue entrenado específicamente\n"
             "en prensa colombiana. Los resultados requieren validación humana."),
            ("narrativas", "Narrativas académicas automáticas",
             "Genera texto interpretativo sobre los hallazgos del análisis.\n\n"
             "Comparación honesta:\n"
             "  GPT-4o        — buena calidad de escritura académica en español;\n"
             "                  puede producir afirmaciones no respaldadas por los\n"
             "                  datos si el prompt no es preciso; envía a OpenAI\n"
             "  Claude Sonnet — calidad similar; igual riesgo de alucinación si los\n"
             "                  datos de entrada son ambiguos; envía a Anthropic\n"
             "  Gemini Pro    — aceptable; más variable en registro historiográfico\n"
             "  Ollama        — calidad inferior; mayor riesgo de texto genérico\n\n"
             "Advertencia: el texto generado debe revisarse antes de citarse\n"
             "en publicaciones académicas. No sustituye el análisis del investigador."),
            ("asistente", "Asistente IA por pestañas",
             "Responde preguntas sobre los resultados visibles en cada pestaña.\n\n"
             "Comparación honesta:\n"
             "  GPT-4o        — respuestas sólidas; ventana de contexto amplia;\n"
             "                  envía los datos mostrados en pantalla a OpenAI\n"
             "  GPT-4o mini   — más económico; calidad algo menor en razonamiento\n"
             "                  sobre datos complejos\n"
             "  Claude Sonnet — desempeño comparable a GPT-4o; envía a Anthropic\n"
             "  Claude Haiku  — respuestas más rápidas y económicas; menor\n"
             "                  profundidad en análisis historiográfico\n"
             "  Ollama mistral— local, sin costo, sin envío de datos; menor\n"
             "                  capacidad de razonamiento sobre datos estructurados\n\n"
             "Privacidad: con cualquier opción en la nube, los datos del corpus\n"
             "se envían al servidor del proveedor al hacer cada consulta."),
        ]

        self._vars_modelo_etapa = {}
        for fila_e, (etapa_id, etapa_label, etapa_ayuda) in enumerate(_ETAPAS, 2):
            tk.Label(c9b, text=f"{etapa_label}:", bg=CARD_BG, fg="#CDD6F4",
                     font=("Segoe UI",9,"bold")).grid(row=fila_e, column=0, sticky="w", pady=4)
            erow = tk.Frame(c9b, bg=CARD_BG)
            erow.grid(row=fila_e, column=1, sticky="w", padx=8)
            var_e = tk.StringVar(value=ST.modelos_etapa.get(etapa_id, "anthropic/claude-sonnet-4-6"))
            self._vars_modelo_etapa[etapa_id] = var_e
            cb_e = ttk.Combobox(erow, textvariable=var_e,
                                 values=_vals_modelo,
                                 state="readonly", width=30, font=("Segoe UI",9))
            # Mostrar etiquetas cortas en el combo (trucar con postcommand)
            cb_e._etiq = _etiq_modelo
            cb_e._vals = _vals_modelo
            cb_e.pack(side="left")
            self._mk_ayuda(erow, etapa_ayuda)

        c10 = card_lf("🔗  Extracción de metadatos desde URL", body8)
        tk.Label(c10, text="URL:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=1,column=0,sticky="w",pady=5)
        url_row = tk.Frame(c10, bg=CARD_BG); url_row.grid(row=1,column=1,columnspan=2,sticky="ew",padx=8)
        self._var_meta_url = tk.StringVar()
        tk.Entry(url_row, textvariable=self._var_meta_url, width=56,
                 font=("Segoe UI",9), relief="solid", bd=1, bg="#1C2128").pack(
                 side="left", fill="x", expand=True)
        ttk.Button(url_row, text="Extraer →", style="S.TButton",
                   command=self._extraer_metadatos_url).pack(side="left", padx=(6,0))
        self._txt_meta_result = scrolledtext.ScrolledText(c10, height=5, width=70, font=("Courier",9), bg="#0D1117", fg="#CDD6F4", relief="solid", bd=1, state="disabled")
        self._txt_meta_result.grid(row=2, column=0, columnspan=3, sticky="ew", pady=(8,0))

        # ── Botón confirmar ────────────────────────────────────────────────────
        btn_frame = tk.Frame(pad, bg=CONTENT_BG)
        btn_frame.grid(row=r, column=0, pady=(16, 24)); r += 1
        ttk.Button(btn_frame, text="✓  Confirmar configuración y continuar →",
                   style="P.TButton", command=self._confirmar_cfg).pack(side="left")
        self._lbl_cfg_ok = tk.Label(btn_frame, text="", fg=VERDE,
                                     bg=CONTENT_BG, font=("Segoe UI",9,"bold"))
        self._lbl_cfg_ok.pack(side="left", padx=16)

    def _on_tipo(self):
        # _lf_ent is now a plain Frame; we update the label separately if needed
        self._lb.delete(0, "end"); self._archivos_disp = []

    def _sel_todos(self, s):
        self._lb.select_set(0,"end") if s else self._lb.select_clear(0,"end")

    def _pick_ent(self):
        d = filedialog.askdirectory(title="Carpeta de entrada")
        if d:
            self._var_ent.set(d); self._poblar_lista(Path(d))

    def _pick_sal(self):
        d = filedialog.askdirectory(title="Carpeta de resultados")
        if d: self._var_sal.set(d)

    def _pick_ref(self):
        d = filedialog.askdirectory(title="Carpeta de publicaciones de referencia")
        if d: self._var_ref.set(d)

    def _extraer_metadatos_url(self):
        url = self._var_meta_url.get().strip()
        if not url:
            messagebox.showwarning("Sin URL", "Introduce una URL en el campo correspondiente.")
            return
        self._txt_meta_result.config(state="normal")
        self._txt_meta_result.delete("1.0", "end")
        self._txt_meta_result.insert("1.0", "⏳ Extrayendo metadatos… (puede tardar unos segundos)\n")
        self._txt_meta_result.config(state="disabled")
        self.update_idletasks()

        def _worker():
            try:
                from core.metadata_fetcher import extraer_metadatos_url
                datos = extraer_metadatos_url(url)
            except Exception as e:
                datos = {"exito": False, "error": str(e)}
            self.after(0, lambda d=datos: self._mostrar_metadatos(d))

        threading.Thread(target=_worker, daemon=True).start()

    def _mostrar_metadatos(self, datos: dict):
        self._txt_meta_result.config(state="normal")
        self._txt_meta_result.delete("1.0", "end")
        if not datos.get("exito"):
            self._txt_meta_result.insert("1.0",
                f"⚠️ No se pudieron extraer metadatos.\n"
                f"Error: {datos.get('error','desconocido')}\n\n"
                f"Sugerencia: verifica que la URL sea accesible desde tu navegador.")
        else:
            campos = [
                ("Título",        datos.get("titulo","")),
                ("Autor/es",      datos.get("autor","")),
                ("Fecha",         datos.get("fecha","")),
                ("Institución",   datos.get("institucion","")),
                ("Tipo",          datos.get("tipo_documento","")),
                ("ISSN",          datos.get("issn","")),
                ("Idioma",        datos.get("idioma","")),
                ("Editorial",     datos.get("editorial","")),
                ("Lugar",         datos.get("lugar","")),
                ("Derechos",      datos.get("derechos","")),
                ("Descripción",   datos.get("descripcion","")[:300]),
                ("Temas",         ", ".join(datos.get("temas",[])[:8])),
                ("Fuente",        datos.get("fuente","")),
            ]
            lineas = []
            for etiq, val in campos:
                if val: lineas.append(f"  {etiq:<16} {val}")
            self._txt_meta_result.insert("1.0",
                "✅ Metadatos extraídos:\n" + "\n".join(lineas))
            # Pre-rellenar nombre de publicación si está vacío
            if datos.get("titulo") and self._var_pub.get() == "Mi publicación":
                self._var_pub.set(datos["titulo"][:60])
        self._txt_meta_result.config(state="disabled")

    def _poblar_lista(self, carpeta):
        from core.ocr_engine import EXTS_IMAGEN, analizar_pdf
        self._lb.delete(0,"end"); self._archivos_disp = []
        tipo = self._var_tipo.get()

        if tipo == "carpetas":
            # Modo subcarpetas: cada subdirectorio es un número de la publicación.
            # Se registra la subcarpeta como "archivo" — el worker OCR la procesará
            # iterando sus PDFs internamente.
            subcarpetas = sorted(p for p in carpeta.iterdir() if p.is_dir())
            if not subcarpetas:
                self._lbl_arch_info.config(
                    text="⚠️ No se encontraron subcarpetas en esa carpeta")
                return
            self._lbl_arch_info.config(
                text=f"Analizando {len(subcarpetas)} subcarpeta(s)…")
            self.update_idletasks()
            for sc in subcarpetas:
                pdfs = sorted(sc.glob("*.pdf"))
                n_pdfs = len(pdfs)
                mb_total = sum(p.stat().st_size for p in pdfs) / 1024 / 1024
                # Analizar primer PDF para detectar si tiene texto
                if pdfs:
                    info = analizar_pdf(pdfs[0])
                    estado = ("✅ texto" if info["tiene_texto"] else "🔍 OCR")
                    estado += f" · {n_pdfs} PDF(s) · {mb_total:.1f} MB"
                else:
                    estado = f"Sin PDFs ({n_pdfs})"
                self._archivos_disp.append(sc)
                self._lb.insert("end", f"{sc.name:<42}  {estado}")
                self._lb.select_set(self._lb.size()-1)
            self._lbl_arch_info.config(
                text=f"{len(subcarpetas)} número(s) · todos seleccionados")
            ST.input_tipo = "carpetas"
            return

        if tipo == "pdf":
            archivos = sorted(carpeta.glob("*.pdf"))
        else:
            archivos = sorted(p for p in carpeta.iterdir()
                              if p.suffix.lower() in EXTS_IMAGEN)

        if not archivos:
            self._lbl_arch_info.config(text="⚠️ No se encontraron archivos")
            return
        self._lbl_arch_info.config(
            text=f"Analizando {len(archivos)} archivo(s)…")
        self.update_idletasks()
        for p in archivos:
            self._archivos_disp.append(p)
            mb = p.stat().st_size / 1024 / 1024
            if tipo == "pdf":
                info  = analizar_pdf(p)
                estado = (f"✅ texto ({info['palabras_promedio']:.0f} pal/pág)"
                          if info["tiene_texto"]
                          else f"🔍 OCR ({info['palabras_promedio']:.0f} pal/pág)")
            else:
                estado = ""
            self._lb.insert("end", f"{p.name:<42} {mb:5.1f} MB  {estado}")
            self._lb.select_set(self._lb.size()-1)
        self._lbl_arch_info.config(
            text=f"{len(archivos)} archivo(s) · todos seleccionados")

    def _cfg_toggle_ia(self):
        habilitada = self._var_ia_habilitada.get()
        ST.ia_habilitada = habilitada
        if hasattr(self, "_lbl_ia_estado"):
            if habilitada:
                self._lbl_ia_estado.config(
                    text="IA EXTERNA: ACTIVADA", fg=VERDE, bg=CARD_BG)
            else:
                self._lbl_ia_estado.config(
                    text="IA EXTERNA: DESACTIVADA  (modo offline)",
                    fg=ROJO, bg=CARD_BG)
        # Sincronizar topbar
        if hasattr(self, "_lbl_ia_topbar"):
            self._lbl_ia_topbar.config(
                text="● IA ON" if habilitada else "○ IA OFF",
                fg=VERDE if habilitada else ROJO)

    def _confirmar_cfg(self):
        ent = Path(self._var_ent.get().strip())
        sal = Path(self._var_sal.get().strip())
        if not ent.exists():
            messagebox.showerror("Error", f"Carpeta no existe: {ent}"); return
        sel = [self._archivos_disp[i] for i in self._lb.curselection()]
        if not sel:
            messagebox.showwarning("Sin selección","Selecciona al menos un archivo."); return
        if not str(sal).strip():
            messagebox.showerror("Error","Define la carpeta de resultados."); return
        sal.mkdir(parents=True, exist_ok=True)
        ST.publicacion = self._var_pub.get().strip() or "Publicación"
        ST.periodo     = self._var_per.get().strip()
        ST.pdf_dir     = ent; ST.out_dir = sal
        ST.archivos_sel = sel; ST.input_tipo = self._var_tipo.get()
        # Guardar switch IA
        ST.ia_habilitada = self._var_ia_habilitada.get()
        # Guardar claves por proveedor
        ST.api_keys["anthropic"] = self._var_key_anthropic.get().strip()
        ST.api_keys["openai"]    = self._var_key_openai.get().strip()
        ST.api_keys["gemini"]    = self._var_key_gemini.get().strip()
        ST.api_keys["ollama"]    = self._var_key_ollama.get().strip()
        # api_key legado = primera clave no vacía (prioridad: anthropic → openai → gemini → ollama)
        ST.api_key = next(
            (v for v in (ST.api_keys["anthropic"], ST.api_keys["openai"],
                         ST.api_keys["gemini"]) if v),
            ""
        )
        self._var_api_key.set(ST.api_key)
        # Guardar modelos por etapa
        for etapa_id, var_e in self._vars_modelo_etapa.items():
            ST.modelos_etapa[etapa_id] = var_e.get()
        ST.max_ia       = self._var_max_ia.get()
        # Parsear campos semánticos
        import json
        try:
            raw = "{" + self._txt_sem.get("1.0","end").strip() + "}"
            ST.campos_semillas = json.loads(raw)
        except Exception:
            ST.campos_semillas = CAMPOS_DEFAULT
        # Guardar configuración lingüística
        if hasattr(self, "_var_lematizar"):
            ST.lematizar = self._var_lematizar.get()
        if hasattr(self, "_txt_stopwords"):
            texto_sw = self._txt_stopwords.get("1.0", "end-1c")
            ST.stopwords_proyecto = [p.strip().lower() for p in texto_sw.splitlines() if p.strip()]
        self._lbl_pub_hdr.config(text=f"{ST.publicacion}  ·  {ST.periodo}")
        self._lbl_cfg_ok.config(text=f"  ✅  {len(sel)} archivo(s) listos")
        self.toast(f"{len(sel)} archivo(s) configurados — continúa con Extracción", tipo="info")
        self._mostrar_pagina("ocr")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 2: EXTRACCIÓN / OCR
    # ══════════════════════════════════════════════════════════════════════════
    def _build_ocr(self):
        f = self._tab_ocr
        self._page_header(f, "Extracción de texto",
                          "Detecta texto digital o aplica OCR automáticamente página a página", "📄")

        self._build_ai_panel(f, "ocr")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True, padx=24, pady=16)

        # ── Tarjetas de métricas ──────────────────────────────────────────────
        ind = tk.Frame(pad, bg=CONTENT_BG); ind.pack(fill="x", pady=(0, 16))
        self._lbl_o_pdf = self._mk_ind(ind, "Archivos",      "—", 0)
        self._lbl_o_pag = self._mk_ind(ind, "Páginas",       "—", 1)
        self._lbl_o_pal = self._mk_ind(ind, "Palabras",      "—", 2)
        self._lbl_o_con = self._mk_ind(ind, "Confianza OCR", "—", 3)
        self._lbl_o_rev = self._mk_ind(ind, "Para revisión", "—", 4)

        # ── Progreso ──────────────────────────────────────────────────────────
        prog_card = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1,
                             highlightbackground=CARD_BOR, highlightthickness=1)
        prog_card.pack(fill="x", pady=(0, 12))
        prog_inner = tk.Frame(prog_card, bg=CARD_BG, padx=16, pady=12)
        prog_inner.pack(fill="x")
        self._lbl_fase = tk.Label(prog_inner, text="Esperando…",
                                   bg=CARD_BG, fg="#8B949E",
                                   font=("Segoe UI", 9, "italic"))
        self._lbl_fase.pack(anchor="w")
        self._prog = ttk.Progressbar(prog_inner, mode="determinate", length=600)
        self._prog.pack(fill="x", pady=(6, 4))
        self._lbl_pct = tk.Label(prog_inner, text="", bg=CARD_BG, fg="#8B949E",
                                  font=("Courier", 8))
        self._lbl_pct.pack(anchor="w")

        # ── Selector de ruta de extracción ───────────────────────────────────
        ruta_card = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1,
                             highlightbackground=CARD_BOR, highlightthickness=1)
        ruta_card.pack(fill="x", pady=(0, 10))
        ruta_inner = tk.Frame(ruta_card, bg=CARD_BG, padx=16, pady=10)
        ruta_inner.pack(fill="x")

        ruta_hdr = tk.Frame(ruta_inner, bg=CARD_BG)
        ruta_hdr.pack(fill="x", pady=(0, 6))
        tk.Label(ruta_hdr, text="Ruta de extracción", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI", 9, "bold")).pack(side="left")
        self._mk_ayuda(ruta_hdr,
            "Elige cómo se obtiene el texto de cada página del PDF.\n\n"
            "Ruta 1 — Tesseract propio (recomendada para corpus BNC):\n"
            "  Convierte el PDF a imágenes y aplica OCR con Tesseract.\n"
            "  Ignora el texto embebido por la BNC, que mezcla columnas.\n"
            "  Resultado: texto limpio, columnas en orden correcto.\n\n"
            "Ruta 2 — Claude Vision (máxima calidad, requiere API key):\n"
            "  Envía cada página como imagen a Claude. Comprende el\n"
            "  layout visual y transcribe respetando columnas, títulos\n"
            "  y pies de foto. Más lento y tiene costo en tokens.\n\n"
            "Ruta 3 — Texto BNC + reconstrucción de líneas (más rápida):\n"
            "  Usa el texto ya extraído por la BNC pero aplica un\n"
            "  algoritmo que detecta y une líneas rotas de columna.\n\n"
            "Ruta 4 — Kraken + CATMuS-Print (★ MEJOR para prensa histórica):\n"
            "  Motor OCR entrenado específicamente en prensa latinoamericana\n"
            "  del siglo XX. 100% offline. Requiere instalar kraken y\n"
            "  descargar el modelo CATMuS-Print (~200 MB, botón abajo).\n"
            "  Confianza típica: 85-92% en textos de los años 30-40.\n\n"
            "Ruta 5 — Ollama Vision (offline con modelo de visión local):\n"
            "  Usa Qwen2.5-VL u otro modelo multimodal instalado en Ollama.\n"
            "  100% offline, sin costo de API. Requiere Ollama instalado\n"
            "  y al menos 8 GB de RAM. Más lento que Kraken.")

        self._var_ruta_ocr = tk.StringVar(value="tesseract")
        rutas = [
            ("tesseract", "Ruta 1 — Tesseract propio  ✓ Recomendada"),
            ("vision_ia", "Ruta 2 — IA de visión  (Claude · GPT-4o · Gemini · Ollama)"),
            ("bnc",       "Ruta 3 — Texto BNC + reconstrucción de líneas"),
            ("kraken",    "Ruta 4 — Kraken + CATMuS-Print  ★ Mejor para prensa histórica (offline)"),
            ("ollama",    "Ruta 5 — Ollama Vision  (offline, requiere Ollama + modelo visión)"),
        ]
        for val, etiq in rutas:
            r = tk.Frame(ruta_inner, bg=CARD_BG)
            r.pack(fill="x", pady=1)
            ttk.Radiobutton(r, text=etiq, variable=self._var_ruta_ocr,
                            value=val).pack(side="left")

        # ── Sub-panel Ruta 2: selector proveedor + modelo ─────────────────────
        from core.zone_labeler import VISION_PROVEEDORES
        self._ocr_vision_frame = tk.Frame(ruta_inner, bg=CARD_BG)
        self._ocr_vision_frame.pack(fill="x", pady=(2, 0), padx=(24, 0))

        self._ocr_vision_prov  = tk.StringVar(value="claude")
        self._ocr_vision_model = tk.StringVar(value="claude-sonnet-4-6")

        tk.Label(self._ocr_vision_frame, text="Proveedor:",
                 bg=CARD_BG, fg=TXT_SEC, font=("Segoe UI", 8)).pack(side="left")
        _cb_ocr_prov = ttk.Combobox(
            self._ocr_vision_frame, textvariable=self._ocr_vision_prov,
            values=list(VISION_PROVEEDORES.keys()), state="readonly", width=10)
        _cb_ocr_prov.pack(side="left", padx=(4, 8))

        tk.Label(self._ocr_vision_frame, text="Modelo:",
                 bg=CARD_BG, fg=TXT_SEC, font=("Segoe UI", 8)).pack(side="left")
        self._cb_ocr_vision_model = ttk.Combobox(
            self._ocr_vision_frame, textvariable=self._ocr_vision_model,
            state="readonly", width=26)
        self._cb_ocr_vision_model.pack(side="left", padx=(4, 0))

        # Tooltip ?
        _VISION_OCR_AYUDA = (
            "Ruta 2 — IA de visión para OCR:\n\n"
            "claude  — Claude Sonnet/Haiku. Alta calidad, entiende layout histórico.\n"
            "openai  — GPT-4o / GPT-4o-mini. Muy buena calidad, gpt-4o-mini es económico.\n"
            "gemini  — Gemini 1.5 Flash. Casi gratuito en el tier libre de Google.\n"
            "ollama  — Llava local, 100% offline. Requiere Ollama + modelo visión.\n\n"
            "Todos usan el mismo prompt interno de transcripción OCR.\n"
            "Costo estimado con Claude Haiku: ~$0.002/página."
        )
        _btn_q = tk.Label(self._ocr_vision_frame, text=" ?", bg=CARD_BG,
                          fg=TXT_SEC, font=("Segoe UI", 8, "bold"), cursor="hand2")
        _btn_q.pack(side="left", padx=(6, 0))
        _btn_q.bind("<Enter>",    lambda e: self._mostrar_tooltip(_VISION_OCR_AYUDA, _btn_q))
        _btn_q.bind("<Leave>",    lambda e: self._ocultar_tooltip())
        _btn_q.bind("<Button-1>", lambda e: messagebox.showinfo("Ruta 2 — IA de visión",
                                                                  _VISION_OCR_AYUDA))

        def _on_ocr_prov(*_):
            prov = self._ocr_vision_prov.get()
            info = VISION_PROVEEDORES.get(prov, {})
            mods = info.get("modelos", [])
            self._cb_ocr_vision_model["values"] = mods
            self._ocr_vision_model.set(info.get("default", mods[0] if mods else ""))

        def _on_ruta_change(*_):
            if self._var_ruta_ocr.get() == "vision_ia":
                self._ocr_vision_frame.pack(fill="x", pady=(2, 0), padx=(24, 0))
            else:
                self._ocr_vision_frame.pack_forget()

        _cb_ocr_prov.bind("<<ComboboxSelected>>", _on_ocr_prov)
        self._var_ruta_ocr.trace_add("write", _on_ruta_change)
        _on_ocr_prov()
        _on_ruta_change()   # estado inicial

        # ── Opción: usar zonas etiquetadas ────────────────────────────────────
        etz_card = tk.Frame(ruta_inner, bg="#0D1117", relief="solid", bd=1)
        etz_card.pack(fill="x", pady=(8, 0))
        etz_inner = tk.Frame(etz_card, bg="#0D1117", padx=10, pady=6)
        etz_inner.pack(fill="x")

        self._var_ocr_usar_etiquetas = tk.BooleanVar(value=False)
        ttk.Checkbutton(
            etz_inner,
            text="Usar zonas etiquetadas del Etiquetador",
            variable=self._var_ocr_usar_etiquetas,
        ).pack(side="left")

        tk.Label(etz_inner,
                 text="— solo procesa las zonas marcadas como texto; mantiene el orden de lectura",
                 bg="#0D1117", fg="#6E7681", font=("Segoe UI", 8)).pack(side="left", padx=(6, 0))

        self._var_ocr_det_auto = tk.BooleanVar(value=False)
        det_row = tk.Frame(etz_card, bg="#0D1117", padx=10, pady=6)
        det_row.pack(fill="x")
        ttk.Checkbutton(
            det_row,
            text="Detección automática IA en páginas sin etiquetar",
            variable=self._var_ocr_det_auto,
        ).pack(side="left")
        tk.Label(det_row,
                 text="(requiere IA activa)",
                 bg="#0D1117", fg="#6E7681", font=("Segoe UI", 8)).pack(side="left", padx=4)

        # ── Sub-opciones Kraken ────────────────────────────────────────────────
        kraken_card = tk.Frame(ruta_inner, bg=CONTENT_BG, relief="solid", bd=1)
        kraken_card.pack(fill="x", pady=(4, 0))
        ki = tk.Frame(kraken_card, bg=CONTENT_BG, padx=10, pady=6); ki.pack(fill="x")
        tk.Label(ki, text="Modelo Kraken:", bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 8, "bold")).pack(side="left")
        self._var_kraken_modelo = tk.StringVar(value="")
        tk.Entry(ki, textvariable=self._var_kraken_modelo, width=42,
                 font=("Segoe UI", 8), relief="solid", bd=1).pack(side="left", padx=6)
        ttk.Button(ki, text="📂", width=3,
                   command=self._ocr_elegir_modelo_kraken).pack(side="left", padx=(0, 6))
        self._btn_catmus = ttk.Button(ki, text="⬇ Descargar CATMuS-Print",
                                       style="S.TButton",
                                       command=self._ocr_descargar_catmus)
        self._btn_catmus.pack(side="left")
        self._lbl_kraken_ok = tk.Label(ki, text="", bg=CONTENT_BG, fg=VERDE,
                                        font=("Segoe UI", 8))
        self._lbl_kraken_ok.pack(side="left", padx=6)
        # Verificar Kraken al construir el panel
        self.after(200, self._ocr_verificar_kraken)

        # ── Paralelismo Kraken ────────────────────────────────────────────────
        kpar = tk.Frame(kraken_card, bg=CONTENT_BG, padx=10, pady=4)
        kpar.pack(fill="x")

        # Fila 1: selector de workers + explicación
        kpar_row1 = tk.Frame(kpar, bg=CONTENT_BG); kpar_row1.pack(fill="x")
        tk.Label(kpar_row1, text="Páginas en paralelo:", bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 8, "bold")).pack(side="left")
        self._var_kraken_workers = tk.IntVar(value=3)
        spin = tk.Spinbox(kpar_row1, from_=1, to=12,
                          textvariable=self._var_kraken_workers,
                          width=3, font=("Segoe UI", 9),
                          command=self._ocr_actualizar_estimacion)
        spin.pack(side="left", padx=(6, 10))
        spin.bind("<FocusOut>", lambda e: self._ocr_actualizar_estimacion())
        spin.bind("<Return>",   lambda e: self._ocr_actualizar_estimacion())

        tk.Label(kpar_row1, text="Timeout por página (seg):", bg=CONTENT_BG,
                 fg="#CDD6F4", font=("Segoe UI", 8)).pack(side="left", padx=(10, 0))
        self._var_kraken_timeout = tk.IntVar(value=600)
        spin_to = tk.Spinbox(kpar_row1, from_=60, to=3600,
                             textvariable=self._var_kraken_timeout,
                             increment=60, width=5, font=("Segoe UI", 9))
        spin_to.pack(side="left", padx=(6, 6))
        self._mk_ayuda(kpar_row1,
            "Tiempo máximo que Bashkar espera a Kraken por página.\n\n"
            "Si Kraken tarda más que este límite, la página se marca\n"
            "como fallida y se intenta con Tesseract como respaldo.\n\n"
            "Valores sugeridos:\n"
            "  Páginas simples (texto solo):       120–180 seg\n"
            "  Páginas mixtas (texto + fotos):     300–600 seg\n"
            "  Páginas muy complejas (publicidad): 600–900 seg\n\n"
            "Si ves muchos errores de timeout, súbelo.\n"
            "Si el corpus es simple, bájalo para detectar\n"
            "páginas problemáticas más rápido.")

        self._mk_ayuda(kpar_row1,
            "Cuántas páginas procesa Kraken al mismo tiempo.\n\n"
            "Cada proceso paralelo carga el modelo en RAM (~400 MB).\n"
            "Recomendaciones según tu equipo:\n\n"
            "  4 GB RAM  →  1 proceso (sin paralelismo)\n"
            "  8 GB RAM  →  2-3 procesos  ← tu equipo (7.3 GB)\n"
            "  16 GB RAM →  4-6 procesos\n"
            "  32 GB RAM →  8-10 procesos\n\n"
            "Más procesos = más rápido, pero si te quedas sin RAM\n"
            "el sistema empieza a usar disco (swap) y se vuelve\n"
            "más lento que con menos procesos.\n\n"
            "Empieza con el valor sugerido y auméntalo si el\n"
            "procesador no llega al 80% de uso durante el OCR.")

        # Fila 2: estimación de tiempo
        kpar_row2 = tk.Frame(kpar, bg=CONTENT_BG); kpar_row2.pack(fill="x", pady=(4, 0))
        tk.Label(kpar_row2, text="Tiempo estimado:", bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 8, "bold")).pack(side="left")
        self._lbl_kraken_est = tk.Label(kpar_row2,
                                         text="— (carga un proyecto para estimar)",
                                         bg=CONTENT_BG, fg="#8B949E",
                                         font=("Segoe UI", 8))
        self._lbl_kraken_est.pack(side="left", padx=(6, 10))

        self._mk_ayuda(kpar_row2,
            "Estimación basada en el corpus del proyecto activo.\n\n"
            "Velocidad de referencia: ~60 seg/página en Ryzen 5 5500U\n"
            "con 1 proceso. Con N procesos: tiempo ÷ N (aprox.).\n\n"
            "La primera vez que corres Kraken en tu equipo puede\n"
            "ser más lento por la carga inicial del modelo.\n"
            "Las corridas siguientes suelen ser más rápidas.\n\n"
            "La estimación mejora si previamente usas 'Calibrar'\n"
            "sobre una página real de tu corpus.")

        ttk.Button(kpar_row2, text="↺ Calibrar",
                   style="S.TButton",
                   command=self._ocr_calibrar_kraken).pack(side="left")

        # Calcular estimación inicial si hay proyecto cargado
        self.after(500, self._ocr_actualizar_estimacion)

        # ── Sub-opciones Ollama ────────────────────────────────────────────────
        ollama_card = tk.Frame(ruta_inner, bg=CONTENT_BG, relief="solid", bd=1)
        ollama_card.pack(fill="x", pady=(4, 0))
        oi = tk.Frame(ollama_card, bg=CONTENT_BG, padx=10, pady=6); oi.pack(fill="x")
        tk.Label(oi, text="Modelo Ollama:", bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 8, "bold")).pack(side="left")
        self._var_ollama_modelo = tk.StringVar(value="qwen2.5vl:7b")
        self._cmb_ollama = ttk.Combobox(oi, textvariable=self._var_ollama_modelo,
                                         values=["qwen2.5vl:7b", "llava:13b", "llava:7b"],
                                         state="normal", width=20, font=("Segoe UI", 8))
        self._cmb_ollama.pack(side="left", padx=6)
        ttk.Button(oi, text="🔄 Detectar modelos", style="S.TButton",
                   command=self._ocr_detectar_ollama).pack(side="left", padx=(0, 6))
        self._lbl_ollama_ok = tk.Label(oi, text="", bg=CONTENT_BG, fg="#8B949E",
                                        font=("Segoe UI", 8))
        self._lbl_ollama_ok.pack(side="left", padx=6)
        self.after(300, self._ocr_detectar_ollama)

        # ── Preprocesamiento de imagen ────────────────────────────────────────
        pre_card = tk.Frame(ruta_inner, bg=CARD_BG, relief="solid", bd=1)
        pre_card.pack(fill="x", pady=(6, 0))
        pi_f = tk.Frame(pre_card, bg=CARD_BG, padx=10, pady=5)
        pi_f.pack(fill="x")
        tk.Label(pi_f, text="Preprocesamiento de imagen:",
                 bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 8, "bold")).pack(side="left")
        # Desactivados por defecto — activar solo si la imagen lo necesita
        self._var_pre_deskew    = tk.BooleanVar(value=False)
        self._var_pre_enhance   = tk.BooleanVar(value=False)
        self._var_pre_despeckle = tk.BooleanVar(value=False)
        for txt, var in [("Deskew (inclinación)", self._var_pre_deskew),
                          ("CLAHE (contraste)",    self._var_pre_enhance),
                          ("Despeckle (ruido)",    self._var_pre_despeckle)]:
            ttk.Checkbutton(pi_f, text=txt, variable=var).pack(side="left", padx=(8, 0))

        # Botón preview — muestra la imagen antes y después del preprocesamiento
        tk.Label(pi_f, text="  ", bg=CARD_BG).pack(side="left")
        _btn_prev = tk.Label(pi_f, text="🔍 Preview", bg=CARD_BG, fg=TXT_SEC,
                              font=("Segoe UI", 8), cursor="hand2")
        _btn_prev.pack(side="left", padx=(4, 0))
        _btn_prev.bind("<Button-1>", lambda e: self._ocr_preview_preprocesamiento())
        _btn_prev.bind("<Enter>", lambda e: _btn_prev.config(fg=TXT_PRI))
        _btn_prev.bind("<Leave>", lambda e: _btn_prev.config(fg=TXT_SEC))

        self._mk_ayuda(pi_f,
            "Preprocesamiento aplicado antes del OCR (desactivado por defecto):\n\n"
            "Deskew: Endereza páginas torcidas 1-4°. Útil para escanes de la BNC.\n"
            "Activar solo si las páginas tienen inclinación visible.\n\n"
            "CLAHE: Mejora contraste en papel amarillado o con iluminación no uniforme.\n\n"
            "Despeckle: Elimina puntos de papel envejecido.\n"
            "Puede borrar puntuación fina — activar con cuidado.\n\n"
            "Las originales a color siempre se preservan en 02_imagenes/.\n"
            "Las imágenes procesadas se guardan en 02_imagenes_ocr/.\n"
            "Usa 🔍 Preview para ver el efecto antes de correr el OCR completo.")

        # ── Opciones avanzadas de salida ──────────────────────────────────────
        def _build_ocr_avanzado(f):
            row1 = tk.Frame(f, bg=CONTENT_BG); row1.pack(fill="x", pady=2)
            tk.Label(row1, text="Umbral confianza para revisión manual (%):",
                     bg=CONTENT_BG, fg=TXT_SEC, font=("Segoe UI", 8)).pack(side="left")
            self._var_ocr_umbral_rev = tk.IntVar(value=70)
            ttk.Spinbox(row1, from_=10, to=99, textvariable=self._var_ocr_umbral_rev,
                        width=5).pack(side="left", padx=6)
            tk.Label(row1, text="(páginas por debajo irán a revisión)",
                     bg=CONTENT_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left")

            row2 = tk.Frame(f, bg=CONTENT_BG); row2.pack(fill="x", pady=2)
            self._var_ocr_guardar_json = tk.BooleanVar(value=False)
            ttk.Checkbutton(row2, text="Guardar metadata de confianza por página (.json)",
                            variable=self._var_ocr_guardar_json).pack(side="left")

            row3 = tk.Frame(f, bg=CONTENT_BG); row3.pack(fill="x", pady=2)
            self._var_ocr_combinar_paginas = tk.BooleanVar(value=True)
            ttk.Checkbutton(row3, text="Combinar páginas en un solo TXT por número",
                            variable=self._var_ocr_combinar_paginas).pack(side="left")

        self._mk_avanzado(pad, "Opciones avanzadas de extracción", _build_ocr_avanzado)

        # ── Botones ───────────────────────────────────────────────────────────
        bf = tk.Frame(pad, bg=CONTENT_BG); bf.pack(fill="x", pady=(0, 8))
        self._btn_ocr = ttk.Button(bf, text="▶  Iniciar extracción",
                                    style="P.TButton", command=self._start_ocr)
        self._btn_ocr.pack(side="left", padx=(0, 12))
        ttk.Button(bf, text="✍ Re-normalizar textos", style="S.TButton",
                   command=self._renormalizar_textos).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📐 Completar costura", style="S.TButton",
                   command=self._gutter_completar_corpus).pack(side="left", padx=(0, 8))
        self._mk_ayuda(bf,
            "Completar costura: detecta palabras cortadas por el pliegue de\n"
            "encuadernación y las reconstruye con IA. Las palabras generadas\n"
            "aparecen marcadas con ⟦⟧ y en rojo en las exportaciones DOCX.\n"
            "Requiere IA habilitada y API key configurada.")
        ttk.Button(bf, text="Mejorar con IA", style="S.TButton",
                   command=self._start_mejorar_ocr).pack(side="left", padx=(0, 4))
        # Selector de proveedor LLM para corrección post-OCR
        self._var_ocr_llm_prov = tk.StringVar(value="claude")
        cmb_prov_ocr = ttk.Combobox(bf, textvariable=self._var_ocr_llm_prov,
                     values=["claude", "openai", "gemini", "ollama", "lmstudio"],
                     state="readonly", width=8,
                     font=("Segoe UI", 9))
        cmb_prov_ocr.pack(side="left", padx=(0, 2))
        # Modelo local (visible solo cuando proveedor=ollama/lmstudio)
        self._var_ocr_ollama_modelo = tk.StringVar(value="latamgpt")
        self._cmb_ocr_ollama_modelo = ttk.Combobox(
            bf, textvariable=self._var_ocr_ollama_modelo,
            values=["latamgpt", "llama3.1", "mistral", "qwen2.5"],
            state="normal", width=10, font=("Segoe UI", 9))
        # mostrar/ocultar según proveedor; para lmstudio consulta qué hay cargado
        def _ocr_prov_changed(*_):
            prov = self._var_ocr_llm_prov.get()
            if prov == "lmstudio":
                from core.ocr_llm import modelos_cargados_lmstudio
                modelos = modelos_cargados_lmstudio()
                self._cmb_ocr_ollama_modelo.config(values=modelos)
                if modelos and self._var_ocr_ollama_modelo.get() not in modelos:
                    self._var_ocr_ollama_modelo.set(modelos[0])
                self._cmb_ocr_ollama_modelo.pack(side="left", padx=(0, 8))
            elif prov == "ollama":
                self._cmb_ocr_ollama_modelo.config(
                    values=["latamgpt", "llama3.1", "mistral", "qwen2.5"])
                self._cmb_ocr_ollama_modelo.pack(side="left", padx=(0, 8))
            else:
                self._cmb_ocr_ollama_modelo.pack_forget()
        self._var_ocr_llm_prov.trace_add("write", _ocr_prov_changed)
        self._mk_ayuda(bf,
            "Mejorar con IA: aplica Vision o corrección post-OCR\n"
            "a páginas con confianza Tesseract por debajo del umbral.\n"
            "Proveedores: claude / openai / gemini / ollama / lmstudio\n"
            "  → ollama: escribe el modelo (ej: latamgpt, llama3.1)\n"
            "  → lmstudio: servidor local (Developer → Start Server),\n"
            "    lista los modelos cargados automáticamente\n"
            "Umbral 60: solo páginas malas. Umbral 40: más agresivo.")
        tk.Label(bf, text="Umbral IA:", bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 9)).pack(side="left", padx=(8, 2))
        self._var_ocr_umbral = tk.IntVar(value=60)
        tk.Spinbox(bf, from_=10, to=90, textvariable=self._var_ocr_umbral,
                   width=4, font=("Segoe UI", 9), relief="solid", bd=1).pack(side="left", padx=(0,12))
        tk.Label(bf, text="⚠  Confirma la configuración antes de empezar",
                 bg=CONTENT_BG, fg=ACENT, font=("Segoe UI", 9)).pack(side="left")

        # ── Log ───────────────────────────────────────────────────────────────
        log_frame = tk.Frame(pad, bg="#0F1B2D", bd=1, relief="solid")
        log_frame.pack(fill="both", expand=True)
        log_hdr = tk.Frame(log_frame, bg="#1A2F4A")
        log_hdr.pack(fill="x")
        tk.Label(log_hdr, text="  📋  Registro de actividad",
                 bg="#1A2F4A", fg="#94A3B8",
                 font=("Segoe UI", 8, "bold")).pack(side="left", pady=4)
        self._log_w = scrolledtext.ScrolledText(log_frame, height=12,
                                                 font=("Consolas", 9),
                                                 bg="#0F1B2D", fg="#7DD3FC",
                                                 relief="flat", insertbackground="white",
                                                 selectbackground="#1D4ED8")
        self._log_w.pack(fill="both", expand=True, padx=1, pady=(0, 1))
        self._log_w.config(state="disabled")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB CONV: CONVERSOR MASIVO PDF → WORD / TXT
    # ══════════════════════════════════════════════════════════════════════════
    def _build_conv(self):
        """Panel de conversión masiva PDF→Word/TXT (ruta rápida, sin re-OCR)."""
        try:
            from core.conversor_pdf_a_word import (
                ConfiguracionConversor,
                ConversorPDFaWord,
            )
            self._conv_disponible = True
        except ImportError as _conv_err:
            self._conv_disponible = False
            self._conv_import_error = str(_conv_err)

        f = self._tab_conv
        self._page_header(
            f,
            "Conversor masivo PDF → Word / TXT",
            "Para PDFs que ya tienen texto: extrae y organiza en carpetas sin volver a hacer OCR.",
            "⚡",
        )

        pad = tk.Frame(f, bg=CONTENT_BG)
        pad.pack(fill="both", expand=True, padx=24, pady=16)

        # ── Dependencias faltantes ────────────────────────────────────────────
        if not self._conv_disponible:
            err_card = tk.Frame(pad, bg="#2D1B00", relief="solid", bd=1,
                                highlightbackground="#F97316", highlightthickness=1)
            err_card.pack(fill="x", pady=(0, 16))
            err_inner = tk.Frame(err_card, bg="#2D1B00", padx=16, pady=14)
            err_inner.pack(fill="x")
            tk.Label(err_inner, text="⚠  Faltan librerías requeridas",
                     bg="#2D1B00", fg="#F97316",
                     font=("Segoe UI", 10, "bold")).pack(anchor="w")
            tk.Label(err_inner,
                     text=f"Error: {getattr(self, '_conv_import_error', '')}\n\n"
                          "Solución: abrí una terminal y ejecutá\n"
                          "    pip install pymupdf python-docx",
                     bg="#2D1B00", fg="#FED7AA",
                     font=("Segoe UI", 9), justify="left").pack(anchor="w", pady=(6, 0))
            return

        # ── Variables de estado ───────────────────────────────────────────────
        self._conv_entrada  = tk.StringVar(value=getattr(ST, "pdf_dir", "") or "")
        self._conv_salida   = tk.StringVar(value=getattr(ST, "out_dir", "") or "")
        self._conv_word_con = tk.BooleanVar(value=True)
        self._conv_word_pag = tk.BooleanVar(value=False)
        self._conv_txt_con  = tk.BooleanVar(value=True)
        self._conv_txt_pag  = tk.BooleanVar(value=False)
        self._conv_frag     = tk.BooleanVar(value=False)
        self._conv_limpiar   = tk.BooleanVar(value=True)
        self._conv_al_norm   = tk.BooleanVar(value=True)
        self._conv_desde     = tk.StringVar(value="")
        self._conv_hasta     = tk.StringVar(value="")
        self._conv_modo      = tk.StringVar(value="texto")
        self._conv_proc      = None

        # ══ PASO 1 — Carpetas ════════════════════════════════════════════════
        c1 = self._card(pad, "  1  Elegí las carpetas")

        def _fila_dir(parent, etiqueta, descripcion, var, comando):
            fila = tk.Frame(parent, bg=CARD_BG)
            fila.pack(fill="x", pady=(0, 8))
            tk.Label(fila, text=etiqueta, bg=CARD_BG, fg=TXT_PRI,
                     font=("Segoe UI", 9, "bold"), width=8,
                     anchor="w").pack(side="left")
            tk.Label(fila, text=descripcion, bg=CARD_BG, fg=TXT_DIM,
                     font=("Segoe UI", 8)).pack(side="left", padx=(0, 8))
            tk.Button(fila, text="📁 Seleccionar", bg=AZ3, fg="#FFFFFF",
                      relief="flat", font=("Segoe UI", 8), padx=8, pady=3,
                      cursor="hand2", command=comando).pack(side="right")
            tk.Entry(fila, textvariable=var, bg=CARD_BG, fg=TXT_SEC,
                     relief="solid", bd=1, font=("Segoe UI", 8),
                     state="readonly").pack(side="right", fill="x",
                                           expand=True, padx=(0, 8))

        def _sel_entrada():
            d = filedialog.askdirectory(title="Carpeta con los PDF originales")
            if d:
                self._conv_entrada.set(d)
                ST.pdf_dir = d

        def _sel_salida():
            d = filedialog.askdirectory(title="Carpeta donde guardar los archivos convertidos")
            if d:
                self._conv_salida.set(d)
                ST.out_dir = d

        _fila_dir(c1, "Entrada", "carpeta con los PDF a convertir",
                  self._conv_entrada, _sel_entrada)
        _fila_dir(c1, "Salida",  "dónde guardar los Word y TXT",
                  self._conv_salida, _sel_salida)

        # ── Nota explicativa ──────────────────────────────────────────────────
        tk.Label(c1,
                 text="ℹ  Cada PDF genera su propia subcarpeta con los archivos organizados.",
                 bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 8),
                 justify="left").pack(anchor="w", pady=(0, 4))

        # ══ PASO 2 — Qué generar ═════════════════════════════════════════════
        c2 = self._card(pad, "  2  Qué archivos querés generar")

        for var, titulo, desc in [
            (self._conv_word_con, "📄  Word (.docx)",
             "Un documento Word por PDF, con todo el texto. Ideal para leer y citar."),
            (self._conv_txt_con,  "📃  Texto plano (.txt)",
             "Un archivo TXT por PDF. Más liviano, ideal para análisis computacional."),
            (self._conv_frag,     "📎  PDF por página",
             "Divide cada PDF en páginas individuales (opcional, ocupa más espacio)."),
        ]:
            row = tk.Frame(c2, bg=CARD_BG)
            row.pack(fill="x", pady=3)
            ttk.Checkbutton(row, text=titulo, variable=var).pack(side="left")
            tk.Label(row, text=desc, bg=CARD_BG, fg=TXT_DIM,
                     font=("Segoe UI", 8)).pack(side="left", padx=(8, 0))

        # ── Separador ─────────────────────────────────────────────────────────
        tk.Frame(c2, bg=CARD_BOR, height=1).pack(fill="x", pady=(8, 8))

        # ── Limpieza de texto ─────────────────────────────────────────────────
        limpiar_row = tk.Frame(c2, bg=CARD_BG)
        limpiar_row.pack(fill="x", pady=2)
        ttk.Checkbutton(limpiar_row,
                        text="🧹  Limpiar el texto automáticamente",
                        variable=self._conv_limpiar).pack(side="left")
        self._mk_ayuda(limpiar_row,
            "Aplica limpieza al texto extraído de cada página:\n\n"
            "  ✓ Elimina los números de coordenadas que aparecen\n"
            "    entre párrafos (artefacto de los PDF de la BNC)\n"
            "  ✓ Elimina el sello 'Digitalizado Biblioteca Nacional'\n"
            "  ✓ Une palabras cortadas al final de columna (pági-\n"
            "    nas → páginas)\n"
            "  ✓ Corrige errores OCR frecuentes (6→á, 11→ll, 1→l)\n"
            "  ✓ Normaliza tildes y caracteres especiales\n\n"
            "Preserva el español de época: habia, fué, Luégo, etc.\n"
            "son grafías históricas legítimas, no errores.\n\n"
            "Desactivá esta opción solo si necesitás el texto\n"
            "completamente sin procesar.")
        tk.Label(limpiar_row,
                 text="  ← recomendado para corpus BNC",
                 bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left")

        norm_row = tk.Frame(c2, bg=CARD_BG)
        norm_row.pack(fill="x", pady=(4, 0))
        ttk.Checkbutton(norm_row,
                        text="📝  Enviar al módulo Normalizar",
                        variable=self._conv_al_norm).pack(side="left")
        self._mk_ayuda(norm_row,
            "Copia el texto de cada página en la estructura que\n"
            "usa el módulo Normalizar (03_ocr/<nombre>/p0001.txt).\n\n"
            "Con esta opción activada, al terminar la conversión\n"
            "podés ir directamente a Normalizar para revisar y\n"
            "corregir el texto página por página antes del NER.\n\n"
            "Usa la carpeta de salida del proyecto (configurada\n"
            "en Configuración). Si no hay proyecto abierto, esta\n"
            "opción se ignora.")
        tk.Label(norm_row,
                 text="  ← conecta con el pipeline de análisis",
                 bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left")

        # ── Rango de páginas (avanzado, colapsado) ────────────────────────────
        rango_row = tk.Frame(c2, bg=CARD_BG)
        rango_row.pack(fill="x", pady=(6, 0))
        tk.Label(rango_row, text="Páginas:", bg=CARD_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left")
        tk.Label(rango_row, text="desde", bg=CARD_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left", padx=(8, 2))
        tk.Entry(rango_row, textvariable=self._conv_desde, width=5,
                 bg=CARD_BG, fg=TXT_PRI, relief="solid", bd=1,
                 font=("Segoe UI", 8)).pack(side="left")
        tk.Label(rango_row, text="hasta", bg=CARD_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left", padx=(6, 2))
        tk.Entry(rango_row, textvariable=self._conv_hasta, width=5,
                 bg=CARD_BG, fg=TXT_PRI, relief="solid", bd=1,
                 font=("Segoe UI", 8)).pack(side="left")
        tk.Label(rango_row, text="  (vacío = todas)",
                 bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(side="left")

        # ══ BOTONES DE ACCIÓN ════════════════════════════════════════════════
        # Deben ir ANTES del widget con expand=True (regla de layout Tkinter)
        btn_row = tk.Frame(pad, bg=CONTENT_BG)
        btn_row.pack(fill="x", pady=(4, 10))

        self._btn_conv_iniciar = tk.Button(
            btn_row, text="⚡  Convertir ahora",
            bg=AZ3, fg="#FFFFFF", relief="flat",
            font=("Segoe UI", 11, "bold"), padx=24, pady=8,
            cursor="hand2", command=self._conv_iniciar)
        self._btn_conv_iniciar.pack(side="left", padx=(0, 10))

        self._btn_conv_cancelar = tk.Button(
            btn_row, text="✖  Detener",
            bg=CARD_BG, fg=TXT_SEC, relief="flat",
            font=("Segoe UI", 10), padx=16, pady=8,
            cursor="hand2", command=self._conv_cancelar,
            state="disabled")
        self._btn_conv_cancelar.pack(side="left", padx=(0, 10))

        self._btn_conv_abrir = tk.Button(
            btn_row, text="📂  Ver archivos generados",
            bg=CARD_BG, fg=TXT_SEC, relief="flat",
            font=("Segoe UI", 10), padx=16, pady=8,
            cursor="hand2", command=self._conv_abrir_salida,
            state="disabled")
        self._btn_conv_abrir.pack(side="left")

        # ══ PROGRESO ═════════════════════════════════════════════════════════
        prog_card = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1,
                             highlightbackground=CARD_BOR, highlightthickness=1)
        prog_card.pack(fill="x", pady=(0, 8))
        prog_inner = tk.Frame(prog_card, bg=CARD_BG, padx=16, pady=12)
        prog_inner.pack(fill="x")

        self._conv_lbl_fase = tk.Label(
            prog_inner, text="Listo para convertir.",
            bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 9, "italic"))
        self._conv_lbl_fase.pack(anchor="w")

        self._conv_prog = ttk.Progressbar(
            prog_inner, mode="determinate", length=600, maximum=100)
        self._conv_prog.pack(fill="x", pady=(6, 0))

        # ══ REGISTRO DE ACTIVIDAD ════════════════════════════════════════════
        log_hdr = tk.Frame(pad, bg=CONTENT_BG)
        log_hdr.pack(fill="x", pady=(4, 2))
        tk.Label(log_hdr, text="Registro de actividad",
                 bg=CONTENT_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8, "bold")).pack(side="left")

        self._conv_log = tk.Text(
            pad, bg=CARD_BG, fg=TXT_SEC,
            font=("Segoe UI", 9), height=8,
            relief="solid", bd=1, state="disabled",
            wrap="word")
        self._conv_log.pack(fill="both", expand=True)

    def _conv_log_append(self, texto: str, color: str = "#94A3B8"):
        try:
            self._conv_log.config(state="normal")
            tag = f"c{abs(hash(color))}"
            self._conv_log.tag_configure(tag, foreground=color)
            self._conv_log.insert("end", texto + "\n", tag)
            self._conv_log.see("end")
            self._conv_log.config(state="disabled")
        except Exception:
            pass

    def _conv_log_clear(self):
        try:
            self._conv_log.config(state="normal")
            self._conv_log.delete("1.0", "end")
            self._conv_log.config(state="disabled")
        except Exception:
            pass

    def _conv_iniciar(self):
        try:
            from core.conversor_pdf_a_word import (
                ConfiguracionConversor,
                ConversorPDFaWord,
            )
        except ImportError as e:
            messagebox.showerror("Conversor", f"Faltan librerías:\n{e}\n\npip install pymupdf python-docx")
            return

        entrada = self._conv_entrada.get().strip()
        salida  = self._conv_salida.get().strip()
        if not entrada:
            messagebox.showwarning("Conversor", "Selecciona la carpeta de entrada.")
            return
        if not salida:
            messagebox.showwarning("Conversor", "Selecciona la carpeta de salida.")
            return
        if not os.path.isdir(entrada):
            messagebox.showwarning("Conversor",
                f"La carpeta de entrada no existe:\n{entrada}")
            return
        pdfs_disponibles = list(Path(entrada).glob("*.pdf"))
        if not pdfs_disponibles:
            messagebox.showwarning("Conversor",
                f"No se encontraron archivos PDF en:\n{entrada}\n\n"
                "Verificá que la carpeta seleccionada contenga los PDF directamente\n"
                "(no en subcarpetas).")
            return

        def _int_o_none(var):
            v = var.get().strip()
            try:
                return int(v) if v else None
            except ValueError:
                return None

        al_norm  = self._conv_al_norm.get()
        out_dir  = getattr(ST, "out_dir", None)
        cfg = ConfiguracionConversor(
            carpeta_entrada          = entrada,
            carpeta_salida           = salida,
            modo_texto               = self._conv_modo.get(),
            fragmentar_pdf           = self._conv_frag.get(),
            word_consolidado         = self._conv_word_con.get(),
            word_por_pagina          = self._conv_word_pag.get(),
            txt_consolidado          = self._conv_txt_con.get(),
            txt_por_pagina           = self._conv_txt_pag.get(),
            limpiar_texto            = self._conv_limpiar.get(),
            exportar_para_normalizar = al_norm and out_dir is not None,
            carpeta_out_dir          = out_dir,
            paginas_desde            = _int_o_none(self._conv_desde),
            paginas_hasta            = _int_o_none(self._conv_hasta),
        )

        self._btn_conv_iniciar.config(state="disabled")
        self._btn_conv_cancelar.config(state="normal")
        self._btn_conv_abrir.config(state="disabled")
        self._conv_prog["value"] = 0
        self._conv_lbl_fase.config(text="Iniciando conversión…")
        self._conv_log_clear()
        self._conv_log_append(f"Carpeta de entrada:  {entrada}", "#484F58")
        self._conv_log_append(f"Carpeta de salida:   {salida}", "#484F58")
        self._conv_log_append("─" * 48, "#2D333B")

        def _on_progress(ev):
            pct  = ev.get("porcentaje", 0)
            npdf = ev.get("indice_pdf", 0)
            tot  = ev.get("total_pdfs", 1)
            pag  = ev.get("pagina", 0)
            tpag = ev.get("total_paginas", 0)
            nombre = ev.get("pdf", "")
            msg  = f"PDF {npdf}/{tot}  ·  {nombre}  ·  página {pag}/{tpag}"
            self._conv_prog.after(0, lambda p=pct, m=msg: (
                self._conv_prog.config(value=p),
                self._conv_lbl_fase.config(text=m),
            ))

        def _worker():
            try:
                proc = ConversorPDFaWord(cfg, callback_progreso=_on_progress)
                self._conv_proc = proc
                reporte = proc.procesar_todo()
                self._conv_prog.after(0, lambda: self._conv_finalizar(reporte))
            except Exception as exc:
                import traceback
                tb = traceback.format_exc()
                self._conv_prog.after(0, lambda e=str(exc), t=tb: self._conv_error(e, t))

        import threading
        threading.Thread(target=_worker, daemon=True).start()

    def _conv_finalizar(self, reporte: dict):
        self._btn_conv_iniciar.config(state="normal")
        self._btn_conv_cancelar.config(state="disabled")
        self._btn_conv_abrir.config(state="normal")

        cancelado = reporte.get("cancelado", False)
        pdfs      = reporte.get("pdfs", [])
        n_ok      = sum(1 for d in pdfs if "error" not in d)
        n_err     = sum(1 for d in pdfs if "error" in d)
        n_pags    = sum(d.get("paginas_procesadas", 0) for d in pdfs if "error" not in d)
        seg       = reporte.get("segundos_totales", 0)

        if cancelado:
            self._conv_prog["value"] = 0
            self._conv_lbl_fase.config(text="Conversión detenida por el usuario.")
            self._conv_log_append("\n⚠  Proceso detenido antes de terminar.", "#F97316")
        else:
            self._conv_prog["value"] = 100
            mins = int(seg) // 60
            secs = int(seg) % 60
            tiempo = f"{mins} min {secs} s" if mins else f"{secs} s"
            self._conv_lbl_fase.config(
                text=f"✓  Conversión completada en {tiempo}")
            self._conv_log_append("─" * 48, "#2D333B")
            self._conv_log_append(
                f"✓  {n_ok} PDF convertidos  ·  {n_pags} páginas procesadas  ·  {tiempo}",
                "#3FB950")

        if n_err:
            self._conv_log_append(
                f"⚠  {n_err} PDF no pudieron procesarse:", "#F97316")
            for d in pdfs:
                if "error" in d:
                    self._conv_log_append(
                        f"   • {d['archivo']}: {d['error']}", "#F97316")
        for d in pdfs:
            if "error" not in d and d.get("paginas_sin_texto", 0) > 0:
                self._conv_log_append(
                    f"   ℹ  {d['archivo']}: {d['paginas_sin_texto']} páginas sin texto detectable",
                    "#8B949E")

        # Refrescar Normalizar y reconstruir corpus_meta si se exportó a 03_ocr/
        if self._conv_al_norm.get() and getattr(ST, "out_dir", None):
            try:
                self._reconstruir_corpus_meta_desde_txt()
                self._norm_refrescar_numeros()
                self._conv_log_append(
                    "📝  Texto disponible en Normalizar y Segmentar.", "#3FB950")
            except Exception:
                pass

    def _conv_error(self, exc_str: str, tb: str):
        self._btn_conv_iniciar.config(state="normal")
        self._btn_conv_cancelar.config(state="disabled")
        self._conv_lbl_fase.config(text="Error durante la conversión.")
        self._conv_log_append(f"\n✗ Error: {exc_str}", "#F97316")
        self._conv_log_append(tb, "#484F58")

    def _conv_cancelar(self):
        if self._conv_proc:
            self._conv_proc.cancelar()
        self._btn_conv_cancelar.config(state="disabled")
        self._conv_lbl_fase.config(text="Cancelando…")

    def _conv_abrir_salida(self):
        salida = self._conv_salida.get().strip()
        if salida and os.path.isdir(salida):
            import subprocess
            subprocess.Popen(["explorer", os.path.normpath(salida)])

    # ══════════════════════════════════════════════════════════════════════════
    # TAB MMX: EXTRACCIÓN MULTIMODAL ESTRUCTURADA (IA de visión → JSON → .md)
    # ══════════════════════════════════════════════════════════════════════════

    # Modelos de visión vigentes por proveedor (jun-2026). El primero es el
    # default. Verificados: Claude vía skill claude-api; Gemini/OpenAI/Ollama vía
    # docs oficiales. El combobox es editable por si el usuario quiere otro.
    _MMX_MODELOS = {
        "gemini": ["gemini-2.5-flash", "gemini-2.5-pro", "gemini-3.1-flash",
                   "gemini-3-pro"],
        "claude": ["claude-opus-4-8", "claude-sonnet-4-6", "claude-haiku-4-5",
                   "claude-opus-4-7", "claude-fable-5"],
        "openai": ["gpt-5.5", "gpt-5.4-mini", "gpt-4o", "gpt-4o-mini"],
        "ollama": ["llava", "llama3.2-vision", "qwen2.5-vl", "minicpm-v"],
        "lmstudio": [],  # sin catálogo fijo: se consulta al servidor local
    }

    def _build_mmx(self):
        """Panel de extracción multimodal: imágenes de página → JSON estricto
        (artículo jerárquico + imágenes con pies + publicidad) → .md, con opción
        de alimentar el corpus textual. Respuesta al hallazgo de auditoría: las
        imágenes BNC requieren IA de visión; aquí su salida queda estructurada."""
        try:
            from core import extractor_multimodal  # noqa: F401
            self._mmx_disponible = True
        except ImportError as _mmx_err:
            self._mmx_disponible = False
            self._mmx_import_error = str(_mmx_err)

        f = self._tab_mmx
        self._page_header(
            f,
            "Extracción multimodal con IA",
            "Convierte imágenes de página en datos estructurados (texto jerárquico, "
            "fotos con pies, publicidad) usando IA de visión.",
            "🧠",
        )

        pad = tk.Frame(f, bg=CONTENT_BG)
        pad.pack(fill="both", expand=True, padx=24, pady=16)

        if not self._mmx_disponible:
            err = tk.Frame(pad, bg="#2D1B00", relief="solid", bd=1)
            err.pack(fill="x", pady=(0, 16))
            tk.Label(err, text="⚠  No se pudo cargar core/extractor_multimodal",
                     bg="#2D1B00", fg="#F97316",
                     font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=14, pady=10)
            tk.Label(err, text=getattr(self, "_mmx_import_error", ""),
                     bg="#2D1B00", fg="#FED7AA", font=("Segoe UI", 9)).pack(
                         anchor="w", padx=14, pady=(0, 10))
            return

        # ── Variables ─────────────────────────────────────────────────────────
        self._mmx_entrada = tk.StringVar(value="")
        self._mmx_salida  = tk.StringVar(value="")
        self._mmx_prov    = tk.StringVar(value="gemini")
        self._mmx_modelo  = tk.StringVar(value=self._MMX_MODELOS["gemini"][0])
        self._mmx_guardar_json = tk.BooleanVar(value=True)
        self._mmx_alimentar    = tk.BooleanVar(value=True)
        self._mmx_corriendo    = False

        # ══ PASO 1 — Carpetas ═══════════════════════════════════════════════
        c1 = self._card(pad, "  1  Carpetas")

        def _fila_dir(parent, etiqueta, descripcion, var, comando):
            fila = tk.Frame(parent, bg=CARD_BG)
            fila.pack(fill="x", pady=(0, 8))
            tk.Label(fila, text=etiqueta, bg=CARD_BG, fg=TXT_PRI,
                     font=("Segoe UI", 9, "bold"), width=8, anchor="w").pack(side="left")
            tk.Label(fila, text=descripcion, bg=CARD_BG, fg=TXT_DIM,
                     font=("Segoe UI", 8)).pack(side="left", padx=(0, 8))
            tk.Button(fila, text="📁 Seleccionar", bg=AZ3, fg="#FFFFFF",
                      relief="flat", font=("Segoe UI", 8), padx=8, pady=3,
                      cursor="hand2", command=comando).pack(side="right")
            tk.Entry(fila, textvariable=var, bg=CARD_BG, fg=TXT_SEC, relief="solid",
                     bd=1, font=("Segoe UI", 8), state="readonly").pack(
                         side="right", fill="x", expand=True, padx=(0, 8))

        def _sel_entrada():
            d = filedialog.askdirectory(title="Carpeta con las imágenes de página")
            if d:
                self._mmx_entrada.set(d)
                if not self._mmx_salida.get():
                    self._mmx_salida.set(os.path.join(d, "extraccion_ia"))

        def _sel_salida():
            d = filedialog.askdirectory(title="Carpeta donde guardar JSON y .md")
            if d:
                self._mmx_salida.set(d)

        _fila_dir(c1, "Entrada", "carpeta con .jpg/.png/.tif de cada página",
                  self._mmx_entrada, _sel_entrada)
        _fila_dir(c1, "Salida", "dónde guardar los .json y .md",
                  self._mmx_salida, _sel_salida)

        # ══ PASO 2 — Proveedor de IA ════════════════════════════════════════
        c2 = self._card(pad, "  2  IA de visión")
        rowp = tk.Frame(c2, bg=CARD_BG)
        rowp.pack(fill="x", pady=2)
        tk.Label(rowp, text="Proveedor:", bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9)).pack(side="left")
        cb_prov = ttk.Combobox(rowp, textvariable=self._mmx_prov, width=12,
                               state="readonly",
                               values=["gemini", "claude", "openai", "ollama", "lmstudio"])
        cb_prov.pack(side="left", padx=(6, 16))
        tk.Label(rowp, text="Modelo:", bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9)).pack(side="left")
        # Combobox editable: muestra los modelos de visión vigentes del proveedor,
        # pero permite escribir uno propio.
        self._mmx_cb_modelo = ttk.Combobox(
            rowp, textvariable=self._mmx_modelo, width=22, state="normal",
            values=self._MMX_MODELOS.get(self._mmx_prov.get(), []))
        self._mmx_cb_modelo.pack(side="left", padx=(6, 0))

        def _prov_cambio(*_a):
            prov = self._mmx_prov.get()
            if prov == "lmstudio":
                from core.ocr_llm import modelos_cargados_lmstudio
                modelos = modelos_cargados_lmstudio()
            else:
                modelos = self._MMX_MODELOS.get(prov, [])
            self._mmx_cb_modelo.config(values=modelos)
            # Al cambiar de proveedor, selecciona su modelo recomendado (1.º).
            self._mmx_modelo.set(modelos[0] if modelos else "")
        self._mmx_prov.trace_add("write", _prov_cambio)

        self._mk_ayuda(rowp,
            "El selector lista los modelos de visión vigentes del proveedor\n"
            "elegido (el 1.º es el recomendado); podés escribir otro si querés.\n"
            "Gemini 2.5 Flash es ideal para lotes (rápido y económico).\n"
            "La clave se toma de ⚙ Configuración → Claves API por proveedor.\n"
            "Ollama corre localmente (modelo de visión, p.ej. llava) sin costo.\n"
            "LM Studio corre localmente (Developer → Start Server en LM Studio);\n"
            "  el selector consulta qué modelo tenés cargado ahora mismo.")

        # ══ PASO 3 — Qué hacer con el resultado ═════════════════════════════
        c3 = self._card(pad, "  3  Salida")
        ttk.Checkbutton(c3, text="💾  Guardar JSON crudo por página (auditable, trazable)",
                        variable=self._mmx_guardar_json).pack(anchor="w", pady=2)
        ttk.Checkbutton(c3, text="📥  Alimentar el corpus textual (para NER, análisis, grafo)",
                        variable=self._mmx_alimentar).pack(anchor="w", pady=2)

        # ══ PASO 4 — Ejecutar ═══════════════════════════════════════════════
        c4 = self._card(pad, "  4  Procesar")
        rowb = tk.Frame(c4, bg=CARD_BG)
        rowb.pack(fill="x", pady=(0, 8))
        tk.Button(rowb, text="🧮 Estimar costo", bg="#1F2937", fg="#E5E7EB",
                  relief="flat", font=("Segoe UI", 9), padx=12, pady=5,
                  cursor="hand2", command=self._mmx_estimar).pack(side="left")
        self._mmx_btn = tk.Button(rowb, text="🧠 Extraer todo", bg=AZ3, fg="#FFFFFF",
                                  relief="flat", font=("Segoe UI", 9, "bold"),
                                  padx=14, pady=5, cursor="hand2",
                                  command=self._mmx_iniciar)
        self._mmx_btn.pack(side="left", padx=(8, 0))
        tk.Button(rowb, text="📂 Abrir salida", bg="#1F2937", fg="#E5E7EB",
                  relief="flat", font=("Segoe UI", 9), padx=12, pady=5,
                  cursor="hand2", command=self._mmx_abrir_salida).pack(side="right")

        self._mmx_prog = ttk.Progressbar(c4, mode="determinate")
        self._mmx_prog.pack(fill="x", pady=(0, 6))
        self._mmx_log = tk.Text(c4, height=11, bg="#0D1117", fg="#9CA3AF",
                                relief="solid", bd=1, font=("Consolas", 8),
                                wrap="word")
        self._mmx_log.pack(fill="both", expand=True)

    def _mmx_log_add(self, texto: str):
        self._mmx_log.insert("end", texto + "\n")
        self._mmx_log.see("end")

    def _mmx_clave(self):
        """Devuelve (api_key, proveedor, modelo) validando que haya clave."""
        prov = self._mmx_prov.get().strip().lower()
        modelo = self._mmx_modelo.get().strip() or None
        if prov == "ollama":
            return ST.api_keys.get("ollama", "http://localhost:11434"), prov, modelo
        if prov == "lmstudio":
            return ST.api_keys.get("lmstudio", "http://localhost:1234"), prov, modelo
        clave_map = {"gemini": "gemini", "claude": "anthropic", "openai": "openai"}
        api_key = ST.api_keys.get(clave_map.get(prov, prov), "") or ST.api_key
        return api_key, prov, modelo

    def _mmx_estimar(self):
        carpeta = self._mmx_entrada.get().strip()
        if not carpeta or not os.path.isdir(carpeta):
            messagebox.showwarning("Sin carpeta", "Elegí la carpeta de imágenes.")
            return
        try:
            from core import extractor_multimodal as em
            _, prov, modelo = self._mmx_clave()
            est = em.estimar_costo_directorio(carpeta, proveedor=prov, modelo=modelo)
            messagebox.showinfo("Estimación de costo", est.resumen())
        except Exception as e:  # noqa: BLE001
            messagebox.showerror("Error", str(e))

    def _mmx_iniciar(self):
        if self._mmx_corriendo:
            return
        carpeta = self._mmx_entrada.get().strip()
        salida = self._mmx_salida.get().strip()
        if not carpeta or not os.path.isdir(carpeta):
            messagebox.showwarning("Sin carpeta", "Elegí la carpeta de imágenes.")
            return
        if not salida:
            messagebox.showwarning("Sin salida", "Elegí la carpeta de salida.")
            return

        from core import extractor_multimodal as em
        try:
            imgs = em.listar_imagenes(carpeta)
        except Exception as e:  # noqa: BLE001
            messagebox.showerror("Error", str(e))
            return
        if not imgs:
            messagebox.showinfo("Sin imágenes", "No hay imágenes en esa carpeta.")
            return

        api_key, prov, modelo = self._mmx_clave()
        if prov != "ollama" and not api_key:
            messagebox.showwarning(
                "Falta API key",
                f"No hay clave para «{prov}».\nConfigurala en ⚙ Configuración.")
            return

        # Estándar de costo IA: estimar y confirmar antes de gastar.
        try:
            est = em.estimar_costo_directorio(carpeta, proveedor=prov, modelo=modelo)
            costo_txt = est.resumen() + "\n\n"
        except Exception:  # noqa: BLE001
            costo_txt = f"{len(imgs)} imagen(es); proveedor {prov}.\n\n"
        if not messagebox.askyesno(
                "Extracción multimodal",
                f"{costo_txt}Se procesarán {len(imgs)} imagen(es) con IA de visión.\n"
                "¿Continuar?"):
            return

        self._mmx_corriendo = True
        self._mmx_btn.config(state="disabled")
        self._mmx_log.delete("1.0", "end")
        self._mmx_prog.config(value=0, maximum=len(imgs))
        threading.Thread(target=self._mmx_worker,
                         args=(carpeta, salida, api_key, prov, modelo),
                         daemon=True).start()

    def _mmx_worker(self, carpeta, salida, api_key, prov, modelo):
        from core import extractor_multimodal as em
        from core import ocr_llm
        self.after(0, lambda: self._mmx_log_add(
            f"Iniciando extracción con {prov} ({modelo or 'default'})…"))

        def cb(i, total, res):
            estado = "✓" if res.ok else "✗"
            detalle = (res.error if not res.ok else
                       f"{len(res.datos.get('imagenes_registro', []))} fotos, "
                       f"{len(res.datos.get('bloque_publicitario', []))} anuncios")
            self.after(0, lambda: (
                self._mmx_prog.config(value=i),
                self._mmx_log_add(f"  [{i}/{total}] {estado} {res.imagen} — {detalle}")))

        try:
            resultados = em.procesar_directorio(
                carpeta, api_key, salida, proveedor=prov, modelo=modelo,
                guardar_json=self._mmx_guardar_json.get(), guardar_md=True,
                callback=cb)
            resumen = em.resumen_lote(resultados)

            # Costo real desde el acumulador de usages (estándar costo-IA).
            costo_str = ""
            try:
                from core import costos
                cr = costos.costo_real_desde_usages(prov, modelo or "", ocr_llm.usages())
                if cr.costo_usd > 0:
                    costo_str = (f"\n💲 Costo real: ${cr.costo_usd:.4f} USD "
                                 f"({cr.tokens_input + cr.tokens_output:,} tokens)")
            except Exception:  # noqa: BLE001
                pass

            # Alimentar el corpus textual existente.
            alim = ""
            if self._mmx_alimentar.get():
                nuevos = [em.json_a_texto_plano(r.datos) for r in resultados
                          if r.ok and em.json_a_texto_plano(r.datos)]
                if nuevos:
                    if getattr(ST, "corpus_txt", None):
                        ST.corpus_txt.extend(nuevos)
                    else:
                        ST.corpus_txt = list(nuevos)
                    alim = f"\n📥 {len(nuevos)} textos añadidos al corpus."
                    self.after(0, self._marcar_modificado)

            msg = (f"Extracción completada\n"
                   f"  Páginas OK:   {resumen['ok']}\n"
                   f"  Fallidas:     {resumen['fallidas']}\n"
                   f"  Fotos:        {resumen['imagenes_detectadas']}\n"
                   f"  Anuncios:     {resumen['anuncios_detectados']}"
                   f"{costo_str}{alim}")
            self.after(0, lambda: self._mmx_log_add("\n" + msg))
            self.after(0, lambda: messagebox.showinfo("Extracción completada", msg))
        except Exception as e:  # noqa: BLE001
            err = str(e)
            self.after(0, lambda: self._mmx_log_add(f"ERROR: {err}"))
            self.after(0, lambda: messagebox.showerror("Error", err))
        finally:
            self._mmx_corriendo = False
            self.after(0, lambda: self._mmx_btn.config(state="normal"))

    def _mmx_abrir_salida(self):
        salida = self._mmx_salida.get().strip()
        if salida and os.path.isdir(salida):
            import subprocess
            subprocess.Popen(["explorer", os.path.normpath(salida)])

    # ══════════════════════════════════════════════════════════════════════════
    # TAB ETZ: ETIQUETADOR DE ZONAS
    # ══════════════════════════════════════════════════════════════════════════
    def _build_etz(self):
        """
        Etiquetador estilo ABBYY FineReader — 4 paneles sincronizados:
          Páginas (izq) | Imagen+Zonas (centro) | Texto OCR (der) | Zoom (inferior)
        """
        f = self._tab_etz

        # ── Estado interno ─────────────────────────────────────────────────────
        self._etz_numero    = tk.StringVar(value="")
        self._etz_pagina    = tk.StringVar(value="")
        self._etz_tipo      = tk.StringVar(value="articulo")
        self._etz_img_orig  = None
        self._etz_img_tk    = None
        self._etz_escala    = 1.0
        self._etz_rect_ini  = None
        self._etz_rect_tmp  = None
        self._etz_zonas     = []
        self._etz_canvas_ids = []
        self._etz_detector  = None
        self._etz_modo          = None
        self._etz_resize_idx    = None
        self._etz_resize_handle = None
        self._etz_move_offset   = None
        self._etz_zona_sel_idx  = None
        self._etz_zoom          = 1.0
        self._etz_pan_start     = None
        self._etz_space_held    = False
        self._etz_modo_det  = tk.StringVar(value="tesseract")

        from core.zone_labeler import TIPOS_ZONA, DetectorZonas
        self._etz_detector = DetectorZonas()

        # Colores FineReader para tipos de zona
        _FR_COLORS = {
            "articulo":   "#22AA22",   # verde texto
            "titulo":     "#0066FF",   # azul titular
            "publicidad": "#CC2222",   # rojo imagen/publicidad
            "foto":       "#CC2222",
            "pie_foto":   "#FF8800",
            "numero_pag": "#888888",
            "cabecera":   "#8855BB",
            "indice":     "#0099CC",
            "colofon":    "#CC6688",
        }
        # Actualizar colores de TIPOS_ZONA con los de FineReader
        for tid, color in _FR_COLORS.items():
            if tid in TIPOS_ZONA:
                TIPOS_ZONA[tid]["color"] = color

        # ── TOOLBAR RIBBON — 2 filas para no desbordar ────────────────────────
        ribbon_wrap = tk.Frame(f, bg="#2D333B")
        ribbon_wrap.pack(fill="x")

        # Fila 1: Navegación + Tipos de zona
        row1 = tk.Frame(ribbon_wrap, bg="#2D333B", height=30)
        row1.pack(fill="x")
        row1.pack_propagate(False)

        def _rb_btn(parent, text, cmd, bg="#2D333B", fg="#CDD6F4", bold=False):
            b = tk.Label(parent, text=text, bg=bg, fg=fg, cursor="hand2",
                         font=("Segoe UI", 7, "bold" if bold else "normal"),
                         padx=5, pady=2)
            b.pack(side="left", padx=1)
            b.bind("<Button-1>", lambda e, c=cmd: c())
            orig_bg = bg
            b.bind("<Enter>", lambda e, w=b: w.config(bg="#444C56"))
            b.bind("<Leave>", lambda e, w=b, ob=orig_bg: w.config(bg=ob))
            return b

        def _rb_sep(parent):
            tk.Frame(parent, bg="#444C56", width=1).pack(
                side="left", fill="y", pady=3, padx=3)

        # Abrir PDF
        _rb_btn(row1, "📂 PDF", self._etz_abrir_pdf_directo,
                bg="#1F6FEB", fg="white", bold=True)
        _rb_sep(row1)

        # Número
        tk.Label(row1, text="N°:", bg="#2D333B", fg="#8B949E",
                 font=("Segoe UI", 7)).pack(side="left", padx=(4, 1))
        self._etz_cb_num = ttk.Combobox(row1, textvariable=self._etz_numero,
                                         width=18, state="readonly", font=("Segoe UI", 8))
        self._etz_cb_num.pack(side="left", padx=(0, 3))
        self._etz_cb_num.bind("<<ComboboxSelected>>", self._etz_on_numero)

        tk.Label(row1, text="Pág:", bg="#2D333B", fg="#8B949E",
                 font=("Segoe UI", 7)).pack(side="left", padx=(2, 1))
        self._etz_cb_pag = ttk.Combobox(row1, textvariable=self._etz_pagina,
                                         width=8, state="readonly", font=("Segoe UI", 8))
        self._etz_cb_pag.pack(side="left", padx=(0, 2))
        self._etz_cb_pag.bind("<<ComboboxSelected>>", self._etz_on_pagina)

        for txt, cmd in [("◀", self._etz_pagina_ant), ("▶", self._etz_pagina_sig)]:
            b = tk.Label(row1, text=txt, bg="#2D333B", fg="#CDD6F4",
                         font=("Segoe UI", 9, "bold"), cursor="hand2", padx=3)
            b.pack(side="left", padx=1)
            b.bind("<Button-1>", lambda e, c=cmd: c())
            b.bind("<Enter>", lambda e, w=b: w.config(bg="#388BFD"))
            b.bind("<Leave>", lambda e, w=b: w.config(bg="#2D333B"))

        _rb_sep(row1)

        # Tipos de zona
        tk.Label(row1, text="Zona:", bg="#2D333B", fg="#8B949E",
                 font=("Segoe UI", 7)).pack(side="left", padx=(3, 2))
        self._etz_tipo_btns = {}
        for tid, meta in TIPOS_ZONA.items():
            color = meta.get("color", "#888")
            lbl   = meta.get("label", tid)[:7]
            btn = tk.Label(row1, text=lbl, bg=color, fg="white",
                           font=("Segoe UI", 7, "bold"), padx=4, pady=1,
                           cursor="hand2", relief="flat")
            btn.pack(side="left", padx=1)
            btn.bind("<Button-1>", lambda e, t=tid: self._etz_sel_tipo(t))
            self._etz_tipo_btns[tid] = btn

        _rb_sep(row1)
        _rb_btn(row1, "＋ Tipo", self._etz_agregar_tipo_custom,
                fg="#3FB950", bold=True)
        _rb_btn(row1, "🔤 Tipogr.", self._etz_deepfont_zona,
                fg="#F0883E", bold=True)
        _rb_sep(row1)
        _rb_btn(row1, "📊 Estadísticas", self._etz_mostrar_estadisticas,
                fg="#A371F7", bold=True)

        # Fila 2: Acciones + Detección + Estado
        row2 = tk.Frame(ribbon_wrap, bg="#1C2128", height=28)
        row2.pack(fill="x")
        row2.pack_propagate(False)

        tk.Frame(row2, bg="#444C56", width=1).pack(side="left", fill="y", pady=2)

        for txt, cmd in [
            ("🗑 Última",   self._etz_borrar_ultima),
            ("🗑 Todo",     self._etz_limpiar_todo),
            ("💾 Guardar",  self._etz_guardar_pagina),
            ("⟳ Inclinar", self._etz_deskew_pagina),
        ]:
            b = tk.Label(row2, text=txt, bg="#1C2128", fg="#8B949E",
                         font=("Segoe UI", 7), cursor="hand2", padx=6, pady=2)
            b.pack(side="left", padx=1)
            b.bind("<Button-1>", lambda e, c=cmd: c())
            b.bind("<Enter>", lambda e, w=b: w.config(bg="#2D333B", fg="#CDD6F4"))
            b.bind("<Leave>", lambda e, w=b: w.config(bg="#1C2128", fg="#8B949E"))

        tk.Frame(row2, bg="#444C56", width=1).pack(side="left", fill="y", pady=2, padx=4)

        tk.Label(row2, text="Detectar:", bg="#1C2128", fg="#8B949E",
                 font=("Segoe UI", 7)).pack(side="left", padx=(4, 3))

        for txt, cmd in [("📄 Esta página", self._etz_detectar_pagina),
                          ("📚 Todo el número", self._etz_detectar_numero)]:
            b = tk.Label(row2, text=txt, bg="#1F6FEB", fg="white",
                         font=("Segoe UI", 7, "bold"), cursor="hand2", padx=6, pady=2)
            b.pack(side="left", padx=2)
            b.bind("<Button-1>", lambda e, c=cmd: c())
            b.bind("<Enter>", lambda e, w=b: w.config(bg="#388BFD"))
            b.bind("<Leave>", lambda e, w=b: w.config(bg="#1F6FEB"))

        # OCR por zonas — reconoce cada zona por separado en orden de lectura
        b_oz = tk.Label(row2, text="👁 OCR zonas", bg="#238636", fg="white",
                        font=("Segoe UI", 7, "bold"), cursor="hand2",
                        padx=6, pady=2)
        b_oz.pack(side="left", padx=2)
        b_oz.bind("<Button-1>", lambda e: self._etz_ocr_zonas_preview())
        b_oz.bind("<Enter>", lambda e: b_oz.config(bg="#2EA043"))
        b_oz.bind("<Leave>", lambda e: b_oz.config(bg="#238636"))

        tk.Frame(row2, bg="#444C56", width=1).pack(side="left", fill="y", pady=2, padx=4)

        # Botón predicción — aplica la plantilla aprendida al resto del número
        self._btn_etz_predecir = tk.Label(
            row2, text="🔮 Predecir resto", bg="#6E40C9", fg="white",
            font=("Segoe UI", 7, "bold"), cursor="hand2", padx=6, pady=2)
        self._btn_etz_predecir.pack(side="left", padx=2)
        self._btn_etz_predecir.bind("<Button-1>", lambda e: self._etz_predecir_numero())
        self._btn_etz_predecir.bind("<Enter>",
            lambda e: self._btn_etz_predecir.config(bg="#8957E5"))
        self._btn_etz_predecir.bind("<Leave>",
            lambda e: self._btn_etz_predecir.config(bg="#6E40C9"))

        tk.Label(row2, text="Motor:", bg="#1C2128", fg="#8B949E",
                 font=("Segoe UI", 7)).pack(side="left", padx=(8, 2))

        _cb_motor = ttk.Combobox(row2, textvariable=self._etz_modo_det,
                     values=["tesseract", "opencv", "yolo", "onnx", "dit", "vision_ia"],
                     state="readonly", width=9,
                     font=("Segoe UI", 7))
        _cb_motor.pack(side="left")

        # ── Selector proveedor + modelo (visible solo con vision_ia) ─────────
        self._etz_vision_frame = tk.Frame(row2, bg="#1C2128")
        self._etz_vision_frame.pack(side="left", padx=(4, 0))

        from core.zone_labeler import VISION_PROVEEDORES
        self._etz_vision_prov  = tk.StringVar(value="claude")
        self._etz_vision_model = tk.StringVar(value="claude-sonnet-4-6")

        _cb_prov = ttk.Combobox(self._etz_vision_frame,
                                  textvariable=self._etz_vision_prov,
                                  values=list(VISION_PROVEEDORES.keys()),
                                  state="readonly", width=8,
                                  font=("Segoe UI", 7))
        _cb_prov.pack(side="left", padx=(0, 2))

        self._etz_cb_vision_model = ttk.Combobox(
            self._etz_vision_frame,
            textvariable=self._etz_vision_model,
            state="readonly", width=18,
            font=("Segoe UI", 7))
        self._etz_cb_vision_model.pack(side="left", padx=(0, 2))

        # Botón editar prompt
        _btn_prompt = tk.Label(self._etz_vision_frame, text="✏ Prompt",
                                bg="#1C2128", fg="#8B949E",
                                font=("Segoe UI", 7), cursor="hand2", padx=4)
        _btn_prompt.pack(side="left")
        _btn_prompt.bind("<Button-1>", lambda e: self._etz_editar_prompt())
        _btn_prompt.bind("<Enter>",    lambda e: _btn_prompt.config(fg="#CDD6F4"))
        _btn_prompt.bind("<Leave>",    lambda e: _btn_prompt.config(fg="#8B949E"))

        # ? tooltip proveedor
        _PROV_AYUDA = "\n\n".join(
            f"{k} — {v['label']}\n  {v['help']}"
            for k, v in VISION_PROVEEDORES.items()
        )
        _btn_qp = tk.Label(self._etz_vision_frame, text="?", bg="#1C2128",
                            fg="#8B949E", font=("Segoe UI", 7, "bold"),
                            cursor="hand2", padx=2)
        _btn_qp.pack(side="left")
        _btn_qp.bind("<Enter>",    lambda e: self._mostrar_tooltip(_PROV_AYUDA, _btn_qp))
        _btn_qp.bind("<Leave>",    lambda e: self._ocultar_tooltip())
        _btn_qp.bind("<Button-1>", lambda e: messagebox.showinfo(
            "Proveedores de visión IA", _PROV_AYUDA))

        def _on_prov_change(*_):
            prov = self._etz_vision_prov.get()
            info = VISION_PROVEEDORES.get(prov, {})
            modelos = info.get("modelos", [])
            self._etz_cb_vision_model["values"] = modelos
            self._etz_vision_model.set(info.get("default", modelos[0] if modelos else ""))

        def _on_motor_change(*_):
            is_vision = self._etz_modo_det.get() == "vision_ia"
            if is_vision:
                self._etz_vision_frame.pack(side="left", padx=(4, 0))
            else:
                self._etz_vision_frame.pack_forget()

        self._etz_vision_prov.trace_add("write", _on_prov_change)
        self._etz_modo_det.trace_add("write", _on_motor_change)
        _on_prov_change()   # inicializar modelos del proveedor por defecto
        _on_motor_change()  # ocultar si el motor inicial no es vision_ia

        # Botón instalar dependencias del motor seleccionado
        self._btn_instalar_motor = tk.Label(
            row2, text="⬇ Instalar", bg="#1C2128", fg="#8B949E",
            font=("Segoe UI", 7), cursor="hand2", padx=4)
        self._btn_instalar_motor.pack(side="left", padx=2)
        self._btn_instalar_motor.bind("<Button-1>", lambda e: self._etz_instalar_motor())

        # Botón ? — explica cada motor
        _MOTOR_AYUDA = (
            "Motores de detección de zonas:\n\n"
            "tesseract — Layout engine de Tesseract + OpenCV (estilo FineReader).\n"
            "            100% local, sin instalación extra. Detecta bloques de\n"
            "            texto reales, títulos, fotos, pies de foto y orden de\n"
            "            lectura por columnas. RECOMENDADO.\n\n"
            "opencv    — OpenCV local, sin IA. Rápido (<1s). Sin instalación extra.\n\n"
            "yolo      — YOLOv8n-DocLayNet (~6 MB). CPU ~2s/pág. 11 tipos de zona.\n"
            "            pip install ultralytics\n\n"
            "onnx      — YOLOS-DocLayNet ONNX (~45 MB). CPU ~1s/pág. Sin torch.\n"
            "            pip install onnxruntime\n\n"
            "dit       — Microsoft DiT (~330 MB). Para el PC nuevo con ≥16 GB RAM.\n"
            "            pip install transformers torch\n\n"
            "vision_ia — Cualquier IA de visión via API: Claude, GPT-4o,\n"
            "            Gemini o Llava (Ollama local). Usa el prompt editable.\n"
            "            El prompt se puede personalizar con ✏ Prompt."
        )
        _btn_q = tk.Label(row2, text="?", bg="#1C2128", fg="#8B949E",
                          font=("Segoe UI", 7, "bold"), cursor="hand2", padx=3)
        _btn_q.pack(side="left")
        _btn_q.bind("<Enter>",    lambda e: self._mostrar_tooltip(_MOTOR_AYUDA, _btn_q))
        _btn_q.bind("<Leave>",    lambda e: self._ocultar_tooltip())
        _btn_q.bind("<Button-1>", lambda e: messagebox.showinfo(
            "Motores de detección", _MOTOR_AYUDA))

        # Status — lado derecho fila 2
        self._etz_lbl_train = tk.Label(row2, text="",
                                        bg="#1C2128", fg="#3FB950",
                                        font=("Segoe UI", 7))
        self._etz_lbl_train.pack(side="right", padx=12)

        # ── BODY: 3 paneles horizontales (PanedWindow) ─────────────────────────
        # Layout FineReader: [Páginas 160px] | [Imagen central] | [Texto OCR 220px]
        body_paned = tk.PanedWindow(f, orient="horizontal",
                                     sashwidth=4, sashpad=0,
                                     bg=CARD_BOR, relief="flat",
                                     handlesize=0)
        body_paned.pack(fill="both", expand=True)

        # ── PANEL IZQUIERDO: miniaturas de páginas (estilo Pages Pane) ─────────
        pages_frm = tk.Frame(body_paned, bg="#161B22", width=160)
        pages_frm.pack_propagate(False)
        body_paned.add(pages_frm, minsize=120, width=160)

        tk.Label(pages_frm, text="PÁGINAS", bg="#161B22", fg=TXT_DIM,
                 font=("Segoe UI", 7, "bold")).pack(anchor="w", padx=8, pady=(8, 4))

        # Canvas scrollable para miniaturas
        pages_canvas = tk.Canvas(pages_frm, bg="#161B22",
                                  highlightthickness=0, width=152)
        pages_sb = tk.Scrollbar(pages_frm, orient="vertical",
                                 command=pages_canvas.yview)
        pages_canvas.configure(yscrollcommand=pages_sb.set)
        pages_sb.pack(side="right", fill="y")
        pages_canvas.pack(fill="both", expand=True)
        pages_inner = tk.Frame(pages_canvas, bg="#161B22")
        pages_win = pages_canvas.create_window((0, 0), window=pages_inner, anchor="nw")
        def _pages_cfg(e): pages_canvas.configure(scrollregion=pages_canvas.bbox("all"))
        def _pages_cw(e): pages_canvas.itemconfig(pages_win, width=e.width)
        def _pages_wheel(e):
            pages_canvas.yview_scroll(-1 if (e.delta > 0 or e.num == 4) else 1, "units")
        pages_inner.bind("<Configure>", _pages_cfg)
        pages_canvas.bind("<Configure>", _pages_cw)
        pages_canvas.bind("<MouseWheel>", _pages_wheel)
        pages_inner.bind("<MouseWheel>", _pages_wheel)
        self._etz_pages_inner  = pages_inner
        self._etz_pages_canvas = pages_canvas
        self._etz_thumb_btns   = {}   # {pagina: frame_miniatura}

        # ── PANEL CENTRAL: imagen con zonas ────────────────────────────────────
        center_frm = tk.Frame(body_paned, bg="#0D1117")
        body_paned.add(center_frm, minsize=400)

        # Sub-PanedWindow vertical: imagen arriba | zoom inferior
        center_paned = tk.PanedWindow(center_frm, orient="vertical",
                                       sashwidth=4, bg=CARD_BOR,
                                       relief="flat", handlesize=0)
        center_paned.pack(fill="both", expand=True)

        # Image Pane (canvas principal)
        img_frm = tk.Frame(center_paned, bg="#0D1117")
        center_paned.add(img_frm, minsize=300)

        # Toolbar mínima del Image Pane (zoom)
        img_tb = tk.Frame(img_frm, bg="#161B22", height=24)
        img_tb.pack(fill="x")
        img_tb.pack_propagate(False)
        tk.Label(img_tb, text="Ctrl+rueda: zoom  ·  Rueda: scroll  ·  "
                              "Medio/Space+drag: pan  ·  Clic der: menú",
                 bg="#161B22", fg=TXT_DIM,
                 font=("Segoe UI", 7)).pack(side="left", padx=8)
        self._etz_lbl_zoom = tk.Label(img_tb, text="100%",
                                       bg="#161B22", fg=TXT_SEC,
                                       font=("Segoe UI", 7, "bold"))
        self._etz_lbl_zoom.pack(side="right", padx=8)

        canvas_wrap = tk.Frame(img_frm, bg="#0D1117")
        canvas_wrap.pack(fill="both", expand=True)

        self._etz_canvas = tk.Canvas(canvas_wrap, bg="#1A1F2B",
                                      cursor="crosshair",
                                      highlightthickness=0)
        etz_scroll_y = tk.Scrollbar(canvas_wrap, orient="vertical",
                                     command=self._etz_canvas.yview)
        etz_scroll_x = tk.Scrollbar(canvas_wrap, orient="horizontal",
                                     command=self._etz_canvas.xview)
        self._etz_canvas.configure(yscrollcommand=etz_scroll_y.set,
                                    xscrollcommand=etz_scroll_x.set)
        etz_scroll_y.pack(side="right", fill="y")
        etz_scroll_x.pack(side="bottom", fill="x")
        self._etz_canvas.pack(fill="both", expand=True)

        # Eventos del canvas (mismos que antes)
        self._etz_canvas.bind("<ButtonPress-1>",      self._etz_on_press)
        self._etz_canvas.bind("<B1-Motion>",           self._etz_on_drag)
        self._etz_canvas.bind("<ButtonRelease-1>",     self._etz_on_release)
        self._etz_canvas.bind("<Motion>",              self._etz_on_motion)
        self._etz_canvas.bind("<ButtonPress-3>",       self._etz_on_click_derecho)
        self._etz_canvas.bind("<Control-MouseWheel>",  self._etz_on_zoom)
        self._etz_canvas.bind("<ButtonPress-2>",       self._etz_pan_start_cb)
        self._etz_canvas.bind("<B2-Motion>",           self._etz_pan_drag_cb)
        self._etz_canvas.bind("<ButtonRelease-2>",     self._etz_pan_end_cb)
        self._etz_canvas.bind("<KeyPress-space>",      self._etz_space_press)
        self._etz_canvas.bind("<KeyRelease-space>",    self._etz_space_release)
        self._etz_canvas.bind("<Delete>",              self._etz_suprimir_sel)
        self._etz_canvas.bind("<BackSpace>",           self._etz_suprimir_sel)
        self._etz_canvas.bind("<ButtonPress-1>",       lambda e: self._etz_canvas.focus_set(), add=True)
        self._etz_canvas.bind("<MouseWheel>",
            lambda e: self._etz_canvas.yview_scroll(
                -1 if e.delta > 0 else 1, "units"))
        self._etz_canvas.focus_set()

        # Zoom Pane (inferior) — detalle ampliado de la zona activa
        zoom_frm = tk.Frame(center_paned, bg="#161B22", height=120)
        center_paned.add(zoom_frm, minsize=80, height=120)

        tk.Label(zoom_frm, text="DETALLE (Zoom Pane)", bg="#161B22", fg=TXT_DIM,
                 font=("Segoe UI", 7, "bold")).pack(anchor="w", padx=8, pady=(4, 2))
        self._etz_zoom_canvas = tk.Canvas(zoom_frm, bg="#1A1F2B",
                                           highlightthickness=0, height=90)
        self._etz_zoom_canvas.pack(fill="both", expand=True, padx=4, pady=(0, 4))
        self._etz_zoom_img_tk = None

        # ── PANEL DERECHO: texto OCR + lista zonas ─────────────────────────────
        right_frm = tk.Frame(body_paned, bg="#161B22", width=220)
        right_frm.pack_propagate(False)
        body_paned.add(right_frm, minsize=160, width=220)

        # Sub-PanedWindow vertical: texto OCR arriba | zonas abajo
        right_paned = tk.PanedWindow(right_frm, orient="vertical",
                                      sashwidth=4, bg=CARD_BOR,
                                      relief="flat", handlesize=0)
        right_paned.pack(fill="both", expand=True)

        # Text Pane — texto OCR reconocido (sincronizado con imagen)
        txt_frm = tk.Frame(right_paned, bg="#161B22")
        right_paned.add(txt_frm, minsize=120)

        tk.Label(txt_frm, text="TEXTO OCR", bg="#161B22", fg=TXT_DIM,
                 font=("Segoe UI", 7, "bold")).pack(anchor="w", padx=8, pady=(6, 2))
        self._etz_txt_ocr = scrolledtext.ScrolledText(
            txt_frm, font=("Consolas", 8),
            bg="#0D1117", fg="#CDD6F4",
            insertbackground="#CDD6F4",
            selectbackground="#1F6FEB",
            relief="flat", padx=6, pady=4,
            wrap="word", state="normal")
        self._etz_txt_ocr.pack(fill="both", expand=True, padx=4)
        # Palabras con baja confianza → subrayado azul (estilo FineReader)
        self._etz_txt_ocr.tag_configure("low_conf",
            foreground="#58A6FF", underline=True)

        # Zones Pane — lista de zonas + acciones
        zones_frm = tk.Frame(right_paned, bg="#161B22")
        right_paned.add(zones_frm, minsize=100)

        tk.Label(zones_frm, text="ZONAS", bg="#161B22", fg=TXT_DIM,
                 font=("Segoe UI", 7, "bold")).pack(anchor="w", padx=8, pady=(6, 2))

        zona_list_frame = tk.Frame(zones_frm, bg="#161B22")
        zona_list_frame.pack(fill="both", expand=True, padx=4)
        self._etz_zona_list = tk.Listbox(zona_list_frame,
                                          font=("Segoe UI", 8),
                                          bg="#0D1117", fg="#CDD6F4",
                                          selectbackground="#1F6FEB",
                                          selectforeground="white",
                                          activestyle="none",
                                          relief="flat",
                                          selectmode="single", height=8)
        zona_scroll = tk.Scrollbar(zona_list_frame,
                                    command=self._etz_zona_list.yview)
        self._etz_zona_list.configure(yscrollcommand=zona_scroll.set)
        zona_scroll.pack(side="right", fill="y")
        self._etz_zona_list.pack(side="left", fill="both", expand=True)
        self._etz_zona_list.bind("<<ListboxSelect>>", self._etz_on_zona_sel)
        self._etz_zona_list.bind("<Delete>",           self._etz_suprimir_sel)
        self._etz_zona_list.bind("<BackSpace>",        self._etz_suprimir_sel)

        # Estado
        self._etz_lbl_estado = tk.Label(zones_frm, text="—",
                                         bg="#161B22", fg=TXT_SEC,
                                         font=("Segoe UI", 7),
                                         wraplength=200, justify="left")
        self._etz_lbl_estado.pack(anchor="w", padx=8, pady=(2, 4))

        # ── STATUS BAR inferior ────────────────────────────────────────────────
        status_bar = tk.Frame(f, bg="#161B22", height=20)
        status_bar.pack(fill="x", side="bottom")
        status_bar.pack_propagate(False)
        self._etz_lbl_coords = tk.Label(status_bar, text="x:— y:—",
                                         bg="#161B22", fg=TXT_DIM,
                                         font=("Segoe UI", 7))
        self._etz_lbl_coords.pack(side="left", padx=8)
        self._etz_lbl_tipo_activo = tk.Label(status_bar, text="",
                                              bg="#161B22", fg="#3FB950",
                                              font=("Segoe UI", 7, "bold"))
        self._etz_lbl_tipo_activo.pack(side="left", padx=16)
        tk.Label(status_bar,
                 text="Dibuja: clic+arrastre  ·  Mover: Ctrl+arrastre  ·  Eliminar: clic der",
                 bg="#161B22", fg=TXT_DIM, font=("Segoe UI", 7)).pack(side="right", padx=8)

        # Inicializar
        self._etz_sel_tipo("articulo")

    # ── Helpers internos del etiquetador ─────────────────────────────────────

    def _etz_sel_tipo(self, tipo: str):
        from core.zone_labeler import TIPOS_ZONA
        self._etz_tipo.set(tipo)
        for tid, btn in self._etz_tipo_btns.items():
            meta = TIPOS_ZONA[tid]
            color = meta["color"]
            if tid == tipo:
                # Activo: borde blanco visible
                btn.config(relief="solid", bd=2,
                           bg=color, fg="white",
                           font=("Segoe UI", 7, "bold"),
                           highlightbackground="white",
                           highlightthickness=2)
            else:
                btn.config(relief="flat", bd=0,
                           bg=color, fg="white",
                           font=("Segoe UI", 7),
                           highlightbackground=color,
                           highlightthickness=0)
        # Status bar
        if hasattr(self, "_etz_lbl_tipo_activo"):
            label = TIPOS_ZONA.get(tipo, {}).get("label", tipo)
            color = TIPOS_ZONA.get(tipo, {}).get("color", "#888")
            self._etz_lbl_tipo_activo.config(
                text=f"● {label}", fg=color)

    def _etz_actualizar_lista_zonas(self):
        from core.zone_labeler import TIPOS_ZONA
        self._etz_zona_list.delete(0, "end")
        for i, z in enumerate(self._etz_zonas):
            meta = TIPOS_ZONA.get(z.tipo, {})
            label = meta.get("label", z.tipo)
            conf_str = f" [{z.confianza:.0%}]" if z.confianza < 1.0 else ""
            orden_str = f"#{z.orden} " if getattr(z, "orden", 0) else ""
            self._etz_zona_list.insert(
                "end",
                f"{i+1}. {orden_str}{label} ({z.x0:.2f},{z.y0:.2f})–({z.x1:.2f},{z.y1:.2f}){conf_str}"
            )

    def _etz_redibujar_zonas(self):
        """Redibuja todos los rectángulos con handles visuales y etiquetas contrastadas."""
        from core.zone_labeler import TIPOS_ZONA
        for cid in self._etz_canvas_ids:
            self._etz_canvas.delete(cid)
        self._etz_canvas_ids.clear()

        if self._etz_img_orig is None:
            return

        R = self._ETZ_HANDLE_D
        idx_sel = getattr(self, "_etz_zona_sel_idx", None)

        for i, z in enumerate(self._etz_zonas):
            x0p, y0p, x1p, y1p = self._etz_zona_canvas_coords(i)
            lx, rx = min(x0p, x1p), max(x0p, x1p)
            ty, by = min(y0p, y1p), max(y0p, y1p)
            color  = TIPOS_ZONA.get(z.tipo, {}).get("color", "#888888")
            label  = TIPOS_ZONA.get(z.tipo, {}).get("label", z.tipo)
            manual = z.confianza >= 1.0
            activa = (i == idx_sel)

            # Relleno semitransparente — zona activa más visible
            stipple = "gray25" if activa else "gray12"
            cid_fill = self._etz_canvas.create_rectangle(
                lx, ty, rx, by,
                outline="", fill=color, stipple=stipple
            )

            # Borde principal
            grosor = 3 if activa else (2 if manual else 1)
            dash   = () if manual else (6, 3)
            cid_rect = self._etz_canvas.create_rectangle(
                lx, ty, rx, by,
                outline=color, width=grosor, dash=dash, fill=""
            )

            # Etiqueta con fondo sólido para contraste
            lbl_x = lx + 5
            lbl_y = ty + 4
            lbl_text = label if manual else f"{label} {z.confianza:.0%}"
            # Fondo negro de la etiqueta
            cid_bg = self._etz_canvas.create_rectangle(
                lbl_x - 2, lbl_y - 1,
                lbl_x + len(lbl_text) * 6 + 2, lbl_y + 13,
                fill="#000000", outline="", stipple=""
            )
            cid_lbl = self._etz_canvas.create_text(
                lbl_x, lbl_y,
                text=lbl_text, anchor="nw",
                font=("Segoe UI", 8, "bold"), fill="white"
            )

            # Badge circular con el orden de lectura (esquina superior derecha)
            ids_orden = []
            if getattr(z, "orden", 0):
                bx, by_ = rx - 12, ty + 12
                cid_oc = self._etz_canvas.create_oval(
                    bx - 9, by_ - 9, bx + 9, by_ + 9,
                    fill="#1F6FEB", outline="white", width=1)
                cid_on = self._etz_canvas.create_text(
                    bx, by_, text=str(z.orden),
                    font=("Segoe UI", 8, "bold"), fill="white")
                ids_orden = [cid_oc, cid_on]

            # Handles en las 4 esquinas + 4 bordes medios (solo zona activa o hover)
            ids_handles = []
            if activa or (rx - lx) > 30:
                mx = (lx + rx) / 2
                my = (ty + by) / 2
                puntos_handle = [
                    (lx, ty), (mx, ty), (rx, ty),
                    (rx, my),
                    (rx, by), (mx, by), (lx, by),
                    (lx, my),
                ]
                for hx, hy in puntos_handle:
                    h_bg = self._etz_canvas.create_rectangle(
                        hx - R, hy - R, hx + R, hy + R,
                        fill="#1E293B", outline=color, width=1
                    )
                    h_sq = self._etz_canvas.create_rectangle(
                        hx - R + 2, hy - R + 2, hx + R - 2, hy + R - 2,
                        fill=color, outline=""
                    )
                    ids_handles.extend([h_bg, h_sq])

            self._etz_canvas_ids.extend(
                [cid_fill, cid_rect, cid_bg, cid_lbl] + ids_orden + ids_handles
            )

        # Línea punteada entre pie de foto y su foto vinculada (z.vinculo = zid).
        por_zid = {z.zid: i for i, z in enumerate(self._etz_zonas)}
        for z in self._etz_zonas:
            if z.tipo != "pie_foto" or not z.vinculo or z.vinculo not in por_zid:
                continue
            x0p, y0p, x1p, y1p = self._etz_zona_canvas_coords(por_zid[z.zid])
            fx0, fy0, fx1, fy1 = self._etz_zona_canvas_coords(por_zid[z.vinculo])
            cid_link = self._etz_canvas.create_line(
                (x0p + x1p) / 2, (y0p + y1p) / 2,
                (fx0 + fx1) / 2, (fy0 + fy1) / 2,
                fill="#FF8800", width=2, dash=(4, 3), arrow="last")
            self._etz_canvas_ids.append(cid_link)

    def _etz_cargar_imagen_pagina(self, numero: str, pagina: str):
        """Carga la imagen de la página desde el PDF o desde caché."""
        if not ST.out_dir or not ST.archivos_sel:
            return

        try:

            from PIL import Image, ImageTk

            # Buscar el PDF del número — primero en archivos seleccionados,
            # luego en el directorio de entrada como fallback
            pdf_path = None
            for a in (ST.archivos_sel or []):
                if a.stem == numero:
                    pdf_path = a; break
            if pdf_path is None and ST.pdf_dir and Path(ST.pdf_dir).exists():
                for ext in ("*.pdf", "*.PDF"):
                    candidatos = list(Path(ST.pdf_dir).glob(ext))
                    for c in candidatos:
                        if c.stem == numero:
                            pdf_path = c; break
                    if pdf_path:
                        break
            if pdf_path is None:
                # Intentar cargar desde imágenes ya extraídas aunque no haya PDF
                img_dir = ST.out_dir / "02_imagenes" / numero if ST.out_dir else None
                if not (img_dir and img_dir.exists()):
                    self._etz_canvas.delete("all")
                    self._etz_canvas.create_text(
                        150, 100,
                        text=f"PDF no encontrado para '{numero}'.\n"
                             "Asegurate de tener el PDF en la carpeta\n"
                             "de entrada y de haberlo seleccionado.",
                        fill="#94A3B8", font=("Segoe UI", 9),
                        anchor="nw")
                    return

            # Número de página — manejar múltiples formatos:
            # "p0001"   → página 1 → índice 0
            # "0001"    → página 1 → índice 0
            # "0001-03" → el segundo número es la página → índice 2
            import re as _re
            _nums = _re.findall(r'\d+', pagina)
            if len(_nums) >= 2:
                # Formato "NNNN-PP": el segundo número es la página
                try:
                    n_pag = int(_nums[-1]) - 1
                except ValueError:
                    n_pag = 0
            elif len(_nums) == 1:
                try:
                    n_pag = int(_nums[0]) - 1
                except ValueError:
                    n_pag = 0
            else:
                n_pag = 0
            n_pag = max(0, n_pag)

            # Intentar caché de imágenes
            img_dir = ST.out_dir / "02_imagenes" / numero
            img_candidatos = list(img_dir.glob(f"*{pagina}*.png")) if img_dir.exists() else []
            if not img_candidatos:
                # También buscar por índice (pdf2image nombra p-N.png)
                img_candidatos = sorted(img_dir.glob("*.png")) if img_dir.exists() else []

            if img_candidatos and n_pag < len(img_candidatos):
                img = Image.open(img_candidatos[n_pag])
            else:
                # Convertir desde PDF en memoria (solo la página solicitada)
                from pdf2image import convert_from_path

                from core.ocr_engine import _get_poppler_path
                poppler = _get_poppler_path()
                kwargs = dict(dpi=120, first_page=n_pag+1, last_page=n_pag+1)
                if poppler:
                    kwargs["poppler_path"] = poppler
                imgs = convert_from_path(str(pdf_path), **kwargs)
                if not imgs:
                    return
                img = imgs[0]

            # Escalar para caber en el canvas (máx 900px de alto)
            max_h = 900
            if img.height > max_h:
                escala = max_h / img.height
                img = img.resize((int(img.width * escala), max_h), Image.LANCZOS)
                self._etz_escala = escala
            else:
                self._etz_escala = 1.0

            self._etz_img_orig_full = Image.open(img_candidatos[n_pag]) if (img_candidatos and n_pag < len(img_candidatos)) else img
            self._etz_img_orig = img
            self._etz_img_tk = ImageTk.PhotoImage(img)
            self._etz_canvas.delete("all")
            self._etz_canvas.create_image(0, 0, anchor="nw", image=self._etz_img_tk)
            self._etz_canvas.configure(
                scrollregion=(0, 0, img.width, img.height)
            )
            self._etz_redibujar_zonas()

        except Exception as ex:
            self._etz_canvas.delete("all")
            self._etz_canvas.create_text(
                150, 100, text=f"Error cargando imagen:\n{ex}",
                fill="white", font=("Segoe UI", 9))

    # ── Tipos de zona extensibles ─────────────────────────────────────────────
    def _etz_agregar_tipo_custom(self):
        """Diálogo para crear un nuevo tipo de zona global."""
        from core.zone_labeler import TIPOS_ZONA, agregar_tipo_zona
        win = tk.Toplevel(self)
        win.title("Nuevo tipo de zona")
        win.geometry("380x280")
        win.configure(bg=CONTENT_BG)
        win.grab_set()
        win.resizable(False, False)

        tk.Label(win, text="Nuevo tipo de zona", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 11, "bold")).pack(pady=(16, 8))
        tk.Label(win, text="El tipo quedará disponible en todos los proyectos futuros.",
                 bg=CONTENT_BG, fg=TXT_SEC, font=("Segoe UI", 8)).pack()

        frm = tk.Frame(win, bg=CONTENT_BG)
        frm.pack(fill="x", padx=24, pady=12)

        def _fila(label, var, row):
            tk.Label(frm, text=label, bg=CONTENT_BG, fg=TXT_SEC,
                     font=("Segoe UI", 9), width=10, anchor="e").grid(
                     row=row, column=0, sticky="e", padx=(0, 8), pady=4)
            tk.Entry(frm, textvariable=var, width=24,
                     bg=CARD_BG, fg=TXT_PRI, insertbackground=TXT_PRI,
                     relief="solid", bd=1, font=("Segoe UI", 9)).grid(
                     row=row, column=1, sticky="w")

        var_id    = tk.StringVar(value="tipo_nuevo")
        var_label = tk.StringVar(value="Mi tipo")
        var_color = tk.StringVar(value="#FF6B35")
        var_ocr   = tk.BooleanVar(value=True)

        _fila("ID interno:", var_id, 0)
        _fila("Etiqueta:",   var_label, 1)

        # Color picker
        tk.Label(frm, text="Color:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9), width=10, anchor="e").grid(
                 row=2, column=0, sticky="e", padx=(0, 8), pady=4)
        color_frm = tk.Frame(frm, bg=CONTENT_BG)
        color_frm.grid(row=2, column=1, sticky="w")
        color_entry = tk.Entry(color_frm, textvariable=var_color, width=10,
                               bg=CARD_BG, fg=TXT_PRI, insertbackground=TXT_PRI,
                               relief="solid", bd=1, font=("Segoe UI", 9))
        color_entry.pack(side="left")
        color_preview = tk.Label(color_frm, text="   ", bg=var_color.get(),
                                  relief="flat", width=3)
        color_preview.pack(side="left", padx=(4, 0))

        def _pick_color():
            from tkinter.colorchooser import askcolor
            res = askcolor(color=var_color.get(), parent=win, title="Elegir color")
            if res and res[1]:
                var_color.set(res[1])
                color_preview.config(bg=res[1])
        tk.Button(color_frm, text="…", command=_pick_color,
                  font=("Segoe UI", 8), bg=CARD_BOR, fg=TXT_PRI,
                  relief="flat").pack(side="left", padx=2)
        var_color.trace_add("write", lambda *_: color_preview.config(
            bg=var_color.get() if var_color.get().startswith("#") else CARD_BG))

        tk.Label(frm, text="Procesar OCR:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9), width=10, anchor="e").grid(
                 row=3, column=0, sticky="e", padx=(0, 8), pady=4)
        ttk.Checkbutton(frm, variable=var_ocr,
                         text="Sí — incluir esta zona en el OCR").grid(
                         row=3, column=1, sticky="w")

        def _guardar():
            id_tipo = var_id.get().strip().replace(" ", "_").lower()
            label   = var_label.get().strip()
            color   = var_color.get().strip()
            if not id_tipo or not label:
                messagebox.showwarning("Incompleto",
                    "Ingresa el ID y la etiqueta.", parent=win); return
            if id_tipo in TIPOS_ZONA:
                messagebox.showwarning("Ya existe",
                    f"Ya existe un tipo con ID '{id_tipo}'.", parent=win); return
            agregar_tipo_zona(id_tipo, label, color, var_ocr.get())
            # Agregar botón al ribbon en tiempo real
            self._etz_refrescar_botones_tipo()
            win.destroy()
            messagebox.showinfo("Tipo creado",
                f"Tipo '{label}' creado y disponible en todos los proyectos.")

        btn_frm = tk.Frame(win, bg=CONTENT_BG)
        btn_frm.pack(pady=12)
        ttk.Button(btn_frm, text="✓ Crear tipo", style="P.TButton",
                   command=_guardar).pack(side="left", padx=8)
        ttk.Button(btn_frm, text="Cancelar", style="S.TButton",
                   command=win.destroy).pack(side="left")

    def _etz_refrescar_botones_tipo(self):
        """Recarga los botones de tipo de zona en el ribbon con los tipos actuales."""
        for tid, btn in list(self._etz_tipo_btns.items()):
            try:
                btn.destroy()
            except Exception:
                pass
        self._etz_tipo_btns.clear()
        # Buscar el frame de tipos en el ribbon (reconstruir)
        # Simplificación: indicar al usuario que reinicie para ver el tipo nuevo en el ribbon
        # El tipo ya está en TIPOS_ZONA y se usa correctamente al dibujar zonas
        self._etz_lbl_train.config(
            text="✅ Tipo nuevo disponible — se mostrará al reiniciar")

    # ── Estadísticas de etiquetas del número ─────────────────────────────────
    def _etz_mostrar_estadisticas(self):
        """Panel con estadísticas de las zonas etiquetadas en el número actual."""
        numero = self._etz_numero.get()
        if not numero or not ST.out_dir:
            messagebox.showwarning("Sin número", "Selecciona un número primero.")
            return

        from collections import Counter

        from core.zone_labeler import (
            TIPOS_ZONA,
            cargar_pagina,
            listar_paginas_etiquetadas,
        )

        etiquetadas = listar_paginas_etiquetadas(ST.out_dir, numero)
        if not etiquetadas:
            messagebox.showinfo("Sin etiquetas",
                "Este número no tiene páginas etiquetadas aún.")
            return

        # Contar zonas por tipo
        conteo_tipo    = Counter()
        conteo_manual  = 0
        conteo_pred    = 0
        total_zonas    = 0

        for pag in etiquetadas:
            pd = cargar_pagina(ST.out_dir, numero, pag)
            if not pd:
                continue
            if pd.manual:
                conteo_manual += 1
            else:
                conteo_pred += 1
            for z in pd.zonas:
                conteo_tipo[z.tipo] += 1
                total_zonas += 1

        # Ventana de estadísticas
        win = tk.Toplevel(self)
        win.title(f"Estadísticas de etiquetas — {numero}")
        win.geometry("480x520")
        win.configure(bg=CONTENT_BG)
        win.resizable(False, True)

        self._page_header(win, f"Etiquetas — {numero}",
                          f"{len(etiquetadas)} páginas · {total_zonas} zonas totales",
                          "📊")

        pad = tk.Frame(win, bg=CONTENT_BG, padx=20, pady=10)
        pad.pack(fill="both", expand=True)

        # Resumen de páginas
        res_f = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1)
        res_f.pack(fill="x", pady=(0, 12))
        ri = tk.Frame(res_f, bg=CARD_BG, padx=12, pady=8)
        ri.pack(fill="x")
        for txt, val, color in [
            ("Páginas etiquetadas:",   len(etiquetadas),  TXT_PRI),
            ("  · Manuales:",          conteo_manual,     VERDE),
            ("  · Predichas (IA):",    conteo_pred,       AZ4),
            ("Zonas totales:",         total_zonas,       TXT_PRI),
        ]:
            fila = tk.Frame(ri, bg=CARD_BG)
            fila.pack(fill="x", pady=1)
            tk.Label(fila, text=txt, bg=CARD_BG, fg=TXT_SEC,
                     font=("Segoe UI", 9), width=24, anchor="w").pack(side="left")
            tk.Label(fila, text=str(val), bg=CARD_BG, fg=color,
                     font=("Segoe UI", 9, "bold")).pack(side="left")

        # Desglose por tipo
        tk.Label(pad, text="Zonas por tipo:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", pady=(0, 4))

        cols_s = ("tipo", "label", "cantidad", "porcentaje")
        tv = ttk.Treeview(pad, columns=cols_s, show="headings", height=14)
        for cid, lbl, w in [("tipo","ID tipo",120),("label","Etiqueta",160),
                              ("cantidad","Zonas",70),("porcentaje","%",60)]:
            tv.heading(cid, text=lbl)
            tv.column(cid, width=w, anchor="w")

        # Ordenar por cantidad descendente
        for tipo, n in conteo_tipo.most_common():
            info  = TIPOS_ZONA.get(tipo, {})
            label = info.get("label", tipo)
            pct   = round(n / max(total_zonas, 1) * 100, 1)
            tv.insert("", "end", values=(tipo, label, n, f"{pct}%"))

        # Tipos con 0 zonas (definidos pero no usados en este número)
        for tipo, info in TIPOS_ZONA.items():
            if tipo not in conteo_tipo:
                tv.insert("", "end", values=(tipo, info.get("label", tipo), 0, "0%"),
                          tags=("vacio",))
        tv.tag_configure("vacio", foreground=TXT_DIM)

        sb = ttk.Scrollbar(pad, orient="vertical", command=tv.yview)
        tv.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        tv.pack(fill="both", expand=True)

        # Botón para ir al módulo de descripción de imágenes
        btn_f = tk.Frame(win, bg=CONTENT_BG)
        btn_f.pack(fill="x", padx=20, pady=(6, 12))
        ttk.Button(btn_f, text="🎨  Ir a Descripción de imágenes",
                   style="S.TButton",
                   command=lambda: (win.destroy(),
                                    self._imgd_var_num.set(numero),
                                    self._mostrar_pagina("imgdesc"),
                                    self._imgd_cargar_db())).pack(side="left")
        ttk.Button(btn_f, text="Cerrar",
                   command=win.destroy).pack(side="right")

    # ── Descripción de imágenes etiquetadas (legado — ahora en módulo Analizar) ─
    def _etz_describir_imagenes(self):
        """Abre el panel de descripción automática de zonas de foto del número actual."""
        numero = self._etz_numero.get()
        if not numero or not ST.out_dir:
            messagebox.showwarning("Sin número", "Selecciona un número primero.")
            return

        from core.image_captioner import (
            buscar_imagenes_similares,
            cargar_descripciones_db,
            describir_numero,
        )
        from core.zone_labeler import VISION_PROVEEDORES

        win = tk.Toplevel(self)
        win.title(f"Descripción de imágenes — {numero}")
        win.geometry("1100x680")
        win.configure(bg=CONTENT_BG)

        # ── Barra de control ─────────────────────────────────────────────────
        ctrl = tk.Frame(win, bg=CONTENT_BG)
        ctrl.pack(fill="x", padx=12, pady=(10, 4))

        tk.Label(ctrl, text="Proveedor:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        var_prov  = tk.StringVar(value="claude")
        var_model = tk.StringVar(value="claude-haiku-4-5-20251001")
        cb_prov = ttk.Combobox(ctrl, textvariable=var_prov,
                                values=list(VISION_PROVEEDORES.keys()),
                                state="readonly", width=10)
        cb_prov.pack(side="left", padx=(4, 8))
        cb_model = ttk.Combobox(ctrl, textvariable=var_model,
                                 state="readonly", width=22)
        cb_model.pack(side="left", padx=(0, 12))

        def _on_prov(*_):
            info = VISION_PROVEEDORES.get(var_prov.get(), {})
            mods = info.get("modelos", [])
            cb_model["values"] = mods
            var_model.set(info.get("default", mods[0] if mods else ""))
        var_prov.trace_add("write", _on_prov)
        _on_prov()

        lbl_estado = tk.Label(ctrl, text="", bg=CONTENT_BG, fg=VERDE,
                               font=("Segoe UI", 9, "bold"))
        lbl_estado.pack(side="left")

        btn_run = ttk.Button(ctrl, text="▶  Describir fotos del número",
                              style="P.TButton",
                              command=lambda: _run_describir())
        btn_run.pack(side="right", padx=(8, 0))
        ttk.Button(ctrl, text="↺ Cargar guardadas",
                   style="S.TButton",
                   command=lambda: _cargar_db()).pack(side="right")

        # ── Split: tabla izquierda + detalle derecha ─────────────────────────
        split = tk.Frame(win, bg=CONTENT_BG)
        split.pack(fill="both", expand=True, padx=12, pady=(4, 8))

        # Tabla
        izq = tk.Frame(split, bg=CONTENT_BG, width=500)
        izq.pack(side="left", fill="both", expand=True, padx=(0, 8))

        cols = ("pagina", "descripcion", "categorias", "texto_visible")
        tv = ttk.Treeview(izq, columns=cols, show="headings", height=20)
        for cid, lbl, w in [("pagina","Página",70),("descripcion","Descripción",240),
                              ("categorias","Categorías",130),("texto_visible","Texto visible",110)]:
            tv.heading(cid, text=lbl)
            tv.column(cid, width=w, anchor="w")
        sb_tv = ttk.Scrollbar(izq, orient="vertical", command=tv.yview)
        tv.configure(yscrollcommand=sb_tv.set)
        sb_tv.pack(side="right", fill="y")
        tv.pack(fill="both", expand=True)

        # Panel derecho: recorte + detalle + búsqueda similitud
        der = tk.Frame(split, bg=CARD_BG, width=340, relief="solid", bd=1)
        der.pack(side="right", fill="y")
        der.pack_propagate(False)

        cv_recorte = tk.Canvas(der, bg="#000", height=200, highlightthickness=0)
        cv_recorte.pack(fill="x", padx=6, pady=6)

        lbl_desc  = tk.Label(der, text="", bg=CARD_BG, fg=TXT_PRI,
                              font=("Segoe UI", 9), wraplength=310, justify="left")
        lbl_desc.pack(anchor="w", padx=8, pady=(0, 4))

        lbl_cats  = tk.Label(der, text="", bg=CARD_BG, fg=AZ4,
                              font=("Segoe UI", 8), wraplength=310, justify="left")
        lbl_cats.pack(anchor="w", padx=8)

        lbl_txt   = tk.Label(der, text="", bg=CARD_BG, fg=TXT_SEC,
                              font=("Courier New", 8), wraplength=310, justify="left")
        lbl_txt.pack(anchor="w", padx=8, pady=(0, 4))

        lbl_ctx   = tk.Label(der, text="", bg=CARD_BG, fg=TXT_DIM,
                              font=("Segoe UI", 8, "italic"), wraplength=310, justify="left")
        lbl_ctx.pack(anchor="w", padx=8, pady=(0, 8))

        # Búsqueda por similitud
        tk.Frame(der, bg=CARD_BOR, height=1).pack(fill="x", padx=6, pady=4)
        tk.Label(der, text="🔍 Buscar imágenes similares:",
                 bg=CARD_BG, fg=TXT_PRI, font=("Segoe UI", 8, "bold")).pack(anchor="w", padx=8)
        var_busqueda = tk.StringVar()
        ttk.Entry(der, textvariable=var_busqueda, width=36).pack(padx=8, pady=4, fill="x")
        lbl_sim = tk.Label(der, text="", bg=CARD_BG, fg=TXT_SEC,
                            font=("Segoe UI", 8), wraplength=310, justify="left")
        lbl_sim.pack(anchor="w", padx=8)

        def _buscar_sim():
            q = var_busqueda.get().strip()
            if not q:
                return
            sims = buscar_imagenes_similares(q, ST.out_dir, numero, top_n=5)
            if sims:
                txt = "\n".join(f"  p.{s['pagina']} — {s['descripcion'][:50]}"
                                for s in sims)
                lbl_sim.config(text=f"Similares:\n{txt}")
            else:
                lbl_sim.config(text="Sin índice FAISS disponible aún.")
        ttk.Button(der, text="Buscar", command=_buscar_sim).pack(padx=8, pady=(0, 8))

        # Estado interno
        _descripciones: list[dict] = []

        def _poblar_tv(descs):
            tv.delete(*tv.get_children())
            for d in descs:
                cats = ", ".join(d.get("categorias", []))[:40]
                tv.insert("", "end", values=(
                    d.get("pagina",""), d.get("descripcion","")[:60],
                    cats, d.get("texto_visible","")[:30]))
            _descripciones.clear()
            _descripciones.extend(descs)
            lbl_estado.config(text=f"✅ {len(descs)} descripciones")

        def _on_select(e):
            sel = tv.selection()
            if not sel:
                return
            idx = tv.index(sel[0])
            if idx >= len(_descripciones):
                return
            d = _descripciones[idx]
            lbl_desc.config(text=d.get("descripcion",""))
            lbl_cats.config(text="📌 " + ", ".join(d.get("categorias",[])))
            lbl_txt.config(text=("📝 " + d.get("texto_visible","")) if d.get("texto_visible") else "")
            lbl_ctx.config(text=d.get("contexto_historico",""))
            # Mostrar recorte
            _mostrar_recorte(d)

        tv.bind("<<TreeviewSelect>>", _on_select)

        def _mostrar_recorte(d):
            cv_recorte.delete("all")
            try:
                from PIL import Image, ImageTk
                img_dir = ST.out_dir / "02_imagenes" / numero
                pagina  = d.get("pagina","")
                hits = sorted(img_dir.glob(f"*{pagina}*.png")) if img_dir.exists() else []
                if not hits:
                    return
                img = Image.open(hits[0]).convert("RGB")
                W, H = img.size
                x0 = int(d.get("x0",0) * W)
                y0 = int(d.get("y0",0) * H)
                x1 = int(d.get("x1",1) * W)
                y1 = int(d.get("y1",1) * H)
                recorte = img.crop((x0,y0,x1,y1))
                recorte.thumbnail((320, 190), Image.LANCZOS)
                tk_img = ImageTk.PhotoImage(recorte)
                cv_recorte._ref = tk_img
                cw = cv_recorte.winfo_width() or 320
                cv_recorte.create_image(cw//2, 100, anchor="center", image=tk_img)
            except Exception:
                pass

        def _run_describir():
            prov  = var_prov.get()
            model = var_model.get()
            api_k = ST.api_keys.get(prov, "") or ST.api_key
            if prov != "ollama" and not api_k:
                messagebox.showwarning("Sin API key",
                    f"Configura la API key de {prov} en Configuración.")
                return
            btn_run.config(state="disabled")
            lbl_estado.config(text="⏳ Describiendo…", fg="#F59E0B")

            def cb(n, t, pag, desc):
                self.after(0, lambda: lbl_estado.config(
                    text=f"⏳ {n}/{t}: {pag} — {desc[:40]}…", fg="#F59E0B"))

            def _worker():
                db = Path(ST.ruta_db) if ST.ruta_db else None
                descs = describir_numero(
                    ST.out_dir, numero, proveedor=prov,
                    api_key=api_k, modelo=model,
                    db_path=db, callback=cb)
                self.after(0, lambda: (
                    _poblar_tv(descs),
                    btn_run.config(state="normal"),
                    lbl_estado.config(text=f"✅ {len(descs)} fotos descritas", fg=VERDE),
                ))
            threading.Thread(target=_worker, daemon=True).start()

        def _cargar_db():
            db = Path(ST.ruta_db) if ST.ruta_db else None
            if not db or not db.exists():
                lbl_estado.config(text="⚠ Sin base de datos", fg=ROJO)
                return
            descs = cargar_descripciones_db(db, numero)
            if descs:
                _poblar_tv(descs)
            else:
                lbl_estado.config(text="Sin descripciones guardadas aún", fg=TXT_SEC)

        # Cargar automáticamente si ya hay descripciones guardadas
        win.after(200, _cargar_db)

    # ── DeepFont — clasificación tipográfica de zona seleccionada ─────────────
    def _etz_deepfont_zona(self):
        """Analiza el estilo tipográfico de la zona seleccionada con DeepFont."""
        if self._etz_zona_sel_idx is None:
            messagebox.showinfo("Sin selección",
                "Selecciona una zona haciendo clic sobre ella."); return
        if self._etz_img_orig is None:
            messagebox.showwarning("Sin imagen", "Carga una página primero."); return

        idx = self._etz_zona_sel_idx
        if idx >= len(self._etz_zonas):
            return
        zona = self._etz_zonas[idx]

        try:
            import os
            import tempfile

            from core.deepfont import clasificar_tipografia

            # Recortar la zona de la imagen original
            iw = self._etz_img_orig.width
            ih = self._etz_img_orig.height
            box = (int(zona.x0 * iw), int(zona.y0 * ih),
                   int(zona.x1 * iw), int(zona.y1 * ih))
            if box[2] <= box[0] or box[3] <= box[1]:
                messagebox.showwarning("Zona inválida",
                    "La zona no tiene dimensiones suficientes."); return

            recorte = self._etz_img_orig.crop(box)
            # Guardar en temp ASCII para evitar problemas con FAISS/Windows
            tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False,
                                              dir="C:\\Windows\\Temp")
            tmp.close()
            recorte.save(tmp.name)

            self._etz_lbl_train.config(text="Analizando tipografía…")

            def _worker():
                try:
                    res = clasificar_tipografia(tmp.name, usar_clip=True)
                finally:
                    try: os.unlink(tmp.name)
                    except OSError: pass
                self.after(0, lambda: self._etz_mostrar_deepfont(res, zona))

            threading.Thread(target=_worker, daemon=True).start()

        except Exception as e:
            messagebox.showerror("Error", f"No se pudo analizar: {e}")

    def _etz_mostrar_deepfont(self, resultado: dict, zona):
        """Muestra el resultado de DeepFont en una ventana."""
        from core.deepfont import COLORES, ETIQUETAS_ES
        cat   = resultado["categoria"]
        eta   = resultado["etiqueta"]
        conf  = resultado["confianza"]
        color = resultado.get("color", "#888")
        metodo = resultado.get("metodo", "?")

        self._etz_lbl_train.config(text=f"🔤 {eta[:30]} ({conf:.0%})")

        win = tk.Toplevel(self)
        win.title("Análisis tipográfico — DeepFont")
        win.geometry("440x300")
        win.configure(bg=CONTENT_BG)

        # Resultado principal
        res_frm = tk.Frame(win, bg=color, pady=12)
        res_frm.pack(fill="x")
        tk.Label(res_frm, text=eta, bg=color, fg="white",
                 font=("Segoe UI", 13, "bold")).pack()
        tk.Label(res_frm, text=f"Confianza: {conf:.0%}  ·  Método: {metodo}",
                 bg=color, fg="white",
                 font=("Segoe UI", 9)).pack()

        # Scores por categoría
        scores = resultado.get("scores", {})
        if isinstance(scores, dict) and scores and all(
                isinstance(v, float) for v in scores.values()):
            tk.Label(win, text="Distribución de probabilidades:",
                     bg=CONTENT_BG, fg=TXT_SEC,
                     font=("Segoe UI", 8, "bold")).pack(anchor="w", padx=16, pady=(12, 4))
            from core.deepfont import CATEGORIAS_TIPOGRAFIA
            for c in CATEGORIAS_TIPOGRAFIA:
                pct = scores.get(c, 0.0)
                row = tk.Frame(win, bg=CONTENT_BG)
                row.pack(fill="x", padx=16, pady=1)
                tk.Label(row, text=ETIQUETAS_ES.get(c, c)[:35],
                         bg=CONTENT_BG, fg=TXT_PRI if c == cat else TXT_SEC,
                         font=("Segoe UI", 8,
                               "bold" if c == cat else "normal"),
                         width=38, anchor="w").pack(side="left")
                bar_frm = tk.Frame(row, bg=CARD_BOR, width=160, height=10)
                bar_frm.pack(side="left", padx=4)
                bar_frm.pack_propagate(False)
                tk.Frame(bar_frm, bg=COLORES.get(c, "#888"),
                         width=int(pct * 160), height=10).pack(side="left")
                tk.Label(row, text=f"{pct:.0%}", bg=CONTENT_BG, fg=TXT_SEC,
                         font=("Segoe UI", 7)).pack(side="left", padx=2)

        ttk.Button(win, text="Cerrar", style="S.TButton",
                   command=win.destroy).pack(pady=12)

    # ── Zoom Pane ─────────────────────────────────────────────────────────────
    def _etz_actualizar_zoom_pane(self, cx: float, cy: float):
        """Muestra un recorte ampliado 4x alrededor del cursor en el Zoom Pane."""
        if self._etz_img_orig is None:
            zc = getattr(self, "_etz_zoom_canvas", None)
            if zc:
                zc.delete("all")
                zc.create_text(80, 40, text="Carga una página para ver el detalle",
                               fill=TXT_DIM, font=("Segoe UI", 8), anchor="w")
            return
        try:
            from PIL import Image, ImageTk
            # Convertir coords canvas a coords imagen original
            escala = self._etz_escala * self._etz_zoom
            if escala <= 0:
                return
            ix = cx / escala
            iy = cy / escala
            radio = 40  # radio en px de imagen original
            box = (max(0, int(ix - radio)), max(0, int(iy - radio)),
                   min(self._etz_img_orig.width,  int(ix + radio)),
                   min(self._etz_img_orig.height, int(iy + radio)))
            if box[2] <= box[0] or box[3] <= box[1]:
                return
            recorte = self._etz_img_orig.crop(box)
            # Ampliar 3x
            nw = min(recorte.width  * 3, 400)
            nh = min(recorte.height * 3, 200)
            if nw < 1 or nh < 1:
                return
            ampliado = recorte.resize((nw, nh), Image.NEAREST)
            self._etz_zoom_img_tk = ImageTk.PhotoImage(ampliado)
            zc = self._etz_zoom_canvas
            zc.delete("all")
            zc.create_image(0, 0, anchor="nw", image=self._etz_zoom_img_tk)
            # Cruz central
            mx, my = nw // 2, nh // 2
            zc.create_line(mx - 8, my, mx + 8, my, fill="#F85149", width=1)
            zc.create_line(mx, my - 8, mx, my + 8, fill="#F85149", width=1)
        except Exception:
            pass

    # ── Miniaturas de páginas (Pages Pane) ────────────────────────────────────
    def _etz_poblar_miniaturas(self, pags: list[str]):
        """Crea/actualiza las miniaturas del Pages Pane."""
        for w in self._etz_pages_inner.winfo_children():
            w.destroy()
        self._etz_thumb_btns.clear()

        for pag in pags:
            frm = tk.Frame(self._etz_pages_inner, bg="#161B22",
                           cursor="hand2", pady=2)
            frm.pack(fill="x", padx=4, pady=1)

            # Número de página
            lbl = tk.Label(frm, text=pag, bg="#161B22", fg=TXT_SEC,
                           font=("Segoe UI", 7), anchor="w")
            lbl.pack(fill="x", padx=4)

            # Indicador de estado (etiquetada / sin etiquetar)
            from core.zone_labeler import cargar_pagina
            num = self._etz_numero.get()
            tiene_zonas = False
            if ST.out_dir and num:
                pd = cargar_pagina(ST.out_dir, num, pag)
                tiene_zonas = bool(pd and pd.zonas)

            dot_color = "#3FB950" if tiene_zonas else TXT_DIM
            dot = tk.Label(frm, text="●", bg="#161B22", fg=dot_color,
                           font=("Segoe UI", 7))
            dot.pack(side="right", padx=4)

            # Click navega a esa página
            def _goto(p=pag):
                self._etz_cb_pag.set(p)
                self._etz_pagina.set(p)
                self._etz_on_pagina()
            for w in (frm, lbl, dot):
                w.bind("<Button-1>", lambda e, fn=_goto: fn())
                w.bind("<Enter>",    lambda e, f=frm: f.config(bg="#1C2128"))
                w.bind("<Leave>",    lambda e, f=frm: f.config(bg="#161B22"))

            self._etz_thumb_btns[pag] = frm

        self._etz_pages_canvas.update_idletasks()
        self._etz_pages_canvas.configure(
            scrollregion=self._etz_pages_canvas.bbox("all"))

    def _etz_resaltar_miniatura(self, pag: str):
        """Resalta la miniatura de la página activa."""
        for p, frm in self._etz_thumb_btns.items():
            frm.config(bg="#1F6FEB" if p == pag else "#161B22")
            for child in frm.winfo_children():
                child.config(bg="#1F6FEB" if p == pag else "#161B22")

    # ── Carga de texto OCR en el Text Pane ────────────────────────────────────
    def _etz_cargar_texto_ocr(self, numero: str, pagina: str):
        """Carga el texto OCR de la página en el Text Pane."""
        self._etz_txt_ocr.config(state="normal")
        self._etz_txt_ocr.delete("1.0", "end")
        if not ST.out_dir:
            return
        txt_path = ST.out_dir / "03_ocr" / numero / f"{pagina}.txt"
        if txt_path.exists():
            texto = txt_path.read_text(encoding="utf-8", errors="replace")
            self._etz_txt_ocr.insert("1.0", texto)
        else:
            self._etz_txt_ocr.insert("1.0",
                f"(Sin texto OCR para {pagina})\n"
                "Extrae el texto en la pestaña OCR primero.")
            self._etz_txt_ocr.tag_add("low_conf", "1.0", "end")

    def _etz_on_numero(self, event=None):
        numero = self._etz_numero.get()
        if not numero:
            return
        # Listar páginas disponibles como .txt
        txt_dir = ST.out_dir / "03_ocr" / numero if ST.out_dir else None
        if txt_dir and txt_dir.exists():
            pags = sorted(p.stem for p in txt_dir.glob("*.txt"))
        else:
            pags = [f"p{i:04d}" for i in range(1, 50)]
        self._etz_cb_pag["values"] = pags
        # Poblar panel de miniaturas
        self._etz_poblar_miniaturas(pags)
        if pags:
            self._etz_cb_pag.set(pags[0])
            self._etz_on_pagina()
        self._etz_actualizar_estado()
        # Re-entrenar detector con las etiquetas del número seleccionado
        self._etz_entrenar_detector(numero)

    def _etz_on_pagina(self, event=None):
        numero  = self._etz_numero.get()
        pagina  = self._etz_pagina.get()
        if not numero or not pagina:
            return

        # Cargar zonas guardadas para esta página
        from core.zone_labeler import cargar_pagina
        if ST.out_dir:
            pag_data = cargar_pagina(ST.out_dir, numero, pagina)
            self._etz_zonas = list(pag_data.zonas) if pag_data else []
        else:
            self._etz_zonas = []

        self._etz_actualizar_lista_zonas()

        # Intentar cargar imagen: primero via caché/PDF del corpus,
        # luego directamente desde PDF si se cargó con "Abrir PDF"
        cargado = False
        if ST.archivos_sel or (ST.out_dir and ST.pdf_dir):
            self._etz_cargar_imagen_pagina(numero, pagina)
            cargado = self._etz_img_orig is not None

        if not cargado:
            # Buscar el PDF directamente en archivos_sel por stem
            pdf = next((a for a in (ST.archivos_sel or [])
                        if a.stem == numero), None)
            if pdf:
                import re as _re
                _nums = _re.findall(r'\d+', pagina)
                n_pag = max(0, int(_nums[-1]) - 1) if _nums else 0
                self._etz_cargar_imagen_desde_pdf(str(pdf), n_pag)

        # Sincronizar Text Pane y Pages Pane
        self._etz_cargar_texto_ocr(numero, pagina)
        self._etz_resaltar_miniatura(pagina)

    def _etz_pagina_ant(self):
        vals = list(self._etz_cb_pag["values"])
        if not vals:
            return
        cur = self._etz_pagina.get()
        idx = vals.index(cur) if cur in vals else 0
        if idx > 0:
            self._etz_cb_pag.set(vals[idx - 1])
            self._etz_on_pagina()

    def _etz_pagina_sig(self):
        vals = list(self._etz_cb_pag["values"])
        if not vals:
            return
        cur = self._etz_pagina.get()
        idx = vals.index(cur) if cur in vals else 0
        if idx < len(vals) - 1:
            self._etz_cb_pag.set(vals[idx + 1])
            self._etz_on_pagina()

    # ── Helpers de detección de handles ──────────────────────────────────────

    _ETZ_HANDLE_R = 7   # radio de detección de esquinas/bordes (px)
    _ETZ_HANDLE_D = 7   # radio visual de los cuadraditos dibujados (px)

    def _etz_canvas_wh(self):
        """Ancho y alto de la imagen en coordenadas canvas (imagen ya escalada)."""
        if self._etz_img_orig is None:
            return 1, 1
        return self._etz_img_orig.width, self._etz_img_orig.height

    def _etz_zona_canvas_coords(self, idx):
        """Devuelve (x0,y0,x1,y1) en coords canvas de la zona[idx]."""
        z = self._etz_zonas[idx]
        w, h = self._etz_canvas_wh()
        return z.x0 * w, z.y0 * h, z.x1 * w, z.y1 * h

    def _etz_canvas_a_norm(self, cx, cy):
        """Convierte coordenadas canvas a normalizado [0,1]."""
        w, h = self._etz_canvas_wh()
        return max(0.0, min(1.0, cx / w)), max(0.0, min(1.0, cy / h))

    def _etz_hit_handle(self, cx, cy):
        """Devuelve (idx, handle) si cx,cy está sobre esquina/borde/interior
        de alguna zona, o None si está en espacio vacío.
        handle: 'nw','n','ne','e','se','s','sw','w','move'
        """
        if not self._etz_zonas or self._etz_img_orig is None:
            return None
        R = self._ETZ_HANDLE_R
        # Última zona dibujada tiene prioridad
        for i in range(len(self._etz_zonas) - 1, -1, -1):
            x0, y0, x1, y1 = self._etz_zona_canvas_coords(i)
            # Normalizar por si x0>x1 o y0>y1
            lx, rx = min(x0, x1), max(x0, x1)
            ty, by = min(y0, y1), max(y0, y1)
            # Esquinas
            if abs(cx - lx) <= R and abs(cy - ty) <= R: return (i, "nw")
            if abs(cx - rx) <= R and abs(cy - ty) <= R: return (i, "ne")
            if abs(cx - lx) <= R and abs(cy - by) <= R: return (i, "sw")
            if abs(cx - rx) <= R and abs(cy - by) <= R: return (i, "se")
            # Bordes medios
            if abs(cx - lx) <= R and ty <= cy <= by:    return (i, "w")
            if abs(cx - rx) <= R and ty <= cy <= by:    return (i, "e")
            if abs(cy - ty) <= R and lx <= cx <= rx:    return (i, "n")
            if abs(cy - by) <= R and lx <= cx <= rx:    return (i, "s")
            # Interior → mover
            if lx <= cx <= rx and ty <= cy <= by:       return (i, "move")
        return None

    _ETZ_CURSOR_MAP = {
        "nw": "size_nw_se", "se": "size_nw_se",
        "ne": "size_ne_sw", "sw": "size_ne_sw",
        "n":  "sb_v_double_arrow", "s": "sb_v_double_arrow",
        "e":  "sb_h_double_arrow", "w": "sb_h_double_arrow",
        "move": "fleur",
    }

    def _etz_on_motion(self, event):
        """Cambia cursor, muestra coordenadas en statusbar y actualiza Zoom Pane."""
        if self._etz_img_orig is None:
            return
        cx = self._etz_canvas.canvasx(event.x)
        cy = self._etz_canvas.canvasy(event.y)
        hit = self._etz_hit_handle(cx, cy)
        cursor = self._ETZ_CURSOR_MAP.get(hit[1], "crosshair") if hit else "crosshair"
        self._etz_canvas.config(cursor=cursor)

        # Status bar — coordenadas normalizadas
        if self._etz_img_orig:
            w, h = self._etz_canvas_wh()
            nx = round(cx / w, 3) if w > 0 else 0
            ny = round(cy / h, 3) if h > 0 else 0
            if hasattr(self, "_etz_lbl_coords"):
                self._etz_lbl_coords.config(text=f"x:{nx:.3f}  y:{ny:.3f}")

        # Zoom Pane — muestra recorte ampliado alrededor del cursor
        self._etz_actualizar_zoom_pane(cx, cy)

    # ── Zoom y Pan ───────────────────────────────────────────────────────────

    def _etz_on_zoom(self, event):
        """Zoom con Ctrl+rueda. Reescala la imagen y redibuja zonas."""
        if self._etz_img_orig is None:
            return
        factor = 1.15 if event.delta > 0 else 1 / 1.15
        nuevo_zoom = max(0.2, min(5.0, self._etz_zoom * factor))
        if nuevo_zoom == self._etz_zoom:
            return
        self._etz_zoom = nuevo_zoom
        self._etz_aplicar_zoom()

    def _etz_aplicar_zoom(self):
        """Redimensiona la imagen según _etz_zoom y redibuja todo."""
        if self._etz_img_orig is None:
            return
        from PIL import Image, ImageTk
        base = getattr(self, "_etz_img_orig_full", self._etz_img_orig)
        new_w = max(1, int(base.width  * self._etz_zoom))
        new_h = max(1, int(base.height * self._etz_zoom))
        img_zoom = base.resize((new_w, new_h), Image.LANCZOS)
        self._etz_img_orig = img_zoom
        self._etz_img_tk   = ImageTk.PhotoImage(img_zoom)
        self._etz_canvas.delete("all")
        self._etz_canvas.create_image(0, 0, anchor="nw", image=self._etz_img_tk)
        self._etz_canvas.configure(scrollregion=(0, 0, new_w, new_h))
        self._etz_redibujar_zonas()

    def _etz_pan_start_cb(self, event):
        self._etz_pan_start = (event.x, event.y)
        self._etz_canvas.config(cursor="fleur")

    def _etz_pan_drag_cb(self, event):
        if self._etz_pan_start is None:
            return
        dx = self._etz_pan_start[0] - event.x
        dy = self._etz_pan_start[1] - event.y
        self._etz_pan_start = (event.x, event.y)
        self._etz_canvas.xview_scroll(int(dx), "units")
        self._etz_canvas.yview_scroll(int(dy), "units")

    def _etz_pan_end_cb(self, event):
        self._etz_pan_start = None
        self._etz_canvas.config(cursor="crosshair")

    def _etz_space_press(self, event):
        self._etz_space_held = True
        self._etz_canvas.config(cursor="fleur")

    def _etz_space_release(self, event):
        self._etz_space_held = False
        self._etz_canvas.config(cursor="crosshair")

    # ── Interacción principal ─────────────────────────────────────────────────

    def _etz_on_press(self, event):
        cx = self._etz_canvas.canvasx(event.x)
        cy = self._etz_canvas.canvasy(event.y)

        # Space presionado → pan con botón izquierdo
        if self._etz_space_held:
            self._etz_modo = "pan"
            self._etz_pan_start = (event.x, event.y)
            return

        if self._etz_rect_tmp:
            self._etz_canvas.delete(self._etz_rect_tmp)
            self._etz_rect_tmp = None

        hit = self._etz_hit_handle(cx, cy)
        if hit:
            idx, handle = hit
            self._etz_zona_sel_idx = idx
            if handle == "move":
                self._etz_modo = "move"
                self._etz_resize_idx = idx
                x0, y0, x1, y1 = self._etz_zona_canvas_coords(idx)
                # offset en coordenadas canvas (imagen ya escalada)
                self._etz_move_offset = (cx - min(x0, x1), cy - min(y0, y1))
            else:
                self._etz_modo = "resize"
                self._etz_resize_idx = idx
                self._etz_resize_handle = handle
            self._etz_redibujar_zonas()
        else:
            self._etz_zona_sel_idx = None
            self._etz_modo = "draw"
            self._etz_rect_ini = (cx, cy)

    def _etz_on_drag(self, event):
        if self._etz_modo == "pan" and self._etz_pan_start:
            dx = self._etz_pan_start[0] - event.x
            dy = self._etz_pan_start[1] - event.y
            self._etz_pan_start = (event.x, event.y)
            self._etz_canvas.xview_scroll(int(dx), "units")
            self._etz_canvas.yview_scroll(int(dy), "units")
            return

        cx = self._etz_canvas.canvasx(event.x)
        cy = self._etz_canvas.canvasy(event.y)

        if self._etz_modo == "draw":
            self._etz_drag_draw(cx, cy)
        elif self._etz_modo == "resize":
            self._etz_drag_resize(cx, cy)
        elif self._etz_modo == "move":
            self._etz_drag_move(cx, cy)

    def _etz_drag_draw(self, cx, cy):
        """Dibuja un nuevo rectángulo desde _etz_rect_ini hasta cx,cy."""
        if not self._etz_rect_ini:
            return
        from core.zone_labeler import TIPOS_ZONA
        x0, y0 = self._etz_rect_ini
        color = TIPOS_ZONA.get(self._etz_tipo.get(), {}).get("color", "#888")
        label = TIPOS_ZONA.get(self._etz_tipo.get(), {}).get("label", "")
        if self._etz_rect_tmp:
            self._etz_canvas.delete(self._etz_rect_tmp)
        self._etz_rect_tmp = self._etz_canvas.create_rectangle(
            x0, y0, cx, cy, outline=color, width=2, dash=(6, 3), fill=""
        )
        escala = self._etz_escala if self._etz_escala > 0 else 1.0
        pct_w = abs(cx - x0) / (self._etz_img_orig.width * escala) * 100 if self._etz_img_orig else 0
        pct_h = abs(cy - y0) / (self._etz_img_orig.height * escala) * 100 if self._etz_img_orig else 0
        dim_txt = f"{label}  {pct_w:.0f}% × {pct_h:.0f}%"
        lx = min(x0, cx) + 4
        ly = min(y0, cy) - 14 if min(y0, cy) > 20 else min(y0, cy) + 4
        if not hasattr(self, "_etz_dim_lbl"):
            self._etz_dim_lbl = None
        if self._etz_dim_lbl:
            self._etz_canvas.delete(self._etz_dim_lbl)
        self._etz_dim_lbl = self._etz_canvas.create_text(
            lx, ly, text=dim_txt, anchor="nw",
            fill="white", font=("Segoe UI", 8, "bold"), tags="dim_label"
        )

    def _etz_drag_resize(self, cx, cy):
        """Mueve la esquina/borde indicado por _etz_resize_handle."""
        idx = self._etz_resize_idx
        if idx is None or idx >= len(self._etz_zonas):
            return
        nx, ny = self._etz_canvas_a_norm(cx, cy)
        z = self._etz_zonas[idx]
        h = self._etz_resize_handle
        if "w" in h: z.x0 = min(nx, z.x1 - 0.01)
        if "e" in h: z.x1 = max(nx, z.x0 + 0.01)
        if "n" in h: z.y0 = min(ny, z.y1 - 0.01)
        if "s" in h: z.y1 = max(ny, z.y0 + 0.01)
        self._etz_redibujar_zonas()
        self._etz_actualizar_lista_zonas()

    def _etz_drag_move(self, cx, cy):
        """Mueve toda la zona manteniendo su tamaño."""
        idx = self._etz_resize_idx
        if idx is None or idx >= len(self._etz_zonas) or self._etz_move_offset is None:
            return
        w, h = self._etz_canvas_wh()
        z = self._etz_zonas[idx]
        dw = z.x1 - z.x0
        dh = z.y1 - z.y0
        ox, oy = self._etz_move_offset
        new_x0 = max(0.0, min(1.0 - dw, (cx - ox) / w))
        new_y0 = max(0.0, min(1.0 - dh, (cy - oy) / h))
        z.x0 = new_x0
        z.x1 = new_x0 + dw
        z.y0 = new_y0
        z.y1 = new_y0 + dh
        self._etz_redibujar_zonas()
        self._etz_actualizar_lista_zonas()

    def _etz_on_release(self, event):
        cx = self._etz_canvas.canvasx(event.x)
        cy = self._etz_canvas.canvasy(event.y)

        if self._etz_rect_tmp:
            self._etz_canvas.delete(self._etz_rect_tmp)
            self._etz_rect_tmp = None
        if hasattr(self, "_etz_dim_lbl") and self._etz_dim_lbl:
            self._etz_canvas.delete(self._etz_dim_lbl)
            self._etz_dim_lbl = None

        if self._etz_modo == "draw":
            self._etz_finish_draw(cx, cy)
        # resize y move ya actualizaron la zona en tiempo real, nada más que hacer

        self._etz_modo = None
        self._etz_rect_ini = None
        self._etz_resize_idx = None
        self._etz_resize_handle = None
        self._etz_move_offset = None

    def _etz_finish_draw(self, cx, cy):
        """Guarda el rectángulo nuevo al soltar el mouse."""
        if not self._etz_rect_ini or self._etz_img_orig is None:
            return
        from core.zone_labeler import Zona
        x0, y0 = self._etz_rect_ini
        nx0, ny0 = self._etz_canvas_a_norm(x0, y0)
        nx1, ny1 = self._etz_canvas_a_norm(cx, cy)
        if abs(nx1 - nx0) < 0.01 or abs(ny1 - ny0) < 0.01:
            return
        zona = Zona(
            tipo=self._etz_tipo.get(),
            x0=min(nx0, nx1), y0=min(ny0, ny1),
            x1=max(nx0, nx1), y1=max(ny0, ny1),
            confianza=1.0,
        )
        self._etz_zonas.append(zona)
        self._etz_zona_sel_idx = len(self._etz_zonas) - 1
        self._etz_actualizar_lista_zonas()
        self._etz_zona_list.selection_clear(0, "end")
        self._etz_zona_list.selection_set(self._etz_zona_sel_idx)
        self._etz_redibujar_zonas()

    def _etz_on_click_derecho(self, event):
        """Click derecho sobre el canvas: elimina la zona bajo el cursor,
        o muestra menú para cambiar el tipo si hay una zona allí."""
        if self._etz_img_orig is None or not self._etz_zonas:
            return
        cx = self._etz_canvas.canvasx(event.x)
        cy = self._etz_canvas.canvasy(event.y)
        nx, ny = self._etz_canvas_a_norm(cx, cy)

        # Buscar zona que contiene el punto (última dibujada tiene prioridad)
        idx_hit = None
        for i in range(len(self._etz_zonas) - 1, -1, -1):
            z = self._etz_zonas[i]
            if z.x0 <= nx <= z.x1 and z.y0 <= ny <= z.y1:
                idx_hit = i
                break

        if idx_hit is None:
            return

        # Menú contextual
        from core.zone_labeler import TIPOS_ZONA
        menu = tk.Menu(self, tearoff=0)
        zona_hit = self._etz_zonas[idx_hit]

        # Submenú: cambiar tipo
        sub = tk.Menu(menu, tearoff=0)
        for tipo_key, info in TIPOS_ZONA.items():
            def _set_tipo(k=tipo_key, i=idx_hit):
                self._etz_zonas[i].tipo = k
                self._etz_actualizar_lista_zonas()
                self._etz_redibujar_zonas()
            marca = "✓ " if tipo_key == zona_hit.tipo else "   "
            sub.add_command(label=f"{marca}{info['label']}", command=_set_tipo)
        menu.add_cascade(label="Cambiar tipo", menu=sub)
        menu.add_separator()

        # Dividir la zona exactamente donde se hizo clic (estilo FineReader)
        frac_h = (ny - zona_hit.y0) / max(zona_hit.y1 - zona_hit.y0, 1e-6)
        frac_v = (nx - zona_hit.x0) / max(zona_hit.x1 - zona_hit.x0, 1e-6)
        menu.add_command(
            label="✂ Dividir aquí (horizontal)",
            command=lambda i=idx_hit, f=frac_h: self._etz_dividir_zona(i, "h", f))
        menu.add_command(
            label="✂ Dividir aquí (vertical)",
            command=lambda i=idx_hit, f=frac_v: self._etz_dividir_zona(i, "v", f))

        # Fusionar con otra zona del mismo tipo o cercana
        sub_fus = tk.Menu(menu, tearoff=0)
        n_fus = 0
        for j, otra in enumerate(self._etz_zonas):
            if j == idx_hit or n_fus >= 15:
                continue
            lbl_o = TIPOS_ZONA.get(otra.tipo, {}).get("label", otra.tipo)
            sub_fus.add_command(
                label=f"{j+1}. {lbl_o} ({otra.x0:.2f},{otra.y0:.2f})",
                command=lambda a=idx_hit, b=j: self._etz_fusionar_zonas(a, b))
            n_fus += 1
        if n_fus:
            menu.add_cascade(label="🔗 Fusionar con…", menu=sub_fus)

        # Vincular pie de foto ↔ foto (Zona.vinculo, por identidad zid estable)
        if zona_hit.tipo == "pie_foto":
            fotos_disp = [(j, o) for j, o in enumerate(self._etz_zonas) if o.tipo == "foto"]
            if fotos_disp:
                sub_vinc = tk.Menu(menu, tearoff=0)
                for j, foto in fotos_disp:
                    marca = "✓ " if zona_hit.vinculo == foto.zid else "   "
                    sub_vinc.add_command(
                        label=f"{marca}Foto {j+1} ({foto.x0:.2f},{foto.y0:.2f})",
                        command=lambda i=idx_hit, fz=foto.zid: self._etz_vincular_foto(i, fz))
                menu.add_cascade(label="🖇 Vincular a foto…", menu=sub_vinc)

        menu.add_command(label="🔢 Recalcular orden de lectura",
                         command=self._etz_recalcular_orden)
        menu.add_separator()
        menu.add_command(label="🗑 Eliminar esta zona",
                         command=lambda i=idx_hit: self._etz_borrar_zona(i))
        try:
            menu.tk_popup(event.x_root, event.y_root)
        finally:
            menu.grab_release()

    def _etz_dividir_zona(self, idx: int, eje: str, frac: float):
        """Divide la zona idx en dos por el eje indicado, en la posición frac."""
        from core.zone_labeler import dividir_zona
        if not (0 <= idx < len(self._etz_zonas)):
            return
        a, b = dividir_zona(self._etz_zonas[idx], eje=eje, frac=frac)
        self._etz_zonas[idx:idx + 1] = [a, b]
        self._etz_limpiar_vinculos_huerfanos()
        self._etz_recalcular_orden(redibujar=False)
        self._etz_zona_sel_idx = idx
        self._etz_actualizar_lista_zonas()
        self._etz_redibujar_zonas()

    def _etz_fusionar_zonas(self, idx_a: int, idx_b: int):
        """Fusiona dos zonas en su bounding box común."""
        from core.zone_labeler import fusionar_zonas
        n = len(self._etz_zonas)
        if not (0 <= idx_a < n and 0 <= idx_b < n) or idx_a == idx_b:
            return
        fusion = fusionar_zonas([self._etz_zonas[idx_a], self._etz_zonas[idx_b]])
        for i in sorted((idx_a, idx_b), reverse=True):
            self._etz_zonas.pop(i)
        self._etz_zonas.append(fusion)
        self._etz_limpiar_vinculos_huerfanos()
        self._etz_recalcular_orden(redibujar=False)
        self._etz_zona_sel_idx = len(self._etz_zonas) - 1
        self._etz_actualizar_lista_zonas()
        self._etz_redibujar_zonas()

    def _etz_limpiar_vinculos_huerfanos(self):
        """Pone en None cualquier Zona.vinculo que apunte a un zid que ya no
        existe (tras dividir/fusionar/borrar la zona vinculada)."""
        zids_vivos = {z.zid for z in self._etz_zonas}
        for z in self._etz_zonas:
            if z.vinculo and z.vinculo not in zids_vivos:
                z.vinculo = None

    def _etz_vincular_foto(self, idx_pie: int, foto_zid: str):
        """Vincula (o desvincula si ya estaba vinculado a esa misma foto) la
        zona pie_foto en idx_pie con la zona foto de zid `foto_zid`."""
        if not (0 <= idx_pie < len(self._etz_zonas)):
            return
        pie = self._etz_zonas[idx_pie]
        pie.vinculo = None if pie.vinculo == foto_zid else foto_zid
        self._etz_redibujar_zonas()

    def _etz_detectar_cabeceras(self):
        """Reporta cabeceras repetidas entre las páginas YA etiquetadas
        manualmente de este número (solo informa; no reetiqueta nada solo).
        """
        numero = self._etz_numero.get() if hasattr(self, "_etz_numero") else ""
        if not numero or not ST.out_dir:
            self.toast("Selecciona un número con páginas etiquetadas primero.", "warn")
            return
        from core.layout_patterns import detectar_cabeceras_repetidas
        from core.zone_labeler import cargar_todas_manual
        paginas_et = cargar_todas_manual(Path(ST.out_dir), numero)
        if len(paginas_et) < 2:
            self.toast("Se necesitan al menos 2 páginas etiquetadas de este "
                       "número para comparar cabeceras.", "warn")
            return
        zonas_por_pagina = [pg.zonas for pg in paginas_et]
        grupos = detectar_cabeceras_repetidas(zonas_por_pagina, min_repeticiones=2)
        if not grupos:
            messagebox.showinfo("Cabeceras repetidas",
                "No se detectaron cabeceras con la misma posición en al menos "
                "2 páginas etiquetadas de este número.")
            return
        lineas = []
        for i, indices_pag in enumerate(grupos.values(), 1):
            nombres = ", ".join(paginas_et[p].pagina for p in indices_pag)
            lineas.append(f"Grupo {i}: {len(indices_pag)} páginas → {nombres}")
        messagebox.showinfo("Cabeceras repetidas",
            f"{len(grupos)} patrón(es) de cabecera detectado(s):\n\n" + "\n".join(lineas))

    def _etz_recalcular_orden(self, redibujar: bool = True):
        """Recalcula el orden de lectura de todas las zonas de la página."""
        from core.zone_labeler import calcular_orden_lectura
        calcular_orden_lectura(self._etz_zonas)
        if redibujar:
            self._etz_actualizar_lista_zonas()
            self._etz_redibujar_zonas()

    def _etz_ocr_zonas_preview(self):
        """OCR por zonas de la página actual: reconoce cada zona recortada
        por separado, en orden de lectura, y muestra el resultado en el
        panel TEXTO OCR. Es el flujo central de FineReader."""
        numero = self._etz_numero.get()
        pagina = self._etz_pagina.get()
        if not numero or not pagina:
            messagebox.showwarning("Sin selección", "Selecciona un número y página.")
            return
        if not self._etz_zonas:
            messagebox.showwarning(
                "Sin zonas",
                "No hay zonas en esta página.\n"
                "Usa 'Detectar → Esta página' o dibuja zonas manualmente.")
            return
        img_path = self._etz_get_img_path(numero, pagina)
        if not img_path:
            messagebox.showwarning("Sin imagen",
                "No se encontró la imagen de esta página.")
            return

        zonas = list(self._etz_zonas)
        self._etz_lbl_train.config(text="⏳ OCR por zonas…", fg="#F59E0B")

        def _worker():
            try:
                from core.layout_tesseract import ocr_por_zonas
                def _cb(m):
                    self.after(0, lambda m=m: self._etz_lbl_train.config(
                        text=m[:80], fg="#F59E0B"))
                res = ocr_por_zonas(img_path, zonas, callback=_cb)
                err = ""
            except Exception as e:
                res, err = None, str(e)

            def _mostrar():
                if err or res is None:
                    self._etz_lbl_train.config(
                        text=f"⚠ Error OCR zonas: {err[:60]}", fg="#EF4444")
                    return
                from core.zone_labeler import TIPOS_ZONA as _TZ
                self._etz_txt_ocr.delete("1.0", "end")
                for rz in res["zonas"]:
                    lbl = _TZ.get(rz["tipo"], {}).get("label", rz["tipo"])
                    self._etz_txt_ocr.insert(
                        "end",
                        f"━━ #{rz['orden']} {lbl} "
                        f"(conf {rz['confianza']:.0f}) ━━\n", "low_conf")
                    self._etz_txt_ocr.insert("end", (rz["texto"] or "—") + "\n\n")
                n_pal = len(res["texto"].split())
                self._etz_lbl_train.config(
                    text=f"✅ OCR zonal: {len(res['zonas'])} zonas, "
                         f"{n_pal} palabras, conf {res['confianza']:.0f}",
                    fg="#22C55E")
                self.toast(f"OCR por zonas: {n_pal} palabras", tipo="ok")
            self.after(0, _mostrar)

        threading.Thread(target=_worker, daemon=True).start()

    def _etz_borrar_zona(self, idx: int):
        """Elimina la zona en la posición idx."""
        if 0 <= idx < len(self._etz_zonas):
            self._etz_zonas.pop(idx)
            self._etz_limpiar_vinculos_huerfanos()
            self._etz_zona_sel_idx = None
            self._etz_actualizar_lista_zonas()
            self._etz_redibujar_zonas()

    def _etz_suprimir_sel(self, event=None):
        """Elimina la zona seleccionada al presionar Suprimir o Retroceso."""
        idx = self._etz_zona_sel_idx
        if idx is not None and 0 <= idx < len(self._etz_zonas):
            self._etz_borrar_zona(idx)

    def _etz_on_zona_sel(self, event=None):
        """Activa la zona seleccionada en la lista (muestra handles)."""
        sel = self._etz_zona_list.curselection()
        if not sel or not self._etz_img_orig:
            self._etz_zona_sel_idx = None
            self._etz_redibujar_zonas()
            return
        idx = sel[0]
        if idx >= len(self._etz_zonas):
            return
        self._etz_zona_sel_idx = idx
        self._etz_redibujar_zonas()
        # Hacer scroll para que la zona sea visible
        x0p, y0p, x1p, y1p = self._etz_zona_canvas_coords(idx)
        _, h = self._etz_canvas_wh()
        frac = min(y0p, y1p) / max(h, 1)
        self._etz_canvas.yview_moveto(max(0.0, frac - 0.1))

    def _etz_borrar_ultima(self):
        if self._etz_zonas:
            self._etz_zonas.pop()
            self._etz_actualizar_lista_zonas()
            self._etz_redibujar_zonas()

    def _etz_limpiar_todo(self):
        self._etz_zonas.clear()
        self._etz_actualizar_lista_zonas()
        self._etz_redibujar_zonas()

    def _etz_guardar_pagina(self):
        from core.zone_labeler import PaginaEtiquetada, guardar_pagina
        numero  = self._etz_numero.get()
        pagina  = self._etz_pagina.get()
        if not numero or not pagina or not ST.out_dir:
            messagebox.showwarning("Sin datos", "Selecciona un número y página primero.")
            return
        ancho = self._etz_img_orig.width if self._etz_img_orig else 1000
        alto  = self._etz_img_orig.height if self._etz_img_orig else 1400
        pag_data = PaginaEtiquetada(
            pagina=pagina,
            ancho_px=ancho,
            alto_px=alto,
            zonas=list(self._etz_zonas),
            manual=True,
        )
        guardar_pagina(ST.out_dir, numero, pag_data)
        # Re-entrenar el detector con todas las etiquetas manuales disponibles
        self._etz_entrenar_detector(numero)
        self.toast(f"Etiquetas de {pagina} guardadas", tipo="ok")
        self._etz_actualizar_estado()

    def _etz_entrenar_detector(self, numero: str):
        """Entrena DetectorZonas con todas las páginas etiquetadas manualmente del número."""
        from core.zone_labeler import DetectorZonas, cargar_todas_manual
        manuales = cargar_todas_manual(ST.out_dir, numero)  # list[PaginaEtiquetada]
        if not manuales:
            return
        if not hasattr(self, "_etz_detector"):
            self._etz_detector = DetectorZonas()
        stats = self._etz_detector.entrenar(manuales)       # acepta lista directamente
        n     = stats.get("n_paginas", 0)
        tipos = list(stats.get("tipos", {}).keys())
        self._etz_lbl_train.config(
            text=f"🧠 Modelo: {n} pág · {len(tipos)} tipos aprendidos",
            fg="#3FB950"
        )
        self._btn_etz_predecir.config(bg="#6E40C9")

    def _etz_predecir_numero(self):
        """Aplica la plantilla aprendida a todas las páginas sin etiquetar del número."""
        numero = self._etz_numero.get()
        if not numero or not ST.out_dir:
            return
        if not hasattr(self, "_etz_detector") or not self._etz_detector.esta_entrenado():
            # Intentar entrenar ahora con lo que haya
            self._etz_entrenar_detector(numero)
            if not hasattr(self, "_etz_detector") or not self._etz_detector.esta_entrenado():
                messagebox.showwarning(
                    "Sin modelo",
                    "Etiqueta al menos 2 páginas manualmente para activar la predicción.")
                return

        from core.zone_labeler import cargar_todas_manual
        # Páginas disponibles desde imágenes (fuente más fiable que txt)
        txt_dir = ST.out_dir / "03_ocr" / numero
        img_dir = ST.out_dir / "02_imagenes" / numero

        paginas_disponibles = []
        if img_dir.exists():
            paginas_disponibles = sorted(
                p.stem for p in img_dir.iterdir()
                if p.suffix.lower() in {".png", ".jpg", ".tif", ".tiff"})
        if not paginas_disponibles and txt_dir.exists():
            paginas_disponibles = sorted(p.stem for p in txt_dir.glob("*.txt"))

        if not paginas_disponibles:
            messagebox.showwarning("Sin páginas",
                "No hay imágenes ni archivos OCR para este número.\n"
                "Ejecuta primero la Extracción OCR.")
            return

        # Páginas ya etiquetadas manualmente (lista de PaginaEtiquetada)
        manuales_list  = cargar_todas_manual(ST.out_dir, numero)
        ya_etiquetadas = [p.pagina for p in manuales_list]   # ← lista, no dict
        sin_etiquetar  = [p for p in paginas_disponibles if p not in ya_etiquetadas]

        if not sin_etiquetar:
            messagebox.showinfo("Completo",
                f"Todas las {len(paginas_disponibles)} páginas ya están etiquetadas.")
            return

        # Tamaño de imagen de referencia
        ancho = self._etz_img_orig.width  if self._etz_img_orig else 1000
        alto  = self._etz_img_orig.height if self._etz_img_orig else 1400

        self._etz_lbl_train.config(
            text=f"🔮 Prediciendo {len(sin_etiquetar)} páginas…", fg="#F59E0B")
        self.update_idletasks()

        def _run():
            n = self._etz_detector.aplicar_a_numero(
                ST.out_dir, numero,
                paginas_disponibles, ya_etiquetadas,
                ancho_px=ancho, alto_px=alto,
                umbral_frecuencia=0.3,
            )
            def _after():
                self._etz_lbl_train.config(
                    text=f"✅ {n} páginas predichas · {len(ya_etiquetadas)} manuales",
                    fg="#3FB950")
                self._etz_actualizar_estado()
                # Recargar la página actual si era una de las predichas
                pagina_actual = self._etz_pagina.get()
                if pagina_actual and pagina_actual not in ya_etiquetadas:
                    self._etz_on_pagina()   # recarga zonas + redibuja canvas
            self.after(0, _after)

        threading.Thread(target=_run, daemon=True).start()

    def _etz_actualizar_estado(self):
        if not ST.out_dir:
            return
        numero = self._etz_numero.get()
        if not numero:
            return
        from core.zone_labeler import cargar_todas_manual, listar_paginas_etiquetadas
        etiq     = listar_paginas_etiquetadas(ST.out_dir, numero)
        manuales = cargar_todas_manual(ST.out_dir, numero)
        txt_dir  = ST.out_dir / "03_ocr" / numero
        img_dir  = ST.out_dir / "02_imagenes" / numero
        if txt_dir.exists():
            total = len(list(txt_dir.glob("*.txt")))
        elif img_dir.exists():
            total = len([p for p in img_dir.iterdir()
                         if p.suffix.lower() in {".png",".jpg",".tif",".tiff"}])
        else:
            total = 0
        n_man  = len(manuales)
        n_pred = len(etiq) - n_man
        sin    = max(0, total - len(etiq))
        self._etz_lbl_estado.config(
            text=f"Total páginas: {total}\n"
                 f"Etiquetadas manual: {n_man}\n"
                 f"Predichas: {n_pred}\n"
                 f"Sin etiquetar: {sin}"
        )
        detector = getattr(self, "_etz_detector", None)
        if detector and detector.esta_entrenado():
            # Ya hay modelo — mostrar estado del modelo
            self._etz_lbl_train.config(
                text=f"🧠 Modelo activo · {n_man} pág · {sin} pendientes",
                fg="#3FB950")
            self._btn_etz_predecir.config(bg="#6E40C9")
        elif n_man >= 2:
            # Hay páginas para entrenar pero no se ha entrenado todavía
            self._etz_lbl_train.config(
                text=f"✅ {n_man} páginas · listo para predecir",
                fg="#059669")
            self._btn_etz_predecir.config(bg="#6E40C9")
            # Entrenar automáticamente
            self._etz_entrenar_detector(numero)
        else:
            falta = 2 - n_man
            self._etz_lbl_train.config(
                text=f"Etiqueta {falta} página(s) más para activar predicción",
                fg="#94A3B8")
            self._btn_etz_predecir.config(bg="#444C56")

    def _etz_get_img_path(self, numero: str, pagina: str):
        """
        Devuelve la ruta de imagen para numero/pagina.
        Si no hay PNG en disco, extrae la página del PDF y la guarda en temp.
        """
        # 1. Buscar en carpeta de imágenes del proyecto
        if ST.out_dir:
            img_dir = ST.out_dir / "02_imagenes" / numero
            if img_dir.exists():
                candidatos = list(img_dir.glob(f"*{pagina}*.png"))
                if not candidatos:
                    candidatos = sorted(img_dir.glob("*.png"))
                if candidatos:
                    import re as _re
                    _nums = _re.findall(r'\d+', pagina)
                    n = max(0, int(_nums[-1]) - 1) if _nums else 0
                    return candidatos[n] if n < len(candidatos) else candidatos[0]

        # 2. Extraer desde PDF directamente (carga directa sin OCR previo)
        pdf = next((a for a in (ST.archivos_sel or [])
                    if a.stem == numero), None)
        if pdf and pdf.exists():
            try:
                import re as _re
                import tempfile

                import fitz
                from PIL import Image
                _nums = _re.findall(r'\d+', pagina)
                n_pag = max(0, int(_nums[-1]) - 1) if _nums else 0
                doc  = fitz.open(str(pdf))
                if n_pag < doc.page_count:
                    page = doc[n_pag]
                    mat  = fitz.Matrix(2.0, 2.0)  # 200 DPI para buena detección
                    pix  = page.get_pixmap(matrix=mat)
                    doc.close()
                    # Guardar en temp ASCII
                    tmp_dir = Path(tempfile.gettempdir()) / "bashkar_etz"
                    tmp_dir.mkdir(exist_ok=True)
                    tmp_path = tmp_dir / f"{numero}_{pagina}.png"
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    img.save(str(tmp_path))
                    return tmp_path
            except Exception:
                pass

        # 3. Si la imagen ya está en canvas, usar la imagen en memoria como temp
        if self._etz_img_orig is not None:
            try:
                import tempfile
                tmp_dir = Path(tempfile.gettempdir()) / "bashkar_etz"
                tmp_dir.mkdir(exist_ok=True)
                tmp_path = tmp_dir / f"{numero}_{pagina}_canvas.png"
                self._etz_img_orig.save(str(tmp_path))
                return tmp_path
            except Exception:
                pass

        return None

    def _etz_detectar_pagina(self):
        """Detecta zonas de la página actual con OpenCV o Claude Vision."""
        numero  = self._etz_numero.get()
        pagina  = self._etz_pagina.get()
        if not numero or not pagina:
            messagebox.showwarning("Sin selección", "Selecciona un número y página.")
            return
        img_path = self._etz_get_img_path(numero, pagina)
        if not img_path:
            messagebox.showwarning(
                "Sin imagen",
                "No se encontró la imagen de esta página.\n"
                "Asegúrate de haber convertido el PDF a imágenes (paso OCR Ruta 1 o 2)."
            )
            return
        modo = getattr(self, "_etz_modo_det", tk.StringVar(value="tesseract")).get()

        if modo == "tesseract":
            # Análisis de layout local (Tesseract + OpenCV) — en thread,
            # tarda varios segundos por página
            self._etz_lbl_train.config(
                text="⏳ Analizando layout (Tesseract local)…", fg="#F59E0B")

            def _worker_layout():
                try:
                    from core.layout_tesseract import analizar_pagina_local
                    zonas_t = analizar_pagina_local(img_path)
                    err = ""
                except Exception as e:
                    zonas_t, err = [], str(e)

                def _aplicar():
                    if err:
                        self._etz_lbl_train.config(
                            text=f"⚠ Error: {err[:70]}", fg="#EF4444")
                        return
                    if not zonas_t:
                        self._etz_lbl_train.config(
                            text="⚠ No se detectaron zonas.", fg="#EF4444")
                        return
                    self._etz_zonas = zonas_t
                    # Recargar la imagen: el deskew automático pudo corregirla
                    self._etz_cargar_imagen_pagina(numero, pagina)
                    self._etz_actualizar_lista_zonas()
                    self._etz_redibujar_zonas()
                    self._etz_lbl_train.config(
                        text=f"✅ {len(zonas_t)} zonas detectadas (tesseract). "
                             "Revisa y guarda.",
                        fg="#22C55E")
                self.after(0, _aplicar)

            threading.Thread(target=_worker_layout, daemon=True).start()
            return

        if modo in ("yolo", "onnx", "dit"):
            from core.layout_neural import detectar_layout, motor_disponible
            ok, msg = motor_disponible(modo)
            if not ok:
                if messagebox.askyesno("Motor no instalado",
                        f"{msg}\n\n¿Instalar ahora?"):
                    self._etz_instalar_motor()
                return
            self._etz_lbl_train.config(
                text=f"⏳ Analizando con {modo.upper()}…", fg="#F59E0B")
            self.update_idletasks()

            def _log(m): self._etz_lbl_train.config(text=m[:80], fg="#F59E0B")
            zonas_raw = detectar_layout(img_path, motor=modo, callback=_log)
            # Convertir dicts a objetos Zona
            from core.zone_labeler import TIPOS_ZONA, Zona
            zonas = []
            for z in zonas_raw:
                tipo = z["tipo"] if z["tipo"] in TIPOS_ZONA else "articulo"
                zonas.append(Zona(tipo=tipo, x0=z["x0"], y0=z["y0"],
                                  x1=z["x1"], y1=z["y1"],
                                  confianza=z["confianza"]))

        elif modo == "vision_ia":
            prov   = getattr(self, "_etz_vision_prov",  tk.StringVar(value="claude")).get()
            modelo = getattr(self, "_etz_vision_model", tk.StringVar(value="")).get()
            api_key = ST.api_keys.get(prov, "") or ST.api_key
            if prov != "ollama" and not api_key:
                messagebox.showwarning("API key faltante",
                    f"Configura la API key de {prov} en la pestaña Configuración.")
                return
            self._etz_lbl_train.config(
                text=f"⏳ Consultando {prov}/{modelo}…", fg="#F59E0B")
            self.update_idletasks()
            from core.zone_labeler import detectar_zonas_vision
            zonas = detectar_zonas_vision(img_path, proveedor=prov,
                                          api_key=api_key, modelo=modelo,
                                          prompt_custom=ST.prompt_deteccion)
        else:
            self._etz_lbl_train.config(text="⏳ Analizando con OpenCV…", fg="#F59E0B")
            self.update_idletasks()
            from core.zone_labeler import detectar_zonas_opencv
            zonas = detectar_zonas_opencv(img_path)

        if not zonas:
            self._etz_lbl_train.config(text="⚠ No se detectaron zonas.", fg="#EF4444")
            return

        self._etz_zonas = zonas
        self._etz_actualizar_lista_zonas()
        self._etz_redibujar_zonas()
        self._etz_lbl_train.config(
            text=f"✅ {len(zonas)} zonas detectadas ({modo}). Revisa y guarda.",
            fg="#22C55E"
        )

    def _etz_detectar_numero(self):
        """Detecta zonas en todas las páginas del número (en hilo separado)."""
        numero = self._etz_numero.get()
        if not numero or not ST.out_dir:
            messagebox.showwarning("Sin selección", "Selecciona un número primero.")
            return
        img_dir = ST.out_dir / "02_imagenes" / numero
        if not img_dir.exists() or not list(img_dir.glob("*.png")):
            messagebox.showwarning(
                "Sin imágenes",
                "No se encontraron imágenes para este número.\n"
                "Usa la Ruta 1 o 2 de OCR para generar las imágenes primero."
            )
            return
        modo = getattr(self, "_etz_modo_det", tk.StringVar(value="opencv")).get()
        api_key = ""
        _vision_prov  = getattr(self, "_etz_vision_prov",  tk.StringVar(value="claude")).get()
        _vision_model = getattr(self, "_etz_vision_model", tk.StringVar(value="")).get()
        if modo == "vision_ia":
            api_key = ST.api_keys.get(_vision_prov, "") or ST.api_key
            if _vision_prov != "ollama" and not api_key:
                messagebox.showwarning("API key faltante",
                    f"Configura la API key de {_vision_prov} en la pestaña Configuración.")
                return

        imagenes = sorted(img_dir.glob("*.png"))
        n_total = len(imagenes)
        if not messagebox.askyesno(
            "Detectar todo el número",
            f"Se analizarán {n_total} páginas con "
            f"{'Vision IA (' + _vision_prov + '/' + _vision_model + ')' if modo == 'vision_ia' else modo.upper()}.\n"
            + ("Esto usará tokens de tu cuenta de IA.\n" if modo == "vision_ia" else "")
            + "¿Continuar?"
        ):
            return

        self._etz_lbl_train.config(text=f"⏳ Detectando 0/{n_total}…", fg="#F59E0B")

        def _worker():
            from core.zone_labeler import (
                TIPOS_ZONA,
                PaginaEtiquetada,
                Zona,
                detectar_zonas_opencv,
                guardar_pagina,
            )
            ok = 0
            for i, img_path in enumerate(imagenes):
                stem = img_path.stem
                import re as _re
                m = _re.search(r'(\d+)', stem)
                pagina = f"p{int(m.group(1)):04d}" if m else stem

                if modo == "tesseract":
                    from core.layout_tesseract import analizar_pagina_local
                    zonas = analizar_pagina_local(img_path)
                elif modo in ("yolo", "onnx", "dit"):
                    from core.layout_neural import detectar_layout
                    zonas_raw = detectar_layout(img_path, motor=modo)
                    zonas = []
                    for z in zonas_raw:
                        tipo = z["tipo"] if z["tipo"] in TIPOS_ZONA else "articulo"
                        zonas.append(Zona(tipo=tipo, x0=z["x0"], y0=z["y0"],
                                          x1=z["x1"], y1=z["y1"],
                                          confianza=z["confianza"]))
                elif modo == "vision_ia":
                    from core.zone_labeler import detectar_zonas_vision
                    zonas = detectar_zonas_vision(img_path, proveedor=_vision_prov,
                                                   api_key=api_key, modelo=_vision_model,
                                                   prompt_custom=ST.prompt_deteccion)
                else:
                    zonas = detectar_zonas_opencv(img_path)

                if zonas:
                    from PIL import Image as _Img
                    try:
                        im = _Img.open(img_path)
                        W, H = im.size
                    except Exception:
                        W, H = 1000, 1400
                    pag = PaginaEtiquetada(
                        pagina=pagina, ancho_px=W, alto_px=H,
                        zonas=zonas, manual=False
                    )
                    guardar_pagina(ST.out_dir, numero, pag)
                    ok += 1

                self.after(0, lambda i=i, ok=ok: self._etz_lbl_train.config(
                    text=f"⏳ Detectando {i+1}/{n_total} ({ok} con zonas)…",
                    fg="#F59E0B"
                ))

            self.after(0, lambda: (
                self._etz_lbl_train.config(
                    text=f"✅ {ok}/{n_total} páginas con zonas detectadas.",
                    fg="#22C55E"
                ),
                self._etz_actualizar_estado(),
            ))

        import threading
        threading.Thread(target=_worker, daemon=True).start()

    def _etz_instalar_motor(self):
        """Instala las dependencias del motor seleccionado en segundo plano."""
        from core.layout_neural import instalar_motor, motor_disponible
        modo = self._etz_modo_det.get()
        if modo in ("tesseract", "opencv", "vision_ia"):
            messagebox.showinfo(
                "Sin instalación",
                f"El motor '{modo}' no requiere instalar nada adicional.")
            return
        ok, _ = motor_disponible(modo)
        if ok:
            messagebox.showinfo("Ya instalado", f"El motor '{modo}' ya está disponible.")
            return
        self._etz_lbl_train.config(text=f"⬇ Instalando {modo}…", fg="#F59E0B")
        def _run():
            def cb(m): self.after(0, lambda: self._etz_lbl_train.config(
                text=m[:80], fg="#F59E0B"))
            exito = instalar_motor(modo, callback=cb)
            msg = f"✅ {modo} instalado. Reinicia la app." if exito \
                  else f"⚠ Error instalando {modo}. Revisa la conexión."
            self.after(0, lambda: (
                self._etz_lbl_train.config(text=msg, fg="#22C55E" if exito else "#EF4444"),
                messagebox.showinfo("Instalación", msg),
            ))
        threading.Thread(target=_run, daemon=True).start()

    # ── Actualizar selector de números en ETZ ────────────────────────────────
    def _etz_deskew_pagina(self):
        """Detecta la inclinación de la página actual, muestra previsualización
        con slider de ajuste fino y guarda la imagen corregida si el usuario confirma."""
        if self._etz_img_orig is None:
            messagebox.showwarning("Sin imagen", "Carga una página primero.")
            return

        numero = self._etz_numero.get()
        pagina = self._etz_pagina.get()
        img_path = self._etz_get_img_path(numero, pagina)
        if not img_path:
            messagebox.showwarning("Sin imagen", "No se encontró el archivo de imagen.")
            return

        from PIL import Image

        from core.image_preprocessor import detectar_angulo_pagina

        img_orig = Image.open(img_path)
        angulo_detectado = detectar_angulo_pagina(img_orig)

        # ── Ventana de previsualización ───────────────────────────────────────
        win = tk.Toplevel(self)
        win.title(f"Corregir inclinación — {pagina}")
        win.geometry("820x640")
        win.configure(bg=CONTENT_BG)
        win.grab_set()

        # Encabezado
        hdr = tk.Frame(win, bg=CONTENT_BG)
        hdr.pack(fill="x", padx=12, pady=(10, 0))
        tk.Label(hdr, text="Corrección de inclinación",
                 bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 10, "bold")).pack(side="left")
        lbl_angulo = tk.Label(hdr,
                              text=f"Ángulo detectado: {angulo_detectado:+.2f}°",
                              bg=CONTENT_BG, fg="#F59E0B" if abs(angulo_detectado) > 0.3 else "#22C55E",
                              font=("Segoe UI", 9))
        lbl_angulo.pack(side="left", padx=12)

        # Canvas de previsualización
        canvas_frame = tk.Frame(win, bg="#1E293B", relief="sunken", bd=1)
        canvas_frame.pack(fill="both", expand=True, padx=12, pady=8)
        prev_canvas = tk.Canvas(canvas_frame, bg="#1E293B", highlightthickness=0)
        prev_canvas.pack(fill="both", expand=True)

        # Estado interno de la ventana
        _state = {"angulo": tk.DoubleVar(value=angulo_detectado), "img_tk": None}

        def _render_preview(angulo_val):
            """Rota la imagen al ángulo dado y la muestra en el canvas."""
            from PIL import Image, ImageTk
            try:
                a = float(angulo_val)
                if abs(a) < 0.1:
                    img_rot = img_orig.copy()
                else:
                    img_rot = img_orig.rotate(
                        -a, expand=True,
                        fillcolor=(255, 255, 255) if img_orig.mode != "L" else 255,
                        resample=Image.BICUBIC,
                    )
                # Escalar para caber en el canvas
                cw = prev_canvas.winfo_width() or 760
                ch = prev_canvas.winfo_height() or 480
                ratio = min(cw / img_rot.width, ch / img_rot.height, 1.0)
                nw = max(1, int(img_rot.width  * ratio))
                nh = max(1, int(img_rot.height * ratio))
                img_small = img_rot.resize((nw, nh), Image.LANCZOS)
                _state["img_tk"] = ImageTk.PhotoImage(img_small)
                prev_canvas.delete("all")
                prev_canvas.create_image(cw // 2, ch // 2, anchor="center",
                                         image=_state["img_tk"])
            except Exception:
                pass

        def _on_slider(val):
            a = round(float(val), 2)
            _state["angulo"].set(a)
            lbl_angulo.config(
                text=f"Ángulo: {a:+.2f}°",
                fg="#F59E0B" if abs(a) > 0.3 else "#22C55E")
            _render_preview(a)

        # Panel de controles
        ctrl = tk.Frame(win, bg=CONTENT_BG)
        ctrl.pack(fill="x", padx=12, pady=(0, 4))

        tk.Label(ctrl, text="Ajuste fino:", bg=CONTENT_BG, fg="#CDD6F4",
                 font=("Segoe UI", 8, "bold")).pack(side="left")

        slider = tk.Scale(ctrl, from_=-15.0, to=15.0, resolution=0.1,
                          orient="horizontal", length=400,
                          variable=_state["angulo"],
                          command=_on_slider,
                          bg=CONTENT_BG, fg="#CDD6F4",
                          highlightthickness=0, troughcolor="#E2E8F0",
                          font=("Segoe UI", 7))
        slider.pack(side="left", padx=(6, 12))

        def _reset():
            _state["angulo"].set(angulo_detectado)
            slider.set(angulo_detectado)
            _on_slider(angulo_detectado)

        tk.Button(ctrl, text="↺ Restablecer detectado",
                  command=_reset,
                  font=("Segoe UI", 8), bg="#1C2128").pack(side="left", padx=4)

        tk.Button(ctrl, text="0° (sin corrección)",
                  command=lambda: (_state["angulo"].set(0), slider.set(0), _on_slider(0)),
                  font=("Segoe UI", 8), bg="#1C2128").pack(side="left", padx=4)

        # Botones finales
        btn_frame = tk.Frame(win, bg=CONTENT_BG)
        btn_frame.pack(fill="x", padx=12, pady=(0, 12))

        lbl_aviso = tk.Label(btn_frame,
                             text="La imagen original se preserva. La versión corregida se guarda en 02_imagenes_ocr/.",
                             bg=CONTENT_BG, fg=TXT_SEC, font=("Segoe UI", 8))
        lbl_aviso.pack(side="left")

        def _aplicar():
            a = float(_state["angulo"].get())
            if abs(a) < 0.1:
                messagebox.showinfo("Sin cambio",
                                    "El ángulo es 0° — no hay nada que corregir.",
                                    parent=win)
                return
            try:
                from PIL import Image
                img_corr = img_orig.rotate(
                    -a, expand=True,
                    fillcolor=(255, 255, 255) if img_orig.mode != "L" else 255,
                    resample=Image.BICUBIC,
                )
                # Guardar SOLO en 02_imagenes_ocr/ — nunca sobreescribir original
                if ST.out_dir:
                    dir_ocr_img = Path(ST.out_dir) / "02_imagenes_ocr" / numero
                    dir_ocr_img.mkdir(parents=True, exist_ok=True)
                    dest = dir_ocr_img / img_path.name
                    img_corr.save(str(dest))
                    messagebox.showinfo("Guardado",
                        f"Imagen corregida guardada en:\n02_imagenes_ocr/{numero}/{img_path.name}\n\n"
                        f"La imagen original permanece intacta.", parent=win)
                win.destroy()
                self._etz_cargar_imagen_pagina(numero, pagina)
                self._etz_zoom = 1.0
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar: {e}", parent=win)

        tk.Button(btn_frame, text="Cancelar",
                  command=win.destroy,
                  font=("Segoe UI", 8), bg="#1C2128").pack(side="right", padx=(4, 0))
        tk.Button(btn_frame, text="✅ Aplicar y guardar",
                  command=_aplicar,
                  font=("Segoe UI", 9, "bold"),
                  bg="#1D4ED8", fg="white").pack(side="right")

        # Renderizar preview inicial tras que el canvas tenga tamaño real
        win.update_idletasks()
        _render_preview(angulo_detectado)

    def _etz_editar_prompt(self):
        """Abre ventana para editar el prompt de detección de zonas por IA."""
        from core.zone_labeler import construir_prompt_deteccion
        win, prompt_content = self._mk_glass_toplevel(
            "⚙ Prompt de detección de zonas", ancho=720, alto=580)
        win.grab_set()

        tk.Label(prompt_content, text="Prompt enviado a la IA de visión para detectar zonas",
                 bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=12, pady=(12, 0))
        tk.Label(prompt_content,
                 text="Incluye automáticamente los tipos personalizados de tu proyecto (★). "
                      "Edita para ajustar instrucciones específicas de Estampa.",
                 bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 8), wraplength=680, justify="left").pack(
                     anchor="w", padx=12, pady=(2, 6))

        frame_txt = tk.Frame(prompt_content, bg=CONTENT_BG)
        frame_txt.pack(fill="both", expand=True, padx=12, pady=(0, 6))
        scroll = tk.Scrollbar(frame_txt)
        scroll.pack(side="right", fill="y")
        txt = tk.Text(frame_txt, wrap="word", font=("Consolas", 9),
                      bg=CARD_BG, fg=TXT_PRI,
                      insertbackground=TXT_PRI,
                      yscrollcommand=scroll.set,
                      relief="solid", bd=1)
        txt.pack(fill="both", expand=True)
        scroll.config(command=txt.yview)

        # Mostrar prompt generado dinámicamente (con tipos custom incluidos)
        prompt_generado = construir_prompt_deteccion(ST.prompt_deteccion)
        txt.insert("1.0", prompt_generado)

        btn_frame = tk.Frame(prompt_content, bg=CONTENT_BG)
        btn_frame.pack(fill="x", padx=12, pady=(0, 12))

        def _restaurar():
            """Regenera el prompt con todos los tipos activos actuales."""
            txt.delete("1.0", "end")
            txt.insert("1.0", construir_prompt_deteccion(""))

        def _guardar():
            nuevo = txt.get("1.0", "end").strip()
            default_generado = construir_prompt_deteccion("").strip()
            # Guardar vacío si es igual al generado (se regenerará siempre al usar)
            ST.prompt_deteccion = "" if nuevo == default_generado else nuevo
            win.destroy()
            self._etz_lbl_train.config(
                text="✅ Prompt personalizado guardado." if ST.prompt_deteccion
                     else "✅ Usando prompt generado automáticamente.",
                fg="#22C55E")

        tk.Button(btn_frame, text="↺ Restaurar por defecto",
                  command=_restaurar,
                  font=("Segoe UI", 8), bg="#1C2128").pack(side="left")
        tk.Button(btn_frame, text="Cancelar",
                  command=win.destroy,
                  font=("Segoe UI", 8), bg="#1C2128").pack(side="right", padx=(4, 0))
        tk.Button(btn_frame, text="💾 Guardar",
                  command=_guardar,
                  font=("Segoe UI", 9, "bold"),
                  bg="#3B82F6", fg="white").pack(side="right")

    def _etz_abrir_pdf_directo(self):
        """Carga un PDF directamente en el etiquetador sin pasar por Configuración."""
        from tkinter import filedialog
        ruta = filedialog.askopenfilename(
            title="Abrir PDF para etiquetar",
            filetypes=[("PDF", "*.pdf"), ("Todos", "*.*")])
        if not ruta:
            return
        p = Path(ruta)

        # Agregar a ST.archivos_sel si no está
        if not hasattr(ST, "archivos_sel") or ST.archivos_sel is None:
            ST.archivos_sel = []
        if p not in ST.archivos_sel:
            ST.archivos_sel.append(p)

        # Configurar out_dir si no está configurado
        if not getattr(ST, "out_dir", None):
            import tempfile
            ST.out_dir = Path(tempfile.gettempdir()) / "bashkar_etz_temp" / p.stem
            ST.out_dir.mkdir(parents=True, exist_ok=True)

        # Actualizar combobox con el nuevo archivo
        nombre = p.stem
        vals = list(self._etz_cb_num["values"])
        if nombre not in vals:
            vals.append(nombre)
            self._etz_cb_num["values"] = vals
        self._etz_cb_num.set(nombre)
        self._etz_numero.set(nombre)

        # Poblar páginas directamente desde el PDF
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(p))
            n_pags = doc.page_count
            doc.close()
            pags = [f"p{i+1:04d}" for i in range(n_pags)]
        except Exception:
            pags = [f"p{i+1:04d}" for i in range(50)]

        self._etz_cb_pag["values"] = pags
        self._etz_poblar_miniaturas(pags)

        if pags:
            self._etz_cb_pag.set(pags[0])
            self._etz_pagina.set(pags[0])
            # Cargar imagen de la primera página directamente desde el PDF
            self._etz_cargar_imagen_desde_pdf(str(p), 0)

        self._etz_lbl_train.config(text=f"✅ {nombre} — {len(pags)} páginas")

    def _etz_cargar_imagen_desde_pdf(self, pdf_path: str, n_pag: int):
        """Carga una página de PDF directamente como imagen en el canvas."""
        try:
            import fitz
            from PIL import Image, ImageTk
            doc  = fitz.open(pdf_path)
            page = doc[n_pag]
            mat  = fitz.Matrix(1.5, 1.5)  # 150% escala para mejor resolución
            pix  = page.get_pixmap(matrix=mat)
            doc.close()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Escalar al canvas
            max_h = 900
            if img.height > max_h:
                escala = max_h / img.height
                img = img.resize((int(img.width * escala), max_h), Image.LANCZOS)
                self._etz_escala = escala
            else:
                self._etz_escala = 1.0

            self._etz_img_orig = img
            self._etz_img_tk   = ImageTk.PhotoImage(img)
            self._etz_canvas.delete("all")
            self._etz_canvas.create_image(0, 0, anchor="nw", image=self._etz_img_tk)
            self._etz_canvas.configure(scrollregion=(0, 0, img.width, img.height))
            self._etz_redibujar_zonas()

        except Exception as ex:
            self._etz_canvas.delete("all")
            self._etz_canvas.create_text(
                150, 100, text=f"Error cargando imagen:\n{ex}",
                fill="#94A3B8", font=("Segoe UI", 9), anchor="nw")

    def _etz_refrescar_numeros(self):
        """Llamado cuando el corpus cambia para actualizar el combobox."""
        if not hasattr(self, "_etz_cb_num"):
            return
        numeros = []
        if ST.corpus_meta is not None and "numero" in ST.corpus_meta.columns:
            numeros = sorted(ST.corpus_meta["numero"].unique().tolist())
        elif ST.archivos_sel:
            numeros = [a.stem for a in ST.archivos_sel]
        self._etz_cb_num["values"] = numeros
        if numeros and not self._etz_numero.get():
            self._etz_cb_num.set(numeros[0])

    # ══════════════════════════════════════════════════════════════════════════
    # NORMALIZAR: revisión y edición del texto OCR por bloques
    # ══════════════════════════════════════════════════════════════════════════
    def _build_norm(self):
        f = self._tab_norm
        self._page_header(f, "Normalizar texto",
                          "Revisa y edita el texto OCR antes de analizar · 4 vistas por bloque", "📝")

        # ── Barra de acciones ─────────────────────────────────────────────────
        bbar = tk.Frame(f, bg=CONTENT_BG)
        bbar.pack(fill="x", padx=24, pady=(0, 6))

        self._norm_var_numero = tk.StringVar()
        self._norm_cb_num = ttk.Combobox(bbar, textvariable=self._norm_var_numero,
                                          width=24, state="readonly")
        self._norm_cb_num.pack(side="left", padx=(0, 8))
        self._norm_cb_num.bind("<<ComboboxSelected>>", lambda e: self._norm_cargar_numero())

        ttk.Button(bbar, text="▶  Normalizar automático", style="P.TButton",
                   command=self._norm_auto).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="🤖  Sugerir con IA", style="S.TButton",
                   command=self._norm_ia).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="💾  Guardar ediciones", style="S.TButton",
                   command=self._norm_guardar).pack(side="left", padx=(0, 8))

        # Re-extraer con Tesseract — rescata páginas con OCR corrupto (ej. Kraken sin RAM)
        ttk.Button(bbar, text="🔄 Re-OCR Tesseract (página)",
                   style="S.TButton",
                   command=self._norm_reocr_pagina).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="🔄 Re-OCR Tesseract (todo)",
                   style="S.TButton",
                   command=self._norm_reocr_numero).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="🖼 Regenerar imágenes",
                   style="S.TButton",
                   command=self._norm_regenerar_imagenes).pack(side="left", padx=(0, 8))

        ttk.Button(bbar, text="📂 Importar .txt", style="S.TButton",
                   command=self._norm_importar_txt).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="⚙ Reconstruir columnas BNC", style="S.TButton",
                   command=self._norm_reconstruir_columnas).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="📖 Diccionario de corpus", style="S.TButton",
                   command=self._norm_diccionario_corpus).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="🔍 Ver cambios", style="S.TButton",
                   command=self._norm_ver_diff).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="🏋 Dataset HTR", style="S.TButton",
                   command=self._norm_exportar_ground_truth).pack(side="left", padx=(0, 8))
        ttk.Button(bbar, text="↺  Actualizar lista", style="S.TButton",
                   command=self._norm_refrescar_numeros).pack(side="right")
        ttk.Button(bbar, text="📓 Nota", style="S.TButton",
                   command=lambda: self._bitacora_nueva_nota("norm")).pack(side="right", padx=(0, 4))

        self._lbl_norm_estado = tk.Label(bbar, text="", bg=CONTENT_BG, fg=VERDE,
                                          font=("Segoe UI", 9, "bold"))
        self._lbl_norm_estado.pack(side="right", padx=8)

        # ── Selector de versión para el pipeline ──────────────────────────────
        vbar = tk.Frame(f, bg="#0D1117", pady=6)
        vbar.pack(fill="x", padx=24, pady=(0, 4))

        tk.Label(vbar,
                 text="Versión que pasa al análisis:",
                 bg="#0D1117", fg=TXT_SEC,
                 font=("Segoe UI", 9, "bold")).pack(side="left", padx=(0, 12))

        self._norm_var_version = tk.StringVar(
            value=getattr(ST, "norm_version", "manual"))

        _opciones = [
            ("crudo",  "🗒 Crudo (OCR sin cambios)",
             "El texto tal como salió del OCR o del conversor.\n"
             "Para investigadores que prefieren trabajar con el\n"
             "original y aplicar sus propias transformaciones."),
            ("manual", "✏ Manual (mis ediciones)",
             "Lo que escribiste o corregiste en el panel izquierdo.\n"
             "Si no editaste una página, usa el texto crudo.\n"
             "La opción más común para trabajo de investigación."),
            ("ia",     "🤖 IA (sugerencia revisada)",
             "La versión generada por el asistente IA.\n"
             "Si no hay sugerencia para una página, cae al manual\n"
             "y luego al crudo.\n"
             "Para quienes quieren modernizar la ortografía\n"
             "o acelerar la corrección masiva."),
        ]

        for val, etiq, ayuda in _opciones:
            rb = ttk.Radiobutton(vbar, text=etiq,
                                  variable=self._norm_var_version,
                                  value=val,
                                  command=self._norm_version_cambio)
            rb.pack(side="left", padx=(0, 4))
            self._mk_ayuda_bg(vbar, ayuda, bg="#0D1117")

        self._norm_lbl_version_info = tk.Label(
            vbar, text="", bg="#0D1117", fg=TXT_DIM,
            font=("Segoe UI", 8, "italic"))
        self._norm_lbl_version_info.pack(side="right", padx=8)

        ttk.Button(vbar, text="✔ Verificar", style="S.TButton",
                   command=self._verif_abrir).pack(side="right", padx=(0, 8))
        self._mk_ayuda_bg(vbar,
            "Verificación palabra por palabra (estilo ABBYY FineReader):\n"
            "recorre las palabras de baja confianza del OCR de esta página,\n"
            "muestra el recorte ampliado y sugerencias, y aplica la\n"
            "corrección elegida al texto manual del bloque actual.",
            bg="#0D1117")

        # ── Selector de bloque (listbox de bloques de la página) ─────────────
        mid = tk.Frame(f, bg=CONTENT_BG)
        mid.pack(fill="both", expand=True, padx=24, pady=(0, 12))

        # Panel izquierdo: lista de páginas y bloques
        izq = tk.Frame(mid, bg=CARD_BG, relief="solid", bd=1, width=220)
        izq.pack(side="left", fill="y", padx=(0, 8))
        izq.pack_propagate(False)

        tk.Label(izq, text="Páginas / bloques", bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(pady=(8, 4), padx=8, anchor="w")

        self._norm_lb = tk.Listbox(izq, bg=CARD_BG, fg=TXT_SEC, selectbackground=AB_SEL,
                                    selectforeground="#FFFFFF", relief="flat",
                                    font=("Segoe UI", 9), activestyle="none",
                                    exportselection=False)
        sb_lb = ttk.Scrollbar(izq, orient="vertical", command=self._norm_lb.yview)
        self._norm_lb.config(yscrollcommand=sb_lb.set)
        sb_lb.pack(side="right", fill="y")
        self._norm_lb.pack(fill="both", expand=True, padx=(4, 0), pady=(0, 8))
        self._norm_lb.bind("<<ListboxSelect>>", lambda e: self._norm_seleccionar_bloque())

        # Panel derecho: 4 vistas
        der = tk.Frame(mid, bg=CONTENT_BG)
        der.pack(side="left", fill="both", expand=True)

        # Fila 1: imagen + OCR crudo
        fila1 = tk.Frame(der, bg=CONTENT_BG)
        fila1.pack(fill="both", expand=True, pady=(0, 6))

        # Vista 1: imagen de la página con zoom/pan
        v1 = tk.LabelFrame(fila1, text=" 🖼  Imagen original ",
                            bg=CARD_BG, fg=TXT_PRI, font=("Segoe UI", 9, "bold"),
                            relief="solid", bd=1)
        v1.pack(side="left", fill="both", expand=True, padx=(0, 6))

        # Toolbar de zoom
        v1_tb = tk.Frame(v1, bg=CARD_BG)
        v1_tb.pack(fill="x", padx=6, pady=(4, 0))
        self._norm_zoom = 1.0
        self._norm_img_orig_full = None   # PIL Image a resolución original
        self._norm_pan_start     = None

        for txt, delta in [("−", -1), ("+", 1)]:
            tk.Button(v1_tb, text=txt, bg=CARD_BG, fg=TXT_SEC, relief="flat",
                      font=("Segoe UI", 10, "bold"), width=2, cursor="hand2",
                      command=lambda d=delta: self._norm_zoom_step(d)
                      ).pack(side="left", padx=1)
        self._norm_lbl_zoom = tk.Label(v1_tb, text="100%", bg=CARD_BG, fg=TXT_DIM,
                                        font=("Segoe UI", 8))
        self._norm_lbl_zoom.pack(side="left", padx=6)
        tk.Label(v1_tb, text="Ctrl+rueda: zoom  ·  Arrastrar: pan",
                 bg=CARD_BG, fg=TXT_DIM, font=("Segoe UI", 7)).pack(side="right")

        # Canvas con scrollbars
        v1_wrap = tk.Frame(v1, bg="#000000")
        v1_wrap.pack(fill="both", expand=True, padx=6, pady=(2, 6))
        _sb_cx = ttk.Scrollbar(v1_wrap, orient="horizontal")
        _sb_cy = ttk.Scrollbar(v1_wrap, orient="vertical")
        self._norm_canvas_img = tk.Canvas(v1_wrap, bg="#000000", width=300, height=180,
                                           highlightthickness=0,
                                           xscrollcommand=_sb_cx.set,
                                           yscrollcommand=_sb_cy.set)
        _sb_cx.config(command=self._norm_canvas_img.xview)
        _sb_cy.config(command=self._norm_canvas_img.yview)
        _sb_cy.pack(side="right",  fill="y")
        _sb_cx.pack(side="bottom", fill="x")
        self._norm_canvas_img.pack(fill="both", expand=True)

        # Bindings zoom y pan
        self._norm_canvas_img.bind("<Control-MouseWheel>", self._norm_on_zoom)
        self._norm_canvas_img.bind("<ButtonPress-1>",      self._norm_pan_start_cb)
        self._norm_canvas_img.bind("<B1-Motion>",          self._norm_pan_drag_cb)
        self._norm_canvas_img.bind("<ButtonRelease-1>",    self._norm_pan_end_cb)

        # Vista 2: OCR crudo (solo lectura)
        v2 = tk.LabelFrame(fila1, text=" 📄  OCR crudo (solo lectura) ",
                            bg=CARD_BG, fg=TXT_PRI, font=("Segoe UI", 9, "bold"),
                            relief="solid", bd=1)
        v2.pack(side="left", fill="both", expand=True)
        self._norm_txt_ocr = scrolledtext.ScrolledText(
            v2, bg="#0D1117", fg="#8B949E", insertbackground="#CDD6F4",
            font=("Courier New", 9), relief="flat", wrap="word", state="disabled")
        self._norm_txt_ocr.pack(fill="both", expand=True, padx=6, pady=6)

        # Fila 2: normalizado usuario + normalizado IA
        fila2 = tk.Frame(der, bg=CONTENT_BG)
        fila2.pack(fill="both", expand=True)

        # Vista 3: edición manual del usuario
        v3 = tk.LabelFrame(fila2, text=" ✏️  Normalizado por usuario ",
                            bg=CARD_BG, fg=TXT_PRI, font=("Segoe UI", 9, "bold"),
                            relief="solid", bd=1)
        v3.pack(side="left", fill="both", expand=True, padx=(0, 6))

        # Barra de herramientas del panel de usuario (dictado)
        v3_bar = tk.Frame(v3, bg=CARD_BG)
        v3_bar.pack(fill="x", padx=6, pady=(4, 0))
        self._btn_dictar = ttk.Button(v3_bar, text="🎙 Dictar",
                                       style="S.TButton",
                                       command=self._norm_dictar_toggle)
        self._btn_dictar.pack(side="left")
        self._lbl_dictar_estado = tk.Label(v3_bar, text="", bg=CARD_BG,
                                            fg=TXT_DIM, font=("Segoe UI", 8))
        self._lbl_dictar_estado.pack(side="left", padx=(8, 0))
        self._dictar_session = None   # DictadoSession activa o None

        self._norm_txt_usuario = scrolledtext.ScrolledText(
            v3, bg="#1C2128", fg="#CDD6F4", insertbackground="#CDD6F4",
            font=("Courier New", 9), relief="flat", wrap="word")
        self._norm_txt_usuario.pack(fill="both", expand=True, padx=6, pady=6)

        # Vista 4: sugerencia de IA (revisable)
        v4 = tk.LabelFrame(fila2, text=" 🤖  Normalizado por IA (revisable) ",
                            bg=CARD_BG, fg=TXT_PRI, font=("Segoe UI", 9, "bold"),
                            relief="solid", bd=1)
        v4.pack(side="left", fill="both", expand=True)
        self._norm_txt_ia = scrolledtext.ScrolledText(
            v4, bg="#1C2128", fg="#CDD6F4", insertbackground="#CDD6F4",
            font=("Courier New", 9), relief="flat", wrap="word")
        self._norm_txt_ia.pack(fill="both", expand=True, padx=6, pady=6)

        # Estado interno del panel
        self._norm_bloques: list[dict] = []   # [{pagina, bloque_idx, ocr_crudo, norm_usuario, norm_ia}]
        self._norm_idx_actual: int = -1

    # ── Helpers de normalización ──────────────────────────────────────────────

    def _norm_refrescar_numeros(self):
        if not ST.out_dir:
            return
        txt_base = Path(ST.out_dir) / "03_ocr"
        if not txt_base.exists():
            return
        nums = sorted(p.name for p in txt_base.iterdir()
                      if p.is_dir() and list(p.glob("*.txt")))
        if not nums:
            return
        self._norm_cb_num["values"] = nums
        # Conservar selección actual si sigue siendo válida; si no, tomar el primero
        actual = self._norm_var_numero.get()
        if actual not in nums:
            actual = nums[0]
        self._norm_var_numero.set(actual)
        # Siempre recargar — puede haber archivos nuevos tras el OCR
        self._norm_cargar_numero()

    def _norm_cargar_numero(self):
        """Carga todas las páginas del número seleccionado."""
        num = self._norm_var_numero.get()
        if not num or not ST.out_dir:
            self._lbl_norm_estado.config(
                text="⚠ Sin proyecto abierto o sin OCR ejecutado")
            return
        txt_dir = Path(ST.out_dir) / "03_ocr" / num
        if not txt_dir.exists():
            self._lbl_norm_estado.config(
                text=f"⚠ Carpeta no encontrada: {txt_dir}")
            return
        db_path = Path(ST.ruta_db) if ST.ruta_db else None

        archivos = sorted(txt_dir.glob("*.txt"))
        if not archivos:
            self._lbl_norm_estado.config(
                text=f"⚠ Sin archivos .txt en {txt_dir}")
            return

        bloques = []
        for txt_path in archivos:
            pagina = txt_path.stem
            try:
                ocr_crudo = txt_path.read_text(encoding="utf-8", errors="replace")
            except Exception:
                ocr_crudo = ""
            norm_usuario, norm_ia = self._norm_leer_db(db_path, num, pagina)
            bloques.append({
                "numero":       num,
                "pagina":       pagina,
                "ocr_crudo":    ocr_crudo,
                "norm_usuario": norm_usuario or "",
                "norm_ia":      norm_ia or "",
                "txt_path":     str(txt_path),
            })

        self._norm_bloques = bloques
        self._norm_lb.delete(0, "end")
        avisos_ocr = getattr(self, "_avisos_ocr", {})
        for b in bloques:
            estado = "✓" if b["norm_usuario"] else "·"
            alerta = " ⚠" if (b["numero"], b["pagina"]) in avisos_ocr else ""
            self._norm_lb.insert("end", f"{estado} {b['pagina']}{alerta}")

        self._lbl_norm_estado.config(
            text=f"{len(bloques)} páginas cargadas — {num}")

        if bloques:
            self._norm_lb.selection_set(0)
            self._norm_idx_actual = 0
            self._norm_mostrar_bloque(0)

    def _norm_seleccionar_bloque(self):
        sel = self._norm_lb.curselection()
        if not sel:
            return
        idx = sel[0]
        if 0 <= idx < len(self._norm_bloques):
            self._norm_guardar_bloque_actual()
            self._norm_idx_actual = idx
            self._norm_mostrar_bloque(idx)

    def _norm_mostrar_bloque(self, idx: int):
        if idx < 0 or idx >= len(self._norm_bloques):
            return
        b = self._norm_bloques[idx]

        # Vista OCR crudo (solo lectura)
        self._norm_txt_ocr.config(state="normal")
        self._norm_txt_ocr.delete("1.0", "end")
        self._norm_txt_ocr.insert("1.0", b["ocr_crudo"])
        self._norm_txt_ocr.config(state="disabled")

        # Vista usuario — vacía si no hay edición manual previa (no rellenar con basura OCR)
        self._norm_txt_usuario.delete("1.0", "end")
        if b["norm_usuario"]:
            self._norm_txt_usuario.insert("1.0", b["norm_usuario"])

        # Vista IA
        self._norm_txt_ia.delete("1.0", "end")
        self._norm_txt_ia.insert("1.0", b["norm_ia"])

        # Imagen del bloque
        self._norm_mostrar_imagen(b)

    def _norm_mostrar_imagen(self, bloque: dict):
        """Carga la imagen ORIGINAL a color de la página."""
        self._norm_canvas_img.delete("all")
        try:
            from PIL import Image, ImageTk
            numero  = bloque["numero"]
            pagina  = bloque["pagina"]
            img_pil = None

            # 1. Buscar en 02_imagenes/ (originales a color, cualquier extensión)
            if ST.out_dir:
                img_dir = Path(ST.out_dir) / "02_imagenes" / numero
                for ext in ("*.png", "*.jpg", "*.tif", "*.tiff"):
                    hits = sorted(img_dir.glob(f"*{pagina}*{ext[1:]}")) if img_dir.exists() else []
                    if not hits:
                        hits = sorted(img_dir.glob(ext)) if img_dir.exists() else []
                    if hits:
                        img_pil = Image.open(hits[0]).convert("RGB")
                        break

            # 2. Renderizar directamente desde el PDF (sin necesitar imágenes extraídas)
            if img_pil is None:
                import re as _re
                m_pag = _re.search(r'\d+', pagina)
                n_pag = max(0, int(m_pag.group()) - 1) if m_pag else 0

                # Candidatos: 01_pdfs/, archivos_sel, pdf_dir de entrada, carpeta entrada conversor
                candidatos_pdf: list[Path] = []
                if ST.out_dir:
                    pdf_dir = Path(ST.out_dir) / "01_pdfs"
                    if pdf_dir.exists():
                        candidatos_pdf += list(pdf_dir.glob(f"{numero}*.pdf"))
                candidatos_pdf += [p for p in getattr(ST, "archivos_sel", [])
                                   if hasattr(p, "suffix") and p.suffix.lower() == ".pdf"]
                # Buscar en pdf_dir de entrada (lo que usó el conversor)
                if ST.pdf_dir and Path(ST.pdf_dir).exists():
                    candidatos_pdf += list(Path(ST.pdf_dir).glob("*.pdf"))
                # Filtrar: preferir el que tenga el número en el stem
                exactos = [p for p in candidatos_pdf if numero in p.stem]
                pdfs = exactos or candidatos_pdf

                for pdf_path in pdfs:
                    try:
                        import io

                        import fitz
                        doc = fitz.open(str(pdf_path))
                        if n_pag < doc.page_count:
                            pix = doc[n_pag].get_pixmap(dpi=120)
                            img_pil = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                        doc.close()
                        if img_pil:
                            break
                    except Exception:
                        continue

            if img_pil is None:
                self._norm_canvas_img.create_text(
                    150, 90,
                    text="Sin imagen\n(configura la carpeta de entrada en Configuración)",
                    fill="#484F58", font=("Segoe UI", 9), anchor="center")
                return

            # Guardar original completa para zoom y resetear nivel
            self._norm_img_orig_full = img_pil.copy()
            self._norm_zoom = 1.0
            if hasattr(self, "_norm_lbl_zoom"):
                self._norm_lbl_zoom.config(text="100%")

            cw = max(self._norm_canvas_img.winfo_width(), 300)
            ch = max(self._norm_canvas_img.winfo_height(), 180)
            img_pil.thumbnail((cw, ch), Image.LANCZOS)
            self._norm_zoom = img_pil.width / max(self._norm_img_orig_full.width, 1)
            if hasattr(self, "_norm_lbl_zoom"):
                self._norm_lbl_zoom.config(text=f"{int(self._norm_zoom*100)}%")
            self._norm_canvas_img._img_ref = ImageTk.PhotoImage(img_pil)
            self._norm_canvas_img.configure(
                scrollregion=(0, 0, img_pil.width, img_pil.height))
            self._norm_canvas_img.create_image(
                cw // 2, ch // 2, anchor="center",
                image=self._norm_canvas_img._img_ref)
        except Exception as e:
            self._norm_canvas_img.create_text(
                150, 90, text=f"Error: {e}",
                fill="#F85149", font=("Segoe UI", 8), anchor="center")

    def _norm_guardar_bloque_actual(self):
        """Guarda el estado del bloque actualmente en pantalla."""
        idx = self._norm_idx_actual
        if idx < 0 or idx >= len(self._norm_bloques):
            return
        b = self._norm_bloques[idx]
        b["norm_usuario"] = self._norm_txt_usuario.get("1.0", "end-1c")
        b["norm_ia"]      = self._norm_txt_ia.get("1.0", "end-1c")
        self._marcar_modificado()

    def _norm_version_cambio(self):
        """Callback al cambiar el selector de versión — actualiza ST y muestra info."""
        ver = self._norm_var_version.get()
        ST.norm_version = ver
        _info = {
            "crudo":  "El texto crudo pasará al análisis tal como salió del OCR.",
            "manual": "Tus ediciones manuales pasarán al análisis (crudo si no editaste).",
            "ia":     "La sugerencia IA pasará al análisis (manual → crudo como fallback).",
        }
        if hasattr(self, "_norm_lbl_version_info"):
            self._norm_lbl_version_info.config(text=_info.get(ver, ""))

    def _norm_texto_para_pipeline(self, b: dict) -> str:
        """Devuelve el texto que debe usarse en el pipeline según ST.norm_version."""
        ver = getattr(ST, "norm_version", "manual")
        if ver == "crudo":
            return b["ocr_crudo"]
        if ver == "ia":
            return (b["norm_ia"].strip()
                    or b["norm_usuario"].strip()
                    or b["ocr_crudo"])
        # "manual" (default)
        return b["norm_usuario"].strip() or b["ocr_crudo"]

    # ══════════════════════════════════════════════════════════════════════════
    # VERIFICACIÓN OCR PALABRA POR PALABRA (estilo ABBYY FineReader)
    # ══════════════════════════════════════════════════════════════════════════
    def _verif_abrir(self):
        """Abre el diálogo de verificación sobre la página actualmente
        mostrada en Normalizar. Recalcula las palabras dudosas bajo demanda
        (nada se persiste aparte del texto corregido al cerrar)."""
        idx = self._norm_idx_actual
        if idx < 0 or idx >= len(self._norm_bloques):
            self.toast("Selecciona primero una página en la lista.", "warn")
            return
        img = getattr(self, "_norm_img_orig_full", None)
        if img is None:
            self.toast("No hay imagen cargada para esta página.", "warn")
            return

        b = self._norm_bloques[idx]
        texto_base = self._norm_txt_usuario.get("1.0", "end-1c") or b["ocr_crudo"]

        win, content = self._mk_glass_toplevel("Verificación OCR", 900, 620)
        self._verif_win = win
        self._verif_bloque_idx = idx
        self._verif_texto = texto_base
        self._verif_palabras = []
        self._verif_pos = 0
        self._verif_img_tk = None  # referencia viva contra el GC de Tk
        self._verif_q = queue.Queue()

        info = tk.Label(content, text="Analizando palabras de baja confianza…",
                         bg=CONTENT_BG, fg=TXT_SEC, font=("Segoe UI", 10))
        info.pack(pady=40)
        self._verif_lbl_info = info

        # Layout principal (se puebla cuando el worker entrega resultados)
        cuerpo = tk.Frame(content, bg=CONTENT_BG)
        cuerpo.pack(fill="both", expand=True, padx=16, pady=(0, 12))
        self._verif_cuerpo = cuerpo

        img_frame = tk.Frame(cuerpo, bg="#0D1117", relief="solid", bd=1, height=180)
        img_frame.pack(fill="x", pady=(0, 10))
        img_frame.pack_propagate(False)
        self._verif_lbl_img = tk.Label(img_frame, bg="#0D1117")
        self._verif_lbl_img.pack(expand=True)

        fila_txt = tk.Frame(cuerpo, bg=CONTENT_BG)
        fila_txt.pack(fill="x", pady=(0, 8))
        tk.Label(fila_txt, text="Corrección:", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(side="left")
        self._verif_var_texto = tk.StringVar()
        entry = tk.Entry(fila_txt, textvariable=self._verif_var_texto,
                          font=("Consolas", 11), bg="#0D1117", fg="#CDD6F4",
                          insertbackground="#CDD6F4", relief="solid", bd=1)
        entry.pack(side="left", fill="x", expand=True, padx=(8, 0))
        self._verif_entry = entry

        tk.Label(cuerpo, text="Sugerencias (doble clic para usar):",
                 bg=CONTENT_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(anchor="w")
        self._verif_lb_sug = tk.Listbox(cuerpo, height=4, bg="#0D1117", fg="#CDD6F4",
                                         relief="solid", bd=1, font=("Segoe UI", 9))
        self._verif_lb_sug.pack(fill="x", pady=(2, 10))
        self._verif_lb_sug.bind("<Double-Button-1>", lambda e: self._verif_usar_sugerencia())

        botones = tk.Frame(cuerpo, bg=CONTENT_BG)
        botones.pack(fill="x", pady=(0, 8))
        ttk.Button(botones, text="Omitir", style="S.TButton",
                   command=self._verif_omitir).pack(side="left", padx=(0, 6))
        ttk.Button(botones, text="Omitir todas", style="S.TButton",
                   command=self._verif_omitir_todas).pack(side="left", padx=(0, 6))
        ttk.Button(botones, text="Reemplazar", style="P.TButton",
                   command=self._verif_reemplazar).pack(side="left", padx=(0, 6))
        ttk.Button(botones, text="Reemplazar todas", style="P.TButton",
                   command=self._verif_reemplazar_todas).pack(side="left", padx=(0, 6))
        ttk.Button(botones, text="📖 Agregar a diccionario", style="S.TButton",
                   command=self._verif_agregar_diccionario).pack(side="left", padx=(0, 6))

        pie = tk.Frame(cuerpo, bg=CONTENT_BG)
        pie.pack(fill="x")
        self._verif_lbl_contador = tk.Label(pie, text="", bg=CONTENT_BG, fg=TXT_DIM,
                                             font=("Segoe UI", 9))
        self._verif_lbl_contador.pack(side="left")
        ttk.Button(pie, text="✅ Terminar y guardar", style="P.TButton",
                   command=self._verif_cerrar).pack(side="right")

        cuerpo.pack_forget()  # se muestra al terminar el worker

        threading.Thread(target=self._verif_worker_analizar, args=(img,), daemon=True).start()
        win.after(100, self._verif_poll)
        win.protocol("WM_DELETE_WINDOW", self._verif_cerrar)

    def _verif_worker_analizar(self, img):
        """Corre en thread: extrae palabras dudosas + prepara diccionario de
        corpus para sugerencias. La imagen ya está en memoria (self._norm_img_orig_full,
        cargada una sola vez por _norm_mostrar_imagen) — no se vuelve a leer disco."""
        from core.word_verifier import extraer_palabras_dudosas
        try:
            palabras = extraer_palabras_dudosas(img)
        except Exception as e:
            self._verif_q.put(("error", str(e)))
            return

        dicc_corpus = None
        try:
            if ST.out_dir:
                from core.ocr_normalizer import construir_diccionario_corpus
                cache_path = Path(ST.out_dir) / "diccionario_corpus.json"
                txt_dir = Path(ST.out_dir) / "03_ocr"
                if txt_dir.exists():
                    dicc_corpus = construir_diccionario_corpus(
                        txt_dir, freq_min=3, cache_path=cache_path)
        except Exception:
            dicc_corpus = None
        self._verif_dicc_corpus = dicc_corpus
        self._verif_q.put(("ok", palabras))

    def _verif_poll(self):
        win = getattr(self, "_verif_win", None)
        if win is None or not win.winfo_exists():
            return
        try:
            tipo, payload = self._verif_q.get_nowait()
        except queue.Empty:
            win.after(100, self._verif_poll)
            return

        self._verif_lbl_info.pack_forget()
        if tipo == "error":
            tk.Label(win, text=f"Error: {payload}", bg=CONTENT_BG, fg="#F85149",
                     font=("Segoe UI", 9)).pack(pady=20)
            return

        self._verif_palabras = payload
        self._verif_pos = 0
        self._verif_cuerpo.pack(fill="both", expand=True, padx=16, pady=(0, 12))
        if not self._verif_palabras:
            self._verif_lbl_contador.config(
                text="Sin palabras de baja confianza en esta página. 🎉")
            for w in (self._verif_entry, self._verif_lb_sug):
                w.config(state="disabled")
            return
        self._verif_mostrar_actual()

    def _verif_mostrar_actual(self):
        from core.word_verifier import recortar_palabra, sugerencias_para
        img = self._norm_img_orig_full
        p = self._verif_palabras[self._verif_pos]

        recorte = recortar_palabra(img, p, margen=8, zoom=2.5)
        recorte.thumbnail((820, 160))
        from PIL import ImageTk
        self._verif_img_tk = ImageTk.PhotoImage(recorte)
        self._verif_lbl_img.config(image=self._verif_img_tk)

        self._verif_var_texto.set(p.texto)

        from core.spell_corrector import obtener_corrector
        corrector = obtener_corrector()
        corrector._cargar_diccionario()
        sugerencias = sugerencias_para(p.texto, corrector, getattr(self, "_verif_dicc_corpus", None))
        self._verif_lb_sug.delete(0, "end")
        for s in sugerencias:
            self._verif_lb_sug.insert("end", s)

        n = len(self._verif_palabras)
        self._verif_lbl_contador.config(
            text=f"Palabra {self._verif_pos + 1} de {n} — confianza {p.conf:.0f}%\n"
                 f"Contexto: …{p.contexto}…")

    def _verif_usar_sugerencia(self):
        sel = self._verif_lb_sug.curselection()
        if sel:
            self._verif_var_texto.set(self._verif_lb_sug.get(sel[0]))

    def _verif_avanzar(self):
        if self._verif_pos + 1 < len(self._verif_palabras):
            self._verif_pos += 1
            self._verif_mostrar_actual()
        else:
            self._verif_lbl_contador.config(text="✅ Última palabra revisada.")
            for w in (self._verif_entry, self._verif_lb_sug):
                w.config(state="disabled")

    def _verif_omitir(self):
        self._verif_avanzar()

    def _verif_omitir_todas(self):
        self._verif_lbl_contador.config(text="✅ Verificación cerrada sin más cambios.")
        for w in (self._verif_entry, self._verif_lb_sug):
            w.config(state="disabled")

    def _verif_reemplazar(self):
        from core.word_verifier import aplicar_reemplazo
        p = self._verif_palabras[self._verif_pos]
        nuevo_valor = self._verif_var_texto.get()
        texto, encontrada = aplicar_reemplazo(
            self._verif_texto, p.texto, nuevo_valor, p.idx_ocurrencia)
        self._verif_texto = texto
        if not encontrada:
            self.toast(f"«{p.texto}» (ocurrencia {p.idx_ocurrencia + 1}) no se "
                       "localizó en el texto — probablemente ya fue editado.", "warn")
        self._verif_avanzar()

    def _verif_reemplazar_todas(self):
        from core.word_verifier import reemplazar_todas
        p = self._verif_palabras[self._verif_pos]
        nuevo_valor = self._verif_var_texto.get()
        texto, n = reemplazar_todas(self._verif_texto, p.texto, nuevo_valor)
        self._verif_texto = texto
        self.toast(f"{n} ocurrencia(s) de «{p.texto}» reemplazadas.", "ok")
        self._verif_avanzar()

    def _verif_agregar_diccionario(self):
        from core.spell_corrector import obtener_corrector
        p = self._verif_palabras[self._verif_pos]
        obtener_corrector().agregar_palabra_usuario(p.texto)
        self.toast(f"«{p.texto}» agregada al vocabulario de usuario.", "ok")
        self._verif_avanzar()

    def _verif_cerrar(self):
        """Vuelca el texto corregido al bloque actual del panel Normalizar
        (mismo flujo de guardado que ya existe: _norm_guardar_bloque_actual
        → _norm_guardar → UPSERT en SQLite). El verificador nunca escribe
        directamente a disco/BD — un solo escritor."""
        win = getattr(self, "_verif_win", None)
        if getattr(self, "_verif_texto", None) is not None and \
           self._verif_bloque_idx == self._norm_idx_actual:
            self._norm_txt_usuario.delete("1.0", "end")
            self._norm_txt_usuario.insert("1.0", self._verif_texto)
            self._norm_guardar_bloque_actual()
        if win is not None and win.winfo_exists():
            win.destroy()
        self._verif_win = None

    def _norm_guardar(self):
        """Persiste todas las ediciones en SQLite y en los archivos .txt."""
        self._norm_guardar_bloque_actual()
        db_path = Path(ST.ruta_db) if ST.ruta_db else None
        ver     = getattr(ST, "norm_version", "manual")
        guardados = n_crudo = n_manual = n_ia = 0

        for b in self._norm_bloques:
            texto_final = self._norm_texto_para_pipeline(b)
            # Contadores por versión efectiva usada
            if ver == "crudo" or not b["norm_usuario"].strip():
                n_crudo += 1
            elif ver == "ia" and b["norm_ia"].strip():
                n_ia += 1
            else:
                n_manual += 1
            try:
                Path(b["txt_path"]).write_text(texto_final, encoding="utf-8")
                guardados += 1
            except Exception:
                pass
            # Persistir todas las versiones en SQLite para no perder nada
            if db_path:
                self._norm_escribir_db(db_path, b["numero"], b["pagina"],
                                       b["ocr_crudo"], b["norm_usuario"], b["norm_ia"])

        ST.norm_done   = True
        ST.norm_version = ver
        # Las etapas posteriores deben re-ejecutarse con el nuevo texto
        ST.marcar_etapa("norm", "ready")
        self._actualizar_badges()

        _etiq = {"crudo": "crudo", "manual": "manual", "ia": "IA"}
        detalle = f"  ({n_manual} manual · {n_ia} IA · {n_crudo} crudo)"
        self._lbl_norm_estado.config(
            text=f"✅ {guardados} páginas guardadas como {_etiq[ver]}{detalle}")

    def _norm_reocr_pagina(self):
        """Re-extrae el texto de la página actual con Tesseract y actualiza el bloque."""
        idx = self._norm_idx_actual
        if idx < 0 or idx >= len(self._norm_bloques):
            messagebox.showwarning("Sin selección", "Selecciona una página primero.")
            return
        b = self._norm_bloques[idx]
        num, pagina = b["numero"], b["pagina"]

        # Buscar imagen original a color
        img_path = None
        if ST.out_dir:
            img_dir = Path(ST.out_dir) / "02_imagenes" / num
            for ext in ("*.png", "*.jpg", "*.tif", "*.tiff"):
                hits = sorted(img_dir.glob(f"*{pagina}*")) if img_dir.exists() else []
                if hits:
                    img_path = hits[0]
                    break
        if img_path is None:
            messagebox.showwarning("Sin imagen",
                "No se encontró la imagen original para esta página.\n"
                "Ejecuta primero la extracción OCR con Ruta 1.")
            return

        self._lbl_norm_estado.config(text=f"⏳ Re-OCR Tesseract: {pagina}…", fg="#F59E0B")
        self.update_idletasks()

        def _run():
            try:
                # Por zonas si la página tiene etiquetas guardadas; si no, completa
                from core.layout_tesseract import ocr_pagina_con_zonas
                texto, conf, con_z = ocr_pagina_con_zonas(
                    img_path, ST.out_dir, num, pagina, lang="spa")
                # Actualizar el bloque en memoria y en disco
                b["ocr_crudo"]    = texto
                b["norm_usuario"] = ""   # limpiar edición anterior (era basura)
                txt_path = Path(b["txt_path"])
                txt_path.write_text(texto, encoding="utf-8")
                _modo = " · por zonas" if con_z else ""
                self.after(0, lambda: (
                    self._norm_mostrar_bloque(idx),
                    self._lbl_norm_estado.config(
                        text=f"✅ {pagina} re-extraído (conf: {conf}%{_modo})", fg=VERDE),
                    self._norm_refrescar_lista(),
                ))
            except Exception as e:
                self.after(0, lambda err=str(e): self._lbl_norm_estado.config(
                    text=f"⚠ Error: {err}", fg=ROJO))

        threading.Thread(target=_run, daemon=True).start()

    def _norm_regenerar_imagenes(self):
        """Regenera las imágenes del número desde el PDF original, respetando rotación."""
        num = self._norm_var_numero.get()
        if not num or not ST.out_dir:
            return
        # Buscar el PDF original
        pdf_path = None
        for archivo in getattr(ST, "archivos_sel", []):
            if archivo.stem == num:
                pdf_path = archivo
                break
        if pdf_path is None:
            # Buscar en carpeta de entrada configurada
            if ST.pdf_dir:
                hits = list(Path(ST.pdf_dir).glob(f"{num}*.pdf"))
                if hits:
                    pdf_path = hits[0]
        if pdf_path is None:
            messagebox.showwarning("Sin PDF",
                f"No se encontró el PDF original para '{num}'.\n"
                "Asegúrate de que la carpeta de entrada está configurada.")
            return

        if not messagebox.askyesno("Regenerar imágenes",
            f"Se regenerarán las imágenes de '{num}' desde el PDF original.\n\n"
            "Esto corrige imágenes rotadas o en mal formato.\n"
            "Las imágenes anteriores se sobreescribirán.\n\n"
            "¿Continuar?"):
            return

        img_dir = Path(ST.out_dir) / "02_imagenes" / num
        self._lbl_norm_estado.config(text="⏳ Regenerando imágenes…", fg="#F59E0B")

        def _run():
            try:
                # Eliminar imágenes anteriores para forzar regeneración
                if img_dir.exists():
                    for f in img_dir.glob("*.png"):
                        f.unlink()
                img_dir.mkdir(parents=True, exist_ok=True)

                from core.ocr_engine import pdf_a_imagenes
                dpi = 150  # DPI recomendado
                imgs = pdf_a_imagenes(pdf_path, img_dir, dpi)

                self.after(0, lambda: (
                    self._lbl_norm_estado.config(
                        text=f"✅ {len(imgs)} imágenes regeneradas — {num}", fg=VERDE),
                    self._norm_mostrar_bloque(self._norm_idx_actual),
                ))
            except Exception as e:
                self.after(0, lambda err=str(e): self._lbl_norm_estado.config(
                    text=f"⚠ Error: {err}", fg=ROJO))

        threading.Thread(target=_run, daemon=True).start()

    def _norm_reocr_numero(self):
        """Re-extrae el texto de TODAS las páginas del número con Tesseract."""
        num = self._norm_var_numero.get()
        if not num or not ST.out_dir:
            return
        if not messagebox.askyesno("Re-OCR completo",
            f"Se re-extraerá el texto de {len(self._norm_bloques)} páginas con Tesseract.\n\n"
            "Esto sobreescribirá los .txt actuales (incluyendo texto basura de Kraken).\n"
            "¿Continuar?"):
            return

        self._lbl_norm_estado.config(text="⏳ Re-OCR Tesseract en curso…", fg="#F59E0B")

        def _run():
            from core.layout_tesseract import ocr_pagina_con_zonas
            total = len(self._norm_bloques)
            ok = 0
            for i, b in enumerate(self._norm_bloques):
                # Buscar imagen
                img_path = None
                img_dir = Path(ST.out_dir) / "02_imagenes" / b["numero"]
                for ext in ("*.png","*.jpg","*.tif","*.tiff"):
                    hits = sorted(img_dir.glob(f"*{b['pagina']}*")) if img_dir.exists() else []
                    if hits:
                        img_path = hits[0]; break

                if img_path is None:
                    continue
                try:
                    texto, conf, _cz = ocr_pagina_con_zonas(
                        img_path, ST.out_dir, b["numero"], b["pagina"], lang="spa")
                    b["ocr_crudo"]    = texto
                    b["norm_usuario"] = ""
                    Path(b["txt_path"]).write_text(texto, encoding="utf-8")
                    ok += 1
                except Exception:
                    pass

                n = i + 1
                self.after(0, lambda n=n, total=total: self._lbl_norm_estado.config(
                    text=f"⏳ Re-OCR {n}/{total}…", fg="#F59E0B"))

            self.after(0, lambda: (
                self._norm_cargar_numero(),
                self._lbl_norm_estado.config(
                    text=f"✅ {ok}/{total} páginas re-extraídas con Tesseract", fg=VERDE),
            ))

        threading.Thread(target=_run, daemon=True).start()

    def _norm_importar_txt(self):
        """
        Importa un .txt externo (exportado desde Acrobat u otro programa)
        para la página actualmente seleccionada.
        El texto importado reemplaza el OCR crudo y se coloca en la vista usuario.
        """
        idx = self._norm_idx_actual
        if idx < 0 or idx >= len(self._norm_bloques):
            # Si no hay página seleccionada, ofrecer importar para todo el número
            self._norm_importar_txt_numero()
            return

        b = self._norm_bloques[idx]
        ruta = filedialog.askopenfilename(
            title=f"Importar texto para {b['pagina']}",
            filetypes=[("Archivos de texto", "*.txt"), ("Todos los archivos", "*.*")],
        )
        if not ruta:
            return

        try:
            texto = Path(ruta).read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            messagebox.showerror("Error al leer", str(e))
            return

        # Guardar como ocr_crudo (fuente externa) y limpiar edición previa
        b["ocr_crudo"]    = texto
        b["norm_usuario"] = ""
        Path(b["txt_path"]).write_text(texto, encoding="utf-8")

        self._norm_mostrar_bloque(idx)
        self._lbl_norm_estado.config(
            text=f"✅ Texto importado en {b['pagina']} ({len(texto.split())} palabras)",
            fg=VERDE)

    def _norm_importar_txt_numero(self):
        """
        Importa una carpeta completa de .txt para reemplazar el OCR de un número.
        Útil cuando se exportó el texto desde Acrobat página por página.
        El nombre de cada archivo debe coincidir con el stem de la página (p0001, p0002, etc.)
        o ser numerado (1.txt, 2.txt, etc.).
        """
        num = self._norm_var_numero.get()
        if not num:
            messagebox.showwarning("Sin número", "Selecciona un número primero.")
            return

        carpeta = filedialog.askdirectory(
            title="Selecciona la carpeta con los archivos .txt exportados de Acrobat")
        if not carpeta:
            return

        txts = sorted(Path(carpeta).glob("*.txt"))
        if not txts:
            messagebox.showwarning("Sin archivos",
                "No se encontraron archivos .txt en esa carpeta.")
            return

        if not messagebox.askyesno("Importar textos externos",
            f"Se importarán {len(txts)} archivos .txt desde:\n{carpeta}\n\n"
            "Reemplazarán el OCR actual de cada página. ¿Continuar?"):
            return

        importados = 0
        for b in self._norm_bloques:
            # Buscar por nombre exacto (p0001.txt) o por número (1.txt, 01.txt)
            import re as _re
            m = _re.search(r'(\d+)', b["pagina"])
            n = int(m.group(1)) if m else -1
            candidatos = [
                Path(carpeta) / f"{b['pagina']}.txt",
                Path(carpeta) / f"{n}.txt",
                Path(carpeta) / f"{n:02d}.txt",
                Path(carpeta) / f"{n:03d}.txt",
                Path(carpeta) / f"{n:04d}.txt",
            ] + [t for t in txts if b["pagina"] in t.stem]

            for cand in candidatos:
                if cand.exists():
                    try:
                        texto = cand.read_text(encoding="utf-8", errors="replace")
                        b["ocr_crudo"]    = texto
                        b["norm_usuario"] = ""
                        Path(b["txt_path"]).write_text(texto, encoding="utf-8")
                        importados += 1
                    except Exception:
                        pass
                    break

        self._norm_cargar_numero()
        self._lbl_norm_estado.config(
            text=f"✅ {importados}/{len(self._norm_bloques)} páginas importadas desde carpeta",
            fg=VERDE)

    def _norm_reconstruir_columnas(self):
        """
        Aplica la reconstrucción de líneas rotas (algoritmo BNC) al bloque actual
        o a todo el número. Corrige el problema de columnas mezcladas del texto BNC:
        une líneas cortas que pertenecen al mismo párrafo y separa las de columnas distintas.
        """
        idx = self._norm_idx_actual
        alcance = "pagina" if (idx >= 0 and idx < len(self._norm_bloques)) else "numero"

        if alcance == "pagina":
            resp = messagebox.askyesnocancel(
                "Reconstruir columnas",
                "¿Reconstruir columnas solo en esta página (Sí)\n"
                "o en todo el número (No)?")
            if resp is None:
                return
            alcance = "pagina" if resp else "numero"

        from core.ocr_normalizer import reconstruir_lineas_rotas

        if alcance == "pagina":
            b = self._norm_bloques[idx]
            base = b["norm_usuario"] if b["norm_usuario"] else b["ocr_crudo"]
            reconstruido = reconstruir_lineas_rotas(base)
            b["norm_usuario"] = reconstruido
            self._norm_mostrar_bloque(idx)
            palabras = len(reconstruido.split())
            self._lbl_norm_estado.config(
                text=f"✅ Columnas reconstruidas en {b['pagina']} ({palabras} palabras)",
                fg=VERDE)
        else:
            n = 0
            for b in self._norm_bloques:
                base = b["norm_usuario"] if b["norm_usuario"] else b["ocr_crudo"]
                if base.strip():
                    b["norm_usuario"] = reconstruir_lineas_rotas(base)
                    n += 1
            self._norm_refrescar_lista()
            self._lbl_norm_estado.config(
                text=f"✅ Columnas reconstruidas en {n} páginas",
                fg=VERDE)

    def _norm_diccionario_corpus(self):
        """Construye el diccionario de frecuencias del corpus completo y lo guarda en JSON."""
        if not ST.out_dir:
            messagebox.showwarning("Sin corpus",
                                   "Carga primero un corpus para construir el diccionario.")
            return

        txt_dir = Path(ST.out_dir) / "03_ocr"
        if not txt_dir.exists():
            messagebox.showwarning("Sin textos OCR",
                                   "No se encontró la carpeta 03_ocr/.\n"
                                   "Ejecuta primero el paso Extracción OCR o el Conversor PDF.")
            return

        cache_path = Path(ST.out_dir) / "diccionario_corpus.json"

        def _run():
            from core.ocr_normalizer import construir_diccionario_corpus
            self.after(0, lambda: self._lbl_norm_estado.config(
                text="⏳ Construyendo diccionario de corpus…", fg=TXT_SEC))

            def _cb(n, total, nombre):
                self.after(0, lambda: self._lbl_norm_estado.config(
                    text=f"⏳ Procesando {n}/{total}: {nombre}", fg=TXT_SEC))

            try:
                dic = construir_diccionario_corpus(
                    txt_dir, freq_min=3, cache_path=cache_path, callback=_cb)
                n_palabras = len(dic)
                top5 = sorted(dic.items(), key=lambda x: -x[1])[:5]
                top5_str = ", ".join(f"{p}({f})" for p, f in top5)
                self.after(0, lambda: self._lbl_norm_estado.config(
                    text=f"✅ Diccionario listo: {n_palabras:,} palabras · top: {top5_str}",
                    fg=VERDE))
                self.after(0, lambda: messagebox.showinfo(
                    "Diccionario de corpus",
                    f"Diccionario construido con {n_palabras:,} palabras (freq ≥ 3).\n\n"
                    f"Top 5: {top5_str}\n\n"
                    f"Guardado en:\n{cache_path}"))
            except Exception as ex:
                self.after(0, lambda err=str(ex): self._lbl_norm_estado.config(
                    text=f"❌ Error: {err}", fg="#F85149"))

        threading.Thread(target=_run, daemon=True).start()

    # ── Dictado por voz ───────────────────────────────────────────────────────

    def _norm_dictar_toggle(self):
        """Inicia o detiene la sesión de dictado por voz."""
        if self._dictar_session is not None:
            # Detener sesión activa
            self._dictar_session.detener()
            self._dictar_session = None
            self._btn_dictar.config(text="🎙 Dictar")
            self._lbl_dictar_estado.config(text="", fg=TXT_DIM)
            return

        try:
            from core.voice_dictation import DictadoSession
        except ImportError:
            messagebox.showerror("Dependencia faltante",
                                 "Instala las dependencias de dictado:\n\n"
                                 "  pip install SpeechRecognition sounddevice\n\n"
                                 "Luego reinicia la aplicación.")
            return

        self._btn_dictar.config(text="⏹ Detener")
        self._lbl_dictar_estado.config(text="⏳ Iniciando micrófono…", fg=TXT_SEC)

        def _on_texto(texto: str):
            # Llamado desde hilo de audio — usar after() para acceder a tkinter
            self.after(0, lambda t=texto: self._norm_dictar_insertar(t))

        self._dictar_session = DictadoSession(
            callback=_on_texto,
            idioma="es-CO",
            modo_online=True,
        )
        self._dictar_session.iniciar()
        # Iniciar polling de estados del hilo de audio
        self.after(200, self._norm_dictar_poll)

    def _norm_dictar_insertar(self, texto: str):
        """Inserta texto transcrito en el textarea de usuario, con espacio separador."""
        if not texto.strip():
            return
        widget = self._norm_txt_usuario
        # Posición actual del cursor; si no hay cursor, insertar al final
        try:
            pos = widget.index("insert")
        except Exception:
            pos = "end"
        # Añadir espacio si el texto previo no termina en espacio o salto
        contenido_actual = widget.get("1.0", pos)
        if contenido_actual and contenido_actual[-1] not in (" ", "\n"):
            texto = " " + texto
        widget.insert(pos, texto)
        widget.see("insert")
        # Indicador visual del último fragmento reconocido
        preview = texto.strip()[:40] + ("…" if len(texto.strip()) > 40 else "")
        self._lbl_dictar_estado.config(
            text=f"🔴 Escuchando · '{preview}'", fg="#F85149")

    def _norm_dictar_poll(self):
        """Polling de mensajes de estado del hilo de dictado (cada 200 ms)."""
        if self._dictar_session is None:
            return
        estado = self._dictar_session.estado()
        if estado:
            if estado == "escuchando":
                self._lbl_dictar_estado.config(text="🔴 Escuchando…", fg="#F85149")
            elif estado == "detenido":
                self._btn_dictar.config(text="🎙 Dictar")
                self._lbl_dictar_estado.config(text="", fg=TXT_DIM)
                self._dictar_session = None
                return
            elif estado.startswith("error:"):
                msg = estado[6:]
                self._btn_dictar.config(text="🎙 Dictar")
                self._lbl_dictar_estado.config(text=f"⚠ {msg}", fg="#F59E0B")
                self._dictar_session = None
                return
        # Continuar polling mientras la sesión esté activa
        self.after(200, self._norm_dictar_poll)

    def _norm_auto(self):
        """Aplica normalización automática a todos los bloques no editados."""
        def _run():
            from core.text_postprocessor import normalizar_bloque
            for b in self._norm_bloques:
                if not b["norm_usuario"]:
                    b["norm_usuario"] = normalizar_bloque(b["ocr_crudo"])
            self.after(0, self._norm_refrescar_lista)
            self.after(0, lambda: self._lbl_norm_estado.config(
                text=f"✅ Normalización automática aplicada a {len(self._norm_bloques)} bloques"))
        threading.Thread(target=_run, daemon=True).start()

    def _norm_ia(self):
        """Solicita sugerencia de IA para el bloque actual."""
        if not ST.ia_habilitada:
            messagebox.showwarning("IA deshabilitada",
                                   "Activa la IA desde el switch en la barra superior.")
            return
        idx = self._norm_idx_actual
        if idx < 0 or idx >= len(self._norm_bloques):
            return
        b = self._norm_bloques[idx]
        texto_base = b["norm_usuario"] or b["ocr_crudo"]

        api_key = _resolver_api_key_modelo("ocr_mejora")[0]
        if not api_key:
            messagebox.showwarning("Sin API key",
                "No hay clave de API configurada.\n\n"
                "Andá a ⚙ Configuración → claves de API\n"
                "y pegá tu clave de Anthropic, OpenAI o Gemini.")
            return

        def _run():
            try:
                from core.ocr_llm import corregir_texto
                sugerencia = corregir_texto(texto_base, api_key)
            except Exception as exc:
                sugerencia = f"[Error: {exc}]"
                self.after(0, lambda m=str(exc): messagebox.showerror(
                    "Error IA", f"No se pudo obtener sugerencia:\n{m}"))
            b["norm_ia"] = sugerencia
            self.after(0, lambda: (
                self._norm_txt_ia.delete("1.0", "end"),
                self._norm_txt_ia.insert("1.0", sugerencia),
            ))
        threading.Thread(target=_run, daemon=True).start()

    # ── Zoom / pan del canvas de imagen en Normalizar ─────────────────────────

    def _norm_on_zoom(self, event):
        factor = 1.15 if event.delta > 0 else 1 / 1.15
        self._norm_zoom = max(0.15, min(6.0, self._norm_zoom * factor))
        self._norm_aplicar_zoom()

    def _norm_zoom_step(self, direction: int):
        factor = 1.25 if direction > 0 else 1 / 1.25
        self._norm_zoom = max(0.15, min(6.0, self._norm_zoom * factor))
        self._norm_aplicar_zoom()

    def _norm_aplicar_zoom(self):
        if self._norm_img_orig_full is None:
            return
        from PIL import ImageTk
        new_w = max(1, int(self._norm_img_orig_full.width  * self._norm_zoom))
        new_h = max(1, int(self._norm_img_orig_full.height * self._norm_zoom))
        img_r = self._norm_img_orig_full.resize((new_w, new_h), 1)  # 1 = LANCZOS
        self._norm_canvas_img._img_ref = ImageTk.PhotoImage(img_r)
        self._norm_canvas_img.delete("all")
        self._norm_canvas_img.create_image(0, 0, anchor="nw",
                                            image=self._norm_canvas_img._img_ref)
        self._norm_canvas_img.configure(scrollregion=(0, 0, new_w, new_h))
        pct = int(self._norm_zoom * 100)
        if hasattr(self, "_norm_lbl_zoom"):
            self._norm_lbl_zoom.config(text=f"{pct}%")

    def _norm_pan_start_cb(self, event):
        self._norm_pan_start = (event.x, event.y)
        self._norm_canvas_img.config(cursor="fleur")

    def _norm_pan_drag_cb(self, event):
        if self._norm_pan_start is None:
            return
        dx = self._norm_pan_start[0] - event.x
        dy = self._norm_pan_start[1] - event.y
        self._norm_pan_start = (event.x, event.y)
        self._norm_canvas_img.xview_scroll(int(dx), "units")
        self._norm_canvas_img.yview_scroll(int(dy), "units")

    def _norm_pan_end_cb(self, event):
        self._norm_pan_start = None
        self._norm_canvas_img.config(cursor="")

    def _norm_refrescar_lista(self):
        """Refresca los indicadores ✓/⚠ en la lista de bloques."""
        avisos_ocr = getattr(self, "_avisos_ocr", {})
        for i, b in enumerate(self._norm_bloques):
            estado = "✓" if b["norm_usuario"] else "·"
            alerta = " ⚠" if (b["numero"], b["pagina"]) in avisos_ocr else ""
            self._norm_lb.delete(i)
            self._norm_lb.insert(i, f"{estado} {b['pagina']}{alerta}")
        if self._norm_idx_actual >= 0:
            self._norm_lb.selection_set(self._norm_idx_actual)
            self._norm_mostrar_bloque(self._norm_idx_actual)

    def _norm_leer_db(self, db_path, numero: str, pagina: str):
        """Lee norm_usuario y norm_ia de SQLite. Retorna (None, None) si no existe."""
        if not db_path or not db_path.exists():
            return None, None
        try:
            import sqlite3
            con = sqlite3.connect(str(db_path))
            cur = con.execute(
                "SELECT norm_usuario, norm_ia FROM normalizaciones "
                "WHERE numero=? AND pagina=? LIMIT 1", (numero, pagina))
            row = cur.fetchone()
            con.close()
            return (row[0], row[1]) if row else (None, None)
        except Exception:
            return None, None

    def _norm_escribir_db(self, db_path, numero: str, pagina: str,
                           ocr_crudo: str, norm_usuario: str, norm_ia: str):
        """Inserta o actualiza la fila en la tabla normalizaciones de SQLite."""
        try:
            import sqlite3
            con = sqlite3.connect(str(db_path))
            con.execute("""CREATE TABLE IF NOT EXISTS normalizaciones (
                numero       TEXT NOT NULL,
                pagina       TEXT NOT NULL,
                ocr_crudo    TEXT,
                norm_usuario TEXT,
                norm_ia      TEXT,
                ts_usuario   TEXT,
                ts_ia        TEXT,
                PRIMARY KEY (numero, pagina)
            )""")
            from datetime import datetime
            ts = datetime.now().isoformat(timespec="seconds")
            con.execute("""INSERT INTO normalizaciones
                (numero, pagina, ocr_crudo, norm_usuario, norm_ia, ts_usuario, ts_ia)
                VALUES (?,?,?,?,?,?,?)
                ON CONFLICT(numero, pagina) DO UPDATE SET
                    ocr_crudo=excluded.ocr_crudo,
                    norm_usuario=excluded.norm_usuario,
                    norm_ia=excluded.norm_ia,
                    ts_usuario=excluded.ts_usuario,
                    ts_ia=excluded.ts_ia
            """, (numero, pagina, ocr_crudo, norm_usuario, norm_ia, ts, ts))
            con.commit()
            con.close()
        except Exception:
            pass

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 3: SEGMENTACIÓN DE ARTÍCULOS
    # ══════════════════════════════════════════════════════════════════════════
    def _build_seg(self):
        f = self._tab_seg
        self._page_header(f, "Segmentación de artículos",
                          "Identifica artículos y asigna autoría por bylines y firmas", "📝")
        self._build_ai_panel(f, "seg")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True, padx=24, pady=16)

        # ── Barra de acción fija (siempre visible) ────────────────────────────
        bf = tk.Frame(pad, bg=CONTENT_BG); bf.pack(fill="x", pady=(0, 8))
        self._btn_seg = ttk.Button(bf, text="▶  Segmentar artículos",
                                    style="P.TButton", command=self._start_seg)
        self._btn_seg.pack(side="left", padx=(0,12))
        self._var_seg_v2 = tk.BooleanVar(value=True)
        ttk.Checkbutton(bf, text="Segmentador avanzado (v2)",
                        variable=self._var_seg_v2).pack(side="left", padx=(0, 12))
        self._lbl_seg_n = tk.Label(bf, text="", bg=CONTENT_BG, fg=VERDE,
                                    font=("Segoe UI",10,"bold"))
        self._lbl_seg_n.pack(side="left", padx=8)
        ttk.Button(bf, text="💾 Exportar CSV", style="S.TButton",
                   command=self._export_seg_csv).pack(side="left", padx=8)
        ttk.Button(bf, text="📄 Exportar DOCX (con costura)", style="S.TButton",
                   command=self._export_seg_docx_costura).pack(side="left", padx=4)
        tk.Label(bf, text="⚠  Requiere extracción completada",
                 bg=CONTENT_BG, fg=ACENT, font=("Segoe UI",9)).pack(side="left", padx=8)
        ttk.Button(bf, text="📓 Nota", style="S.TButton",
                   command=lambda: self._bitacora_nueva_nota("seg")).pack(side="right")

        # Progressive disclosure — opciones avanzadas de segmentación
        def _build_seg_avanzado(f):
            row1 = tk.Frame(f, bg=CONTENT_BG); row1.pack(fill="x", pady=2)
            tk.Label(row1, text="Umbral confianza autoría:", bg=CONTENT_BG,
                     fg=TXT_SEC, font=("Segoe UI",8)).pack(side="left")
            self._var_seg_umbral = tk.DoubleVar(value=0.4)
            ttk.Scale(row1, from_=0.1, to=0.9, variable=self._var_seg_umbral,
                      orient="horizontal", length=120).pack(side="left", padx=6)
            tk.Label(row1, textvariable=self._var_seg_umbral, bg=CONTENT_BG,
                     fg=TXT_SEC, font=("Segoe UI",8), width=4).pack(side="left")
            row2 = tk.Frame(f, bg=CONTENT_BG); row2.pack(fill="x", pady=2)
            self._var_seg_max_art = tk.IntVar(value=0)
            tk.Label(row2, text="Máx. artículos por número (0=sin límite):",
                     bg=CONTENT_BG, fg=TXT_SEC, font=("Segoe UI",8)).pack(side="left")
            ttk.Spinbox(row2, from_=0, to=200, textvariable=self._var_seg_max_art,
                        width=5).pack(side="left", padx=6)
        self._mk_avanzado(pad, "Opciones avanzadas de segmentación", _build_seg_avanzado)

        cols = ("Número", "Título", "Autor", "Confianza", "Sección", "Páginas", "Palabras")
        anchos = {"Número":140,"Título":340,"Autor":180,"Confianza":75,
                  "Sección":110,"Páginas":110,"Palabras":75}
        tv_outer = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1)
        tv_outer.pack(fill="both", expand=True)
        self._tv_seg_outer = tv_outer   # para skeleton
        self._seg_skeleton = None
        sb_y = ttk.Scrollbar(tv_outer, orient="vertical")
        sb_x = ttk.Scrollbar(tv_outer, orient="horizontal")
        self._tv_seg = ttk.Treeview(tv_outer, columns=cols, show="headings",
                                     yscrollcommand=sb_y.set, xscrollcommand=sb_x.set, height=13)
        for col in cols:
            self._tv_seg.heading(col, text=col, anchor="w",
                                  command=lambda c=col: self._sort_tv_seg(c))
            self._tv_seg.column(col, width=anchos[col], minwidth=50, stretch=False)
        self._tv_seg.tag_configure("alta",    background="#1A3A2A", foreground="#3FB950")
        self._tv_seg.tag_configure("media",   background="#2D2210", foreground="#F0883E")
        self._tv_seg.tag_configure("anonimo", background="#1C2128", foreground="#8B949E")
        sb_y.config(command=self._tv_seg.yview)
        sb_x.config(command=self._tv_seg.xview)
        sb_y.pack(side="right", fill="y")
        sb_x.pack(side="bottom", fill="x")
        self._tv_seg.pack(fill="both", expand=True)
        self._tv_seg.bind("<<TreeviewSelect>>", self._on_seg_select)
        self._tv_seg_sort_rev = {c: False for c in cols}

        ley_f = tk.Frame(pad, bg=CONTENT_BG); ley_f.pack(anchor="w", pady=(4, 0))
        for bg_, fg_, txt in [("#F0FDF4","#166534","  ✓ Autoría identificada  "),
                               ("#FFFBEB","#92400E","  ≈ Confianza media  "),
                               ("#F8FAFC","#64748B","  — Anónimo  ")]:
            tk.Label(ley_f, text=txt, bg=bg_, fg=fg_,
                     font=("Segoe UI",8), relief="solid", bd=1).pack(side="left", padx=3)

        det_outer = tk.Frame(pad, bg=CARD_BOR, relief="solid", bd=1)
        det_outer.pack(fill="x", pady=(10,0))
        det_hdr = tk.Frame(det_outer, bg="#1C2128"); det_hdr.pack(fill="x")
        tk.Label(det_hdr, text="  📄  Texto del artículo seleccionado",
                 bg="#1C2128", fg=TXT_PRI, font=("Segoe UI",8,"bold")).pack(side="left", pady=4)
        self._txt_seg_art = scrolledtext.ScrolledText(det_outer, height=6, font=("Consolas",9), bg="#0D1117", fg="#CDD6F4", relief="flat", state="disabled", wrap="word")
        self._txt_seg_art.pack(fill="x", padx=1, pady=(0,1))


    def _sort_tv_seg(self, col: str):
        """Ordena la tabla de artículos por columna al hacer clic en el encabezado."""
        if ST.df_articulos is None or ST.df_articulos.empty: return
        rev = self._tv_seg_sort_rev.get(col, False)
        self._tv_seg_sort_rev[col] = not rev
        col_map = {"Número":"numero","Título":"titulo","Autor":"autor",
                   "Confianza":"confianza_autor","Sección":"seccion",
                   "Páginas":"pagina","Palabras":"palabras"}
        df_col = col_map.get(col)
        if df_col and df_col in ST.df_articulos.columns:
            df_sorted = ST.df_articulos.sort_values(df_col, ascending=not rev)
            self._tv_seg.delete(*self._tv_seg.get_children())
            for _, row in df_sorted.iterrows():
                self._insertar_fila_seg(row)

    def _insertar_fila_seg(self, row):
        """Inserta una fila en la tabla de artículos con color según confianza."""
        conf = float(row.get("confianza_autor", 0))
        if conf >= 0.65:   tag = "alta"
        elif conf >= 0.30: tag = "media"
        else:              tag = "anonimo"
        titulo_corto = str(row.get("titulo",""))[:80]
        vals = (
            str(row.get("numero",""))[:30],
            titulo_corto,
            str(row.get("autor",""))[:50],
            f"{conf:.2f}",
            str(row.get("seccion",""))[:20],
            str(row.get("pagina",""))[:15],
            str(row.get("palabras",""))
        )
        self._tv_seg.insert("", "end", values=vals, tags=(tag,))

    def _on_seg_select(self, _event=None):
        sel = self._tv_seg.selection()
        if not sel or ST.df_articulos is None: return
        item = self._tv_seg.item(sel[0])
        vals = item["values"]
        if not vals: return
        mask = ((ST.df_articulos["numero"] == str(vals[0])) &
                (ST.df_articulos["titulo"] == str(vals[1])))
        rows = ST.df_articulos[mask]
        texto = rows.iloc[0]["texto"] if not rows.empty else "(no disponible)"
        self._txt_seg_art.config(state="normal")
        self._txt_seg_art.delete("1.0","end")
        self._txt_seg_art.insert("1.0", texto[:3000])
        self._txt_seg_art.config(state="disabled")

    def _export_seg_csv(self):
        if ST.df_articulos is None or ST.df_articulos.empty:
            messagebox.showwarning("Sin datos","Ejecuta la segmentación primero."); return
        dest = filedialog.asksaveasfilename(defaultextension=".csv",
               filetypes=[("CSV","*.csv")], initialfile="articulos_segmentados.csv")
        if dest:
            ST.df_articulos.drop(columns=["texto"],errors="ignore").to_csv(dest, index=False, encoding="utf-8-sig")
            self.toast(f"CSV guardado → {Path(dest).name}", tipo="ok")

    def _export_seg_docx_costura(self):
        """Exporta artículos a DOCX con palabras de costura marcadas en rojo."""
        articulos = getattr(ST, "articulos", None) or []
        if not articulos:
            messagebox.showwarning("Sin datos", "Segmenta el corpus primero."); return
        dest = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word DOCX", "*.docx")],
            initialfile="corpus_con_costura.docx",
            title="Exportar DOCX con marcas de costura")
        if not dest:
            return
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor

            from core.gutter_completion import RE_GENERADO, exportar_docx_con_marcas
        except ImportError:
            messagebox.showerror("Falta python-docx", "pip install python-docx"); return

        doc = Document()
        n_marcadas = 0
        for art in articulos:
            texto = art.get("texto", "") or ""
            titulo = art.get("titulo", "Sin título")
            doc.add_heading(titulo, level=2)
            # Detectar si tiene marcas de costura
            if RE_GENERADO.search(texto):
                # Renderizar con marcas rojas
                segmentos = RE_GENERADO.split(texto)
                p = doc.add_paragraph()
                for j, seg in enumerate(segmentos):
                    if not seg:
                        continue
                    run = p.add_run(seg)
                    if j % 2 == 1:
                        run.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
                        run.bold = True
                        n_marcadas += 1
            else:
                doc.add_paragraph(texto)
            doc.add_paragraph()

        doc.save(dest)
        messagebox.showinfo("Exportado",
            f"DOCX guardado en:\n{dest}\n\n"
            f"Palabras reconstruidas marcadas en rojo: {n_marcadas}")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 4: ANÁLISIS TEXTUAL
    # ══════════════════════════════════════════════════════════════════════════
    def _build_anal(self):
        f = self._tab_anal
        self._page_header(f, "Análisis textual y semántico",
                          "NER · LDA · campos semánticos · Word2Vec · red de autoría", "🔍")
        self._build_ai_panel(f, "anal")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True, padx=24, pady=16)

        # ── Barra de acción fija ──────────────────────────────────────────────
        bf_an = tk.Frame(pad, bg=CONTENT_BG); bf_an.pack(fill="x", pady=(0, 8))
        self._btn_anal = ttk.Button(bf_an, text="▶  Iniciar análisis textual",
                                     style="P.TButton", command=self._start_anal)
        self._btn_anal.pack(side="left", padx=(0,12))
        tk.Label(bf_an, text="⚠  Requiere extracción completada",
                 bg=CONTENT_BG, fg=ACENT, font=("Segoe UI",9)).pack(side="left")
        ttk.Button(bf_an, text="📓 Nota", style="S.TButton",
                   command=lambda: self._bitacora_nueva_nota("anal")).pack(side="right")

        self._lbl_fase_a = tk.Label(pad, text="Esperando…",
                                     bg=CONTENT_BG, fg="#8B949E",
                                     font=("Segoe UI",9,"italic"))
        self._lbl_fase_a.pack(anchor="w")
        self._prog_a = ttk.Progressbar(pad, mode="determinate", length=600)
        self._prog_a.pack(fill="x", pady=(6,4))
        self._lbl_pct_a = tk.Label(pad, text="", bg=CONTENT_BG, fg="#8B949E",
                                    font=("Courier",8))
        self._lbl_pct_a.pack(anchor="w")

        log_frame = tk.Frame(pad, bg="#0F1B2D", bd=1, relief="solid")
        log_frame.pack(fill="both", expand=True, pady=(12,0))
        log_hdr = tk.Frame(log_frame, bg="#1A2F4A"); log_hdr.pack(fill="x")
        tk.Label(log_hdr, text="  📋  Registro", bg="#1A2F4A", fg="#94A3B8",
                 font=("Segoe UI",8,"bold")).pack(side="left", pady=4)
        self._log_a = scrolledtext.ScrolledText(log_frame, height=14,
                                                 font=("Consolas",9), bg="#0F1B2D",
                                                 fg="#86EFAC", relief="flat",
                                                 insertbackground="white")
        self._log_a.pack(fill="both", expand=True, padx=1, pady=(0,1))
        self._log_a.config(state="disabled")

        # Inicializar vars antes de que _mk_avanzado las construya lazy
        self._var_nt         = tk.IntVar(value=2)
        self._var_mf         = tk.IntVar(value=3)
        self._var_wv         = tk.BooleanVar(value=False)
        self._var_campo_exp  = tk.StringVar(value=list(CAMPOS_DEFAULT.keys())[0])
        self._txt_exp_res    = None  # se crea dentro del lazy builder

        def _build_anal_params(frame):
            p2 = tk.Frame(frame, bg=CARD_BG, padx=16, pady=10)
            p2.pack(fill="x")
            # N-gramas
            row_ng = tk.Frame(p2, bg=CARD_BG); row_ng.pack(anchor="w", pady=3)
            tk.Label(row_ng, text="N-gramas máx.:", bg=CARD_BG, fg="#CDD6F4",
                     font=("Segoe UI",9,"bold"), width=16, anchor="w").pack(side="left")
            tk.Spinbox(row_ng, from_=1, to=4, textvariable=self._var_nt,
                       width=4, font=("Segoe UI",10), relief="solid", bd=1).pack(side="left", padx=6)
            tk.Label(row_ng, text="(1=unigramas, 2=bigramas, etc.)", bg=CARD_BG,
                     fg="#6E7681", font=("Segoe UI",8)).pack(side="left", padx=4)
            # Frecuencia mínima
            row_mf = tk.Frame(p2, bg=CARD_BG); row_mf.pack(anchor="w", pady=3)
            tk.Label(row_mf, text="Frec. mínima:", bg=CARD_BG, fg="#CDD6F4",
                     font=("Segoe UI",9,"bold"), width=16, anchor="w").pack(side="left")
            tk.Spinbox(row_mf, from_=1, to=20, textvariable=self._var_mf,
                       width=4, font=("Segoe UI",10), relief="solid", bd=1).pack(side="left", padx=6)
            tk.Label(row_mf, text="apariciones mínimas para incluir en vocabulario", bg=CARD_BG,
                     fg="#6E7681", font=("Segoe UI",8)).pack(side="left", padx=4)
            # Word2Vec
            row_wv = tk.Frame(p2, bg=CARD_BG); row_wv.pack(anchor="w", pady=3)
            ttk.Checkbutton(row_wv, text="🧠  Entrenar Word2Vec (expansión semántica automática de campos)",
                            variable=self._var_wv).pack(side="left")

        def _build_anal_exp(frame):
            exp_inner = tk.Frame(frame, bg=CARD_BG, padx=16, pady=10)
            exp_inner.pack(fill="x")
            exp_ctrl = tk.Frame(exp_inner, bg=CARD_BG); exp_ctrl.pack(anchor="w")
            tk.Label(exp_ctrl, text="Campo:", bg=CARD_BG, fg="#CDD6F4",
                     font=("Segoe UI",9,"bold")).pack(side="left")
            campo_cmb = ttk.Combobox(exp_ctrl, textvariable=self._var_campo_exp,
                                      values=list(CAMPOS_DEFAULT.keys()),
                                      state="readonly", width=20, font=("Segoe UI",9))
            campo_cmb.pack(side="left", padx=8)
            ttk.Button(exp_ctrl, text="Explorar →", style="S.TButton",
                       command=self._explorar_expansion).pack(side="left")
            self._txt_exp_res = scrolledtext.ScrolledText(
                exp_inner, height=4, font=("Consolas",9),
                bg="#0D1117", fg="#CDD6F4", relief="solid", bd=1, state="disabled")
            self._txt_exp_res.pack(fill="x", pady=(8,0))

        self._mk_avanzado(pad, "⚙  Parámetros del análisis", _build_anal_params)
        self._mk_avanzado(pad, "🔎  Expansión semántica (requiere Word2Vec)", _build_anal_exp)



    def _explorar_expansion(self):
        if ST.word_model is None:
            messagebox.showinfo("Sin modelo","Ejecuta el análisis textual con Word2Vec activado primero."); return
        campo = self._var_campo_exp.get()
        campos = getattr(ST,"campos_semillas", CAMPOS_DEFAULT)
        semillas = campos.get(campo, [])
        if not semillas:
            messagebox.showinfo("Sin semillas",f"No hay semillas para '{campo}'."); return
        from core.word_vectors import expandir_campo_semantico
        res = expandir_campo_semantico(semillas, ST.word_model, topn=20)
        exp_txt = (f"Semillas encontradas: {res['semillas_encontradas']}\n"
                   f"Expansiones top-15:\n" +
                   "\n".join(f"  {p}  ({s:.3f})" for p,s in res["expansiones"][:15]))
        if self._txt_exp_res is None:
            messagebox.showinfo("Expansión semántica",
                "Abre 'Expansión semántica' en las opciones avanzadas para ver el resultado.\n\n"
                + exp_txt); return
        self._txt_exp_res.config(state="normal")
        self._txt_exp_res.delete("1.0","end")
        self._txt_exp_res.insert("1.0", exp_txt)
        self._txt_exp_res.config(state="disabled")

    # ══════════════════════════════════════════════════════════════════════════
    # ══════════════════════════════════════════════════════════════════════════
    # DESCRIPCIÓN E ICONOGRAFÍA DE IMÁGENES
    # ══════════════════════════════════════════════════════════════════════════
    def _build_imgdesc(self):
        f = self._frames_pagina["imgdesc"]
        self._page_header(f, "Descripción e iconografía de imágenes",
                          "Describe, categoriza y busca imágenes etiquetadas · "
                          "Claude · GPT-4o · Gemini · Ollama", "🎨")

        from core.image_captioner import CATEGORIAS_TEMATICAS
        from core.zone_labeler import VISION_PROVEEDORES

        # ── Barra de control ─────────────────────────────────────────────────
        ctrl = tk.Frame(f, bg=CONTENT_BG)
        ctrl.pack(fill="x", padx=24, pady=(0, 6))

        # Selector de número
        tk.Label(ctrl, text="Número:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        self._imgd_var_num = tk.StringVar()
        self._imgd_cb_num  = ttk.Combobox(ctrl, textvariable=self._imgd_var_num,
                                           width=28, state="readonly")
        self._imgd_cb_num.pack(side="left", padx=(4, 12))
        self._imgd_cb_num.bind("<<ComboboxSelected>>",
                                lambda e: self._imgd_cargar_db())

        # Proveedor + modelo
        tk.Label(ctrl, text="IA:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        self._imgd_var_prov  = tk.StringVar(value="claude")
        self._imgd_var_model = tk.StringVar(value="claude-haiku-4-5-20251001")
        cb_prov = ttk.Combobox(ctrl, textvariable=self._imgd_var_prov,
                                values=list(VISION_PROVEEDORES.keys()),
                                state="readonly", width=9)
        cb_prov.pack(side="left", padx=(4, 4))
        self._imgd_cb_model = ttk.Combobox(ctrl, textvariable=self._imgd_var_model,
                                            state="readonly", width=24)
        self._imgd_cb_model.pack(side="left", padx=(0, 8))

        def _on_prov(*_):
            info = VISION_PROVEEDORES.get(self._imgd_var_prov.get(), {})
            mods = info.get("modelos", [])
            self._imgd_cb_model["values"] = mods
            self._imgd_var_model.set(info.get("default", mods[0] if mods else ""))
        cb_prov.bind("<<ComboboxSelected>>", _on_prov)
        _on_prov()

        # Botones principales
        self._imgd_btn_run = ttk.Button(ctrl, text="▶  Describir fotos",
                                         style="P.TButton",
                                         command=self._imgd_describir)
        self._imgd_btn_run.pack(side="left", padx=(0, 6))
        ttk.Button(ctrl, text="↺ Cargar guardadas", style="S.TButton",
                   command=self._imgd_cargar_db).pack(side="left", padx=(0, 6))
        ttk.Button(ctrl, text="📥 Exportar CSV", style="S.TButton",
                   command=self._imgd_exportar).pack(side="left", padx=(0, 6))
        ttk.Button(ctrl, text="↺ Actualizar números", style="S.TButton",
                   command=self._imgd_refrescar_numeros).pack(side="right")

        self._imgd_lbl_estado = tk.Label(ctrl, text="", bg=CONTENT_BG,
                                          fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._imgd_lbl_estado.pack(side="right", padx=8)

        # ── Filtros por categoría ─────────────────────────────────────────────
        filt_f = tk.Frame(f, bg=CARD_BG, relief="solid", bd=1)
        filt_f.pack(fill="x", padx=24, pady=(0, 6))
        fi = tk.Frame(filt_f, bg=CARD_BG, padx=10, pady=6)
        fi.pack(fill="x")
        tk.Label(fi, text="Filtrar por categoría:", bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 8, "bold")).pack(side="left")
        self._imgd_var_cat = tk.StringVar(value="todas")
        cats = ["todas"] + sorted(CATEGORIAS_TEMATICAS)
        ttk.Combobox(fi, textvariable=self._imgd_var_cat,
                     values=cats, state="readonly", width=22).pack(
                         side="left", padx=(6, 12))
        self._imgd_var_cat.trace_add("write", lambda *_: self._imgd_filtrar())

        tk.Label(fi, text="Buscar similitud:", bg=CARD_BG, fg=TXT_SEC,
                 font=("Segoe UI", 8)).pack(side="left")
        self._imgd_var_busq = tk.StringVar()
        ttk.Entry(fi, textvariable=self._imgd_var_busq, width=28).pack(
            side="left", padx=(4, 6))
        ttk.Button(fi, text="Buscar", style="S.TButton",
                   command=self._imgd_buscar_similitud).pack(side="left")

        # ── Split: tabla + detalle ────────────────────────────────────────────
        split = tk.Frame(f, bg=CONTENT_BG)
        split.pack(fill="both", expand=True, padx=24, pady=(0, 12))

        # Tabla izquierda
        izq = tk.Frame(split, bg=CONTENT_BG)
        izq.pack(side="left", fill="both", expand=True, padx=(0, 8))

        cols = ("pagina", "descripcion", "categorias", "texto_visible", "contexto")
        self._imgd_tv = ttk.Treeview(izq, columns=cols, show="headings", height=22)
        for cid, lbl, w in [
            ("pagina",       "Página",        70),
            ("descripcion",  "Descripción",   260),
            ("categorias",   "Categorías",    140),
            ("texto_visible","Texto visible",  110),
            ("contexto",     "Contexto hist.", 160),
        ]:
            self._imgd_tv.heading(cid, text=lbl,
                command=lambda c=cid: self._imgd_ordenar(c))
            self._imgd_tv.column(cid, width=w, anchor="w")
        sv = ttk.Scrollbar(izq, orient="vertical",   command=self._imgd_tv.yview)
        sh = ttk.Scrollbar(izq, orient="horizontal",  command=self._imgd_tv.xview)
        self._imgd_tv.configure(yscrollcommand=sv.set, xscrollcommand=sh.set)
        sv.pack(side="right", fill="y")
        sh.pack(side="bottom", fill="x")
        self._imgd_tv.pack(fill="both", expand=True)
        self._imgd_tv.bind("<<TreeviewSelect>>", self._imgd_on_sel)

        # Panel derecho: recorte + metadatos
        der = tk.Frame(split, bg=CARD_BG, width=320, relief="solid", bd=1)
        der.pack(side="right", fill="y")
        der.pack_propagate(False)

        self._imgd_canvas = tk.Canvas(der, bg="#000", height=210, highlightthickness=0)
        self._imgd_canvas.pack(fill="x", padx=6, pady=6)

        self._imgd_lbl_desc = tk.Label(der, text="", bg=CARD_BG, fg=TXT_PRI,
                                        font=("Segoe UI", 9, "bold"),
                                        wraplength=300, justify="left")
        self._imgd_lbl_desc.pack(anchor="w", padx=8, pady=(4, 2))

        self._imgd_lbl_cats = tk.Label(der, text="", bg=CARD_BG, fg=AZ4,
                                        font=("Segoe UI", 8),
                                        wraplength=300, justify="left")
        self._imgd_lbl_cats.pack(anchor="w", padx=8)

        self._imgd_lbl_txt = tk.Label(der, text="", bg=CARD_BG, fg=TXT_SEC,
                                       font=("Courier New", 8),
                                       wraplength=300, justify="left")
        self._imgd_lbl_txt.pack(anchor="w", padx=8, pady=(2, 0))

        self._imgd_lbl_ctx = tk.Label(der, text="", bg=CARD_BG, fg=TXT_DIM,
                                       font=("Segoe UI", 8, "italic"),
                                       wraplength=300, justify="left")
        self._imgd_lbl_ctx.pack(anchor="w", padx=8, pady=(2, 8))

        # Búsqueda por similitud — resultados
        tk.Frame(der, bg=CARD_BOR, height=1).pack(fill="x", padx=6, pady=4)
        tk.Label(der, text="Imágenes similares:", bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 8, "bold")).pack(anchor="w", padx=8)
        self._imgd_lbl_sim = tk.Label(der, text="", bg=CARD_BG, fg=TXT_SEC,
                                       font=("Segoe UI", 8),
                                       wraplength=300, justify="left")
        self._imgd_lbl_sim.pack(anchor="w", padx=8, pady=(2, 8))

        # Estado interno
        self._imgd_datos: list[dict] = []
        self._imgd_sort_col = ""
        self._imgd_sort_rev = False

        # Cargar al entrar al panel
        self.after(300, self._imgd_refrescar_numeros)

    # ── Helpers del panel imgdesc ─────────────────────────────────────────────

    def _imgd_refrescar_numeros(self):
        if not ST.out_dir:
            return
        etiq_dir = Path(ST.out_dir) / "05_etiquetas"
        if not etiq_dir.exists():
            return
        nums = sorted(p.name for p in etiq_dir.iterdir() if p.is_dir())
        self._imgd_cb_num["values"] = nums
        if nums and not self._imgd_var_num.get():
            self._imgd_var_num.set(nums[0])
            self._imgd_cargar_db()

    def _imgd_cargar_db(self):
        num = self._imgd_var_num.get()
        if not num or not ST.ruta_db:
            return
        from core.image_captioner import cargar_descripciones_db
        db = Path(ST.ruta_db)
        if not db.exists():
            self._imgd_lbl_estado.config(
                text="Sin descripciones guardadas aún — usa ▶ Describir fotos", fg=TXT_SEC)
            return
        datos = cargar_descripciones_db(db, num)
        self._imgd_datos = datos
        self._imgd_poblar_tv(datos)
        self._imgd_lbl_estado.config(
            text=f"{len(datos)} imágenes descritas — {num}", fg=VERDE)

    def _imgd_poblar_tv(self, datos: list[dict]):
        self._imgd_tv.delete(*self._imgd_tv.get_children())
        for d in datos:
            cats = ", ".join(d.get("categorias", []))[:45]
            self._imgd_tv.insert("", "end", values=(
                d.get("pagina", ""),
                d.get("descripcion", "")[:70],
                cats,
                d.get("texto_visible", "")[:35],
                d.get("contexto_historico", "")[:50],
            ))

    def _imgd_filtrar(self, *_):
        cat = self._imgd_var_cat.get()
        if cat == "todas":
            self._imgd_poblar_tv(self._imgd_datos)
        else:
            filtrados = [d for d in self._imgd_datos
                         if cat in d.get("categorias", [])]
            self._imgd_poblar_tv(filtrados)
        self._imgd_lbl_estado.config(
            text=f"Filtrando: {cat}", fg=TXT_SEC)

    def _imgd_ordenar(self, col: str):
        if self._imgd_sort_col == col:
            self._imgd_sort_rev = not self._imgd_sort_rev
        else:
            self._imgd_sort_col = col
            self._imgd_sort_rev = False
        datos_ord = sorted(self._imgd_datos,
                           key=lambda d: str(d.get(col, "")),
                           reverse=self._imgd_sort_rev)
        self._imgd_poblar_tv(datos_ord)

    def _imgd_on_sel(self, event=None):
        sel = self._imgd_tv.selection()
        if not sel:
            return
        idx = self._imgd_tv.index(sel[0])
        datos_visibles = [
            self._imgd_tv.item(iid)["values"]
            for iid in self._imgd_tv.get_children()
        ]
        # Encontrar el dict completo que corresponde a esta fila
        pag_sel = datos_visibles[idx][0] if datos_visibles else ""
        d = next((x for x in self._imgd_datos
                  if x.get("pagina") == pag_sel), None)
        if d is None:
            return
        self._imgd_lbl_desc.config(text=d.get("descripcion", ""))
        cats = d.get("categorias", [])
        self._imgd_lbl_cats.config(
            text=("📌 " + " · ".join(cats)) if cats else "")
        tv = d.get("texto_visible", "")
        self._imgd_lbl_txt.config(
            text=("📝 " + tv) if tv else "")
        self._imgd_lbl_ctx.config(
            text=d.get("contexto_historico", ""))
        self._imgd_mostrar_recorte(d)

    def _imgd_mostrar_recorte(self, d: dict):
        self._imgd_canvas.delete("all")
        try:
            from PIL import Image, ImageTk
            num     = self._imgd_var_num.get()
            pagina  = d.get("pagina", "")
            img_dir = Path(ST.out_dir) / "02_imagenes" / num
            hits    = sorted(img_dir.glob(f"*{pagina}*.png")) if img_dir.exists() else []
            if not hits:
                return
            img = Image.open(hits[0]).convert("RGB")
            W, H = img.size
            x0 = int(d.get("x0", 0) * W)
            y0 = int(d.get("y0", 0) * H)
            x1 = int(d.get("x1", 1) * W)
            y1 = int(d.get("y1", 1) * H)
            recorte = img.crop((x0, y0, x1, y1))
            recorte.thumbnail((306, 200), Image.LANCZOS)
            tk_img = ImageTk.PhotoImage(recorte)
            self._imgd_canvas._ref = tk_img
            cw = self._imgd_canvas.winfo_width() or 306
            self._imgd_canvas.create_image(cw // 2, 105, anchor="center", image=tk_img)
        except Exception:
            pass

    def _imgd_describir(self):
        num = self._imgd_var_num.get()
        if not num or not ST.out_dir:
            messagebox.showwarning("Sin número",
                "Selecciona un número con imágenes etiquetadas.")
            return
        prov  = self._imgd_var_prov.get()
        model = self._imgd_var_model.get()
        api_k = ST.api_keys.get(prov, "") or ST.api_key
        if prov != "ollama" and not api_k:
            messagebox.showwarning("Sin API key",
                f"Configura la API key de {prov} en ⚙ Configuración.")
            return

        self._imgd_btn_run.config(state="disabled")
        self._imgd_lbl_estado.config(text="⏳ Describiendo…", fg="#F59E0B")

        from core.image_captioner import describir_numero

        def cb(n, t, pag, desc):
            self.after(0, lambda: self._imgd_lbl_estado.config(
                text=f"⏳ {n}/{t}: {pag} — {desc[:45]}…", fg="#F59E0B"))

        def _run():
            db = Path(ST.ruta_db) if ST.ruta_db else None
            datos = describir_numero(ST.out_dir, num, proveedor=prov,
                                     api_key=api_k, modelo=model,
                                     db_path=db, callback=cb)
            self.after(0, lambda: (
                self._imgd_datos.__setitem__(slice(None), datos),
                self._imgd_poblar_tv(datos),
                self._imgd_btn_run.config(state="normal"),
                self._imgd_lbl_estado.config(
                    text=f"✅ {len(datos)} imágenes descritas", fg=VERDE),
            ))
        threading.Thread(target=_run, daemon=True).start()

    def _imgd_buscar_similitud(self):
        q = self._imgd_var_busq.get().strip()
        if not q:
            return
        num = self._imgd_var_num.get()
        if not num or not ST.out_dir:
            return
        from core.image_captioner import buscar_imagenes_similares
        sims = buscar_imagenes_similares(q, ST.out_dir, num, top_n=5)
        if sims:
            txt = "\n".join(
                f"  p.{s['pagina']} — {s['descripcion'][:55]}"
                for s in sims)
            self._imgd_lbl_sim.config(text=txt)
        else:
            self._imgd_lbl_sim.config(
                text="Sin índice FAISS. Describe las imágenes primero.")

    def _imgd_exportar(self):
        if not self._imgd_datos:
            messagebox.showwarning("Sin datos", "Describe las imágenes primero.")
            return
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("JSON", "*.json")],
            initialfile=f"imagenes_{self._imgd_var_num.get()}.csv")
        if not ruta:
            return
        import pandas as _pd
        df = _pd.DataFrame(self._imgd_datos)
        if ruta.endswith(".json"):
            import json as _json
            Path(ruta).write_text(
                _json.dumps(self._imgd_datos, ensure_ascii=False, indent=2),
                encoding="utf-8")
        else:
            df.to_csv(ruta, index=False, encoding="utf-8-sig")
        self.toast(f"Exportado → {Path(str(ruta)).name}", tipo="ok")

    # TAB 5: ANÁLISIS VISUAL Y TIPOGRÁFICO
    # ══════════════════════════════════════════════════════════════════════════
    def _build_vis(self):
        f = self._tab_vis
        self._page_header(f, "Análisis visual y tipográfico",
                          "Fuentes tipográficas · imágenes detectadas · diagrama de layout", "🖼")
        self._build_ai_panel(f, "vis")
        # Sub-pestañas con botones propios (no ttk.Notebook para mantener estilo)
        top = tk.Frame(f, bg=CONTENT_BG); top.pack(fill="x", padx=24, pady=(12,0))
        self._vis_tabs_btns = {}
        for i, (tid, label) in enumerate([("tip","🔤 Tipografía"),
                                           ("ele","📷 Imágenes"),
                                           ("diag","📐 Diagrama")]):
            btn = tk.Label(top, text=f"  {label}  ", bg=CARD_BOR if i > 0 else AZ3,
                           fg="white" if i == 0 else "#64748B",
                           font=("Segoe UI",9,"bold"), cursor="hand2",
                           padx=10, pady=6, relief="flat")
            btn.pack(side="left", padx=(0,2))
            self._vis_tabs_btns[tid] = btn
        tk.Frame(f, bg=CARD_BOR, height=1).pack(fill="x")

        # Frames de sub-contenido
        self._vis_frames = {}
        for tid in ("tip","ele","diag"):
            frm = tk.Frame(f, bg=CONTENT_BG)
            self._vis_frames[tid] = frm

        self._tab_vis_tip  = self._vis_frames["tip"]
        self._tab_vis_ele  = self._vis_frames["ele"]
        self._tab_vis_diag = self._vis_frames["diag"]

        self._build_vis_tip(); self._build_vis_ele(); self._build_vis_diag()

        for tid, btn in self._vis_tabs_btns.items():
            btn.bind("<Button-1>", lambda e, t=tid: self._vis_switch(t))

        # Botones de acción
        bf = tk.Frame(f, bg=CONTENT_BG); bf.pack(fill="x", padx=24, pady=10)
        self._btn_vis = ttk.Button(bf, text="▶  Analizar visual y tipografía",
                                    style="P.TButton", command=self._start_vis)
        self._btn_vis.pack(side="left", padx=(0,12))
        self._lbl_vis_ok = tk.Label(bf, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI",10,"bold"))
        self._lbl_vis_ok.pack(side="left")
        tk.Label(bf, text="⚠  Requiere extracción completada",
                 bg=CONTENT_BG, fg=ACENT, font=("Segoe UI",9)).pack(side="left", padx=8)

        # Mostrar primera sub-pestaña
        self._vis_switch("tip")

    def _vis_switch(self, tid: str):
        for t, frm in self._vis_frames.items():
            frm.pack_forget()
        self._vis_frames[tid].pack(fill="both", expand=True)
        for t, btn in self._vis_tabs_btns.items():
            if t == tid:
                btn.config(bg=AZ3, fg="white")
            else:
                btn.config(bg=CARD_BOR, fg="#8B949E")

    def _build_vis_tip(self):
        pad = self._tab_vis_tip
        cols = ("Número","Fuente principal","Clasificación","N fuentes","Cuerpo (pt)",
                "Título (pt)","Ratio T/C","Interlineado","Columnas","% Negrita","% Cursiva","Imgs.")
        widths = [110,160,130,70,80,80,60,80,70,70,70,50]
        tv_f = tk.Frame(pad, bg=CONTENT_BG); tv_f.pack(fill="both", expand=True, padx=8, pady=8)
        sbv = ttk.Scrollbar(tv_f, orient="vertical")
        sbh = ttk.Scrollbar(tv_f, orient="horizontal")
        self._tv_tip = ttk.Treeview(tv_f, columns=cols, show="headings",
                                     yscrollcommand=sbv.set, xscrollcommand=sbh.set, height=14)
        for col, w in zip(cols, widths):
            self._tv_tip.heading(col, text=col, anchor="w")
            self._tv_tip.column(col, width=w, minwidth=40)
        sbv.config(command=self._tv_tip.yview); sbh.config(command=self._tv_tip.xview)
        sbv.pack(side="right", fill="y"); sbh.pack(side="bottom", fill="x")
        self._tv_tip.pack(fill="both", expand=True)
        # detalle de fuentes al hacer clic
        det_f = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1); tk.Label(det_f, text="  Detalle de fuentes del número seleccionado", bg="#1C2128", fg=TXT_PRI, font=("Segoe UI",8,"bold")).pack(fill="x")
        det_f.pack(fill="x", padx=8, pady=(0,6))
        self._txt_tip_det = scrolledtext.ScrolledText(det_f, height=5, font=("Courier",9), bg="#0D1117", fg="#CDD6F4", relief="flat", state="disabled")
        self._txt_tip_det.pack(fill="x")
        self._tv_tip.bind("<<TreeviewSelect>>", self._on_tip_select)

    def _build_vis_ele(self):
        pad = self._tab_vis_ele
        cols = ("Número","Página","Tipo","Confianza","Ancho cm","Alto cm","Área cm²",
                "Pos X%","Pos Y%","Autor imagen","Pie de foto","Descripción IA")
        widths = [110,80,130,70,70,70,70,60,60,120,160,280]
        tv_f = tk.Frame(pad, bg=CONTENT_BG); tv_f.pack(fill="both", expand=True, padx=8, pady=8)
        sbv = ttk.Scrollbar(tv_f, orient="vertical")
        sbh = ttk.Scrollbar(tv_f, orient="horizontal")
        self._tv_ele = ttk.Treeview(tv_f, columns=cols, show="headings",
                                     yscrollcommand=sbv.set, xscrollcommand=sbh.set, height=16)
        for col, w in zip(cols, widths):
            self._tv_ele.heading(col, text=col, anchor="w")
            self._tv_ele.column(col, width=w, minwidth=40)
        sbv.config(command=self._tv_ele.yview); sbh.config(command=self._tv_ele.xview)
        sbv.pack(side="right", fill="y"); sbh.pack(side="bottom", fill="x")
        self._tv_ele.pack(fill="both", expand=True)
        # Etiquetas de color por tipo
        self._tv_ele.tag_configure("foto",        background="#0D2137", foreground="#58A6FF")
        self._tv_ele.tag_configure("ilustracion", background="#0D2318", foreground="#3FB950")
        self._tv_ele.tag_configure("publicidad",  background="#2D1A00", foreground="#F0883E")
        self._tv_ele.tag_configure("mixto",       background="#1A1030", foreground="#BC8CFF")
        # Contador
        cnt_f = tk.Frame(pad, bg=CONTENT_BG); cnt_f.pack(anchor="w", padx=8, pady=(0,4))
        self._lbl_ele_cnt = ttk.Label(cnt_f, text="", foreground="#8B949E", font=("Segoe UI",9))
        self._lbl_ele_cnt.pack(side="left")
        ttk.Button(cnt_f, text="📥 Exportar CSV", style="S.TButton",
                   command=self._exportar_csv_imagenes).pack(side="left", padx=12)
        ttk.Button(cnt_f, text="🖼 Exportar imágenes recortadas", style="A.TButton",
                   command=self._exportar_imagenes_carpeta).pack(side="left", padx=4)

    def _build_vis_diag(self):
        pad = self._tab_vis_diag
        ctrl = tk.Frame(pad, bg=CONTENT_BG); ctrl.pack(fill="x", padx=12, pady=8)
        ttk.Label(ctrl, text="Número:", font=("Segoe UI",10)).pack(side="left")
        self._var_diag_num = tk.StringVar()
        self._cmb_diag = ttk.Combobox(ctrl, textvariable=self._var_diag_num,
                                       state="readonly", width=30, font=("Segoe UI",10))
        self._cmb_diag.pack(side="left", padx=8)
        ttk.Label(ctrl, text="Página:", font=("Segoe UI",10)).pack(side="left")
        self._var_diag_pag = tk.StringVar()
        self._cmb_diag_pag = ttk.Combobox(ctrl, textvariable=self._var_diag_pag,
                                            state="readonly", width=14, font=("Segoe UI",10))
        self._cmb_diag_pag.pack(side="left", padx=8)
        self._cmb_diag.bind("<<ComboboxSelected>>", self._on_diag_num_sel)
        self._cmb_diag_pag.bind("<<ComboboxSelected>>", self._on_diag_pag_sel)
        ttk.Button(ctrl, text="📐 Mostrar diagrama", style="S.TButton",
                   command=self._mostrar_diagrama).pack(side="left", padx=8)
        # Canvas para la imagen
        self._canvas_diag = tk.Canvas(pad, bg=CONTENT_BG, highlightthickness=0)
        self._canvas_diag.pack(fill="both", expand=True, padx=12, pady=4)
        self._diag_img_ref = None   # evitar GC de la imagen

    def _on_tip_select(self, _event):
        """Muestra detalle de fuentes cuando el usuario hace clic en una fila."""
        sel = self._tv_tip.selection()
        if not sel: return
        vals = self._tv_tip.item(sel[0])["values"]
        if not vals: return
        nombre = str(vals[0])
        tip = ST.datos_visual.get("tipografia", {}).get(nombre, {})
        fuentes = tip.get("fuentes_resumen", [])
        self._txt_tip_det.config(state="normal")
        self._txt_tip_det.delete("1.0", "end")
        if fuentes:
            self._txt_tip_det.insert("end",
                f"{'Fuente':<35} {'Clasificación':<22} {'Cuerpo(pt)':<12} {'Negrita%':<10} {'Cursiva%'}\n")
            self._txt_tip_det.insert("end", "─"*100 + "\n")
            for fd in fuentes[:10]:
                self._txt_tip_det.insert("end",
                    f"{fd['fuente']:<35} {fd.get('clasificacion',''):<22} "
                    f"{fd.get('tam_mediano',0):<12.1f} {fd.get('negrita_pct',0):<10.1f} {fd.get('cursiva_pct',0):.1f}\n")
        else:
            self._txt_tip_det.insert("end", "(sin detalle disponible)")
        self._txt_tip_det.config(state="disabled")

    def _on_diag_num_sel(self, _event):
        """Rellena la lista de páginas cuando el usuario selecciona un número."""
        nombre = self._var_diag_num.get()
        datos = ST.datos_imagenes.get(nombre, {})
        paginas = [p["pagina"] for p in datos.get("paginas", []) if p.get("elementos")]
        self._cmb_diag_pag["values"] = paginas
        if paginas:
            self._cmb_diag_pag.set(paginas[0])
            self._mostrar_diagrama()

    def _on_diag_pag_sel(self, _event):
        self._mostrar_diagrama()

    def _mostrar_diagrama(self):
        """Genera y muestra el diagrama de layout de la página seleccionada."""
        from core.image_analyzer import generar_diagrama_layout
        nombre = self._var_diag_num.get()
        pagina = self._var_diag_pag.get()
        datos  = ST.datos_imagenes.get(nombre, {})
        pag_datos = next((p for p in datos.get("paginas", []) if p["pagina"] == pagina), None)
        if not pag_datos:
            return
        png_bytes = generar_diagrama_layout(pag_datos, titulo=f"{nombre}")
        self._pegar_imagen_canvas(self._canvas_diag, png_bytes)

    def _pegar_imagen_canvas(self, canvas, png_bytes: bytes):
        """Pega bytes PNG en un canvas Tkinter."""
        import io

        from PIL import Image, ImageTk
        img = Image.open(io.BytesIO(png_bytes))
        cw  = canvas.winfo_width()  or 500
        ch  = canvas.winfo_height() or 600
        img.thumbnail((cw, ch), Image.LANCZOS)
        tk_img = ImageTk.PhotoImage(img)
        canvas.delete("all")
        canvas.create_image(cw//2, ch//2, anchor="center", image=tk_img)
        self._diag_img_ref = tk_img   # evitar GC

    def _exportar_csv_imagenes(self):
        """Exporta la tabla de imágenes detectadas a CSV."""
        import csv
        if not ST.datos_imagenes:
            messagebox.showwarning("Sin datos", "Ejecuta el análisis visual primero."); return
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile="imagenes_detectadas.csv")
        if not dest: return
        cols = ["numero","pagina","tipo","confianza","w_cm","h_cm","area_cm2",
                "pos_x_pct","pos_y_pct","autor_imagen","pie_de_foto","descripcion_ia"]
        with open(dest, "w", newline="", encoding="utf-8") as f:
            w = csv.DictWriter(f, fieldnames=cols, extrasaction="ignore")
            w.writeheader()
            for nombre, datos in ST.datos_imagenes.items():
                for pag in datos.get("paginas", []):
                    for el in pag.get("elementos", []):
                        if el["tipo"] == "Texto": continue
                        row = {k: el.get(k,"") for k in cols}
                        row["numero"] = nombre
                        row["pagina"] = pag["pagina"]
                        w.writerow(row)
        self.toast(f"Exportado → {Path(dest).name}", tipo="ok")

    def _exportar_imagenes_carpeta(self):
        """Recorta y guarda cada imagen detectada en una carpeta organizada."""
        if not ST.datos_imagenes:
            messagebox.showwarning("Sin datos", "Ejecuta el análisis visual primero."); return

        destino = filedialog.askdirectory(title="Carpeta de destino para las imágenes")
        if not destino: return
        destino = Path(destino) / "imagenes_extraidas"

        # Diálogo de progreso
        prog_win = tk.Toplevel(self)
        prog_win.title("Exportando imágenes…")
        prog_win.geometry("420x140")
        prog_win.resizable(False, False)
        prog_win.grab_set()
        tk.Label(prog_win, text="Exportando imágenes recortadas",
                 font=("Segoe UI", 10, "bold")).pack(pady=(16, 4))
        lbl_prog = tk.Label(prog_win, text="Preparando…", font=("Segoe UI", 9))
        lbl_prog.pack()
        bar = ttk.Progressbar(prog_win, mode="determinate", length=380)
        bar.pack(pady=8)

        def worker():
            from core.image_exporter import exportar_imagenes
            img_dir_raiz = ST.out_dir / "imgs" if hasattr(ST, "out_dir") else Path("imgs")

            def cb(n, total, desc):
                pct = int(n / max(total, 1) * 100)
                self.after(0, lambda: bar.config(value=pct))
                self.after(0, lambda d=desc: lbl_prog.config(text=d))

            try:
                stats = exportar_imagenes(
                    datos_imagenes=ST.datos_imagenes,
                    img_dir_raiz=img_dir_raiz,
                    destino=destino,
                    publicacion=getattr(ST, "publicacion", "Publicacion"),
                    callback=cb,
                )
                self.after(0, prog_win.destroy)
                msg = (f"✅ {stats['exportadas']} imágenes exportadas\n"
                       f"   Omitidas (demasiado pequeñas): {stats['omitidas']}\n"
                       f"   Errores: {stats['errores']}\n\n"
                       f"   Carpeta:\n   {destino}")
                self.after(0, lambda: messagebox.showinfo("Exportación completada", msg))
            except Exception as e:
                self.after(0, prog_win.destroy)
                self.after(0, lambda err=str(e): messagebox.showerror("Error", err))

        threading.Thread(target=worker, daemon=True).start()

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 7: METADATOS URL
    # ══════════════════════════════════════════════════════════════════════════
    def _build_meta(self):
        f = self._tab_meta
        self._page_header(f, "Metadatos desde URL",
                          "Extrae título, autor, fecha y datos del archivo digital", "🔗")
        self._build_ai_panel(f, "meta")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True, padx=24, pady=16)

        # Tarjeta de entrada
        url_card = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1,
                             highlightbackground=CARD_BOR)
        url_card.pack(fill="x", pady=(0,12))
        url_inner = tk.Frame(url_card, bg=CARD_BG, padx=16, pady=14)
        url_inner.pack(fill="x")
        url_inner.columnconfigure(1, weight=1)

        tk.Label(url_inner, text="URL del catálogo:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI",9,"bold")).grid(row=0,column=0,sticky="w",pady=4)
        self._var_meta_url = tk.StringVar()
        url_row = tk.Frame(url_inner, bg=CARD_BG)
        url_row.grid(row=0, column=1, columnspan=2, sticky="ew", padx=8)
        tk.Entry(url_row, textvariable=self._var_meta_url, width=60,
                 font=("Segoe UI",9), relief="solid", bd=1, bg="#1C2128").pack(
                 side="left", fill="x", expand=True)
        ttk.Button(url_row, text="Extraer metadatos →", style="P.TButton",
                   command=self._extraer_meta).pack(side="left", padx=(8,0))
        ttk.Button(url_row, text="Ejemplos", style="S.TButton",
                   command=self._ejemplos_url).pack(side="left", padx=(4,0))

        tk.Label(url_inner,
                 text="Compatible con BNCO · Archive.org · Europeana · BNE · HathiTrust",
                 bg=CARD_BG, fg="#6E7681", font=("Segoe UI",8)).grid(
                 row=1, column=0, columnspan=3, sticky="w", pady=(0,4))

        # Resultado
        res_f = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1)
        res_f.pack(fill="both", expand=True)
        res_hdr = tk.Frame(res_f, bg="#1C2128"); res_hdr.pack(fill="x")
        tk.Label(res_hdr, text="  📋  Metadatos extraídos",
                 bg="#1C2128", fg=TXT_PRI, font=("Segoe UI",8,"bold")).pack(
                 side="left", pady=4)
        self._txt_meta = scrolledtext.ScrolledText(res_f, height=22, font=("Consolas",10),
                                                    bg="#1C2128", fg="#CDD6F4",
                                                    relief="flat", state="disabled")
        self._txt_meta.pack(fill="both", expand=True, padx=1, pady=(0,1))

        # Botones de acción
        act_f = tk.Frame(pad, bg=CONTENT_BG); act_f.pack(fill="x", pady=(8,0))
        ttk.Button(act_f, text="📥 Guardar JSON", style="S.TButton",
                   command=self._guardar_meta_json).pack(side="left", padx=(0,8))
        ttk.Button(act_f, text="📄 Añadir a Excel", style="S.TButton",
                   command=self._meta_a_excel).pack(side="left")
        self._lbl_meta_ok = tk.Label(act_f, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI",9,"bold"))
        self._lbl_meta_ok.pack(side="left", padx=12)
        self._meta_actual = {}

    def _extraer_meta(self):
        url = self._var_meta_url.get().strip()
        if not url:
            messagebox.showwarning("URL vacía", "Escribe o pega la URL primero."); return
        self._txt_meta.config(state="normal")
        self._txt_meta.delete("1.0", "end")
        self._txt_meta.insert("end", "⏳ Consultando biblioteca digital…\n")
        self._txt_meta.insert("end", "🌐 Buscando metadatos adicionales en la web…\n")
        self._txt_meta.config(state="disabled")
        self._lbl_meta_ok.config(text="")
        def worker():
            from core.metadata_extractor import (
                enriquecer_con_busqueda_web,
                extraer_metadata_url,
                formatear_metadatos,
            )
            try:
                meta = extraer_metadata_url(url)
                # Enriquecer con búsqueda web si faltan campos
                campos_vacios = [k for k in ("titulo","creador","fecha","descripcion")
                                 if not meta.get(k)]
                if campos_vacios:
                    meta = enriquecer_con_busqueda_web(meta, url)
                texto = formatear_metadatos(meta)
                self._meta_actual = meta
                self.after(0, lambda t=texto, m=meta: self._mostrar_meta(t, m))
            except Exception as e:
                self.after(0, lambda err=str(e): self._mostrar_meta(f"⚠️ Error: {err}", {}))
        threading.Thread(target=worker, daemon=True).start()

    def _mostrar_meta(self, texto: str, meta: dict):
        self._txt_meta.config(state="normal")
        self._txt_meta.delete("1.0", "end")
        self._txt_meta.insert("end", texto)
        self._txt_meta.config(state="disabled")
        fuente = meta.get("fuente_metadata", "—")
        if meta.get("titulo"):
            self._lbl_meta_ok.config(
                text=f"✅ Metadatos obtenidos · Fuente: {fuente}")
        else:
            self._lbl_meta_ok.config(text="⚠️ Sin metadatos encontrados")

    def _ejemplos_url(self):
        ejemplos = (
            "URLs de ejemplo para probar:\n\n"
            "• Biblioteca Nacional de Colombia (BNCO / SirsiDynix):\n"
            "  https://bnco.ent.sirsi.net/custom/web/content/conservacion/"
            "html/visorFicheros.html?idFichero=190988\n\n"
            "• Archive.org:\n"
            "  https://archive.org/details/estampa_colombia_1939\n\n"
            "• Europeana:\n"
            "  https://www.europeana.eu/item/9200579/BibliographicResource_3000126484305\n\n"
            "• Biblioteca Nacional de España (BNE):\n"
            "  https://hemerotecadigital.bne.es/hd/es/viewer?id=...\n\n"
            "El sistema intentará OAI-PMH, JSON-LD, meta etiquetas y scraping "
            "específico en ese orden."
        )
        messagebox.showinfo("URLs de ejemplo", ejemplos)

    def _guardar_meta_json(self):
        if not self._meta_actual:
            messagebox.showwarning("Sin datos", "Extrae metadatos primero."); return
        import json
        dest = filedialog.asksaveasfilename(
            defaultextension=".json", filetypes=[("JSON","*.json")],
            initialfile="metadatos.json")
        if not dest: return
        with open(dest, "w", encoding="utf-8") as f:
            json.dump(self._meta_actual, f, ensure_ascii=False, indent=2)
        self._lbl_meta_ok.config(text=f"✅ Guardado: {Path(dest).name}")

    def _meta_a_excel(self):
        messagebox.showinfo("Próximamente",
                            "Esta función añadirá los metadatos como hoja adicional "
                            "al Excel en la próxima ejecución de exportación.")

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 6: ANÁLISIS COMPARATIVO
    # ══════════════════════════════════════════════════════════════════════════
    def _build_comp(self):
        f = self._tab_comp
        self._page_header(f, "Análisis comparativo",
                          "Compara el perfil temático con otras publicaciones del período", "📊")
        self._build_ai_panel(f, "comp")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True)
        # Sub-pestañas
        top = tk.Frame(pad, bg=CONTENT_BG); top.pack(fill="x", padx=24, pady=(12,0))
        self._comp_tabs_btns = {}
        for i,(tid,label) in enumerate([("sim","🔁 Similaridad"),
                                         ("dist","🏷 Términos distintivos"),
                                         ("cam","📊 Campos semánticos")]):
            btn = tk.Label(top, text=f"  {label}  ",
                           bg=AZ3 if i==0 else CARD_BOR,
                           fg="white" if i==0 else "#64748B",
                           font=("Segoe UI",9,"bold"), cursor="hand2",
                           padx=10, pady=6)
            btn.pack(side="left", padx=(0,2))
            self._comp_tabs_btns[tid] = btn
        tk.Frame(pad, bg=CARD_BOR, height=1).pack(fill="x")

        # ── Barra de acción fija ──────────────────────────────────────────────
        bf_comp = tk.Frame(pad, bg=CONTENT_BG); bf_comp.pack(fill="x", padx=0, pady=(0,8))
        self._btn_comp = ttk.Button(bf_comp, text="▶  Ejecutar análisis comparativo",
                                     style="P.TButton", command=self._start_comp)
        self._btn_comp.pack(side="left", padx=(0,12))
        self._lbl_comp_ok = tk.Label(bf_comp, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI",10,"bold"))
        self._lbl_comp_ok.pack(side="left")
        tk.Label(bf_comp, text="⚠  Requiere corpus de referencia configurado",
                 bg=CONTENT_BG, fg=ACENT, font=("Segoe UI",9)).pack(side="left",padx=8)

        self._comp_frames = {}
        for tid in ("sim","dist","cam"):
            frm = tk.Frame(pad, bg=CONTENT_BG)
            self._comp_frames[tid] = frm
        self._tab_c_sim  = self._comp_frames["sim"]
        self._tab_c_dist = self._comp_frames["dist"]
        self._tab_c_cam  = self._comp_frames["cam"]

        self._build_comp_sim(); self._build_comp_dist(); self._build_comp_cam()

        for tid, btn in self._comp_tabs_btns.items():
            btn.bind("<Button-1>", lambda e, t=tid: self._comp_switch(t))

        self._comp_switch("sim")

    def _comp_switch(self, tid: str):
        for t, frm in self._comp_frames.items():
            frm.pack_forget()
        self._comp_frames[tid].pack(fill="both", expand=True)
        for t, btn in self._comp_tabs_btns.items():
            btn.config(bg=AZ3 if t==tid else CARD_BOR,
                       fg="white" if t==tid else "#64748B")

    def _build_comp_sim(self):
        pad = self._tab_c_sim
        tk.Label(pad,text="Matriz de similaridad coseno (TF-IDF)",bg=CONTENT_BG,fg=TXT_PRI,font=("Segoe UI",10,"bold")).pack(anchor="w",padx=8,pady=6)
        self._txt_sim = scrolledtext.ScrolledText(pad, height=12, font=("Consolas",9), bg=CARD_BG, fg="#CDD6F4", relief="flat")
        self._txt_sim.pack(fill="both",expand=True,padx=8,pady=4)

    def _build_comp_dist(self):
        pad = self._tab_c_dist
        tk.Label(pad,text="Palabras más distintivas de la publicación analizada",bg=CONTENT_BG,fg=TXT_PRI,font=("Segoe UI",10,"bold")).pack(anchor="w",padx=8,pady=6)
        self._txt_dist = scrolledtext.ScrolledText(pad, height=14, font=("Consolas",9), bg=CARD_BG, fg="#CDD6F4", relief="flat")
        self._txt_dist.pack(fill="both",expand=True,padx=8,pady=4)

    def _build_comp_cam(self):
        pad = self._tab_c_cam
        tk.Label(pad,text="Perfil de campos semánticos por publicación",bg=CONTENT_BG,fg=TXT_PRI,font=("Segoe UI",10,"bold")).pack(anchor="w",padx=8,pady=6)
        # Figura embebida
        self._fig_cam_frame = tk.Frame(pad, bg=CONTENT_BG); self._fig_cam_frame.pack(fill="both",expand=True,padx=8,pady=4)

    # ══════════════════════════════════════════════════════════════════════════
    # TAB 7: RESULTADOS
    # ══════════════════════════════════════════════════════════════════════════
    def _build_res(self):
        f = self._tab_res
        self._page_header(f, "Resultados y exportación",
                          "Gráficas interactivas, análisis de red y exportación Excel", "📈")
        self._build_ai_panel(f, "res")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True, padx=24, pady=12)
        # Métricas
        ind = tk.Frame(pad, bg=CONTENT_BG); ind.pack(fill="x", pady=(0,12))
        self._lbl_r_num  = self._mk_ind(ind,"Números","—",0)
        self._lbl_r_pag  = self._mk_ind(ind,"Páginas","—",1)
        self._lbl_r_pal  = self._mk_ind(ind,"Palabras","—",2)
        self._lbl_r_art  = self._mk_ind(ind,"Artículos","—",3)
        self._lbl_r_aut  = self._mk_ind(ind,"Autores","—",4)
        self._lbl_r_fir  = self._mk_ind(ind,"Firmas NER","—",5)
        # Notebook de gráficas
        nb_r = ttk.Notebook(pad); nb_r.pack(fill="both",expand=True,pady=6)
        self._figs_tabs = {}
        for key, label in [
            ("secciones","📋 Secciones"),("firmas","✍️ Firmas"),
            ("campos","🔤 Campos sem."),("articulos","📝 Artículos"),
            ("lda","🧩 Temas LDA"),("red","🕸️ Red"),
            ("visual","🖼️ Visual"),("comparativo","📊 Comparativo"),("layout","📐 Layout"),
        ]:
            tab = tk.Frame(nb_r, bg=CONTENT_BG); nb_r.add(tab, text=f"  {label}  ")
            self._figs_tabs[key] = tab
        # Botones de exportación
        bb = ttk.Frame(pad,padding=8); bb.pack(fill="x")
        ttk.Button(bb, text="📊  Excel completo (10 hojas)",
                   style="P.TButton", command=self._gen_excel).pack(side="left",padx=8)
        ttk.Button(bb, text="🕸️  Red .graphml",
                   style="S.TButton", command=self._guardar_graphml).pack(side="left",padx=8)
        ttk.Button(bb, text="📁  Abrir carpeta",
                   style="S.TButton", command=self._abrir_carpeta).pack(side="left",padx=8)
        ttk.Button(bb, text="📄  XML-TEI",
                   style="S.TButton", command=self._res_exportar_tei).pack(side="left",padx=4)
        ttk.Button(bb, text="✓  Validar TEI",
                   style="S.TButton", command=self._res_validar_tei).pack(side="left",padx=2)
        ttk.Button(bb, text="📚  BibTeX",
                   style="S.TButton", command=self._res_exportar_bibtex).pack(side="left",padx=4)
        ttk.Button(bb, text="📦  Paquete publicación",
                   style="S.TButton", command=self._res_paquete_publicacion).pack(side="left",padx=4)
        ttk.Button(bb, text="📋  METHODS.md",
                   style="S.TButton", command=self._res_generar_methods).pack(side="left",padx=4)
        ttk.Button(bb, text="🔗  JSON Observable",
                   style="S.TButton", command=self._res_exportar_json_observable).pack(side="left",padx=4)
        ttk.Button(bb, text="📊  PowerPoint",
                   style="S.TButton", command=self._res_exportar_pptx).pack(side="left",padx=4)
        ttk.Button(bb, text="💾  Guardar como…",
                   style="P.TButton", command=self._exp_abrir_dialogo).pack(side="left",padx=4)
        self._lbl_excel = tk.Label(bb, text="", bg=CONTENT_BG, fg=VERDE, font=("Segoe UI",10,"bold"))
        self._lbl_excel.pack(side="left",padx=12)

    # ══════════════════════════════════════════════════════════════════════════
    # HELPERS
    # ══════════════════════════════════════════════════════════════════════════
    def _mk_ind(self, parent, label, val, col):
        bg   = CARD_BG
        box  = tk.Frame(parent, bg=bg, width=138, height=70,
                        relief="solid", bd=1, highlightbackground=CARD_BOR)
        box.grid(row=0, column=col, padx=5)
        box.grid_propagate(False)
        # accent bar on top
        tk.Frame(box, bg=AZ3, height=3).pack(fill="x")
        lv = tk.Label(box, text=val, bg=bg, fg=TXT_PRI,
                      font=("Segoe UI",18,"bold"))
        lv.pack(expand=True)
        tk.Label(box, text=label, bg=bg, fg="#8B949E",
                 font=("Segoe UI",8)).pack(pady=(0,6))
        return lv

    def _log(self, msg, color="#9DC3E6"):
        cur    = getattr(self, "_current_page", "ocr")
        target = self._log_a if cur == "anal" else self._log_w
        target.config(state="normal")
        target.insert("end", f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n")
        target.see("end"); target.config(state="disabled")
        self._lbl_status.config(text=f"  {msg[:90]}")
    def _log_a_write(self, msg):
        self._log_a.config(state="normal")
        self._log_a.insert("end", f"[{datetime.now().strftime('%H:%M:%S')}] {msg}\n")
        self._log_a.see("end"); self._log_a.config(state="disabled")
        self._lbl_status.config(text=f"  {msg[:90]}")

    def _set_prog(self, v, txt=""):
        self._prog["value"] = v; self._lbl_pct.config(text=f"{v}%  {txt}")
    def _set_prog_a(self, v, txt=""):
        self._prog_a["value"] = v; self._lbl_pct_a.config(text=f"{v}%  {txt}")
        self._lbl_fase_a.config(text=txt)

    # ══════════════════════════════════════════════════════════════════════════
    # WORKERS — OCR
    # ══════════════════════════════════════════════════════════════════════════
    def _gutter_completar_corpus(self):
        """Detecta y reconstruye palabras cortadas por costura en todos los .txt."""
        if not getattr(ST, "out_dir", None):
            messagebox.showwarning("Sin carpeta",
                "Confirma la configuración y ejecuta la extracción primero."); return
        api_key, _ = _resolver_api_key_modelo("ocr_mejora")
        if not api_key:
            messagebox.showwarning("Sin IA",
                "Activa la IA y configura una API key para usar esta función.\n"
                "La reconstrucción de palabras requiere un modelo de lenguaje."); return

        ocr_dir = ST.out_dir / "03_ocr"
        if ocr_dir.exists():
            txt_dirs = [d for d in ocr_dir.iterdir() if d.is_dir()]
        else:
            txt_dirs = list({f.parent for f in ST.out_dir.rglob("*.txt") if f.is_file()})

        if not txt_dirs:
            messagebox.showinfo("Sin archivos",
                "No se encontraron archivos de texto OCR."); return

        total = sum(len(list(d.glob("*.txt"))) for d in txt_dirs)
        if not messagebox.askyesno("Completar costura",
            f"Se analizarán {total} páginas en busca de palabras cortadas por costura.\n\n"
            "Las palabras reconstruidas se marcarán con ⟦palabra⟧ en el texto\n"
            "y aparecerán en rojo en las exportaciones DOCX.\n\n"
            "¿Continuar?"): return

        self._put(tipo="fase", txt="Completando palabras cortadas por costura…")
        threading.Thread(target=self._worker_gutter, args=(txt_dirs, api_key),
                         daemon=True).start()

    def _worker_gutter(self, txt_dirs, api_key):
        from core.gutter_completion import estadisticas, reconstruir_texto
        modelo = "claude-haiku-4-5-20251001"
        total_fragmentos = 0
        total_reconstruidos = 0

        for txt_dir in txt_dirs:
            for txt_path in sorted(txt_dir.glob("*.txt")):
                try:
                    texto = txt_path.read_text("utf-8", errors="replace")
                    texto_rec, frags = reconstruir_texto(texto, api_key, modelo)
                    stats = estadisticas(frags)
                    if stats["reconstruidos"] > 0:
                        # Guardar backup
                        backup = txt_path.with_suffix(".txt.orig")
                        if not backup.exists():
                            backup.write_text(texto, encoding="utf-8")
                        txt_path.write_text(texto_rec, encoding="utf-8")
                        total_fragmentos   += stats["total_fragmentos"]
                        total_reconstruidos += stats["reconstruidos"]
                        self._put(tipo="log",
                                  texto=f"  {txt_path.stem}: "
                                        f"{stats['reconstruidos']}/{stats['total_fragmentos']} "
                                        f"palabras reconstruidas")
                except Exception as e:
                    self._put(tipo="log",
                              texto=f"  ⚠ Error en {txt_path.stem}: {e}")

        self._put(tipo="fase", txt="")
        self.after(0, lambda: messagebox.showinfo(
            "Costura completada",
            f"Análisis completado.\n\n"
            f"Fragmentos detectados: {total_fragmentos}\n"
            f"Palabras reconstruidas: {total_reconstruidos}\n\n"
            f"Los archivos originales se guardaron con extensión .txt.orig\n"
            f"Las palabras generadas están marcadas con ⟦⟧"))

    def _renormalizar_textos(self):
        """Re-aplica la normalización OCR a todos los .txt ya extraídos."""
        if not getattr(ST, "out_dir", None):
            messagebox.showwarning("Sin carpeta",
                "Confirma la configuración y ejecuta la extracción primero."); return

        # Buscar en 03_ocr/ (estructura estándar) o cualquier subcarpeta con .txt
        ocr_dir = ST.out_dir / "03_ocr"
        if ocr_dir.exists():
            txt_dirs = [d for d in ocr_dir.iterdir() if d.is_dir()]
        else:
            txt_dirs = [d for d in ST.out_dir.rglob("*.txt")
                        if d.is_file()]
            # Convertir a carpetas únicas
            txt_dirs = list({f.parent for f in txt_dirs})

        if not txt_dirs:
            messagebox.showinfo("Sin archivos",
                "No se encontraron archivos de texto.\n"
                "Extrae el OCR primero en la pestaña OCR."); return

        total_txt = sum(len(list(d.glob("*.txt"))) for d in txt_dirs)
        msg_norm = (f"Se normalizarán {total_txt} archivos .txt en "
                    f"{len(txt_dirs)} carpeta(s).\n\n"
                    "Se guardará copia del original como .txt.orig\n"
                    "¿Continuar?")
        if not messagebox.askyesno("Re-normalizar", msg_norm): return

        prog_win = tk.Toplevel(self)
        prog_win.title("Normalizando textos OCR…")
        prog_win.geometry("420x130")
        prog_win.resizable(False, False)
        prog_win.grab_set()
        tk.Label(prog_win, text="Normalizando textos OCR",
                 font=("Segoe UI", 10, "bold")).pack(pady=(14, 4))
        lbl_p = tk.Label(prog_win, text="Preparando…", font=("Segoe UI", 9))
        lbl_p.pack()
        bar = ttk.Progressbar(prog_win, mode="determinate", length=380)
        bar.pack(pady=8)

        def worker():
            total_archivos = 0
            total_cambios  = 0
            all_txts = []
            for d in txt_dirs:
                all_txts.extend(sorted(d.glob("*.txt")))

            for i, txt_path in enumerate(all_txts, 1):
                pct = int(i / max(len(all_txts), 1) * 100)
                self.after(0, lambda p=pct: bar.config(value=p))
                self.after(0, lambda n=txt_path.name: lbl_p.config(text=n))
                try:
                    from core.ocr_normalizer import normalizar_archivo
                    stats = normalizar_archivo(txt_path, guardar_original=True)
                    total_archivos += 1
                    total_cambios  += stats["chars_cambiados"]
                except Exception:
                    pass

            self.after(0, prog_win.destroy)
            msg_done = (f"Archivos normalizados: {total_archivos}\n"
                        f"Caracteres corregidos: ~{total_cambios:,}\n\n"
                        "Los originales se conservan como .txt.orig")
            self.after(0, lambda m=msg_done: messagebox.showinfo(
                "Normalización completada", m))

        threading.Thread(target=worker, daemon=True).start()

    # ── Kraken handlers ───────────────────────────────────────────────────────
    def _ocr_verificar_kraken(self):
        """Verifica si Kraken + modelo están disponibles y actualiza el label."""
        try:
            from core.ocr_kraken import _buscar_modelo, kraken_disponible
            if kraken_disponible():
                modelo = _buscar_modelo()
                nombre = Path(modelo).name if modelo else "modelo"
                self._lbl_kraken_ok.config(text=f"✅ {nombre}", fg=VERDE)
                if modelo:
                    self._var_kraken_modelo.set(str(modelo))
            else:
                # Verificar via subprocess (Kraken está en venv separado, no en Python principal)
                import subprocess

                from core.ocr_kraken import _python_kraken
                chk = subprocess.run([_python_kraken(), "-c", "import kraken"],
                                     capture_output=True, timeout=10)
                if chk.returncode == 0:
                    self._lbl_kraken_ok.config(text="⚠ Sin modelo — descarga CATMuS-Print", fg="#D97706")
                else:
                    self._lbl_kraken_ok.config(text="✗ Kraken no instalado en venv", fg=ROJO)
        except Exception as e:
            self._lbl_kraken_ok.config(text=f"✗ {e}", fg=ROJO)

    def _ocr_elegir_modelo_kraken(self):
        """Abre diálogo para seleccionar un archivo .mlmodel de Kraken."""
        from tkinter import filedialog
        ruta = filedialog.askopenfilename(
            title="Seleccionar modelo Kraken (.mlmodel)",
            filetypes=[("Kraken model", "*.mlmodel"), ("Todos", "*.*")],
        )
        if ruta:
            self._var_kraken_modelo.set(ruta)
            self._lbl_kraken_ok.config(text=f"✅ {Path(ruta).name}", fg=VERDE)

    def _ocr_descargar_catmus(self):
        """Descarga el modelo CATMuS-Print Large en un thread."""
        self._lbl_kraken_ok.config(text="⏳ Descargando (~300 MB)…", fg="#8B949E")

        def _worker():
            try:
                from core.ocr_kraken import descargar_modelo_catmus
                def cb(msg):
                    self.after(0, lambda m=msg: self._lbl_kraken_ok.config(text=m, fg="#8B949E"))
                ruta = descargar_modelo_catmus(callback=cb)
                self.after(0, lambda r=ruta: (
                    self._var_kraken_modelo.set(r),
                    self._lbl_kraken_ok.config(text=f"✅ {Path(r).name}", fg=VERDE),
                ))
            except Exception as e:
                msg = str(e)
                if "No module named kraken" in msg:
                    msg = ("✗ Kraken no instalado — incompatible con torch 2.x.\n"
                           "Para instalar: pip install kraken (requiere torch<1.11)")
                elif "Error descargando" in msg:
                    msg = "✗ " + msg.split("\n")[0]
                else:
                    msg = f"✗ {msg[:120]}"
                self.after(0, lambda m=msg: self._lbl_kraken_ok.config(
                    text=m, fg=ROJO, wraplength=320))

        threading.Thread(target=_worker, daemon=True).start()

    # ── Ollama handlers ───────────────────────────────────────────────────────
    # Velocidad de referencia en segundos por página (1 worker).
    # Se actualiza con _ocr_calibrar_kraken().
    _KRAKEN_SEG_PAG = 60.0

    def _ocr_contar_paginas_corpus(self) -> int:
        """Cuenta páginas pendientes de OCR en el corpus activo."""
        if not ST.out_dir or not ST.archivos_sel:
            return 0
        total = 0
        dir_img = ST.out_dir / "02_imagenes"
        dir_ocr = ST.out_dir / "03_ocr"
        for pdf in ST.archivos_sel:
            nombre = pdf.stem
            img_dir = dir_img / nombre
            ocr_dir = dir_ocr / nombre
            if img_dir.exists():
                imgs = list(img_dir.glob("*.png"))
                # Contar solo las que aún no tienen .txt
                if ocr_dir.exists():
                    hechas = {p.stem for p in ocr_dir.glob("*.txt")}
                    total += sum(1 for i in imgs if i.stem not in hechas)
                else:
                    total += len(imgs)
            else:
                # Sin imágenes aún: estimamos por tamaño del PDF (1 pág ≈ 150 KB)
                try:
                    total += max(1, pdf.stat().st_size // 150_000)
                except Exception:
                    total += 100
        return total

    def _ocr_actualizar_estimacion(self, *_):
        """Recalcula y muestra la estimación de tiempo según workers y corpus."""
        if not hasattr(self, "_lbl_kraken_est"):
            return
        try:
            workers = max(1, int(self._var_kraken_workers.get()))
        except (ValueError, tk.TclError):
            return
        paginas = self._ocr_contar_paginas_corpus()
        if paginas == 0:
            self._lbl_kraken_est.config(
                text="— (carga un proyecto para estimar)", fg="#8B949E")
            return
        seg_total = (paginas * self._KRAKEN_SEG_PAG) / workers
        horas  = int(seg_total // 3600)
        minutos = int((seg_total % 3600) // 60)
        if horas > 0:
            tiempo_txt = f"~{horas}h {minutos:02d}min"
        else:
            tiempo_txt = f"~{minutos} min"
        color = "#22C55E" if seg_total < 3600 else ("#F59E0B" if seg_total < 7200 else "#EF4444")
        self._lbl_kraken_est.config(
            text=f"{tiempo_txt}  ({paginas} páginas, {workers} proceso{'s' if workers>1 else ''})",
            fg=color)

    def _ocr_calibrar_kraken(self):
        """Mide el tiempo real de Kraken en una página del corpus y actualiza la estimación."""
        if not ST.out_dir or not ST.archivos_sel:
            messagebox.showwarning("Sin proyecto", "Carga un proyecto primero.")
            return
        # Buscar la primera imagen disponible
        img_path = None
        for pdf in ST.archivos_sel:
            img_dir = ST.out_dir / "02_imagenes" / pdf.stem
            if img_dir.exists():
                imgs = sorted(img_dir.glob("*.png"))
                if imgs:
                    img_path = imgs[0]
                    break
        if not img_path:
            messagebox.showwarning(
                "Sin imágenes",
                "No hay imágenes generadas.\nProcesa primero al menos 1 página con Ruta 1.")
            return

        self._lbl_kraken_est.config(text="⏳ Calibrando (1 página)…", fg="#F59E0B")
        self.update_idletasks()

        def _worker():
            import time

            from core.ocr_kraken import ocr_kraken
            modelo = getattr(self, "_var_kraken_modelo", tk.StringVar()).get() or None
            t0 = time.perf_counter()
            try:
                ocr_kraken(str(img_path), modelo)
                seg = time.perf_counter() - t0
                # Actualizar referencia y recalcular
                BashkarApp._KRAKEN_SEG_PAG = round(seg, 1)
                self.after(0, lambda: self._lbl_kraken_est.config(
                    text=f"✅ Calibrado: {seg:.0f} seg/página — recalculando…",
                    fg="#22C55E"))
                self.after(200, self._ocr_actualizar_estimacion)
            except Exception as e:
                self.after(0, lambda err=str(e): self._lbl_kraken_est.config(
                    text=f"✗ Error calibrando: {err}", fg="#EF4444"))

        threading.Thread(target=_worker, daemon=True).start()

    def _ocr_detectar_ollama(self):
        """Detecta modelos de visión disponibles en Ollama y puebla el combobox."""
        def _worker():
            try:
                from core.ocr_ollama_local import listar_modelos_vision
                modelos = listar_modelos_vision()
                def _update(ms=modelos):
                    if ms:
                        self._cmb_ollama["values"] = ms
                        if not self._var_ollama_modelo.get() or \
                                self._var_ollama_modelo.get() not in ms:
                            self._var_ollama_modelo.set(ms[0])
                        self._lbl_ollama_ok.config(
                            text=f"✅ {len(ms)} modelo(s)", fg=VERDE)
                    else:
                        self._lbl_ollama_ok.config(
                            text="⚠ Ollama sin modelos de visión", fg="#D97706")
                self.after(0, _update)
            except Exception as e:
                self.after(0, lambda err=e: self._lbl_ollama_ok.config(
                    text=f"✗ {err}", fg=ROJO))

        threading.Thread(target=_worker, daemon=True).start()

    # ─────────────────────────────────────────────────────────────────────────
    def _ocr_aplicar_zonas(self, texto: str, numero: str, pagina: str,
                            det_auto: bool = False, api_key: str = "") -> str:
        """
        Filtra el texto de una página usando las zonas etiquetadas.

        - Toma solo las zonas con ocr=True, en orden de arriba a abajo (y0).
        - Si la página no tiene etiquetas y det_auto=True, detecta automáticamente.
        - El texto de cada zona se extrae por posición relativa en las líneas del OCR.
        - Retorna el texto ordenado según el flujo de lectura de las zonas.
        """
        from core.zone_labeler import TIPOS_ZONA, cargar_pagina

        if not ST.out_dir:
            return texto

        pag_data = cargar_pagina(ST.out_dir, numero, pagina)
        zonas_ocr = []

        if pag_data and pag_data.zonas:
            # Filtrar solo zonas que deben procesarse con OCR, ordenadas por y0
            zonas_ocr = sorted(
                [z for z in pag_data.zonas if TIPOS_ZONA.get(z.tipo, {}).get("ocr", True)],
                key=lambda z: (z.y0, z.x0)
            )
        elif det_auto and api_key:
            # Detección automática en páginas sin etiquetar
            try:
                img_path = self._etz_get_img_path(numero, pagina)
                if img_path:
                    from core.zone_labeler import detectar_zonas_claude
                    zonas_detectadas = detectar_zonas_claude(img_path, api_key)
                    zonas_ocr = sorted(
                        [z for z in zonas_detectadas if TIPOS_ZONA.get(z.tipo, {}).get("ocr", True)],
                        key=lambda z: (z.y0, z.x0)
                    )
                    self._put(tipo="log",
                              texto=f"    🤖 {len(zonas_ocr)} zonas detectadas automáticamente")
            except Exception:
                pass

        if not zonas_ocr:
            return texto  # sin zonas → devolver texto completo

        # Dividir el texto en líneas y asignar cada línea a la zona más cercana
        lineas = texto.splitlines()
        total_lineas = len(lineas) or 1
        bloques = {i: [] for i in range(len(zonas_ocr))}

        for i_linea, linea in enumerate(lineas):
            pos_rel = i_linea / total_lineas  # posición 0.0-1.0 en la página
            # Buscar la zona que más se superpone con esta posición vertical
            mejor_zona = None
            mejor_dist = float("inf")
            for i_zona, zona in enumerate(zonas_ocr):
                if zona.y0 <= pos_rel <= zona.y1:
                    # Dentro de la zona — asignar directamente
                    mejor_zona = i_zona
                    break
                dist = min(abs(pos_rel - zona.y0), abs(pos_rel - zona.y1))
                if dist < mejor_dist:
                    mejor_dist = dist
                    mejor_zona = i_zona
            if mejor_zona is not None:
                bloques[mejor_zona].append(linea)

        # Reconstruir texto en orden de lectura de las zonas
        partes = []
        for i_zona in range(len(zonas_ocr)):
            bloque = "\n".join(bloques[i_zona]).strip()
            if bloque:
                partes.append(bloque)

        return "\n\n".join(partes) if partes else texto

    def _ocr_preview_preprocesamiento(self):
        """Abre ventana con antes/después del preprocesamiento sobre la primera imagen disponible."""
        if not ST.out_dir or not ST.archivos_sel:
            messagebox.showwarning("Sin archivos", "Configura el corpus primero.")
            return
        # Buscar primera imagen disponible
        img_path = None
        for archivo in ST.archivos_sel:
            img_dir = ST.out_dir / "02_imagenes" / archivo.stem
            if img_dir.exists():
                hits = sorted(img_dir.glob("*.png"))
                if hits:
                    img_path = hits[0]
                    break
        if img_path is None:
            messagebox.showwarning("Sin imágenes",
                "Ejecuta primero la extracción OCR para generar las imágenes.")
            return

        from PIL import Image, ImageTk

        from core.image_preprocessor import preprocesar_para_ocr

        win = tk.Toplevel(self)
        win.title(f"Preview preprocesamiento — {img_path.name}")
        win.geometry("900x500")
        win.configure(bg=CONTENT_BG)
        win.grab_set()

        tk.Label(win, text="Original", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).grid(row=0, column=0, pady=(8, 2))
        tk.Label(win, text="Procesada", bg=CONTENT_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).grid(row=0, column=1, pady=(8, 2))

        cv_orig = tk.Canvas(win, bg="#000", width=420, height=420, highlightthickness=0)
        cv_orig.grid(row=1, column=0, padx=8, pady=4)
        cv_proc = tk.Canvas(win, bg="#000", width=420, height=420, highlightthickness=0)
        cv_proc.grid(row=1, column=1, padx=8, pady=4)

        def _mostrar():
            img_o = Image.open(img_path).convert("RGB")
            img_p = preprocesar_para_ocr(
                img_o.copy(),
                deskew_en    = self._var_pre_deskew.get(),
                enhance_en   = self._var_pre_enhance.get(),
                despeckle_en = self._var_pre_despeckle.get(),
            )
            for cv, im in [(cv_orig, img_o), (cv_proc, img_p)]:
                im.thumbnail((420, 420), Image.LANCZOS)
                tk_img = ImageTk.PhotoImage(im)
                cv._ref = tk_img
                cv.delete("all")
                cv.create_image(210, 210, anchor="center", image=tk_img)

        _mostrar()

        # Checkboxes dentro del preview para ajustar en tiempo real
        ctrl = tk.Frame(win, bg=CONTENT_BG)
        ctrl.grid(row=2, column=0, columnspan=2, pady=6)
        for txt, var in [("Deskew", self._var_pre_deskew),
                          ("CLAHE",  self._var_pre_enhance),
                          ("Despeckle", self._var_pre_despeckle)]:
            ttk.Checkbutton(ctrl, text=txt, variable=var,
                             command=_mostrar).pack(side="left", padx=10)
        ttk.Button(ctrl, text="↺ Actualizar", command=_mostrar).pack(side="left", padx=10)
        ttk.Button(ctrl, text="Cerrar", command=win.destroy).pack(side="left", padx=10)

    def _start_ocr(self):
        if ST.pdf_dir is None:
            messagebox.showwarning("Sin config","Confirma la configuración primero."); return
        self._btn_ocr.config(state="disabled")
        threading.Thread(target=self._worker_ocr, daemon=True).start()

    def _worker_ocr(self):
        from core.ocr_engine import (
            EXTS_IMAGEN,
            analizar_pdf,
            imagenes_a_texto,
            ocr_pagina,
            pdf_a_imagenes,
        )

        # Modo subcarpetas: cada "archivo" en archivos_sel es una carpeta con PDFs
        # Expandimos la lista: cada PDF de la subcarpeta se procesa como una página
        if ST.input_tipo == "carpetas":
            archivos_orig = ST.archivos_sel
            archivos_expandidos = []
            for sc in archivos_orig:
                pdfs = sorted(Path(sc).glob("*.pdf"))
                archivos_expandidos.append({
                    "nombre": Path(sc).name,   # nombre del número
                    "pdfs":   pdfs,             # lista de PDFs de la carpeta
                    "carpeta": Path(sc),
                })
            self._put(tipo="log",
                texto=f"📁 Modo subcarpetas: {len(archivos_expandidos)} número(s) detectado(s)")
            self._worker_ocr_carpetas(archivos_expandidos)
            return

        archivos   = ST.archivos_sel; total = len(archivos)
        dpi, lang  = self._var_dpi.get(), self._var_lang.get()
        ruta_ocr   = getattr(self, "_var_ruta_ocr", None)
        ruta_ocr      = ruta_ocr.get() if ruta_ocr else "tesseract"
        usar_etiq     = getattr(self, "_var_ocr_usar_etiquetas", None)
        usar_etiq     = usar_etiq.get() if usar_etiq else False
        det_auto      = getattr(self, "_var_ocr_det_auto", None)
        det_auto      = det_auto.get() if det_auto else False
        out           = ST.out_dir
        dir_img       = out/"02_imagenes"; dir_ocr = out/"03_ocr"
        dir_img.mkdir(exist_ok=True); dir_ocr.mkdir(exist_ok=True)
        meta_rows, errores = [], []

        # Leer proveedor/modelo de visión si se seleccionó Ruta 2
        _ocr_vision_prov  = getattr(self, "_ocr_vision_prov",
                                     tk.StringVar(value="claude")).get()
        _ocr_vision_model = getattr(self, "_ocr_vision_model",
                                     tk.StringVar(value="claude-sonnet-4-6")).get()

        RUTA_LABELS = {
            "tesseract": "Ruta 1 · Tesseract propio",
            "vision_ia": f"Ruta 2 · {_ocr_vision_prov}/{_ocr_vision_model}",
            "claude":    "Ruta 2 · Claude Vision (legado)",
            "bnc":       "Ruta 3 · Texto BNC + reconstrucción",
            "kraken":    "Ruta 4 · Kraken CATMuS-Print",
            "ollama":    "Ruta 5 · Ollama Vision",
        }
        etiq_label = " · con zonas etiquetadas" if usar_etiq else ""
        self._put(tipo="log", texto=f"📄 {total} archivo(s) · {RUTA_LABELS.get(ruta_ocr, ruta_ocr)}{etiq_label}")
        self._put(tipo="fase", txt="Detectando tipos…")

        # En Ruta 1 (tesseract), Ruta 2 (vision_ia) forzamos re-OCR.
        # En Ruta 3 (bnc) usamos el texto embebido con reconstrucción de líneas.
        modos = {}
        for p in archivos:
            if ST.input_tipo == "img" or p.suffix.lower() in EXTS_IMAGEN:
                modos[p.name] = "imagen"
            elif ruta_ocr in ("tesseract", "vision_ia", "claude", "kraken", "ollama"):
                # Forzar re-OCR desde imágenes, ignorar texto BNC
                modos[p.name] = "escaneado"
            else:
                # Ruta 3: usar texto embebido si existe
                info = analizar_pdf(p)
                modos[p.name] = "digital" if info["tiene_texto"] else "escaneado"

        n_dig = sum(1 for m in modos.values() if m=="digital")
        n_ocr = sum(1 for m in modos.values() if m=="escaneado")
        n_img = sum(1 for m in modos.values() if m=="imagen")
        parts = []
        if n_dig: parts.append(f"⚡ {n_dig} digital(es)")
        if n_ocr: parts.append(f"🔍 {n_ocr} escaneado(s)")
        if n_img: parts.append(f"🖼️ {n_img} imagen(es)")
        self._put(tipo="log", texto="   "+"|".join(parts))
        ST.modos_detec = modos

        for idx, archivo in enumerate(archivos):
            nombre = archivo.stem; modo = modos.get(archivo.name,"escaneado")
            icono  = {"digital":"⚡","escaneado":"🔍","imagen":"🖼️"}.get(modo,"📄")
            self._put(tipo="log", texto=f"{icono} {archivo.name} ({idx+1}/{total})")
            txt_dir = dir_ocr/nombre; txt_dir.mkdir(exist_ok=True)
            try:
                if modo == "digital":
                    # Ruta 3: texto BNC con reconstrucción por coordenadas (alto_reconstructor)
                    # Usa las coordenadas X/Y de cada span para reconstruir el orden de columnas
                    # en vez de confiar en el orden lineal del PDF (que mezcla columnas).
                    import fitz

                    from core.alto_reconstructor import reconstruir_texto_pagina
                    from core.ocr_normalizer import normalizar_texto_ocr

                    _MARCA_BNC = "Digitalizado Biblioteca Nacional de Colombia"

                    rows = []
                    doc_bnc = fitz.open(str(archivo))
                    for i, page in enumerate(doc_bnc):
                        pagina_id = f"p{i+1:04d}"
                        tp = txt_dir / f"{pagina_id}.txt"
                        if not tp.exists():
                            try:
                                resultado = reconstruir_texto_pagina(
                                    page,
                                    ignorar_ocr_basura=True,
                                )
                                texto = resultado.get("texto", "")
                                # Eliminar marca de digitalización de la BNC
                                texto = texto.replace(_MARCA_BNC, "").strip()
                                texto = normalizar_texto_ocr(texto)
                            except Exception:
                                texto = page.get_text("text")
                            tp.write_text(texto, encoding="utf-8")
                        else:
                            texto = tp.read_text("utf-8", errors="replace")
                        rows.append({
                            "numero":   nombre,
                            "pagina":   pagina_id,
                            "txt_path": str(tp),
                            "palabras": len(texto.split()),
                            "confianza": None,
                            "revision": False,
                            "metodo":   "bnc_coordenadas",
                        })
                    doc_bnc.close()
                    meta_rows.extend(rows)
                    self._put(tipo="log",
                        texto=f"  ✅ {len(rows)} pág · BNC + reconstrucción por coordenadas")

                elif modo == "escaneado":
                    img_dir_n = dir_img/nombre
                    try:
                        imgs = pdf_a_imagenes(archivo, img_dir_n, dpi)
                    except Exception as e:
                        if "poppler" in str(e).lower():
                            try:
                                from instalar import instalar_poppler_windows
                                instalar_poppler_windows()
                                from core.ocr_engine import _get_poppler_path as _gp
                                nr = _gp()
                                if nr: os.environ["PATH"]=nr+os.pathsep+os.environ.get("PATH","")
                                imgs = pdf_a_imagenes(archivo, img_dir_n, dpi)
                            except Exception as e2:
                                self._put(tipo="log",texto=f"  ❌ Poppler: {e2}"); errores.append(archivo.name); continue
                        else:
                            self._put(tipo="log",texto=f"  ❌ {e}"); errores.append(archivo.name); continue
                    np_total = len(imgs)

                    # ── Preprocesamiento opcional — guarda en carpeta separada ──
                    # Las originales a color se preservan en 02_imagenes/<nombre>/
                    # Las procesadas para OCR van a 02_imagenes_ocr/<nombre>/
                    pre_deskew    = getattr(self, "_var_pre_deskew",    None)
                    pre_enhance   = getattr(self, "_var_pre_enhance",   None)
                    pre_despeckle = getattr(self, "_var_pre_despeckle", None)
                    do_pre = (
                        (pre_deskew    and pre_deskew.get()) or
                        (pre_enhance   and pre_enhance.get()) or
                        (pre_despeckle and pre_despeckle.get())
                    )
                    if do_pre:
                        try:
                            from PIL import Image as _PIL

                            from core.image_preprocessor import preprocesar_para_ocr
                            dir_ocr_imgs = out / "02_imagenes_ocr" / nombre
                            dir_ocr_imgs.mkdir(parents=True, exist_ok=True)
                            imgs_procesadas = []
                            for ip in imgs:
                                img_pil = _PIL.open(ip).convert("RGB")
                                img_proc = preprocesar_para_ocr(
                                    img_pil,
                                    deskew_en    = bool(pre_deskew    and pre_deskew.get()),
                                    enhance_en   = bool(pre_enhance   and pre_enhance.get()),
                                    despeckle_en = bool(pre_despeckle and pre_despeckle.get()),
                                )
                                # Guardar procesada en carpeta OCR, original intacta
                                dest = dir_ocr_imgs / ip.name
                                img_proc.save(str(dest))
                                imgs_procesadas.append(dest)
                            imgs = imgs_procesadas
                            self._put(tipo="log", texto=f"  🔧 Preprocesamiento aplicado a {len(imgs)} imágenes")
                        except Exception as ep:
                            self._put(tipo="log", texto=f"  ⚠ Preprocesamiento omitido: {ep}")

                    if ruta_ocr in ("vision_ia", "claude"):
                        # Ruta 2: IA de visión multiproveedor (Claude/GPT-4o/Gemini/Ollama)
                        prov  = _ocr_vision_prov  if ruta_ocr == "vision_ia" else "claude"
                        model = _ocr_vision_model if ruta_ocr == "vision_ia" else ""
                        api_key = ST.api_keys.get(prov, "") or ST.api_key
                        if prov != "ollama" and not api_key:
                            self._put(tipo="log",
                                texto=f"  ❌ Ruta 2 requiere API key de {prov}. Configúrala en ⚙ Configuración.")
                            errores.append(archivo.name); continue

                        for pi, ip in enumerate(imgs):
                            tp = txt_dir/(ip.stem+".txt")
                            if tp.exists():
                                texto = tp.read_text("utf-8", errors="replace"); conf = 95.0
                            else:
                                try:
                                    # Usar el proveedor elegido
                                    texto = _ocr_vision_multiproveedor(
                                        ip, prov, api_key, model)
                                    conf  = 95.0
                                except Exception as ec:
                                    self._put(tipo="log",
                                        texto=f"    ⚠ {prov} falló en {ip.stem}: {ec}. Usando Tesseract.")
                                    texto, conf = ocr_pagina(ip, lang=lang)
                                tp.write_text(texto, "utf-8")
                            meta_rows.append({"numero":nombre,"pagina":ip.stem,"txt_path":str(tp),
                                              "palabras":len(texto.split()),"confianza":conf,
                                              "revision":False,"metodo":f"vision_{prov}"})
                            pct=int(((idx*np_total+pi+1)/(total*max(np_total,1)))*100)
                            self._put(tipo="prog",val=pct,txt=f"{nombre}·pg{pi+1}/{np_total}")
                        self._put(tipo="log",texto=f"  ✅ {np_total} pág · {prov}/{model}")

                    elif ruta_ocr == "kraken":
                        # Ruta 4: Kraken CATMuS-Print
                        from core.ocr_kraken import ocr_kraken_lote
                        modelo_k = self._var_kraken_modelo.get() or None
                        try:
                            workers_k = max(1, int(self._var_kraken_workers.get()))
                        except Exception:
                            workers_k = 3
                        try:
                            timeout_k = max(60, int(self._var_kraken_timeout.get()))
                        except Exception:
                            timeout_k = 600
                        rutas_imgs = [str(ip) for ip in imgs]
                        resultados_k = ocr_kraken_lote(
                            rutas_imgs, modelo_path=modelo_k,
                            workers=workers_k, timeout=timeout_k,
                            callback=lambda i,t,r,ok: self._put(
                                tipo="prog",
                                val=int(((idx*t+i)/(total*max(t,1)))*100),
                                txt=f"{nombre}·pg{i}/{t}")
                        )
                        for pi, (ip, res_k) in enumerate(zip(imgs, resultados_k)):
                            tp = txt_dir/(ip.stem+".txt")
                            if res_k["ok"]:
                                texto = res_k["texto"]; conf = round(res_k["confianza"]*100, 1)
                                tp.write_text(texto, "utf-8")
                            else:
                                err_msg = res_k['error'] or ""
                                self._put(tipo="log", texto=f"    ⚠ Kraken falló en {ip.stem}: {err_msg[:120]}")
                                # Intentar fallback Tesseract; si no está disponible, guardar vacío
                                try:
                                    texto, conf = ocr_pagina(ip, lang=lang)
                                    tp.write_text(texto, "utf-8")
                                    self._put(tipo="log", texto="      → Tesseract usado como respaldo")
                                except Exception as ef:
                                    texto, conf = "", 0.0
                                    tp.write_text("", "utf-8")
                                    self._put(tipo="log", texto=f"      → Sin respaldo disponible ({ef}). Página marcada para revisión manual.")
                            meta_rows.append({"numero":nombre,"pagina":ip.stem,"txt_path":str(tp),
                                              "palabras":len(texto.split()),"confianza":conf,
                                              "revision":bool(conf and conf<60),"metodo":"kraken"})
                        self._put(tipo="log",texto=f"  ✅ {np_total} pág · Kraken CATMuS-Print")

                    elif ruta_ocr == "ollama":
                        # Ruta 5: Ollama Vision local
                        from core.ocr_ollama_local import ocr_ollama_lote
                        modelo_o = getattr(self, "_var_ollama_modelo",
                                           tk.StringVar(value="qwen2.5vl:7b")).get()
                        rutas_imgs = [str(ip) for ip in imgs]
                        resultados_o = ocr_ollama_lote(
                            rutas_imgs, modelo=modelo_o,
                            callback=lambda i,t,r,ok: self._put(
                                tipo="prog",
                                val=int(((idx*t+i)/(total*max(t,1)))*100),
                                txt=f"{nombre}·pg{i}/{t}")
                        )
                        for pi, (ip, res_o) in enumerate(zip(imgs, resultados_o)):
                            tp = txt_dir/(ip.stem+".txt")
                            if res_o["ok"]:
                                texto = res_o["texto"]; conf = round(res_o["confianza"]*100, 1)
                                tp.write_text(texto, "utf-8")
                            else:
                                self._put(tipo="log", texto=f"    ⚠ Ollama falló en {ip.stem}: {res_o['error']}. Usando Tesseract.")
                                texto, conf = ocr_pagina(ip, lang=lang)
                                tp.write_text(texto, "utf-8")
                            meta_rows.append({"numero":nombre,"pagina":ip.stem,"txt_path":str(tp),
                                              "palabras":len(texto.split()),"confianza":conf,
                                              "revision":bool(conf and conf<60),"metodo":"ollama"})
                        self._put(tipo="log",texto=f"  ✅ {np_total} pág · Ollama Vision")

                    else:
                        # Ruta 1: Tesseract propio (ignora texto BNC).
                        # Si la página tiene zonas etiquetadas (05_etiquetas/),
                        # se OCR-ea por zonas en orden de lectura (estilo FineReader).
                        from core.layout_tesseract import ocr_pagina_con_zonas
                        n_zonal = 0
                        for pi, ip in enumerate(imgs):
                            tp = txt_dir/(ip.stem+".txt")
                            metodo_pag = "ocr"
                            # Forzar re-OCR si la fuente era texto BNC (sobreescribir)
                            if tp.exists() and ruta_ocr == "tesseract":
                                # Solo reusar si el metodo ya fue tesseract (no BNC)
                                meta_csv = out/"04_analisis"/"ocr_metadatos.csv"
                                fue_tesseract = False
                                if meta_csv.exists():
                                    try:
                                        import pandas as _pd
                                        _m = _pd.read_csv(meta_csv)
                                        row_m = _m[(_m["numero"]==nombre) & (_m["pagina"]==ip.stem)]
                                        if not row_m.empty and row_m.iloc[0].get("metodo","") in ("ocr","ocr_zonas"):
                                            fue_tesseract = True
                                            metodo_pag = row_m.iloc[0].get("metodo","ocr")
                                    except Exception:
                                        pass
                                if fue_tesseract:
                                    texto=tp.read_text("utf-8",errors="replace"); conf=None
                                else:
                                    texto, conf, con_z = ocr_pagina_con_zonas(
                                        ip, out, nombre, ip.stem, lang=lang)
                                    tp.write_text(texto,"utf-8")
                                    if con_z:
                                        metodo_pag = "ocr_zonas"; n_zonal += 1
                            else:
                                texto, conf, con_z = ocr_pagina_con_zonas(
                                    ip, out, nombre, ip.stem, lang=lang)
                                tp.write_text(texto,"utf-8")
                                if con_z:
                                    metodo_pag = "ocr_zonas"; n_zonal += 1
                            meta_rows.append({"numero":nombre,"pagina":ip.stem,"txt_path":str(tp),
                                              "palabras":len(texto.split()),"confianza":conf,
                                              "revision":bool(conf and conf<60),"metodo":metodo_pag})
                            pct=int(((idx*np_total+pi+1)/(total*max(np_total,1)))*100)
                            self._put(tipo="prog",val=pct,txt=f"{nombre}·pg{pi+1}/{np_total}")
                        extra_z = f" ({n_zonal} por zonas)" if n_zonal else ""
                        self._put(tipo="log",texto=f"  ✅ {np_total} pág · Tesseract{extra_z}")
                else:
                    rows=imagenes_a_texto(archivo.parent,txt_dir,lang)
                    for r in rows: r["numero"]=nombre
                    meta_rows.extend(rows)
                    self._put(tipo="log",texto=f"  ✅ {len(rows)} imágenes")
            except Exception as e:
                self._put(tipo="log",texto=f"  ❌ {e}"); errores.append(archivo.name)
            gc.collect()
            self._put(tipo="prog",val=int((idx+1)/total*100),txt=f"{idx+1}/{total}")

        # ── Postprocesamiento: aplicar zonas etiquetadas si está activo ──────────
        if usar_etiq and meta_rows:
            api_key_det = ""
            if det_auto:
                api_key_det, _ = _resolver_api_key_modelo("deteccion")
            n_filtradas = 0
            self._put(tipo="fase", txt="Aplicando zonas etiquetadas…")
            for row in meta_rows:
                try:
                    tp = Path(row["txt_path"])
                    if tp.exists():
                        texto_orig = tp.read_text("utf-8", errors="replace")
                        texto_filt = self._ocr_aplicar_zonas(
                            texto_orig, row["numero"], row["pagina"],
                            det_auto=det_auto, api_key=api_key_det)
                        if texto_filt != texto_orig:
                            tp.write_text(texto_filt, "utf-8")
                            row["palabras"] = len(texto_filt.split())
                            n_filtradas += 1
                except Exception:
                    pass
            if n_filtradas:
                self._put(tipo="log",
                          texto=f"✂️ {n_filtradas} páginas filtradas por zonas etiquetadas")

        # ── Avisos de calidad por página (estilo FineReader) ─────────────────────
        # Se calculan desde los datos YA disponibles en meta_rows (texto, conteo
        # de palabras, confianza) sin reabrir ninguna imagen — el DPI real solo se
        # chequea bajo demanda cuando el investigador abre una página puntual en
        # Normalizar, para no sumar una relectura de disco/Drive por cada página
        # del lote (que puede ser de cientos de páginas).
        if meta_rows:
            from core.page_quality import es_pagina_vacia
            self._avisos_ocr = {}
            for row in meta_rows:
                claves = []
                texto = ""
                tp = Path(row["txt_path"]) if row.get("txt_path") else None
                if tp and tp.exists():
                    try:
                        texto = tp.read_text("utf-8", errors="replace")
                    except Exception:
                        texto = ""
                if es_pagina_vacia(texto, n_tokens=row.get("palabras")):
                    claves.append("Página posiblemente en blanco o sin texto útil")
                conf = row.get("confianza")
                if conf is not None and conf < 30:
                    claves.append(f"Confianza OCR muy baja: {conf:.0f}%")
                if claves:
                    self._avisos_ocr[(row["numero"], row["pagina"])] = claves
            if self._avisos_ocr:
                self._put(tipo="log",
                    texto=f"⚠ {len(self._avisos_ocr)} página(s) con avisos de calidad "
                          "(ver detalle en Normalizar)")

        if not meta_rows:
            self._put(tipo="err",txt="No se pudo procesar ningún archivo.")
            self.after(0, lambda: self._btn_ocr.config(state="normal")); return

        COLS=["numero","pagina","txt_path","palabras","confianza","revision"]
        df=pd.DataFrame(meta_rows)
        for c in COLS:
            if c not in df.columns: df[c]=None
        df["palabras"]=pd.to_numeric(df["palabras"],errors="coerce").fillna(0).astype(int)
        df["confianza"]=pd.to_numeric(df["confianza"],errors="coerce")
        df["revision"]=df["confianza"].apply(lambda c: bool(pd.notna(c) and c<60))
        ad=out/"04_analisis"; ad.mkdir(exist_ok=True)
        df.to_csv(ad/"ocr_metadatos.csv", index=False)
        ST.corpus_meta=df; ST.ocr_done=True
        ST.marcar_etapa("ocr", "ready")

        n_rev=int(df["revision"].sum()); cm=df["confianza"].dropna().mean(); tp=int(df["palabras"].sum())
        n_dir=int(df.get("metodo",pd.Series()).eq("texto_embebido").sum()) if "metodo" in df.columns else 0
        self.after(0,lambda: self._lbl_o_pdf.config(text=str(total-len(errores))))
        self.after(0,lambda: self._lbl_o_pag.config(text=f"{len(df):,}"))
        self.after(0,lambda: self._lbl_o_pal.config(text=f"{tp:,}"))
        self.after(0,lambda: self._lbl_o_con.config(text=f"{cm:.0f}%" if pd.notna(cm) else "—"))
        self.after(0,lambda: self._lbl_o_rev.config(text=str(n_rev)))
        if errores: self._put(tipo="log",texto=f"⚠️ Errores: {', '.join(errores)}")
        self._put(tipo="prog",val=100,txt="✅"); self._put(tipo="fase",txt="✅ Extracción completada")
        self._put(tipo="log",texto=f"🎉 {len(df):,} páginas · {tp:,} palabras")
        self._put(tipo="ok",res="ocr")

    def _worker_ocr_carpetas(self, numeros: list[dict]):
        """
        Worker OCR para el modo 'subcarpetas': cada elemento de numeros es
        {'nombre': str, 'pdfs': list[Path], 'carpeta': Path}.
        Cada PDF de la subcarpeta es una página del número.
        Usa alto_reconstructor para reconstruir el orden de columnas.
        """
        import fitz
        import pandas as pd

        from core.alto_reconstructor import reconstruir_texto_pagina
        from core.ocr_normalizer import normalizar_texto_ocr

        out  = ST.out_dir
        ruta_ocr = getattr(self, "_var_ruta_ocr",
                            type("V", (), {"get": lambda s: "bnc"})()).get()
        lang = self._var_dpi and self._var_lang.get() or "spa"

        _MARCA_BNC = "Digitalizado Biblioteca Nacional de Colombia"
        meta_rows  = []
        errores    = []
        total_nums = len(numeros)

        for n_idx, num_info in enumerate(numeros):
            nombre = num_info["nombre"]
            pdfs   = num_info["pdfs"]
            self._put(tipo="log", texto=f"📁 {nombre} ({n_idx+1}/{total_nums}) — {len(pdfs)} páginas")

            txt_dir = out / "03_ocr" / nombre
            img_dir = out / "02_imagenes" / nombre
            txt_dir.mkdir(parents=True, exist_ok=True)
            img_dir.mkdir(parents=True, exist_ok=True)

            for p_idx, pdf_path in enumerate(pdfs):
                pagina_id = f"p{p_idx+1:04d}"
                tp = txt_dir / f"{pagina_id}.txt"
                pct = int(((n_idx * len(pdfs) + p_idx + 1) /
                           max(total_nums * len(pdfs), 1)) * 100)
                self._put(tipo="prog", val=pct,
                          txt=f"{nombre}·{pagina_id}")

                palabras = 0
                confianza = None
                metodo = "bnc_coordenadas"

                try:
                    if ruta_ocr in ("bnc", "tesseract") or not tp.exists():
                        doc = fitz.open(str(pdf_path))
                        page = doc[0]

                        # Guardar imagen original a color
                        img_dest = img_dir / f"{pagina_id}.png"
                        if not img_dest.exists():
                            import io

                            from PIL import Image as _PIL
                            zoom = 150 / 72.0
                            mat  = fitz.Matrix(zoom, zoom)
                            pix  = page.get_pixmap(matrix=mat, alpha=False)
                            img  = _PIL.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                            img.save(str(img_dest), "PNG")

                        if not tp.exists():
                            if ruta_ocr == "tesseract":
                                # OCR con Tesseract sobre la imagen extraída
                                from core.ocr_engine import ocr_pagina
                                texto, confianza = ocr_pagina(img_dest, lang=lang)
                                metodo = "tesseract"
                            else:
                                # Ruta 3 BNC: reconstrucción por coordenadas
                                resultado = reconstruir_texto_pagina(page,
                                                ignorar_ocr_basura=True)
                                texto = resultado.get("texto", "")
                                texto = texto.replace(_MARCA_BNC, "").strip()
                                texto = normalizar_texto_ocr(texto)

                            tp.write_text(texto, encoding="utf-8")
                        else:
                            texto = tp.read_text("utf-8", errors="replace")

                        doc.close()
                        palabras = len(texto.split())

                except Exception as e:
                    self._put(tipo="log",
                        texto=f"  ⚠ {pagina_id}: {str(e)[:80]}")
                    errores.append(f"{nombre}/{pagina_id}")
                    continue

                meta_rows.append({
                    "numero":    nombre,
                    "pagina":    pagina_id,
                    "txt_path":  str(tp),
                    "palabras":  palabras,
                    "confianza": confianza,
                    "revision":  False,
                    "metodo":    metodo,
                })

            n_pag = len(pdfs)
            pal   = sum(r["palabras"] for r in meta_rows if r["numero"] == nombre)
            self._put(tipo="log",
                texto=f"  ✅ {n_pag} pág · {pal:,} palabras · {nombre}")

        # Consolidar metadatos
        COLS = ["numero","pagina","txt_path","palabras","confianza","revision","metodo"]
        df = pd.DataFrame(meta_rows) if meta_rows else pd.DataFrame(columns=COLS)
        for c in COLS:
            if c not in df.columns: df[c] = None
        df["palabras"]  = pd.to_numeric(df["palabras"],  errors="coerce").fillna(0).astype(int)
        df["confianza"] = pd.to_numeric(df["confianza"], errors="coerce")
        df["revision"]  = df["confianza"].apply(
            lambda c: bool(pd.notna(c) and c < 60))

        ad = out / "04_analisis"
        ad.mkdir(exist_ok=True)
        df.to_csv(ad / "ocr_metadatos.csv", index=False)
        ST.corpus_meta = df
        ST.ocr_done    = True
        ST.marcar_etapa("ocr", "ready")

        tp_total = int(df["palabras"].sum())
        self._put(tipo="prog", val=100, txt="✅")
        self._put(tipo="fase", txt="✅ Extracción completada")
        self._put(tipo="log",
            texto=f"🎉 {total_nums} números · {len(df):,} páginas · {tp_total:,} palabras")
        if errores:
            self._put(tipo="log", texto=f"⚠ Errores: {len(errores)} páginas")
        self._put(tipo="ok", res="ocr")

    # ══════════════════════════════════════════════════════════════════════════
    # WORKERS — SEGMENTACIÓN
    # ══════════════════════════════════════════════════════════════════════════
    def _reconstruir_corpus_meta_desde_txt(self) -> bool:
        """Construye ST.corpus_meta leyendo los TXT de 03_ocr/ cuando no hay OCR previo.
        Retorna True si encontró archivos, False si no hay nada."""
        import pandas as pd
        if not ST.out_dir:
            return False
        txt_base = Path(ST.out_dir) / "03_ocr"
        if not txt_base.exists():
            return False
        meta_rows = []
        for num_dir in sorted(txt_base.iterdir()):
            if not num_dir.is_dir():
                continue
            for txt_path in sorted(num_dir.glob("*.txt")):
                try:
                    texto = txt_path.read_text(encoding="utf-8", errors="replace")
                    palabras = len(texto.split())
                except Exception:
                    palabras = 0
                meta_rows.append({
                    "numero":    num_dir.name,
                    "pagina":    txt_path.stem,
                    "txt_path":  str(txt_path),
                    "palabras":  palabras,
                    "confianza": None,
                    "revision":  False,
                    "metodo":    "conversor",
                })
        if not meta_rows:
            return False
        df = pd.DataFrame(meta_rows)
        df["palabras"] = pd.to_numeric(df["palabras"], errors="coerce").fillna(0).astype(int)
        ad = Path(ST.out_dir) / "04_analisis"
        ad.mkdir(exist_ok=True)
        df.to_csv(ad / "ocr_metadatos.csv", index=False)
        ST.corpus_meta = df
        ST.ocr_done    = True
        ST.marcar_etapa("ocr", "ready")
        return True

    def _start_seg(self):
        if not ST.ocr_done:
            # Intentar reconstruir desde TXT del conversor antes de bloquear
            if not self._reconstruir_corpus_meta_desde_txt():
                messagebox.showwarning("Sin texto extraído",
                    "No se encontró texto en 03_ocr/.\n\n"
                    "Usá primero el Conversor PDF o la Extracción OCR.")
                return
        self._btn_seg.config(state="disabled")
        self._seg_skeleton = self._skeleton_show(self._tv_seg_outer, n_filas=8)
        threading.Thread(target=self._worker_seg, daemon=True).start()

    def _worker_seg(self):
        from core.article_segmenter import segmentar_numero
        from core.zone_labeler import (
            filtrar_texto_con_etiquetas,
            listar_paginas_etiquetadas,
        )
        usar_v2 = getattr(self, "_var_seg_v2", None)
        usar_v2 = usar_v2.get() if usar_v2 else False

        out     = ST.out_dir; txt_dir = out / "03_ocr"

        # Obtener lista de números desde corpus_meta si existe,
        # o desde las carpetas en 03_ocr/ si el texto vino del conversor.
        if ST.corpus_meta is not None and "numero" in ST.corpus_meta.columns:
            numeros = sorted(ST.corpus_meta["numero"].unique().tolist())
        elif txt_dir.exists():
            numeros = sorted(p.name for p in txt_dir.iterdir()
                             if p.is_dir() and list(p.glob("*.txt")))
        elif ST.archivos_sel:
            numeros = [a.stem for a in ST.archivos_sel]
        else:
            self._put(tipo="err",
                      txt="No hay texto extraído. Ejecutá primero Extracción OCR o el Conversor.")
            return

        if not numeros:
            self._put(tipo="err",
                      txt="No se encontraron números con texto en 03_ocr/.")
            return

        total   = len(numeros); todos = []

        for k, nombre in enumerate(numeros):
            self._put(tipo="fase", txt=f"Segmentando {nombre}…")
            etiq_pags = listar_paginas_etiquetadas(out, nombre)
            if etiq_pags:
                self._put(tipo="log", texto=f"  📌 {len(etiq_pags)} páginas con etiquetas de zona")
                carpeta_txt = txt_dir / nombre
                if carpeta_txt.exists():
                    for tf in sorted(carpeta_txt.glob("*.txt")):
                        texto_orig = tf.read_text("utf-8", errors="replace")
                        texto_filt = filtrar_texto_con_etiquetas(out, nombre, tf.stem, texto_orig)
                        if texto_filt != texto_orig:
                            tf.write_text(texto_filt, encoding="utf-8")

            pdf_path = None
            for a in ST.archivos_sel:
                if a.stem == nombre: pdf_path = a; break

            if usar_v2:
                # Segmentador avanzado v2 — grafo de continuidad
                from core.article_segmenter_v2 import (
                    comparar_segmentaciones,
                    segmentar_avanzado,
                )
                carpeta_txt = txt_dir / nombre
                paginas_txt = []
                if carpeta_txt.exists():
                    paginas_txt = [
                        tf.read_text("utf-8", errors="replace")
                        for tf in sorted(carpeta_txt.glob("*.txt"))
                    ]
                arts_v2 = segmentar_avanzado(paginas_txt, numero=nombre)
                # Convertir a formato dict compatible con el resto del pipeline
                arts = []
                for a in arts_v2:
                    arts.append({
                        "numero":    nombre,
                        "titulo":    a.titulo,
                        "autor":     a.autor,
                        "tipo":      a.seccion,
                        "paginas":   str(a.paginas),
                        "palabras":  a.palabras,
                        "confianza": a.confianza,
                        "metodo_seg": a.metodo,
                        "texto":     a.texto,
                    })
                # Log comparativo con v1
                arts_v1 = segmentar_numero(txt_dir, nombre, pdf_path)
                comp = comparar_segmentaciones(arts_v1, arts_v2)
                self._put(tipo="log", texto=
                    f"  {nombre}: {comp['v2_articulos']} arts (v2, conf. media "
                    f"{comp['v2_confianza_media']:.2f}) vs {comp['v1_articulos']} arts (v1)")
            else:
                arts = segmentar_numero(txt_dir, nombre, pdf_path)
                self._put(tipo="log", texto=f"  {nombre}: {len(arts)} artículos detectados")

            todos.extend(arts)

        ST.df_articulos = pd.DataFrame(todos) if todos else pd.DataFrame()
        # Poblar ST.articulos para módulos que lo esperan como lista de dicts
        ST.articulos = todos
        if not ST.df_articulos.empty:
            ad = out/"04_analisis"; ad.mkdir(exist_ok=True)
            ST.df_articulos.drop(columns=["texto"], errors="ignore").to_csv(
                ad/"articulos_segmentados.csv", index=False, encoding="utf-8-sig")
        # Construir lista plana de textos para módulo de Lingüística y búsqueda semántica
        if not ST.df_articulos.empty and "texto" in ST.df_articulos.columns:
            ST.corpus_txt = ST.df_articulos["texto"].dropna().tolist()
        elif todos:
            ST.corpus_txt = [
                str(a.get("texto", "") or a.get("contenido", ""))
                for a in todos if a.get("texto") or a.get("contenido")
            ]

        ST.seg_done = True
        ST.marcar_etapa("seg", "ready")
        self.after(0, self._marcar_modificado)
        self._put(tipo="ok", res="seg")

    # ══════════════════════════════════════════════════════════════════════════
    # WORKERS — ANÁLISIS TEXTUAL
    # ══════════════════════════════════════════════════════════════════════════
    def _start_anal(self):
        if not ST.ocr_done and not self._reconstruir_corpus_meta_desde_txt():
            messagebox.showwarning("Sin texto extraído",
                "No se encontró texto en 03_ocr/.\n"
                "Usá primero el Conversor PDF o la Extracción OCR."); return
        self._btn_anal.config(state="disabled")
        threading.Thread(target=self._worker_anal, daemon=True).start()

    def _worker_anal(self):
        import spacy

        from core.analysis_engine import (
            analizar_layout_pagina,
            analizar_numero_con_campos_expandidos,
            construir_red,
            leer_numero,
            run_lda,
        )
        from core.word_vectors import (
            entrenar_word2vec,
            expandir_campo_semantico,
        )
        def log(m): self.after(0, lambda msg=m: self._log_a_write(msg))
        def prg(v,t=""): self.after(0, lambda: self._set_prog_a(v,t))

        colabs  = [c.strip() for c in self._txt_col.get("1.0","end").strip().splitlines() if c.strip()]
        modelo  = self._var_spacy.get(); n_t=self._var_nt.get(); min_f=self._var_mf.get()
        out     = ST.out_dir; txt_dir=out/"03_ocr"; img_dir=out/"02_imagenes"
        ad      = out/"04_analisis"; ad.mkdir(exist_ok=True)
        campos_semillas = getattr(ST,"campos_semillas",CAMPOS_DEFAULT)

        log(f"🧠 Cargando spaCy {modelo}…"); prg(2,"Cargando modelo NLP…")
        try:
            nlp = spacy.load(modelo)
        except OSError:
            self.after(0, lambda: messagebox.showerror("Modelo no encontrado",
                f"Ejecuta: python -m spacy download {modelo}"))
            self.after(0, lambda: self._btn_anal.config(state="normal")); return
        stopwords = spacy.lang.es.stop_words.STOP_WORDS | {
            "año","años","así","vez","día","días","hacer","gran","mismo",
            "todo","todos","toda","todas","estampa","revista","número","hace","sido"
        }

        if ST.corpus_meta is None:
            self._reconstruir_corpus_meta_desde_txt()
        if ST.corpus_meta is None:
            self._put(tipo="err", txt="No hay texto extraído. Usá el Conversor o la Extracción OCR primero.")
            return
        numeros = sorted(ST.corpus_meta["numero"].unique()); total=len(numeros)
        fr_rows,se_rows,ca_rows,lay_rows,lema_docs,lema_nms = [],[],[],[],[],[]

        # --- Paso 1: Word2Vec (si activado)
        word_model = None
        if self._var_wv.get():
            log("🔢 Recopilando corpus para Word2Vec…"); prg(5,"Corpus Word2Vec…")
            corpus_txt = []
            for nombre in numeros:
                from core.analysis_engine import leer_numero as _lr
                corpus_txt.append(_lr(txt_dir, nombre))
            log("📐 Entrenando Word2Vec…"); prg(8,"Entrenando vectores…")
            model_path = ad/"word2vec.model"
            word_model = entrenar_word2vec(corpus_txt, model_path)
            if word_model:
                log(f"  ✅ Modelo entrenado ({len(word_model.wv)} términos)")
            else:
                log("  ⚠️ Corpus insuficiente para Word2Vec — omitido")
            gc.collect()
        ST.word_model = word_model

        # --- Expandir campos semánticos con Word2Vec
        campos_expandidos = {}
        if word_model:
            log("🔤 Expandiendo campos semánticos con Word2Vec…")
            for campo, semillas in campos_semillas.items():
                res = expandir_campo_semantico(semillas, word_model, topn=20, umbral_sim=0.3)
                campos_expandidos[campo] = res["campo_expandido"]
                log(f"  {campo}: {len(res['campo_expandido'])} términos ({len(semillas)} semillas + {len(res['expansiones'])} expansiones)")
        else:
            campos_expandidos = campos_semillas
        ST.campos_expandidos = campos_expandidos

        # Actualizar combobox de campos
        self.after(0, lambda: self._cb_campos.config(values=list(campos_expandidos.keys())))

        # --- Paso 2: Análisis texto número a número
        for k, nombre in enumerate(numeros):
            log(f"── {nombre} ({k+1}/{total})"); prg(int(10+k/total*65),f"Analizando {nombre}…")
            texto = leer_numero(txt_dir, nombre)
            if not texto.strip(): log("  ⚠️ Sin texto"); continue
            firmas,secciones,campos,lema = analizar_numero_con_campos_expandidos(
                nombre, texto, colabs, nlp, stopwords, campos_expandidos)
            for fi in firmas: fr_rows.append({"numero":nombre,"firma":fi})
            for s,c in secciones.items(): se_rows.append({"numero":nombre,"seccion":s,"menciones":c})
            campos["numero"]=nombre; ca_rows.append(campos)
            if lema.strip(): lema_docs.append(lema); lema_nms.append(nombre)
            del texto; gc.collect()
            if self._var_layout.get():
                imgd = img_dir/nombre
                if imgd.exists():
                    for ip in sorted(imgd.glob("*.png")):
                        r=analizar_layout_pagina(ip)
                        if r: r["numero"]=nombre; r["pagina"]=ip.stem; lay_rows.append(r)

        log("🧠 Liberando modelo NLP…"); del nlp; gc.collect()

        # --- LDA
        log(f"🧩 LDA ({n_t} temas)…"); prg(78,"Modelado LDA…")
        df_temas=pd.DataFrame(); df_doc_temas=pd.DataFrame()
        if lema_docs:
            try:
                df_temas,df_doc_temas=run_lda(lema_docs,lema_nms,n_t)
                log(f"  ✅ {n_t} temas")
            except Exception as e: log(f"  ⚠️ LDA: {e}")

        # --- Red de autoría
        graph_path=None
        if self._var_red.get() and fr_rows:
            log("🕸️ Red de autoría…"); prg(85,"Red…")
            import networkx as nx
            df_f_tmp=pd.DataFrame(fr_rows)
            G=construir_red(df_f_tmp,min_f); graph_path=ad/"red_autoria.graphml"
            nx.write_graphml(G,str(graph_path))
            log(f"  ✅ {G.number_of_nodes()} nodos · {G.number_of_edges()} aristas"); del G; gc.collect()

        # --- Guardar
        log("💾 Guardando resultados…"); prg(92,"Guardando…")
        df_firmas    = pd.DataFrame(fr_rows)
        df_secciones = pd.DataFrame(se_rows)
        df_campos    = pd.DataFrame(ca_rows)
        df_layout    = pd.DataFrame(lay_rows)
        for df_i,fn in [(df_firmas,"firmas.csv"),(df_secciones,"secciones.csv"),
                         (df_campos,"campos_semanticos.csv"),(df_layout,"layout.csv"),
                         (df_temas,"lda_temas.csv")]:
            if not df_i.empty: df_i.to_csv(ad/fn,index=False)
        if not df_doc_temas.empty: df_doc_temas.to_csv(ad/"lda_distribucion.csv")
        ST.df_firmas=df_firmas; ST.df_secciones=df_secciones; ST.df_campos=df_campos
        ST.df_layout=df_layout; ST.df_temas=df_temas; ST.df_doc_temas=df_doc_temas
        ST.graph_path=graph_path; ST.anal_done=True
        ST.marcar_etapa("anal", "ready")
        self.after(0, self._marcar_modificado)
        prg(100,"✅"); log("🎉 Análisis textual completado.")
        self._put(tipo="ok",res="anal")

    # ══════════════════════════════════════════════════════════════════════════
    # WORKERS — VISUAL Y TIPOGRAFÍA
    # ══════════════════════════════════════════════════════════════════════════
    def _start_vis(self):
        if not ST.ocr_done and not self._reconstruir_corpus_meta_desde_txt():
            messagebox.showwarning("Sin texto extraído",
                "No se encontró texto en 03_ocr/.\n"
                "Usá primero el Conversor PDF o la Extracción OCR."); return
        self._btn_vis.config(state="disabled")
        threading.Thread(target=self._worker_vis, daemon=True).start()

    def _worker_vis(self):
        from core.image_analyzer import analizar_numero_imagenes
        from core.visual_analyzer import analizar_tipografia_numero
        out     = ST.out_dir
        img_dir = out / "02_imagenes"
        ocr_dir = out / "03_ocr"
        if ST.corpus_meta is None:
            self._reconstruir_corpus_meta_desde_txt()
        if ST.corpus_meta is None:
            self._put(tipo="err", txt="No hay texto extraído. Usá el Conversor o la Extracción OCR primero.")
            return
        numeros = sorted(ST.corpus_meta["numero"].unique())
        tip_global = {}

        self.after(0, lambda: self._tv_tip.delete(*self._tv_tip.get_children()))
        self.after(0, lambda: self._tv_ele.delete(*self._tv_ele.get_children()))

        api_key, _modelo_vis = _resolver_api_key_modelo("asistente")
        max_ia  = getattr(ST, "max_ia", 15)
        dpi     = self._var_dpi.get()

        for k, nombre in enumerate(numeros):
            self._put(tipo="log",  texto=f"🖼️  {nombre} ({k+1}/{len(numeros)})…")
            self._put(tipo="prog", val=int(k/len(numeros)*100), txt=f"{nombre}…")
            pdf_path = None
            for a in ST.archivos_sel:
                if a.stem == nombre and a.suffix.lower() == ".pdf":
                    pdf_path = a; break

            if pdf_path and pdf_path.exists():
                try:
                    tip = analizar_tipografia_numero(pdf_path)
                    if tip and tip.get("n_fuentes", 0) > 0:
                        tip["numero"] = nombre; tip_global[nombre] = tip
                        self._put(tipo="log",
                            texto=f"  ✅ Tipografía: {tip.get('n_fuentes',0)} fuentes · "
                                  f"{tip.get('fuente_principal','N/D')} · cuerpo {tip.get('tam_cuerpo_medio',0)} pt")
                    else:
                        self._put(tipo="log", texto="  ⚠️ Tipografía: sin datos")
                except Exception as e:
                    self._put(tipo="log", texto=f"  ⚠️ Tipografía: {e}")

            imgd = img_dir / nombre
            if imgd.exists():
                def _cb(msg): self._put(tipo="log", texto=f"  {msg}")
                datos_img = analizar_numero_imagenes(
                    imgd, ocr_dir, nombre,
                    dpi=dpi, api_key=api_key, max_ia=max_ia, callback=_cb)
                if datos_img:
                    ST.datos_imagenes[nombre] = datos_img
                    self._put(tipo="log",
                        texto=f"  ✅ Visual: {datos_img['total_fotos']} fotos · "
                              f"{datos_img['total_ilustraciones']} ilustraciones · "
                              f"{datos_img['total_publicidades']} avisos · "
                              f"{datos_img['area_visual_media']}% área visual promedio")
            else:
                self._put(tipo="log", texto="  ℹ️ Sin imágenes PNG (solo tipografía disponible)")
            gc.collect()

        ST.datos_visual = {"tipografia": tip_global, "visual_elementos": ST.datos_imagenes}
        ST.vis_done = True

        # ── Descripción con IA (opcional) ─────────────────────────────────────
        api_key, _modelo_vis2 = _resolver_api_key_modelo("asistente")
        if api_key and ST.datos_imagenes:
            self._put(tipo="log", texto="🤖 Describiendo imágenes con Claude AI… (esto puede tardar)")
            try:
                from core.image_describer import describir_pagina
                img_dir_base = ST.out_dir / "02_imagenes"
                n_descritos  = 0
                for nombre, datos_num in ST.datos_imagenes.items():
                    imgd = img_dir_base / nombre
                    if not imgd.exists(): continue
                    nuevas_pags = []
                    for pag_datos in datos_num.get("paginas", []):
                        pagina_id = pag_datos.get("pagina","")
                        ip = imgd / f"{pagina_id}.png"
                        if not ip.exists():
                            nuevas_pags.append(pag_datos); continue
                        pag_desc = describir_pagina(ip, pag_datos, api_key, max_elementos=6)
                        nuevas_pags.append(pag_desc)
                        n_desc_pag = sum(1 for el in pag_desc.get("elementos",[]) if el.get("descrito_por_ai"))
                        if n_desc_pag: n_descritos += n_desc_pag
                        self._put(tipo="log", texto=f"  🤖 {nombre}/{pagina_id}: {n_desc_pag} elemento(s) descritos")
                    datos_num["paginas"] = nuevas_pags
                self._put(tipo="log", texto=f"  ✅ IA: {n_descritos} elemento(s) descritos en total")
            except Exception as e:
                self._put(tipo="log", texto=f"  ⚠️ Descripción IA: {e}")

        for nombre, tip in tip_global.items():
            vals = (
                nombre,
                tip.get("fuente_principal","N/D"),
                tip.get("clasificacion_fuente",""),
                tip.get("n_fuentes","—"),
                tip.get("tam_cuerpo_medio","—"),
                tip.get("tam_titulo_medio","—"),
                tip.get("ratio_titulo_cuerpo","—"),
                tip.get("interlineado_rel","—"),
                tip.get("columnas_moda", tip.get("columnas_prom","—")),
                tip.get("negrita_pct","—"),
                tip.get("cursiva_pct","—"),
                tip.get("imagenes_total","—"),
            )
            self.after(0, lambda v=vals: self._tv_tip.insert("","end",values=v))

        tag_map = {
            "Fotografía":            "foto",
            "Ilustración/Caricatura":"ilustracion",
            "Publicidad/Aviso":      "publicidad",
            "Mixto":                 "mixto",
        }
        total_el = 0
        for nombre, datos in ST.datos_imagenes.items():
            for pag in datos.get("paginas", []):
                pagina = pag.get("pagina","")
                for el in pag.get("elementos", []):
                    if el.get("tipo") == "Texto": continue
                    vals = (
                        nombre, pagina,
                        el.get("tipo",""),
                        f"{el.get('confianza',0):.2f}",
                        f"{el.get('w_cm',0):.1f}",
                        f"{el.get('h_cm',0):.1f}",
                        f"{el.get('area_cm2',0):.1f}",
                        f"{el.get('pos_x_pct',0):.0f}",
                        f"{el.get('pos_y_pct',0):.0f}",
                        el.get("autor_imagen",""),
                        el.get("pie_de_foto","")[:80],
                        el.get("descripcion_ia","")[:120],
                    )
                    tag = tag_map.get(el.get("tipo",""), "")
                    self.after(0, lambda v=vals, t=tag: self._tv_ele.insert("","end",values=v,tags=(t,)))
                    total_el += 1

        numeros_con_imgs = [n for n in numeros if n in ST.datos_imagenes]
        if numeros_con_imgs:
            self.after(0, lambda ns=numeros_con_imgs: self._actualizar_combos_diag(ns))

        cnt_msg = f"{total_el} elementos visuales detectados"
        self.after(0, lambda m=cnt_msg: self._lbl_ele_cnt.config(text=m))
        resumen = f"✅ {len(tip_global)} número(s) con tipografía · {total_el} elementos visuales"
        self._put(tipo="log", texto=f"🎉 {resumen}")
        self.after(0, lambda: self._lbl_vis_ok.config(text=resumen))
        self.after(0, lambda: self._btn_vis.config(state="normal"))
        self._put(tipo="prog", val=100, txt="✅")

    def _actualizar_combos_diag(self, numeros: list):
        self._cmb_diag["values"] = numeros
        if numeros:
            self._cmb_diag.set(numeros[0])
            self._on_diag_num_sel(None)

    # ══════════════════════════════════════════════════════════════════════════
    # WORKERS — ANÁLISIS COMPARATIVO
    # ══════════════════════════════════════════════════════════════════════════
    def _start_comp(self):
        ref = self._var_ref.get().strip()
        if not ref or not Path(ref).exists():
            messagebox.showwarning("Sin referencia",
                "Configura la carpeta de publicaciones de referencia en ⚙️ Config."); return
        if not ST.ocr_done and not self._reconstruir_corpus_meta_desde_txt():
            messagebox.showwarning("Sin texto extraído",
                "No se encontró texto en 03_ocr/.\n"
                "Usá primero el Conversor PDF o la Extracción OCR."); return
        self._btn_comp.config(state="disabled")
        threading.Thread(target=self._worker_comp, args=(Path(ref),), daemon=True).start()

    def _worker_comp(self, ref_dir):
        from core.comparative_analyzer import (
            cargar_corpora,
            generar_reporte_comparativo,
        )
        self._put(tipo="log",texto="📚 Cargando corpora de referencia…")
        # Corpus principal: concatenar todos los textos OCR
        out=ST.out_dir; txt_dir=out/"03_ocr"
        if ST.corpus_meta is None:
            self._reconstruir_corpus_meta_desde_txt()
        if ST.corpus_meta is None:
            self._put(tipo="err", txt="No hay texto extraído. Usá el Conversor o la Extracción OCR primero.")
            return
        texto_principal = ""
        for nombre in sorted(ST.corpus_meta["numero"].unique()):
            carpeta = txt_dir/nombre
            if carpeta.exists():
                for tf in sorted(carpeta.glob("*.txt")):
                    texto_principal += tf.read_text("utf-8",errors="replace") + "\n"
        pub_nombre = ST.publicacion
        corpus_principal = {pub_nombre: texto_principal}
        # Cargar referencia
        corpora = cargar_corpora(ref_dir, corpus_principal)
        self._put(tipo="log",texto=f"  {len(corpora)} publicaciones: {', '.join(corpora.keys())}")
        if len(corpora) < 2:
            self._put(tipo="log",texto="⚠️ Se necesitan al menos 2 publicaciones para comparar.")
            self.after(0, lambda: self._btn_comp.config(state="normal")); return
        # Campos con expansión si disponible
        campos = ST.campos_expandidos if ST.campos_expandidos else getattr(ST,"campos_semillas",CAMPOS_DEFAULT)
        campos_simple = {k: v for k,v in campos.items()}
        self._put(tipo="log",texto="📊 Calculando similaridad y palabras distintivas…")
        rep = generar_reporte_comparativo(pub_nombre, corpora, campos_simple)
        ST.datos_comparativo=rep; ST.comp_done=True
        # Mostrar en UI
        sim=rep.get("similaridad"); dist=rep.get("palabras_distintivas",{}); pc=rep.get("perfil_campos")
        if sim is not None:
            sim_txt=sim.to_string()
            def _show_sim(t=sim_txt):
                self._txt_sim.delete("1.0","end"); self._txt_sim.insert("1.0",t)
            self.after(0,_show_sim)
        if dist:
            lines=[]
            for ref_n,pals in dist.items():
                lines.append(f"\n── vs. {ref_n} ──")
                for p,s in pals[:15]: lines.append(f"  {p:<25} G²={s:.1f}")
            def _show_dist(t="\n".join(lines)):
                self._txt_dist.delete("1.0","end"); self._txt_dist.insert("1.0",t)
            self.after(0,_show_dist)
        if pc is not None and not pc.empty:
            def _show_campos_fig():
                for w in self._fig_cam_frame.winfo_children(): w.destroy()
                fig,ax=plt.subplots(figsize=(10,4))
                x=np.arange(len(pc.columns)); w_bar=0.8/max(len(pc),1)
                cols_pal=PALETTE
                for i,(pub,row) in enumerate(pc.iterrows()):
                    ax.bar(x+i*w_bar,row.values,width=w_bar,label=pub,color=cols_pal[i%len(cols_pal)],alpha=0.85)
                ax.set_xticks(x+w_bar*(len(pc)-1)/2)
                ax.set_xticklabels(pc.columns,rotation=30,ha="right",fontsize=9)
                ax.set_ylabel("Menciones/1000 pal."); ax.legend(fontsize=8)
                ax.set_title("Campos semánticos — comparativo",fontsize=11,fontweight="bold")
                plt.tight_layout()
                canvas=FigureCanvasTkAgg(fig,master=self._fig_cam_frame)
                canvas.draw(); canvas.get_tk_widget().pack(fill="both",expand=True)
                plt.close(fig)
            self.after(0,_show_campos_fig)
        self._put(tipo="log",texto="🎉 Análisis comparativo completado.")
        self.after(0,lambda: self._lbl_comp_ok.config(text="✅ Comparativo completado"))
        self.after(0,lambda: self._btn_comp.config(state="normal"))

    # ══════════════════════════════════════════════════════════════════════════
    # RESULTADOS Y EXPORTACIÓN
    # ══════════════════════════════════════════════════════════════════════════
    def _cargar_resultados(self):
        if not ST.anal_done:
            # Mostrar estado pendiente en los indicadores
            for attr in ("_lbl_r_num","_lbl_r_pag","_lbl_r_pal",
                         "_lbl_r_art","_lbl_r_aut","_lbl_r_fir"):
                if hasattr(self, attr):
                    getattr(self, attr).config(text="—")
            return
        from core.excel_export import generar_figuras_completas

        meta=ST.corpus_meta
        self._lbl_r_num.config(text=str(meta["numero"].nunique()))
        self._lbl_r_pag.config(text=f"{len(meta):,}")
        self._lbl_r_pal.config(text=f"{int(meta['palabras'].sum()):,}")
        n_art=len(ST.df_articulos) if ST.df_articulos is not None else 0
        n_aut=(ST.df_articulos[ST.df_articulos["autor"]!="Anónimo / Sin atribuir"]["autor"].nunique()
               if ST.df_articulos is not None and not ST.df_articulos.empty else 0)
        n_fir=ST.df_firmas["firma"].nunique() if ST.df_firmas is not None and not ST.df_firmas.empty else 0
        self._lbl_r_art.config(text=str(n_art))
        self._lbl_r_aut.config(text=str(n_aut))
        self._lbl_r_fir.config(text=str(n_fir))

        def gen():
            datos={
                "df_secciones":ST.df_secciones,"df_firmas":ST.df_firmas,"df_campos":ST.df_campos,
                "df_temas":ST.df_temas,"df_doc_temas":ST.df_doc_temas,"df_layout":ST.df_layout,
                "graph_path":ST.graph_path,"colaboradores":set(),
                "df_articulos":ST.df_articulos,
                "datos_visual":ST.datos_visual,
                "datos_comparativo":ST.datos_comparativo,
            }
            figs = generar_figuras_completas(datos)
            ST.figuras=figs; self.after(0, self._pintar_graficas)
        threading.Thread(target=gen, daemon=True).start()

    def _pintar_graficas(self):
        import io

        from PIL import Image as PILImage
        from PIL import ImageTk
        for key, tab in self._figs_tabs.items():
            if key not in ST.figuras: continue
            for w in tab.winfo_children(): w.destroy()
            buf = io.BytesIO(ST.figuras[key])
            img = PILImage.open(buf); img.thumbnail((940,500),PILImage.LANCZOS)
            tk_img=ImageTk.PhotoImage(img)
            lbl=tk.Label(tab,image=tk_img,bg=CONTENT_BG); lbl.image=tk_img; lbl.pack(expand=True)

    def _gen_excel(self):
        if not ST.anal_done:
            messagebox.showwarning("Sin datos","Completa el análisis textual primero."); return
        if not ST.figuras:
            messagebox.showwarning("Gráficas pendientes","Espera a que se generen las gráficas."); return
        pub_safe = re.sub(r"[^\w]","_",ST.publicacion)
        out_path = ST.out_dir/f"{pub_safe}_Analisis_Editorial.xlsx"
        try:
            from core.excel_export import construir_excel_completo
            datos={"corpus_meta":ST.corpus_meta,"df_firmas":ST.df_firmas,"df_secciones":ST.df_secciones,
                   "df_campos":ST.df_campos,"df_layout":ST.df_layout,"df_temas":ST.df_temas,
                   "df_doc_temas":ST.df_doc_temas,"graph_path":ST.graph_path,
                   "df_articulos":ST.df_articulos,"datos_visual":ST.datos_visual,
                   "datos_comparativo":ST.datos_comparativo}
            construir_excel_completo(datos, ST.figuras, out_path)
            ST.xlsx_path=out_path
            self._lbl_excel.config(text=f"✅ {out_path.name} ({out_path.stat().st_size/1024:.0f} KB)")
            if messagebox.askyesno("Excel generado",f"✅ Guardado:\n{out_path}\n\n¿Abrir carpeta?"):
                self._abrir_carpeta()
        except Exception as e:
            messagebox.showerror("Error",f"No se pudo generar el Excel:\n{e}"); raise

    def _guardar_graphml(self):
        if not ST.graph_path or not ST.graph_path.exists():
            messagebox.showwarning("Sin red","La red no fue generada."); return
        dest=filedialog.asksaveasfilename(defaultextension=".graphml",
             filetypes=[("GraphML","*.graphml")], initialfile="red_autoria.graphml")
        if dest:
            import shutil; shutil.copy2(str(ST.graph_path),dest)
            self.toast(f"Red guardada → {Path(dest).name}", tipo="ok")

    def _abrir_carpeta(self):
        if not ST.out_dir: return
        import subprocess
        p=str(ST.out_dir)
        if platform.system()=="Windows": subprocess.Popen(f'explorer "{p}"')
        elif platform.system()=="Darwin": subprocess.Popen(["open",p])
        else: subprocess.Popen(["xdg-open",p])

    # ══════════════════════════════════════════════════════════════════════════
    # DISPATCHER
    # ══════════════════════════════════════════════════════════════════════════

    # ══════════════════════════════════════════════════════════════════════════
    # OCR LLM — mejora páginas de baja confianza con Claude
    # ══════════════════════════════════════════════════════════════════════════
    def _start_mejorar_ocr(self):
        if not ST.ocr_done or ST.corpus_meta is None:
            messagebox.showwarning("OCR pendiente",
                "Ejecuta la extraccion OCR primero."); return
        if not ST.api_key:
            messagebox.showwarning("Sin API key",
                "Configura la API key en Configuracion."); return
        umbral_val = getattr(self, "_var_ocr_umbral", None)
        umbral_val = umbral_val.get() if umbral_val else 60
        proveedor = getattr(self, "_var_ocr_llm_prov", None)
        proveedor = proveedor.get() if proveedor else "claude"
        ollama_modelo = getattr(self, "_var_ocr_ollama_modelo", None)
        ollama_modelo = ollama_modelo.get() if ollama_modelo else "latamgpt"
        if ST.corpus_meta is None:
            self._reconstruir_corpus_meta_desde_txt()
        if ST.corpus_meta is None:
            messagebox.showwarning("Sin datos", "No hay texto extraído aún."); return
        candidatas = ST.corpus_meta[
            ST.corpus_meta["confianza"].notna() &
            (ST.corpus_meta["confianza"] < umbral_val)
        ]
        n = len(candidatas)
        if n == 0:
            messagebox.showinfo("Sin candidatas",
                f"No hay paginas con confianza < {umbral_val}."); return
        # Estándar de costo IA: estimar volumen→tokens→USD y pedir confirmación
        # antes de gastar. Las páginas con conf < 30 van por visión (más caras,
        # tokens de imagen); el resto por corrección de texto. El modelo de
        # referencia es el de visión/ texto por defecto de cada proveedor.
        try:
            from core.costos import estimar_lote_ocr

            n_vision = int((candidatas["confianza"] < 30).sum())
            modelo_ref = {
                "claude": "claude-sonnet-4-6",
                "openai": "gpt-4o",
            }.get(proveedor, ollama_modelo)
            est = estimar_lote_ocr(n, proveedor, modelo_ref, n_vision=n_vision)
            costo_txt = est.resumen() + "\n\n"
        except Exception:
            # Si la estimación falla por cualquier causa, no bloquear: avisar sin cifra.
            costo_txt = (f"{n} página(s); proveedor {proveedor}. "
                         "(No se pudo estimar el costo con precisión.)\n\n")
        if not messagebox.askyesno("Mejorar OCR con IA",
                f"{costo_txt}"
                "El LLM corregira errores de digitalizacion.\n"
                "Continuamos?"):
            return
        self._btn_ocr.config(state="disabled")
        threading.Thread(target=self._worker_mejorar_ocr,
                         args=(umbral_val, proveedor, ollama_modelo), daemon=True).start()

    def _worker_mejorar_ocr(self, umbral: float, proveedor: str = "claude",
                             ollama_modelo: str = "latamgpt"):
        from core.ocr_llm import mejorar_lote
        img_dir_raiz = ST.out_dir / "02_imagenes" if ST.out_dir else None
        modelo_txt = f" [{ollama_modelo}]" if proveedor in ("ollama", "lmstudio") else ""
        self._put(tipo="log", texto=f"Iniciando mejora OCR con {proveedor}{modelo_txt}...")
        # Para proveedores locales la api_key es la URL del servidor; para otros
        # es la clave real.
        if proveedor == "ollama":
            api_key = ST.api_keys.get("ollama", "http://localhost:11434")
        elif proveedor == "lmstudio":
            api_key = ST.api_keys.get("lmstudio", "http://localhost:1234")
        else:
            api_key = ST.api_key

        def cb(n_actual, n_total, desc):
            self._put(tipo="prog", val=int(n_actual / max(n_total, 1) * 100),
                      txt=f"{n_actual}/{n_total}")
            self._put(tipo="log", texto=f"  [{n_actual}/{n_total}] {desc}")

        try:
            stats = mejorar_lote(
                corpus_meta=ST.corpus_meta,
                api_key=api_key,
                umbral_confianza=umbral,
                img_dir_raiz=img_dir_raiz,
                modo="auto",
                proveedor=proveedor,
                modelo_ollama=ollama_modelo,
                callback=cb,
            )
            msg = (
                f"Mejora completada\n"
                f"  Mejoradas:  {stats['mejoradas']}\n"
                f"  Omitidas:   {stats['omitidas']}\n"
                f"  Errores:    {stats['errores']}\n"
                f"  Candidatas: {stats['total_candidatas']}"
            )
            if "costo_real_usd" in stats:
                msg += (
                    f"\n\n💲 Costo real: ${stats['costo_real_usd']:.4f} USD"
                    f"  ({stats.get('tokens_reales', 0):,} tokens)"
                )
            self._put(tipo="log", texto=msg)
            self.after(0, lambda m=msg: messagebox.showinfo("Mejora completada", m))
        except Exception as e:
            err = str(e)
            self._put(tipo="log", texto=f"Error: {err}")
            self.after(0, lambda e=err: messagebox.showerror("Error", e))
        finally:
            self.after(0, lambda: self._btn_ocr.config(state="normal"))

    def _on_ok(self, res):
        self._actualizar_badges()
        if res == "ocr":
            self._btn_ocr.config(state="normal")
            self.toast("OCR completado — revisa y normaliza el texto", tipo="ok")
            self._etz_refrescar_numeros()
            self.after(200, self._norm_refrescar_numeros)
        elif res == "seg":
            self._skeleton_hide(getattr(self, "_seg_skeleton", None))
            self._seg_skeleton = None
            self._btn_seg.config(state="normal")
            n=len(ST.df_articulos) if ST.df_articulos is not None else 0
            self._lbl_seg_n.config(text=f"✅ {n} artículos detectados")
            self._tv_seg.delete(*self._tv_seg.get_children())
            if ST.df_articulos is not None and not ST.df_articulos.empty:
                for _, row in ST.df_articulos.head(600).iterrows():
                    self._insertar_fila_seg(row)
            self.toast(f"Segmentación completada — {n} artículos detectados", tipo="ok")
        elif res == "anal":
            self._btn_anal.config(state="normal")
            self.toast("Análisis completado — pasando a Resultados", tipo="ok")
            self._mostrar_pagina("res"); self._cargar_resultados()

    # ══════════════════════════════════════════════════════════════════════════
    # PANEL LATERAL DE PARÁMETROS — generado dinámicamente desde PARAMS_SCHEMA
    # ══════════════════════════════════════════════════════════════════════════

    def _build_params_panel(self, parent, schema: dict, storage: dict) -> tk.Frame:
        """
        Construye un panel lateral de parámetros a partir de un PARAMS_SCHEMA.

        Args:
            parent:  Frame contenedor (normalmente el lado derecho de un split).
            schema:  dict PARAMS_SCHEMA del módulo core.
            storage: dict mutable donde se guardan las tkVar de cada parámetro.
                     Las claves son los nombres del schema; los valores son tkVar.

        Returns:
            El Frame construido (ya empaquetado en parent).
        """
        panel = tk.Frame(parent, bg=CARD_BG, relief="solid", bd=1, width=220)
        panel.pack(side="right", fill="y", padx=(8, 0))
        panel.pack_propagate(False)

        tk.Label(panel, text="⚙ Parámetros", bg=CARD_BG, fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(pady=(10, 6), padx=10, anchor="w")

        # Canvas + scrollbar para muchos parámetros
        cv = tk.Canvas(panel, bg=CARD_BG, highlightthickness=0)
        sb = ttk.Scrollbar(panel, orient="vertical", command=cv.yview)
        cv.configure(yscrollcommand=sb.set)
        sb.pack(side="right", fill="y")
        cv.pack(side="left", fill="both", expand=True)
        inner = tk.Frame(cv, bg=CARD_BG)
        win_id = cv.create_window((0, 0), window=inner, anchor="nw")

        def _on_resize(e):
            cv.itemconfig(win_id, width=e.width)
        cv.bind("<Configure>", _on_resize)
        inner.bind("<Configure>", lambda e: cv.configure(scrollregion=cv.bbox("all")))

        # Scroll con rueda del mouse — propagar a todos los hijos
        def _wheel(e):
            cv.yview_scroll(-1 if (e.delta > 0 or e.num == 4) else 1, "units")

        def _bind_wheel(w):
            w.bind("<MouseWheel>", _wheel)
            w.bind("<Button-4>",   _wheel)
            w.bind("<Button-5>",   _wheel)
            for child in w.winfo_children():
                _bind_wheel(child)

        cv.bind("<MouseWheel>", _wheel)
        cv.bind("<Button-4>",   _wheel)
        cv.bind("<Button-5>",   _wheel)
        inner.bind("<MouseWheel>", _wheel)
        # Diferir el bind a hijos porque aún no existen
        panel.after(300, lambda: _bind_wheel(inner))

        for key, spec in schema.items():
            tipo  = spec.get("type", "str")
            label = spec.get("label", key)
            help_ = spec.get("help", "")
            default = spec.get("default")

            grp = tk.Frame(inner, bg=CARD_BG)
            grp.pack(fill="x", padx=10, pady=(0, 10))

            lbl = tk.Label(grp, text=label, bg=CARD_BG, fg=TXT_PRI,
                           font=("Segoe UI", 8, "bold"), anchor="w")
            lbl.pack(fill="x")
            if help_:
                tk.Label(grp, text=help_, bg=CARD_BG, fg=TXT_DIM,
                         font=("Segoe UI", 7), wraplength=180, anchor="w",
                         justify="left").pack(fill="x")

            if tipo == "bool":
                var = tk.BooleanVar(value=bool(default))
                ttk.Checkbutton(grp, variable=var).pack(anchor="w")
                storage[key] = var

            elif tipo == "choice":
                var = tk.StringVar(value=str(default))
                cb = ttk.Combobox(grp, textvariable=var,
                                   values=spec.get("options", []),
                                   state="readonly", width=18)
                cb.pack(fill="x")
                storage[key] = var

            elif tipo == "multicheck":
                var = {}
                defaults_set = set(default or [])
                for opt in spec.get("options", []):
                    bv = tk.BooleanVar(value=(opt in defaults_set))
                    ttk.Checkbutton(grp, text=opt, variable=bv).pack(anchor="w")
                    var[opt] = bv
                storage[key] = var

            elif tipo in ("int", "float"):
                var = tk.DoubleVar(value=float(default or 0))
                mn  = spec.get("min", 0)
                mx  = spec.get("max", 100)
                step = spec.get("step", 1)
                sl = ttk.Scale(grp, from_=mn, to=mx, variable=var, orient="horizontal")
                sl.pack(fill="x")
                lbl_val = tk.Label(grp, bg=CARD_BG, fg=TXT_SEC,
                                   font=("Segoe UI", 8))
                lbl_val.pack(anchor="e")
                def _update_lbl(v, lv=lbl_val, tp=tipo, st=step):
                    val = float(v)
                    lv.config(text=f"{int(val)}" if tp == "int" else f"{val:.2f}")
                sl.config(command=_update_lbl)
                _update_lbl(var.get())
                storage[key] = var

            else:  # str / text
                var = tk.StringVar(value=str(default or ""))
                ttk.Entry(grp, textvariable=var, width=20).pack(fill="x")
                storage[key] = var

        return panel

    def _params_get_values(self, storage: dict) -> dict:
        """Extrae los valores actuales de un storage generado por _build_params_panel."""
        result = {}
        for key, var in storage.items():
            if isinstance(var, dict):
                result[key] = [opt for opt, bv in var.items() if bv.get()]
            elif hasattr(var, "get"):
                result[key] = var.get()
        return result

    # ══════════════════════════════════════════════════════════════════════════
    # TAB NER: ÍNDICE DE ENTIDADES NOMBRADAS
    # ══════════════════════════════════════════════════════════════════════════
    def _build_ner(self):
        f = self._tab_ner
        self._page_header(f, "Índice de Entidades Nombradas",
                          "Personas · Lugares · Organizaciones · Fechas · Obras · Eventos", "🔍")
        self._build_ai_panel(f, "ner")

        # Contenedor con split horizontal (contenido izq + panel params der)
        main_split = tk.Frame(f, bg=CONTENT_BG)
        main_split.pack(fill="both", expand=True, padx=24, pady=16)

        # Panel de parámetros NER (derecha, colapsable)
        self._ner_params: dict = {}
        try:
            from core.ner_engine import PARAMS_SCHEMA as _NER_SCHEMA
            self._build_params_panel(main_split, _NER_SCHEMA, self._ner_params)
        except Exception:
            pass

        pad = tk.Frame(main_split, bg=CONTENT_BG)
        pad.pack(side="left", fill="both", expand=True)

        # ── Barra de acciones ─────────────────────────────────────────────────
        bf = tk.Frame(pad, bg=CONTENT_BG); bf.pack(fill="x", pady=(0, 8))
        self._btn_ner_art = ttk.Button(bf, text="▶  Analizar artículo actual",
                                        style="P.TButton", command=self._ner_articulo_actual)
        self._btn_ner_art.pack(side="left", padx=(0, 8))
        self._btn_ner_corpus = ttk.Button(bf, text="📚  Analizar corpus completo",
                                           style="S.TButton", command=self._ner_corpus_completo)
        self._btn_ner_corpus.pack(side="left", padx=(0, 8))
        self._var_ner_llm = tk.BooleanVar(value=False)
        ttk.Checkbutton(bf, text="🤖 Usar IA:", variable=self._var_ner_llm).pack(side="left", padx=(8, 2))
        self._var_ner_prov = tk.StringVar(value="claude")
        self._cmb_ner_prov = ttk.Combobox(bf, textvariable=self._var_ner_prov,
                                           values=["claude", "ollama", "lmstudio"], state="readonly",
                                           width=8, font=("Segoe UI", 9))
        self._cmb_ner_prov.pack(side="left", padx=(0, 4))
        self._var_ner_ollama_modelo = tk.StringVar(value="latamgpt")
        self._ent_ner_ollama_modelo = ttk.Entry(bf, textvariable=self._var_ner_ollama_modelo,
                                                width=10, font=("Segoe UI", 9))
        def _ner_prov_changed(*_):
            prov = self._var_ner_prov.get()
            if prov == "lmstudio":
                from core.ocr_llm import modelos_cargados_lmstudio
                modelos = modelos_cargados_lmstudio()
                if modelos:
                    self._var_ner_ollama_modelo.set(modelos[0])
                self._ent_ner_ollama_modelo.pack(side="left", padx=(0, 8))
            elif prov == "ollama":
                self._ent_ner_ollama_modelo.pack(side="left", padx=(0, 8))
            else:
                self._ent_ner_ollama_modelo.pack_forget()
        self._var_ner_prov.trace_add("write", _ner_prov_changed)
        ttk.Button(bf, text="📥 Exportar CSV", style="S.TButton",
                   command=self._ner_exportar_csv).pack(side="right")
        ttk.Button(bf, text="🌐 Enlazar Wikidata", style="S.TButton",
                   command=self._ner_enlazar_wikidata).pack(side="right", padx=(0, 6))
        ttk.Button(bf, text="📓 Nota", style="S.TButton",
                   command=lambda: self._bitacora_nueva_nota("ner")).pack(side="right", padx=(0, 6))

        self._lbl_ner_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_ner_ok.pack(anchor="w", pady=(0, 4))

        # ── Split: treeview izquierda + detalle derecha ───────────────────────
        split = tk.Frame(pad, bg=CONTENT_BG); split.pack(fill="both", expand=True)

        izq = tk.Frame(split, bg=CONTENT_BG, width=440)
        izq.pack(side="left", fill="both", expand=True, padx=(0, 8))
        izq.pack_propagate(False)

        # Filtro categoría + búsqueda
        filt_f = tk.Frame(izq, bg=CARD_BG, relief="solid", bd=1)
        filt_f.pack(fill="x", pady=(0, 6))
        fi = tk.Frame(filt_f, bg=CARD_BG, padx=10, pady=6); fi.pack(fill="x")
        tk.Label(fi, text="Categoría:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI", 9, "bold")).pack(side="left")
        _cats = ["Todas", "personas", "lugares", "organizaciones",
                 "fechas", "obras_publicaciones", "eventos_historicos"]
        self._var_ner_cat = tk.StringVar(value="Todas")
        self._cmb_ner_cat = ttk.Combobox(fi, textvariable=self._var_ner_cat, values=_cats,
                                          state="readonly", width=22, font=("Segoe UI", 9))
        self._cmb_ner_cat.pack(side="left", padx=8)
        self._cmb_ner_cat.bind("<<ComboboxSelected>>", lambda e: self._ner_refrescar_tv())
        tk.Label(fi, text="Buscar:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI", 9, "bold")).pack(side="left", padx=(12, 0))
        self._var_ner_buscar = tk.StringVar()
        self._var_ner_buscar.trace_add("write", lambda *_: self._ner_refrescar_tv())
        tk.Entry(fi, textvariable=self._var_ner_buscar, width=18,
                 font=("Segoe UI", 9), relief="solid", bd=1).pack(side="left", padx=6)

        # Treeview
        tv_frame = tk.Frame(izq, bg=CARD_BG, relief="solid", bd=1)
        tv_frame.pack(fill="both", expand=True)
        cols = ("entidad", "tipo", "n_arts", "wikidata")
        self._tv_ner = ttk.Treeview(tv_frame, columns=cols, show="headings", height=22)
        self._tv_ner.heading("entidad",  text="Entidad")
        self._tv_ner.heading("tipo",     text="Tipo")
        self._tv_ner.heading("n_arts",   text="# Arts.")
        self._tv_ner.heading("wikidata", text="Wikidata")
        self._tv_ner.column("entidad",  width=220, anchor="w")
        self._tv_ner.column("tipo",     width=120, anchor="w")
        self._tv_ner.column("n_arts",   width=60,  anchor="center")
        self._tv_ner.column("wikidata", width=110, anchor="w")
        vsb = ttk.Scrollbar(tv_frame, orient="vertical", command=self._tv_ner.yview)
        self._tv_ner.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._tv_ner.pack(fill="both", expand=True)
        self._tv_ner.bind("<<TreeviewSelect>>", self._ner_on_select)
        self._tv_ner.bind("<Double-1>", self._ner_abrir_wikidata)

        # Panel detalle
        der = tk.Frame(split, bg=CONTENT_BG, width=260)
        der.pack(side="left", fill="both")
        der.pack_propagate(False)
        det_card = tk.Frame(der, bg=CARD_BG, relief="solid", bd=1)
        det_card.pack(fill="both", expand=True)
        det_hdr = tk.Frame(det_card, bg="#1C2128"); det_hdr.pack(fill="x")
        tk.Label(det_hdr, text="  📋  Detalle", bg="#1C2128", fg=TXT_PRI,
                 font=("Segoe UI", 8, "bold")).pack(side="left", pady=4)
        self._txt_ner_det = scrolledtext.ScrolledText(det_card, font=("Consolas", 9),
                                                       bg="#1C2128", fg="#CDD6F4",
                                                       relief="flat", state="disabled")
        self._txt_ner_det.pack(fill="both", expand=True, padx=1, pady=(0, 1))

        # Log
        log_f = tk.Frame(pad, bg="#0F1B2D", bd=1, relief="solid")
        log_f.pack(fill="x", pady=(8, 0))
        log_hdr = tk.Frame(log_f, bg="#1A2F4A"); log_hdr.pack(fill="x")
        tk.Label(log_hdr, text="  📋  Registro NER", bg="#1A2F4A", fg="#94A3B8",
                 font=("Segoe UI", 8, "bold")).pack(side="left", pady=4)
        self._log_ner = scrolledtext.ScrolledText(log_f, height=5, font=("Consolas", 9),
                                                   bg="#0F1B2D", fg="#86EFAC",
                                                   relief="flat", state="disabled")
        self._log_ner.pack(fill="x", padx=1, pady=(0, 1))

    def _ner_log(self, msg: str):
        def _do():
            self._log_ner.config(state="normal")
            self._log_ner.insert("end", msg + "\n")
            self._log_ner.see("end")
            self._log_ner.config(state="disabled")
        self.after(0, _do)

    def _ner_articulo_actual(self):
        texto = ""
        art_id = "articulo_actual"
        if ST.df_articulos is not None and not ST.df_articulos.empty:
            sel = self._tv_seg.selection() if hasattr(self, "_tv_seg") else []
            if sel:
                idx = self._tv_seg.index(sel[0])
                row = ST.df_articulos.iloc[idx]
                texto = str(row.get("texto", row.get("contenido", "")))
                art_id = str(row.get("id", row.get("titulo", f"art_{idx}")))
        if not texto and ST.out_dir:
            txt_dir = ST.out_dir / "03_ocr"
            if txt_dir.exists():
                # Buscar el primer txt con contenido real (>100 palabras)
                # para evitar procesar archivos vacíos o de log
                for tf in sorted(txt_dir.rglob("*.txt")):
                    contenido = tf.read_text("utf-8", errors="replace")
                    if len(contenido.split()) > 100:
                        texto = contenido
                        art_id = tf.stem
                        break
        if not texto:
            messagebox.showwarning("Sin texto",
                "No hay texto disponible para analizar.\n\n"
                "Opciones:\n"
                "1. Selecciona un artículo en la pestaña Segmentar\n"
                "2. Completa la extracción OCR primero")
            return
        self._btn_ner_art.config(state="disabled")
        threading.Thread(target=self._worker_ner_articulo, args=(texto, art_id), daemon=True).start()

    def _worker_ner_articulo(self, texto: str, art_id: str):
        import spacy

        from core.ner_engine import actualizar_indice_global, pipeline_ner
        p = self._params_get_values(self._ner_params) if self._ner_params else {}
        usar_ia = p.get("usar_ia", self._var_ner_llm.get())
        proveedor_llm = p.get("proveedor_llm", self._var_ner_prov.get())
        modelo_ollama = p.get("modelo_ollama_ner", self._var_ner_ollama_modelo.get())
        if usar_ia:
            if proveedor_llm == "ollama":
                api_key = ST.api_keys.get("ollama", "http://localhost:11434")
            elif proveedor_llm == "lmstudio":
                api_key = ST.api_keys.get("lmstudio", "http://localhost:1234")
            else:
                api_key, _modelo_ner = _resolver_api_key_modelo("ner")
        else:
            api_key = None
        min_palabras = int(p.get("min_longitud_texto", 100))
        if len(texto.split()) < min_palabras:
            self._ner_log(f"⚠ {art_id}: menos de {min_palabras} palabras, omitido")
            self.after(0, lambda: self._btn_ner_art.config(state="normal"))
            return
        self._ner_log(f"▶ Analizando: {art_id}")
        motor = p.get("motor", "auto")
        try:
            nlp = spacy.load("es_core_news_lg") if motor in ("auto", "spacy") else None
        except OSError:
            nlp = None
            if motor == "spacy":
                self._ner_log("⚠ spaCy no instalado. Ejecuta: python -m spacy download es_core_news_lg")
                self.after(0, lambda: self._btn_ner_art.config(state="normal"))
                return
        umbral = float(p.get("umbral_confianza", 0.7))
        cats   = p.get("categorias") or None
        ner = pipeline_ner(texto, nlp, api_key=api_key, callback=self._ner_log,
                           umbral_confianza=umbral, categorias=cats,
                           proveedor_llm=proveedor_llm, modelo_ollama=modelo_ollama)
        if not getattr(ST, "indice_ner_global", None):
            ST.indice_ner_global = {}
        actualizar_indice_global(ST.indice_ner_global, art_id, ner)
        ST.ner_done = True
        total = sum(len(v) for v in ner.values())
        self._ner_log(f"✅ {art_id}: {total} entidades")
        self.after(0, self._ner_refrescar_tv)
        self.after(0, lambda: self._btn_ner_art.config(state="normal"))
        self.after(0, lambda: self._lbl_ner_ok.config(text=f"✅ {art_id} analizado"))
        self.after(0, self._actualizar_badges)

    def _ner_corpus_completo(self):
        if not ST.ocr_done:
            messagebox.showwarning("OCR pendiente", "Completa la extracción OCR primero."); return
        if not messagebox.askyesno("Analizar corpus",
                "Esto analizará TODOS los textos del corpus.\n"
                "Con IA activada puede generar costo de API.\n\n¿Continuar?"):
            return
        self._btn_ner_corpus.config(state="disabled")
        self._btn_ner_art.config(state="disabled")
        threading.Thread(target=self._worker_ner_corpus, daemon=True).start()

    def _worker_ner_corpus(self):
        import spacy

        from core.ner_engine import (
            actualizar_indice_global,
            indice_global_vacio,
            pipeline_ner,
        )
        p = self._params_get_values(self._ner_params) if self._ner_params else {}
        usar_ia = p.get("usar_ia", self._var_ner_llm.get())
        proveedor_llm = p.get("proveedor_llm", self._var_ner_prov.get())
        modelo_ollama = p.get("modelo_ollama_ner", self._var_ner_ollama_modelo.get())
        if usar_ia:
            if proveedor_llm == "ollama":
                api_key = ST.api_keys.get("ollama", "http://localhost:11434")
            elif proveedor_llm == "lmstudio":
                api_key = ST.api_keys.get("lmstudio", "http://localhost:1234")
            else:
                api_key = _resolver_api_key_modelo("ner")[0]
        else:
            api_key = None
        motor   = p.get("motor", "auto")
        umbral  = float(p.get("umbral_confianza", 0.7))
        cats    = p.get("categorias") or None
        min_palabras = int(p.get("min_longitud_texto", 100))
        try:
            nlp = spacy.load("es_core_news_lg") if motor in ("auto", "spacy") else None
        except OSError:
            nlp = None
            if motor == "spacy":
                self._ner_log("⚠ spaCy no instalado. Ejecuta: python -m spacy download es_core_news_lg")
                self.after(0, lambda: self._btn_ner_corpus.config(state="normal"))
                self.after(0, lambda: self._btn_ner_art.config(state="normal"))
                return

        ST.indice_ner_global = indice_global_vacio()
        articulos = []
        if ST.df_articulos is not None and not ST.df_articulos.empty:
            for i, row in ST.df_articulos.iterrows():
                txt = str(row.get("texto", row.get("contenido", "")))
                aid = str(row.get("id", row.get("titulo", f"art_{i}")))
                if txt.strip() and len(txt.split()) >= min_palabras:
                    articulos.append((aid, txt))
        elif ST.out_dir:
            txt_dir = ST.out_dir / "03_ocr"
            if txt_dir.exists():
                for tf in sorted(txt_dir.rglob("*.txt")):
                    txt = tf.read_text("utf-8", errors="replace")
                    if txt.strip() and len(txt.split()) >= min_palabras:
                        articulos.append((tf.stem, txt))

        total = len(articulos)
        self._ner_log(f"📚 Corpus: {total} textos (umbral ≥{min_palabras} palabras)")
        for i, (aid, txt) in enumerate(articulos, 1):
            self._ner_log(f"[{i}/{total}] {aid}")
            ner = pipeline_ner(txt, nlp, api_key=api_key,
                               umbral_confianza=umbral, categorias=cats,
                               proveedor_llm=proveedor_llm, modelo_ollama=modelo_ollama)
            actualizar_indice_global(ST.indice_ner_global, aid, ner)
            if i % 5 == 0:
                self.after(0, self._ner_refrescar_tv)

        ST.ner_done = True
        n_ents = sum(len(v) for v in ST.indice_ner_global.values())
        self._ner_log(f"✅ Corpus completo: {n_ents} entidades únicas")
        self.after(0, self._ner_refrescar_tv)
        self.after(0, lambda: self._btn_ner_corpus.config(state="normal"))
        self.after(0, lambda: self._btn_ner_art.config(state="normal"))
        self.after(0, lambda: self._lbl_ner_ok.config(text=f"✅ {n_ents} entidades en índice global"))
        self.after(0, self._actualizar_badges)

    def _ner_refrescar_tv(self):
        self._tv_ner.delete(*self._tv_ner.get_children())
        if not getattr(ST, "indice_ner_global", None):
            return
        cat_filtro = self._var_ner_cat.get()
        buscar = self._var_ner_buscar.get().lower().strip()
        filas = []
        for cat, entidades in ST.indice_ner_global.items():
            if cat_filtro != "Todas" and cat != cat_filtro:
                continue
            for ent, arts in entidades.items():
                if buscar and buscar not in ent.lower():
                    continue
                filas.append((ent, cat, len(arts)))
        filas.sort(key=lambda r: (-r[2], r[1], r[0]))
        # Mapa wikidata: {cat: {entidad: {id, url, label}}}
        wiki = getattr(ST, "wikidata_enlaces", {}) or {}
        for ent, cat, n in filas:
            wiki_info = wiki.get(cat, {}).get(ent, {})
            wiki_id   = wiki_info.get("id", "")
            wiki_lbl  = wiki_info.get("label", wiki_id)
            wiki_cell = wiki_lbl if wiki_lbl else ""
            self._tv_ner.insert("", "end", values=(ent, cat, n, wiki_cell),
                                 tags=(cat,), iid=f"{cat}|{ent}")
        paleta = {"personas": "#DBEAFE", "lugares": "#D1FAE5",
                  "organizaciones": "#FEF3C7", "fechas": "#FCE7F3",
                  "obras_publicaciones": "#EDE9FE", "eventos_historicos": "#FEE2E2"}
        for cat, color in paleta.items():
            self._tv_ner.tag_configure(cat, background=color, foreground="#0D1117")

    def _ner_abrir_wikidata(self, event=None):
        """Doble click en la tabla NER: abre el enlace Wikidata en el browser."""
        sel = self._tv_ner.selection()
        if not sel:
            return
        iid = sel[0]
        if "|" not in iid:
            return
        cat, ent = iid.split("|", 1)
        wiki = getattr(ST, "wikidata_enlaces", {}) or {}
        wiki_info = wiki.get(cat, {}).get(ent, {})
        url = wiki_info.get("url", "")
        if url:
            import webbrowser
            webbrowser.open(url)
        elif wiki_info.get("id"):
            import webbrowser
            webbrowser.open(f"https://www.wikidata.org/wiki/{wiki_info['id']}")
        else:
            messagebox.showinfo("Sin enlace",
                                f"No hay enlace Wikidata para «{ent}».\n"
                                f"Usa 'Enlazar Wikidata' para buscar coincidencias.")

    def _ner_on_select(self, event=None):
        sel = self._tv_ner.selection()
        if not sel:
            return
        # El iid puede ser "cat|ent" o simplemente el índice según la versión del treeview
        iid = sel[0]
        if "|" in iid:
            cat, ent = iid.split("|", 1)
        else:
            vals = self._tv_ner.item(iid)["values"]
            if not vals or len(vals) < 2:
                return
            ent, cat = str(vals[0]), str(vals[1])
        arts = ST.indice_ner_global.get(cat, {}).get(ent, [])
        txt = f"Entidad: {ent}\nTipo: {cat}\nArtículos ({len(arts)}):\n"
        for a in sorted(arts)[:30]:
            txt += f"  • {a}\n"
        if len(arts) > 30:
            txt += f"  … y {len(arts)-30} más\n"
        # Añadir enlace Wikidata si está disponible
        wikidata = getattr(ST, "wikidata_enlaces", {})
        enlace = wikidata.get(cat, {}).get(ent)
        if enlace:
            txt += "\n── Wikidata ──\n"
            txt += f"  ID:          {enlace.get('id','')}\n"
            txt += f"  Nombre:      {enlace.get('label','')}\n"
            txt += f"  Descripción: {enlace.get('description','')}\n"
            txt += f"  URL:         {enlace.get('url','')}\n"
            txt += f"  Confianza:   {enlace.get('confianza', 0):.0%}\n"
        self._txt_ner_det.config(state="normal")
        self._txt_ner_det.delete("1.0", "end")
        self._txt_ner_det.insert("1.0", txt)
        self._txt_ner_det.config(state="disabled")

    def _ner_exportar_csv(self):
        if not getattr(ST, "indice_ner_global", None):
            messagebox.showwarning("Sin datos", "Ejecuta el análisis NER primero."); return
        from pathlib import Path as _Path
        from tkinter import filedialog
        pub = getattr(ST, "publicacion", "corpus").replace(" ", "_")
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile=f"indice_NER_{pub}.csv",
            title="Guardar índice NER")
        if not dest:
            return
        from core.ner_engine import exportar_csv
        n = exportar_csv(ST.indice_ner_global, _Path(dest))
        messagebox.showinfo("Exportado", f"✅ {n} entidades exportadas a:\n{dest}")

    def _ner_enlazar_wikidata(self):
        """Enlaza todas las entidades del índice NER con Wikidata en un thread."""
        if not getattr(ST, "indice_ner_global", None) or not any(ST.indice_ner_global.values()):
            messagebox.showwarning("Sin datos", "Ejecuta el análisis NER primero."); return
        try:
            from core.entity_linker import enlazar_indice_ner
        except ImportError:
            messagebox.showerror("Módulo no disponible", "core/entity_linker.py no encontrado."); return

        # Contar total de entidades
        total = sum(len(ents) for ents in ST.indice_ner_global.values())
        if total == 0:
            messagebox.showwarning("Sin entidades", "El índice NER está vacío."); return

        # Ventana de progreso
        win = tk.Toplevel(self)
        win.title("Enlazando con Wikidata…")
        win.resizable(False, False)
        win.grab_set()
        tk.Label(win, text="Enlazando entidades con Wikidata",
                 font=("Segoe UI", 11, "bold"), pady=12).pack()
        tk.Label(win, text=f"{total} entidades · puede tardar varios minutos",
                 font=("Segoe UI", 9), fg="#8B949E").pack()
        var_prog = tk.StringVar(value="0 / 0")
        lbl_prog = tk.Label(win, textvariable=var_prog, font=("Consolas", 9), pady=6)
        lbl_prog.pack()
        pb = ttk.Progressbar(win, mode="determinate", length=340, maximum=total)
        pb.pack(padx=20, pady=(0, 8))
        txt_log = scrolledtext.ScrolledText(win, height=8, width=52,
                                             font=("Consolas", 8), state="disabled")
        txt_log.pack(padx=12, pady=(0, 12))
        btn_cerrar = ttk.Button(win, text="Cerrar", state="disabled",
                                command=win.destroy)
        btn_cerrar.pack(pady=(0, 12))

        def _log(msg):
            txt_log.config(state="normal")
            txt_log.insert("end", msg + "\n")
            txt_log.see("end")
            txt_log.config(state="disabled")

        def _worker():
            encontradas = 0
            def _cb(n, tot):
                nonlocal encontradas
                self.after(0, lambda: pb.config(value=n))
                self.after(0, lambda: var_prog.set(f"{n} / {tot}"))

            # Mapa {art_id: texto} para desambiguación contextual de Wikidata.
            # Mismo criterio de id que el pipeline NER (id/titulo de df_articulos).
            textos_articulos = {}
            try:
                if ST.df_articulos is not None and not ST.df_articulos.empty:
                    for _i, _row in ST.df_articulos.iterrows():
                        _txt = str(_row.get("texto", _row.get("contenido", "")))
                        _aid = str(_row.get("id", _row.get("titulo", f"art_{_i}")))
                        if _txt.strip():
                            textos_articulos[_aid] = _txt
            except Exception:
                textos_articulos = {}

            try:
                resultado = enlazar_indice_ner(
                    ST.indice_ner_global,
                    sin_red=False,
                    callback=_cb,
                    textos_articulos=textos_articulos or None,
                )
                # Guardar resultado en ST y mostrar resumen
                ST.wikidata_enlaces = resultado
                for cat, ents in resultado.items():
                    for texto, enlace in ents.items():
                        if enlace:
                            encontradas += 1
                            self.after(0, lambda t=texto, e=enlace:
                                _log(f"✓ {t} → {e['label']} ({e['id']})"))
                self.after(0, lambda: _log(
                    f"\n✅ {encontradas}/{total} entidades enlazadas con Wikidata."))
                # Actualizar detalle si hay entidad seleccionada
                self.after(0, self._ner_refrescar_tv)
            except Exception as exc:
                self.after(0, lambda err=str(exc): _log(f"⚠️ Error: {err}"))
            finally:
                self.after(0, lambda: btn_cerrar.config(state="normal"))

        import threading
        threading.Thread(target=_worker, daemon=True).start()


    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: BÚSQUEDA SEMÁNTICA FAISS
    # ══════════════════════════════════════════════════════════════════════════

    def _build_busqueda_semantica(self):
        f = self._tab_bsem
        self._page_header(f, "Búsqueda Semántica",
                          "Encuentra artículos por similitud de significado · powered by FAISS + sentence-transformers",
                          "🔍")
        pad = tk.Frame(f, bg=CONTENT_BG); pad.pack(fill="both", expand=True, padx=24, pady=16)

        # ── Card: Construir índice ────────────────────────────────────────────
        card_idx = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1)
        card_idx.pack(fill="x", pady=(0, 12))
        hdr_idx = tk.Frame(card_idx, bg="#1C2128"); hdr_idx.pack(fill="x")
        tk.Label(hdr_idx, text="  📦  Índice vectorial", bg="#1C2128", fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(side="left", pady=4)
        self._lbl_bsem_estado = tk.Label(hdr_idx, text="Sin índice", bg="#1C2128",
                                          fg="#EF4444", font=("Segoe UI", 8))
        self._lbl_bsem_estado.pack(side="right", padx=10)

        body_idx = tk.Frame(card_idx, bg=CARD_BG, padx=12, pady=8)
        body_idx.pack(fill="x")
        tk.Label(body_idx,
                 text="Genera embeddings de todos los artículos del corpus y construye el índice FAISS.",
                 bg=CARD_BG, fg="#CDD6F4", font=("Segoe UI", 9), wraplength=580, justify="left"
                 ).pack(anchor="w", pady=(0, 6))
        bf_idx = tk.Frame(body_idx, bg=CARD_BG); bf_idx.pack(fill="x")
        self._btn_bsem_construir = ttk.Button(bf_idx, text="▶  Construir índice",
                                               style="P.TButton",
                                               command=self._bsem_construir)
        self._btn_bsem_construir.pack(side="left", padx=(0, 8))
        ttk.Button(bf_idx, text="💾  Guardar índice", style="S.TButton",
                   command=self._bsem_guardar).pack(side="left", padx=(0, 8))
        ttk.Button(bf_idx, text="📂  Cargar índice", style="S.TButton",
                   command=self._bsem_cargar).pack(side="left")
        self._lbl_bsem_n = tk.Label(bf_idx, text="", bg=CARD_BG, fg=VERDE,
                                     font=("Segoe UI", 8, "bold"))
        self._lbl_bsem_n.pack(side="right")

        # ── Card: Consulta ────────────────────────────────────────────────────
        card_q = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1)
        card_q.pack(fill="x", pady=(0, 12))
        hdr_q = tk.Frame(card_q, bg="#1C2128"); hdr_q.pack(fill="x")
        tk.Label(hdr_q, text="  🔎  Consulta", bg="#1C2128", fg=TXT_PRI,
                 font=("Segoe UI", 9, "bold")).pack(side="left", pady=4)

        body_q = tk.Frame(card_q, bg=CARD_BG, padx=12, pady=8)
        body_q.pack(fill="x")
        row_q = tk.Frame(body_q, bg=CARD_BG); row_q.pack(fill="x")
        tk.Label(row_q, text="Consulta:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI", 9, "bold"), width=9).pack(side="left")
        self._var_bsem_q = tk.StringVar()
        self._ent_bsem_q = tk.Entry(row_q, textvariable=self._var_bsem_q,
                                     font=("Segoe UI", 10), relief="solid", bd=1)
        self._ent_bsem_q.pack(side="left", fill="x", expand=True, padx=(6, 8))
        self._ent_bsem_q.bind("<Return>", lambda e: self._bsem_buscar())

        tk.Label(row_q, text="K:", bg=CARD_BG, fg="#CDD6F4",
                 font=("Segoe UI", 9, "bold")).pack(side="left")
        self._var_bsem_k = tk.IntVar(value=10)
        ttk.Spinbox(row_q, from_=1, to=50, textvariable=self._var_bsem_k,
                    width=4).pack(side="left", padx=(4, 8))
        self._btn_bsem_buscar = ttk.Button(row_q, text="Buscar", style="P.TButton",
                                            command=self._bsem_buscar)
        self._btn_bsem_buscar.pack(side="left")

        # ── Opciones de resultado ────────────────────────────────────────────
        opt_f = tk.Frame(body_q, bg=CARD_BG); opt_f.pack(fill="x", pady=(6, 0))
        ttk.Button(opt_f, text="💡  Explicar seleccionado", style="S.TButton",
                   command=self._bsem_explicar_seleccionado).pack(side="left", padx=(0, 8))
        ttk.Button(opt_f, text="📋  Explicar todos", style="S.TButton",
                   command=self._bsem_explicar_todos).pack(side="left")

        # ── Resultados ────────────────────────────────────────────────────────
        res_f = tk.Frame(pad, bg=CONTENT_BG); res_f.pack(fill="both", expand=True)

        # Treeview resultados
        tv_frame = tk.Frame(res_f, bg=CARD_BG, relief="solid", bd=1)
        tv_frame.pack(fill="both", expand=True)
        cols = ("rank", "articulo_id", "similitud", "titulo")
        self._tv_bsem = ttk.Treeview(tv_frame, columns=cols, show="headings", height=11)
        self._tv_bsem.heading("rank",        text="#")
        self._tv_bsem.heading("articulo_id", text="ID artículo")
        self._tv_bsem.heading("similitud",   text="Similitud")
        self._tv_bsem.heading("titulo",      text="Título / fragmento")
        self._tv_bsem.column("rank",        width=40,  anchor="center")
        self._tv_bsem.column("articulo_id", width=160, anchor="w")
        self._tv_bsem.column("similitud",   width=90,  anchor="center")
        self._tv_bsem.column("titulo",      width=460, anchor="w")
        vsb = ttk.Scrollbar(tv_frame, orient="vertical", command=self._tv_bsem.yview)
        self._tv_bsem.configure(yscrollcommand=vsb.set)
        vsb.pack(side="right", fill="y")
        self._tv_bsem.pack(fill="both", expand=True)
        self._bsem_resultados_raw = []  # lista de dicts con datos completos

        # Log
        log_bsem = tk.Frame(pad, bg="#0F1B2D", bd=1, relief="solid")
        log_bsem.pack(fill="x", pady=(8, 0))
        log_hdr = tk.Frame(log_bsem, bg="#1A2F4A"); log_hdr.pack(fill="x")
        tk.Label(log_hdr, text="  📋  Registro", bg="#1A2F4A", fg="#94A3B8",
                 font=("Segoe UI", 8, "bold")).pack(side="left", pady=4)
        self._log_bsem = scrolledtext.ScrolledText(log_bsem, height=4,
                                                    font=("Consolas", 8),
                                                    bg="#0F1B2D", fg="#86EFAC",
                                                    relief="flat", state="disabled")
        self._log_bsem.pack(fill="x", padx=1, pady=(0, 1))

        # Estado interno
        self._bsem_indice = None  # instancia IndiceSemantico

    def _bsem_explicar_seleccionado(self):
        sel = self._tv_bsem.selection()
        if not sel:
            messagebox.showinfo("Sin selección", "Selecciona un resultado primero.")
            return
        idx = self._tv_bsem.index(sel[0])
        if idx >= len(self._bsem_resultados_raw):
            return
        resultado = self._bsem_resultados_raw[idx]
        query = self._var_bsem_q.get().strip()
        self._bsem_mostrar_explicacion(query, [resultado])

    def _bsem_explicar_todos(self):
        if not self._bsem_resultados_raw:
            messagebox.showwarning("Sin resultados", "Realiza una búsqueda primero.")
            return
        query = self._var_bsem_q.get().strip()
        self._bsem_mostrar_explicacion(query, self._bsem_resultados_raw)

    def _bsem_mostrar_explicacion(self, query: str, resultados: list):
        from core.explainer import explicar_lote, resumir_busqueda
        explicados = explicar_lote(query, resultados)
        resumen    = resumir_busqueda(query, explicados)

        win = tk.Toplevel(self)
        win.title(f"Explicación de resultados: '{query}'")
        win.geometry("780x540")
        win.configure(bg=CONTENT_BG)

        tk.Label(win, text="Síntesis de la búsqueda", bg=CONTENT_BG,
                 fg="#CDD6F4", font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=16, pady=(12, 4))
        lbl_res = tk.Label(win, text=resumen, bg=CONTENT_BG, fg="#94A3B8",
                           font=("Segoe UI", 9), wraplength=740, justify="left")
        lbl_res.pack(anchor="w", padx=16, pady=(0, 10))

        txt = scrolledtext.ScrolledText(win, font=("Consolas", 9),
                                         bg="#1E1E2E", fg="#CDD6F4",
                                         relief="flat", padx=10, pady=8)
        txt.pack(fill="both", expand=True, padx=12, pady=(0, 12))

        for i, exp in enumerate(explicados, 1):
            e = exp.get("explicacion", {})
            proc = e.get("procedencia", {})
            txt.insert("end", f"── Resultado {i}: {proc.get('titulo','')[:60]} ──\n", "titulo")
            txt.insert("end", f"  {e.get('resumen_explicacion','')}\n")
            if e.get("terminos_relevantes"):
                tops = ", ".join(t["termino"] for t in e["terminos_relevantes"][:5])
                txt.insert("end", f"  Términos clave: {tops}\n")
            if e.get("fragmento_relevante"):
                txt.insert("end", f"  Fragmento: «{e['fragmento_relevante'][:120]}»\n")
            txt.insert("end", "\n")

        txt.tag_configure("titulo", foreground="#3B82F6", font=("Consolas", 9, "bold"))
        txt.config(state="disabled")

    def _bsem_log(self, msg: str):
        def _do():
            self._log_bsem.config(state="normal")
            self._log_bsem.insert("end", msg + "\n")
            self._log_bsem.see("end")
            self._log_bsem.config(state="disabled")
        self.after(0, _do)

    def _bsem_construir(self):
        """Genera embeddings de todos los artículos y construye el índice FAISS."""
        try:
            from core.busqueda_semantica import IndiceSemantico, faiss_disponible
            from core.embeddings_local import (
                generar_embeddings,
                sentence_transformers_disponible,
            )
        except ImportError as e:
            messagebox.showerror("Módulo no disponible", str(e)); return

        if not sentence_transformers_disponible():
            messagebox.showerror("sentence-transformers no instalado",
                                 "Ejecuta: pip install sentence-transformers"); return
        if not faiss_disponible():
            messagebox.showerror("faiss-cpu no instalado",
                                 "Ejecuta: pip install faiss-cpu"); return

        # Recopilar textos desde df_articulos o archivos OCR
        textos, ids = [], []
        if ST.df_articulos is not None and not ST.df_articulos.empty:
            for i, row in ST.df_articulos.iterrows():
                t = str(row.get("texto", row.get("contenido", "")))
                if t and t != "nan" and len(t.strip()) > 10:
                    textos.append(t[:2000])   # truncar a 2000 chars para velocidad
                    ids.append(str(row.get("id", row.get("titulo", f"art_{i}"))))
        elif ST.out_dir:
            txt_dir = ST.out_dir / "03_ocr"
            if txt_dir.exists():
                for tf in sorted(txt_dir.rglob("*.txt")):
                    t = tf.read_text("utf-8", errors="replace")
                    if len(t.strip()) > 10:
                        textos.append(t[:2000])
                        ids.append(tf.stem)

        if not textos:
            messagebox.showwarning("Sin textos",
                "No hay textos disponibles. Ejecuta el OCR primero."); return

        self._btn_bsem_construir.config(state="disabled")
        self._bsem_log(f"Construyendo índice para {len(textos)} artículos…")

        def _worker():
            try:
                self.after(0, lambda: self._bsem_log("  Generando embeddings…"))
                embs = generar_embeddings(textos, mostrar_progreso=False)
                self.after(0, lambda: self._bsem_log(f"  Embeddings: {embs.shape}"))

                indice = IndiceSemantico(dimension=embs.shape[1])
                indice.construir(embs, ids)
                self._bsem_indice = indice

                self.after(0, lambda: self._lbl_bsem_estado.config(
                    text=f"✓ {indice.n_articulos} artículos indexados", fg=VERDE))
                self.after(0, lambda: self._lbl_bsem_n.config(
                    text=f"{indice.n_articulos} artículos"))
                self.after(0, lambda: self._bsem_log(
                    f"✅ Índice listo · {indice.n_articulos} artículos"))
            except Exception as exc:
                self.after(0, lambda err=str(exc): self._bsem_log(f"⚠️ Error: {err}"))
            finally:
                self.after(0, lambda: self._btn_bsem_construir.config(state="normal"))

        import threading
        threading.Thread(target=_worker, daemon=True).start()

    def _bsem_buscar(self):
        """Busca artículos similares a la consulta. Si no hay índice FAISS, usa búsqueda léxica."""
        consulta = self._var_bsem_q.get().strip()
        if not consulta:
            return

        # Fallback léxico si no hay índice semántico
        if self._bsem_indice is None or not getattr(self._bsem_indice, "construido", False):
            self._bsem_buscar_lexico(consulta)
            return

        k = max(1, min(self._var_bsem_k.get(), self._bsem_indice.n_articulos))

        try:
            from core.embeddings_local import generar_embeddings
        except ImportError as e:
            messagebox.showerror("Módulo no disponible", str(e)); return

        self._btn_bsem_buscar.config(state="disabled")

        def _worker():
            try:
                embs = generar_embeddings([consulta])
                resultados = self._bsem_indice.buscar(embs[0], k=k)

                # Enriquecer con título si está disponible
                titulos = {}
                if ST.df_articulos is not None and not ST.df_articulos.empty:
                    for _, row in ST.df_articulos.iterrows():
                        aid = str(row.get("id", row.get("titulo", "")))
                        titulos[aid] = str(row.get("titulo", aid))

                def _actualizar():
                    for item in self._tv_bsem.get_children():
                        self._tv_bsem.delete(item)
                    for r in resultados:
                        aid = r["articulo_id"]
                        self._tv_bsem.insert("", "end", values=(
                            r["rank"],
                            aid,
                            f"{r['similitud']:.3f}",
                            titulos.get(aid, aid),
                        ))
                    self._bsem_log(
                        f"🔍 '{consulta[:60]}' → {len(resultados)} resultados")

                self.after(0, _actualizar)
            except Exception as exc:
                self.after(0, lambda err=str(exc): self._bsem_log(f"⚠️ Error: {err}"))
            finally:
                self.after(0, lambda: self._btn_bsem_buscar.config(state="normal"))

        import threading
        threading.Thread(target=_worker, daemon=True).start()

    def _bsem_buscar_lexico(self, consulta: str):
        """
        Búsqueda léxica de fallback — no requiere FAISS.
        Busca la consulta en los TXT del corpus con coincidencias exactas e inexactas.
        Muestra resultados en la tabla de búsqueda semántica.
        """
        corpus_txt = getattr(ST, "corpus_txt", None) or []
        corpus_meta = getattr(ST, "corpus_meta", None)
        if corpus_meta is None:
            corpus_meta = {}

        if not corpus_txt:
            self._bsem_log("⚠ Sin corpus cargado. Ejecuta primero la extracción OCR.")
            return

        self._bsem_log(f"🔍 Búsqueda léxica (sin índice FAISS): '{consulta}'")
        k = self._var_bsem_k.get() if hasattr(self, "_var_bsem_k") else 10
        terminos = [t.lower() for t in consulta.split() if len(t) > 2]

        def _worker():
            resultados = []
            for i, texto in enumerate(corpus_txt):
                if not texto:
                    continue
                texto_lower = texto.lower()
                # Score = número de términos encontrados / total términos
                encontrados = sum(1 for t in terminos if t in texto_lower)
                if encontrados == 0:
                    continue
                # Bonus por cercanía de términos (todos en la misma frase)
                score = encontrados / max(len(terminos), 1)
                if len(terminos) > 1 and all(t in texto_lower for t in terminos):
                    score = min(score * 1.5, 1.0)
                art_id = str(i)
                titulo = ""
                if isinstance(corpus_meta, dict):
                    meta = corpus_meta.get(art_id, {})
                    titulo = meta.get("titulo", "") if isinstance(meta, dict) else ""
                resultados.append({"rank": 0, "articulo_id": art_id,
                                   "similitud": score, "titulo": titulo})

            resultados.sort(key=lambda r: -r["similitud"])
            for j, r in enumerate(resultados[:k], 1):
                r["rank"] = j

            def _mostrar():
                for item in self._tv_bsem.get_children():
                    self._tv_bsem.delete(item)
                for r in resultados[:k]:
                    self._tv_bsem.insert("", "end", values=(
                        r["rank"], r["articulo_id"],
                        f"{r['similitud']:.3f}", r["titulo"] or r["articulo_id"]))
                self._bsem_log(
                    f"✅ {len(resultados[:k])} resultados léxicos "
                    f"(construye el índice FAISS para búsqueda semántica)")
            self.after(0, _mostrar)

        threading.Thread(target=_worker, daemon=True).start()

    def _bsem_guardar(self):
        """Guarda el índice FAISS en disco."""
        if self._bsem_indice is None or not self._bsem_indice.construido:
            messagebox.showwarning("Sin índice", "Construye el índice primero."); return
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension="",
            filetypes=[("Índice FAISS", "*.faiss"), ("Todos", "*.*")],
            initialfile="indice_semantico",
            title="Guardar índice semántico")
        if not dest:
            return
        # Quitar extensión si el usuario la puso manualmente
        ruta_base = dest.replace(".faiss", "")
        try:
            self._bsem_indice.guardar(ruta_base)
            messagebox.showinfo("Guardado",
                f"✅ Índice guardado:\n  {ruta_base}.faiss\n  {ruta_base}.ids.json")
            self._bsem_log(f"💾 Índice guardado en: {ruta_base}")
        except Exception as e:
            messagebox.showerror("Error al guardar", str(e))

    def _bsem_cargar(self):
        """Carga un índice FAISS desde disco."""
        try:
            from core.busqueda_semantica import IndiceSemantico
        except ImportError as e:
            messagebox.showerror("Módulo no disponible", str(e)); return
        from tkinter import filedialog
        src = filedialog.askopenfilename(
            filetypes=[("Índice FAISS", "*.faiss"), ("Todos", "*.*")],
            title="Abrir índice semántico")
        if not src:
            return
        ruta_base = src.replace(".faiss", "")
        try:
            indice = IndiceSemantico()
            ok = indice.cargar(ruta_base)
            if not ok:
                messagebox.showerror("Error", "No se pudo cargar el índice."); return
            self._bsem_indice = indice
            self._lbl_bsem_estado.config(
                text=f"✓ {indice.n_articulos} artículos (cargado)", fg=VERDE)
            self._lbl_bsem_n.config(text=f"{indice.n_articulos} artículos")
            self._bsem_log(f"📂 Índice cargado: {indice.n_articulos} artículos")
        except Exception as e:
            messagebox.showerror("Error al cargar", str(e))

    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: COLLOCATES Y REDES LÉXICAS
    # ══════════════════════════════════════════════════════════════════════════

    def _build_coloc(self):
        self._page_header(self._tab_coloc, "Collocates y Redes Léxicas",
                          "Palabras que co-ocurren con una clave · KWIC · Dispersión léxica", "🔤")
        outer = tk.Frame(self._tab_coloc, bg=CONTENT_BG)
        outer.pack(fill="both", expand=True, padx=16, pady=8)

        self._coloc_params: dict = {}
        try:
            from core.collocation_engine import PARAMS_SCHEMA as _COLOC_SCHEMA
            self._build_params_panel(outer, _COLOC_SCHEMA, self._coloc_params)
        except Exception:
            pass

        pad = tk.Frame(outer, bg=CONTENT_BG)
        pad.pack(side="left", fill="both", expand=True)

        # Botón bitácora en la barra superior
        bbar_coloc = tk.Frame(pad, bg=CONTENT_BG)
        bbar_coloc.pack(fill="x", pady=(0, 4))
        ttk.Button(bbar_coloc, text="📓 Nota", style="S.TButton",
                   command=lambda: self._bitacora_nueva_nota("coloc")).pack(side="right")

        nb = ttk.Notebook(pad)
        nb.pack(fill="both", expand=True)

        # ── Sub-pestaña: Collocates ──
        frm_col = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_col, text="  Collocates  ")
        pad_col = tk.Frame(frm_col, bg=CONTENT_BG, padx=10, pady=8); pad_col.pack(fill="both", expand=True)

        bf = tk.Frame(pad_col, bg=CONTENT_BG); bf.pack(fill="x", pady=(0, 6))
        tk.Label(bf, text="Palabra clave:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        self._var_coloc_kw = tk.StringVar()
        tk.Entry(bf, textvariable=self._var_coloc_kw, width=20,
                 font=("Segoe UI", 10), relief="solid", bd=1,
                 bg="#0D1B2A", fg="#CDD6F4").pack(side="left", padx=(0, 8))
        tk.Label(bf, text="Ventana:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_coloc_vent = tk.IntVar(value=5)
        ttk.Spinbox(bf, from_=2, to=15, textvariable=self._var_coloc_vent,
                    width=4).pack(side="left", padx=(0, 8))
        tk.Label(bf, text="Top N:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_coloc_n = tk.IntVar(value=20)
        ttk.Spinbox(bf, from_=5, to=50, textvariable=self._var_coloc_n,
                    width=4).pack(side="left", padx=(0, 8))
        self._btn_coloc = ttk.Button(bf, text="▶  Calcular", style="P.TButton",
                                      command=self._coloc_calcular)
        self._btn_coloc.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📊  Graficar red", style="S.TButton",
                   command=self._coloc_graficar_red).pack(side="left")
        ttk.Button(bf, text="💾 CSV", style="S.TButton",
                   command=self._coloc_exportar_csv).pack(side="left", padx=(8, 0))

        self._lbl_coloc_ok = tk.Label(pad_col, text="", bg=CONTENT_BG, fg=VERDE,
                                       font=("Segoe UI", 9, "bold"))
        self._lbl_coloc_ok.pack(anchor="w", pady=(0, 4))

        cols = ("palabra", "frecuencia", "pmi")
        self._tv_coloc = ttk.Treeview(pad_col, columns=cols, show="headings", height=14)
        heads = [("palabra", "Palabra collocate", 200),
                 ("frecuencia", "Frecuencia", 100),
                 ("pmi", "PMI (asociación)", 120)]
        for cid, txt, w in heads:
            self._tv_coloc.heading(cid, text=txt)
            self._tv_coloc.column(cid, width=w, anchor="w")
        sv = ttk.Scrollbar(pad_col, orient="vertical", command=self._tv_coloc.yview)
        self._tv_coloc.configure(yscrollcommand=sv.set)
        self._tv_coloc.pack(side="left", fill="both", expand=True)
        sv.pack(side="left", fill="y")

        # ── Sub-pestaña: KWIC ──
        frm_kwic = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_kwic, text="  KWIC  ")
        pad_kwic = tk.Frame(frm_kwic, bg=CONTENT_BG, padx=10, pady=8); pad_kwic.pack(fill="both", expand=True)

        bk = tk.Frame(pad_kwic, bg=CONTENT_BG); bk.pack(fill="x", pady=(0, 6))
        tk.Label(bk, text="Palabra:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        self._var_kwic_kw = tk.StringVar()
        tk.Entry(bk, textvariable=self._var_kwic_kw, width=22,
                 font=("Segoe UI", 10), relief="solid", bd=1,
                 bg="#0D1B2A", fg="#CDD6F4").pack(side="left", padx=(0, 8))
        ttk.Button(bk, text="▶  Buscar concordancias", style="P.TButton",
                   command=self._coloc_kwic).pack(side="left")
        ttk.Button(bk, text="💾 Exportar CSV", style="S.TButton",
                   command=self._coloc_kwic_exportar_csv).pack(side="left", padx=(8, 0))
        self._kwic_resultados: list[dict] = []

        self._txt_kwic = scrolledtext.ScrolledText(pad_kwic, font=("Consolas", 9),
                                                    bg="#0D1B2A", fg="#CDD6F4",
                                                    height=20, relief="flat")
        self._txt_kwic.pack(fill="both", expand=True)
        self._txt_kwic.tag_configure("kw", foreground="#F59E0B", font=("Consolas", 9, "bold"))

        # ── Sub-pestaña: Frecuencias ──
        frm_freq = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_freq, text="  Frecuencias  ")
        pad_freq = tk.Frame(frm_freq, bg=CONTENT_BG, padx=10, pady=8); pad_freq.pack(fill="both", expand=True)

        bfr = tk.Frame(pad_freq, bg=CONTENT_BG); bfr.pack(fill="x", pady=(0, 6))
        self._var_freq_n = tk.IntVar(value=30)
        tk.Label(bfr, text="Top N:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        ttk.Spinbox(bfr, from_=10, to=100, textvariable=self._var_freq_n,
                    width=4).pack(side="left", padx=(0, 8))
        ttk.Button(bfr, text="▶  Calcular frecuencias", style="P.TButton",
                   command=self._coloc_frecuencias).pack(side="left", padx=(0, 8))
        ttk.Button(bfr, text="📊  Graficar", style="S.TButton",
                   command=self._coloc_graficar_freq).pack(side="left")
        ttk.Button(bfr, text="💾 CSV", style="S.TButton",
                   command=self._coloc_freq_exportar_csv).pack(side="left", padx=(8, 0))
        self._var_freq_relativa = tk.BooleanVar(value=False)
        ttk.Checkbutton(bfr, text="Relativa (/10.000)", variable=self._var_freq_relativa).pack(
            side="left", padx=(12, 0))

        cols_f = ("rank", "palabra", "freq", "df")
        self._tv_freq = ttk.Treeview(pad_freq, columns=cols_f, show="headings", height=16)
        for cid, txt, w in [("rank","#",40),("palabra","Palabra",200),
                             ("freq","Frecuencia",100),("df","En N docs",100)]:
            self._tv_freq.heading(cid, text=txt)
            self._tv_freq.column(cid, width=w, anchor="w")
        svf = ttk.Scrollbar(pad_freq, orient="vertical", command=self._tv_freq.yview)
        self._tv_freq.configure(yscrollcommand=svf.set)
        self._tv_freq.pack(side="left", fill="both", expand=True)
        svf.pack(side="left", fill="y")

        # ── Sub-pestaña: N-gramas ─────────────────────────────────────────────
        frm_ng = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_ng, text="  N-gramas  ")
        pad_ng = tk.Frame(frm_ng, bg=CONTENT_BG, padx=10, pady=8)
        pad_ng.pack(fill="both", expand=True)

        bng = tk.Frame(pad_ng, bg=CONTENT_BG); bng.pack(fill="x", pady=(0, 6))
        tk.Label(bng, text="N:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_ng_n = tk.IntVar(value=2)
        ttk.Spinbox(bng, from_=2, to=5, textvariable=self._var_ng_n,
                    width=3).pack(side="left", padx=(0, 8))
        tk.Label(bng, text="Top:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_ng_top = tk.IntVar(value=30)
        ttk.Spinbox(bng, from_=10, to=100, textvariable=self._var_ng_top,
                    width=4).pack(side="left", padx=(0, 8))
        self._var_ng_sw = tk.BooleanVar(value=False)
        ttk.Checkbutton(bng, text="Filtrar stopwords", variable=self._var_ng_sw).pack(
            side="left", padx=(0, 8))
        ttk.Button(bng, text="▶  Calcular", style="P.TButton",
                   command=self._coloc_ngramas).pack(side="left")
        ttk.Button(bng, text="💾 CSV", style="S.TButton",
                   command=self._coloc_ngramas_csv).pack(side="left", padx=(8, 0))

        cols_ng = ("rank", "ngrama", "freq")
        self._tv_ng = ttk.Treeview(pad_ng, columns=cols_ng, show="headings", height=16)
        for cid, txt, w in [("rank","#",40),("ngrama","N-grama",320),("freq","Frecuencia",100)]:
            self._tv_ng.heading(cid, text=txt)
            self._tv_ng.column(cid, width=w, anchor="w")
        sv_ng = ttk.Scrollbar(pad_ng, orient="vertical", command=self._tv_ng.yview)
        self._tv_ng.configure(yscrollcommand=sv_ng.set)
        self._tv_ng.pack(side="left", fill="both", expand=True)
        sv_ng.pack(side="left", fill="y")

        # ── Sub-pestaña: Dispersión léxica ────────────────────────────────────
        frm_disp = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_disp, text="  Dispersión  ")
        pad_disp = tk.Frame(frm_disp, bg=CONTENT_BG, padx=10, pady=8)
        pad_disp.pack(fill="both", expand=True)

        bdisp = tk.Frame(pad_disp, bg=CONTENT_BG); bdisp.pack(fill="x", pady=(0, 6))
        tk.Label(bdisp, text="Palabras (coma):", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        self._var_disp_words = tk.StringVar()
        tk.Entry(bdisp, textvariable=self._var_disp_words, width=40,
                 font=("Segoe UI", 9), bg="#0D1B2A", fg="#CDD6F4",
                 relief="solid", bd=1).pack(side="left", padx=(0, 8))
        ttk.Button(bdisp, text="▶  Graficar", style="P.TButton",
                   command=self._coloc_dispersion).pack(side="left")

        self._frm_disp_canvas = tk.Frame(pad_disp, bg=CONTENT_BG)
        self._frm_disp_canvas.pack(fill="both", expand=True)

        # ── Sub-pestaña: Stopwords del proyecto ───────────────────────────────
        frm_sw = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_sw, text="  Stopwords  ")
        pad_sw = tk.Frame(frm_sw, bg=CONTENT_BG, padx=10, pady=8)
        pad_sw.pack(fill="both", expand=True)

        tk.Label(pad_sw,
                 text="Stopwords adicionales para este proyecto "
                      "(una por línea, se suman a la lista base en español):",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 4))
        self._txt_stopwords = scrolledtext.ScrolledText(
            pad_sw, height=10, bg="#1C2128", fg=TXT_PRI,
            insertbackground=TXT_PRI, font=("Courier New", 9),
            relief="solid", bd=1, wrap="word")
        self._txt_stopwords.pack(fill="both", expand=True)
        # Cargar stopwords del proyecto si existen
        sw_guardadas = getattr(ST, "stopwords_proyecto", [])
        if sw_guardadas:
            self._txt_stopwords.insert("1.0", "\n".join(sw_guardadas))
        bsw = tk.Frame(pad_sw, bg=CONTENT_BG); bsw.pack(fill="x", pady=(6, 0))
        ttk.Button(bsw, text="💾 Guardar stopwords del proyecto", style="P.TButton",
                   command=self._coloc_guardar_stopwords).pack(side="left")
        tk.Label(bsw,
                 text="Se aplican a Collocates, Frecuencias y N-gramas al activar 'Filtrar stopwords'",
                 bg=CONTENT_BG, fg=TXT_DIM, font=("Segoe UI", 8)).pack(
                 side="left", padx=10)

    def _coloc_calcular(self):
        from core.collocation_engine import collocates
        kw = self._var_coloc_kw.get().strip()
        if not kw:
            messagebox.showwarning("Sin palabra clave", "Escribe una palabra clave."); return
        corpus = getattr(ST, "corpus_txt", None) or []
        if not corpus:
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero."); return
        self._btn_coloc.config(state="disabled")
        self._lbl_coloc_ok.config(text="Calculando…")

        def _worker():
            res = collocates(corpus, kw,
                             ventana=self._var_coloc_vent.get(),
                             top_n=self._var_coloc_n.get())
            def _show():
                for row in self._tv_coloc.get_children():
                    self._tv_coloc.delete(row)
                for r in res:
                    self._tv_coloc.insert("", "end", values=(
                        r["palabra"], r["frecuencia"], f"{r['pmi']:.3f}"))
                self._lbl_coloc_ok.config(text=f"✅ {len(res)} collocates de «{kw}»")
                self._btn_coloc.config(state="normal")
            self.after(0, _show)

        threading.Thread(target=_worker, daemon=True).start()

    def _coloc_graficar_red(self):

        from core.collocation_engine import red_lexica
        corpus = getattr(ST, "corpus_txt", None) or []
        if not corpus:
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero."); return
        kw = self._var_coloc_kw.get().strip()
        palabras_clave = [kw] if kw else None

        def _worker():
            red = red_lexica(corpus, palabras_clave=palabras_clave, top_n_nodos=25)
            self.after(0, lambda: self._coloc_mostrar_red(red))

        threading.Thread(target=_worker, daemon=True).start()

    def _coloc_mostrar_red(self, red):
        try:
            import matplotlib.pyplot as plt
            import networkx as nx
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        except ImportError:
            messagebox.showerror("Falta networkx", "pip install networkx"); return

        G = nx.Graph()
        for n in red["nodos"]:
            G.add_node(n["id"], size=n["size"])
        for a in red["aristas"]:
            G.add_edge(a["source"], a["target"], weight=a["weight"])

        win = tk.Toplevel(self)
        win.title("Red léxica")
        win.geometry("700x560")
        win.configure(bg="#1E1E2E")

        fig, ax = plt.subplots(figsize=(8, 6))
        fig.patch.set_facecolor("#1E1E2E")
        ax.set_facecolor("#1E1E2E")
        pos = nx.spring_layout(G, seed=42, k=1.2)
        sizes = [G.nodes[n].get("size", 20) * 15 for n in G.nodes]
        nx.draw_networkx(G, pos=pos, ax=ax,
                         node_color="#3B82F6", node_size=sizes,
                         font_color="#CDD6F4", font_size=7,
                         edge_color="#45475A", width=0.8, alpha=0.85)
        ax.axis("off")
        plt.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    def _coloc_kwic(self):
        from core.collocation_engine import concordancias
        kw = self._var_kwic_kw.get().strip()
        if not kw:
            messagebox.showwarning("Sin palabra", "Escribe una palabra."); return
        corpus = getattr(ST, "corpus_txt", None) or []
        if not corpus:
            messagebox.showwarning("Sin corpus", "Extrae el texto primero."); return

        def _worker():
            res = concordancias(corpus, kw, max_resultados=200)
            self._kwic_resultados = res

            def _show():
                self._txt_kwic.config(state="normal")
                self._txt_kwic.delete("1.0", "end")
                for r in res:
                    self._txt_kwic.insert("end", r["izquierda"] + " ")
                    self._txt_kwic.insert("end", r["kwic"], "kw")
                    self._txt_kwic.insert("end", " " + r["derecha"] + "\n")
                self._txt_kwic.config(state="disabled")
            self.after(0, _show)

        threading.Thread(target=_worker, daemon=True).start()

    def _coloc_kwic_exportar_csv(self):
        """Exporta las concordancias KWIC actuales a CSV."""
        if not getattr(self, "_kwic_resultados", None):
            messagebox.showwarning("Sin resultados", "Busca concordancias primero.")
            return
        import csv
        from tkinter import filedialog
        kw = self._var_kwic_kw.get().strip() or "kwic"
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile=f"concordancias_{kw}.csv",
            title="Exportar concordancias KWIC")
        if not dest:
            return
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=["doc_idx", "izquierda", "kwic", "derecha", "posicion"])
            w.writeheader()
            w.writerows(self._kwic_resultados)
        messagebox.showinfo("Exportado", f"✅ {len(self._kwic_resultados)} concordancias exportadas:\n{dest}")

    def _coloc_frecuencias(self):
        from core.collocation_engine import frecuencias
        corpus = getattr(ST, "corpus_txt", None) or []
        if not corpus:
            messagebox.showwarning("Sin corpus", "Extrae el texto primero."); return
        relativa = getattr(self, "_var_freq_relativa", None)
        usar_relativa = relativa.get() if relativa else False

        def _worker():
            res = frecuencias(corpus, top_n=self._var_freq_n.get())
            # Total de tokens para normalización
            total_tokens = sum(r["freq"] for r in res) or 1
            self._freq_resultados = res
            self._freq_total_tokens = total_tokens

            def _show():
                for row in self._tv_freq.get_children():
                    self._tv_freq.delete(row)
                for i, r in enumerate(res, 1):
                    if usar_relativa:
                        freq_display = f"{r['freq'] / total_tokens * 10000:.1f}"
                    else:
                        freq_display = r["freq"]
                    self._tv_freq.insert("", "end", values=(
                        i, r["palabra"], freq_display, r.get("df", "")))
            self.after(0, _show)

        threading.Thread(target=_worker, daemon=True).start()

    def _coloc_exportar_csv(self):
        """Exporta collocates a CSV."""
        items = self._tv_coloc.get_children() if hasattr(self, "_tv_coloc") else []
        if not items:
            messagebox.showwarning("Sin datos", "Calcula collocates primero."); return
        import csv
        from tkinter import filedialog
        kw = self._var_coloc_kw.get().strip() or "collocates"
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile=f"collocates_{kw}.csv")
        if not dest: return
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            w.writerow(["palabra", "frecuencia", "pmi"])
            for iid in items:
                w.writerow(self._tv_coloc.item(iid)["values"])
        messagebox.showinfo("Exportado", f"✅ Collocates exportados:\n{dest}")

    def _coloc_freq_exportar_csv(self):
        """Exporta frecuencias a CSV."""
        items = self._tv_freq.get_children() if hasattr(self, "_tv_freq") else []
        if not items:
            messagebox.showwarning("Sin datos", "Calcula frecuencias primero."); return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile="frecuencias.csv")
        if not dest: return
        relativa = getattr(self, "_var_freq_relativa", None)
        usar_rel  = relativa.get() if relativa else False
        total     = getattr(self, "_freq_total_tokens", 1) or 1
        res       = getattr(self, "_freq_resultados", [])
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            if usar_rel:
                w.writerow(["rank", "palabra", "freq_relativa_x10000", "freq_absoluta", "n_documentos"])
                for i, r in enumerate(res, 1):
                    w.writerow([i, r["palabra"],
                                round(r["freq"]/total*10000, 2),
                                r["freq"], r.get("df","")])
            else:
                w.writerow(["rank", "palabra", "frecuencia", "n_documentos"])
                for i, r in enumerate(res, 1):
                    w.writerow([i, r["palabra"], r["freq"], r.get("df","")])
        messagebox.showinfo("Exportado", f"✅ Frecuencias exportadas:\n{dest}")

    def _coloc_ngramas(self):
        from core.collocation_engine import ngramas
        corpus = getattr(ST, "corpus_txt", None) or []
        if not corpus:
            messagebox.showwarning("Sin corpus", "Extrae el texto primero."); return
        n   = self._var_ng_n.get()
        top = self._var_ng_top.get()
        sw  = self._var_ng_sw.get()

        def _worker():
            res = ngramas(corpus, n=n, top_n=top, stopwords=sw)
            self._ng_resultados = res
            def _show():
                for row in self._tv_ng.get_children():
                    self._tv_ng.delete(row)
                for i, r in enumerate(res, 1):
                    self._tv_ng.insert("", "end", values=(i, r["ngrama"], r["frecuencia"]))
            self.after(0, _show)
        threading.Thread(target=_worker, daemon=True).start()

    def _coloc_ngramas_csv(self):
        items = self._tv_ng.get_children() if hasattr(self, "_tv_ng") else []
        if not items:
            messagebox.showwarning("Sin datos", "Calcula n-gramas primero."); return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile=f"ngramas_n{self._var_ng_n.get()}.csv")
        if not dest: return
        res = getattr(self, "_ng_resultados", [])
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            w.writerow(["rank", "ngrama", "frecuencia"])
            for i, r in enumerate(res, 1):
                w.writerow([i, r["ngrama"], r["frecuencia"]])
        messagebox.showinfo("Exportado", f"✅ N-gramas exportados:\n{dest}")

    def _coloc_dispersion(self):
        corpus = getattr(ST, "corpus_txt", None) or []
        if not corpus:
            messagebox.showwarning("Sin corpus", "Extrae el texto primero."); return
        palabras_raw = self._var_disp_words.get()
        palabras = [p.strip() for p in palabras_raw.split(",") if p.strip()]
        if not palabras:
            messagebox.showwarning("Sin palabras", "Escribe al menos una palabra."); return

        def _worker():
            from core.collocation_engine import dispersion
            res = dispersion(corpus, palabras)
            self.after(0, lambda: self._coloc_mostrar_dispersion(res, palabras))
        threading.Thread(target=_worker, daemon=True).start()

    def _coloc_mostrar_dispersion(self, resultado: dict, palabras: list[str]):
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            from core.chart_builder import _FONDO, _TEXTO
        except ImportError:
            messagebox.showerror("Falta matplotlib", "pip install matplotlib"); return

        # Limpiar canvas anterior
        for w in self._frm_disp_canvas.winfo_children():
            w.destroy()

        n = len(palabras)
        fig, axes = plt.subplots(n, 1, figsize=(10, max(2, n * 1.4)),
                                  facecolor=_FONDO, sharex=True)
        if n == 1:
            axes = [axes]

        colores = ["#58A6FF","#F59E0B","#3FB950","#F85149","#D2A8FF","#FFA657"]
        for ax, palabra, color in zip(axes, palabras, colores * 10):
            posiciones = resultado.get(palabra, [])
            ax.vlines(posiciones, 0, 1, linewidth=0.8, alpha=0.7, color=color)
            ax.set_yticks([])
            ax.set_ylabel(palabra, rotation=0, labelpad=50,
                          fontsize=9, color=_TEXTO, ha="right", va="center")
            ax.set_facecolor(_FONDO)
            for spine in ax.spines.values():
                spine.set_edgecolor("#30363D")

        axes[-1].set_xlabel("Posición en el corpus (0 = inicio, 1 = final)",
                             color=_TEXTO, fontsize=8)
        plt.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=self._frm_disp_canvas)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    def _coloc_guardar_stopwords(self):
        """Guarda las stopwords personalizadas en ST y en el proyecto."""
        texto = self._txt_stopwords.get("1.0", "end-1c")
        palabras = [p.strip().lower() for p in texto.splitlines() if p.strip()]
        ST.stopwords_proyecto = palabras
        # Actualizar la lista global en collocation_engine para esta sesión
        try:
            import core.collocation_engine as _ce
            _ce.STOPWORDS_ES = _ce.STOPWORDS_ES | frozenset(palabras)
        except Exception:
            pass
        messagebox.showinfo("Guardadas",
                            f"✅ {len(palabras)} stopwords guardadas para este proyecto.")

    def _coloc_graficar_freq(self):
        items = self._tv_freq.get_children()
        if not items:
            messagebox.showwarning("Sin datos", "Calcula las frecuencias primero."); return
        palabras, freqs = [], []
        for iid in list(items)[:25]:
            vals = self._tv_freq.item(iid)["values"]
            palabras.append(str(vals[1]))
            freqs.append(int(vals[2]))

        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            from core.chart_builder import _FONDO, _TEXTO, _fig
        except ImportError:
            messagebox.showerror("Falta matplotlib", "pip install matplotlib"); return

        win = tk.Toplevel(self)
        win.title("Frecuencias léxicas")
        win.geometry("720x480")
        win.configure(bg=CONTENT_BG)
        fig, ax = _fig(8, 5)
        ax.barh(palabras[::-1], freqs[::-1], color="#3B82F6", alpha=0.85)
        ax.set_xlabel("Frecuencia")
        ax.set_title("Palabras más frecuentes del corpus")
        plt.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: ANOTACIÓN SEMÁNTICA REVISABLE
    # ══════════════════════════════════════════════════════════════════════════

    def _build_anot(self):
        self._page_header(self._tab_anot, "Anotación Semántica",
                          "Revisa y corrige entidades detectadas automáticamente · historial de cambios", "✍️")
        pad = tk.Frame(self._tab_anot, bg=CONTENT_BG, padx=16, pady=8)
        pad.pack(fill="both", expand=True)

        # ── Controles ──
        bf = tk.Frame(pad, bg=CONTENT_BG); bf.pack(fill="x", pady=(0, 8))
        ttk.Button(bf, text="📥  Importar NER automático", style="P.TButton",
                   command=self._anot_importar_ner).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="⏳  Ver pendientes", style="S.TButton",
                   command=self._anot_ver_pendientes).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="💾  Exportar JSON-LD", style="S.TButton",
                   command=self._anot_exportar).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📊  Estadísticas", style="S.TButton",
                   command=self._anot_stats).pack(side="left")
        ttk.Button(bf, text="📓 Nota", style="S.TButton",
                   command=lambda: self._bitacora_nueva_nota("anot")).pack(side="right")
        self._lbl_anot_ok = tk.Label(pad, text="Sin anotaciones cargadas",
                                      bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9))
        self._lbl_anot_ok.pack(anchor="w", pady=(0, 6))

        # ── Filtros ──
        ff = tk.Frame(pad, bg=CONTENT_BG); ff.pack(fill="x", pady=(0, 6))
        tk.Label(ff, text="Categoría:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_anot_cat = tk.StringVar(value="Todas")
        self._cb_anot_cat = ttk.Combobox(ff, textvariable=self._var_anot_cat,
                                          values=["Todas","PER","LOC","ORG","OBRA","EVE","CARGO"],
                                          state="readonly", width=10)
        self._cb_anot_cat.pack(side="left", padx=(0, 12))
        tk.Label(ff, text="Estado:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_anot_estado = tk.StringVar(value="Todos")
        ttk.Combobox(ff, textvariable=self._var_anot_estado,
                     values=["Todos","auto","confirmada","corregida","rechazada","pendiente"],
                     state="readonly", width=12).pack(side="left", padx=(0, 8))
        ttk.Button(ff, text="Filtrar", style="S.TButton",
                   command=self._anot_refrescar).pack(side="left")

        # ── Tabla ──
        cols = ("id", "art_id", "texto_norm", "categoria", "estado", "confianza")
        self._tv_anot = ttk.Treeview(pad, columns=cols, show="headings", height=12)
        heads = [("id","ID",50),("art_id","Artículo",120),("texto_norm","Entidad",180),
                 ("categoria","Categoría",90),("estado","Estado",100),("confianza","Conf.",70)]
        for cid, txt, w in heads:
            self._tv_anot.heading(cid, text=txt)
            self._tv_anot.column(cid, width=w, anchor="w")
        sv = ttk.Scrollbar(pad, orient="vertical", command=self._tv_anot.yview)
        self._tv_anot.configure(yscrollcommand=sv.set)
        self._tv_anot.pack(side="left", fill="both", expand=True)
        sv.pack(side="left", fill="y")

        # Tags de color por estado
        self._tv_anot.tag_configure("confirmada", foreground="#22C55E")
        self._tv_anot.tag_configure("corregida",  foreground="#3B82F6")
        self._tv_anot.tag_configure("rechazada",  foreground="#EF4444")
        self._tv_anot.tag_configure("pendiente",  foreground="#F59E0B")
        self._tv_anot.tag_configure("auto",       foreground="#94A3B8")

        # Botones de acción sobre selección
        ab = tk.Frame(pad, bg=CONTENT_BG); ab.pack(fill="x", pady=(6, 0))
        for label, estado in [("✅ Confirmar","confirmada"),("✏️ Corregir","corregida"),
                               ("❌ Rechazar","rechazada")]:
            ttk.Button(ab, text=label, style="S.TButton",
                       command=lambda e=estado: self._anot_cambiar_estado(e)
                       ).pack(side="left", padx=(0, 8))
        ttk.Button(ab, text="🕐 Ver historial", style="S.TButton",
                   command=self._anot_ver_historial).pack(side="left")

        self._anot_gestor = None
        self._anot_db_ruta = None

    def _anot_gestor_activo(self):
        from core.annotation_engine import GestorAnotaciones
        if self._anot_gestor is None:
            ruta = getattr(ST, "ruta_db", "") or ""
            if ruta:
                db = ruta.replace(".db", "_anotaciones.db")
            else:
                db = str(Path.home() / ".bashkar" / "anotaciones.db")
            Path(db).parent.mkdir(parents=True, exist_ok=True)
            self._anot_db_ruta = db
            self._anot_gestor = GestorAnotaciones(db)
        return self._anot_gestor

    def _anot_importar_ner(self):
        ner = getattr(ST, "indice_ner_global", {})
        if not ner:
            messagebox.showwarning("Sin NER", "Ejecuta el análisis NER primero."); return

        # ST.indice_ner_global tiene forma {cat: {entidad: [art_ids]}}
        # importar_desde_ner espera         {art_id: {cat: [{texto, inicio, fin, confianza}]}}
        # Convertir antes de pasar.
        ner_por_art: dict = {}
        for cat, entidades in ner.items():
            for entidad, art_ids in entidades.items():
                for art_id in (art_ids if isinstance(art_ids, list) else [art_ids]):
                    ner_por_art.setdefault(art_id, {}).setdefault(cat, []).append({
                        "texto": entidad,
                        "inicio": 0,
                        "fin": len(str(entidad)),
                        "confianza": 0.9,
                    })

        g = self._anot_gestor_activo()
        n = g.importar_desde_ner(ner_por_art, reemplazar=True)
        self._lbl_anot_ok.config(text=f"✅ {n} anotaciones importadas")
        self._anot_refrescar()

    def _anot_refrescar(self):
        g = self._anot_gestor_activo()
        cat    = self._var_anot_cat.get()
        estado = self._var_anot_estado.get()
        rows   = g.por_articulo(
            art_id="",  # truco: buscar todas
            categoria=None if cat == "Todas" else cat,
            estado=None if estado == "Todos" else estado,
        ) if False else []
        # Consulta directa para todas las anotaciones con filtros
        import sqlite3
        con = sqlite3.connect(self._anot_db_ruta or ":memory:")
        con.row_factory = sqlite3.Row
        sql = "SELECT * FROM anotaciones WHERE 1=1"
        params = []
        if cat != "Todas":
            sql += " AND categoria=?"; params.append(cat)
        if estado != "Todos":
            sql += " AND estado=?"; params.append(estado)
        sql += " ORDER BY id DESC LIMIT 500"
        rows = [dict(r) for r in con.execute(sql, params).fetchall()]
        con.close()

        for row in self._tv_anot.get_children():
            self._tv_anot.delete(row)
        for r in rows:
            self._tv_anot.insert("", "end", tags=(r["estado"],), values=(
                r["id"], r["art_id"], r["texto_norm"],
                r["categoria"], r["estado"], f"{r['confianza']:.2f}"))
        stats = self._anot_gestor_activo().estadisticas()
        self._lbl_anot_ok.config(
            text=f"{stats['total']} anotaciones  ·  "
                 f"confirmadas: {stats['por_estado'].get('confirmada',0)}  "
                 f"pendientes: {stats['por_estado'].get('auto',0)}")

    def _anot_cambiar_estado(self, nuevo_estado: str):
        sel = self._tv_anot.selection()
        if not sel:
            messagebox.showinfo("Sin selección", "Selecciona una anotación."); return
        g = self._anot_gestor_activo()
        for iid in sel:
            vals = self._tv_anot.item(iid)["values"]
            g.actualizar(int(vals[0]), {"estado": nuevo_estado}, razon="revisión manual")
        self._anot_refrescar()

    def _anot_ver_historial(self):
        sel = self._tv_anot.selection()
        if not sel:
            messagebox.showinfo("Sin selección", "Selecciona una anotación."); return
        vals   = self._tv_anot.item(sel[0])["values"]
        anot_id = int(vals[0])
        hist   = self._anot_gestor_activo().historial(anot_id)
        win = tk.Toplevel(self)
        win.title(f"Historial — anotación {anot_id}")
        win.geometry("580x300")
        win.configure(bg=CONTENT_BG)
        txt = scrolledtext.ScrolledText(win, font=("Consolas", 9),
                                         bg="#1E1E2E", fg="#CDD6F4", relief="flat")
        txt.pack(fill="both", expand=True, padx=10, pady=10)
        if not hist:
            txt.insert("end", "Sin cambios registrados.")
        for h in hist:
            txt.insert("end",
                f"{h['ts']}  [{h['campo']}]  {h['valor_ant']} → {h['valor_nuevo']}"
                f"  ({h['razon']})\n")
        txt.config(state="disabled")

    def _anot_ver_pendientes(self):
        g = self._anot_gestor_activo()
        pend = g.pendientes()
        self._var_anot_estado.set("auto")
        self._anot_refrescar()
        messagebox.showinfo("Pendientes", f"{len(pend)} anotaciones sin revisar.")

    def _anot_exportar(self):
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON-LD", "*.json"), ("Todos", "*.*")],
            initialfile="anotaciones_bashkar.json")
        if not dest: return
        n = self._anot_gestor_activo().exportar_json(dest, solo_confirmadas=False)
        messagebox.showinfo("Exportado", f"{n} anotaciones exportadas a:\n{dest}")

    def _anot_stats(self):
        stats = self._anot_gestor_activo().estadisticas()
        msg = f"Total: {stats['total']}\n\nPor estado:\n"
        for k, v in stats.get("por_estado", {}).items():
            msg += f"  {k}: {v}\n"
        msg += "\nPor categoría:\n"
        for k, v in stats.get("por_categoria", {}).items():
            msg += f"  {k}: {v}\n"
        messagebox.showinfo("Estadísticas de anotación", msg)

    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: DETECCIÓN DE NOVEDAD Y CAMBIO DISCURSIVO
    # ══════════════════════════════════════════════════════════════════════════

    def _build_nov(self):
        self._page_header(self._tab_nov, "Novedad y Cambio Discursivo",
                          "Palabras nuevas · cambio de vocabulario entre períodos · eventos temáticos", "🆕")
        pad = tk.Frame(self._tab_nov, bg=CONTENT_BG, padx=16, pady=8)
        pad.pack(fill="both", expand=True)

        nb = ttk.Notebook(pad)
        nb.pack(fill="both", expand=True)

        # ── Sub-pestaña: Cambio discursivo ──
        frm_cd = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_cd, text="  Cambio discursivo  ")
        pad_cd = tk.Frame(frm_cd, bg=CONTENT_BG, padx=10, pady=8); pad_cd.pack(fill="both", expand=True)

        tk.Label(pad_cd,
                 text="Mide cuánto cambia el vocabulario entre números consecutivos. "
                      "Alta distancia = cambio abrupto de tema o tono.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9),
                 wraplength=800, justify="left").pack(anchor="w", pady=(0, 8))

        bf_cd = tk.Frame(pad_cd, bg=CONTENT_BG); bf_cd.pack(fill="x", pady=(0, 6))
        self._btn_nov_cd = ttk.Button(bf_cd, text="▶  Calcular cambio discursivo",
                                       style="P.TButton", command=self._nov_cambio)
        self._btn_nov_cd.pack(side="left", padx=(0, 8))
        ttk.Button(bf_cd, text="📊  Graficar", style="S.TButton",
                   command=self._nov_graficar_cambio).pack(side="left")
        self._lbl_nov_ok = tk.Label(pad_cd, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_nov_ok.pack(anchor="w", pady=(0, 4))

        cols_cd = ("periodo_a", "periodo_b", "distancia", "ganadas", "perdidas")
        self._tv_nov_cd = ttk.Treeview(pad_cd, columns=cols_cd, show="headings", height=8)
        for cid, txt, w in [("periodo_a","De",100),("periodo_b","A",100),
                              ("distancia","Distancia",90),
                              ("ganadas","Palabras ganadas",250),
                              ("perdidas","Palabras perdidas",250)]:
            self._tv_nov_cd.heading(cid, text=txt)
            self._tv_nov_cd.column(cid, width=w, anchor="w")
        sv_cd = ttk.Scrollbar(pad_cd, orient="vertical", command=self._tv_nov_cd.yview)
        self._tv_nov_cd.configure(yscrollcommand=sv_cd.set)
        self._tv_nov_cd.pack(side="left", fill="both", expand=True)
        sv_cd.pack(side="left", fill="y")
        self._nov_cambio_data = []

        # ── Sub-pestaña: Palabras nuevas ──
        frm_pn = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_pn, text="  Palabras nuevas  ")
        pad_pn = tk.Frame(frm_pn, bg=CONTENT_BG, padx=10, pady=8); pad_pn.pack(fill="both", expand=True)

        bf_pn = tk.Frame(pad_pn, bg=CONTENT_BG); bf_pn.pack(fill="x", pady=(0, 6))
        tk.Label(bf_pn, text="Freq. mínima:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_nov_freq = tk.IntVar(value=3)
        ttk.Spinbox(bf_pn, from_=1, to=20, textvariable=self._var_nov_freq,
                    width=4).pack(side="left", padx=(0, 8))
        ttk.Button(bf_pn, text="▶  Detectar palabras nuevas", style="P.TButton",
                   command=self._nov_palabras_nuevas).pack(side="left")

        self._txt_nov_pn = scrolledtext.ScrolledText(pad_pn, font=("Consolas", 9),
                                                      bg="#0D1B2A", fg="#CDD6F4",
                                                      height=18, relief="flat")
        self._txt_nov_pn.pack(fill="both", expand=True)

        # ── Sub-pestaña: Tendencia de términos ──
        frm_td = tk.Frame(nb, bg=CONTENT_BG); nb.add(frm_td, text="  Tendencia de términos  ")
        pad_td = tk.Frame(frm_td, bg=CONTENT_BG, padx=10, pady=8); pad_td.pack(fill="both", expand=True)

        tk.Label(pad_td, text="Términos a seguir (separados por coma):",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w")
        self._var_nov_terms = tk.StringVar(value="radio, cine, mujer, guerra, colombia")
        tk.Entry(pad_td, textvariable=self._var_nov_terms, width=60,
                 font=("Segoe UI", 9), relief="solid", bd=1,
                 bg="#0D1B2A", fg="#CDD6F4").pack(anchor="w", pady=(4, 8))
        bf_td = tk.Frame(pad_td, bg=CONTENT_BG); bf_td.pack(fill="x", pady=(0, 6))
        ttk.Button(bf_td, text="▶  Calcular tendencia", style="P.TButton",
                   command=self._nov_tendencia).pack(side="left", padx=(0, 8))
        ttk.Button(bf_td, text="📊  Graficar", style="S.TButton",
                   command=self._nov_graficar_tendencia).pack(side="left")
        self._nov_tendencia_data = {}
        self._lbl_nov_td = tk.Label(pad_td, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_nov_td.pack(anchor="w", pady=(0, 4))
        self._txt_nov_td = scrolledtext.ScrolledText(pad_td, font=("Consolas", 9),
                                                      bg="#0D1B2A", fg="#CDD6F4",
                                                      height=14, relief="flat")
        self._txt_nov_td.pack(fill="both", expand=True)

    def _nov_corpus_por_periodo(self) -> dict:
        """Agrupa textos del corpus por número."""
        from collections import defaultdict
        por_num = defaultdict(list)
        articulos = getattr(ST, "articulos", None) or []
        corpus_txt = getattr(ST, "corpus_txt", None) or []
        if articulos:
            for art in articulos:
                num  = str(art.get("numero", "sin_número"))
                txt  = art.get("texto", "") or ""
                if txt.strip():
                    por_num[num].append(txt)
        elif corpus_txt:
            for i, txt in enumerate(corpus_txt):
                por_num[f"pag_{i:04d}"].append(txt or "")
        return dict(por_num)

    def _nov_cambio(self):
        from core.novelty_engine import cambio_discursivo
        corpus = self._nov_corpus_por_periodo()
        if len(corpus) < 2:
            messagebox.showwarning("Pocos datos",
                "Se necesitan al menos 2 períodos. Segmenta el corpus primero."); return
        self._btn_nov_cd.config(state="disabled")
        self._lbl_nov_ok.config(text="Calculando…")

        def _worker():
            res = cambio_discursivo(corpus)
            self._nov_cambio_data = res
            def _show():
                for row in self._tv_nov_cd.get_children():
                    self._tv_nov_cd.delete(row)
                for r in res:
                    ganadas  = ", ".join(g["palabra"] for g in r["palabras_ganadas"][:5])
                    perdidas = ", ".join(p["palabra"] for p in r["palabras_perdidas"][:5])
                    self._tv_nov_cd.insert("", "end", values=(
                        r["periodo_a"], r["periodo_b"],
                        f"{r['distancia']:.3f}", ganadas, perdidas))
                self._lbl_nov_ok.config(
                    text=f"✅ {len(res)} transiciones analizadas")
                self._btn_nov_cd.config(state="normal")
            self.after(0, _show)

        threading.Thread(target=_worker, daemon=True).start()

    def _nov_graficar_cambio(self):
        if not self._nov_cambio_data:
            messagebox.showwarning("Sin datos", "Calcula el cambio discursivo primero."); return
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            from core.chart_builder import _fig
        except ImportError:
            messagebox.showerror("Falta matplotlib", "pip install matplotlib"); return

        pares = [f"{r['periodo_a']}→{r['periodo_b']}" for r in self._nov_cambio_data]
        dists = [r["distancia"] for r in self._nov_cambio_data]
        pares_ord = [x for _, x in sorted(zip(
            [r["periodo_a"] for r in self._nov_cambio_data], pares))]
        dists_ord = [d for _, d in sorted(zip(
            [r["periodo_a"] for r in self._nov_cambio_data], dists))]

        win = tk.Toplevel(self)
        win.title("Cambio discursivo por período")
        win.geometry("720x400")
        win.configure(bg=CONTENT_BG)
        fig, ax = _fig(8, 4)
        colores = ["#EF4444" if d > 0.5 else "#3B82F6" for d in dists_ord]
        ax.bar(pares_ord, dists_ord, color=colores, alpha=0.85)
        ax.axhline(0.5, color="#F59E0B", linestyle="--", linewidth=1, label="umbral alto")
        ax.set_ylabel("Distancia coseno")
        ax.set_title("Cambio discursivo entre períodos consecutivos")
        ax.legend(facecolor="#313244", labelcolor="#CDD6F4", fontsize=8)
        plt.xticks(rotation=35, ha="right")
        plt.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    def _nov_palabras_nuevas(self):
        from core.novelty_engine import palabras_nuevas
        corpus = self._nov_corpus_por_periodo()
        if not corpus:
            messagebox.showwarning("Sin corpus", "Segmenta el corpus primero."); return

        def _worker():
            res = palabras_nuevas(corpus, min_freq=self._var_nov_freq.get())
            def _show():
                self._txt_nov_pn.config(state="normal")
                self._txt_nov_pn.delete("1.0", "end")
                for periodo in sorted(res.keys()):
                    nuevas = res[periodo]
                    self._txt_nov_pn.insert("end",
                        f"── {periodo} ({len(nuevas)} palabras nuevas) ──\n",
                        "titulo")
                    if nuevas:
                        self._txt_nov_pn.insert("end",
                            "  " + ", ".join(nuevas[:40]) + "\n\n")
                    else:
                        self._txt_nov_pn.insert("end", "  (ninguna)\n\n")
                self._txt_nov_pn.tag_configure(
                    "titulo", foreground="#3B82F6", font=("Consolas", 9, "bold"))
                self._txt_nov_pn.config(state="disabled")
            self.after(0, _show)

        threading.Thread(target=_worker, daemon=True).start()

    def _nov_tendencia(self):
        from core.novelty_engine import tendencia_vocabulario
        corpus = self._nov_corpus_por_periodo()
        if not corpus:
            messagebox.showwarning("Sin corpus", "Segmenta el corpus primero."); return
        palabras = [p.strip() for p in self._var_nov_terms.get().split(",") if p.strip()]
        if not palabras:
            messagebox.showwarning("Sin términos", "Escribe términos a seguir."); return

        def _worker():
            res = tendencia_vocabulario(corpus, palabras)
            self._nov_tendencia_data = res
            def _show():
                self._txt_nov_td.config(state="normal")
                self._txt_nov_td.delete("1.0", "end")
                periodos = sorted(corpus.keys())
                header = f"{'Término':<20}" + "".join(f"{p:<14}" for p in periodos) + "\n"
                self._txt_nov_td.insert("end", header, "header")
                self._txt_nov_td.insert("end", "─" * len(header) + "\n")
                for palabra in palabras:
                    row = f"{palabra:<20}"
                    for p in periodos:
                        v = res.get(palabra, {}).get(p, 0)
                        row += f"{v:<14.1f}"
                    self._txt_nov_td.insert("end", row + "\n")
                self._txt_nov_td.tag_configure(
                    "header", foreground="#F59E0B", font=("Consolas", 9, "bold"))
                self._lbl_nov_td.config(text=f"✅ Tendencia de {len(palabras)} términos")
                self._txt_nov_td.config(state="disabled")
            self.after(0, _show)

        threading.Thread(target=_worker, daemon=True).start()

    def _nov_graficar_tendencia(self):
        if not self._nov_tendencia_data:
            messagebox.showwarning("Sin datos", "Calcula la tendencia primero."); return
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            from core.chart_builder import PALETA, _fig
        except ImportError:
            messagebox.showerror("Falta matplotlib", "pip install matplotlib"); return

        win = tk.Toplevel(self)
        win.title("Tendencia de vocabulario")
        win.geometry("760x460")
        win.configure(bg=CONTENT_BG)
        fig, ax = _fig(9, 5)
        corpus = self._nov_corpus_por_periodo()
        periodos = sorted(corpus.keys())
        for i, (palabra, por_periodo) in enumerate(self._nov_tendencia_data.items()):
            vals = [por_periodo.get(p, 0) for p in periodos]
            ax.plot(periodos, vals, marker="o", label=palabra,
                    color=PALETA[i % len(PALETA)], linewidth=2)
        ax.set_ylabel("Frecuencia relativa (×10.000)")
        ax.set_title("Tendencia de términos a lo largo del corpus")
        ax.legend(facecolor="#313244", labelcolor="#CDD6F4", fontsize=8)
        plt.xticks(rotation=30, ha="right")
        plt.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: REDES DE CO-OCURRENCIA (v12)
    # ══════════════════════════════════════════════════════════════════════════

    def _build_red(self):
        pad = tk.Frame(self._tab_red, bg=CONTENT_BG, padx=16, pady=12)
        pad.pack(fill="both", expand=True)
        tk.Label(pad, text="Redes de co-ocurrencia", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad, text="Construye un grafo de entidades que co-ocurren en los mismos artículos.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 10))

        # ── Controles ────────────────────────────────────────────────────────
        ctrl = tk.Frame(pad, bg=CONTENT_BG)
        ctrl.pack(fill="x", pady=(0, 8))

        # Categorías
        tk.Label(ctrl, text="Categorías:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).grid(row=0, column=0, sticky="w", padx=(0, 4))
        self._var_red_cats = {}
        cats_frame = tk.Frame(ctrl, bg=CONTENT_BG)
        cats_frame.grid(row=0, column=1, sticky="w")
        cat_labels = {
            "personas": "Personas",
            "lugares": "Lugares",
            "organizaciones": "Organizaciones",
            "obras_publicaciones": "Obras",
            "eventos_historicos": "Eventos",
        }
        for cat, lbl in cat_labels.items():
            v = tk.BooleanVar(value=True)
            self._var_red_cats[cat] = v
            ttk.Checkbutton(cats_frame, text=lbl, variable=v).pack(side="left", padx=4)

        # Peso mínimo
        tk.Label(ctrl, text="Co-ocurrencias mínimas:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).grid(row=1, column=0, sticky="w", padx=(0, 4), pady=(6,0))
        self._var_red_peso = tk.IntVar(value=2)
        ttk.Spinbox(ctrl, from_=1, to=20, textvariable=self._var_red_peso,
                    width=5).grid(row=1, column=1, sticky="w", pady=(6,0))

        # Botones
        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 8))
        self._btn_red_construir = ttk.Button(bf, text="▶  Construir red",
                                              style="P.TButton",
                                              command=self._red_construir)
        self._btn_red_construir.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🔬  Métricas avanzadas",
                   style="S.TButton",
                   command=self._red_metricas_avanzadas).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📅  Evolución temporal",
                   style="S.TButton",
                   command=self._red_evolucion_temporal).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🌐  Ver en navegador",
                   style="S.TButton",
                   command=self._red_abrir_html).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📤  Exportar Gephi",
                   style="S.TButton",
                   command=self._red_exportar_gephi).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="💾  CSV métricas",
                   style="S.TButton",
                   command=self._red_exportar_csv).pack(side="left", padx=(0, 8))
        self._lbl_red_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_red_ok.pack(anchor="w", pady=(0, 6))

        # ── Notebook con pestañas de análisis ────────────────────────────────
        nb_red = ttk.Notebook(pad)
        nb_red.pack(fill="both", expand=True, pady=(4, 0))

        # ── Pestaña: Métricas globales + top centralidad ──────────────────────
        frm_met = tk.Frame(nb_red, bg=CONTENT_BG)
        nb_red.add(frm_met, text="  Métricas  ")

        met_frame = tk.Frame(frm_met, bg=CONTENT_BG)
        met_frame.pack(fill="x", pady=(6, 4))
        cols_met = ("metrica", "valor")
        self._tv_red_met = ttk.Treeview(met_frame, columns=cols_met,
                                         show="headings", height=7)
        self._tv_red_met.heading("metrica", text="Métrica")
        self._tv_red_met.heading("valor",   text="Valor")
        self._tv_red_met.column("metrica", width=240, anchor="w")
        self._tv_red_met.column("valor",   width=160, anchor="e")
        self._tv_red_met.pack(fill="x", padx=6)

        tk.Label(frm_met, text="Top nodos por centralidad de grado",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9, "bold")).pack(
                 anchor="w", padx=6, pady=(6, 2))
        cols_top = ("rango", "entidad", "categoria", "grado")
        self._tv_red_top = ttk.Treeview(frm_met, columns=cols_top,
                                         show="headings", height=8)
        self._tv_red_top.heading("rango",    text="#")
        self._tv_red_top.heading("entidad",  text="Entidad")
        self._tv_red_top.heading("categoria",text="Categoría")
        self._tv_red_top.heading("grado",    text="Centralidad grado")
        self._tv_red_top.column("rango",     width=35,  anchor="center")
        self._tv_red_top.column("entidad",   width=220, anchor="w")
        self._tv_red_top.column("categoria", width=140, anchor="w")
        self._tv_red_top.column("grado",     width=130, anchor="e")
        sv0 = ttk.Scrollbar(frm_met, orient="vertical", command=self._tv_red_top.yview)
        self._tv_red_top.configure(yscrollcommand=sv0.set)
        sv0.pack(side="right", fill="y", padx=(0, 6))
        self._tv_red_top.pack(fill="both", expand=True, padx=6)

        # ── Pestaña: Métricas avanzadas (betweenness, PageRank, closeness) ────
        frm_av = tk.Frame(nb_red, bg=CONTENT_BG)
        nb_red.add(frm_av, text="  Centralidad avanzada  ")

        bav = tk.Frame(frm_av, bg=CONTENT_BG)
        bav.pack(fill="x", padx=6, pady=(6, 4))
        tk.Label(bav, text="Tipo de centralidad:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        self._var_red_cent_tipo = tk.StringVar(value="betweenness")
        for val, lbl in [("betweenness","Betweenness (puentes)"),
                          ("pagerank","PageRank (influencia)"),
                          ("closeness","Closeness (proximidad)")]:
            ttk.Radiobutton(bav, text=lbl, variable=self._var_red_cent_tipo,
                            value=val,
                            command=self._red_refrescar_avanzadas).pack(side="left", padx=4)

        cols_av = ("rango", "entidad", "categoria", "valor")
        self._tv_red_av = ttk.Treeview(frm_av, columns=cols_av,
                                        show="headings", height=16)
        self._tv_red_av.heading("rango",    text="#")
        self._tv_red_av.heading("entidad",  text="Entidad")
        self._tv_red_av.heading("categoria",text="Categoría")
        self._tv_red_av.heading("valor",    text="Valor")
        self._tv_red_av.column("rango",     width=35,  anchor="center")
        self._tv_red_av.column("entidad",   width=230, anchor="w")
        self._tv_red_av.column("categoria", width=140, anchor="w")
        self._tv_red_av.column("valor",     width=120, anchor="e")
        sv1 = ttk.Scrollbar(frm_av, orient="vertical", command=self._tv_red_av.yview)
        self._tv_red_av.configure(yscrollcommand=sv1.set)
        sv1.pack(side="right", fill="y", padx=(0, 6))
        self._tv_red_av.pack(fill="both", expand=True, padx=6)
        self._red_metricas_av_cache: dict = {}

        # ── Pestaña: Comunidades ──────────────────────────────────────────────
        frm_com = tk.Frame(nb_red, bg=CONTENT_BG)
        nb_red.add(frm_com, text="  Comunidades  ")

        tk.Label(frm_com,
                 text="Comunidades detectadas por algoritmo Louvain\n"
                      "(grupos de entidades fuertemente conectadas entre sí)",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 8)).pack(
                 anchor="w", padx=6, pady=(6, 4))

        split_com = tk.Frame(frm_com, bg=CONTENT_BG)
        split_com.pack(fill="both", expand=True)

        # Lista de comunidades (izquierda)
        izq_com = tk.Frame(split_com, bg=CONTENT_BG, width=180)
        izq_com.pack(side="left", fill="y", padx=(6, 0))
        izq_com.pack_propagate(False)
        tk.Label(izq_com, text="Comunidades", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 8, "bold")).pack(anchor="w")
        self._lb_red_comunidades = tk.Listbox(
            izq_com, bg=CARD_BG, fg=TXT_SEC, selectbackground=AB_SEL,
            font=("Segoe UI", 9), relief="flat", activestyle="none")
        self._lb_red_comunidades.pack(fill="both", expand=True)
        self._lb_red_comunidades.bind("<<ListboxSelect>>",
                                      self._red_mostrar_comunidad)

        # Miembros de la comunidad (derecha)
        der_com = tk.Frame(split_com, bg=CONTENT_BG)
        der_com.pack(side="left", fill="both", expand=True, padx=6)
        tk.Label(der_com, text="Miembros", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 8, "bold")).pack(anchor="w")
        cols_com = ("entidad", "categoria", "grado")
        self._tv_red_miembros = ttk.Treeview(der_com, columns=cols_com,
                                              show="headings", height=18)
        self._tv_red_miembros.heading("entidad",  text="Entidad")
        self._tv_red_miembros.heading("categoria",text="Categoría")
        self._tv_red_miembros.heading("grado",    text="Grado")
        self._tv_red_miembros.column("entidad",   width=200, anchor="w")
        self._tv_red_miembros.column("categoria", width=130, anchor="w")
        self._tv_red_miembros.column("grado",     width=70,  anchor="e")
        sv2 = ttk.Scrollbar(der_com, orient="vertical",
                             command=self._tv_red_miembros.yview)
        self._tv_red_miembros.configure(yscrollcommand=sv2.set)
        sv2.pack(side="right", fill="y")
        self._tv_red_miembros.pack(fill="both", expand=True)
        self._red_comunidades_cache: list = []

        # ── Pestaña: Evolución temporal ───────────────────────────────────────
        frm_evo = tk.Frame(nb_red, bg=CONTENT_BG)
        nb_red.add(frm_evo, text="  Evolución temporal  ")

        tk.Label(frm_evo,
                 text="Cómo cambia la red entre números del corpus "
                      "(requiere haber procesado varios números).",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 8)).pack(
                 anchor="w", padx=6, pady=(6, 4))

        bevo = tk.Frame(frm_evo, bg=CONTENT_BG)
        bevo.pack(fill="x", padx=6, pady=(0, 4))
        ttk.Button(bevo, text="▶  Calcular evolución", style="P.TButton",
                   command=self._red_calcular_evolucion).pack(side="left")
        ttk.Button(bevo, text="📊  Graficar", style="S.TButton",
                   command=self._red_graficar_evolucion).pack(side="left", padx=(8, 0))
        ttk.Button(bevo, text="💾  CSV", style="S.TButton",
                   command=self._red_evolucion_csv).pack(side="left", padx=(8, 0))

        cols_evo = ("numero", "nodos", "aristas", "densidad", "top_nodo")
        self._tv_red_evo = ttk.Treeview(frm_evo, columns=cols_evo,
                                         show="headings", height=16)
        for cid, txt, w in [("numero","Número",150),("nodos","Nodos",70),
                              ("aristas","Aristas",70),("densidad","Densidad",80),
                              ("top_nodo","Nodo central",200)]:
            self._tv_red_evo.heading(cid, text=txt)
            self._tv_red_evo.column(cid, width=w, anchor="w")
        sv3 = ttk.Scrollbar(frm_evo, orient="vertical",
                             command=self._tv_red_evo.yview)
        self._tv_red_evo.configure(yscrollcommand=sv3.set)
        sv3.pack(side="right", fill="y", padx=(0, 6))
        self._tv_red_evo.pack(fill="both", expand=True, padx=6)
        self._red_evolucion_cache: list = []

        # ── Pestaña: Grafo canónico (entidades fundidas + relaciones) ─────────
        frm_can = tk.Frame(nb_red, bg=CONTENT_BG)
        nb_red.add(frm_can, text="  Grafo canónico  ")

        tk.Label(frm_can,
                 text="Funde las menciones NER en entidades canónicas (id estable) y\n"
                      "modela tripletas sujeto–predicado–objeto con procedencia y confianza.\n"
                      "Capa de grafo de conocimiento en el SQLite del proyecto.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 8), justify="left").pack(
                 anchor="w", padx=6, pady=(6, 4))

        bcan = tk.Frame(frm_can, bg=CONTENT_BG)
        bcan.pack(fill="x", padx=6, pady=(0, 4))
        self._btn_can_fundir = ttk.Button(
            bcan, text="▶  Fundir menciones → entidades canónicas",
            style="P.TButton", command=self._can_fundir)
        self._btn_can_fundir.pack(side="left")
        ttk.Button(bcan, text="🔗  Generar tripletas «mencionado_en»",
                   style="S.TButton",
                   command=self._can_generar_menciones).pack(side="left", padx=(8, 0))
        ttk.Button(bcan, text="💾  GEXF",
                   style="S.TButton",
                   command=self._can_exportar_gexf).pack(side="left", padx=(8, 0))
        ttk.Button(bcan, text="🌐  RDF/Turtle",
                   style="S.TButton",
                   command=self._can_exportar_rdf).pack(side="left", padx=(8, 0))

        # segunda fila: exploradores + editor de relaciones
        bcan2 = tk.Frame(frm_can, bg=CONTENT_BG)
        bcan2.pack(fill="x", padx=6, pady=(0, 4))
        ttk.Button(bcan2, text="🗺  Mapa de lugares",
                   style="S.TButton",
                   command=self._can_mapa_lugares).pack(side="left")
        ttk.Button(bcan2, text="📅  Timeline de números",
                   style="S.TButton",
                   command=self._can_timeline).pack(side="left", padx=(8, 0))
        ttk.Button(bcan2, text="📖  Vocabulario controlado",
                   style="S.TButton",
                   command=self._can_vocabulario).pack(side="left", padx=(8, 0))
        ttk.Button(bcan2, text="➕  Añadir relación…",
                   style="S.TButton",
                   command=self._can_editor_relacion).pack(side="left", padx=(8, 0))

        self._lbl_can_ok = tk.Label(frm_can, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_can_ok.pack(anchor="w", padx=6, pady=(2, 4))

        cols_can = ("id", "tipo", "nombre", "menciones", "wikidata", "fuente")
        self._tv_can = ttk.Treeview(frm_can, columns=cols_can,
                                    show="headings", height=16)
        for cid, txt, w, anc in [("id","id estable",200,"w"),("tipo","Tipo",90,"w"),
                                  ("nombre","Nombre",200,"w"),("menciones","Menc.",55,"e"),
                                  ("wikidata","Wikidata",90,"w"),("fuente","Fuente",90,"w")]:
            self._tv_can.heading(cid, text=txt)
            self._tv_can.column(cid, width=w, anchor=anc)
        svc = ttk.Scrollbar(frm_can, orient="vertical", command=self._tv_can.yview)
        self._tv_can.configure(yscrollcommand=svc.set)
        svc.pack(side="right", fill="y", padx=(0, 6))
        self._tv_can.pack(fill="both", expand=True, padx=6)

        # ── Log ───────────────────────────────────────────────────────────────
        self._txt_red_log = scrolledtext.ScrolledText(
            pad, height=4, font=("Consolas", 8),
            bg="#0D1B2A", fg="#94A3B8", state="disabled", wrap="word")
        self._txt_red_log.pack(fill="x", pady=(6, 0))

        self._grafo_actual = None
        self._html_red_path = None

    def _red_log(self, msg: str):
        self._txt_red_log.config(state="normal")
        self._txt_red_log.insert("end", msg + "\n")
        self._txt_red_log.see("end")
        self._txt_red_log.config(state="disabled")

    # ── Grafo canónico (entidades + relaciones) ───────────────────────────────

    def _can_repo(self):
        """Repositorio del proyecto activo, o None con aviso si no hay DB."""
        if not getattr(ST, "ruta_db", ""):
            messagebox.showwarning(
                "Sin proyecto",
                "Abre o guarda un proyecto con base de datos para usar el grafo canónico.")
            return None
        from datos.repositorio import Repositorio
        return Repositorio(ST.ruta_db)

    def _can_fundir(self):
        repo = self._can_repo()
        if repo is None:
            return
        self._btn_can_fundir.config(state="disabled")
        self._lbl_can_ok.config(text="Fundiendo menciones…")
        threading.Thread(target=self._worker_can_fundir,
                         args=(repo,), daemon=True).start()

    def _worker_can_fundir(self, repo):
        try:
            res = repo.fundir_menciones_en_canonicas(fuente="ner")
            cans = repo.listar_entidades_canonicas()
        except Exception as e:
            self.after(0, lambda err=e: messagebox.showerror("Error", str(err)))
            self.after(0, lambda: self._btn_can_fundir.config(state="normal"))
            return

        def _ui():
            self._btn_can_fundir.config(state="normal")
            self._can_poblar_tabla(cans)
            self._lbl_can_ok.config(
                text=f"✓ {res['canonicas']} entidades canónicas "
                     f"({res['menciones_vinculadas']} menciones fundidas)")
            self._red_log(f"[grafo] {res['canonicas']} canónicas, "
                          f"{res['menciones_vinculadas']} menciones vinculadas")
            self.toast("Menciones fundidas en entidades canónicas", "ok")
        self.after(0, _ui)

    def _can_poblar_tabla(self, cans):
        for row in self._tv_can.get_children():
            self._tv_can.delete(row)
        for c in cans:
            self._tv_can.insert("", "end", values=(
                c.get("id", ""), c.get("tipo", ""), c.get("nombre", ""),
                c.get("n_menciones", 0), c.get("wikidata_id") or "—",
                c.get("fuente", ""),
            ))

    def _can_generar_menciones(self):
        repo = self._can_repo()
        if repo is None:
            return
        threading.Thread(target=self._worker_can_menciones,
                         args=(repo,), daemon=True).start()

    def _worker_can_menciones(self, repo):
        import sqlite3
        try:
            con = sqlite3.connect(ST.ruta_db)
            con.row_factory = sqlite3.Row
            filas = con.execute(
                "SELECT mc.canonica_id AS cid, e.articulo_id AS art "
                "FROM menciones_canonicas mc JOIN entidades e ON e.id = mc.mencion_id"
            ).fetchall()
            con.close()
            if not filas:
                self.after(0, lambda: messagebox.showinfo(
                    "Sin menciones",
                    "Funde primero las menciones en entidades canónicas."))
                return
            n = 0
            for f in filas:
                repo.guardar_relacion(f["cid"], "mencionado_en",
                                      destino_pagina=f["art"], evidencia=f["art"],
                                      fuente="ner")
                n += 1
        except Exception as e:
            self.after(0, lambda err=e: messagebox.showerror("Error", str(err)))
            return

        def _ui():
            self._lbl_can_ok.config(text=f"✓ {n} tripletas «mencionado_en» generadas")
            self._red_log(f"[grafo] {n} tripletas mencionado_en")
            self.toast(f"{n} tripletas generadas", "ok")
        self.after(0, _ui)

    def _can_exportar_gexf(self):
        repo = self._can_repo()
        if repo is None:
            return
        graf = repo.grafo_entidades()
        if not graf["nodos"]:
            messagebox.showinfo("Grafo vacío",
                                "Funde menciones y genera tripletas primero.")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".gexf",
            filetypes=[("Graph Exchange XML", "*.gexf")],
            initialfile="grafo_canonico.gexf")
        if not dest:
            return
        try:
            self._can_escribir_gexf(graf, dest)
        except Exception as e:
            messagebox.showerror("Error", str(e))
            return
        self._red_log(f"[grafo] GEXF exportado: {dest}")
        self.toast("Grafo canónico exportado (GEXF)", "ok")

    @staticmethod
    def _can_escribir_gexf(graf, ruta):
        """Escribe el grafo canónico a GEXF (Gephi). Sin dependencias externas."""
        from xml.sax.saxutils import escape
        nodos, aristas = graf["nodos"], graf["aristas"]
        lineas = ['<?xml version="1.0" encoding="UTF-8"?>',
                  '<gexf xmlns="http://gexf.net/1.3" version="1.3">',
                  '<graph mode="static" defaultedgetype="directed">',
                  '<attributes class="node"><attribute id="0" title="tipo" type="string"/>'
                  '<attribute id="1" title="wikidata" type="string"/></attributes>',
                  '<nodes>']
        for n in nodos:
            nid = escape(str(n["id"]))
            lbl = escape(str(n.get("nombre", n["id"])))
            tipo = escape(str(n.get("tipo", "")))
            wd = escape(str(n.get("wikidata_id") or ""))
            lineas.append(f'<node id="{nid}" label="{lbl}">'
                          f'<attvalues><attvalue for="0" value="{tipo}"/>'
                          f'<attvalue for="1" value="{wd}"/></attvalues></node>')
        lineas.append('</nodes>')
        lineas.append('<edges>')
        for i, a in enumerate(aristas):
            src = escape(str(a["origen_id"]))
            dst = escape(str(a["destino_id"]))
            pred = escape(str(a.get("predicado", "")))
            w = a.get("confianza", 1.0)
            lineas.append(f'<edge id="{i}" source="{src}" target="{dst}" '
                          f'label="{pred}" weight="{w}"/>')
        lineas.append('</edges></graph></gexf>')
        with open(ruta, "w", encoding="utf-8") as f:
            f.write("\n".join(lineas))

    # ── Fase 4: export RDF ─────────────────────────────────────────────────────

    def _can_exportar_rdf(self):
        repo = self._can_repo()
        if repo is None:
            return
        graf = repo.grafo_entidades()
        if not graf["nodos"]:
            messagebox.showinfo("Grafo vacío", "Funde menciones primero.")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".ttl",
            filetypes=[("RDF Turtle", "*.ttl")], initialfile="grafo_canonico.ttl")
        if not dest:
            return
        from core.exploradores import exportar_rdf
        res = exportar_rdf(graf, dest)
        self._red_log(f"[grafo] RDF ({res['motor']}): {res['n_tripletas']} tripletas")
        self.toast(f"RDF exportado ({res['motor']})", "ok")

    # ── Fase 3: exploradores ───────────────────────────────────────────────────

    def _can_mapa_lugares(self):
        repo = self._can_repo()
        if repo is None:
            return
        from core.exploradores import geocodificar_lugares, mapa_lugares_html
        cans = repo.listar_entidades_canonicas(tipo="lugar")
        lug = geocodificar_lugares(cans)
        if not lug:
            messagebox.showinfo(
                "Sin lugares",
                "No hay entidades tipo «lugar» georreferenciables.\n"
                "Funde menciones primero (los lugares se geocodifican con el "
                "gazetteer local de Colombia).")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".html",
            filetypes=[("HTML", "*.html")], initialfile="mapa_lugares.html")
        if not dest:
            return
        res = mapa_lugares_html(lug, dest)
        self._red_log(f"[mapa] {res['n']} lugares ({res['motor']})")
        self.toast(f"Mapa de {res['n']} lugares generado", "ok")
        try:
            import webbrowser
            webbrowser.open(Path(dest).as_uri())
        except Exception:
            pass

    def _can_timeline(self):
        repo = self._can_repo()
        if repo is None:
            return
        from collections import defaultdict
        arts = repo.listar_articulos()
        if not arts:
            messagebox.showinfo("Sin artículos", "El proyecto no tiene artículos.")
            return
        por_num = defaultdict(lambda: {"n": 0, "fecha": ""})
        for a in arts:
            num = a.get("numero") or "—"
            por_num[num]["n"] += 1
            if not por_num[num]["fecha"] and a.get("fecha_publicacion"):
                por_num[num]["fecha"] = a["fecha_publicacion"]
        numeros = [{"numero": k, "fecha": v["fecha"], "n_articulos": v["n"]}
                   for k, v in por_num.items()]
        dest = filedialog.asksaveasfilename(
            defaultextension=".html",
            filetypes=[("HTML", "*.html")], initialfile="timeline_numeros.html")
        if not dest:
            return
        from core.exploradores import timeline_numeros_html
        res = timeline_numeros_html(numeros, dest)
        self._red_log(f"[timeline] {res['n']} números")
        self.toast(f"Timeline de {res['n']} números generada", "ok")
        try:
            import webbrowser
            webbrowser.open(Path(dest).as_uri())
        except Exception:
            pass

    # ── Fase 2: vocabulario controlado ─────────────────────────────────────────

    def _can_vocabulario(self):
        from core import vocabulario_controlado as vc
        ruta_db = getattr(ST, "ruta_db", "") or None
        vocab = vc.construir_vocabulario(ruta_db_proyecto=ruta_db,
                                         incluir_entidades=bool(ruta_db))
        st = vc.estadisticas(vocab)
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("JSON", "*.json")],
            initialfile="vocabulario_controlado.csv")
        if not dest:
            return
        if dest.lower().endswith(".json"):
            n = vc.exportar_json(vocab, dest)
        else:
            n = vc.exportar_csv(vocab, dest)
        cats = ", ".join(f"{k}:{v}" for k, v in st["por_categoria"].items())
        self._red_log(f"[vocab] {n} términos ({cats})")
        self.toast(f"Vocabulario controlado: {n} términos exportados", "ok")

    # ── Editor manual de relaciones ────────────────────────────────────────────

    def _can_editor_relacion(self):
        repo = self._can_repo()
        if repo is None:
            return
        cans = repo.listar_entidades_canonicas()
        if len(cans) < 2:
            messagebox.showinfo(
                "Pocas entidades",
                "Funde menciones primero: necesitas al menos 2 entidades "
                "canónicas para crear una relación.")
            return

        etiquetas = [f"{c['nombre']} [{c['tipo']}]" for c in cans]
        ids = [c["id"] for c in cans]

        win = tk.Toplevel(self)
        win.title("Añadir relación (tripleta)")
        win.configure(bg=CONTENT_BG)
        win.transient(self)
        win.grab_set()

        tk.Label(win, text="Crear una aserción sujeto — predicado — objeto",
                 bg=CONTENT_BG, fg="#FFFFFF",
                 font=("Segoe UI", 11, "bold")).grid(row=0, column=0, columnspan=2,
                                                     sticky="w", padx=12, pady=(12, 8))

        tk.Label(win, text="Sujeto:", bg=CONTENT_BG, fg=GRIS2).grid(
            row=1, column=0, sticky="e", padx=8, pady=4)
        cb_suj = ttk.Combobox(win, values=etiquetas, state="readonly", width=38)
        cb_suj.grid(row=1, column=1, sticky="w", padx=8, pady=4); cb_suj.current(0)

        tk.Label(win, text="Predicado:", bg=CONTENT_BG, fg=GRIS2).grid(
            row=2, column=0, sticky="e", padx=8, pady=4)
        cb_pred = ttk.Combobox(win, width=38, values=[
            "colaboro_con", "dirigio", "publico_en", "aliado_de", "opositor_de",
            "ubicado_en", "miembro_de", "fundador_de", "co_aparece_con"])
        cb_pred.grid(row=2, column=1, sticky="w", padx=8, pady=4)
        cb_pred.set("colaboro_con")

        tk.Label(win, text="Objeto:", bg=CONTENT_BG, fg=GRIS2).grid(
            row=3, column=0, sticky="e", padx=8, pady=4)
        cb_obj = ttk.Combobox(win, values=etiquetas, state="readonly", width=38)
        cb_obj.grid(row=3, column=1, sticky="w", padx=8, pady=4)
        cb_obj.current(1 if len(ids) > 1 else 0)

        tk.Label(win, text="Confianza:", bg=CONTENT_BG, fg=GRIS2).grid(
            row=4, column=0, sticky="e", padx=8, pady=4)
        var_conf = tk.DoubleVar(value=1.0)
        ttk.Spinbox(win, from_=0.1, to=1.0, increment=0.1, textvariable=var_conf,
                    width=6).grid(row=4, column=1, sticky="w", padx=8, pady=4)

        def _guardar():
            i_s, i_o = cb_suj.current(), cb_obj.current()
            pred = (cb_pred.get() or "").strip()
            if i_s == i_o:
                messagebox.showwarning("Inválido",
                                       "Sujeto y objeto deben ser distintos.",
                                       parent=win)
                return
            if not pred:
                messagebox.showwarning("Inválido", "Indica un predicado.",
                                       parent=win)
                return
            repo.guardar_relacion(ids[i_s], pred, destino_id=ids[i_o],
                                  confianza=var_conf.get(), fuente="revision_manual")
            self._red_log(f"[grafo] relación manual: {ids[i_s]} {pred} {ids[i_o]}")
            self.toast("Relación añadida al grafo", "ok")
            win.destroy()

        bf = tk.Frame(win, bg=CONTENT_BG)
        bf.grid(row=5, column=0, columnspan=2, pady=12)
        ttk.Button(bf, text="Guardar", style="P.TButton",
                   command=_guardar).pack(side="left", padx=6)
        ttk.Button(bf, text="Cancelar", style="S.TButton",
                   command=win.destroy).pack(side="left", padx=6)

    def _red_construir(self):
        if not getattr(ST, "indice_ner_global", None) or not any(ST.indice_ner_global.values()):
            messagebox.showwarning("Sin datos",
                "Analiza el corpus en la pestaña Índice NER primero.")
            return
        cats_sel = [c for c, v in self._var_red_cats.items() if v.get()]
        if not cats_sel:
            messagebox.showwarning("Sin categorías", "Selecciona al menos una categoría.")
            return
        peso = self._var_red_peso.get()
        self._btn_red_construir.config(state="disabled")
        self._lbl_red_ok.config(text="Construyendo red…")
        threading.Thread(target=self._worker_red_construir,
                         args=(cats_sel, peso), daemon=True).start()

    def _worker_red_construir(self, cats_sel, peso_min):
        try:
            from core.network_engine import construir_grafo, metricas_red
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Error", err))
            self.after(0, lambda: self._btn_red_construir.config(state="normal"))
            return

        def log(m):
            self.after(0, lambda msg=m: self._red_log(msg))

        try:
            G = construir_grafo(
                ST.indice_ner_global,
                categorias=cats_sel,
                peso_minimo=peso_min,
                callback=log,
            )
            self._grafo_actual = G
            met = metricas_red(G)

            def _actualizar_ui():
                # Tabla métricas
                for row in self._tv_red_met.get_children():
                    self._tv_red_met.delete(row)
                pares = [
                    ("Nodos", met.get("nodos", 0)),
                    ("Aristas", met.get("aristas", 0)),
                    ("Densidad", met.get("densidad", "—")),
                    ("Componentes conexas", met.get("componentes_conexas", "—")),
                    ("Comunidades (Louvain)", met.get("comunidades_louvain", "N/A")),
                    ("Modularidad", met.get("modularidad", "—")),
                ]
                for m, v in pares:
                    self._tv_red_met.insert("", "end", values=(m, v))

                # Tabla top nodos
                for row in self._tv_red_top.get_children():
                    self._tv_red_top.delete(row)
                top = met.get("top_centralidad", [])
                for rk, (nodo, cent) in enumerate(top, 1):
                    cat = G.nodes[nodo].get("categoria", "") if G.has_node(nodo) else ""
                    self._tv_red_top.insert("", "end",
                                             values=(rk, nodo, cat, f"{cent:.4f}"))

                n = met.get("nodos", 0)
                a = met.get("aristas", 0)
                self._lbl_red_ok.config(text=f"✅ Red construida: {n} nodos, {a} aristas")
                self._btn_red_construir.config(state="normal")

            self.after(0, _actualizar_ui)

            # Auto-generar HTML
            self.after(0, self._red_generar_html)

        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._red_log(f"⚠ Error: {err}"))
            self.after(0, lambda: self._btn_red_construir.config(state="normal"))
            self.after(0, lambda: self._lbl_red_ok.config(text=f"⚠ Error: {err}"))

    def _red_generar_html(self):
        if not self._grafo_actual:
            return
        try:
            from pathlib import Path as _PPath

            from core.network_engine import exportar_pyvis
            ruta_html = _PPath.home() / "Documents" / "BashkarStation" / "redes" / "red_entidades.html"
            ruta_html.parent.mkdir(parents=True, exist_ok=True)
            exportar_pyvis(self._grafo_actual, ruta_html)
            self._html_red_path = ruta_html
            self._red_log(f"HTML generado: {ruta_html}")
        except ImportError:
            self._red_log("pyvis no instalado. Para visualización HTML: pip install pyvis>=0.3.2")
        except Exception as e:
            self._red_log(f"HTML no generado: {e}")

    def _red_abrir_html(self):
        if not self._html_red_path or not self._html_red_path.exists():
            messagebox.showinfo("Sin HTML", "Construye la red primero. El HTML se genera automáticamente.")
            return
        import webbrowser
        webbrowser.open(str(self._html_red_path))

    def _red_exportar_gephi(self):
        if not self._grafo_actual:
            messagebox.showwarning("Sin red", "Construye la red primero.")
            return
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".gexf",
            filetypes=[("GEXF Gephi", "*.gexf"), ("Todos", "*.*")],
            initialfile="red_estampa.gexf",
            title="Exportar red para Gephi")
        if not dest:
            return
        try:
            from pathlib import Path as _PPath

            from core.network_engine import exportar_gephi
            exportar_gephi(self._grafo_actual, _PPath(dest))
            messagebox.showinfo("Exportado", f"✅ Red exportada a:\n{dest}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _red_metricas_avanzadas(self):
        """Calcula betweenness, PageRank y closeness en worker thread."""
        if not self._grafo_actual:
            messagebox.showwarning("Sin red", "Construye la red primero."); return
        self._lbl_red_ok.config(text="⏳ Calculando métricas avanzadas…")

        def _worker():
            from core.network_engine import metricas_avanzadas
            av = metricas_avanzadas(self._grafo_actual)
            self._red_metricas_av_cache = av
            # También actualiza comunidades
            coms = av.get("comunidades", [])
            self._red_comunidades_cache = coms
            self.after(0, lambda: self._red_refrescar_avanzadas())
            self.after(0, lambda: self._red_refrescar_comunidades(coms))
            n_coms = len(coms)
            self.after(0, lambda: self._lbl_red_ok.config(
                text=f"✅ Métricas avanzadas calculadas · {n_coms} comunidades"))

        threading.Thread(target=_worker, daemon=True).start()

    def _red_refrescar_avanzadas(self):
        """Puebla la tabla de centralidad avanzada según el radio seleccionado."""
        if not hasattr(self, "_tv_red_av"):
            return
        tipo = self._var_red_cent_tipo.get()
        datos = self._red_metricas_av_cache.get(tipo, [])
        import networkx as nx
        cat_map = (nx.get_node_attributes(self._grafo_actual, "categoria")
                   if self._grafo_actual else {})
        for row in self._tv_red_av.get_children():
            self._tv_red_av.delete(row)
        for i, (nodo, valor) in enumerate(datos[:30], 1):
            self._tv_red_av.insert("", "end",
                values=(i, nodo, cat_map.get(nodo, ""), f"{valor:.6f}"))

    def _red_refrescar_comunidades(self, coms: list):
        """Puebla la lista de comunidades."""
        if not hasattr(self, "_lb_red_comunidades"):
            return
        self._lb_red_comunidades.delete(0, "end")
        for com_id, miembros in coms:
            self._lb_red_comunidades.insert("end",
                f"Comunidad {com_id}  ({len(miembros)} miembros)")

    def _red_mostrar_comunidad(self, event=None):
        """Muestra los miembros de la comunidad seleccionada."""
        sel = self._lb_red_comunidades.curselection()
        if not sel or not self._red_comunidades_cache:
            return
        idx = sel[0]
        if idx >= len(self._red_comunidades_cache):
            return
        com_id, miembros = self._red_comunidades_cache[idx]
        import networkx as nx
        G = self._grafo_actual
        cat_map = nx.get_node_attributes(G, "categoria") if G else {}
        grado_map = dict(G.degree()) if G else {}
        for row in self._tv_red_miembros.get_children():
            self._tv_red_miembros.delete(row)
        for nodo in sorted(miembros, key=lambda n: -grado_map.get(n, 0)):
            self._tv_red_miembros.insert("", "end",
                values=(nodo, cat_map.get(nodo, ""), grado_map.get(nodo, 0)))

    def _red_evolucion_temporal(self):
        """Abre pestaña de evolución temporal — la construcción es bajo demanda."""
        messagebox.showinfo(
            "Evolución temporal",
            "Usa el botón '▶ Calcular evolución' en la pestaña 'Evolución temporal'.\n\n"
            "Requiere que el corpus_meta tenga varios números procesados.")

    def _red_calcular_evolucion(self):
        """Calcula la evolución temporal de la red número por número."""
        # Esta función espera corpus_meta como dict {art_id: {...}} (flujo del
        # conversor); si es el DataFrame que arma _worker_ocr, se trata como
        # "no disponible" en vez de intentar iterarlo mal (.items() de un
        # DataFrame recorre columnas, no artículos).
        corpus_meta = getattr(ST, "corpus_meta", None)
        if not isinstance(corpus_meta, dict):
            corpus_meta = {}
        indice = getattr(ST, "indice_ner_global", {}) or {}

        if not corpus_meta and not indice:
            messagebox.showwarning("Sin corpus",
                "Procesa al menos dos números del corpus primero."); return

        # Construir índice_por_numero desde corpus_meta + indice NER
        # Si hay corpus_meta, usamos los art_ids para separar por número
        indice_por_numero: dict = {}
        if corpus_meta:
            for art_id, meta in corpus_meta.items():
                numero = (meta.get("numero") or meta.get("pagina", "")[:7]
                          if isinstance(meta, dict) else "sin_numero")
                if numero not in indice_por_numero:
                    indice_por_numero[numero] = {cat: {} for cat in indice}
                for cat, ents in indice.items():
                    if not isinstance(ents, dict):
                        continue
                    for ent, arts in ents.items():
                        if art_id in arts:
                            indice_por_numero[numero][cat].setdefault(ent, []).append(art_id)

        if len(indice_por_numero) < 2:
            messagebox.showwarning(
                "Pocos números",
                f"Solo se detectó {len(indice_por_numero)} número(s).\n"
                "La evolución temporal requiere al menos 2 números procesados."); return

        self._lbl_red_ok.config(text="⏳ Calculando evolución temporal…")

        def _worker():
            from core.network_engine import evolucion_temporal
            cats = [c for c, v in self._var_red_cats.items() if v.get()]
            peso  = self._var_red_peso.get()

            def _cb(n, total, numero):
                self.after(0, lambda: self._lbl_red_ok.config(
                    text=f"⏳ {n}/{total}: {numero}"))

            serie = evolucion_temporal(indice_por_numero, categorias=cats,
                                       peso_minimo=peso, callback=_cb)
            self._red_evolucion_cache = serie

            def _mostrar():
                for row in self._tv_red_evo.get_children():
                    self._tv_red_evo.delete(row)
                for r in serie:
                    self._tv_red_evo.insert("", "end", values=(
                        r["numero"], r["nodos"], r["aristas"],
                        f"{r['densidad']:.4f}", r["top_nodo"]))
                self._lbl_red_ok.config(
                    text=f"✅ Evolución calculada: {len(serie)} números")
            self.after(0, _mostrar)

        threading.Thread(target=_worker, daemon=True).start()

    def _red_graficar_evolucion(self):
        """Grafica la evolución temporal de nodos/aristas/densidad."""
        if not self._red_evolucion_cache:
            messagebox.showwarning("Sin datos", "Calcula la evolución primero."); return
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

            from core.chart_builder import _FONDO, _TEXTO
        except ImportError:
            messagebox.showerror("Falta matplotlib", "pip install matplotlib"); return

        serie = self._red_evolucion_cache
        numeros  = [r["numero"] for r in serie]
        nodos    = [r["nodos"]   for r in serie]
        aristas  = [r["aristas"] for r in serie]
        densidad = [r["densidad"] for r in serie]

        win = tk.Toplevel(self)
        win.title("Evolución temporal de la red")
        win.geometry("800x520")
        win.configure(bg=CONTENT_BG)

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6),
                                        facecolor=_FONDO, sharex=True)
        x = range(len(numeros))

        ax1.plot(x, nodos,   "o-", color="#58A6FF", label="Nodos",   linewidth=2)
        ax1.plot(x, aristas, "s-", color="#F59E0B", label="Aristas",  linewidth=2)
        ax1.set_ylabel("Cantidad", color=_TEXTO); ax1.legend(fontsize=8)
        ax1.set_facecolor(_FONDO); ax1.tick_params(colors=_TEXTO)
        for spine in ax1.spines.values(): spine.set_edgecolor("#30363D")

        ax2.plot(x, densidad, "^-", color="#3FB950", linewidth=2)
        ax2.set_ylabel("Densidad", color=_TEXTO)
        ax2.set_xticks(list(x))
        ax2.set_xticklabels(numeros, rotation=30, ha="right", fontsize=7, color=_TEXTO)
        ax2.set_facecolor(_FONDO); ax2.tick_params(colors=_TEXTO)
        for spine in ax2.spines.values(): spine.set_edgecolor("#30363D")

        plt.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)

    def _red_evolucion_csv(self):
        """Exporta la tabla de evolución temporal a CSV."""
        if not self._red_evolucion_cache:
            messagebox.showwarning("Sin datos", "Calcula la evolución primero."); return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile="evolucion_red.csv")
        if not dest: return
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=["numero","nodos","aristas",
                                               "densidad","top_nodo","top_cent"])
            w.writeheader(); w.writerows(self._red_evolucion_cache)
        messagebox.showinfo("Exportado", f"✅ Evolución exportada:\n{dest}")

    def _red_exportar_csv(self):
        """Exporta métricas por nodo (grado, betweenness, PageRank, closeness) a CSV."""
        if not self._grafo_actual:
            messagebox.showwarning("Sin red", "Construye la red primero."); return
        if not self._red_metricas_av_cache:
            messagebox.showwarning("Sin métricas avanzadas",
                "Calcula primero las métricas avanzadas con el botón '🔬 Métricas avanzadas'.")
            return
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV","*.csv")],
            initialfile="metricas_red.csv")
        if not dest: return
        try:
            from core.network_engine import exportar_metricas_csv
            exportar_metricas_csv(self._grafo_actual,
                                   self._red_metricas_av_cache, Path(dest))
            messagebox.showinfo("Exportado", f"✅ Métricas exportadas:\n{dest}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: SEMÁNTICO (v13) — tono editorial + léxico + estilometría
    # ══════════════════════════════════════════════════════════════════════════

    def _build_sem(self):
        outer = tk.Frame(self._tab_sem, bg=CONTENT_BG)
        outer.pack(fill="both", expand=True, padx=16, pady=12)

        self._sem_params: dict = {}
        try:
            from core.sentiment_engine import PARAMS_SCHEMA as _SEM_SCHEMA
            self._build_params_panel(outer, _SEM_SCHEMA, self._sem_params)
        except Exception:
            pass

        pad = tk.Frame(outer, bg=CONTENT_BG)
        pad.pack(side="left", fill="both", expand=True)

        tk.Label(pad, text="Análisis semántico profundo", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad, text="Tono editorial, léxico histórico y estilometría del corpus Estampa.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 10))

        # ── Pestañas internas ─────────────────────────────────────────────────
        nb = ttk.Notebook(pad)
        nb.pack(fill="both", expand=True)

        # Tab Tono
        frm_tono = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_tono, text="  Tono editorial  ")
        self._build_sem_tono(frm_tono)

        # Tab Léxico
        frm_lex = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_lex, text="  Léxico histórico  ")
        self._build_sem_lexico(frm_lex)

        # Tab Estilometría
        frm_estilo = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_estilo, text="  Estilometría  ")
        self._build_sem_estilo(frm_estilo)

    # ── Sub-panel: Tono editorial ─────────────────────────────────────────────
    def _build_sem_tono(self, parent):
        from core.sentiment_engine import COLORES_TONO
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        # ── Botones principales ──
        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 4))
        self._btn_tono_art = ttk.Button(bf, text="▶  Artículo actual",
                                         style="P.TButton",
                                         command=self._sem_tono_articulo)
        self._btn_tono_art.pack(side="left", padx=(0, 8))
        self._btn_tono_corpus = ttk.Button(bf, text="📚  Corpus completo",
                                            style="S.TButton",
                                            command=self._sem_tono_corpus)
        self._btn_tono_corpus.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📊  Ver evolución",
                   style="S.TButton",
                   command=self._sem_tono_ver_evolucion).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="📝  Síntesis narrativa",
                   style="S.TButton",
                   command=self._sem_tono_narrativa).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="💾  Exportar CSV",
                   style="S.TButton",
                   command=self._sem_tono_exportar).pack(side="right")

        self._lbl_tono_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI", 9, "bold"))
        self._lbl_tono_ok.pack(anchor="w", pady=(0, 4))

        # ── Chips de distribución (se actualizan al terminar el análisis) ──
        self._frm_tono_chips = tk.Frame(pad, bg=CONTENT_BG)
        self._frm_tono_chips.pack(fill="x", pady=(0, 6))
        self._tono_chips = {}
        tonos_orden = ("celebratorio", "crítico", "neutro", "elegíaco", "polémico", "informativo")
        for tono in tonos_orden:
            color = COLORES_TONO.get(tono, "#6B7280")
            frm = tk.Frame(self._frm_tono_chips, bg=color, padx=6, pady=2)
            frm.pack(side="left", padx=(0, 4))
            lbl = tk.Label(frm, text=f"{tono}: —", bg=color, fg="white",
                           font=("Segoe UI", 8, "bold"))
            lbl.pack()
            self._tono_chips[tono] = lbl

        # ── Filtro por campo ──
        ff = tk.Frame(pad, bg=CONTENT_BG)
        ff.pack(fill="x", pady=(0, 4))
        tk.Label(ff, text="Filtrar por:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_tono_filtro_campo = tk.StringVar(value="todos")
        self._var_tono_filtro_valor = tk.StringVar(value="")
        ttk.Combobox(ff, textvariable=self._var_tono_filtro_campo,
                     values=["todos", "tono_principal", "intensidad", "numero"],
                     state="readonly", width=14).pack(side="left", padx=(0, 4))
        ttk.Entry(ff, textvariable=self._var_tono_filtro_valor,
                  width=16).pack(side="left", padx=(0, 8))
        ttk.Button(ff, text="Filtrar", style="S.TButton",
                   command=self._sem_tono_refrescar).pack(side="left")

        # ── Tabla resultados ──
        cols = ("articulo", "tono_principal", "tono_sec", "intensidad",
                "confianza", "numero", "resumen")
        self._tv_tono = ttk.Treeview(pad, columns=cols, show="headings", height=11)
        heads = [("articulo",      "Artículo",       130),
                 ("tono_principal","Tono principal",  110),
                 ("tono_sec",      "Secundario",       90),
                 ("intensidad",    "Intensidad",       75),
                 ("confianza",     "Confianza",        70),
                 ("numero",        "Número",           80),
                 ("resumen",       "Resumen",         330)]
        for cid, txt, w in heads:
            self._tv_tono.heading(cid, text=txt)
            self._tv_tono.column(cid, width=w, anchor="w")
        sv = ttk.Scrollbar(pad, orient="vertical", command=self._tv_tono.yview)
        self._tv_tono.configure(yscrollcommand=sv.set)
        self._tv_tono.pack(side="left", fill="both", expand=True)
        sv.pack(side="left", fill="y")

        # Tags de color por tono
        for tono, color in COLORES_TONO.items():
            self._tv_tono.tag_configure(tono, foreground=color)

        self._tono_resultados = {}

    # ── Sub-panel: Léxico histórico ───────────────────────────────────────────
    def _build_sem_lexico(self, parent):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 6))
        self._btn_lex_art = ttk.Button(bf, text="▶  Artículo actual",
                                        style="P.TButton",
                                        command=self._sem_lex_articulo)
        self._btn_lex_art.pack(side="left", padx=(0, 8))
        self._btn_lex_corpus = ttk.Button(bf, text="📚  Corpus completo",
                                           style="S.TButton",
                                           command=self._sem_lex_corpus)
        self._btn_lex_corpus.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="💾  Exportar glosario",
                   style="S.TButton",
                   command=self._sem_lex_exportar).pack(side="right")
        self._lbl_lex_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_lex_ok.pack(anchor="w", pady=(0, 4))

        # Filtro categoría
        fi = tk.Frame(pad, bg=CONTENT_BG)
        fi.pack(fill="x", pady=(0, 4))
        tk.Label(fi, text="Categoría:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_lex_cat = tk.StringVar(value="Todas")
        cats = ["Todas", "arcaismos", "neologismos", "colombianismos", "tecnicismos"]
        ttk.Combobox(fi, textvariable=self._var_lex_cat, values=cats,
                     state="readonly", width=14).pack(side="left")

        # Tabla
        cols = ("categoria", "palabra", "n_arts", "info")
        self._tv_lex = ttk.Treeview(pad, columns=cols, show="headings", height=12)
        heads = [("categoria", "Categoría", 110), ("palabra", "Palabra", 140),
                 ("n_arts", "Artículos", 70), ("info", "Información", 380)]
        for cid, txt, w in heads:
            self._tv_lex.heading(cid, text=txt)
            self._tv_lex.column(cid, width=w, anchor="w")
        sv = ttk.Scrollbar(pad, orient="vertical", command=self._tv_lex.yview)
        self._tv_lex.configure(yscrollcommand=sv.set)
        self._tv_lex.pack(side="left", fill="both", expand=True)
        sv.pack(side="left", fill="y")

        self._glosario_data = {}

    # ── Sub-panel: Estilometría ───────────────────────────────────────────────
    def _build_sem_estilo(self, parent):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        tk.Label(pad, text="Agrupa artículos anónimos por similitud estilística (TF-IDF n-gramas de caracteres).",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 6))

        cf = tk.Frame(pad, bg=CONTENT_BG)
        cf.pack(fill="x", pady=(0, 6))
        tk.Label(cf, text="N° clusters:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_estilo_clusters = tk.IntVar(value=5)
        ttk.Spinbox(cf, from_=2, to=15, textvariable=self._var_estilo_clusters,
                    width=4).pack(side="left")

        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 6))
        self._btn_estilo = ttk.Button(bf, text="▶  Calcular clusters",
                                       style="P.TButton",
                                       command=self._sem_estilo_calcular)
        self._btn_estilo.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="💾  Exportar CSV",
                   style="S.TButton",
                   command=self._sem_estilo_exportar).pack(side="left")
        self._lbl_estilo_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                        font=("Segoe UI", 9, "bold"))
        self._lbl_estilo_ok.pack(anchor="w", pady=(0, 4))

        cols = ("articulo", "cluster")
        self._tv_estilo = ttk.Treeview(pad, columns=cols, show="headings", height=14)
        self._tv_estilo.heading("articulo", text="Artículo")
        self._tv_estilo.heading("cluster",  text="Cluster")
        self._tv_estilo.column("articulo", width=300, anchor="w")
        self._tv_estilo.column("cluster",  width=80,  anchor="center")
        sv = ttk.Scrollbar(pad, orient="vertical", command=self._tv_estilo.yview)
        self._tv_estilo.configure(yscrollcommand=sv.set)
        self._tv_estilo.pack(side="left", fill="both", expand=True)
        sv.pack(side="left", fill="y")

        self._estilo_resultados = {}

    # ── Workers: Tono ─────────────────────────────────────────────────────────
    def _sem_tono_articulo(self):
        api_key, _m = _resolver_api_key_modelo("tono")
        if not api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API en Configuración.")
            return
        texto, art_id = self._ner_articulo_actual()
        if not texto:
            return
        self._btn_tono_art.config(state="disabled")
        threading.Thread(target=self._worker_tono, args=({art_id: texto},),
                         daemon=True).start()

    def _sem_tono_corpus(self):
        api_key, _m = _resolver_api_key_modelo("tono")
        if not api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API en Configuración.")
            return
        if not getattr(ST, "corpus_txt", None):
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        self._btn_tono_corpus.config(state="disabled")
        # Enriquecer entradas con metadatos disponibles en ST
        articulos = ST.articulos if getattr(ST, "articulos", None) else []
        entradas = {}
        for i, texto in enumerate(ST.corpus_txt):
            if not texto or not texto.strip():
                continue
            art_id = str(i)
            entrada = {"texto": texto}
            if i < len(articulos):
                art = articulos[i]
                entrada["seccion"]  = art.get("tipo", "")
                entrada["numero"]   = art.get("numero", "")
                entrada["autor"]    = art.get("autor", "")
            entradas[art_id] = entrada
        threading.Thread(target=self._worker_tono, args=(entradas,),
                         daemon=True).start()

    def _worker_tono(self, entradas):
        from core.sentiment_engine import analizar_corpus_tono, estadisticas_tono
        p = self._params_get_values(self._sem_params) if getattr(self, "_sem_params", None) else {}
        motor   = p.get("motor", "lexicon")
        workers = int(p.get("workers", 4))
        tonos_activos = p.get("tonos_activos") or None
        api_key = _resolver_api_key_modelo("tono")[0] if motor == "ia" else None
        total = len(entradas)

        def cb(n, t, art_id):
            self.after(0, lambda: self._lbl_tono_ok.config(
                text=f"Analizando {n}/{t}: {art_id}"))

        nuevos = analizar_corpus_tono(entradas, api_key, callback=cb, workers=workers)
        self._tono_resultados.update(nuevos)

        stats = estadisticas_tono(self._tono_resultados)
        self.after(0, lambda s=stats: self._sem_tono_actualizar_chips(s))
        self.after(0, self._sem_tono_refrescar)
        self.after(0, lambda: self._btn_tono_art.config(state="normal"))
        self.after(0, lambda: self._btn_tono_corpus.config(state="normal"))
        n = len(self._tono_resultados)
        self.after(0, lambda: self._lbl_tono_ok.config(
            text=f"✅ {n} artículos — dominante: {stats.get('tono_dominante','?')} · "
                 f"polarización: {stats.get('indice_polarizacion', 0)}%"))

    def _sem_tono_actualizar_chips(self, stats):
        dist = stats.get("distribucion", {})
        for tono, lbl in self._tono_chips.items():
            pct = dist.get(tono, {}).get("porcentaje", 0.0)
            n   = dist.get(tono, {}).get("n", 0)
            lbl.config(text=f"{tono}: {pct}% ({n})")

    def _sem_tono_refrescar(self):
        for row in self._tv_tono.get_children():
            self._tv_tono.delete(row)

        campo  = self._var_tono_filtro_campo.get()
        valor  = self._var_tono_filtro_valor.get().strip().lower()

        for art_id, res in self._tono_resultados.items():
            tono = res.get("tono_principal", "neutro")
            # Filtro
            if campo != "todos" and valor:
                val_campo = str(res.get(campo, "")).lower()
                if valor not in val_campo:
                    continue
            self._tv_tono.insert("", "end",
                tags=(tono,),
                values=(
                    art_id,
                    tono,
                    res.get("tono_secundario", "") or "",
                    res.get("intensidad", ""),
                    f"{res.get('confianza', 0):.2f}",
                    res.get("numero", ""),
                    res.get("resumen", ""),
                ))

    def _sem_tono_ver_evolucion(self):
        from core.sentiment_engine import (
            COLORES_TONO,
            evolucion_temporal,
            tendencia_tono,
        )
        if not self._tono_resultados:
            messagebox.showwarning("Sin datos", "Analiza el tono del corpus primero.")
            return

        evol = evolucion_temporal(self._tono_resultados, campo_periodo="numero")
        if len(evol) < 2:
            messagebox.showinfo("Evolución temporal",
                "Se necesitan al menos 2 números para ver la evolución temporal.\n"
                "Analiza el corpus completo primero.")
            return

        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        except ImportError:
            messagebox.showerror("Falta matplotlib",
                "Instala matplotlib para ver el gráfico.")
            return

        win = tk.Toplevel(self)
        win.title("Evolución temporal del tono editorial")
        win.geometry("860x520")
        win.configure(bg=CONTENT_BG)

        periodos = sorted(evol.keys())
        fig, ax = plt.subplots(figsize=(9, 4.5))
        fig.patch.set_facecolor("#1E1E2E")
        ax.set_facecolor("#1E1E2E")

        tonos_a_mostrar = ("celebratorio", "crítico", "elegíaco", "polémico")
        for tono in tonos_a_mostrar:
            vals = [evol[p].get(tono, 0.0) for p in periodos]
            t_info = tendencia_tono(evol, tono)
            dir_arrow = {"sube": " ↑", "baja": " ↓", "estable": ""}.get(
                t_info["direccion"], "")
            ax.plot(periodos, vals,
                    marker="o", linewidth=2,
                    color=COLORES_TONO.get(tono, "#888"),
                    label=f"{tono}{dir_arrow}")

        ax.set_xlabel("Número", color="#CDD6F4", fontsize=9)
        ax.set_ylabel("% artículos", color="#CDD6F4", fontsize=9)
        ax.set_title("Evolución del tono editorial por número", color="#CDD6F4", fontsize=10)
        ax.tick_params(colors="#CDD6F4", labelsize=8)
        ax.legend(facecolor="#313244", labelcolor="#CDD6F4", fontsize=8)
        for spine in ax.spines.values():
            spine.set_edgecolor("#45475A")
        plt.xticks(rotation=30, ha="right")
        plt.tight_layout()

        canvas = FigureCanvasTkAgg(fig, master=win)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True, padx=10, pady=10)

        # Tabla de tendencias
        tf = tk.Frame(win, bg=CONTENT_BG)
        tf.pack(fill="x", padx=10, pady=(0, 8))
        for tono in tonos_a_mostrar:
            t_info = tendencia_tono(evol, tono)
            color = COLORES_TONO.get(tono, "#888")
            icono = {"sube": "↑", "baja": "↓", "estable": "→"}.get(t_info["direccion"], "")
            tk.Label(tf, text=f"{icono} {tono}  (pend. {t_info['pendiente']:+.2f})",
                     bg=CONTENT_BG, fg=color,
                     font=("Segoe UI", 9, "bold")).pack(side="left", padx=8)

    def _sem_tono_narrativa(self):
        from core.sentiment_engine import (
            estadisticas_tono,
            evolucion_temporal,
            resumen_narrativo,
        )
        if not self._tono_resultados:
            messagebox.showwarning("Sin datos", "Analiza el tono del corpus primero.")
            return
        api_key, _m = _resolver_api_key_modelo("tono")
        if not api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API.")
            return

        win = tk.Toplevel(self)
        win.title("Síntesis narrativa del tono editorial")
        win.geometry("680x360")
        win.configure(bg=CONTENT_BG)

        lbl = tk.Label(win, text="Generando síntesis…", bg=CONTENT_BG, fg=GRIS2,
                       font=("Segoe UI", 9))
        lbl.pack(anchor="w", padx=16, pady=(12, 4))
        txt = tk.Text(win, bg="#1E1E2E", fg="#CDD6F4", font=("Segoe UI", 10),
                      wrap="word", padx=12, pady=10, relief="flat")
        txt.pack(fill="both", expand=True, padx=12, pady=(0, 12))

        def _generar():
            stats = estadisticas_tono(self._tono_resultados)
            evol  = evolucion_temporal(self._tono_resultados, "numero")
            nombre = getattr(ST, "nombre_proyecto", "el corpus")
            parrafo = resumen_narrativo(stats, evol, api_key, nombre_corpus=nombre)
            self.after(0, lambda: lbl.config(text="Síntesis generada"))
            self.after(0, lambda: txt.insert("1.0", parrafo or "No se pudo generar la síntesis."))

        threading.Thread(target=_generar, daemon=True).start()

    def _sem_tono_exportar(self):
        if not self._tono_resultados:
            messagebox.showwarning("Sin datos", "Analiza el tono del corpus primero.")
            return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("Todos", "*.*")],
            initialfile="tono_editorial.csv",
            title="Guardar análisis de tono")
        if not dest:
            return
        fieldnames = ["articulo", "tono_principal", "tono_secundario", "intensidad",
                      "confianza", "numero", "seccion", "autor", "resumen", "indicadores"]
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
            w.writeheader()
            for art_id, res in self._tono_resultados.items():
                inds = res.get("indicadores", [])
                w.writerow({
                    "articulo":        art_id,
                    "tono_principal":  res.get("tono_principal", ""),
                    "tono_secundario": res.get("tono_secundario", "") or "",
                    "intensidad":      res.get("intensidad", ""),
                    "confianza":       res.get("confianza", 0),
                    "numero":          res.get("numero", ""),
                    "seccion":         res.get("seccion", ""),
                    "autor":           res.get("autor", ""),
                    "resumen":         res.get("resumen", ""),
                    "indicadores":     "; ".join(inds) if isinstance(inds, list) else str(inds),
                })
        messagebox.showinfo("Exportado", f"Tono exportado a:\n{dest}")

    # ── Workers: Léxico ───────────────────────────────────────────────────────
    def _sem_lex_articulo(self):
        api_key, _m = _resolver_api_key_modelo("tono")
        if not api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API en Configuración.")
            return
        texto, art_id = self._ner_articulo_actual()
        if not texto:
            return
        self._btn_lex_art.config(state="disabled")
        threading.Thread(target=self._worker_lexico, args=({art_id: texto},),
                         daemon=True).start()

    def _sem_lex_corpus(self):
        api_key, _m = _resolver_api_key_modelo("tono")
        if not api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API en Configuración.")
            return
        if not getattr(ST, "corpus_txt", None):
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        self._btn_lex_corpus.config(state="disabled")
        textos = {str(i): t for i, t in enumerate(ST.corpus_txt) if t and t.strip()}
        threading.Thread(target=self._worker_lexico, args=(textos,), daemon=True).start()

    def _worker_lexico(self, textos):
        from core.lexicon_engine import construir_glosario
        api_key, _m = _resolver_api_key_modelo("tono")
        total = len(textos)

        def cb(i, t, aid):
            self.after(0, lambda n=i, tot=t, a=aid:
                       self._lbl_lex_ok.config(text=f"Procesando {n}/{tot}: {a}"))

        self._glosario_data = construir_glosario(textos, api_key, callback=cb)
        self.after(0, self._sem_lex_refrescar)
        self.after(0, lambda: self._btn_lex_art.config(state="normal"))
        self.after(0, lambda: self._btn_lex_corpus.config(state="normal"))
        n = sum(len(v) for v in self._glosario_data.values())
        self.after(0, lambda: self._lbl_lex_ok.config(text=f"✅ {n} entradas en glosario"))

    def _sem_lex_refrescar(self):
        for row in self._tv_lex.get_children():
            self._tv_lex.delete(row)
        cat_filtro = self._var_lex_cat.get()
        for cat, entradas in self._glosario_data.items():
            if cat_filtro not in ("Todas", cat):
                continue
            for palabra, info in sorted(entradas.items()):
                extra = info.get("definicion") or info.get("significado") or info.get("origen") or ""
                n_arts = len(info.get("articulos", []))
                self._tv_lex.insert("", "end", values=(cat, palabra, n_arts, extra))

    def _sem_lex_exportar(self):
        if not self._glosario_data:
            messagebox.showwarning("Sin datos", "Construye el glosario primero.")
            return
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv"), ("Todos", "*.*")],
            initialfile="glosario_estampa.csv",
            title="Guardar glosario")
        if not dest:
            return
        from pathlib import Path as _PPath

        from core.lexicon_engine import exportar_glosario_csv
        n = exportar_glosario_csv(self._glosario_data, _PPath(dest))
        messagebox.showinfo("Exportado", f"✅ {n} entradas exportadas a:\n{dest}")

    # ── Workers: Estilometría ─────────────────────────────────────────────────
    def _sem_estilo_calcular(self):
        if not getattr(ST, "corpus_txt", None):
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        n_clusters = self._var_estilo_clusters.get()
        textos = {str(i): t for i, t in enumerate(ST.corpus_txt) if t and t.strip()}
        if len(textos) < 2:
            messagebox.showwarning("Datos insuficientes", "Se necesitan al menos 2 artículos.")
            return
        self._btn_estilo.config(state="disabled")
        threading.Thread(target=self._worker_estilo, args=(textos, n_clusters),
                         daemon=True).start()

    def _worker_estilo(self, textos, n_clusters):
        from core.stylometry_engine import cluster_tematico
        self.after(0, lambda: self._lbl_estilo_ok.config(text="Calculando clusters…"))
        try:
            result = cluster_tematico(textos, n_clusters=n_clusters)
            self._estilo_resultados = result
            self.after(0, self._sem_estilo_refrescar)
            n = len(result)
            self.after(0, lambda: self._lbl_estilo_ok.config(text=f"✅ {n} artículos agrupados"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_estilo_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_estilo.config(state="normal"))

    def _sem_estilo_refrescar(self):
        for row in self._tv_estilo.get_children():
            self._tv_estilo.delete(row)
        for art_id, cluster in sorted(self._estilo_resultados.items(),
                                       key=lambda x: (x[1], x[0])):
            self._tv_estilo.insert("", "end", values=(art_id, cluster))

    def _sem_estilo_exportar(self):
        if not self._estilo_resultados:
            messagebox.showwarning("Sin datos", "Calcula los clusters primero.")
            return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="estilometria_clusters.csv",
            title="Exportar estilometría")
        if not dest:
            return
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=["articulo", "cluster"])
            w.writeheader()
            for art_id, cluster in self._estilo_resultados.items():
                w.writerow({"articulo": art_id, "cluster": cluster})
        messagebox.showinfo("Exportado", f"Clusters exportados a:\n{dest}")


    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: VISUALIZAR (v14) — nubes, heatmap, mapa, timeline
    # ══════════════════════════════════════════════════════════════════════════

    def _build_viz(self):
        from core.chart_builder import CATALOGO
        pad = tk.Frame(self._tab_viz, bg=CONTENT_BG, padx=16, pady=12)
        pad.pack(fill="both", expand=True)

        tk.Label(pad, text="Constructor de visualizaciones", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad,
                 text="Selecciona qué dato graficar y con qué tipo de gráfico. "
                      "Cada opción incluye una descripción de cuándo usarla.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 10))

        nb = ttk.Notebook(pad)
        nb.pack(fill="both", expand=True)

        # ── Pestaña: Constructor interactivo ──────────────────────────────────
        frm_build = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_build, text="  Constructor  ")
        self._build_viz_constructor(frm_build, CATALOGO)

        # ── Pestañas legacy (se mantienen para compatibilidad) ────────────────
        frm_nube = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_nube, text="  Nube de palabras  ")
        self._build_viz_nube(frm_nube)

        frm_heat = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_heat, text="  Heatmap términos  ")
        self._build_viz_heatmap(frm_heat)

        frm_mapa = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_mapa, text="  Mapa  ")
        self._build_viz_mapa(frm_mapa)

        frm_tl = tk.Frame(nb, bg=CONTENT_BG)
        nb.add(frm_tl, text="  Timeline  ")
        self._build_viz_timeline(frm_tl)

    # ── Constructor interactivo de gráficos ───────────────────────────────────
    def _build_viz_constructor(self, parent, catalogo):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        # ── Fila de selectores ──
        sel_frm = tk.Frame(pad, bg=CONTENT_BG)
        sel_frm.pack(fill="x", pady=(0, 6))

        tk.Label(sel_frm, text="Fuente de datos:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).grid(row=0, column=0, sticky="w", padx=(0, 6))
        self._var_viz_fuente = tk.StringVar()
        fuentes = list(catalogo.keys())
        cb_fuente = ttk.Combobox(sel_frm, textvariable=self._var_viz_fuente,
                                  values=fuentes, state="readonly", width=20)
        cb_fuente.grid(row=0, column=1, padx=(0, 16))
        cb_fuente.current(0)

        tk.Label(sel_frm, text="Tipo de gráfico:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).grid(row=0, column=2, sticky="w", padx=(0, 6))
        self._var_viz_tipo = tk.StringVar()
        self._cb_viz_tipo = ttk.Combobox(sel_frm, textvariable=self._var_viz_tipo,
                                          values=[], state="readonly", width=28)
        self._cb_viz_tipo.grid(row=0, column=3, padx=(0, 16))

        tk.Label(sel_frm, text="Título (opcional):", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).grid(row=0, column=4, sticky="w", padx=(0, 6))
        self._var_viz_titulo = tk.StringVar()
        tk.Entry(sel_frm, textvariable=self._var_viz_titulo,
                 width=22, font=("Segoe UI", 9),
                 relief="solid", bd=1, bg="#0D1B2A", fg="#CDD6F4"
                 ).grid(row=0, column=5, padx=(0, 8))

        self._btn_viz_gen = ttk.Button(sel_frm, text="▶  Generar",
                                        style="P.TButton",
                                        command=self._viz_generar)
        self._btn_viz_gen.grid(row=0, column=6, padx=(0, 8))
        ttk.Button(sel_frm, text="💾  Guardar PNG",
                   style="S.TButton",
                   command=self._viz_guardar).grid(row=0, column=7)

        # ── Descripción del gráfico seleccionado ──
        self._lbl_viz_desc = tk.Label(pad, text="",
                                       bg=CONTENT_BG, fg="#94A3B8",
                                       font=("Segoe UI", 8), wraplength=900,
                                       justify="left")
        self._lbl_viz_desc.pack(anchor="w", pady=(0, 4))

        self._lbl_viz_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_viz_ok.pack(anchor="w", pady=(0, 4))

        # ── Canvas para el gráfico ──
        self._frm_viz_canvas = tk.Frame(pad, bg=CONTENT_BG)
        self._frm_viz_canvas.pack(fill="both", expand=True)
        self._viz_canvas_widget = None
        self._viz_fig_actual    = None

        # ── Conectar eventos ──
        self._viz_catalogo = catalogo
        cb_fuente.bind("<<ComboboxSelected>>", self._viz_actualizar_tipos)
        self._cb_viz_tipo.bind("<<ComboboxSelected>>", self._viz_actualizar_desc)
        self._viz_actualizar_tipos()

    def _viz_actualizar_tipos(self, _event=None):
        fuente = self._var_viz_fuente.get()
        opciones = self._viz_catalogo.get(fuente, [])
        labels = [o["label"] for o in opciones]
        self._cb_viz_tipo["values"] = labels
        if labels:
            self._cb_viz_tipo.current(0)
        self._viz_actualizar_desc()

    def _viz_actualizar_desc(self, _event=None):
        fuente = self._var_viz_fuente.get()
        tipo   = self._var_viz_tipo.get()
        opciones = self._viz_catalogo.get(fuente, [])
        for op in opciones:
            if op["label"] == tipo:
                self._lbl_viz_desc.config(text=f"ℹ  {op['desc']}")
                return
        self._lbl_viz_desc.config(text="")

    def _viz_generar(self):
        fuente = self._var_viz_fuente.get()
        tipo   = self._var_viz_tipo.get()
        titulo = self._var_viz_titulo.get().strip()
        opciones = self._viz_catalogo.get(fuente, [])
        op = next((o for o in opciones if o["label"] == tipo), None)
        if not op:
            return

        self._btn_viz_gen.config(state="disabled")
        self._lbl_viz_ok.config(text="Generando…")
        threading.Thread(target=self._viz_worker, args=(op, titulo),
                         daemon=True).start()

    def _viz_worker(self, op, titulo):
        try:
            fig = self._viz_obtener_datos_y_graficar(op, titulo)
            self.after(0, lambda: self._viz_mostrar(fig))
            self.after(0, lambda: self._lbl_viz_ok.config(text="✅ Gráfico generado"))
        except Exception as e:
            self.after(0, lambda err=str(e): self._lbl_viz_ok.config(
                text=f"⚠ Error: {err}"))
        finally:
            self.after(0, lambda: self._btn_viz_gen.config(state="normal"))

    def _viz_obtener_datos_y_graficar(self, op, titulo):
        param = op["param"]
        fn    = op["fn"]

        if param == "resultados":
            datos = getattr(self, "_tono_resultados", {})
            if not datos:
                raise ValueError("Analiza el tono del corpus primero.")
            return fn(datos, titulo=titulo)

        elif param == "confianza":
            datos = getattr(ST, "confianza_corpus", {})
            if not datos:
                raise ValueError("No hay datos de confianza OCR. Extrae el corpus primero.")
            return fn(datos, titulo=titulo)

        elif param == "ner":
            datos = getattr(ST, "indice_ner_global", {})
            if not datos:
                raise ValueError("Ejecuta el análisis NER primero.")
            # Para frecuencia necesita categoría — usa la primera disponible
            cat = next(iter(datos.keys()), "PER")
            import inspect
            sig = inspect.signature(fn)
            if "categoria" in sig.parameters:
                return fn(datos, categoria=cat, titulo=titulo)
            return fn(datos, titulo=titulo)

        elif param == "articulos":
            datos = getattr(ST, "articulos", None) or []
            if not datos:
                raise ValueError("Segmenta los artículos del corpus primero.")
            return fn(datos, titulo=titulo)

        elif param == "corpus_txt":
            datos = getattr(ST, "corpus_txt", None) or []
            if not datos:
                raise ValueError("Extrae el texto del corpus primero.")
            return fn(datos, titulo=titulo)

        elif param == "delta":
            raise ValueError(
                "El gráfico comparativo requiere seleccionar dos números. "
                "Usa la pestaña Comparativo.")

        elif param == "resultados_por_numero":
            datos = getattr(self, "_tono_resultados", {})
            if not datos:
                raise ValueError("Analiza el tono del corpus completo primero.")
            from collections import defaultdict
            por_num = defaultdict(dict)
            for aid, res in datos.items():
                num = res.get("numero", "sin_número")
                por_num[num][aid] = res
            return fn(dict(por_num), titulo=titulo)

        raise ValueError(f"Fuente de datos desconocida: {param}")

    def _viz_mostrar(self, fig):
        try:
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
        except ImportError:
            messagebox.showerror("Error", "Falta matplotlib.")
            return

        # Limpiar canvas anterior
        for widget in self._frm_viz_canvas.winfo_children():
            widget.destroy()

        self._viz_fig_actual = fig
        canvas = FigureCanvasTkAgg(fig, master=self._frm_viz_canvas)
        canvas.draw()
        canvas.get_tk_widget().pack(fill="both", expand=True)
        self._viz_canvas_widget = canvas

    def _viz_guardar(self):
        if self._viz_fig_actual is None:
            messagebox.showwarning("Sin gráfico", "Genera un gráfico primero.")
            return
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".png",
            filetypes=[("PNG", "*.png"), ("SVG", "*.svg"), ("PDF", "*.pdf")],
            initialfile="grafico_bashkar.png",
            title="Guardar gráfico")
        if not dest:
            return
        self._viz_fig_actual.savefig(dest, dpi=150, bbox_inches="tight",
                                      facecolor=self._viz_fig_actual.get_facecolor())
        self.toast(f"Gráfico guardado → {Path(dest).name}", tipo="ok")

    # ── Nube de palabras ───────────────────────────────────────────────────────
    def _build_viz_nube(self, parent):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 8))
        self._btn_nube = ttk.Button(bf, text="▶  Generar nube",
                                     style="P.TButton",
                                     command=self._viz_nube)
        self._btn_nube.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🌐  Abrir imagen", style="S.TButton",
                   command=self._viz_nube_abrir).pack(side="left")
        self._lbl_nube_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI", 9, "bold"))
        self._lbl_nube_ok.pack(anchor="w", pady=(0, 6))

        # Preview de imagen
        self._lbl_nube_img = tk.Label(pad, bg=CONTENT_BG,
                                       text="(La imagen aparecerá aquí después de generarla)",
                                       fg=GRIS2, font=("Segoe UI", 9))
        self._lbl_nube_img.pack(fill="both", expand=True)
        self._nube_path = None

    # ── Heatmap ────────────────────────────────────────────────────────────────
    def _build_viz_heatmap(self, parent):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        tk.Label(pad, text="Términos a seguir (uno por línea):",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w")
        self._txt_heat_terms = scrolledtext.ScrolledText(
            pad, height=6, font=("Consolas", 9),
            bg="#0D1B2A", fg="#94A3B8", wrap="word")
        self._txt_heat_terms.pack(fill="x", pady=(0, 6))
        default_terms = "colombia\nbogotá\nmedellín\nmujer\ncine\nradio\npolítica\ncultura"
        self._txt_heat_terms.insert("1.0", default_terms)

        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 6))
        self._btn_heat = ttk.Button(bf, text="▶  Generar heatmap",
                                     style="P.TButton",
                                     command=self._viz_heatmap)
        self._btn_heat.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🌐  Abrir imagen", style="S.TButton",
                   command=self._viz_heat_abrir).pack(side="left")
        self._lbl_heat_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI", 9, "bold"))
        self._lbl_heat_ok.pack(anchor="w")
        self._heat_path = None

    # ── Mapa ──────────────────────────────────────────────────────────────────
    def _build_viz_mapa(self, parent):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        tk.Label(pad,
                 text="Genera un mapa HTML interactivo con los lugares del índice NER.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 8))

        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 6))
        self._btn_mapa = ttk.Button(bf, text="▶  Generar mapa",
                                     style="P.TButton",
                                     command=self._viz_mapa)
        self._btn_mapa.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🌐  Abrir en navegador", style="S.TButton",
                   command=self._viz_mapa_abrir).pack(side="left")
        self._lbl_mapa_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI", 9, "bold"))
        self._lbl_mapa_ok.pack(anchor="w")
        self._mapa_path = None

        tk.Label(pad,
                 text="Nota: se mapean automáticamente ciudades colombianas conocidas del período.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 8)).pack(anchor="w", pady=(8, 0))

    # ── Timeline ──────────────────────────────────────────────────────────────
    def _build_viz_timeline(self, parent):
        pad = tk.Frame(parent, bg=CONTENT_BG, padx=10, pady=8)
        pad.pack(fill="both", expand=True)

        tk.Label(pad,
                 text="Genera timeline HTML con personas y eventos del índice NER.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 8))

        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 6))
        self._btn_tl = ttk.Button(bf, text="▶  Generar timeline",
                                   style="P.TButton",
                                   command=self._viz_timeline)
        self._btn_tl.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🌐  Abrir en navegador", style="S.TButton",
                   command=self._viz_tl_abrir).pack(side="left")
        self._lbl_tl_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                    font=("Segoe UI", 9, "bold"))
        self._lbl_tl_ok.pack(anchor="w")
        self._tl_path = None

    # ── Workers: Nube ─────────────────────────────────────────────────────────
    def _viz_nube(self):
        if not getattr(ST, "corpus_txt", None):
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        self._btn_nube.config(state="disabled")
        textos = [t for t in ST.corpus_txt if t and t.strip()]
        threading.Thread(target=self._worker_nube, args=(textos,), daemon=True).start()

    def _worker_nube(self, textos):
        from pathlib import Path as _PPath

        from core.viz_engine import nube_palabras
        self.after(0, lambda: self._lbl_nube_ok.config(text="Generando nube…"))
        try:
            ruta = _PPath.home() / "Documents" / "BashkarStation" / "viz" / "nube_palabras.png"
            nube_palabras(textos, ruta, titulo="Corpus Estampa 1930-1940")
            self._nube_path = ruta
            self.after(0, self._viz_nube_mostrar_preview)
            self.after(0, lambda: self._lbl_nube_ok.config(text=f"✅ Nube generada: {ruta}"))
        except ImportError as e:
            err = str(e)
            self.after(0, lambda: self._lbl_nube_ok.config(text=f"⚠ {err}"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_nube_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_nube.config(state="normal"))

    def _viz_nube_mostrar_preview(self):
        if not self._nube_path or not self._nube_path.exists():
            return
        try:
            from PIL import Image, ImageTk
            img = Image.open(str(self._nube_path))
            img.thumbnail((700, 350))
            self._nube_tk = ImageTk.PhotoImage(img)
            self._lbl_nube_img.config(image=self._nube_tk, text="")
        except Exception:
            pass

    def _viz_nube_abrir(self):
        if not self._nube_path or not self._nube_path.exists():
            messagebox.showinfo("Sin imagen", "Genera la nube de palabras primero.")
            return
        import os
        os.startfile(str(self._nube_path))

    # ── Workers: Heatmap ──────────────────────────────────────────────────────
    def _viz_heatmap(self):
        if not getattr(ST, "corpus_txt", None):
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        terminos_raw = self._txt_heat_terms.get("1.0", "end").strip()
        terminos = [t.strip() for t in terminos_raw.splitlines() if t.strip()]
        if not terminos:
            messagebox.showwarning("Sin términos", "Ingresa al menos un término.")
            return
        self._btn_heat.config(state="disabled")
        threading.Thread(target=self._worker_heatmap, args=(terminos,), daemon=True).start()

    def _worker_heatmap(self, terminos):
        from pathlib import Path as _PPath

        import pandas as pd

        from core.viz_engine import heatmap_temporal
        self.after(0, lambda: self._lbl_heat_ok.config(text="Generando heatmap…"))
        try:
            textos = ST.corpus_txt or []
            df = pd.DataFrame({"texto": textos,
                                "fecha": [f"1935-{(i%12)+1:02d}-01" for i in range(len(textos))]})
            ruta = _PPath.home() / "Documents" / "BashkarStation" / "viz" / "heatmap_temporal.png"
            heatmap_temporal(df, terminos, ruta=ruta)
            self._heat_path = ruta
            self.after(0, lambda: self._lbl_heat_ok.config(text=f"✅ Heatmap generado: {ruta}"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_heat_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_heat.config(state="normal"))

    def _viz_heat_abrir(self):
        if not self._heat_path or not self._heat_path.exists():
            messagebox.showinfo("Sin imagen", "Genera el heatmap primero.")
            return
        import os
        os.startfile(str(self._heat_path))

    # ── Workers: Mapa ─────────────────────────────────────────────────────────
    def _viz_mapa(self):
        if not getattr(ST, "indice_ner_global", None):
            messagebox.showwarning("Sin NER", "Analiza el corpus en la pestaña Índice NER primero.")
            return
        self._btn_mapa.config(state="disabled")
        threading.Thread(target=self._worker_mapa, daemon=True).start()

    def _worker_mapa(self):
        from pathlib import Path as _PPath

        from core.viz_engine import mapa_lugares
        self.after(0, lambda: self._lbl_mapa_ok.config(text="Generando mapa…"))
        try:
            ruta = _PPath.home() / "Documents" / "BashkarStation" / "viz" / "mapa_lugares.html"
            mapa_lugares(ST.indice_ner_global, ruta)
            self._mapa_path = ruta
            self.after(0, lambda: self._lbl_mapa_ok.config(text=f"✅ Mapa generado: {ruta}"))
        except ImportError as e:
            err = str(e)
            self.after(0, lambda: self._lbl_mapa_ok.config(text=f"⚠ {err}"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_mapa_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_mapa.config(state="normal"))

    def _viz_mapa_abrir(self):
        if not self._mapa_path or not self._mapa_path.exists():
            messagebox.showinfo("Sin mapa", "Genera el mapa primero.")
            return
        import webbrowser
        webbrowser.open(str(self._mapa_path))

    # ── Workers: Timeline ─────────────────────────────────────────────────────
    def _viz_timeline(self):
        if not getattr(ST, "indice_ner_global", None):
            messagebox.showwarning("Sin NER", "Analiza el corpus en la pestaña Índice NER primero.")
            return
        self._btn_tl.config(state="disabled")
        threading.Thread(target=self._worker_timeline, daemon=True).start()

    def _worker_timeline(self):
        from pathlib import Path as _PPath

        from core.timeline_engine import generar_timeline_html
        self.after(0, lambda: self._lbl_tl_ok.config(text="Generando timeline editorial…"))
        try:
            # Construir lista de artículos desde corpus_meta o df_articulos
            articulos = []
            if ST.df_articulos is not None and not ST.df_articulos.empty:
                for _, row in ST.df_articulos.iterrows():
                    articulos.append({
                        "art_id":  str(row.get("id", "")),
                        "titulo":  str(row.get("titulo", "") or ""),
                        "autor":   str(row.get("autor", "") or ""),
                        "seccion": str(row.get("seccion", "") or ""),
                        "fecha":   str(row.get("fecha_publicacion", "") or ""),
                        "numero":  str(row.get("numero", "") or ""),
                        "tono":    str(row.get("tono", "") or ""),
                    })
            elif ST.corpus_meta:
                for art_id, meta in ST.corpus_meta.items():
                    articulos.append({
                        "art_id":  art_id,
                        "titulo":  meta.get("titulo", ""),
                        "autor":   meta.get("autor", ""),
                        "seccion": meta.get("seccion", ""),
                        "fecha":   meta.get("fecha", ""),
                        "numero":  meta.get("numero", ""),
                    })

            if not articulos:
                self.after(0, lambda: self._lbl_tl_ok.config(
                    text="⚠ Sin artículos segmentados — ejecuta Segmentar primero"))
                self.after(0, lambda: self._btn_tl.config(state="normal"))
                return

            ruta = _PPath.home() / "Documents" / "BashkarStation" / "viz" / "timeline_editorial.html"
            ruta.parent.mkdir(parents=True, exist_ok=True)

            def _cb(n, total, msg):
                self.after(0, lambda: self._lbl_tl_ok.config(
                    text=f"⏳ {n}/{total} — {msg[:40]}"))

            generar_timeline_html(
                articulos, ruta,
                titulo_corpus=getattr(ST, "publicacion", "Corpus editorial"),
                callback=_cb,
            )
            self._tl_path = ruta
            n = len(articulos)
            self.after(0, lambda: self._lbl_tl_ok.config(
                text=f"✅ Timeline generada: {n} artículos — {ruta}"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_tl_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_tl.config(state="normal"))

    def _viz_tl_abrir(self):
        if not self._tl_path or not self._tl_path.exists():
            messagebox.showinfo("Sin timeline", "Genera la timeline primero.")
            return
        import webbrowser
        webbrowser.open(str(self._tl_path))


    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: REPORTE (v15) — narrativas académicas + HTML + Word
    # ══════════════════════════════════════════════════════════════════════════

    def _build_rep(self):
        pad = tk.Frame(self._tab_rep, bg=CONTENT_BG, padx=16, pady=12)
        pad.pack(fill="both", expand=True)
        tk.Label(pad, text="Reporte narrativo del corpus", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad,
                 text="Genera narrativas académicas con IA y exporta el reporte completo del análisis.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 10))

        # ── Narrativas ────────────────────────────────────────────────────────
        nar_frame = tk.LabelFrame(pad, text=" Narrativas académicas (Claude) ",
                                   bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9))
        nar_frame.pack(fill="x", pady=(0, 10))

        bf = tk.Frame(nar_frame, bg=CONTENT_BG)
        bf.pack(fill="x", padx=8, pady=6)
        self._btn_rep_nar = ttk.Button(bf, text="▶  Generar narrativas IA",
                                        style="P.TButton",
                                        command=self._rep_generar_narrativas)
        self._btn_rep_nar.pack(side="left", padx=(0, 8))
        self._lbl_rep_nar = tk.Label(bf, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI", 9))
        self._lbl_rep_nar.pack(side="left")

        self._txt_rep_nar = scrolledtext.ScrolledText(
            nar_frame, height=8, font=("Georgia", 9),
            bg="#0D1B2A", fg="#CBD5E1", wrap="word", state="disabled")
        self._txt_rep_nar.pack(fill="x", padx=8, pady=(0, 8))

        # ── Exportar HTML ─────────────────────────────────────────────────────
        html_frame = tk.LabelFrame(pad, text=" Reporte HTML scrollytelling ",
                                    bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9))
        html_frame.pack(fill="x", pady=(0, 10))
        hf = tk.Frame(html_frame, bg=CONTENT_BG)
        hf.pack(fill="x", padx=8, pady=6)
        self._btn_rep_html = ttk.Button(hf, text="▶  Generar reporte HTML",
                                         style="P.TButton",
                                         command=self._rep_generar_html)
        self._btn_rep_html.pack(side="left", padx=(0, 8))
        ttk.Button(hf, text="🌐  Abrir en navegador", style="S.TButton",
                   command=self._rep_abrir_html).pack(side="left", padx=(0, 8))
        self._lbl_rep_html = tk.Label(hf, text="", bg=CONTENT_BG, fg=VERDE,
                                       font=("Segoe UI", 9))
        self._lbl_rep_html.pack(side="left")

        # ── Exportar Word ─────────────────────────────────────────────────────
        word_frame = tk.LabelFrame(pad, text=" Exportar Word (.docx) ",
                                    bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9))
        word_frame.pack(fill="x", pady=(0, 10))
        wf = tk.Frame(word_frame, bg=CONTENT_BG)
        wf.pack(fill="x", padx=8, pady=6)
        self._btn_rep_word = ttk.Button(wf, text="📄  Exportar Word",
                                         style="S.TButton",
                                         command=self._rep_exportar_word)
        self._btn_rep_word.pack(side="left", padx=(0, 8))
        self._lbl_rep_word = tk.Label(wf, text="", bg=CONTENT_BG, fg=VERDE,
                                       font=("Segoe UI", 9))
        self._lbl_rep_word.pack(side="left")

        self._narrativas_data = {}
        self._rep_html_path = None

    def _rep_generar_narrativas(self):
        api_key, _m = _resolver_api_key_modelo("narrativas")
        if not api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API en Configuración.")
            return
        self._btn_rep_nar.config(state="disabled")
        self._lbl_rep_nar.config(text="Generando narrativas…")
        threading.Thread(target=self._worker_narrativas, daemon=True).start()

    def _worker_narrativas(self):
        from core.storytelling_engine import generar_narrativa
        # Bug F821: api_key no estaba definida — NameError al generar narrativas.
        api_key = ST.api_keys.get("anthropic", "") or ST.api_key
        narrativas = {}

        # Narrativa corpus
        stats_corpus = self._rep_stats_corpus()
        narrativas["corpus"] = generar_narrativa(stats_corpus, api_key, seccion="corpus")

        # Narrativa NER si disponible
        if getattr(ST, "indice_ner_global", None):
            top_ner = {}
            for cat, ents in ST.indice_ner_global.items():
                top = sorted(ents.items(), key=lambda x: len(x[1]), reverse=True)[:5]
                top_ner[cat] = [(e, len(a)) for e, a in top]
            narrativas["ner"] = generar_narrativa({"top_entidades": top_ner}, api_key, seccion="corpus")

        self._narrativas_data = narrativas
        texto = "\n\n---\n\n".join(f"[{s.upper()}]\n{t}" for s, t in narrativas.items())

        def _mostrar():
            self._txt_rep_nar.config(state="normal")
            self._txt_rep_nar.delete("1.0", "end")
            self._txt_rep_nar.insert("1.0", texto)
            self._txt_rep_nar.config(state="disabled")
            self._lbl_rep_nar.config(text=f"✅ {len(narrativas)} narrativas generadas")
            self._btn_rep_nar.config(state="normal")

        self.after(0, _mostrar)

    def _rep_generar_html(self):
        self._btn_rep_html.config(state="disabled")
        self._lbl_rep_html.config(text="Generando HTML…")
        threading.Thread(target=self._worker_rep_html, daemon=True).start()

    def _worker_rep_html(self):
        from pathlib import Path as _PPath

        from core.storytelling_engine import generar_reporte_html
        try:
            nombre = getattr(ST, "proyecto_nombre", "Corpus Estampa")
            ruta = _PPath.home() / "Documents" / "BashkarStation" / "reporte" / "reporte_bashkar.html"
            generar_reporte_html(
                proyecto_nombre=nombre,
                stats_corpus=self._rep_stats_corpus(),
                indice_ner=getattr(ST, "indice_ner_global", {}),
                stats_tono=getattr(self, "_tono_resultados", None),
                metricas_red=getattr(self, "_metricas_red_cache", None),
                narrativas=self._narrativas_data,
                ruta=ruta,
            )
            self._rep_html_path = ruta
            self.after(0, lambda: self._lbl_rep_html.config(text=f"✅ HTML generado: {ruta}"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_rep_html.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_rep_html.config(state="normal"))

    def _rep_abrir_html(self):
        if not self._rep_html_path or not self._rep_html_path.exists():
            messagebox.showinfo("Sin reporte", "Genera el reporte HTML primero.")
            return
        import webbrowser
        webbrowser.open(str(self._rep_html_path))

    def _rep_exportar_word(self):
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word", "*.docx"), ("Todos", "*.*")],
            initialfile="reporte_estampa.docx",
            title="Guardar reporte Word")
        if not dest:
            return
        self._btn_rep_word.config(state="disabled")
        threading.Thread(target=self._worker_rep_word, args=(dest,), daemon=True).start()

    def _worker_rep_word(self, dest):
        from pathlib import Path as _PPath

        from core.storytelling_engine import exportar_word
        try:
            nombre = getattr(ST, "proyecto_nombre", "Corpus Estampa")
            exportar_word(
                proyecto_nombre=nombre,
                stats_corpus=self._rep_stats_corpus(),
                indice_ner=getattr(ST, "indice_ner_global", {}),
                narrativas=self._narrativas_data,
                ruta=_PPath(dest),
            )
            self.after(0, lambda: self._lbl_rep_word.config(text=f"✅ Exportado: {dest}"))
            self.after(0, lambda: messagebox.showinfo("Exportado", f"Word guardado en:\n{dest}"))
        except ImportError:
            self.after(0, lambda: self._lbl_rep_word.config(
                text="⚠ python-docx no instalado. pip install python-docx>=1.1.0"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_rep_word.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_rep_word.config(state="normal"))

    def _rep_stats_corpus(self) -> dict:
        return {
            "n_pdfs": len(getattr(ST, "pdf_files", []) or []),
            "n_paginas": len(_cm) if (_cm := getattr(ST, "corpus_meta", None)) is not None else 0,
            "n_articulos": len(getattr(ST, "articulos", []) or []),
            "n_palabras_total": sum(
                len((t or "").split())
                for t in (getattr(ST, "corpus_txt", []) or [])
            ),
            "proyecto": getattr(ST, "proyecto_nombre", "Corpus Estampa"),
        }


    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: DASHBOARD (v16) — resumen ejecutivo + exportador ZIP
    # ══════════════════════════════════════════════════════════════════════════

    def _build_dash(self):
        pad = tk.Frame(self._tab_dash, bg=CONTENT_BG, padx=16, pady=12)
        pad.pack(fill="both", expand=True)
        tk.Label(pad, text="Dashboard ejecutivo", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad, text="Resumen del estado del proyecto y exportación del paquete completo.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 10))

        # ── Grid de indicadores ───────────────────────────────────────────────
        grid = tk.Frame(pad, bg=CONTENT_BG)
        grid.pack(fill="x", pady=(0, 12))
        self._dash_cards = {}
        indicadores = [
            ("pdfs",      "PDFs cargados",    "0"),
            ("paginas",   "Páginas OCR",      "0"),
            ("articulos", "Artículos",         "0"),
            ("palabras",  "Palabras",          "0"),
            ("entidades", "Entidades NER",     "0"),
            ("red_nodos", "Nodos en red",      "0"),
        ]
        for col, (key, lbl, val) in enumerate(indicadores):
            frm = tk.Frame(grid, bg="#1a1a2e", padx=12, pady=8, relief="flat", bd=0)
            frm.grid(row=0, column=col, padx=4, pady=2, sticky="nsew")
            grid.columnconfigure(col, weight=1)
            lbl_num = tk.Label(frm, text=val, bg="#1a1a2e", fg="#a78bfa",
                                font=("Segoe UI", 20, "bold"))
            lbl_num.pack()
            tk.Label(frm, text=lbl, bg="#1a1a2e", fg="#94a3b8",
                     font=("Segoe UI", 8)).pack()
            self._dash_cards[key] = lbl_num

        # ── Progreso por módulo ───────────────────────────────────────────────
        prog_frame = tk.LabelFrame(pad, text=" Estado del análisis ",
                                    bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9))
        prog_frame.pack(fill="x", pady=(0, 10))
        self._dash_prog_labels = {}
        modulos = [
            ("ocr",  "OCR / Extracción",    "ocr_done"),
            ("seg",  "Segmentación",         "seg_done"),
            ("ner",  "Índice NER",           "ner_done"),
            ("red",  "Redes",                None),
            ("sem",  "Semántico",            None),
            ("rep",  "Reporte",              None),
        ]
        for i, (mid, mlabel, badge) in enumerate(modulos):
            row = i // 3
            col = i % 3
            frm = tk.Frame(prog_frame, bg=CONTENT_BG)
            frm.grid(row=row, column=col, padx=8, pady=4, sticky="w")
            lbl = tk.Label(frm, text=f"◦ {mlabel}", bg=CONTENT_BG,
                           fg=GRIS2, font=("Segoe UI", 9))
            lbl.pack(side="left")
            self._dash_prog_labels[mid] = (lbl, badge)

        # ── Botón actualizar ──────────────────────────────────────────────────
        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 8))
        ttk.Button(bf, text="↻  Actualizar dashboard",
                   style="S.TButton",
                   command=self._dash_actualizar).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="▶  Generar reporte completo",
                   style="P.TButton",
                   command=self._dash_reporte_completo).pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="🚀  Paquete completo (pipeline maestro)",
                   style="P.TButton",
                   command=self._dash_pipeline_maestro).pack(side="left", padx=(0, 8))
        self._btn_dash_zip = ttk.Button(bf, text="📦  Exportar ZIP",
                                         style="S.TButton",
                                         command=self._dash_exportar_zip)
        self._btn_dash_zip.pack(side="right")
        self._lbl_dash_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                      font=("Segoe UI", 9, "bold"))
        self._lbl_dash_ok.pack(anchor="w", pady=(0, 6))

        # ── Log ───────────────────────────────────────────────────────────────
        self._txt_dash_log = scrolledtext.ScrolledText(
            pad, height=7, font=("Consolas", 8),
            bg="#0D1B2A", fg="#94A3B8", state="disabled", wrap="word")
        self._txt_dash_log.pack(fill="x")

        # Actualizar al abrir
        self.after(500, self._dash_actualizar)

    def _dash_log(self, msg: str):
        self._txt_dash_log.config(state="normal")
        self._txt_dash_log.insert("end", msg + "\n")
        self._txt_dash_log.see("end")
        self._txt_dash_log.config(state="disabled")

    def _dash_actualizar(self):
        # Indicadores numéricos
        def set_card(key, val):
            if key in self._dash_cards:
                self._dash_cards[key].config(text=str(val))

        set_card("pdfs",      len(getattr(ST, "pdf_files", []) or []))
        _cm = getattr(ST, "corpus_meta", None)
        set_card("paginas",   len(_cm) if _cm is not None else 0)
        set_card("articulos", len(getattr(ST, "articulos", []) or []))
        n_words = sum(len((t or "").split()) for t in (getattr(ST, "corpus_txt", []) or []))
        set_card("palabras",  f"{n_words:,}")
        ner = getattr(ST, "indice_ner_global", {}) or {}
        set_card("entidades", sum(len(v) for v in ner.values()))
        grafo = getattr(self, "_grafo_actual", None)
        set_card("red_nodos", grafo.number_of_nodes() if grafo else 0)

        # Estado módulos
        for mid, (lbl, badge) in self._dash_prog_labels.items():
            if badge:
                done = getattr(ST, badge, False)
            else:
                # Inferir por existencia de datos
                if mid == "red":
                    done = getattr(self, "_grafo_actual", None) is not None
                elif mid == "sem":
                    done = bool(getattr(self, "_tono_resultados", {}))
                elif mid == "rep":
                    done = bool(getattr(self, "_narrativas_data", {}))
                else:
                    done = False
            color = VERDE if done else GRIS2
            ico = "✅" if done else "◦"
            lbl.config(text=f"{ico} {mid.upper()}", fg=color)

        self._lbl_dash_ok.config(text="Dashboard actualizado")

    def _dash_reporte_completo(self):
        self._dash_log("Iniciando generación de reporte completo…")
        self._dash_log("  1/3 Narrativas IA…")
        self._rep_generar_narrativas()
        self.after(3000, lambda: (
            self._dash_log("  2/3 Reporte HTML…"),
            self._rep_generar_html(),
        ))
        self.after(6000, lambda: (
            self._dash_log("  3/3 Dashboard actualizado"),
            self._dash_actualizar(),
            self._lbl_dash_ok.config(text="✅ Reporte completo generado"),
        ))

    def _dash_exportar_zip(self):
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".zip",
            filetypes=[("ZIP", "*.zip"), ("Todos", "*.*")],
            initialfile="bashkar_export.zip",
            title="Exportar paquete ZIP")
        if not dest:
            return
        self._btn_dash_zip.config(state="disabled")
        self._lbl_dash_ok.config(text="Empaquetando…")
        threading.Thread(target=self._worker_zip, args=(dest,), daemon=True).start()

    def _worker_zip(self, dest):
        import zipfile
        from pathlib import Path as _PPath
        base = _PPath.home() / "Documents" / "BashkarStation"
        try:
            archivos = list(base.rglob("*")) if base.exists() else []
            archivos = [f for f in archivos if f.is_file()]
            with zipfile.ZipFile(dest, "w", zipfile.ZIP_DEFLATED) as zf:
                for f in archivos:
                    zf.write(f, f.relative_to(base.parent))
            n = len(archivos)
            self.after(0, lambda: self._lbl_dash_ok.config(
                text=f"✅ ZIP exportado: {n} archivos → {dest}"))
            self.after(0, lambda: messagebox.showinfo("Exportado",
                f"Paquete ZIP creado con {n} archivos:\n{dest}"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_dash_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_dash_zip.config(state="normal"))


    # ══════════════════════════════════════════════════════════════════════════
    # PESTAÑA: TÓPICOS (visión definitiva)
    # ══════════════════════════════════════════════════════════════════════════

    def _build_top(self):
        outer = tk.Frame(self._tab_top, bg=CONTENT_BG)
        outer.pack(fill="both", expand=True, padx=16, pady=12)

        self._top_params: dict = {}
        try:
            from core.topic_engine import PARAMS_SCHEMA as _TOP_SCHEMA
            self._build_params_panel(outer, _TOP_SCHEMA, self._top_params)
        except Exception:
            pass

        pad = tk.Frame(outer, bg=CONTENT_BG)
        pad.pack(side="left", fill="both", expand=True)

        tk.Label(pad, text="Topic modeling del corpus", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad, text="Detecta temas recurrentes y su distribución en el corpus.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(anchor="w", pady=(0, 10))

        # Controles
        cf = tk.Frame(pad, bg=CONTENT_BG)
        cf.pack(fill="x", pady=(0, 6))
        tk.Label(cf, text="N° tópicos:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 4))
        self._var_top_n = tk.IntVar(value=8)
        ttk.Spinbox(cf, from_=3, to=20, textvariable=self._var_top_n,
                    width=4).pack(side="left", padx=(0, 12))
        self._var_top_llm = tk.BooleanVar(value=True)
        ttk.Checkbutton(cf, text="Etiquetar con IA (Claude)",
                        variable=self._var_top_llm).pack(side="left", padx=(0, 12))
        self._var_top_bertopic = tk.BooleanVar(value=False)
        ttk.Checkbutton(cf, text="BERTopic (requiere GPU/modelos)",
                        variable=self._var_top_bertopic).pack(side="left")

        # Botones
        bf = tk.Frame(pad, bg=CONTENT_BG)
        bf.pack(fill="x", pady=(0, 6))
        self._btn_top = ttk.Button(bf, text="▶  Modelar tópicos",
                                    style="P.TButton",
                                    command=self._top_ejecutar)
        self._btn_top.pack(side="left", padx=(0, 8))
        ttk.Button(bf, text="💾  Exportar CSV", style="S.TButton",
                   command=self._top_exportar).pack(side="left", padx=(0, 8))
        self._lbl_top_ok = tk.Label(pad, text="", bg=CONTENT_BG, fg=VERDE,
                                     font=("Segoe UI", 9, "bold"))
        self._lbl_top_ok.pack(anchor="w", pady=(0, 4))

        # Tabla de tópicos
        tk.Label(pad, text="Tópicos detectados", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 10, "bold")).pack(anchor="w", pady=(4, 2))
        cols_top = ("id", "nombre", "n_docs", "porcentaje", "palabras_clave")
        self._tv_top = ttk.Treeview(pad, columns=cols_top, show="headings", height=10)
        heads_top = [("id", "#", 40), ("nombre", "Nombre/Tema", 160),
                     ("n_docs", "Artículos", 70), ("porcentaje", "%", 60),
                     ("palabras_clave", "Palabras clave", 380)]
        for cid, txt, w in heads_top:
            self._tv_top.heading(cid, text=txt)
            self._tv_top.column(cid, width=w, anchor="w")
        sv = ttk.Scrollbar(pad, orient="vertical", command=self._tv_top.yview)
        self._tv_top.configure(yscrollcommand=sv.set)
        self._tv_top.pack(side="left", fill="both", expand=True)
        sv.pack(side="left", fill="y")

        self._top_resultado = {}

    def _top_ejecutar(self):
        if not getattr(ST, "corpus_txt", None):
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        # Leer params del panel lateral (prioridad sobre spinboxes legacy)
        p = self._params_get_values(self._top_params) if getattr(self, "_top_params", None) else {}
        n             = int(p.get("n_topics",          self._var_top_n.get()))
        usar_llm      = bool(p.get("etiquetar_ia",     self._var_top_llm.get()))
        usar_bertopic = p.get("backend", "nmf") == "bertopic"
        min_df        = int(p.get("min_df", 2))
        max_df        = float(p.get("max_df", 0.95))
        n_words       = int(p.get("palabras_por_topic", 10))
        if usar_llm and not ST.api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API o desactiva 'Etiquetar con IA'.")
            return
        self._btn_top.config(state="disabled")
        self._lbl_top_ok.config(text="Modelando tópicos…")
        textos = [t for t in ST.corpus_txt if t and t.strip()]
        threading.Thread(target=self._worker_top,
                         args=(textos, n, usar_llm, usar_bertopic, min_df, max_df, n_words),
                         daemon=True).start()

    def _worker_top(self, textos, n, usar_llm, usar_bertopic,
                    min_df=2, max_df=0.95, n_words=10):
        from core.topic_engine import estadisticas_topicos, modelar_topicos
        def cb(m):
            self.after(0, lambda msg=m: self._lbl_top_ok.config(text=str(msg)[:80]))
        try:
            resultado = modelar_topicos(
                textos, n_topicos=n,
                api_key=ST.api_key if usar_llm else None,
                usar_bertopic=usar_bertopic,
                min_df=min_df, max_df=max_df,
                n_palabras=n_words,
                callback=cb,
            )
            self._top_resultado = resultado
            stats = estadisticas_topicos(resultado)
            self.after(0, lambda s=stats: self._top_refrescar(s))
            backend = resultado.get("backend", "?")
            n_top = len(resultado.get("topicos", {}))
            self.after(0, lambda: self._lbl_top_ok.config(
                text=f"✅ {n_top} tópicos detectados ({backend})"))
        except Exception as e:
            err = str(e)
            self.after(0, lambda: self._lbl_top_ok.config(text=f"⚠ Error: {err}"))
        self.after(0, lambda: self._btn_top.config(state="normal"))

    def _top_refrescar(self, stats):
        for row in self._tv_top.get_children():
            self._tv_top.delete(row)
        for tid, info in sorted(stats.get("topicos", {}).items(), key=lambda x: -x[1].get("n_docs", 0)):
            self._tv_top.insert("", "end", values=(
                tid,
                info.get("nombre", f"Tópico {tid}"),
                info.get("n_docs", 0),
                f"{info.get('porcentaje', 0)}%",
                ", ".join(info.get("palabras_clave", [])[:6]),
            ))

    def _top_exportar(self):
        if not self._top_resultado:
            messagebox.showwarning("Sin datos", "Ejecuta el topic modeling primero.")
            return
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="topicos_corpus.csv",
            title="Exportar tópicos")
        if not dest:
            return
        from pathlib import Path as _PPath

        from core.topic_engine import exportar_topicos_csv
        n = exportar_topicos_csv(self._top_resultado, _PPath(dest))
        messagebox.showinfo("Exportado", f"✅ {n} entradas exportadas a:\n{dest}")

    # ══════════════════════════════════════════════════════════════════════════
    # PIPELINE MAESTRO — botón "Generar paquete completo" en Dashboard
    # ══════════════════════════════════════════════════════════════════════════

    def _dash_pipeline_maestro(self):
        """Lanza el PipelineMaestro completo desde el Dashboard."""
        if not ST.api_key:
            messagebox.showwarning("Sin API key", "Configura tu clave Claude API en Configuración.")
            return
        if not self._proyecto_ruta:
            messagebox.showwarning("Sin proyecto",
                "Guarda el proyecto primero (Archivo > Guardar proyecto).")
            return
        if not messagebox.askyesno("Generar paquete completo",
            "Esto ejecutará todo el análisis y generará el paquete ZIP de investigación.\n"
            "Puede tardar varios minutos según el tamaño del corpus.\n\n"
            "¿Continuar?"):
            return

        self._lbl_dash_ok.config(text="🚀 Pipeline Maestro iniciado…")
        self._dash_log("=== PIPELINE MAESTRO INICIADO ===")
        self._dash_log(f"Proyecto: {self._proyecto_ruta}")

        articulos = self._pipeline_maestro_articulos()

        from core.pipeline_maestro import PipelineMaestro
        pm = PipelineMaestro(
            bashkar_path=str(self._proyecto_ruta),
            api_key=ST.api_key,
            callback_progreso=lambda p, m: self.after(0, lambda pct=p, msg=m: (
                self._lbl_dash_ok.config(text=f"[{pct}%] {msg}"),
                self._dash_log(f"[{pct}%] {msg}"),
            )),
            callback_log=lambda m: self.after(0, lambda msg=m: self._dash_log(msg)),
            repositorio=ST.repo,
        )
        pm.ejecutar_en_hilo(articulos_existentes=articulos if articulos else None)

    def _pipeline_maestro_articulos(self) -> list:
        """Construye lista de artículos desde el estado actual del proyecto."""
        articulos = []
        corpus_txt = getattr(ST, "corpus_txt", []) or []
        corpus_meta = getattr(ST, "corpus_meta", None)
        if corpus_meta is None:
            corpus_meta = {}
        ner_global = getattr(ST, "indice_ner_global", {}) or {}

        if hasattr(corpus_meta, "iterrows"):
            for _, row in corpus_meta.iterrows():
                art_id = str(row.get("numero", "?")) + "_" + str(row.get("pagina", "?"))
                txt_path = row.get("txt_path", "")
                texto = ""
                if txt_path:
                    from pathlib import Path as _PPath
                    p = _PPath(str(txt_path))
                    if p.exists():
                        texto = p.read_text("utf-8", errors="replace")
                articulos.append({
                    "id": art_id,
                    "texto": texto,
                    "titulo": None,
                    "autor": None,
                    "fecha": None,
                    "ner": {},
                })
        elif corpus_txt:
            for i, t in enumerate(corpus_txt):
                articulos.append({
                    "id": f"art_{i:04d}",
                    "texto": t or "",
                    "titulo": None,
                    "autor": None,
                    "fecha": None,
                    "ner": {},
                })
        return articulos

    # ══════════════════════════════════════════════════════════════════════════
    # «GUARDAR COMO…» — presets de exportación estilo ABBYY FineReader
    # ══════════════════════════════════════════════════════════════════════════

    def _exp_abrir_dialogo(self):
        """Diálogo con 4 presets: Copia exacta (PDF buscable) / Edición
        académica (TEI+BibTeX) / Datos de análisis (Excel) / Texto plano.
        Reusa los exportadores YA existentes para los 3 últimos; el único
        camino nuevo es el PDF buscable."""
        from core.user_prefs import guardar_pref, obtener_pref

        win, content = self._mk_glass_toplevel("Guardar como…", 480, 380)
        pad = tk.Frame(content, bg=CONTENT_BG)
        pad.pack(fill="both", expand=True, padx=20, pady=16)

        var_abrir = tk.BooleanVar(value=obtener_pref("exp_abrir_al_terminar", True))
        ttk.Checkbutton(pad, text="Abrir el archivo al terminar",
                         variable=var_abrir,
                         command=lambda: guardar_pref("exp_abrir_al_terminar", var_abrir.get())
                         ).pack(anchor="w", pady=(0, 12))

        def _tarjeta(titulo, descripcion, comando):
            c = tk.Frame(pad, bg=CARD_BG, relief="solid", bd=1, cursor="hand2")
            c.pack(fill="x", pady=4)
            tk.Label(c, text=titulo, bg=CARD_BG, fg=TXT_PRI,
                     font=("Segoe UI", 10, "bold")).pack(anchor="w", padx=12, pady=(8, 0))
            tk.Label(c, text=descripcion, bg=CARD_BG, fg=TXT_DIM,
                     font=("Segoe UI", 8), wraplength=400, justify="left").pack(
                         anchor="w", padx=12, pady=(0, 8))

            def _click(_e=None, fn=comando):
                win.destroy()
                fn(abrir_al_terminar=var_abrir.get())
            c.bind("<Button-1>", _click)
            for w in c.winfo_children():
                w.bind("<Button-1>", _click)

        _tarjeta("📄 Copia exacta", "PDF buscable: imagen de cada página + capa de "
                 "texto invisible (búsqueda y copiar/pegar). Estilo FineReader.",
                 self._exp_pdf_buscable)
        _tarjeta("🎓 Edición académica", "XML-TEI P5 + BibTeX del corpus (pide destino "
                 "para cada uno).", self._exp_edicion_academica)
        _tarjeta("📊 Datos de análisis", "Excel completo (10 hojas) con todo lo "
                 "generado en Análisis/Visualizar.", self._exp_datos_analisis)
        _tarjeta("📝 Texto plano", "Todo el corpus concatenado en un único .md/.txt.",
                 self._exp_texto_plano)

    def _exp_resolver_imagen(self, numero: str, pagina: str) -> "Path | None":
        """Busca la imagen original de una página en 02_imagenes/<numero>/."""
        if not ST.out_dir:
            return None
        img_dir = Path(ST.out_dir) / "02_imagenes" / str(numero)
        if not img_dir.exists():
            return None
        for ext in ("png", "jpg", "jpeg", "tif", "tiff"):
            hits = sorted(img_dir.glob(f"*{pagina}*.{ext}"))
            if hits:
                return hits[0]
        return None

    def _exp_pdf_buscable(self, abrir_al_terminar: bool = True):
        if ST.corpus_meta is None or ST.corpus_meta.empty:
            messagebox.showwarning("Sin datos",
                "No hay páginas con OCR. Ejecuta Extracción OCR primero.")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".pdf", filetypes=[("PDF", "*.pdf")],
            initialfile="corpus_buscable.pdf", title="Guardar PDF buscable")
        if not dest:
            return
        threading.Thread(target=self._exp_pdf_worker,
                         args=(dest, abrir_al_terminar), daemon=True).start()

    def _exp_pdf_worker(self, dest: str, abrir_al_terminar: bool):
        from core.pdf_export import exportar_pdf_buscable

        paginas = []
        for _, row in ST.corpus_meta.iterrows():
            numero, pagina = str(row.get("numero", "")), str(row.get("pagina", ""))
            img_path = self._exp_resolver_imagen(numero, pagina)
            if not img_path:
                continue
            texto = ""
            tp = row.get("txt_path")
            if tp and Path(tp).exists():
                try:
                    texto = Path(tp).read_text("utf-8", errors="replace")
                except Exception:
                    texto = ""
            paginas.append({"img_path": img_path, "texto": texto})

        if not paginas:
            self.after(0, lambda: messagebox.showwarning("Sin imágenes",
                "No se encontraron imágenes de página en 02_imagenes/ para exportar."))
            return

        prog_win = tk.Toplevel(self)
        prog_win.title("Generando PDF buscable…")
        prog_win.geometry("360x90")
        tk.Label(prog_win, text="Generando PDF buscable…").pack(pady=(14, 4))
        bar = ttk.Progressbar(prog_win, mode="determinate", maximum=len(paginas))
        bar.pack(fill="x", padx=16)
        lbl = tk.Label(prog_win, text=f"0/{len(paginas)}")
        lbl.pack(pady=6)

        def cb(n, total):
            self.after(0, lambda: (bar.config(value=n), lbl.config(text=f"{n}/{total}")))

        try:
            exportar_pdf_buscable(paginas, dest, callback=cb)
            self.after(0, prog_win.destroy)
            self.after(0, lambda: self.toast(
                f"✅ PDF buscable exportado: {len(paginas)} páginas", "ok"))
            if abrir_al_terminar:
                self.after(0, lambda: os.startfile(dest))
        except Exception as e:
            self.after(0, prog_win.destroy)
            self.after(0, lambda err=str(e): messagebox.showerror("Error", err))

    def _exp_edicion_academica(self, abrir_al_terminar: bool = True):
        self._res_exportar_tei()
        self._res_exportar_bibtex()

    def _exp_datos_analisis(self, abrir_al_terminar: bool = True):
        self._gen_excel()

    def _exp_texto_plano(self, abrir_al_terminar: bool = True):
        corpus_txt = getattr(ST, "corpus_txt", []) or []
        if not corpus_txt:
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        dest = filedialog.asksaveasfilename(
            defaultextension=".md", filetypes=[("Markdown", "*.md"), ("Texto", "*.txt")],
            initialfile="corpus_completo.md", title="Guardar texto plano")
        if not dest:
            return
        try:
            contenido = "\n\n---\n\n".join(
                f"## Artículo {i:04d}\n\n{t}" for i, t in enumerate(corpus_txt))
            Path(dest).write_text(contenido, encoding="utf-8")
            self.toast(f"✅ Texto plano exportado: {len(corpus_txt)} artículos", "ok")
            if abrir_al_terminar:
                os.startfile(dest)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    # ══════════════════════════════════════════════════════════════════════════
    # EXPORTACIONES ESPECIALIZADAS — agrega botones a panel Resultados
    # ══════════════════════════════════════════════════════════════════════════

    def _res_exportar_tei(self):
        """Exporta corpus XML-TEI P5 desde el panel Resultados."""
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".xml",
            filetypes=[("XML-TEI", "*.xml"), ("Todos", "*.*")],
            initialfile="corpus_estampa_tei.xml",
            title="Exportar XML-TEI")
        if not dest:
            return
        corpus_txt = getattr(ST, "corpus_txt", []) or []
        ner_global = getattr(ST, "indice_ner_global", {}) or {}
        if not corpus_txt:
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        articulos = []
        for i, t in enumerate(corpus_txt):
            art_id = f"art_{i:04d}"
            ner_art = {}
            for cat, ents in ner_global.items():
                ner_art[cat] = [e for e, arts in ents.items() if art_id in arts]
            articulos.append({"id": art_id, "texto": t or "", "ner": ner_art})
        try:
            from pathlib import Path as _PPath

            from core.tei_engine import exportar_corpus_tei
            proyecto = getattr(ST, "proyecto_nombre", "Corpus Estampa")
            exportar_corpus_tei(articulos, _PPath(dest),
                                proyecto_nombre=proyecto)
            messagebox.showinfo("Exportado", f"✅ XML-TEI exportado:\n{dest}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _res_exportar_pptx(self):
        """Exporta presentación PowerPoint con resultados del corpus."""
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            title="Guardar presentación PowerPoint",
            defaultextension=".pptx",
            filetypes=[("PowerPoint", "*.pptx"), ("Todos", "*.*")],
        )
        if not dest:
            return
        import threading
        threading.Thread(target=self._res_exportar_pptx_worker, args=(dest,), daemon=True).start()

    def _res_exportar_pptx_worker(self, dest):
        try:
            from pathlib import Path

            from exportadores.exportar_pptx import exportar_presentacion
            datos = {
                "articulos": {},
                "indice_ner_global": ST.indice_ner_global,
                "topicos": getattr(ST, "topicos_resultado", {}),
                "metricas_red": getattr(ST, "metricas_red", {}),
                "estadisticas_tono": getattr(ST, "estadisticas_tono", {}),
                "narrativa": getattr(ST, "narrativa_corpus", ""),
            }
            if ST.df_articulos is not None:
                for _, row in ST.df_articulos.iterrows():
                    aid = str(row.get("id", row.name))
                    datos["articulos"][aid] = {"texto_limpio": str(row.get("texto", ""))}
            exportar_presentacion(
                datos, Path(dest),
                titulo_proyecto=ST.publicacion,
                investigador=getattr(ST, "investigador", ""),
                institucion=getattr(ST, "institucion", "Instituto Caro y Cuervo"),
            )
            ST.pptx_path = dest
            self.after(0, lambda: __import__("tkinter.messagebox", fromlist=["showinfo"]).showinfo(
                "PowerPoint", f"✅ Presentación guardada:\n{dest}"))
        except ImportError:
            self.after(0, lambda: __import__("tkinter.messagebox", fromlist=["showerror"]).showerror(
                "PowerPoint", "Instala python-pptx: pip install python-pptx"))
        except Exception as e:
            self.after(0, lambda err=e: __import__("tkinter.messagebox", fromlist=["showerror"]).showerror(
                "PowerPoint", str(err)))

    def _res_exportar_bibtex(self):
        """Exporta BibTeX del corpus."""
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".bib",
            filetypes=[("BibTeX", "*.bib"), ("Todos", "*.*")],
            initialfile="corpus_estampa.bib",
            title="Exportar BibTeX")
        if not dest:
            return
        corpus_txt = getattr(ST, "corpus_txt", []) or []
        if not corpus_txt:
            messagebox.showwarning("Sin corpus", "Extrae el texto del corpus primero.")
            return
        articulos = [{"id": f"art_{i:04d}", "texto": t} for i, t in enumerate(corpus_txt)]
        try:
            from pathlib import Path as _PPath

            from core.tei_engine import exportar_bibtex
            exportar_bibtex(articulos, _PPath(dest))
            messagebox.showinfo("Exportado", f"✅ BibTeX exportado:\n{dest}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _res_exportar_json_observable(self):
        """Exporta datos en JSON para Observable / Flourish."""
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("JSON", "*.json"), ("Todos", "*.*")],
            initialfile="corpus_observable.json",
            title="Exportar JSON para Observable/Flourish")
        if not dest:
            return
        import json as _json
        ner = getattr(ST, "indice_ner_global", {}) or {}
        grafo = getattr(self, "_grafo_actual", None)
        datos = {
            "proyecto": getattr(ST, "proyecto_nombre", "Corpus Estampa"),
            "n_articulos": len(getattr(ST, "corpus_txt", []) or []),
            "entidades": {
                cat: [{"entidad": e, "n_articulos": len(arts)}
                      for e, arts in sorted(ents.items(), key=lambda x: -len(x[1]))[:50]]
                for cat, ents in ner.items()
            },
            "red": {
                "nodos": [{"id": n, **d} for n, d in grafo.nodes(data=True)]
                if grafo else [],
                "aristas": [{"source": u, "target": v, "peso": d.get("weight", 1)}
                             for u, v, d in grafo.edges(data=True)]
                if grafo else [],
            },
        }
        try:
            with open(dest, "w", encoding="utf-8") as f:
                _json.dump(datos, f, ensure_ascii=False, indent=2, default=str)
            messagebox.showinfo("Exportado", f"✅ JSON exportado:\n{dest}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    # ── Validar TEI ───────────────────────────────────────────────────────────

    def _res_generar_methods(self):
        """Genera METHODS.md con descripción metodológica completa."""
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".md",
            filetypes=[("Markdown","*.md"),("Texto","*.txt")],
            initialfile="METHODS.md",
            title="Guardar sección de metodología")
        if not dest:
            return
        try:
            from core.methods_reporter import generar_methods_md
            cfg = {
                "publicacion":    getattr(ST, "publicacion", ""),
                "periodo":        getattr(ST, "periodo", ""),
                "investigador":   getattr(ST, "investigador", ""),
                "institucion":    getattr(ST, "institucion", ""),
                "bashkar_version": APP_VERSION,
                "dpi":            getattr(ST, "dpi", "150"),
                "lang":           getattr(ST, "lang", "spa"),
                "lematizar":      getattr(ST, "lematizar", True),
                "modelos_etapa":  getattr(ST, "modelos_etapa", {}),
                "archivos_sel":   getattr(ST, "archivos_sel", []),
            }
            stats = {
                "n_paginas":   len(getattr(ST, "archivos_sel", []) or []),
                "n_palabras":  0,
                "n_articulos": (len(ST.df_articulos)
                               if ST.df_articulos is not None else 0),
                "n_entidades": sum(len(v) for v in
                                   (getattr(ST, "indice_ner_global", {}) or {}).values()
                                   if isinstance(v, dict)),
            }
            ruta = generar_methods_md(cfg, stats, Path(dest))
            messagebox.showinfo("METHODS.md generado",
                                f"✅ Sección de metodología guardada:\n{ruta}\n\n"
                                f"Revisa y complementa antes de publicar.")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def _res_validar_tei(self):
        """Valida el último XML-TEI exportado."""
        from tkinter import filedialog
        ruta = filedialog.askopenfilename(
            title="Seleccionar XML-TEI para validar",
            filetypes=[("XML-TEI", "*.xml"), ("Todos", "*.*")])
        if not ruta:
            return
        try:
            from core.tei_engine import validar_tei
            errores = validar_tei(Path(ruta))
        except Exception as e:
            messagebox.showerror("Error", str(e))
            return
        if not errores:
            messagebox.showinfo("TEI válido ✅",
                                f"El archivo es un XML-TEI P5 válido.\n\n{ruta}")
        else:
            resumen = "\n".join(f"• {e}" for e in errores[:10])
            messagebox.showwarning("Problemas encontrados",
                                   f"{len(errores)} problema(s):\n\n{resumen}\n\n{ruta}")

    # ── Paquete para publicación ───────────────────────────────────────────────

    def _res_paquete_publicacion(self):
        """Genera ZIP con todos los artefactos para publicación académica."""
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".zip",
            filetypes=[("ZIP", "*.zip")],
            initialfile=f"paquete_publicacion_{datetime.now().strftime('%Y%m%d')}.zip",
            title="Guardar paquete de publicación")
        if not dest:
            return
        self._lbl_excel.config(text="⏳ Generando paquete…")
        threading.Thread(
            target=self._worker_paquete_publicacion,
            args=(dest,), daemon=True).start()

    def _worker_paquete_publicacion(self, dest_zip: str):
        import json as _json
        import shutil
        import tempfile
        import zipfile
        from pathlib import Path as _PPath

        tmp = _PPath(tempfile.mkdtemp(prefix="bashkar_pub_"))
        errores = []

        def _log(msg):
            self.after(0, lambda m=msg: self._lbl_excel.config(text=f"⏳ {m}"))

        try:
            # 1. XML-TEI
            _log("Exportando TEI…")
            try:
                from core.tei_engine import exportar_corpus_tei
                corpus_txt = getattr(ST, "corpus_txt", []) or []
                arts_tei = [{"id": f"art_{i:04d}", "texto": t}
                            for i, t in enumerate(corpus_txt)]
                if arts_tei:
                    exportar_corpus_tei(
                        arts_tei, tmp / "corpus.xml",
                        titulo=getattr(ST, "publicacion", ""),
                        fecha=getattr(ST, "periodo", ""),
                    )
            except Exception as e:
                errores.append(f"TEI: {e}")

            # 2. BibTeX
            _log("Exportando BibTeX…")
            try:
                from core.tei_engine import exportar_bibtex
                corpus_txt = getattr(ST, "corpus_txt", []) or []
                arts_bib = [{"id": f"art_{i:04d}", "texto": t}
                            for i, t in enumerate(corpus_txt)]
                if arts_bib:
                    exportar_bibtex(arts_bib, tmp / "corpus.bib")
            except Exception as e:
                errores.append(f"BibTeX: {e}")

            # 3. CSV entidades
            _log("Exportando entidades CSV…")
            try:
                from core.ner_engine import exportar_csv as _ner_csv
                ner = getattr(ST, "indice_ner_global", {}) or {}
                if ner:
                    _ner_csv(ner, tmp / "entidades.csv")
            except Exception as e:
                errores.append(f"CSV NER: {e}")

            # 4. Bitácora Markdown
            _log("Exportando bitácora…")
            try:
                eng = self._bitacora_engine()
                if eng is not None:
                    eng.exportar_markdown(
                        tmp / "bitacora.md",
                        publicacion=getattr(ST, "publicacion", ""))
            except Exception as e:
                errores.append(f"Bitácora: {e}")

            # 5. Metadatos JSON
            _log("Escribiendo metadatos…")
            meta = {
                "publicacion":   getattr(ST, "publicacion", ""),
                "periodo":       getattr(ST, "periodo", ""),
                "investigador":  getattr(ST, "investigador", ""),
                "institucion":   getattr(ST, "institucion", ""),
                "fecha_export":  datetime.now().isoformat(),
                "bashkar_version": APP_VERSION,
                "n_archivos":    len(getattr(ST, "archivos_sel", []) or []),
                "modulos_usados": [k for k, v in ST.estado_etapas.items() if v == "ready"],
            }
            (tmp / "metadatos.json").write_text(
                _json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

            # 6. METHODS.md
            _log("Generando METHODS.md…")
            try:
                from core.methods_reporter import generar_methods_md
                cfg_methods = {
                    "publicacion":    getattr(ST, "publicacion", ""),
                    "periodo":        getattr(ST, "periodo", ""),
                    "investigador":   getattr(ST, "investigador", ""),
                    "institucion":    getattr(ST, "institucion", ""),
                    "bashkar_version": APP_VERSION,
                    "dpi":            getattr(ST, "dpi", "150"),
                    "lang":           getattr(ST, "lang", "spa"),
                    "lematizar":      getattr(ST, "lematizar", True),
                    "modelos_etapa":  getattr(ST, "modelos_etapa", {}),
                    "archivos_sel":   getattr(ST, "archivos_sel", []),
                }
                stats_methods = {
                    "n_paginas":   meta.get("n_archivos", 0),
                    "n_palabras":  0,
                    "n_articulos": (len(ST.df_articulos)
                                   if ST.df_articulos is not None else 0),
                    "n_entidades": sum(len(v) for v in
                                      (getattr(ST, "indice_ner_global", {}) or {}).values()
                                      if isinstance(v, dict)),
                }
                generar_methods_md(cfg_methods, stats_methods, tmp / "METHODS.md")
            except Exception as e:
                errores.append(f"METHODS.md: {e}")

            # 7. Empaquetar ZIP
            _log("Comprimiendo…")
            with zipfile.ZipFile(dest_zip, "w", zipfile.ZIP_DEFLATED) as zf:
                for f in tmp.rglob("*"):
                    if f.is_file():
                        zf.write(f, f.relative_to(tmp))

        finally:
            shutil.rmtree(tmp, ignore_errors=True)

        def _fin():
            if errores:
                self._lbl_excel.config(
                    text=f"⚠ Paquete con {len(errores)} advertencias")
                messagebox.showwarning(
                    "Paquete generado (con advertencias)",
                    f"ZIP guardado en:\n{dest_zip}\n\n"
                    f"Advertencias:\n" + "\n".join(f"• {e}" for e in errores))
            else:
                self._lbl_excel.config(text="✅ Paquete de publicación listo")
                messagebox.showinfo(
                    "Paquete listo ✅",
                    f"Paquete generado exitosamente:\n{dest_zip}\n\n"
                    f"Contiene: corpus.xml · corpus.bib · entidades.csv · "
                    f"bitacora.md · METHODS.md · metadatos.json")
            import os
            os.startfile(str(Path(dest_zip).parent))
        self.after(0, _fin)

    # ── Diff visual Normalizar ────────────────────────────────────────────────

    def _norm_ver_diff(self):
        """Abre ventana con diff coloreado entre OCR crudo y versión manual."""
        idx = self._norm_idx_actual
        if idx < 0 or idx >= len(self._norm_bloques):
            messagebox.showwarning("Sin selección", "Selecciona una página primero.")
            return
        import difflib
        b = self._norm_bloques[idx]
        crudo  = b.get("ocr_crudo", "").splitlines(keepends=True)
        manual = (b.get("norm_usuario") or b.get("ocr_crudo", "")).splitlines(keepends=True)

        diff = list(difflib.unified_diff(crudo, manual,
                                          fromfile="OCR crudo",
                                          tofile="Manual",
                                          lineterm="", n=2))
        if not diff:
            messagebox.showinfo("Sin cambios",
                                "No hay diferencias entre el texto crudo y el manual.")
            return

        win, diff_content = self._mk_glass_toplevel(
            f"🔍 Cambios — {b.get('pagina', '')}", ancho=700, alto=500)

        txt = scrolledtext.ScrolledText(
            diff_content, bg="#0D1117", fg=TXT_PRI, font=("Courier New", 9),
            relief="flat", wrap="none")
        txt.pack(fill="both", expand=True, padx=8, pady=8)
        txt.tag_configure("add", foreground="#3FB950", background="#0F2B1A")
        txt.tag_configure("del", foreground="#F85149", background="#2B0F0F")
        txt.tag_configure("hdr", foreground="#58A6FF")
        txt.tag_configure("ctx", foreground="#8B949E")

        for line in diff:
            if line.startswith("+++") or line.startswith("---"):
                txt.insert("end", line + "\n", "hdr")
            elif line.startswith("+"):
                txt.insert("end", line + "\n", "add")
            elif line.startswith("-"):
                txt.insert("end", line + "\n", "del")
            elif line.startswith("@@"):
                txt.insert("end", line + "\n", "hdr")
            else:
                txt.insert("end", line + "\n", "ctx")
        txt.config(state="disabled")

    # ── Dataset HTR Normalizar ────────────────────────────────────────────────

    def _norm_exportar_ground_truth(self):
        """Exporta pares (imagen, texto) para reentrenamiento Kraken."""
        if not ST.out_dir:
            messagebox.showwarning("Sin corpus", "Carga un corpus primero.")
            return
        numero = self._norm_var_numero.get() if hasattr(self, "_norm_var_numero") else ""
        if not numero:
            messagebox.showwarning("Sin número", "Selecciona un número en Normalizar.")
            return
        from tkinter import filedialog
        dest = filedialog.askdirectory(title="Carpeta destino del dataset HTR")
        if not dest:
            return

        txt_dir = Path(ST.out_dir) / "03_ocr" / numero
        img_dir = Path(ST.out_dir) / "02_imagenes" / numero
        out_dir = Path(dest) / f"ground_truth_{numero}"

        self._lbl_norm_estado.config(text="⏳ Exportando dataset HTR…", fg=TXT_SEC)

        def _run():
            from core.kraken_trainer import exportar_ground_truth
            def _cb(n, total, msg):
                self.after(0, lambda: self._lbl_norm_estado.config(
                    text=f"⏳ {n}/{total}: {msg}", fg=TXT_SEC))
            try:
                res = exportar_ground_truth(txt_dir, img_dir, out_dir, callback=_cb)
                self.after(0, lambda r=res: (
                    self._lbl_norm_estado.config(
                        text=f"✅ {r['pares']} pares exportados a {r['out_dir']}",
                        fg=VERDE),
                    messagebox.showinfo(
                        "Dataset HTR exportado",
                        f"Pares exportados: {r['pares']}\n"
                        f"Omitidos: {r['omitidos']}\n\n"
                        f"Carpeta:\n{r['out_dir']}\n\n"
                        f"Usa con:\n"
                        f"  ketos train -f binary "
                        f"--load {r['out_dir']}/manifest.txt\n"
                        f"  (en D:\\kraken_env)")
                ))
            except Exception as e:
                self.after(0, lambda err=str(e): self._lbl_norm_estado.config(
                    text=f"❌ {err}", fg="#F85149"))
        threading.Thread(target=_run, daemon=True).start()


# ══════════════════════════════════════════════════════════════════════════════
# COMPARAR (v17) — comparación multi-proyecto
# ══════════════════════════════════════════════════════════════════════════════

    def _build_comp2(self):
        import tkinter as tk
        from tkinter import ttk
        self._tab_comp2.columnconfigure(0, weight=1)
        self._tab_comp2.rowconfigure(2, weight=1)
        ttk.Label(self._tab_comp2, text="Comparación multi-proyecto", style="H.TLabel").grid(
            row=0, column=0, sticky="w", padx=20, pady=(18, 4))
        ttk.Label(self._tab_comp2, text="Compara entidades, vocabulario y tópicos entre distintos proyectos .bashkar",
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", padx=20, pady=(0, 8))

        cuerpo = ttk.Frame(self._tab_comp2)
        cuerpo.grid(row=2, column=0, sticky="nsew", padx=20, pady=4)
        cuerpo.columnconfigure(0, weight=1)
        cuerpo.rowconfigure(3, weight=1)

        # Lista de proyectos
        ttk.Label(cuerpo, text="Proyectos a comparar:").grid(row=0, column=0, sticky="w")
        self._comp2_lista_var = tk.Variable(value=[])
        self._comp2_listbox = tk.Listbox(cuerpo, listvariable=self._comp2_lista_var,
                                          height=6, bg="#1e293b", fg="#e0e0e0",
                                          selectmode=tk.MULTIPLE, font=("Consolas", 10))
        self._comp2_listbox.grid(row=1, column=0, sticky="ew", pady=4)

        bframe = ttk.Frame(cuerpo)
        bframe.grid(row=2, column=0, sticky="w", pady=4)
        ttk.Button(bframe, text="+ Agregar proyecto",
                   command=self._comp2_agregar).pack(side="left", padx=2)
        ttk.Button(bframe, text="✕ Quitar selección",
                   command=self._comp2_quitar).pack(side="left", padx=2)
        ttk.Button(bframe, text="▶ Comparar",
                   command=self._comp2_ejecutar).pack(side="left", padx=10)
        ttk.Button(bframe, text="🌐 Ver HTML",
                   command=self._comp2_ver_html).pack(side="left", padx=2)

        self._comp2_log = tk.Text(cuerpo, height=14, state="disabled",
                                   bg="#0f172a", fg="#94a3b8",
                                   font=("Consolas", 9), wrap="word")
        self._comp2_log.grid(row=3, column=0, sticky="nsew", pady=8)

    def _comp2_agregar(self):
        from tkinter import filedialog
        ruta = filedialog.askopenfilename(
            title="Seleccionar proyecto .bashkar",
            filetypes=[("Bashkar", "*.bashkar"), ("Todos", "*.*")]
        )
        if ruta and ruta not in ST.comparar_rutas:
            ST.comparar_rutas.append(ruta)
            self._comp2_lista_var.set(ST.comparar_rutas)

    def _comp2_quitar(self):
        sel = self._comp2_listbox.curselection()
        for i in reversed(sel):
            ST.comparar_rutas.pop(i)
        self._comp2_lista_var.set(ST.comparar_rutas)

    def _comp2_ejecutar(self):
        if len(ST.comparar_rutas) < 2:
            from tkinter import messagebox
            messagebox.showwarning("Comparar", "Agrega al menos 2 proyectos.")
            return
        import threading
        self._comp2_log.config(state="normal")
        self._comp2_log.delete("1.0", "end")
        self._comp2_log.config(state="disabled")
        threading.Thread(target=self._comp2_worker, daemon=True).start()

    def _comp2_worker(self):
        def _log(msg):
            self.after(0, lambda m=msg: self._comp2_log_insert(m))
        try:
            from pathlib import Path

            from core.comparador import (
                exportar_reporte_html,
                generar_reporte_comparativo,
            )
            nombres = [Path(r).stem for r in ST.comparar_rutas]
            _log("Generando reporte comparativo...")
            rep = generar_reporte_comparativo(ST.comparar_rutas, nombres, callback=_log)
            ST.reporte_comparativo = rep
            out = Path.home() / "Documents" / "BashkarStation" / "comparativa.html"
            exportar_reporte_html(rep, out)
            _log(f"\n✅ Reporte guardado: {out}")
            self.after(0, lambda: setattr(self, "_comp2_html_path", str(out)))
        except Exception as e:
            _log(f"\n❌ Error: {e}")

    def _comp2_log_insert(self, msg):
        self._comp2_log.config(state="normal")
        self._comp2_log.insert("end", msg + "\n")
        self._comp2_log.see("end")
        self._comp2_log.config(state="disabled")

    def _comp2_ver_html(self):
        import os
        import webbrowser
        ruta = getattr(self, "_comp2_html_path", None)
        if ruta and os.path.exists(ruta):
            webbrowser.open(f"file:///{ruta.replace(chr(92), '/')}")
        else:
            from tkinter import messagebox
            messagebox.showinfo("Comparar", "Ejecuta la comparación primero.")


# ══════════════════════════════════════════════════════════════════════════════
# INTERTEXTUALIDAD (v17)
# ══════════════════════════════════════════════════════════════════════════════

    def _build_intxt(self):
        import tkinter as tk
        from tkinter import ttk
        self._tab_intxt.columnconfigure(0, weight=1)
        self._tab_intxt.rowconfigure(2, weight=1)
        ttk.Label(self._tab_intxt, text="Análisis intertextual", style="H.TLabel").grid(
            row=0, column=0, sticky="w", padx=20, pady=(18, 4))
        ttk.Label(self._tab_intxt, text="Detecta citas compartidas, similitud textual y conexiones entre artículos",
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", padx=20, pady=(0, 8))

        cuerpo = ttk.Frame(self._tab_intxt)
        cuerpo.grid(row=2, column=0, sticky="nsew", padx=20, pady=4)
        cuerpo.columnconfigure(1, weight=1)
        cuerpo.rowconfigure(3, weight=1)

        # Parámetros
        ttk.Label(cuerpo, text="Umbral similitud (0–1):").grid(row=0, column=0, sticky="w", pady=4)
        self._intxt_umbral = ttk.Spinbox(cuerpo, from_=0.1, to=0.9, increment=0.05,
                                          width=8, format="%.2f")
        self._intxt_umbral.set("0.30")
        self._intxt_umbral.grid(row=0, column=1, sticky="w", padx=8)

        self._intxt_usar_llm = tk.BooleanVar(value=True)
        ttk.Checkbutton(cuerpo, text="Analizar pares con LLM (Claude)",
                        variable=self._intxt_usar_llm).grid(row=1, column=0, columnspan=2,
                                                              sticky="w", pady=2)

        bframe = ttk.Frame(cuerpo)
        bframe.grid(row=2, column=0, columnspan=2, sticky="w", pady=8)
        ttk.Button(bframe, text="▶ Analizar intertextualidad",
                   command=self._intxt_ejecutar).pack(side="left", padx=2)
        ttk.Button(bframe, text="🕸 Ver grafo HTML",
                   command=self._intxt_ver_grafo).pack(side="left", padx=4)

        self._intxt_log = tk.Text(cuerpo, height=16, state="disabled",
                                   bg="#0f172a", fg="#94a3b8",
                                   font=("Consolas", 9), wrap="word")
        self._intxt_log.grid(row=3, column=0, columnspan=2, sticky="nsew", pady=4)

    def _intxt_ejecutar(self):
        if not ST.seg_done:
            from tkinter import messagebox
            messagebox.showwarning("Intertextualidad", "Segmenta los artículos primero.")
            return
        import threading
        self._intxt_log.config(state="normal")
        self._intxt_log.delete("1.0", "end")
        self._intxt_log.config(state="disabled")
        threading.Thread(target=self._intxt_worker, daemon=True).start()

    def _intxt_worker(self):
        def _log(msg):
            self.after(0, lambda m=msg: self._intxt_log_insert(m))
        try:
            from pathlib import Path

            from core.intertextual_engine import (
                analizar_intertextualidad,
                exportar_grafo_intertextual,
            )
            articulos = {}
            if ST.df_articulos is not None:
                for _, row in ST.df_articulos.iterrows():
                    aid = str(row.get("id", row.name))
                    articulos[aid] = {
                        "texto_limpio": str(row.get("texto", "")),
                        "titulo": str(row.get("titulo", aid)),
                    }
            umbral = float(self._intxt_umbral.get())
            usar_llm = self._intxt_usar_llm.get()
            resultado = analizar_intertextualidad(
                articulos,
                api_key=ST.api_key,
                umbral_similitud=umbral,
                usar_llm=usar_llm,
                callback=_log,
            )
            ST.intertex_resultado = resultado
            citas = len(resultado.get("citas_compartidas", {}))
            pares = len(resultado.get("pares_similares", []))
            conn  = len(resultado.get("conexiones_llm", []))
            _log(f"\n✅ {citas} citas compartidas · {pares} pares similares · {conn} conexiones LLM")
            # Exportar grafo
            out = Path.home() / "Documents" / "BashkarStation" / "intertextual.html"
            exportar_grafo_intertextual(resultado, out)
            _log(f"Grafo: {out}")
            self.after(0, lambda: setattr(self, "_intxt_grafo_path", str(out)))
        except Exception as e:
            _log(f"\n❌ Error: {e}")

    def _intxt_log_insert(self, msg):
        self._intxt_log.config(state="normal")
        self._intxt_log.insert("end", msg + "\n")
        self._intxt_log.see("end")
        self._intxt_log.config(state="disabled")

    def _intxt_ver_grafo(self):
        import os
        import webbrowser
        ruta = getattr(self, "_intxt_grafo_path", None)
        if ruta and os.path.exists(ruta):
            webbrowser.open(f"file:///{ruta.replace(chr(92), '/')}")
        else:
            from tkinter import messagebox
            messagebox.showinfo("Intertextualidad", "Ejecuta el análisis primero.")


# ══════════════════════════════════════════════════════════════════════════════
# VALIDACIÓN HUMANA + CONFIANZA (v18)
# ══════════════════════════════════════════════════════════════════════════════

    def _build_valid(self):
        import tkinter as tk
        from tkinter import ttk
        self._tab_valid.columnconfigure(0, weight=1)
        self._tab_valid.rowconfigure(3, weight=1)
        ttk.Label(self._tab_valid, text="Validación humana y confianza", style="H.TLabel").grid(
            row=0, column=0, sticky="w", padx=20, pady=(18, 4))
        ttk.Label(self._tab_valid, text="Revisa y valida entidades por nivel de confianza (🟢 confiable · 🟡 revisar · 🔴 obligatorio)",
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", padx=20, pady=(0, 8))

        bframe = ttk.Frame(self._tab_valid)
        bframe.grid(row=2, column=0, sticky="w", padx=20, pady=4)
        ttk.Button(bframe, text="🔄 Calcular confianza del corpus",
                   command=self._valid_calcular).pack(side="left", padx=2)
        ttk.Button(bframe, text="✅ Marcar seleccionada como verificada",
                   command=self._valid_verificar).pack(side="left", padx=4)
        ttk.Button(bframe, text="✅✅ Marcar todas 🟢 como verificadas",
                   command=self._valid_verificar_todas_verdes).pack(side="left", padx=4)
        ttk.Button(bframe, text="💾 Guardar en base de conocimiento",
                   command=self._valid_guardar_kb).pack(side="left", padx=4)
        ttk.Button(bframe, text="📥 Exportar CSV",
                   command=self._valid_exportar_csv).pack(side="left", padx=4)

        cuerpo = ttk.Frame(self._tab_valid)
        cuerpo.grid(row=3, column=0, sticky="nsew", padx=20, pady=4)
        cuerpo.columnconfigure(0, weight=1)
        cuerpo.rowconfigure(0, weight=1)

        cols = ("Entidad", "Categoría", "Score", "Nivel", "Verificada")
        self._valid_tree = ttk.Treeview(cuerpo, columns=cols, show="headings", height=20)
        for col in cols:
            self._valid_tree.heading(col, text=col)
            self._valid_tree.column(col, width=150 if col != "Score" else 80)
        vsb = ttk.Scrollbar(cuerpo, orient="vertical", command=self._valid_tree.yview)
        self._valid_tree.configure(yscrollcommand=vsb.set)
        self._valid_tree.grid(row=0, column=0, sticky="nsew")
        vsb.grid(row=0, column=1, sticky="ns")

        # Tag colors
        self._valid_tree.tag_configure("green",  background="#052e16", foreground="#22c55e")
        self._valid_tree.tag_configure("yellow", background="#431407", foreground="#f59e0b")
        self._valid_tree.tag_configure("red",    background="#450a0a", foreground="#ef4444")

        self._valid_stat_var = tk.StringVar(value="")
        ttk.Label(self._tab_valid, textvariable=self._valid_stat_var,
                  style="Sub.TLabel").grid(row=4, column=0, sticky="w", padx=20, pady=4)

    def _valid_calcular(self):
        from core.confianza_engine import nivel_confianza, score_ner_entidad
        try:
            from conocimiento.base_conocimiento import buscar_entidad, inicializar_db
            inicializar_db()
            _kb = True
        except Exception:
            _kb = False

        self._valid_tree.delete(*self._valid_tree.get_children())
        n_verde = n_amarillo = n_rojo = 0

        # Construir fuente: primero intentar desde Repositorio (DB), luego memoria
        fuente_ner: dict = {}

        if ST.repo:
            try:
                rows = ST.repo.buscar_entidades()
                CAT_DISPLAY = {
                    "PER": "personas", "LOC": "lugares", "ORG": "organizaciones",
                    "EVE": "eventos_historicos", "OBRA": "obras_publicaciones",
                    "CARGO": "personas",
                }
                for row in rows:
                    cat = CAT_DISPLAY.get(row.get("categoria", ""), row.get("categoria", "otros"))
                    ent_text = str(row.get("texto", "")).strip()
                    art_id = row.get("articulo_id", "?")
                    conf = float(row.get("confianza", 0.75))
                    if cat not in fuente_ner:
                        fuente_ner[cat] = {}
                    if ent_text not in fuente_ner[cat]:
                        fuente_ner[cat][ent_text] = {"arts": [], "confianza": conf}
                    if art_id not in fuente_ner[cat][ent_text]["arts"]:
                        fuente_ner[cat][ent_text]["arts"].append(art_id)
            except Exception:
                fuente_ner = {}

        # Si el repo no tiene datos, usar el índice en memoria
        if not fuente_ner and ST.indice_ner_global:
            for cat, ents in ST.indice_ner_global.items():
                if not isinstance(ents, dict):
                    continue
                fuente_ner[cat] = {ent: {"arts": arts, "confianza": 0.75}
                                   for ent, arts in ents.items()}

        for cat, ents in fuente_ner.items():
            for ent, meta in ents.items():
                conf_base = meta.get("confianza", 0.75) if isinstance(meta, dict) else 0.75
                en_kb = bool(_kb and buscar_entidad(ent, cat)) if _kb else False
                sc = score_ner_entidad(
                    en_kb=en_kb,
                    verificada=False,
                    spacy_conf=conf_base,
                    llm_conf=min(conf_base + 0.05, 1.0),
                )
                nivel = nivel_confianza(sc)
                emojis = {"green": "🟢", "yellow": "🟡", "red": "🔴"}
                self._valid_tree.insert("", "end",
                    values=(ent, cat, f"{sc:.2f}", emojis.get(nivel, nivel), "No"),
                    tags=(nivel,))
                if nivel == "green": n_verde += 1
                elif nivel == "yellow": n_amarillo += 1
                else: n_rojo += 1

        origen = "DB" if ST.repo and fuente_ner else "memoria"
        self._valid_stat_var.set(
            f"🟢 {n_verde} confiables · 🟡 {n_amarillo} revisar · 🔴 {n_rojo} validar  [fuente: {origen}]")

    def _valid_verificar(self):
        sel = self._valid_tree.selection()
        if not sel:
            return
        for item in sel:
            vals = list(self._valid_tree.item(item, "values"))
            vals[4] = "Sí"
            vals[2] = "1.00"
            vals[3] = "🟢"
            self._valid_tree.item(item, values=vals, tags=("green",))

    def _valid_guardar_kb(self):
        try:
            from conocimiento.base_conocimiento import inicializar_db, registrar_entidad
            inicializar_db()
        except Exception as e:
            from tkinter import messagebox
            messagebox.showerror("KB", str(e))
            return
        n = 0
        for item in self._valid_tree.get_children():
            vals = self._valid_tree.item(item, "values")
            ent, cat, score, _, verif = vals
            registrar_entidad(
                nombre=ent,
                categoria=cat,
                proyecto=str(self._proyecto_ruta or ""),
                confianza=float(score),
                verificada=(verif == "Sí"),
            )
            n += 1
        from tkinter import messagebox
        messagebox.showinfo("KB", f"✅ {n} entidades guardadas en la base de conocimiento.")
        # Marcar etapa validación como lista en semáforos
        ST.marcar_etapa("anal", "ready")
        self._actualizar_badges()

    def _valid_verificar_todas_verdes(self):
        """Marca como verificadas todas las entidades con nivel 🟢."""
        n = 0
        for item in self._valid_tree.get_children():
            vals = list(self._valid_tree.item(item, "values"))
            if vals[3] == "🟢" and vals[4] != "Sí":
                vals[4] = "Sí"
                self._valid_tree.item(item, values=vals)
                n += 1
        if hasattr(self, "_valid_stat_var"):
            prev = self._valid_stat_var.get()
            self._valid_stat_var.set(f"{prev}  ·  +{n} verificadas manualmente")

    def _valid_exportar_csv(self):
        """Exporta la tabla de validación a CSV."""
        items = self._valid_tree.get_children()
        if not items:
            messagebox.showwarning("Sin datos", "Calcula la confianza primero.")
            return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")],
            initialfile="entidades_validadas.csv",
            title="Exportar validación")
        if not dest:
            return
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            w.writerow(["entidad", "categoria", "score", "nivel", "verificada"])
            for item in items:
                w.writerow(self._valid_tree.item(item, "values"))
        messagebox.showinfo("Exportado", f"✅ {len(items)} entidades exportadas:\n{dest}")


# ══════════════════════════════════════════════════════════════════════════════
# COLABORACIÓN (v19)
# ══════════════════════════════════════════════════════════════════════════════

    def _build_colab(self):
        import tkinter as tk
        from tkinter import ttk
        self._tab_colab.columnconfigure(0, weight=1)
        self._tab_colab.rowconfigure(3, weight=1)
        ttk.Label(self._tab_colab, text="Colaboración y trazabilidad", style="H.TLabel").grid(
            row=0, column=0, sticky="w", padx=20, pady=(18, 4))
        ttk.Label(self._tab_colab, text="Exporta parches de tus cambios · importa cambios de colegas · visualiza contribuciones",
                  style="Sub.TLabel").grid(row=1, column=0, sticky="w", padx=20, pady=(0, 8))

        bframe = ttk.Frame(self._tab_colab)
        bframe.grid(row=2, column=0, sticky="w", padx=20, pady=4)
        ttk.Button(bframe, text="📤 Exportar parche",
                   command=self._colab_exportar).pack(side="left", padx=2)
        ttk.Button(bframe, text="📥 Importar parche",
                   command=self._colab_importar).pack(side="left", padx=4)
        ttk.Button(bframe, text="📋 Ver trazabilidad",
                   command=self._colab_trazabilidad).pack(side="left", padx=4)
        ttk.Button(bframe, text="🌐 HTML trazabilidad",
                   command=self._colab_html_trazabilidad).pack(side="left", padx=4)

        self._colab_log = tk.Text(self._tab_colab, height=18, state="disabled",
                                   bg="#0f172a", fg="#94a3b8",
                                   font=("Consolas", 9), wrap="word")
        self._colab_log.grid(row=3, column=0, sticky="nsew", padx=20, pady=8)

    def _colab_log_insert(self, msg):
        self._colab_log.config(state="normal")
        self._colab_log.insert("end", msg + "\n")
        self._colab_log.see("end")
        self._colab_log.config(state="disabled")

    def _colab_exportar(self):
        if not self._proyecto_ruta:
            from tkinter import messagebox
            messagebox.showwarning("Colaborar", "Abre o guarda un proyecto primero.")
            return
        from tkinter import filedialog, simpledialog
        investigador = simpledialog.askstring("Investigador", "Tu nombre:",
                                               initialvalue="") or "anon"
        notas = simpledialog.askstring("Notas", "Notas del parche:", initialvalue="") or ""
        dest = filedialog.asksaveasfilename(
            title="Guardar parche",
            defaultextension=".bashkar.patch",
            filetypes=[("Parche Bashkar", "*.bashkar.patch"), ("Todos", "*.*")],
        )
        if not dest:
            return
        import threading
        threading.Thread(
            target=self._colab_exportar_worker,
            args=(investigador, notas, dest), daemon=True
        ).start()

    def _colab_exportar_worker(self, investigador, notas, dest):
        def _log(msg): self.after(0, lambda m=msg: self._colab_log_insert(m))
        try:
            import json
            from pathlib import Path

            from core.colaboracion import crear_parche, exportar_parche
            actual = json.loads(Path(self._proyecto_ruta).read_text(encoding="utf-8"))
            # Construir estado actual con NER actualizado
            actual_mod = dict(actual)
            actual_mod["indice_ner_global"] = ST.indice_ner_global
            parche = crear_parche(actual, actual_mod, investigador, notas)
            exportar_parche(parche, Path(dest))
            _log(f"✅ Parche exportado: {dest}")
        except Exception as e:
            _log(f"❌ Error: {e}")

    def _colab_importar(self):
        from tkinter import filedialog
        ruta = filedialog.askopenfilename(
            title="Seleccionar parche",
            filetypes=[("Parche Bashkar", "*.bashkar.patch"), ("Todos", "*.*")],
        )
        if not ruta:
            return
        import threading
        threading.Thread(
            target=self._colab_importar_worker, args=(ruta,), daemon=True
        ).start()

    def _colab_importar_worker(self, ruta):
        def _log(msg): self.after(0, lambda m=msg: self._colab_log_insert(m))
        try:
            import json
            from pathlib import Path

            from core.colaboracion import aplicar_parche, cargar_parche
            parche = cargar_parche(Path(ruta))
            investigador = parche.get("_investigador", "?")
            fecha        = parche.get("_fecha", "")[:16]
            notas        = parche.get("_notas", "")
            n_cambios    = len([k for k in parche if not k.startswith("_")])

            _log("─── Parche recibido ───────────────────────────────────")
            _log(f"  De:     {investigador}")
            _log(f"  Fecha:  {fecha}")
            _log(f"  Notas:  {notas}")
            _log(f"  Secciones modificadas: {n_cambios}")
            _log("───────────────────────────────────────────────────────")

            # Vista previa de claves modificadas
            for k in list(parche.keys())[:10]:
                if not k.startswith("_"):
                    _log(f"  · {k}")

            if not self._proyecto_ruta:
                _log("⚠ Abre un proyecto primero para aplicar el parche.")
                return

            # Confirmación desde hilo principal
            def _confirmar():
                from tkinter import messagebox
                ok = messagebox.askyesno(
                    "Aplicar parche",
                    f"Parche de: {investigador}\n"
                    f"Fecha: {fecha}\n"
                    f"Notas: {notas}\n"
                    f"Secciones: {n_cambios}\n\n"
                    f"Se creará un backup automático del proyecto.\n"
                    f"¿Aplicar?")
                if ok:
                    import threading
                    threading.Thread(target=_aplicar, daemon=True).start()
                else:
                    self._colab_log_insert("Aplicación cancelada.")

            def _aplicar():
                try:
                    actual = json.loads(Path(self._proyecto_ruta).read_text(encoding="utf-8"))
                    # Backup automático antes de aplicar
                    backup = Path(self._proyecto_ruta).with_suffix(".bashkar.bak")
                    backup.write_text(json.dumps(actual, indent=2, ensure_ascii=False),
                                      encoding="utf-8")
                    _log(f"Backup guardado: {backup.name}")

                    actualizado = aplicar_parche(actual, parche, callback=_log)
                    ST.indice_ner_global = actualizado.get(
                        "indice_ner_global", ST.indice_ner_global)
                    Path(self._proyecto_ruta).write_text(
                        json.dumps(actualizado, indent=2, ensure_ascii=False),
                        encoding="utf-8")
                    _log("✅ Parche aplicado y proyecto guardado.")
                except Exception as ex:
                    _log(f"❌ Error al aplicar: {ex}")

            self.after(0, _confirmar)

        except Exception as e:
            _log(f"❌ Error al cargar parche: {e}")

    def _colab_trazabilidad(self):
        if not self._proyecto_ruta:
            from tkinter import messagebox
            messagebox.showwarning("Colaborar", "Abre un proyecto primero.")
            return
        import json
        from pathlib import Path

        from core.colaboracion import reporte_trazabilidad
        data = json.loads(Path(self._proyecto_ruta).read_text(encoding="utf-8"))
        rep = reporte_trazabilidad(data)
        self._colab_log.config(state="normal")
        self._colab_log.delete("1.0", "end")
        self._colab_log.insert("end", rep)
        self._colab_log.config(state="disabled")

    def _colab_html_trazabilidad(self):
        if not self._proyecto_ruta:
            return
        import json
        import webbrowser
        from pathlib import Path

        from core.colaboracion import exportar_trazabilidad_html
        data = json.loads(Path(self._proyecto_ruta).read_text(encoding="utf-8"))
        out = Path.home() / "Documents" / "BashkarStation" / "trazabilidad.html"
        exportar_trazabilidad_html(data, out)
        webbrowser.open(f"file:///{str(out).replace(chr(92), '/')}")



    # ══════════════════════════════════════════════════════════════════════════
    # LINGÜÍSTICA COMPUTACIONAL — sintaxis, correferencia, morfología, emociones
    # ══════════════════════════════════════════════════════════════════════════

    def _build_ling(self):
        pad = tk.Frame(self._frames_pagina["ling"], bg=CONTENT_BG, padx=16, pady=12)
        pad.pack(fill="both", expand=True)
        tk.Label(pad, text="Lingüística Computacional", bg=CONTENT_BG,
                 fg="#FFFFFF", font=("Segoe UI", 14, "bold")).pack(anchor="w")
        tk.Label(pad,
                 text="Sintaxis, correferencia, morfología histórica, emociones, encuadre, "
                      "polaridad, revisión NER y validación metodológica (Kappa).",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9),
                 wraplength=760, justify="left").pack(anchor="w", pady=(0, 10))

        nb_ling = ttk.Notebook(pad)
        nb_ling.pack(fill="both", expand=True)
        self._nb_ling = nb_ling

        # ── Pestaña 1: Análisis sintáctico ────────────────────────────────────
        frm_sint = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_sint, text="  🌲 Sintaxis  ")

        ctrl_sint = tk.Frame(frm_sint, bg=CONTENT_BG)
        ctrl_sint.pack(fill="x", padx=8, pady=(8, 4))
        tk.Label(ctrl_sint, text="Patrón sintáctico:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        self._var_ling_patron = tk.StringVar(value="verbo_sujeto")
        _patrones = [
            ("verbo_sujeto",  "Verbo + sujeto"),
            ("verbo_objeto",  "Verbo + objeto"),
            ("sustantivo_adj","Sust. + adjetivo"),
            ("entidad_verbo", "Entidad como sujeto"),
            ("negacion",      "Negaciones"),
        ]
        for val, lbl in _patrones:
            ttk.Radiobutton(ctrl_sint, text=lbl, variable=self._var_ling_patron,
                            value=val).pack(side="left", padx=3)

        tk.Label(ctrl_sint, text="  Máx. resultados:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(12, 4))
        self._var_ling_max = tk.IntVar(value=200)
        ttk.Spinbox(ctrl_sint, from_=50, to=1000, textvariable=self._var_ling_max,
                    width=6).pack(side="left")

        bf_sint = tk.Frame(frm_sint, bg=CONTENT_BG)
        bf_sint.pack(fill="x", padx=8, pady=(0, 6))
        self._btn_ling_sint = ttk.Button(bf_sint, text="▶  Analizar concordancias sintácticas",
                                          style="P.TButton",
                                          command=self._ling_concordancias)
        self._btn_ling_sint.pack(side="left", padx=(0, 8))
        ttk.Button(bf_sint, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_sint_csv).pack(side="left", padx=(0, 8))
        ttk.Button(bf_sint, text="🔗  Relaciones SVO", style="S.TButton",
                   command=self._ling_relaciones).pack(side="left", padx=(0, 8))
        self._lbl_ling_sint = tk.Label(frm_sint, text="", bg=CONTENT_BG,
                                        fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_sint.pack(anchor="w", padx=8)

        cols_sint = ("patron", "principal", "secundario", "descripcion", "fragmento")
        self._tv_ling_sint = ttk.Treeview(frm_sint, columns=cols_sint,
                                           show="headings", height=16)
        for cid, txt, w in [("patron","Patrón",110),("principal","Elem. 1",140),
                              ("secundario","Elem. 2",140),
                              ("descripcion","Descripción",200),
                              ("fragmento","Fragmento",340)]:
            self._tv_ling_sint.heading(cid, text=txt)
            self._tv_ling_sint.column(cid, width=w, anchor="w")
        sv_s = ttk.Scrollbar(frm_sint, orient="vertical",
                              command=self._tv_ling_sint.yview)
        self._tv_ling_sint.configure(yscrollcommand=sv_s.set)
        sv_s.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_sint.pack(fill="both", expand=True, padx=6, pady=(0, 4))
        self._ling_sint_resultados: list = []

        # ── Pestaña 2: Relaciones SVO ─────────────────────────────────────────
        frm_svo = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_svo, text="  🔗 Relaciones SVO  ")

        ctrl_svo = tk.Frame(frm_svo, bg=CONTENT_BG)
        ctrl_svo.pack(fill="x", padx=8, pady=(8, 4))
        self._var_ling_solo_ent = tk.BooleanVar(value=True)
        ttk.Checkbutton(ctrl_svo, text="Solo entidades nombradas como sujeto/objeto",
                        variable=self._var_ling_solo_ent).pack(side="left")
        tk.Label(ctrl_svo, text="  Confianza mínima:", bg=CONTENT_BG, fg=GRIS2,
                 font=("Segoe UI", 9)).pack(side="left", padx=(12, 4))
        self._var_ling_conf = tk.DoubleVar(value=0.5)
        ttk.Spinbox(ctrl_svo, from_=0.1, to=1.0, increment=0.1,
                    textvariable=self._var_ling_conf, width=5,
                    format="%.1f").pack(side="left")

        bf_svo = tk.Frame(frm_svo, bg=CONTENT_BG)
        bf_svo.pack(fill="x", padx=8, pady=(0, 6))
        self._btn_ling_svo = ttk.Button(bf_svo, text="▶  Extraer relaciones",
                                         style="P.TButton",
                                         command=self._ling_relaciones)
        self._btn_ling_svo.pack(side="left", padx=(0, 8))
        ttk.Button(bf_svo, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_svo_csv).pack(side="left", padx=(0, 8))
        ttk.Button(bf_svo, text="📊  Agrupar por verbo", style="S.TButton",
                   command=self._ling_svo_agrupar).pack(side="left", padx=(0, 8))
        self._lbl_ling_svo = tk.Label(frm_svo, text="", bg=CONTENT_BG,
                                       fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_svo.pack(anchor="w", padx=8)

        cols_svo = ("sujeto", "suj_tipo", "relacion", "objeto", "obj_tipo",
                    "confianza", "oracion")
        self._tv_ling_svo = ttk.Treeview(frm_svo, columns=cols_svo,
                                          show="headings", height=16)
        for cid, txt, w in [("sujeto","Sujeto",150),("suj_tipo","Tipo S",80),
                              ("relacion","Relación",110),("objeto","Objeto",150),
                              ("obj_tipo","Tipo O",80),("confianza","Conf.",60),
                              ("oracion","Oración",320)]:
            self._tv_ling_svo.heading(cid, text=txt)
            self._tv_ling_svo.column(cid, width=w, anchor="w")
        sv_svo = ttk.Scrollbar(frm_svo, orient="vertical",
                                command=self._tv_ling_svo.yview)
        self._tv_ling_svo.configure(yscrollcommand=sv_svo.set)
        sv_svo.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_svo.pack(fill="both", expand=True, padx=6, pady=(0, 4))
        self._ling_svo_resultados: list = []

        # ── Pestaña 3: Correferencia ──────────────────────────────────────────
        frm_coref = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_coref, text="  🔁 Correferencia  ")

        ctrl_coref = tk.Frame(frm_coref, bg=CONTENT_BG)
        ctrl_coref.pack(fill="x", padx=8, pady=(8, 4))
        tk.Label(ctrl_coref, text="Entidad a rastrear (vacío = todas):",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9)).pack(side="left", padx=(0, 6))
        self._ent_coref = ttk.Entry(ctrl_coref, width=24)
        self._ent_coref.pack(side="left", padx=(0, 12))

        bf_coref = tk.Frame(frm_coref, bg=CONTENT_BG)
        bf_coref.pack(fill="x", padx=8, pady=(0, 6))
        self._btn_ling_coref = ttk.Button(bf_coref, text="▶  Resolver correferencias",
                                           style="P.TButton",
                                           command=self._ling_coref)
        self._btn_ling_coref.pack(side="left", padx=(0, 8))
        ttk.Button(bf_coref, text="📊  Estadísticas corpus", style="S.TButton",
                   command=self._ling_coref_stats).pack(side="left", padx=(0, 8))
        self._lbl_ling_coref = tk.Label(frm_coref, text="", bg=CONTENT_BG,
                                         fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_coref.pack(anchor="w", padx=8)

        # Split: lista de cadenas izquierda, menciones derecha
        split_coref = tk.Frame(frm_coref, bg=CONTENT_BG)
        split_coref.pack(fill="both", expand=True, padx=6)

        izq_coref = tk.Frame(split_coref, bg=CONTENT_BG, width=220)
        izq_coref.pack(side="left", fill="y", padx=(0, 6))
        izq_coref.pack_propagate(False)
        tk.Label(izq_coref, text="Cadenas referenciales", bg=CONTENT_BG,
                 fg=GRIS2, font=("Segoe UI", 8, "bold")).pack(anchor="w")
        self._lb_coref = tk.Listbox(
            izq_coref, bg=CARD_BG, fg=TXT_SEC, selectbackground=AB_SEL,
            font=("Segoe UI", 9), relief="flat", activestyle="none")
        self._lb_coref.pack(fill="both", expand=True)
        self._lb_coref.bind("<<ListboxSelect>>", self._ling_coref_mostrar_cadena)

        der_coref = tk.Frame(split_coref, bg=CONTENT_BG)
        der_coref.pack(side="left", fill="both", expand=True)
        tk.Label(der_coref, text="Menciones", bg=CONTENT_BG,
                 fg=GRIS2, font=("Segoe UI", 8, "bold")).pack(anchor="w")
        cols_coref = ("texto", "tipo", "oracion")
        self._tv_coref_men = ttk.Treeview(der_coref, columns=cols_coref,
                                           show="headings", height=18)
        for cid, txt, w in [("texto","Mención",160),("tipo","Tipo",100),
                              ("oracion","Contexto",440)]:
            self._tv_coref_men.heading(cid, text=txt)
            self._tv_coref_men.column(cid, width=w, anchor="w")
        sv_c = ttk.Scrollbar(der_coref, orient="vertical",
                              command=self._tv_coref_men.yview)
        self._tv_coref_men.configure(yscrollcommand=sv_c.set)
        sv_c.pack(side="right", fill="y")
        self._tv_coref_men.pack(fill="both", expand=True)
        self._ling_coref_cadenas: list = []

        # ── Pestaña 4: Morfología histórica ──────────────────────────────────
        frm_morf = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_morf, text="  📜 Morfología histórica  ")

        bf_morf = tk.Frame(frm_morf, bg=CONTENT_BG)
        bf_morf.pack(fill="x", padx=8, pady=(8, 4))
        self._btn_ling_morf = ttk.Button(bf_morf, text="▶  Analizar corpus",
                                          style="P.TButton",
                                          command=self._ling_morf_analizar)
        self._btn_ling_morf.pack(side="left", padx=(0, 8))
        ttk.Button(bf_morf, text="📖  Ver glosario", style="S.TButton",
                   command=self._ling_morf_glosario).pack(side="left", padx=(0, 8))
        ttk.Button(bf_morf, text="🔍  Ver detalle", style="S.TButton",
                   command=self._ling_morf_detalle).pack(side="left", padx=(0, 8))
        ttk.Button(bf_morf, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_morf_csv).pack(side="left", padx=(0, 8))
        self._var_morf_normalizar = tk.BooleanVar(value=True)
        ttk.Checkbutton(bf_morf, text="Normalizar grafías históricas antes de analizar",
                        variable=self._var_morf_normalizar).pack(side="left", padx=(8, 0))
        self._lbl_ling_morf = tk.Label(frm_morf, text="", bg=CONTENT_BG,
                                        fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_morf.pack(anchor="w", padx=8)

        # Resumen del corpus
        frm_morf_stats = tk.Frame(frm_morf, bg=CONTENT_BG)
        frm_morf_stats.pack(fill="x", padx=8, pady=(0, 4))
        self._lbl_morf_resumen = tk.Label(frm_morf_stats,
                                           text="Analiza el corpus para ver la densidad de formas históricas.",
                                           bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9))
        self._lbl_morf_resumen.pack(anchor="w")

        cols_morf = ("doc", "tokens", "arcaismos", "score", "marcadores", "top_formas")
        self._tv_ling_morf = ttk.Treeview(frm_morf, columns=cols_morf,
                                           show="headings", height=16)
        for cid, txt, w in [("doc","Doc.",50),("tokens","Tokens",70),
                              ("arcaismos","Arcaísmos",80),("score","Score hist.",80),
                              ("marcadores","Marcadores",200),
                              ("top_formas","Formas top",320)]:
            self._tv_ling_morf.heading(cid, text=txt)
            self._tv_ling_morf.column(cid, width=w, anchor="w" if w > 100 else "center")
        # Color por densidad
        self._tv_ling_morf.tag_configure("alta",  background="#1A2F1A", foreground="#4EC9B0")
        self._tv_ling_morf.tag_configure("media", background="#2D2210", foreground="#F59E0B")
        self._tv_ling_morf.tag_configure("baja",  background=CONTENT_BG, foreground=TXT_PRI)
        sv_m = ttk.Scrollbar(frm_morf, orient="vertical",
                              command=self._tv_ling_morf.yview)
        self._tv_ling_morf.configure(yscrollcommand=sv_m.set)
        sv_m.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_morf.pack(fill="both", expand=True, padx=6, pady=(0, 4))
        self._ling_morf_datos: list = []

        # ── Pestaña 5: Árbol de dependencias ─────────────────────────────────
        frm_dep = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_dep, text="  🌳 Árbol dep.  ")

        bf_dep = tk.Frame(frm_dep, bg=CONTENT_BG)
        bf_dep.pack(fill="x", padx=8, pady=(8, 4))
        self._btn_ling_dep = ttk.Button(bf_dep, text="▶  Analizar oración",
                                         style="P.TButton",
                                         command=self._ling_dep_analizar)
        self._btn_ling_dep.pack(side="left", padx=(0, 8))
        ttk.Button(bf_dep, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_dep_csv).pack(side="left", padx=(0, 8))
        tk.Label(bf_dep, text="Máx. oraciones:", bg=CONTENT_BG,
                 fg=TXT_SEC, font=("Segoe UI", 9)).pack(side="left", padx=(12, 4))
        self._var_dep_max = tk.IntVar(value=20)
        ttk.Spinbox(bf_dep, from_=1, to=100, textvariable=self._var_dep_max,
                    width=5).pack(side="left")
        self._lbl_ling_dep = tk.Label(frm_dep, text="", bg=CONTENT_BG,
                                       fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_dep.pack(anchor="w", padx=8)

        # Selector de artículo
        sel_dep = tk.Frame(frm_dep, bg=CONTENT_BG)
        sel_dep.pack(fill="x", padx=8, pady=(0, 4))
        tk.Label(sel_dep, text="Artículo:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        self._var_dep_art = tk.IntVar(value=0)
        self._spn_dep_art = ttk.Spinbox(sel_dep, from_=0, to=9999,
                                         textvariable=self._var_dep_art, width=6)
        self._spn_dep_art.pack(side="left", padx=4)
        tk.Label(sel_dep, text="(índice en corpus)", bg=CONTENT_BG, fg=TXT_DIM,
                 font=("Segoe UI", 8)).pack(side="left")

        # Panel split: lista oraciones | detalle tokens
        split_dep = tk.Frame(frm_dep, bg=CONTENT_BG)
        split_dep.pack(fill="both", expand=True, padx=6)

        izq_dep = tk.Frame(split_dep, bg=CONTENT_BG)
        izq_dep.pack(side="left", fill="both", expand=False)
        tk.Label(izq_dep, text="Oraciones", bg=CONTENT_BG,
                 fg=GRIS2, font=("Segoe UI", 8, "bold")).pack(anchor="w")
        self._lb_dep = tk.Listbox(izq_dep, bg=CARD_BG, fg=TXT_PRI,
                                   selectbackground=AZ3, font=("Segoe UI", 9),
                                   width=42, height=18, relief="flat")
        sv_dep_lb = ttk.Scrollbar(izq_dep, orient="vertical",
                                   command=self._lb_dep.yview)
        self._lb_dep.configure(yscrollcommand=sv_dep_lb.set)
        sv_dep_lb.pack(side="right", fill="y")
        self._lb_dep.pack(fill="y", expand=True)
        self._lb_dep.bind("<<ListboxSelect>>", self._ling_dep_mostrar_tokens)

        der_dep = tk.Frame(split_dep, bg=CONTENT_BG)
        der_dep.pack(side="left", fill="both", expand=True, padx=(8, 0))
        tk.Label(der_dep, text="Tokens y dependencias", bg=CONTENT_BG,
                 fg=GRIS2, font=("Segoe UI", 8, "bold")).pack(anchor="w")
        cols_dep = ("texto", "lemma", "pos", "dep_es", "cabeza")
        self._tv_dep_tok = ttk.Treeview(der_dep, columns=cols_dep,
                                         show="headings", height=18)
        for cid, txt, w in [("texto","Token",100),("lemma","Lema",100),
                              ("pos","POS",60),("dep_es","Relación",140),
                              ("cabeza","Cabeza",100)]:
            self._tv_dep_tok.heading(cid, text=txt)
            self._tv_dep_tok.column(cid, width=w, anchor="w")
        # Tags POS
        self._tv_dep_tok.tag_configure("VERB", foreground="#4EC9B0")
        self._tv_dep_tok.tag_configure("NOUN", foreground="#4FC1FF")
        self._tv_dep_tok.tag_configure("PROPN", foreground="#CE9178")
        sv_dep_tok = ttk.Scrollbar(der_dep, orient="vertical",
                                    command=self._tv_dep_tok.yview)
        self._tv_dep_tok.configure(yscrollcommand=sv_dep_tok.set)
        sv_dep_tok.pack(side="right", fill="y")
        self._tv_dep_tok.pack(fill="both", expand=True)
        self._ling_dep_datos: list = []

        # ── Pestaña 6: Emociones y subjetividad ──────────────────────────────
        frm_emo = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_emo, text="  💭 Emociones  ")

        bf_emo = tk.Frame(frm_emo, bg=CONTENT_BG)
        bf_emo.pack(fill="x", padx=8, pady=(8, 4))
        self._btn_ling_emo = ttk.Button(bf_emo, text="▶  Analizar emociones",
                                         style="P.TButton",
                                         command=self._ling_emociones)
        self._btn_ling_emo.pack(side="left", padx=(0, 8))
        ttk.Button(bf_emo, text="📊  Graficar distribución", style="S.TButton",
                   command=self._ling_emo_graficar).pack(side="left", padx=(0, 8))
        ttk.Button(bf_emo, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_emo_csv).pack(side="left", padx=(0, 8))
        self._lbl_ling_emo = tk.Label(frm_emo, text="", bg=CONTENT_BG,
                                       fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_emo.pack(anchor="w", padx=8)

        # Resumen global
        frm_emo_resumen = tk.Frame(frm_emo, bg=CONTENT_BG)
        frm_emo_resumen.pack(fill="x", padx=8, pady=(0, 4))
        self._lbl_emo_resumen = tk.Label(frm_emo_resumen, text="",
                                          bg=CONTENT_BG, fg=TXT_SEC,
                                          font=("Segoe UI", 9), wraplength=700,
                                          justify="left")
        self._lbl_emo_resumen.pack(anchor="w")

        cols_emo = ("art_id", "emocion_dom", "subjetividad", "tipo_discurso",
                    "intensidad", "palabras_emo")
        self._tv_ling_emo = ttk.Treeview(frm_emo, columns=cols_emo,
                                          show="headings", height=15)
        for cid, txt, w in [("art_id","Artículo",90),
                              ("emocion_dom","Emoción dom.",110),
                              ("subjetividad","Subjetividad",95),
                              ("tipo_discurso","Tipo discurso",110),
                              ("intensidad","Intensidad",90),
                              ("palabras_emo","Palabras emocionales",340)]:
            self._tv_ling_emo.heading(cid, text=txt)
            self._tv_ling_emo.column(cid, width=w, anchor="w")
        sv_e = ttk.Scrollbar(frm_emo, orient="vertical",
                              command=self._tv_ling_emo.yview)
        self._tv_ling_emo.configure(yscrollcommand=sv_e.set)
        sv_e.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_emo.pack(fill="both", expand=True, padx=6, pady=(0, 4))
        self._ling_emo_datos: list = []

        # ── Pestaña 7: Encuadre (framing) ─────────────────────────────────────
        frm_frame = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_frame, text="  🖼 Encuadre  ")

        bf_frame = tk.Frame(frm_frame, bg=CONTENT_BG)
        bf_frame.pack(fill="x", padx=8, pady=(8, 4))
        self._btn_ling_frame = ttk.Button(
            bf_frame, text="▶  Analizar encuadres del corpus",
            style="P.TButton", command=self._ling_frames)
        self._btn_ling_frame.pack(side="left", padx=(0, 8))
        ttk.Button(bf_frame, text="📊  Graficar distribución", style="S.TButton",
                   command=self._ling_frames_graficar).pack(side="left", padx=(0, 8))
        ttk.Button(bf_frame, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_frames_csv).pack(side="left", padx=(0, 8))
        self._lbl_ling_frame = tk.Label(frm_frame, text="", bg=CONTENT_BG,
                                         fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_frame.pack(anchor="w", padx=8)

        tk.Label(frm_frame,
                 text="Encuadre periodístico (Media Frames Corpus adaptado a prensa "
                      "ilustrada 1930s): desde qué ÁNGULO cubre cada artículo su tema.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9),
                 wraplength=720, justify="left").pack(anchor="w", padx=8, pady=(0, 4))
        self._lbl_frame_resumen = tk.Label(frm_frame, text="",
                                            bg=CONTENT_BG, fg=TXT_SEC,
                                            font=("Segoe UI", 9), wraplength=720,
                                            justify="left")
        self._lbl_frame_resumen.pack(anchor="w", padx=8, pady=(0, 4))

        cols_frame = ("art_id", "frame", "etiqueta", "porcentaje", "secundario",
                      "marcadores")
        self._tv_ling_frame = ttk.Treeview(frm_frame, columns=cols_frame,
                                            show="headings", height=15)
        for cid, txt, w in [("art_id","Artículo",90),("frame","Frame dom.",110),
                              ("etiqueta","Etiqueta",230),
                              ("porcentaje","% dom.",70),
                              ("secundario","Frame 2º",110),
                              ("marcadores","N marcadores",90)]:
            self._tv_ling_frame.heading(cid, text=txt)
            self._tv_ling_frame.column(cid, width=w, anchor="w")
        sv_f = ttk.Scrollbar(frm_frame, orient="vertical",
                             command=self._tv_ling_frame.yview)
        self._tv_ling_frame.configure(yscrollcommand=sv_f.set)
        sv_f.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_frame.pack(fill="both", expand=True, padx=6, pady=(0, 4))
        self._ling_frame_datos: list = []
        self._ling_frame_corpus: dict = {}

        # ── Pestaña 8: Polaridad discriminante ────────────────────────────────
        frm_pol = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_pol, text="  ⚖ Polaridad  ")

        bf_pol = tk.Frame(frm_pol, bg=CONTENT_BG)
        bf_pol.pack(fill="x", padx=8, pady=(8, 4))
        self._btn_ling_pol = ttk.Button(
            bf_pol, text="▶  Analizar polaridad del corpus",
            style="P.TButton", command=self._ling_polaridad)
        self._btn_ling_pol.pack(side="left", padx=(0, 8))
        ttk.Button(bf_pol, text="💾  Exportar CSV", style="S.TButton",
                   command=self._ling_pol_csv).pack(side="left", padx=(0, 8))
        self._lbl_ling_pol = tk.Label(frm_pol, text="", bg=CONTENT_BG,
                                      fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_pol.pack(anchor="w", padx=8)

        tk.Label(frm_pol,
                 text="Polaridad pos/neg/neutro discriminante (complemento al análisis "
                      "de 8 emociones, que sesga a «confianza»).",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9),
                 wraplength=720, justify="left").pack(anchor="w", padx=8, pady=(0, 4))

        # Polaridad hacia una entidad concreta
        pol_ent = tk.Frame(frm_pol, bg=CONTENT_BG)
        pol_ent.pack(fill="x", padx=8, pady=(0, 4))
        tk.Label(pol_ent, text="Polaridad hacia entidad (formas separadas por «;»):",
                 bg=CONTENT_BG, fg=TXT_SEC, font=("Segoe UI", 9)).pack(side="left")
        self._ent_pol_entidad = ttk.Entry(pol_ent, width=30)
        self._ent_pol_entidad.pack(side="left", padx=(6, 6))
        ttk.Button(pol_ent, text="🎯  Calcular", style="S.TButton",
                   command=self._ling_pol_hacia).pack(side="left")
        self._lbl_pol_hacia = tk.Label(pol_ent, text="", bg=CONTENT_BG,
                                       fg="#4FC1FF", font=("Segoe UI", 9, "bold"))
        self._lbl_pol_hacia.pack(side="left", padx=(10, 0))

        self._lbl_pol_resumen = tk.Label(frm_pol, text="",
                                         bg=CONTENT_BG, fg=TXT_SEC,
                                         font=("Segoe UI", 9), wraplength=720,
                                         justify="left")
        self._lbl_pol_resumen.pack(anchor="w", padx=8, pady=(0, 4))

        cols_pol = ("art_id", "polaridad", "score", "n_pos", "n_neg", "intensidad")
        self._tv_ling_pol = ttk.Treeview(frm_pol, columns=cols_pol,
                                          show="headings", height=14)
        for cid, txt, w in [("art_id","Artículo",90),("polaridad","Polaridad",100),
                              ("score","Score",80),("n_pos","Pos.",60),
                              ("n_neg","Neg.",60),("intensidad","Intensidad %",100)]:
            self._tv_ling_pol.heading(cid, text=txt)
            self._tv_ling_pol.column(cid, width=w, anchor="w")
        self._tv_ling_pol.tag_configure("positivo", foreground="#4EC9B0")
        self._tv_ling_pol.tag_configure("negativo", foreground="#F48771")
        self._tv_ling_pol.tag_configure("neutro",   foreground=TXT_DIM)
        sv_p = ttk.Scrollbar(frm_pol, orient="vertical",
                             command=self._tv_ling_pol.yview)
        self._tv_ling_pol.configure(yscrollcommand=sv_p.set)
        sv_p.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_pol.pack(fill="both", expand=True, padx=6, pady=(0, 4))
        self._ling_pol_datos: list = []

        # ── Pestaña 9: Revisión NER (human-in-the-loop) ───────────────────────
        frm_rev = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_rev, text="  🔍 Revisión NER  ")

        bf_rev = tk.Frame(frm_rev, bg=CONTENT_BG)
        bf_rev.pack(fill="x", padx=8, pady=(8, 4))
        self._btn_ling_rev = ttk.Button(
            bf_rev, text="▶  Construir cola de revisión",
            style="P.TButton", command=self._ling_rev_construir)
        self._btn_ling_rev.pack(side="left", padx=(0, 8))
        ttk.Button(bf_rev, text="✓  Verificar", style="S.TButton",
                   command=lambda: self._ling_rev_decidir("verificada")).pack(side="left", padx=(0, 4))
        ttk.Button(bf_rev, text="✗  Descartar", style="S.TButton",
                   command=lambda: self._ling_rev_decidir("descartada")).pack(side="left", padx=(0, 4))
        ttk.Button(bf_rev, text="✎  Renombrar…", style="S.TButton",
                   command=self._ling_rev_renombrar).pack(side="left", padx=(0, 8))
        self._lbl_ling_rev = tk.Label(frm_rev, text="", bg=CONTENT_BG,
                                      fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_rev.pack(anchor="w", padx=8)

        tk.Label(frm_rev,
                 text="Entidades dudosas (1 artículo = rojo, 2 = ámbar) priorizadas para "
                      "validar a mano. Las decisiones se guardan y se re-aplican al índice.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9),
                 wraplength=720, justify="left").pack(anchor="w", padx=8, pady=(0, 4))

        cols_rev = ("nombre", "categoria", "n_articulos", "nivel", "etiqueta")
        self._tv_ling_rev = ttk.Treeview(frm_rev, columns=cols_rev,
                                          show="headings", height=16)
        for cid, txt, w in [("nombre","Entidad",220),("categoria","Categoría",120),
                              ("n_articulos","N arts.",70),("nivel","Nivel",80),
                              ("etiqueta","Estado",130)]:
            self._tv_ling_rev.heading(cid, text=txt)
            self._tv_ling_rev.column(cid, width=w, anchor="w")
        self._tv_ling_rev.tag_configure("red",    foreground="#F48771")
        self._tv_ling_rev.tag_configure("yellow", foreground="#F59E0B")
        sv_r = ttk.Scrollbar(frm_rev, orient="vertical",
                             command=self._tv_ling_rev.yview)
        self._tv_ling_rev.configure(yscrollcommand=sv_r.set)
        sv_r.pack(side="right", fill="y", padx=(0, 6))
        self._tv_ling_rev.pack(fill="both", expand=True, padx=6, pady=(0, 4))

        # ── Pestaña 10: Validación (Kappa) ────────────────────────────────────
        frm_val = tk.Frame(nb_ling, bg=CONTENT_BG)
        nb_ling.add(frm_val, text="  ✔ Validación  ")

        tk.Label(frm_val, text="Validación metodológica (fiabilidad inter-codificador)",
                 bg=CONTENT_BG, fg="#FFFFFF", font=("Segoe UI", 11, "bold")).pack(
                     anchor="w", padx=8, pady=(8, 2))
        tk.Label(frm_val,
                 text="1) Exporta una muestra aleatoria (semilla fija → reproducible) con la "
                      "clasificación automática y columnas vacías para codificar a mano.\n"
                      "2) Tras codificar el CSV, calcula el % de acuerdo y el Kappa de Cohen.",
                 bg=CONTENT_BG, fg=GRIS2, font=("Segoe UI", 9),
                 justify="left").pack(anchor="w", padx=8, pady=(0, 8))

        val_cfg = tk.Frame(frm_val, bg=CONTENT_BG)
        val_cfg.pack(fill="x", padx=8, pady=(0, 4))
        tk.Label(val_cfg, text="Dimensión:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        self._var_val_dim = tk.StringVar(value="polaridad")
        ttk.Combobox(val_cfg, textvariable=self._var_val_dim, width=12,
                     state="readonly",
                     values=["polaridad", "frame", "emocion"]).pack(side="left", padx=(4, 12))
        tk.Label(val_cfg, text="Tamaño muestra:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        self._var_val_n = tk.IntVar(value=30)
        ttk.Spinbox(val_cfg, from_=10, to=300, textvariable=self._var_val_n,
                    width=6).pack(side="left", padx=(4, 12))
        tk.Label(val_cfg, text="Semilla:", bg=CONTENT_BG, fg=TXT_SEC,
                 font=("Segoe UI", 9)).pack(side="left")
        self._var_val_semilla = tk.IntVar(value=42)
        ttk.Spinbox(val_cfg, from_=0, to=9999, textvariable=self._var_val_semilla,
                    width=6).pack(side="left", padx=(4, 0))

        bf_val = tk.Frame(frm_val, bg=CONTENT_BG)
        bf_val.pack(fill="x", padx=8, pady=(6, 4))
        ttk.Button(bf_val, text="💾  Exportar muestra para codificar…",
                   style="P.TButton",
                   command=self._ling_val_exportar).pack(side="left", padx=(0, 8))
        ttk.Button(bf_val, text="📐  Calcular concordancia (Kappa)…",
                   style="S.TButton",
                   command=self._ling_val_concordancia).pack(side="left", padx=(0, 8))
        self._lbl_ling_val = tk.Label(frm_val, text="", bg=CONTENT_BG,
                                      fg=VERDE, font=("Segoe UI", 9, "bold"))
        self._lbl_ling_val.pack(anchor="w", padx=8, pady=(4, 0))

        self._txt_val_res = scrolledtext.ScrolledText(
            frm_val, height=14, font=("Consolas", 9),
            bg=CARD_BG, fg=TXT_PRI, state="disabled", wrap="word")
        self._txt_val_res.pack(fill="both", expand=True, padx=8, pady=(8, 4))

        # ── Log compartido ────────────────────────────────────────────────────
        self._txt_ling_log = scrolledtext.ScrolledText(
            pad, height=3, font=("Consolas", 8),
            bg="#0D1B2A", fg="#94A3B8", state="disabled", wrap="word")
        self._txt_ling_log.pack(fill="x", pady=(6, 0))

    # ── Helpers internos de lingüística ──────────────────────────────────────

    def _ir_a_ling_pestania(self, indice: int):
        """Navega a Lingüística y selecciona la pestaña por índice (0-based)."""
        self._mostrar_pagina("ling")
        nb = getattr(self, "_nb_ling", None)
        if nb is not None:
            try:
                nb.select(indice)
            except Exception:
                pass

    def _ling_log(self, msg: str):
        self._txt_ling_log.config(state="normal")
        self._txt_ling_log.insert("end", msg + "\n")
        self._txt_ling_log.see("end")
        self._txt_ling_log.config(state="disabled")

    def _ling_corpus_txt(self) -> list[str]:
        """Devuelve el corpus como lista de textos planos (3 fuentes, en orden de prioridad)."""
        # 1. Ya construido por _worker_seg o modo ad-hoc
        if getattr(ST, "corpus_txt", None):
            return ST.corpus_txt
        # 2. df_articulos tiene columna "texto"
        df_art = getattr(ST, "df_articulos", None)
        if df_art is not None:
            try:
                import pandas as pd
                if "texto" in df_art.columns:
                    return df_art["texto"].dropna().tolist()
            except Exception:
                pass
        # 3. corpus_meta — puede ser DataFrame (OCR normal) o dict (modo ad-hoc)
        cm = getattr(ST, "corpus_meta", None)
        if cm is not None:
            txts = []
            try:
                import pandas as pd
                if isinstance(cm, pd.DataFrame):
                    # corpus_meta DataFrame tiene columna txt_path; leer TXTs desde disco
                    if "txt_path" in cm.columns:
                        from pathlib import Path
                        for p in cm["txt_path"].dropna():
                            try:
                                txts.append(Path(p).read_text(encoding="utf-8",
                                                               errors="replace"))
                            except Exception:
                                pass
                elif isinstance(cm, dict):
                    for num_data in cm.values():
                        if isinstance(num_data, dict):
                            for art in num_data.get("articulos", []):
                                t = art.get("texto", "") or art.get("contenido", "")
                                if t:
                                    txts.append(str(t))
            except Exception:
                pass
            if txts:
                return txts
        return []

    # ── Concordancias sintácticas ─────────────────────────────────────────────

    def _ling_concordancias(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus",
                "Procesa primero el corpus en Normalizar o Segmentar.")
            return
        self._btn_ling_sint.config(state="disabled")
        self._lbl_ling_sint.config(text="Analizando…")
        patron = self._var_ling_patron.get()
        max_r = self._var_ling_max.get()
        threading.Thread(target=self._worker_ling_sint,
                         args=(corpus, patron, max_r), daemon=True).start()

    def _worker_ling_sint(self, corpus, patron, max_r):
        try:
            from core.sintaxis_engine import concordancias_sintaticas
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_sint.config(state="normal"))
            return

        def cb(i, total):
            self.after(0, lambda: self._ling_log(f"Parseando doc {i}/{total}…"))

        try:
            res = concordancias_sintaticas(corpus, patron=patron,
                                           max_resultados=max_r, callback=cb)
            self._ling_sint_resultados = res
            self.after(0, lambda: self._poblar_tv_sint(res))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error sintaxis", err))
        finally:
            self.after(0, lambda: self._btn_ling_sint.config(state="normal"))

    def _poblar_tv_sint(self, res):
        for row in self._tv_ling_sint.get_children():
            self._tv_ling_sint.delete(row)
        for r in res:
            frag = r.get("texto_completo", "")[:80]
            self._tv_ling_sint.insert("", "end", values=(
                r.get("patron", ""),
                r.get("match_principal", ""),
                r.get("match_secundario", ""),
                r.get("descripcion", ""),
                frag,
            ))
        self._lbl_ling_sint.config(
            text=f"✓ {len(res)} concordancias sintácticas encontradas.")

    def _ling_sint_csv(self):
        if not self._ling_sint_resultados:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        import csv
        from tkinter import filedialog
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="concordancias_sintaticas.csv",
        )
        if not ruta:
            return
        campos = ["patron", "match_principal", "match_secundario",
                  "descripcion", "texto_completo", "doc_idx"]
        with open(ruta, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.DictWriter(f, fieldnames=campos, extrasaction="ignore")
            w.writeheader()
            w.writerows(self._ling_sint_resultados)
        messagebox.showinfo("Exportado", f"CSV guardado en:\n{ruta}")

    # ── Relaciones SVO ────────────────────────────────────────────────────────

    def _ling_relaciones(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus",
                "Procesa primero el corpus en Normalizar o Segmentar.")
            return
        self._btn_ling_svo.config(state="disabled")
        self._btn_ling_sint.config(state="disabled")
        self._lbl_ling_svo.config(text="Extrayendo relaciones…")
        solo_ent = self._var_ling_solo_ent.get()
        conf = self._var_ling_conf.get()
        threading.Thread(target=self._worker_ling_svo,
                         args=(corpus, solo_ent, conf), daemon=True).start()

    def _worker_ling_svo(self, corpus, solo_ent, conf_min):
        try:
            from core.sintaxis_engine import extraer_relaciones
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_svo.config(state="normal"))
            self.after(0, lambda: self._btn_ling_sint.config(state="normal"))
            return

        def cb(i, total):
            self.after(0, lambda: self._ling_log(f"Extrayendo relaciones doc {i}/{total}…"))

        try:
            res = extraer_relaciones(corpus, solo_entidades=solo_ent,
                                     min_confianza=conf_min, callback=cb)
            self._ling_svo_resultados = res
            self.after(0, lambda: self._poblar_tv_svo(res))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error relaciones", err))
        finally:
            self.after(0, lambda: self._btn_ling_svo.config(state="normal"))
            self.after(0, lambda: self._btn_ling_sint.config(state="normal"))

    def _poblar_tv_svo(self, res):
        for row in self._tv_ling_svo.get_children():
            self._tv_ling_svo.delete(row)
        for r in res:
            self._tv_ling_svo.insert("", "end", values=(
                r.get("sujeto", ""),
                r.get("sujeto_tipo", ""),
                r.get("relacion", ""),
                r.get("objeto", ""),
                r.get("objeto_tipo", ""),
                r.get("confianza", 0),
                r.get("oracion", "")[:80],
            ))
        # Cambiar a pestaña SVO
        self._nb_ling.select(1)
        self._lbl_ling_svo.config(
            text=f"✓ {len(res)} relaciones sujeto-verbo-objeto extraídas.")

    def _ling_svo_csv(self):
        if not self._ling_svo_resultados:
            messagebox.showinfo("Sin datos", "Ejecuta la extracción primero.")
            return
        from tkinter import filedialog

        from core.sintaxis_engine import exportar_relaciones_csv
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="relaciones_svo.csv",
        )
        if not ruta:
            return
        exportar_relaciones_csv(self._ling_svo_resultados, ruta)
        messagebox.showinfo("Exportado", f"CSV guardado en:\n{ruta}")

    def _ling_svo_agrupar(self):
        if not self._ling_svo_resultados:
            messagebox.showinfo("Sin datos", "Ejecuta la extracción primero.")
            return
        from core.sintaxis_engine import agrupar_relaciones
        agrup = agrupar_relaciones(self._ling_svo_resultados)
        # Mostrar en ventana simple
        win, content_agrup = self._mk_glass_toplevel(
            "🔗 Relaciones agrupadas por verbo", ancho=600, alto=500)
        txt = scrolledtext.ScrolledText(content_agrup, font=("Consolas", 9),
                                         bg="#0D1B2A", fg="#E2E8F0", relief="flat")
        txt.pack(fill="both", expand=True, padx=8, pady=8)
        for verbo, info in agrup.items():
            txt.insert("end", f"\n{'─'*50}\n{verbo.upper()} ({info['n']} ocurrencias)\n")
            for r in info["relaciones"][:5]:
                txt.insert("end",
                    f"  {r.get('sujeto','—')} → {verbo} → {r.get('objeto','—')}\n")
        txt.config(state="disabled")

    # ── Correferencia ─────────────────────────────────────────────────────────

    def _ling_coref(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus",
                "Procesa primero el corpus en Normalizar o Segmentar.")
            return
        entidad = self._ent_coref.get().strip()
        self._btn_ling_coref.config(state="disabled")
        self._lbl_ling_coref.config(text="Resolviendo correferencias…")
        threading.Thread(target=self._worker_ling_coref,
                         args=(corpus, entidad), daemon=True).start()

    def _worker_ling_coref(self, corpus, entidad_filtro):
        try:
            from core.coref_engine import cadena_referencial, resolver_correferencias
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_coref.config(state="normal"))
            return

        try:
            todas_cadenas = []
            for i, texto in enumerate(corpus[:50]):  # límite para velocidad
                self.after(0, lambda i=i: self._ling_log(
                    f"Correferencia texto {i+1}/{min(len(corpus),50)}…"))
                cadenas = resolver_correferencias(texto)
                for c in cadenas:
                    c["doc_idx"] = i
                    todas_cadenas.append(c)

            # Filtrar si hay entidad específica
            if entidad_filtro:
                filtradas = [c for c in todas_cadenas
                             if entidad_filtro.lower() in
                             c["entidad_principal"].lower()]
            else:
                filtradas = todas_cadenas

            # Ordenar por nº de menciones
            filtradas.sort(key=lambda x: -x.get("n_menciones", 0))

            self._ling_coref_cadenas = filtradas
            self.after(0, lambda: self._poblar_coref(filtradas))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error coref", err))
        finally:
            self.after(0, lambda: self._btn_ling_coref.config(state="normal"))

    def _poblar_coref(self, cadenas):
        self._lb_coref.delete(0, "end")
        for c in cadenas:
            self._lb_coref.insert("end",
                f"{c['entidad_principal']} ({c['n_menciones']} menciones)")
        self._lbl_ling_coref.config(
            text=f"✓ {len(cadenas)} cadenas referenciales.")

    def _ling_coref_mostrar_cadena(self, event=None):
        sel = self._lb_coref.curselection()
        if not sel:
            return
        idx = sel[0]
        if idx >= len(self._ling_coref_cadenas):
            return
        cadena = self._ling_coref_cadenas[idx]
        for row in self._tv_coref_men.get_children():
            self._tv_coref_men.delete(row)
        for m in cadena.get("menciones", []):
            self._tv_coref_men.insert("", "end", values=(
                m.get("texto", ""),
                m.get("tipo", ""),
                m.get("oracion", "")[:120],
            ))

    def _ling_coref_stats(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus",
                "Procesa primero el corpus.")
            return
        self._btn_ling_coref.config(state="disabled")
        threading.Thread(target=self._worker_coref_stats,
                         args=(corpus,), daemon=True).start()

    def _worker_coref_stats(self, corpus):
        try:
            from core.coref_engine import estadisticas_coref
            def cb(i, t): self.after(0, lambda: self._ling_log(
                f"Stats coref {i}/{t}…"))
            stats = estadisticas_coref(corpus[:30], callback=cb)
            msg = (
                f"Cadenas referenciales totales: {stats['total_cadenas']}\n"
                f"Menciones promedio por cadena: {stats['promedio_menciones']}\n"
                f"Densidad referencial: {stats['densidad_referencial']} "
                f"({stats['total_pronombres']} pronombres / {stats['total_tokens']} tokens)\n\n"
                "Top entidades referidas:\n"
            )
            for e in stats["entidades_mas_referidas"][:10]:
                msg += f"  {e['entidad']}: {e['n_menciones']} menciones\n"
            self.after(0, lambda: messagebox.showinfo("Estadísticas de correferencia", msg))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error", err))
        finally:
            self.after(0, lambda: self._btn_ling_coref.config(state="normal"))

    # ── Morfología histórica ──────────────────────────────────────────────────

    def _ling_morf_analizar(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus",
                "Procesa primero el corpus.")
            return
        normalizar = getattr(self, "_var_morf_normalizar",
                             tk.BooleanVar(value=True)).get()
        self._btn_ling_morf.config(state="disabled")
        self._lbl_ling_morf.config(text="Analizando formas históricas…")
        threading.Thread(target=self._worker_ling_morf,
                         args=(corpus, normalizar), daemon=True).start()

    def _worker_ling_morf(self, corpus, normalizar: bool = True):
        try:
            from core.morfologia_historica import (
                enriquecer_corpus_con_lemas,
                normalizar_formas_historicas,
            )
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_morf.config(state="normal"))
            return

        def cb(i, t): self.after(0, lambda: self._ling_log(
            f"Morfología doc {i}/{t}…"))

        try:
            corpus_proc = ([normalizar_formas_historicas(t) for t in corpus]
                           if normalizar else corpus)
            datos = enriquecer_corpus_con_lemas(corpus_proc, callback=cb)
            self._ling_morf_datos = datos
            self.after(0, lambda: self._poblar_tv_morf(datos))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error morfología", err))
        finally:
            self.after(0, lambda: self._btn_ling_morf.config(state="normal"))

    def _poblar_tv_morf(self, datos):
        for row in self._tv_ling_morf.get_children():
            self._tv_ling_morf.delete(row)
        total_arc = sum(d["n_arcaismos"] for d in datos)
        total_tok = sum(d["n_tokens"] for d in datos)
        score_med = round(sum(d["score"] for d in datos) / len(datos), 4) if datos else 0
        self._lbl_morf_resumen.config(
            text=(f"Corpus: {len(datos)} docs | "
                  f"Arcaísmos: {total_arc} / {total_tok} tokens | "
                  f"Score histórico medio: {score_med:.4f}")
        )
        for d in datos:
            top = ", ".join(f"{x['forma']}×{x['n']}"
                            for x in d.get("top_arcaismos", [])[:4])
            marc = d.get("marcadores", {})
            marc_str = "  ".join(f"{k}({v})" for k, v in marc.items()) if marc else "—"
            score = d["score"]
            tag = "alta" if score >= 0.05 else ("media" if score >= 0.01 else "baja")
            self._tv_ling_morf.insert("", "end", tags=(tag,), values=(
                d["doc_idx"] + 1,
                d["n_tokens"],
                d["n_arcaismos"],
                f"{score:.4f}",
                marc_str,
                top,
            ))
        self._lbl_ling_morf.config(
            text=f"✓ {total_arc} formas históricas en {len(datos)} docs. Score medio: {score_med:.4f}")

    def _ling_morf_detalle(self):
        """Muestra detalle de marcadores y ejemplos del documento seleccionado."""
        sel = self._tv_ling_morf.selection()
        if not sel or not self._ling_morf_datos:
            self.toast("Selecciona un documento en la tabla primero", tipo="warn")
            return
        idx = self._tv_ling_morf.index(sel[0])
        if idx >= len(self._ling_morf_datos):
            return
        d = self._ling_morf_datos[idx]

        win, content = self._mk_glass_toplevel(
            f"📜 Detalle morfológico — Doc {d['doc_idx']+1}", ancho=560, alto=480)

        # Resumen numérico
        info = tk.Frame(content, bg=CONTENT_BG, pady=10)
        info.pack(fill="x", padx=16)
        for lbl, val in [("Tokens:", d["n_tokens"]),
                         ("Arcaísmos:", d["n_arcaismos"]),
                         ("Score histórico:", f"{d['score']:.4f}")]:
            row = tk.Frame(info, bg=CONTENT_BG); row.pack(fill="x", pady=1)
            tk.Label(row, text=lbl, bg=CONTENT_BG, fg=TXT_SEC,
                     font=("Segoe UI", 9), width=20, anchor="w").pack(side="left")
            tk.Label(row, text=str(val), bg=CONTENT_BG, fg=TXT_PRI,
                     font=("Segoe UI", 9, "bold")).pack(side="left")

        tk.Frame(content, bg=CARD_BOR, height=1).pack(fill="x", padx=8)

        # Marcadores morfosintácticos
        tk.Label(content, text="Marcadores morfosintácticos detectados",
                 bg=CONTENT_BG, fg=AZ4, font=("Segoe UI", 9, "bold")).pack(
                 anchor="w", padx=16, pady=(8, 2))
        marc = d.get("marcadores", {})
        if marc:
            for tipo, n in marc.items():
                row = tk.Frame(content, bg=CONTENT_BG); row.pack(fill="x", padx=24, pady=1)
                tk.Label(row, text=f"• {tipo}", bg=CONTENT_BG, fg=VERDE,
                         font=("Segoe UI", 9), width=30, anchor="w").pack(side="left")
                tk.Label(row, text=f"{n} ocurrencias", bg=CONTENT_BG, fg=TXT_SEC,
                         font=("Segoe UI", 9)).pack(side="left")
        else:
            tk.Label(content, text="  (ninguno)", bg=CONTENT_BG, fg=TXT_DIM,
                     font=("Segoe UI", 9, "italic")).pack(anchor="w", padx=24)

        tk.Frame(content, bg=CARD_BOR, height=1).pack(fill="x", padx=8, pady=(6, 0))

        # Ejemplos
        tk.Label(content, text="Ejemplos de formas históricas",
                 bg=CONTENT_BG, fg=AZ4, font=("Segoe UI", 9, "bold")).pack(
                 anchor="w", padx=16, pady=(8, 2))
        ejs = d.get("ejemplos", [])
        if ejs:
            cols_ej = ("token", "tipo")
            tv_ej = ttk.Treeview(content, columns=cols_ej, show="headings", height=8)
            tv_ej.heading("token", text="Forma"); tv_ej.column("token", width=200, anchor="w")
            tv_ej.heading("tipo",  text="Tipo");  tv_ej.column("tipo",  width=280, anchor="w")
            for ej in ejs:
                tv_ej.insert("", "end", values=(ej["token"], ej["tipo"]))
            tv_ej.pack(fill="both", expand=True, padx=16, pady=(0, 12))
        else:
            tk.Label(content, text="  (ninguno)", bg=CONTENT_BG, fg=TXT_DIM,
                     font=("Segoe UI", 9, "italic")).pack(anchor="w", padx=24)

    def _ling_morf_glosario(self):
        from core.morfologia_historica import glosario_arcaismos
        glos = glosario_arcaismos()
        win, content = self._mk_glass_toplevel(
            f"📖 Glosario de arcaísmos ({len(glos)} entradas)", ancho=500, alto=600)

        # ── Pie con botón exportar ────────────────────────────────────────────
        pie = tk.Frame(content, bg=CONTENT_BG); pie.pack(side="bottom", fill="x", padx=6, pady=(4, 6))

        def _exportar_glosario():
            import csv
            import pathlib
            out = pathlib.Path(ST.datos_dir) / "glosario_arcaismos.csv" if ST.datos_dir else None
            if out is None:
                from tkinter import filedialog
                ruta = filedialog.asksaveasfilename(
                    defaultextension=".csv", filetypes=[("CSV","*.csv")],
                    initialfile="glosario_arcaismos.csv")
                if not ruta: return
                out = pathlib.Path(ruta)
            with open(out, "w", newline="", encoding="utf-8-sig") as f:
                w = csv.writer(f)
                w.writerow(["forma_historica", "lema_moderno"])
                for g in glos:
                    w.writerow([g["forma_historica"], g["lema_moderno"]])
            self.toast(f"Glosario exportado → {out.name}", tipo="ok")

        ttk.Button(pie, text="💾  Exportar CSV", style="S.TButton",
                   command=_exportar_glosario).pack(side="right")

        # ── Treeview ─────────────────────────────────────────────────────────
        cols = ("forma", "lema")
        tv = ttk.Treeview(content, columns=cols, show="headings", height=26)
        tv.heading("forma", text="Forma histórica")
        tv.heading("lema",  text="Lema moderno")
        tv.column("forma",  width=200, anchor="w")
        tv.column("lema",   width=200, anchor="w")
        sv = ttk.Scrollbar(content, orient="vertical", command=tv.yview)
        tv.configure(yscrollcommand=sv.set)
        sv.pack(side="right", fill="y")
        tv.pack(fill="both", expand=True, padx=6, pady=(6, 0))
        for g in glos:
            tv.insert("", "end", values=(g["forma_historica"], g["lema_moderno"]))

    def _ling_morf_csv(self):
        if not self._ling_morf_datos:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        import csv
        from tkinter import filedialog
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="morfologia_historica.csv",
        )
        if not ruta:
            return
        with open(ruta, "w", newline="", encoding="utf-8-sig") as f:
            campos = ["doc_idx", "n_tokens", "n_arcaismos", "score", "top_arcaismos"]
            w = csv.writer(f)
            w.writerow(campos)
            for d in self._ling_morf_datos:
                top = " | ".join(f"{x['forma']}x{x['n']}"
                                  for x in d.get("top_arcaismos", []))
                w.writerow([d["doc_idx"], d["n_tokens"],
                             d["n_arcaismos"], d["score"], top])
        messagebox.showinfo("Exportado", f"CSV guardado en:\n{ruta}")

    # ── Árbol de dependencias ─────────────────────────────────────────────────

    def _ling_dep_analizar(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus", "Procesa primero el corpus."); return
        art_idx = self._var_dep_art.get()
        if art_idx >= len(corpus):
            messagebox.showwarning("Índice inválido",
                f"El corpus tiene {len(corpus)} artículos (0–{len(corpus)-1})."); return
        texto = corpus[art_idx]
        max_or = self._var_dep_max.get()
        self._btn_ling_dep.config(state="disabled")
        self._lbl_ling_dep.config(text="Analizando dependencias…")
        threading.Thread(target=self._worker_ling_dep,
                         args=(texto, max_or), daemon=True).start()

    def _worker_ling_dep(self, texto: str, max_or: int):
        try:
            from core.sintaxis_engine import analizar_dependencias, resumir_arbol_dep
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_dep.config(state="normal"))
            return
        try:
            datos = analizar_dependencias(texto, max_oraciones=max_or)
            self._ling_dep_datos = datos
            self.after(0, lambda: self._poblar_dep(datos))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error árbol dep.", err))
        finally:
            self.after(0, lambda: self._btn_ling_dep.config(state="normal"))

    def _poblar_dep(self, datos: list):
        from core.sintaxis_engine import resumir_arbol_dep
        self._lb_dep.delete(0, "end")
        for i, d in enumerate(datos):
            resumen = resumir_arbol_dep(d)
            preview = d["oracion"][:60].replace("\n", " ")
            self._lb_dep.insert("end", f"{i+1}. {preview}")
        self._lbl_ling_dep.config(
            text=f"✓ {len(datos)} oraciones analizadas")
        if datos:
            self._lb_dep.selection_set(0)
            self._ling_dep_mostrar_tokens(None)

    def _ling_dep_mostrar_tokens(self, event):
        sel = self._lb_dep.curselection()
        if not sel or not self._ling_dep_datos:
            return
        idx = sel[0]
        if idx >= len(self._ling_dep_datos):
            return
        d = self._ling_dep_datos[idx]
        for row in self._tv_dep_tok.get_children():
            self._tv_dep_tok.delete(row)
        for tok in d.get("tokens", []):
            tag = tok.get("pos", "")
            self._tv_dep_tok.insert("", "end", tags=(tag,), values=(
                tok["texto"], tok["lemma"], tok["pos"],
                tok["dep_es"], tok["cabeza"],
            ))

    def _ling_dep_csv(self):
        if not self._ling_dep_datos:
            self.toast("Ejecuta el análisis primero", tipo="warn"); return
        import csv
        from tkinter import filedialog
        dest = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="arbol_dependencias.csv")
        if not dest:
            return
        with open(dest, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            w.writerow(["oracion_idx", "oracion", "sujeto", "verbo", "objeto",
                         "token", "lemma", "pos", "dep_es", "cabeza"])
            for i, d in enumerate(self._ling_dep_datos):
                for tok in d.get("tokens", []):
                    w.writerow([i + 1, d["oracion"][:80],
                                 d.get("sujeto", ""), d.get("verbo", ""),
                                 d.get("objeto", ""),
                                 tok["texto"], tok["lemma"], tok["pos"],
                                 tok["dep_es"], tok["cabeza"]])
        self.toast(f"CSV exportado → {Path(dest).name}", tipo="ok")

    # ── Emociones y subjetividad ──────────────────────────────────────────────

    def _ling_emociones(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus",
                "Procesa primero el corpus.")
            return
        self._btn_ling_emo.config(state="disabled")
        self._lbl_ling_emo.config(text="Analizando emociones…")
        threading.Thread(target=self._worker_ling_emo,
                         args=(corpus,), daemon=True).start()

    def _worker_ling_emo(self, corpus):
        try:
            from core.sentiment_engine import analisis_completo_emocion
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_emo.config(state="normal"))
            return

        try:
            resultados = []
            total = len(corpus)
            for i, texto in enumerate(corpus):
                self.after(0, lambda i=i: self._ling_log(
                    f"Emociones doc {i+1}/{total}…"))
                res = analisis_completo_emocion(texto)
                res["doc_idx"] = i
                res["art_id"]  = f"doc_{i+1:04d}"
                resultados.append(res)
            self._ling_emo_datos = resultados
            self.after(0, lambda: self._poblar_tv_emo(resultados))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error emociones", err))
        finally:
            self.after(0, lambda: self._btn_ling_emo.config(state="normal"))

    def _poblar_tv_emo(self, resultados):
        from collections import Counter
        for row in self._tv_ling_emo.get_children():
            self._tv_ling_emo.delete(row)

        cnt_emo: Counter = Counter()
        cnt_disc: Counter = Counter()
        for r in resultados:
            emo_data = r.get("emociones", {}).get("emociones", {})
            subj_data = r.get("subjetividad", {})
            intens    = r.get("intensidad", {})

            emo_dom = r.get("emociones", {}).get("emocion_dominante") or "—"
            cnt_emo[emo_dom] += 1
            tipo_disc = subj_data.get("tipo_discurso", "—")
            cnt_disc[tipo_disc] += 1

            palabras = [p["palabra"] for p in
                        r.get("emociones", {}).get("palabras_detectadas", [])[:6]]
            score_i = intens.get("score_intensidad", 0)

            self._tv_ling_emo.insert("", "end", values=(
                r.get("art_id", ""),
                emo_dom,
                f"{subj_data.get('score_subjetividad', 0):.2f}",
                tipo_disc,
                f"{score_i:.2f}",
                ", ".join(palabras),
            ))

        # Resumen
        resumen = (
            f"Emoción dominante más frecuente: "
            f"{cnt_emo.most_common(1)[0][0] if cnt_emo else '—'} | "
            f"Discurso subjetivo: {cnt_disc.get('subjetivo', 0)} docs | "
            f"Factual: {cnt_disc.get('factual', 0)} docs | "
            f"Mixto: {cnt_disc.get('mixto', 0)} docs"
        )
        self._lbl_emo_resumen.config(text=resumen)
        self._lbl_ling_emo.config(
            text=f"✓ {len(resultados)} artículos analizados.")

    def _ling_emo_graficar(self):
        if not self._ling_emo_datos:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        from collections import Counter
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            messagebox.showerror("Error", "Instala matplotlib para graficar.")
            return

        cnt: Counter = Counter()
        for r in self._ling_emo_datos:
            emo = r.get("emociones", {}).get("emocion_dominante")
            if emo:
                cnt[emo] += 1

        if not cnt:
            messagebox.showinfo("Sin datos", "No se detectaron emociones.")
            return

        emociones = list(cnt.keys())
        valores = [cnt[e] for e in emociones]
        colores = ["#22C55E","#3B82F6","#EF4444","#F59E0B",
                   "#8B5CF6","#EC4899","#14B8A6","#F97316"]

        fig, ax = plt.subplots(figsize=(8, 4))
        ax.bar(emociones, valores,
               color=colores[:len(emociones)], edgecolor="none")
        ax.set_title("Distribución de emociones dominantes en el corpus",
                     pad=12, fontsize=11)
        ax.set_ylabel("Artículos")
        ax.set_facecolor("#0D1B2A")
        fig.patch.set_facecolor("#0D1B2A")
        ax.tick_params(colors="#94A3B8")
        ax.yaxis.label.set_color("#94A3B8")
        ax.title.set_color("#E2E8F0")
        for spine in ax.spines.values():
            spine.set_edgecolor("#1E3A5F")
        plt.tight_layout()
        plt.show()

    def _ling_emo_csv(self):
        if not self._ling_emo_datos:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        import csv
        from tkinter import filedialog
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv",
            filetypes=[("CSV", "*.csv")],
            initialfile="emociones_corpus.csv",
        )
        if not ruta:
            return
        with open(ruta, "w", newline="", encoding="utf-8-sig") as f:
            campos = ["art_id", "emocion_dominante", "score_subjetividad",
                      "tipo_discurso", "score_intensidad", "palabras_emocionales"]
            w = csv.writer(f)
            w.writerow(campos)
            for r in self._ling_emo_datos:
                emo_data = r.get("emociones", {})
                subj     = r.get("subjetividad", {})
                intens   = r.get("intensidad", {})
                palabras = " | ".join(
                    p["palabra"] for p in emo_data.get("palabras_detectadas", [])[:10]
                )
                w.writerow([
                    r.get("art_id", ""),
                    emo_data.get("emocion_dominante", ""),
                    subj.get("score_subjetividad", 0),
                    subj.get("tipo_discurso", ""),
                    intens.get("score_intensidad", 0),
                    palabras,
                ])
        messagebox.showinfo("Exportado", f"CSV guardado en:\n{ruta}")

    # ── Encuadre (framing) ────────────────────────────────────────────────────

    def _ling_frames(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus", "Procesa primero el corpus.")
            return
        self._btn_ling_frame.config(state="disabled")
        self._lbl_ling_frame.config(text="Analizando encuadres…")
        threading.Thread(target=self._worker_ling_frames,
                         args=(corpus,), daemon=True).start()

    def _worker_ling_frames(self, corpus):
        try:
            from core import frame_engine
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_frame.config(state="normal"))
            return
        try:
            corpus_dict = {f"doc_{i+1:04d}": t for i, t in enumerate(corpus)}
            resumen = frame_engine.analizar_corpus_frames(corpus_dict)
            self._ling_frame_corpus = resumen
            datos = []
            for art_id, r in resumen["por_articulo"].items():
                datos.append({"art_id": art_id, **r})
            self._ling_frame_datos = datos
            self.after(0, lambda: self._poblar_tv_frame(resumen, datos))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error encuadre", err))
        finally:
            self.after(0, lambda: self._btn_ling_frame.config(state="normal"))

    def _poblar_tv_frame(self, resumen, datos):
        for row in self._tv_ling_frame.get_children():
            self._tv_ling_frame.delete(row)
        for d in datos:
            dist = d.get("distribucion", [])
            pct = f"{dist[0]['porcentaje']:.0f}%" if dist else "—"
            sec = dist[1]["frame"] if len(dist) > 1 else "—"
            self._tv_ling_frame.insert("", "end", values=(
                d.get("art_id", ""),
                d.get("frame_dominante") or "—",
                d.get("etiqueta") or "—",
                pct, sec, d.get("total_marcadores", 0),
            ))
        dist_corpus = resumen.get("distribucion_corpus", {})
        top = " · ".join(f"{k} ({v})" for k, v in list(dist_corpus.items())[:5])
        self._lbl_frame_resumen.config(
            text=f"Encuadre dominante del corpus: "
                 f"{resumen.get('etiqueta_dominante') or '—'}   |   Top: {top}")
        self._lbl_ling_frame.config(
            text=f"✓ {resumen.get('n_articulos', 0)} artículos analizados.")

    def _ling_frames_graficar(self):
        if not self._ling_frame_corpus:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        try:
            import matplotlib.pyplot as plt
        except ImportError:
            messagebox.showerror("Error", "Instala matplotlib para graficar.")
            return
        dist = self._ling_frame_corpus.get("distribucion_corpus", {})
        if not dist:
            messagebox.showinfo("Sin datos", "No se detectaron encuadres.")
            return
        frames = list(dist.keys())
        valores = [dist[f] for f in frames]
        fig, ax = plt.subplots(figsize=(9, 4.5))
        ax.barh(frames[::-1], valores[::-1], color="#4FC1FF", edgecolor="none")
        ax.set_title("Encuadres dominantes en el corpus", pad=12, fontsize=11)
        ax.set_xlabel("Artículos")
        fig.tight_layout()
        plt.show()

    def _ling_frames_csv(self):
        if not self._ling_frame_datos:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        import csv
        from tkinter import filedialog
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")],
            initialfile="encuadres_corpus.csv")
        if not ruta:
            return
        with open(ruta, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            w.writerow(["art_id", "frame_dominante", "etiqueta", "porcentaje_dom",
                        "total_marcadores", "distribucion"])
            for d in self._ling_frame_datos:
                dist = d.get("distribucion", [])
                pct = dist[0]["porcentaje"] if dist else 0
                resumen_dist = "; ".join(
                    f"{x['frame']}:{x['porcentaje']}%" for x in dist)
                w.writerow([d.get("art_id", ""), d.get("frame_dominante") or "",
                            d.get("etiqueta") or "", pct,
                            d.get("total_marcadores", 0), resumen_dist])
        messagebox.showinfo("Exportado", f"CSV guardado en:\n{ruta}")

    # ── Polaridad discriminante ───────────────────────────────────────────────

    def _ling_polaridad(self):
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus", "Procesa primero el corpus.")
            return
        self._btn_ling_pol.config(state="disabled")
        self._lbl_ling_pol.config(text="Analizando polaridad…")
        threading.Thread(target=self._worker_ling_pol,
                         args=(corpus,), daemon=True).start()

    def _worker_ling_pol(self, corpus):
        try:
            from core import sentimiento_discriminante as sd
        except ImportError as e:
            self.after(0, lambda err=str(e): messagebox.showerror("Import error", err))
            self.after(0, lambda: self._btn_ling_pol.config(state="normal"))
            return
        try:
            datos = []
            for i, texto in enumerate(corpus):
                r = sd.analizar_polaridad(texto)
                r["art_id"] = f"doc_{i+1:04d}"
                datos.append(r)
            self._ling_pol_datos = datos
            self.after(0, lambda: self._poblar_tv_pol(datos))
        except Exception as ex:
            self.after(0, lambda err=str(ex): messagebox.showerror("Error polaridad", err))
        finally:
            self.after(0, lambda: self._btn_ling_pol.config(state="normal"))

    def _poblar_tv_pol(self, datos):
        from collections import Counter
        for row in self._tv_ling_pol.get_children():
            self._tv_ling_pol.delete(row)
        cnt: Counter = Counter()
        for r in datos:
            pol = r.get("polaridad", "neutro")
            cnt[pol] += 1
            self._tv_ling_pol.insert("", "end", tags=(pol,), values=(
                r.get("art_id", ""), pol, f"{r.get('score', 0):+.3f}",
                r.get("n_pos", 0), r.get("n_neg", 0),
                f"{r.get('intensidad', 0):.2f}",
            ))
        try:
            from core import sentimiento_discriminante as sd
            ipa = sd.indice_polarizacion_afectiva(dict(cnt))
        except Exception:
            ipa = 0.0
        self._lbl_pol_resumen.config(
            text=f"Positivo: {cnt.get('positivo', 0)}  ·  "
                 f"Negativo: {cnt.get('negativo', 0)}  ·  "
                 f"Neutro: {cnt.get('neutro', 0)}   |   "
                 f"Índice de polarización afectiva: {ipa:.3f}")
        self._lbl_ling_pol.config(text=f"✓ {len(datos)} artículos analizados.")

    def _ling_pol_hacia(self):
        formas = [s.strip() for s in self._ent_pol_entidad.get().split(";")
                  if s.strip()]
        if not formas:
            messagebox.showinfo("Entidad vacía",
                "Escribe una o más formas de la entidad, separadas por «;».")
            return
        corpus = self._ling_corpus_txt()
        if not corpus:
            messagebox.showwarning("Sin corpus", "Procesa primero el corpus.")
            return
        try:
            from core import sentimiento_discriminante as sd
        except ImportError as e:
            messagebox.showerror("Import error", str(e))
            return
        r = sd.polaridad_hacia_corpus(corpus, formas)
        n_docs = r.get("n_documentos", 0)
        extra = f" en {n_docs} art." if n_docs else ""
        self._lbl_pol_hacia.config(
            text=f"→ {r['polaridad']}  (score {r.get('score', 0):+.3f}, "
                 f"{r.get('n_menciones', 0)} menciones{extra})")

    def _ling_pol_csv(self):
        if not self._ling_pol_datos:
            messagebox.showinfo("Sin datos", "Ejecuta el análisis primero.")
            return
        import csv
        from tkinter import filedialog
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")],
            initialfile="polaridad_corpus.csv")
        if not ruta:
            return
        with open(ruta, "w", newline="", encoding="utf-8-sig") as f:
            w = csv.writer(f)
            w.writerow(["art_id", "polaridad", "score", "n_pos", "n_neg",
                        "intensidad"])
            for r in self._ling_pol_datos:
                w.writerow([r.get("art_id", ""), r.get("polaridad", ""),
                            r.get("score", 0), r.get("n_pos", 0),
                            r.get("n_neg", 0), r.get("intensidad", 0)])
        messagebox.showinfo("Exportado", f"CSV guardado en:\n{ruta}")

    # ── Revisión NER (human-in-the-loop) ──────────────────────────────────────

    def _ling_rev_con(self):
        """Conexión SQLite para la cola de revisión (la del proyecto, o una junto a él)."""
        import sqlite3
        from pathlib import Path
        ruta = getattr(ST, "ruta_db", "") or ""
        if not ruta:
            base = getattr(ST, "out_dir", None) or Path.cwd()
            ruta = str(Path(base) / "revision_ner.db")
        con = sqlite3.connect(ruta, timeout=30)
        con.row_factory = sqlite3.Row
        return con

    def _ling_rev_construir(self):
        indice = getattr(ST, "indice_ner_global", None)
        if not indice:
            messagebox.showwarning("Sin índice NER",
                "Ejecuta primero el NER del corpus (panel Entidades).")
            return
        try:
            from core import revision_engine
            con = self._ling_rev_con()
            cola = revision_engine.construir_cola(indice)
            revision_engine.guardar_cola(con, cola)
            pend = revision_engine.pendientes(con)
            con.close()
        except Exception as ex:
            messagebox.showerror("Error revisión", str(ex))
            return
        self._poblar_tv_rev(pend)
        self._lbl_ling_rev.config(
            text=f"✓ {len(pend)} entidades pendientes de revisar.")

    def _poblar_tv_rev(self, pendientes):
        for row in self._tv_ling_rev.get_children():
            self._tv_ling_rev.delete(row)
        for d in pendientes:
            nivel = d.get("nivel", "")
            self._tv_ling_rev.insert("", "end", iid=f"{d['categoria']}|{d['nombre']}",
                                     tags=(nivel,), values=(
                d.get("nombre", ""), d.get("categoria", ""),
                d.get("n_articulos", 0), nivel,
                d.get("etiqueta", "") or
                ("○ VALIDAR" if nivel == "red" else "◐ REVISAR"),
            ))

    def _ling_rev_sel(self):
        sel = self._tv_ling_rev.selection()
        if not sel:
            messagebox.showinfo("Sin selección",
                "Selecciona una entidad de la lista.")
            return None
        vals = self._tv_ling_rev.item(sel[0], "values")
        return vals[0], vals[1]  # nombre, categoria

    def _ling_rev_aplicar_al_indice(self, con):
        """Re-aplica TODAS las decisiones tomadas al índice NER en memoria
        (descarta lo rechazado, fusiona renombres). Mantiene el índice limpio
        para exportación y para que la cola no vuelva a mostrar lo ya resuelto."""
        from core import revision_engine
        decisiones = revision_engine.cargar_decisiones(con)
        revision_engine.aplicar_revisiones(
            getattr(ST, "indice_ner_global", {}) or {}, decisiones)

    def _ling_rev_decidir(self, decision):
        s = self._ling_rev_sel()
        if not s:
            return
        nombre, categoria = s
        try:
            from core import revision_engine
            con = self._ling_rev_con()
            revision_engine.decidir(con, nombre, categoria, decision)
            self._ling_rev_aplicar_al_indice(con)
            pend = revision_engine.pendientes(con)
            con.close()
        except Exception as ex:
            messagebox.showerror("Error revisión", str(ex))
            return
        self._poblar_tv_rev(pend)
        self._lbl_ling_rev.config(
            text=f"«{nombre}» → {decision}.  {len(pend)} pendientes.")

    def _ling_rev_renombrar(self):
        s = self._ling_rev_sel()
        if not s:
            return
        nombre, categoria = s
        from tkinter import simpledialog
        nuevo = simpledialog.askstring(
            "Renombrar entidad",
            f"Nuevo nombre canónico para «{nombre}»:", initialvalue=nombre)
        if not nuevo or nuevo == nombre:
            return
        try:
            from core import revision_engine
            con = self._ling_rev_con()
            revision_engine.decidir(con, nombre, categoria, "renombrada",
                                    nombre_nuevo=nuevo)
            self._ling_rev_aplicar_al_indice(con)
            pend = revision_engine.pendientes(con)
            con.close()
        except Exception as ex:
            messagebox.showerror("Error revisión", str(ex))
            return
        self._poblar_tv_rev(pend)
        self._lbl_ling_rev.config(text=f"«{nombre}» → «{nuevo}».")

    # ── Validación (Kappa de Cohen) ───────────────────────────────────────────

    def _ling_val_articulos(self):
        """Construye la lista de artículos para la muestra de validación."""
        corpus = self._ling_corpus_txt()
        return [{"art_id": f"doc_{i+1:04d}", "texto": t}
                for i, t in enumerate(corpus)]

    def _ling_val_etiquetador(self, dim):
        try:
            if dim == "polaridad":
                from core import sentimiento_discriminante as sd
                return lambda a: sd.analizar_polaridad(a.get("texto", "")).get("polaridad", "")
            elif dim == "emocion":
                from core.sentiment_engine import analizar_emociones
                return lambda a: (analizar_emociones(a.get("texto", "")) or {}).get(
                    "emocion_dominante") or ""
            else:  # frame
                from core import frame_engine
                return lambda a: frame_engine.analizar_frame(
                    a.get("texto", "")).get("frame_dominante") or ""
        except ImportError:
            return None

    def _ling_val_exportar(self):
        arts = self._ling_val_articulos()
        if not arts:
            messagebox.showwarning("Sin corpus", "Procesa primero el corpus.")
            return
        from tkinter import filedialog
        dim = self._var_val_dim.get()
        ruta = filedialog.asksaveasfilename(
            defaultextension=".csv", filetypes=[("CSV", "*.csv")],
            initialfile=f"muestra_validacion_{dim}.csv")
        if not ruta:
            return
        try:
            from core import validacion_engine
            etiquetador = self._ling_val_etiquetador(dim)
            validacion_engine.exportar_muestra(
                arts, ruta, n=self._var_val_n.get(),
                semilla=self._var_val_semilla.get(),
                etiqueta_auto=etiquetador, nombre_etiqueta=dim)
        except Exception as ex:
            messagebox.showerror("Error validación", str(ex))
            return
        self._lbl_ling_val.config(
            text=f"✓ Muestra exportada. Codifica la columna «{dim}_manual» y vuelve.")
        self._val_log(
            f"Muestra de {self._var_val_n.get()} artículos (semilla "
            f"{self._var_val_semilla.get()}) → {ruta}\n"
            f"Codifica a mano la columna '{dim}_manual' (mismas etiquetas que "
            f"'{dim}_auto') y luego usa «Calcular concordancia».")

    def _ling_val_concordancia(self):
        from tkinter import filedialog
        dim = self._var_val_dim.get()
        ruta = filedialog.askopenfilename(
            filetypes=[("CSV", "*.csv")], title="CSV codificado a mano")
        if not ruta:
            return
        try:
            from core import validacion_engine
            r = validacion_engine.calcular_concordancia(ruta, nombre_etiqueta=dim)
        except Exception as ex:
            messagebox.showerror("Error validación", str(ex))
            return
        if r.get("error"):
            self._val_log("⚠ " + r["error"])
            self._lbl_ling_val.config(text="Sin filas codificadas a mano.")
            return
        lineas = [
            f"Concordancia para «{dim}»  (n = {r['n']} artículos codificados)",
            "─" * 56,
            f"  Acuerdo observado : {r['acuerdo']*100:.1f}%",
            f"  Kappa de Cohen    : {r['kappa']:.3f}  ({r['interpretacion']})",
            "",
            "Matriz de confusión (manual ↓ / auto →):",
        ]
        matriz = r.get("matriz_confusion", {})
        autos = sorted({a for fila in matriz.values() for a in fila})
        lineas.append("            " + "  ".join(f"{a[:8]:>8}" for a in autos))
        for man, fila in matriz.items():
            celdas = "  ".join(f"{fila.get(a, 0):>8}" for a in autos)
            lineas.append(f"  {man[:10]:<10}{celdas}")
        self._val_log("\n".join(lineas))
        self._lbl_ling_val.config(
            text=f"Kappa = {r['kappa']:.3f} ({r['interpretacion']}), "
                 f"acuerdo {r['acuerdo']*100:.1f}%")

    def _val_log(self, msg):
        self._txt_val_res.config(state="normal")
        self._txt_val_res.delete("1.0", "end")
        self._txt_val_res.insert("end", msg + "\n")
        self._txt_val_res.config(state="disabled")


import re as re

# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    app = BashkarApp()
    app.mainloop()

